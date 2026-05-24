"""web/app8508.py — FastAPI app port 8508, layout giống Streamlit
Chạy: python -X utf8 -m uvicorn web.app8508:app --port 8508 --reload
"""
import sqlite3, json, time
from pathlib import Path
from typing import Optional
from fastapi import FastAPI, Request
from fastapi.responses import HTMLResponse, JSONResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

ROOT = Path(__file__).parent.parent
DB   = ROOT / "data" / "TTHC.sqlite"

app = FastAPI(title="TTHC 8508")
app.mount("/static", StaticFiles(directory=str(ROOT / "web" / "static")), name="static")
templates = Jinja2Templates(directory=str(ROOT / "web" / "templates"))

def _db():
    con = sqlite3.connect(str(DB))
    con.row_factory = sqlite3.Row
    return con

# ── Pages ──────────────────────────────────────────────────────────────────────
@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse(request=request, name="app.html")

# ── API: Địa chất ──────────────────────────────────────────────────────────────
@app.get("/api/boreholes")
async def api_boreholes(zone: Optional[str] = None):
    con = _db()
    q = """SELECT name, elevation_m, x_coord_m, y_coord_m,
               CASE WHEN name LIKE 'KE-%'  THEN 'KE'
                    WHEN name LIKE 'BXN-%' THEN 'BXN'
                    WHEN name LIKE 'NHC-%' THEN 'NHC'
                    WHEN name LIKE 'ND-%'  THEN 'QTT'
                    ELSE 'OTHER' END AS zone
           FROM boreholes WHERE x_coord_m > 1000000 ORDER BY name"""
    rows = con.execute(q).fetchall()
    con.close()
    bhs = [dict(r) for r in rows]
    if zone and zone != "all":
        bhs = [b for b in bhs if b["zone"] == zone]
    return JSONResponse({"boreholes": bhs})

@app.get("/api/borehole/{name}/layers")
async def api_layers(name: str):
    con = _db()
    bh = con.execute("SELECT id,elevation_m FROM boreholes WHERE name=?", (name,)).fetchone()
    if not bh:
        con.close()
        return JSONResponse({"layers": [], "spt": [], "vst": []})
    bid, elev = bh["id"], bh["elevation_m"] or 0

    # layers: KHÔNG có cột layer_no — chỉ có id/symbol/description/depth_*
    layers = con.execute(
        "SELECT id AS layer_no, symbol, description, depth_top_m, depth_bot_m "
        "FROM layers WHERE borehole_id=? ORDER BY depth_top_m",
        (bid,)).fetchall()

    spt = con.execute(
        "SELECT depth_m, N FROM spt_values WHERE borehole_id=? ORDER BY depth_m",
        (bid,)).fetchall()

    # vane_shear_tests link qua vst_locations.name (KHÔNG có FK trực tiếp borehole_id).
    # Cột su là Su_kPa (chữ S hoa).
    vst = con.execute(
        "SELECT v.depth_m, v.Su_kPa AS su_kPa "
        "FROM vane_shear_tests v JOIN vst_locations l ON l.id = v.vst_loc_id "
        "WHERE l.name = ? ORDER BY v.depth_m",
        (name,)).fetchall()
    con.close()
    return JSONResponse({
        "elevation_m": elev,
        "layers": [dict(r) for r in layers],
        "spt":    [dict(r) for r in spt],
        "vst":    [dict(r) for r in vst],
    })

# ── API: CDM tọa độ ────────────────────────────────────────────────────────────
@app.get("/api/cdm/toado")
async def api_cdm_toado(zone: Optional[str] = None):
    con = _db()
    q = "SELECT zone, point_name, northing_m, easting_m FROM cdm_toado"
    params = []
    if zone and zone != "all":
        q += " WHERE zone=?"; params.append(zone)
    rows = con.execute(q, params).fetchall()
    con.close()
    return JSONResponse({"n": len(rows),
        "points": [{"z":r[0],"n":r[1],"x":r[2],"y":r[3]} for r in rows]})

@app.get("/api/cdm/stats")
async def api_cdm_stats():
    con = _db()
    zone_rows = con.execute("""
        SELECT zone, COUNT(*) n_coc,
               ROUND(AVG(L_design_m),2) L_avg, ROUND(MIN(L_design_m),2) L_min,
               ROUND(MAX(L_design_m),2) L_max,
               ROUND(SUM(3.14159265/4*D_m*D_m*L_design_m),0) V_m3
        FROM cdm_thi_cong WHERE L_design_m IS NOT NULL GROUP BY zone ORDER BY zone""").fetchall()
    bh_rows = con.execute("""
        SELECT bh_dieu_hanh, zone, COUNT(*) n_coc,
               ROUND(AVG(L_design_m),2) L_avg,
               ROUND(SUM(3.14159265/4*D_m*D_m*L_design_m),0) V_m3
        FROM cdm_thi_cong WHERE L_design_m IS NOT NULL AND bh_dieu_hanh IS NOT NULL
        GROUP BY bh_dieu_hanh, zone ORDER BY zone, bh_dieu_hanh""").fetchall()
    status_rows = con.execute("""
        SELECT zone,status,COUNT(*) n FROM cdm_thi_cong GROUP BY zone,status ORDER BY zone,status""").fetchall()
    con.close()
    return JSONResponse({"by_zone":[dict(r) for r in zone_rows],
                         "by_bh":  [dict(r) for r in bh_rows],
                         "status": [dict(r) for r in status_rows]})

@app.get("/api/cdm/tien-do")
async def api_tien_do():
    con = _db()
    daily = con.execute("""
        SELECT ngay_thi_cong, zone, COUNT(*) n FROM cdm_thi_cong
        WHERE ngay_thi_cong IS NOT NULL AND status IN ('hoan_thanh','dang_thi_cong')
        GROUP BY ngay_thi_cong, zone ORDER BY ngay_thi_cong""").fetchall()
    todoi = con.execute("""
        SELECT to_doi, zone, COUNT(*) n,
               SUM(CASE WHEN status='hoan_thanh' THEN 1 ELSE 0 END) done
        FROM cdm_thi_cong WHERE to_doi IS NOT NULL GROUP BY to_doi,zone ORDER BY to_doi""").fetchall()
    con.close()
    return JSONResponse({"daily":[dict(r) for r in daily],
                         "todoi":[dict(r) for r in todoi]})

@app.get("/api/cdm/kiem-tra")
async def api_kiem_tra():
    con = _db()
    rows = con.execute("""
        SELECT zone,point_name,loai_tn,ngay_thu_nghiem,tuoi_ngay,
               do_sau_tu_dinh_m,z_sample_m,qu_kPa,qu_yc_kPa,dat_yeu_cau,
               ham_luong_xi_mang_pct,ghi_chu
        FROM cdm_kiem_tra ORDER BY zone,ngay_thu_nghiem,z_sample_m""").fetchall()
    con.close()
    return JSONResponse({"rows":[dict(r) for r in rows]})

@app.get("/api/ke-binhdo")
async def api_ke_binhdo():
    con = _db()
    try:
        rows = con.execute(
            "SELECT polyline_id,x_m,y_m FROM ke_binhdo_toadoke ORDER BY polyline_id,vertex_idx"
        ).fetchall()
    except Exception:
        rows = []
    con.close()
    polys: dict = {}
    for pid, x, y in rows:
        polys.setdefault(pid, []).append({"x": x, "y": y})
    return JSONResponse({"polylines": list(polys.values())})

@app.get("/api/settlement/compare")
async def api_settlement(bh_name: str = "BXN-CV-HK1", H_fill: float = 3.0):
    """Tính lún so sánh phương án — wrap settlement_calc.py"""
    try:
        import sys; sys.path.insert(0, str(ROOT / "scripts"))
        from settlement_calc import compare_methods
        result = compare_methods(bh_name, "BXN", H_fill)
        return JSONResponse({"ok": True, "result": result})
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)})


@app.get("/api/settlement/analyze")
async def api_settlement_analyze(
    bh_name: str = "BXN-CV-HK1",
    zone: str = "BXN",
    H_fill: float = 3.0,
    lim_cm: float = 30.0,
    tc_months: float = 6.0,
):
    """Phân tích lún đầy đủ — zone + lim + tc params."""
    try:
        import sys; sys.path.insert(0, str(ROOT / "scripts"))
        from settlement_calc import compare_methods
        result = compare_methods(bh_name, zone, H_fill,
                                 residual_limit_cm=lim_cm,
                                 t_construction_months=tc_months)
        return JSONResponse({"ok": True, "result": result})
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)})


@app.get("/api/ke-sw/catalog")
async def api_ke_sw_catalog():
    """Catalog cọc ván SW từ JSON."""
    p = ROOT / "data" / "sw_pile_catalog.json"
    if not p.exists():
        return JSONResponse({"piles": []})
    data = json.loads(p.read_text(encoding="utf-8"))
    piles = data.get("piles", data) if isinstance(data, dict) else data
    return JSONResponse({"piles": piles})


@app.get("/api/ke-sw/design")
async def api_ke_sw_design():
    """Kết quả thiết kế kè KE: ke_sw_nt_detail (SQLite) + ke_sw JSON config."""
    con = _db()
    rows = con.execute("SELECT * FROM ke_sw_nt_detail ORDER BY bh_name").fetchall()
    con.close()
    nt_detail = [dict(r) for r in rows]
    cfg_path = ROOT / "data" / "ke_sw_202605_TTHC.json"
    cfg = json.loads(cfg_path.read_text(encoding="utf-8")) if cfg_path.exists() else {}
    return JSONResponse({"nt_detail": nt_detail, "config": cfg})


@app.get("/api/ke-sw/winkler")
async def api_ke_sw_winkler():
    """Kết quả nội lực Winkler cọc ván SW: ke_sw_winkler_results."""
    con = _db()
    try:
        rows = con.execute(
            "SELECT bh_name, pile_type, L_m, load_case, u_top_mm, u_max_mm, "
            "M_max_kNm, Mcr_kNm, Q_max_kN, mcr_ratio, u_ok, mcr_ok, solver "
            "FROM ke_sw_winkler_results ORDER BY bh_name, pile_type, load_case"
        ).fetchall()
        return JSONResponse({"rows": [dict(r) for r in rows]})
    except Exception as e:
        return JSONResponse({"rows": [], "error": str(e)})
    finally:
        con.close()


@app.get("/api/ke-sw/stability")
async def api_ke_sw_stability():
    """Kết quả ổn định tổng thể Fellenius: ke_sw_stability."""
    con = _db()
    try:
        rows = con.execute(
            "SELECT bh_name, pile_type, L_m, method, Fs_slip, Fs_overturning, "
            "top_elev, Z_m, Zb_m, wlvl_front, wlvl_back, q_kPa, "
            "cdm_a, cdm_c_col, su_front, su_back, H1_m "
            "FROM ke_sw_stability ORDER BY bh_name, pile_type"
        ).fetchall()
        return JSONResponse({"rows": [dict(r) for r in rows]})
    except Exception as e:
        return JSONResponse({"rows": [], "error": str(e)})
    finally:
        con.close()


@app.get("/api/ke-sw/profile-chainage")
async def api_ke_sw_profile_chainage():
    """Trắc dọc địa chất + bình đồ HK theo tuyến kè SW.

    Trả về:
      - alignment_polyline: vertex tuyến kè (x_m=Easting, y_m=Northing)
      - boreholes: 8 HK on_sw_alignment với chainage_m (PCA SVD chiếu lên trục chính)
      - layers per BH: symbol, depth_top, depth_bot (m từ mặt đất), elev_top, elev_bot
      - spt per BH: depth_m, N
      - vst per BH: depth_m, Su_kPa
      - distances: pairwise giữa các HK theo thứ tự chainage
    """
    cfg_path = ROOT / "data" / "ke_sw_202605_TTHC.json"
    if not cfg_path.exists():
        return JSONResponse({"error": "Thiếu ke_sw_202605_TTHC.json"}, status_code=404)
    cfg = json.loads(cfg_path.read_text(encoding="utf-8"))
    align_names = ["KE-" + b["name"] for b in cfg.get("boreholes", [])
                   if b.get("on_sw_alignment")]
    if not align_names:
        return JSONResponse({"alignment_polyline": [], "boreholes": [], "distances": []})

    con = _db()
    try:
        # HK trên tuyến
        placeholders = ",".join("?" * len(align_names))
        bh_rows = con.execute(
            f"SELECT id, name, elevation_m, x_coord_m, y_coord_m "
            f"FROM boreholes WHERE name IN ({placeholders})",
            align_names,
        ).fetchall()
        bh_by_id = {r["id"]: dict(r) for r in bh_rows}

        # Layers + SPT + VST cho 8 HK
        layers_by_bh: dict = {}
        for r in con.execute(
            f"SELECT borehole_id, symbol, description, depth_top_m, depth_bot_m "
            f"FROM layers WHERE borehole_id IN ({','.join(str(i) for i in bh_by_id)}) "
            f"ORDER BY borehole_id, depth_top_m"
        ).fetchall():
            layers_by_bh.setdefault(r["borehole_id"], []).append(dict(r))
        spt_by_bh: dict = {}
        for r in con.execute(
            f"SELECT borehole_id, depth_m, N FROM spt_values "
            f"WHERE borehole_id IN ({','.join(str(i) for i in bh_by_id)}) "
            f"ORDER BY borehole_id, depth_m"
        ).fetchall():
            spt_by_bh.setdefault(r["borehole_id"], []).append(dict(r))
        vst_by_bh: dict = {}
        for r in con.execute(
            f"SELECT l.name AS bh_name, v.depth_m, v.Su_kPa "
            f"FROM vane_shear_tests v JOIN vst_locations l ON l.id=v.vst_loc_id "
            f"WHERE l.name IN ({placeholders}) ORDER BY l.name, v.depth_m",
            align_names,
        ).fetchall():
            # map name → bh_id
            for bid, b in bh_by_id.items():
                if b["name"] == r["bh_name"]:
                    vst_by_bh.setdefault(bid, []).append({
                        "depth_m": r["depth_m"], "Su_kPa": r["Su_kPa"],
                    })
                    break

        # Tuyến kè polyline
        try:
            poly_rows = con.execute(
                "SELECT polyline_id, x_m, y_m FROM ke_binhdo_toadoke "
                "ORDER BY polyline_id, vertex_idx"
            ).fetchall()
            alignment = [{"polyline_id": r["polyline_id"],
                          "x_m": r["x_m"], "y_m": r["y_m"]} for r in poly_rows]
        except Exception:
            alignment = []
    finally:
        con.close()

    # PCA SVD: chiếu (Northing, Easting) HK lên trục chính → chainage_m
    # boreholes: x_coord_m=Northing, y_coord_m=Easting
    pts = [(b["x_coord_m"], b["y_coord_m"]) for b in bh_by_id.values()]
    cx = sum(p[0] for p in pts) / len(pts)
    cy = sum(p[1] for p in pts) / len(pts)
    centered = [(p[0] - cx, p[1] - cy) for p in pts]
    # 2x2 covariance + eigenvector tay (SVD đơn giản)
    sxx = sum(p[0] * p[0] for p in centered)
    sxy = sum(p[0] * p[1] for p in centered)
    syy = sum(p[1] * p[1] for p in centered)
    tr = sxx + syy
    det = sxx * syy - sxy * sxy
    disc = max(tr * tr / 4 - det, 0)
    lam = tr / 2 + disc ** 0.5
    # Eigenvector cho lam
    if abs(sxy) > 1e-9:
        vx, vy = lam - syy, sxy
    else:
        vx, vy = (1.0, 0.0) if sxx >= syy else (0.0, 1.0)
    nrm = (vx * vx + vy * vy) ** 0.5 or 1
    vx, vy = vx / nrm, vy / nrm
    # Chainage = projection lên trục chính
    chainage_by_id: dict = {}
    for bid, b in bh_by_id.items():
        ch = (b["x_coord_m"] - cx) * vx + (b["y_coord_m"] - cy) * vy
        chainage_by_id[bid] = ch
    # Normalize chainage so min = 0
    ch_min = min(chainage_by_id.values())
    for bid in chainage_by_id:
        chainage_by_id[bid] -= ch_min

    # Build response
    boreholes_out = []
    for bid, b in bh_by_id.items():
        elev = b["elevation_m"] or 0
        lyrs = layers_by_bh.get(bid, [])
        layers_out = []
        for ly in lyrs:
            dtop = ly["depth_top_m"] or 0
            dbot = ly["depth_bot_m"] or 0
            layers_out.append({
                "symbol": ly["symbol"],
                "description": ly["description"],
                "depth_top_m": dtop, "depth_bot_m": dbot,
                "elev_top_m": elev - dtop, "elev_bot_m": elev - dbot,
            })
        boreholes_out.append({
            "name": b["name"], "short_name": b["name"].replace("KE-", ""),
            "elevation_m": elev,
            "x_coord_m": b["x_coord_m"], "y_coord_m": b["y_coord_m"],
            "chainage_m": round(chainage_by_id[bid], 2),
            "layers": layers_out,
            "spt": spt_by_bh.get(bid, []),
            "vst": vst_by_bh.get(bid, []),
        })
    boreholes_out.sort(key=lambda b: b["chainage_m"])

    # Khoảng cách giữa các HK theo thứ tự chainage
    distances = []
    for i in range(1, len(boreholes_out)):
        a = boreholes_out[i - 1]
        b = boreholes_out[i]
        dx = a["x_coord_m"] - b["x_coord_m"]
        dy = a["y_coord_m"] - b["y_coord_m"]
        distances.append({
            "from": a["short_name"], "to": b["short_name"],
            "distance_m": round((dx * dx + dy * dy) ** 0.5, 1),
            "chainage_diff_m": round(b["chainage_m"] - a["chainage_m"], 1),
        })

    return JSONResponse({
        "alignment_polyline": alignment,
        "boreholes": boreholes_out,
        "distances": distances,
        "pca_axis": {"cx": cx, "cy": cy, "vx": vx, "vy": vy},
    })


# ── API: Bản đồ vị trí (VN-2000 → WGS-84) ────────────────────────────────────

@app.get("/api/geo/locations")
async def api_geo_locations(epsg: int = 9210):
    """
    Trả về tọa độ WGS-84 (lat/lon) của hố khoan, ranh kè và CDM centroids.
    epsg: 9210 (mặc định TTHC Q1/Thủ Thiêm) | 9209 | 3405
    """
    try:
        from pyproj import Transformer
        tr = Transformer.from_crs(epsg, 4326, always_xy=True)
    except Exception as e:
        return JSONResponse({"error": f"pyproj: {e}"}, status_code=500)

    con = _db()

    # Hố khoan
    bh_rows = con.execute("""
        SELECT name, elevation_m, x_coord_m, y_coord_m,
               CASE WHEN name LIKE 'KE-%'  THEN 'KE'
                    WHEN name LIKE 'BXN-%' THEN 'BXN'
                    WHEN name LIKE 'NHC-%' THEN 'NHC'
                    WHEN name LIKE 'ND-%'  THEN 'QTT'
                    ELSE 'OTHER' END AS zone
        FROM boreholes WHERE x_coord_m > 1000000 ORDER BY name
    """).fetchall()

    boreholes = []
    for r in bh_rows:
        try:
            # x_coord_m = Northing, y_coord_m = Easting → transform(E, N)
            lon, lat = tr.transform(float(r["y_coord_m"]), float(r["x_coord_m"]))
            boreholes.append({
                "name": r["name"], "zone": r["zone"],
                "elevation_m": r["elevation_m"],
                "lat": round(lat, 7), "lon": round(lon, 7),
            })
        except Exception:
            pass

    # Ranh kè polylines → lat/lon
    ke_lines = []
    try:
        poly_rows = con.execute(
            "SELECT polyline_id, x_m, y_m FROM ke_binhdo_toadoke ORDER BY polyline_id, vertex_idx"
        ).fetchall()
        polys: dict = {}
        for pid, xm, ym in poly_rows:
            polys.setdefault(pid, []).append((xm, ym))
        for pid, pts in polys.items():
            ll = []
            for xm, ym in pts:
                try:
                    # x_m = Easting, y_m = Northing
                    lon2, lat2 = tr.transform(float(xm), float(ym))
                    ll.append([round(lat2, 7), round(lon2, 7)])
                except Exception:
                    pass
            if ll:
                ke_lines.append(ll)
    except Exception:
        pass

    # CDM — centroid per zone (không load 27k điểm cho bản đồ nền)
    cdm_zones = []
    try:
        cdm_rows = con.execute("""
            SELECT zone,
                   AVG(northing_m) cx, AVG(easting_m) cy,
                   MIN(northing_m) nx_min, MAX(northing_m) nx_max,
                   MIN(easting_m)  ey_min, MAX(easting_m)  ey_max,
                   COUNT(*) n
            FROM cdm_toado GROUP BY zone
        """).fetchall()
        for r in cdm_rows:
            lon_c, lat_c = tr.transform(float(r[2]), float(r[1]))
            # Bounding box corners
            corners = []
            for nx, ey in [
                (r[3], r[4]), (r[3], r[5]),
                (r[6], r[4]), (r[6], r[5]),
            ]:
                lo, la = tr.transform(float(ey), float(nx))
                corners.append([round(la, 7), round(lo, 7)])
            cdm_zones.append({
                "zone": r[0], "n": r[7],
                "lat": round(lat_c, 7), "lon": round(lon_c, 7),
                "bbox": corners,
            })
    except Exception:
        pass

    con.close()
    return JSONResponse({
        "epsg": epsg,
        "boreholes": boreholes,
        "ke_lines": ke_lines,
        "cdm_zones": cdm_zones,
    })


# ── CDM Design endpoints ───────────────────────────────────────────────────────
_SOFT_SYMBOLS = {'1', '1b', 'XMD'}


@app.get("/api/cdm/design/clay-params")
async def cdm_clay_params(bh: str):
    """Trả thông số lớp bùn + Su_avg + gamma_avg cho một hố khoan."""
    con = _db()
    row = con.execute("SELECT id, elevation_m FROM boreholes WHERE name=?", (bh,)).fetchone()
    if not row:
        con.close()
        return JSONResponse({"error": f"Không tìm thấy {bh}"}, status_code=404)
    bid, elev = row["id"], float(row["elevation_m"] or 0)

    layers = con.execute(
        "SELECT symbol, description, depth_top_m, depth_bot_m "
        "FROM layers WHERE borehole_id=? ORDER BY depth_top_m", (bid,),
    ).fetchall()

    soft_layers, hard_started = [], False
    for ly in layers:
        sym = (ly["symbol"] or "").strip()
        if not hard_started and sym in _SOFT_SYMBOLS:
            soft_layers.append(ly)
        elif sym not in _SOFT_SYMBOLS and soft_layers:
            hard_started = True
    if not soft_layers and layers:
        soft_layers = [layers[0]]

    depth_top = float(soft_layers[0]["depth_top_m"]) if soft_layers else 0.0
    depth_bot = float(soft_layers[-1]["depth_bot_m"]) if soft_layers else 0.0
    h_clay    = depth_bot - depth_top
    top_clay_m = elev - depth_top

    vst = con.execute(
        "SELECT AVG(v.Su_kPa) AS su FROM vane_shear_tests v "
        "JOIN vst_locations l ON l.id=v.vst_loc_id "
        "WHERE l.name=? AND v.depth_m BETWEEN ? AND ?",
        (bh, depth_top, depth_bot),
    ).fetchone()
    Su_avg = round(float(vst["su"]), 1) if vst and vst["su"] else 11.0

    gam = con.execute(
        "SELECT AVG(gamma_kNm3) AS g FROM lab_tests lt "
        "JOIN boreholes b ON b.id=lt.borehole_id WHERE b.name=? AND gamma_kNm3>0", (bh,),
    ).fetchone()
    gamma_avg = round(float(gam["g"]), 2) if gam and gam["g"] else 15.0

    all_layers = [{"symbol": ly["symbol"],
                   "depth_top_m": ly["depth_top_m"],
                   "depth_bot_m": ly["depth_bot_m"],
                   "thickness_m": round(float(ly["depth_bot_m"] or 0) - float(ly["depth_top_m"] or 0), 2)}
                  for ly in layers]
    con.close()
    return JSONResponse({
        "bh": bh, "elevation_m": elev,
        "top_clay_m": round(top_clay_m, 2), "depth_top_m": round(depth_top, 2),
        "depth_bot_m": round(depth_bot, 2), "h_clay_m": round(h_clay, 2),
        "Su_avg_kPa": Su_avg, "gamma_avg_kNm3": gamma_avg, "layers": all_layers,
    })


@app.post("/api/cdm/design/export-excel")
async def cdm_export_excel(request: Request):
    """scenarios + params JSON → file Excel."""
    body      = await request.json()
    scenarios = body.get("scenarios", [])
    params    = body.get("params",    {})
    try:
        import io, openpyxl
        from openpyxl.styles import PatternFill, Font, Alignment
        wb  = openpyxl.Workbook()
        ws1 = wb.active; ws1.title = "Thong so"
        ws1.append(["Thông số", "Giá trị", "Đơn vị"])
        units = {"D":"m","e":"m","Lc":"m","qu":"kPa","Su":"kPa","gamma":"kN/m³",
                 "q_total":"kN/m²","q_static":"kN/m²","top_clay":"m","h_clay":"m"}
        for k, v in params.items():
            ws1.append([k, v, units.get(k,"")])
        ws2 = wb.create_sheet("So sanh PA")
        if scenarios:
            hdrs = list(scenarios[0].keys())
            ws2.append(hdrs)
            hdr_fill = PatternFill("solid", fgColor="1A3A6B")
            for cell in ws2[1]:
                cell.fill = hdr_fill
                cell.font = Font(bold=True, color="FFFFFF")
                cell.alignment = Alignment(horizontal="center")
            for i, sc in enumerate(scenarios):
                ws2.append([sc.get(h) for h in hdrs])
                if i % 2 == 1:
                    for cell in ws2[ws2.max_row]:
                        cell.fill = PatternFill("solid", fgColor="EEF3FF")
        buf = io.BytesIO(); wb.save(buf); buf.seek(0)
        return Response(
            content=buf.read(),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": 'attachment; filename="CDM_TinhToan.xlsx"'},
        )
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


# ── CDM design: 5 endpoints mới ────────────────────────────────────────────────

_SOFT_SYMBOLS_SET = {"1", "1a", "1b", "2", "2a", "2b", "XMD", "b", "bùn"}

@app.get("/api/cdm/design/layers")
async def cdm_design_layers(bh: str, zone: str = ""):
    """Layer data với Cc/Cs/e0/PC/Cv/γ per layer, fallback zone avg khi thiếu."""
    con = _db()
    try:
        row = con.execute("SELECT id FROM boreholes WHERE name=?", (bh,)).fetchone()
        if not row:
            return JSONResponse({"error": f"Không tìm thấy {bh}"}, status_code=404)
        bid = row["id"]
        layers_raw = con.execute(
            "SELECT symbol, description, depth_top_m, depth_bot_m "
            "FROM layers WHERE borehole_id=? ORDER BY depth_top_m", (bid,)
        ).fetchall()
        result = []
        for ly in layers_raw:
            sym  = (ly["symbol"] or "").strip()
            dt   = float(ly["depth_top_m"] or 0)
            db   = float(ly["depth_bot_m"]  or 0)
            lab  = con.execute("""
                SELECT AVG(gamma_kNm3) g, AVG(e0) e0, AVG(Cc) Cc, AVG(Cs) Cs,
                       AVG(PC_kPa) PC, AVG(Cv_cm2s) Cv
                FROM lab_tests WHERE borehole_id=?
                  AND depth_from_m >= ? AND depth_to_m <= ?
            """, (bid, dt, db)).fetchone()
            rec = {
                "symbol": sym,
                "description": (ly["description"] or "")[:40],
                "depth_top_m": round(dt, 2),
                "depth_bot_m": round(db, 2),
                "thickness_m": round(db - dt, 2),
                "gamma_kNm3": round(float(lab["g"]),  2) if lab and lab["g"]  else None,
                "e0":          round(float(lab["e0"]), 3) if lab and lab["e0"] else None,
                "Cc":          round(float(lab["Cc"]), 4) if lab and lab["Cc"] else None,
                "Cs":          round(float(lab["Cs"]), 4) if lab and lab["Cs"] else None,
                "PC_kPa":      round(float(lab["PC"]), 1) if lab and lab["PC"] else None,
                "Cv_cm2s":     float(lab["Cv"]) if lab and lab["Cv"] else None,
                "source": bh, "borrowed": False,
            }
            if zone and (rec["Cc"] is None or rec["e0"] is None):
                zl = con.execute("""
                    SELECT AVG(lt.gamma_kNm3) g, AVG(lt.e0) e0, AVG(lt.Cc) Cc,
                           AVG(lt.Cs) Cs, AVG(lt.PC_kPa) PC, AVG(lt.Cv_cm2s) Cv,
                           COUNT(DISTINCT lt.borehole_id) n
                    FROM lab_tests lt JOIN boreholes b ON b.id=lt.borehole_id
                    JOIN zones z ON z.id=b.zone_id
                    WHERE z.code=? AND lt.symbol_tcvn=? AND lt.borehole_id!=?
                """, (zone, sym, bid)).fetchone()
                if zl and zl["Cc"]:
                    if rec["Cc"]    is None: rec["Cc"]    = round(float(zl["Cc"]),  4)
                    if rec["Cs"]    is None: rec["Cs"]    = round(float(zl["Cs"]),  4) if zl["Cs"] else None
                    if rec["e0"]    is None: rec["e0"]    = round(float(zl["e0"]),  3) if zl["e0"] else None
                    if rec["PC_kPa"] is None: rec["PC_kPa"] = round(float(zl["PC"]), 1) if zl["PC"] else None
                    if rec["Cv_cm2s"] is None: rec["Cv_cm2s"] = float(zl["Cv"]) if zl["Cv"] else None
                    rec["source"] = f"Zone {zone} avg (n={zl['n']})"
                    rec["borrowed"] = True
            result.append(rec)
        return JSONResponse({"bh": bh, "layers": result})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)
    finally:
        con.close()


@app.get("/api/cdm/design/geo-params-full")
async def cdm_geo_params_full(bh: str, zone: str = "KE"):
    """Clay params với γ nearest-BH fallback + source info."""
    base = await cdm_clay_params(bh)
    data = json.loads(base.body)
    if "error" in data:
        return JSONResponse(data, status_code=404)
    con = _db()
    try:
        brow = con.execute(
            "SELECT x_coord_m, y_coord_m FROM boreholes WHERE name=?", (bh,)
        ).fetchone()
        data["gamma_source"] = bh
        data["gamma_dist_m"] = None
        if data.get("gamma_avg_kNm3", 15.0) == 15.0 and brow and brow[0] is not None:
            x0, y0 = float(brow[0]), float(brow[1])
            cands = con.execute("""
                SELECT b.name, b.x_coord_m, b.y_coord_m, AVG(lt.gamma_kNm3) g
                FROM lab_tests lt JOIN boreholes b ON b.id=lt.borehole_id
                JOIN zones z ON z.id=b.zone_id
                WHERE z.code=? AND b.name!=? AND lt.gamma_kNm3>0
                  AND b.x_coord_m IS NOT NULL
                GROUP BY b.id
            """, (zone, bh)).fetchall()
            if cands:
                best = min(cands, key=lambda r: (float(r[1])-x0)**2+(float(r[2])-y0)**2)
                d = ((float(best[1])-x0)**2+(float(best[2])-y0)**2)**0.5
                if best[3]:
                    data["gamma_avg_kNm3"] = round(float(best[3]), 2)
                    data["gamma_source"]   = best[0]
                    data["gamma_dist_m"]   = round(d, 1)
        data["Su_source"] = f"VST zone {zone}"
        return JSONResponse(data)
    finally:
        con.close()


@app.post("/api/cdm/design/save-scenario")
async def cdm_save_scenario(request: Request):
    """Lưu kịch bản thiết kế CDM vào SQLite."""
    body   = await request.json()
    params = body.get("params", {})
    result = body.get("result", {})
    con    = _db()
    try:
        con.execute("""
            CREATE TABLE IF NOT EXISTS cdm_design (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                zone TEXT, bh_name TEXT,
                D_m REAL, e_m REAL, arrangement TEXT, Lc_m REAL,
                qu_kPa REAL, kEc REAL, Fs_mat REAL,
                Su_kPa REAL, gamma_kNm3 REAL,
                h_fill_m REAL, gamma_fill REAL, q_traffic_kPa REAL,
                a_pct REAL, Etb_kPa REAL, S1_cm REAL,
                Qa_kN REAL, Pcol_kN REAL, ok_sct INTEGER,
                tau_se_kPa REAL, tau_ase_kPa REAL, ok_punch INTEGER,
                ts TEXT
            )
        """)
        con.execute("""
            INSERT INTO cdm_design
            (zone,bh_name,D_m,e_m,arrangement,Lc_m,qu_kPa,kEc,Fs_mat,
             Su_kPa,gamma_kNm3,h_fill_m,gamma_fill,q_traffic_kPa,
             a_pct,Etb_kPa,S1_cm,Qa_kN,Pcol_kN,ok_sct,
             tau_se_kPa,tau_ase_kPa,ok_punch,ts)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
        """, (
            params.get("zone",""), params.get("bh",""),
            params.get("D"),params.get("e"),params.get("arrangement"),params.get("Lc"),
            params.get("qu"),params.get("kEc"),params.get("Fs_mat"),
            params.get("Su"),params.get("gamma"),
            params.get("h_fill"),params.get("gamma_fill"),params.get("q_traffic"),
            result.get("a_pct"),result.get("Etb"),result.get("S1_cm"),
            result.get("Qa"),result.get("Pcol"),1 if result.get("ok") else 0,
            result.get("tse"),result.get("tase"),1 if result.get("pass_punch") else 0,
            time.strftime("%Y-%m-%dT%H:%M:%S"),
        ))
        con.commit()
        lid = con.execute("SELECT last_insert_rowid()").fetchone()[0]
        return JSONResponse({"ok": True, "id": lid})
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)
    finally:
        con.close()


@app.get("/api/cdm/design/scenarios-list")
async def cdm_scenarios_list(zone: str = "", bh: str = ""):
    """Danh sách kịch bản thiết kế đã lưu."""
    con = _db()
    try:
        q = "SELECT * FROM cdm_design WHERE 1=1"
        args: list = []
        if zone: q += " AND zone=?";    args.append(zone)
        if bh:   q += " AND bh_name=?"; args.append(bh)
        q += " ORDER BY ts DESC LIMIT 50"
        rows = con.execute(q, args).fetchall()
        return JSONResponse({"rows": [dict(r) for r in rows]})
    except Exception as e:
        return JSONResponse({"rows": [], "error": str(e)})
    finally:
        con.close()


@app.post("/api/cdm/design/export-word")
async def cdm_export_word(request: Request):
    """Xuất báo cáo Word thiết kế CDM."""
    body      = await request.json()
    params    = body.get("params", {})
    scenarios = body.get("scenarios", [])
    rec_idx   = int(body.get("rec_idx", 0))
    try:
        import io
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from lxml import etree

        doc = Document()
        sec = doc.sections[0]
        for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
            setattr(sec, attr, Cm(2.0))

        t = doc.add_heading("THIẾT KẾ TRỤ ĐẤT XI MĂNG CDM", 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(
            f"Hố khoan: {params.get('bh','—')}  |  Khu vực: {params.get('zone','—')}"
            f"  |  Ngày: {time.strftime('%d/%m/%Y')}"
        )

        doc.add_heading("1. Thông số thiết kế", 1)
        tbl = doc.add_table(rows=1, cols=3); tbl.style = "Table Grid"
        for i, h in enumerate(("Thông số","Giá trị","Đơn vị")):
            tbl.rows[0].cells[i].text = h
        for name, val, unit in [
            ("Đường kính cọc D", f"{params.get('D','')} m", "m"),
            ("Khoảng cách tim e", f"{params.get('e','')} m", "m"),
            ("Sơ đồ bố trí", params.get("arrangement",""), ""),
            ("Chiều dài cọc Lc", f"{params.get('Lc','')} m", "m"),
            ("qu thiết kế", f"{params.get('qu','')} kPa", "kPa"),
            ("Su đất yếu TB", f"{params.get('Su','')} kPa", "kPa"),
            ("Chiều dày lớp bùn Hc", f"{params.get('h_clay','')} m", "m"),
            ("Tải trọng tĩnh q_st", f"{params.get('q_static','')} kN/m²", "kN/m²"),
            ("Tải trọng tổng q", f"{params.get('q_total','')} kN/m²", "kN/m²"),
        ]:
            r = tbl.add_row().cells
            r[0].text = name; r[1].text = str(val); r[2].text = unit

        doc.add_heading("2. So sánh phương án e = 1.2 → 2.4 m", 1)
        if scenarios:
            ks = list(scenarios[0].keys())
            t2 = doc.add_table(rows=1, cols=len(ks)); t2.style = "Table Grid"
            for i, k in enumerate(ks): t2.rows[0].cells[i].text = k
            for i, sc in enumerate(scenarios):
                row = t2.add_row().cells
                for j, k in enumerate(ks):
                    v = sc.get(k, "")
                    row[j].text = f"{v:.2f}" if isinstance(v, float) else str(v)

        if 0 <= rec_idx < len(scenarios):
            rec = scenarios[rec_idx]
            doc.add_heading(f"3. Phương án kiến nghị: e = {rec.get('e',0):.1f} m", 1)
            for line in [
                f"Tỷ lệ thay thế a = {rec.get('a_pct',rec.get('a (%)',0)):.1f} %",
                f"Mô đun tổng hợp Etb = {rec.get('Etb',rec.get('Etb (kN/m²)',0)):,.0f} kN/m²",
                f"Độ lún sơ cấp S₁ = {rec.get('S1_cm',rec.get('S₁ (cm)',0)):.2f} cm",
                f"Kiểm tra sức chịu tải: {'Đạt' if rec.get('ok',rec.get('ok_sct',True)) else 'Không đạt'}",
                f"Kiểm tra xuyên thủng: {'Đạt' if rec.get('pass_punch',rec.get('pass',True)) else 'Không đạt'}",
            ]:
                doc.add_paragraph(line)

        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        return Response(
            content=buf.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": 'attachment; filename="CDM_ThietKe.docx"'},
        )
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


# ── PDF endpoints ──────────────────────────────────────────────────────────────
@app.post("/api/pdf/weasyprint")
async def pdf_weasyprint(request: Request):
    """Nhận HTML string → trả về PDF bytes (WeasyPrint)."""
    body = await request.json()
    html_content = body.get("html", "<p>empty</p>")
    page_size    = body.get("page_size", "A4")
    margin_cm    = body.get("margin_cm", 2)
    try:
        from weasyprint import HTML, CSS
        css = CSS(string=f"@page {{ size: {page_size}; margin: {margin_cm}cm; }}")
        pdf_bytes = HTML(string=html_content, base_url=str(ROOT)).write_pdf(stylesheets=[css])
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="report_weasyprint.pdf"'},
        )
    except ImportError:
        return JSONResponse({"error": "WeasyPrint chưa cài — chạy: pip install weasyprint"}, status_code=503)
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@app.get("/api/pdf/weasyprint/status")
async def pdf_weasyprint_status():
    try:
        import weasyprint
        return JSONResponse({"available": True, "version": weasyprint.__version__})
    except ImportError:
        return JSONResponse({"available": False})


# ── API: Kiểm tra mẫu thí nghiệm ──────────────────────────────────────────────
@app.get("/api/sample-check/tccs41")
async def api_sample_check_tccs41():
    """TCCS 41:2022 Điều 5.3.7 — kiểm tra số mẫu nén cố kết Cc/Cs/Cv/PC
    cho 3 khu vực NHC/BXN/KE."""
    import sys
    sys.path.insert(0, str(ROOT / "scripts"))
    try:
        from settlement_calc import check_samples_vs_tccs41
    except Exception as e:
        return JSONResponse({"error": f"settlement_calc: {e}"}, status_code=500)

    out = {}
    for zc in ("NHC", "BXN", "KE"):
        try:
            out[zc] = check_samples_vs_tccs41(zc)
        except Exception as e:
            out[zc] = {"error": str(e)}
    return JSONResponse({"zones": out})


@app.get("/api/sample-check/qc-9403")
async def api_sample_check_qc_9403(n_col: int = 0, n_lab: int = 0, n_fld: int = 0):
    """TCVN 9403:2012 Bảng B.1 — kiểm tra số mẫu QC CDM theo số cột."""
    import sys
    sys.path.insert(0, str(ROOT / "scripts"))
    try:
        from cdm_column_calc import check_qc_adequacy
    except Exception as e:
        return JSONResponse({"error": f"cdm_column_calc: {e}"}, status_code=500)
    if n_col <= 0:
        return JSONResponse({"error": "n_col phải > 0"}, status_code=400)
    try:
        res = check_qc_adequacy(n_col, n_lab, n_fld)
        return JSONResponse({"result": res})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@app.get("/api/sample-check/cdm-tests")
async def api_sample_check_cdm_tests():
    """Kết quả thí nghiệm CDM 7 ngày — qu R7 + E50 R7 từ bảng cdm_tests."""
    con = _db()
    try:
        rows = con.execute("""
            SELECT b.name AS bh, z.code AS zone,
                   ct.group_id, ct.test_id,
                   ct.cement_type, ct.WC_ratio, ct.dosage_kgm3,
                   ct.qu_R7_kPa, ct.E50_R7_MPa
            FROM cdm_tests ct
            JOIN boreholes b ON ct.borehole_id = b.id
            JOIN zones z ON b.zone_id = z.id
            ORDER BY z.code, ct.group_id, ct.test_id
        """).fetchall()
    except Exception as e:
        con.close()
        return JSONResponse({"error": f"cdm_tests: {e}"}, status_code=500)
    con.close()
    data = [dict(r) for r in rows]

    qus = [r["qu_R7_kPa"] for r in data if r.get("qu_R7_kPa") is not None]
    stats = {}
    if qus:
        stats = {
            "n":   len(qus),
            "avg": round(sum(qus) / len(qus), 1),
            "min": round(min(qus), 1),
            "max": round(max(qus), 1),
        }
    return JSONResponse({"rows": data, "stats": stats})
