"""Xuat 2 bai so sanh model (bac 3 va bac 7) ra 1 file Word co hinh anh minh hoa."""
import sys
import io
import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

BASE = os.path.dirname(os.path.abspath(__file__))
CUBIC_DIR = os.path.join(BASE, "model_compare_cubic")
BAC7_DIR = os.path.join(BASE, "model_compare_bac7")
OUT_PATH = os.path.join(BASE, "So_sanh_3_model_AI_Toan_VeBieuDo.docx")

NAVY = RGBColor(0x0B, 0x2E, 0x59)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
GREY = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0xB7, 0x1C, 0x1C)
GREEN = RGBColor(0x1B, 0x5E, 0x20)


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def no_table_border(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "DDDDDD")
        borders.append(el)
    tblPr.append(borders)


def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = NAVY
    return h


def add_para(doc, text, size=11, bold=False, italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_image_row(doc, images_with_captions, width_cm=5.6):
    """images_with_captions: list of (path, caption). Xep ngang tren 1 bang 1 hang."""
    n = len(images_with_captions)
    table = doc.add_table(rows=2, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (img_path, caption) in enumerate(images_with_captions):
        cell_img = table.cell(0, i)
        cell_img.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if img_path and os.path.exists(img_path):
            run = cell_img.paragraphs[0].add_run()
            run.add_picture(img_path, width=Cm(width_cm))
        else:
            cell_img.paragraphs[0].add_run("(không có ảnh)").italic = True

        cell_cap = table.cell(1, i)
        cell_cap.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cell_cap.paragraphs[0].add_run(caption)
        r.font.size = Pt(9.5)
        r.bold = True
    return table


def add_compare_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        r = hdr_cells[i].paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[i], "0B2E59")
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            if val == "✓ Đạt" or (isinstance(val, str) and val.startswith("✓")):
                r.font.color.rgb = GREEN
                r.bold = True
            elif val == "✗ Lỗi" or (isinstance(val, str) and val.startswith("✗")):
                r.font.color.rgb = RED
                r.bold = True
    return table


def add_code_block(doc, code_text, max_lines=None):
    lines = code_text.splitlines()
    if max_lines and len(lines) > max_lines:
        lines = lines[:max_lines] + ["... (rút gọn)"]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    for i, line in enumerate(lines):
        run = p.add_run((line if line else " ") + ("\n" if i < len(lines) - 1 else ""))
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x8A, 0x1F, 0x1F)
    set_cell_shading  # no-op reference to avoid lints
    return p


def read_text(path):
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    return ""


# ============================================================
# DOCUMENT
# ============================================================
doc = Document()
section = doc.sections[0]
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)

# ---- Trang bìa ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("SO SÁNH KHẢ NĂNG TOÁN HỌC & VẼ BIỂU ĐỒ")
r.font.size = Pt(22)
r.bold = True
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Claude  ·  Nemotron 3.5 Lightning  ·  Gemma4:26b")
r.font.size = Pt(14)
r.font.color.rgb = GOLD
sub.paragraph_format.space_after = Pt(4)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Hai bài kiểm tra: phương trình bậc 3 (một lần chạy) và phương trình bậc 7 (vòng lặp tự sửa lỗi)")
r.font.size = Pt(11)
r.italic = True
r.font.color.rgb = GREY

doc.add_paragraph()

# ============================================================
# PHẦN 1 — BẬC 3
# ============================================================
add_heading(doc, "Phần 1 — Phương trình bậc 3 (một lần chạy, không cho sửa lỗi)", level=1)
add_para(doc, "Đề bài:  x³ − 6x² + 11x − 6 = 0", size=12, bold=True)
add_para(doc, "Nghiệm đúng (tính bằng numpy, không đoán): x = 1, 2, 3  —  phân tích (x−1)(x−2)(x−3)=0",
         size=10.5, color=GREY)
add_para(doc, "Mỗi model chỉ được hỏi 1 lần, code sinh ra được chạy nguyên văn — không sửa hộ lỗi nào, "
              "để phản ánh đúng năng lực thật ở lần thử đầu tiên.", size=10, italic=True, color=GREY)

add_heading(doc, "Bảng so sánh", level=2)
add_compare_table(
    doc,
    headers=["Tiêu chí", "Claude", "Nemotron 3.5 Lightning", "Gemma4:26b"],
    rows=[
        ["Nghiệm đúng?", "✓ Đạt", "✓ Đạt", "✓ Đạt (đúng trong văn bản)"],
        ["Code chạy được ngay?", "✓ Đạt", "✓ Đạt", "✗ Lỗi — SyntaxError"],
        ["Ảnh xuất ra", "Đầy đủ, có chú thích", "Đầy đủ, có toạ độ từng nghiệm", "0 ảnh — dừng ở bước biên dịch"],
        ["Nguyên nhân lỗi", "—", "—", "Để sót dòng \"nháp\" tư duy giữa code: "
                                        "\"x**3 - 6*x** import x**2...\" chưa xoá"],
    ],
)

doc.add_paragraph()
add_heading(doc, "Hình minh họa — 3 ảnh xuất ra (đúng nguyên trạng)", level=2)
add_image_row(doc, [
    (os.path.join(CUBIC_DIR, "claude_output.png"), "Claude"),
    (os.path.join(CUBIC_DIR, "nemotron_output.png"), "Nemotron 3.5 Lightning"),
    (os.path.join(CUBIC_DIR, "gemma_output.png"), "Gemma4:26b — thất bại (thẻ ghi lỗi)"),
], width_cm=5.4)

doc.add_page_break()

# ============================================================
# PHẦN 2 — BẬC 7
# ============================================================
add_heading(doc, "Phần 2 — Phương trình bậc 7 (có vòng lặp tự sửa lỗi, đo thời gian)", level=1)
add_para(doc, "Đề bài:  2x⁷ − 4x⁵ + 3x⁴ − x² + 5 = 0", size=12, bold=True)
add_para(doc, "Nghiệm đúng (tính bằng numpy.roots, không đoán): CHỈ 1 nghiệm thực x ≈ −1,715656; "
              "6 nghiệm còn lại là số phức — phép thử khó hơn hẳn bậc 3, dễ lộ ảo giác toán học.",
         size=10.5, color=GREY)
add_para(doc, "Mỗi model được phép lặp tối đa 6 lần: nếu code lỗi, lỗi thật (traceback) được gửi ngược "
              "lại cho model tự sửa, lặp đến khi thành công hoặc hết lượt. Đo thời gian từng vòng và tổng thời gian.",
         size=10, italic=True, color=GREY)

add_heading(doc, "Bảng so sánh", level=2)
add_compare_table(
    doc,
    headers=["Tiêu chí", "Claude", "Nemotron 3.5 Lightning", "Gemma4:26b"],
    rows=[
        ["Nghiệm thực đúng?", "✓ Đạt (−1,715656)", "✓ Đạt — nhưng chỉ đúng ở vòng 2", "✓ Đạt ngay vòng 1"],
        ["Số vòng lặp cần", "1 (viết & chạy trực tiếp)", "2", "1"],
        ["Tổng thời gian", "0,25 giây", "399,6 giây (~6,7 phút)", "183,9 giây (~3,1 phút)"],
        ["Lỗi vòng đầu", "—", "2 lỗi cùng lúc: hệ số sai (dùng nguyên số ví dụ "
                              "để trong comment, giải nhầm phương trình khác) + quên plt.savefig()", "Không có"],
        ["Chất lượng ảnh cuối", "Sạch, rõ ràng", "Sạch, rõ ràng", "Đúng nhưng legend bị chồng chữ"],
    ],
)

doc.add_paragraph()
add_heading(doc, "Điểm đáng chú ý — lỗi \"âm thầm\" của Nemotron ở vòng 1", level=2)
add_para(doc, "Code vòng 1 chạy trơn tru, không báo lỗi cú pháp — nhưng lại giải nhầm phương trình vì "
              "quên thay bộ hệ số mẫu để sẵn trong comment (dành cho một ví dụ bậc 5 khác):", size=10.5)
add_code_block(doc,
    "coeffs = [1, 0, -4, 3, 0, -1]   # <-- replace with your coefficients\n"
    "# (mẫu ví dụ chưa được thay bằng hệ số thật của bài toán bậc 7)\n"
    "...\n"
    "plt.show()   # quên gọi plt.savefig('output.png') như đề yêu cầu")
add_para(doc, "Khi được báo lỗi thật (\"không tạo ra file output.png\"), model không chỉ sửa lỗi được nêu "
              "mà còn tự phát hiện và sửa luôn cả lỗi hệ số sai — dấu hiệu tự-kiểm-chứng vượt ra ngoài "
              "thông báo lỗi được đưa.", size=10.5, italic=True, color=GREY)

doc.add_paragraph()
add_heading(doc, "Hình minh họa — kết quả cuối cùng (sau khi tự sửa xong)", level=2)
add_image_row(doc, [
    (os.path.join(BAC7_DIR, "claude_output.png"), "Claude"),
    (os.path.join(BAC7_DIR, "nemotron_output_final.png"), "Nemotron — sau 2 vòng"),
    (os.path.join(BAC7_DIR, "gemma_output_final.png"), "Gemma4:26b — vòng 1 (legend lỗi layout)"),
], width_cm=5.4)

doc.add_page_break()

# ============================================================
# KẾT LUẬN
# ============================================================
add_heading(doc, "Kết luận chung", level=1)
bullets = [
    "Toán học: cả 3 model đều ra đúng nghiệm cuối cùng ở cả 2 bài — không model nào bịa nghiệm đẹp cho "
    "phương trình bậc 7 dù dễ gây ảo giác.",
    "Độ tin cậy code lần đầu: Claude và Nemotron chạy được ngay ở bậc 3; Gemma lỗi cú pháp ở bậc 3 nhưng "
    "lại đúng ngay lần đầu ở bậc 7 — cho thấy độ ổn định của model nhỏ dao động theo từng lần sinh, "
    "không nên tin tưởng tuyệt đối dù một lần đúng.",
    "Lỗi nguy hiểm nhất không phải là lỗi dừng chương trình (dễ phát hiện) mà là lỗi \"âm thầm\" như "
    "Nemotron vòng 1 — code chạy xong, có kết quả, nhưng sai đề bài. Đây đúng là loại lỗi mà quy tắc "
    "chống hallucination của dự án (CLAUDE.md) muốn ngăn chặn.",
    "Tốc độ: Gemma4:26b nhanh hơn Nemotron khoảng 2 lần trong bài có vòng lặp; Claude không có độ trễ "
    "round-trip API riêng nên nhanh hơn nhiều bậc so với cả hai.",
    "Khuyến nghị áp dụng cho dự án: giữ nguyên kiến trúc hiện tại — Claude đảm nhiệm suy luận kỹ thuật, "
    "sinh/sửa code; Gemma/Nemotron chỉ dùng để diễn giải kết quả đã tính sẵn bằng tiếng Việt, không giao "
    "việc tính toán hoặc sinh code kỹ thuật cho model local.",
]
for b in bullets:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(b)
    r.font.size = Pt(10.5)

doc.save(OUT_PATH)
print(f"Đã xuất file: {OUT_PATH}")
