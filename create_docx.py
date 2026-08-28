import sys
try:
    from docx import Document
    from docx.shared import Inches, Pt
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt

doc = Document()
doc.add_heading('Báo Cáo: Công Nghệ Mã Nguồn Mở Cho Mô Phỏng Đường Sắt Tốc Độ Cao (HSR)', 0)

doc.add_heading('1. Giới Thiệu', level=1)
doc.add_paragraph('Tài liệu này tổng hợp các công cụ mã nguồn mở hàng đầu thế giới được sử dụng trong nghiên cứu và thiết kế hệ thống Đường sắt tốc độ cao (HSR), nhằm thay thế các giải pháp thương mại đắt đỏ.')

doc.add_heading('2. Sơ Đồ Khối Sự Làm Việc (Workflow Kèm BIM)', level=1)
doc.add_paragraph('Sơ đồ dưới đây mô phỏng sự kết hợp toàn diện của hệ thống mã nguồn mở (Dream Team) thông qua thư viện ghép nối đa vật lý preCICE, được tích hợp trong chu trình luân chuyển dữ liệu số của mô hình BIM (BIM Lifecycle):')

try:
    doc.add_picture(r'G:\My Drive\AI-SUC TAI COC THEO DAT NEN\hsr_workflow.png', width=Inches(6.0))
except Exception as e:
    doc.add_paragraph(f'[Không chèn được ảnh, vui lòng chạy lại script sinh ảnh trước. Lỗi: {e}]')

doc.add_heading('3. Tích Hợp Mô Hình BIM (Digital Twin Lifecycle)', level=1)
doc.add_paragraph('Để biến quy trình mô phỏng thành một Bản sao kỹ thuật số (Digital Twin) phục vụ vòng đời dự án, luồng dữ liệu BIM được đưa vào hệ thống theo cơ chế hai chiều (Closed-loop):')

p4 = doc.add_paragraph()
p4.add_run('Từ BIM sang Mô phỏng (Extraction): ').bold = True
p4.add_run('Dữ liệu hình học và vật liệu từ tệp chuẩn IFC được tự động trích xuất bằng các thư viện như IfcOpenShell. Cột trụ bê tông, dầm cầu thép trong BIM sẽ được tự động chuyển đổi thành lưới hình học 3D (STL/STEP) đẩy vào OpenFOAM và Code_Aster làm đầu vào tính toán.')

p5 = doc.add_paragraph()
p5.add_run('Từ Mô phỏng về BIM (Update for Operations): ').bold = True
p5.add_run('Sau khi preCICE điều phối việc giải bài toán Tương tác Đa vật lý, các thông số trọng yếu (như ứng suất cục bộ, độ lún nền, hệ số mỏi...) sẽ được máy tính ghi ngược trở lại tệp IFC. Cụ thể, các thông tin này được đính kèm vào các tập thuộc tính tùy chỉnh (Property Sets - Pset) của từng đối tượng dầm/cột. Điều này cho phép Ban Quản lý Vận hành Đường sắt lập lịch bảo trì dự đoán chính xác điểm có nguy cơ nứt gãy cao nhất.')

doc.add_heading('4. Các Phân Khúc Phần Mềm Hàng Đầu', level=1)
doc.add_heading('Động lực học Đa vật thể (MBD)', level=2)
p1 = doc.add_paragraph()
p1.add_run('Project Chrono (Top 1):').bold = True
p1.add_run(' Nền tảng MBD C++ mạnh nhất hiện nay, hỗ trợ động lực học phương tiện cực tốt.\n')

doc.add_heading('Kết cấu & Hạ tầng (FEA)', level=2)
p2 = doc.add_paragraph()
p2.add_run('Code_Aster (Top 1):').bold = True
p2.add_run(' Vị Vua tuyệt đối. Được tối ưu bởi SNCF cho tương tác Động lực học Xe - Cầu (VBI) và nứt gãy.\n')

doc.add_heading('Khí động học & Thủy lực (CFD)', level=2)
p3 = doc.add_paragraph()
p3.add_run('OpenFOAM (Top 1):').bold = True
p3.add_run(' Tiêu chuẩn công nghiệp. Tối ưu hình dáng khí động học và sức cản gió.\n')

doc.add_heading('5. Cơ Chế Nội Suy Lưới Qua preCICE', level=1)
doc.add_paragraph('Trong hệ thống trên, các phần mềm KHÔNG dùng chung lưới phần tử hữu hạn. Sự kỳ diệu nằm ở preCICE: Công cụ này hoạt động như một cỗ máy nội suy toán học (Mapping Engine). Nó tự động phân bổ lực từ lưới OpenFOAM lên lưới của Code_Aster một cách trơn tru, đảm bảo định luật bảo toàn năng lượng mà không ép các kỹ sư phải chia lưới trùng khớp (Non-matching meshes).')

doc.add_heading('6. Ứng Dụng Đám Mây Điểm (Scan to BIM)', level=1)
doc.add_paragraph('Việc tích hợp dữ liệu quét 3D (Point Cloud) mang lại giá trị cốt lõi để biến hệ thống thành Bản sao kỹ thuật số (Digital Twin) thực thụ. Thay vì sử dụng bản vẽ thiết kế (As-designed), hệ thống mô phỏng trực tiếp trên bề mặt thực tế đã bị lún, võng (As-is).')
p6 = doc.add_paragraph()
p6.add_run('Các công cụ xử lý thủ công: ').bold = True
p6.add_run('CloudCompare (Xử lý mây điểm, bọc lưới Poisson), MeshLab (Vá lỗ thủng bề mặt).')
p7 = doc.add_paragraph()
p7.add_run('Thư viện tự động hóa (API): ').bold = True
p7.add_run('Open3D / PCL (Lọc nhiễu, tái tạo bề mặt), IfcOpenShell (Đóng gói IFC).')
p8 = doc.add_paragraph()
p8.add_run('Chia lưới thể tích (Volumetric Meshing): ').bold = True
p8.add_run('Gmsh (Đổ khối tứ diện cho Code_Aster), SnappyHexMesh (Chia lưới lục diện cho OpenFOAM).')

doc.save(r'G:\My Drive\AI-SUC TAI COC THEO DAT NEN\Bao_Cao_Mo_Phong_HSR.docx')
