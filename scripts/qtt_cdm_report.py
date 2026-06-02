"""
qtt_cdm_report.py — Xuất báo cáo Word "Phân tích lựa chọn CDM cho QTT".

Bố cục Word docx:
  1. Trang bìa + thông số đầu vào
  2. Cơ sở lý thuyết (S1, S2, Bjerrum, smoothness)
  3. Thuật toán chi tiết (pseudo-code 5 bước)
  4. Tiêu chí thiết kế đã chọn
  5. Lc tối ưu cho mỗi HK + heatmap + bar chart
  6. S(Lc) curves cho 6 HK + 4 ngưỡng ΔS
  7. Kiểm tra độ bằng phẳng pair-wise
  8. Phân vùng grid + Lc heatmap + zoning map
  9. Kết luận và khuyến nghị

Charts dùng matplotlib (PNG embed) — kaleido-free, hoạt động trên mọi môi trường.

Tuân thủ memory:
  - feedback-no-emoji
  - feedback-bold-table-headers
  - feedback-report-typography (12pt + zebra)
  - feedback-show-values-on-charts (label tại điểm dữ liệu)
  - feedback-table-style-ui-docx-parity
"""
from __future__ import annotations
import io
from datetime import datetime
from pathlib import Path
from typing import Optional

_ROOT = Path(__file__).resolve().parent.parent


_ROAD_CLASS_DESC = {
    "cao_toc": "Cao tốc / ≥ 80 km/h / cấp cao A1",
    "cap_I_IV": "Đường ô tô cấp I–IV",
}
_STRUCT_DESC = {"cau": "Cầu", "cong": "Cống"}

# ──────────────────────────────────────────────────────────────────────
# Helpers — Word
# ──────────────────────────────────────────────────────────────────────

def _set_normal_style(doc) -> None:
    """Body default size = 12pt — memory feedback-report-typography."""
    from docx.shared import Pt
    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = "Times New Roman"


def _add_heading(doc, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"


def _add_para_kv(doc, key: str, val: str) -> None:
    from docx.shared import Pt
    p = doc.add_paragraph()
    r1 = p.add_run(f"{key}: "); r1.bold = True; r1.font.size = Pt(12)
    r2 = p.add_run(val); r2.font.size = Pt(12)


def _add_para(doc, text: str, italic: bool = False) -> None:
    from docx.shared import Pt
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(12); r.italic = italic


def _add_data_table(doc, columns: list[str], rows: list[list[str]]) -> None:
    """Thêm bảng với header bold + zebra — dùng helper chung."""
    import sys as _sys
    _sys.path.insert(0, str(_ROOT / "scripts"))
    from core.report_style import style_tbl_for_docx
    tbl = doc.add_table(rows=1, cols=len(columns))
    tbl.style = "Table Grid"
    for j, c in enumerate(columns):
        tbl.cell(0, j).text = c
    for row in rows:
        cells = tbl.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = str(v)
    style_tbl_for_docx(tbl)


def _add_picture(doc, png_buf: io.BytesIO, width_cm: float = 15.0) -> None:
    from docx.shared import Cm
    png_buf.seek(0)
    doc.add_picture(png_buf, width=Cm(width_cm))


# ──────────────────────────────────────────────────────────────────────
# Helpers — Matplotlib charts
# ──────────────────────────────────────────────────────────────────────

def _setup_mpl():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    plt.rcParams["font.family"] = ["DejaVu Sans"]
    plt.rcParams["font.size"] = 11
    plt.rcParams["axes.titlesize"] = 12
    plt.rcParams["axes.labelsize"] = 11
    return plt


def _formula_png(latex_lines: list[tuple[str, str]],
                  title: str = "", height_cm: float = 4.0) -> io.BytesIO:
    """Render danh sách công thức LaTeX (mathtext) → PNG.

    latex_lines: list of (label, latex_expression)
    """
    plt = _setup_mpl()
    fig, ax = plt.subplots(figsize=(13, height_cm))
    ax.axis("off")
    if title:
        ax.set_title(title, fontsize=13, weight="bold", loc="left")
    n = len(latex_lines)
    y_step = 0.85 / max(n, 1)
    y = 0.88 - y_step / 2
    for label, latex_expr in latex_lines:
        if label:
            ax.text(0.02, y, label, fontsize=11, va="center", weight="bold")
        ax.text(0.25, y, latex_expr, fontsize=14, va="center")
        y -= y_step
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=160, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_heatmap_lc(lc_matrix: list[dict], dS_list: list[float]) -> io.BytesIO:
    plt = _setup_mpl()
    fig, ax = plt.subplots(figsize=(10, 5))
    names = [h["name"] for h in lc_matrix]
    Z = [[h["by_dS"].get(d, {}).get("Lc_m") for d in dS_list] for h in lc_matrix]
    Zn = [[v if v is not None else float("nan") for v in r] for r in Z]
    import numpy as np
    arr = np.array(Zn, dtype=float)
    masked = np.ma.masked_invalid(arr)
    im = ax.imshow(masked, aspect="auto", cmap="YlOrRd", origin="upper")
    ax.set_xticks(range(len(dS_list)))
    ax.set_xticklabels([f"ΔS ≤ {int(d)} cm" for d in dS_list], fontsize=11)
    ax.set_yticks(range(len(names)))
    ax.set_yticklabels(names, fontsize=11)
    for i, row in enumerate(Z):
        for j, v in enumerate(row):
            txt = f"{v:.1f}" if v is not None else "—"
            color = "white" if (v is not None and v > 20) else "black"
            ax.text(j, i, txt, ha="center", va="center",
                     fontsize=11, color=color, weight="bold")
    cb = plt.colorbar(im, ax=ax, fraction=0.04)
    cb.set_label("Lc (m)", fontsize=11)
    ax.set_title("Heatmap Lc tối ưu — HK ND × ΔS", fontsize=12, weight="bold")
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=160, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_bar_lc(lc_matrix: list[dict], dS_list: list[float]) -> io.BytesIO:
    plt = _setup_mpl()
    fig, ax = plt.subplots(figsize=(11, 5))
    import numpy as np
    names = [h["name"] for h in lc_matrix]
    x = np.arange(len(names))
    width = 0.2
    colors = {10: "#991b1b", 20: "#dc2626", 30: "#ea580c", 40: "#f59e0b"}
    for k, dS in enumerate(dS_list):
        ys = []
        for h in lc_matrix:
            r = h["by_dS"].get(dS, {})
            ys.append(r["Lc_m"] if r.get("ok") and r.get("Lc_m") else 0)
        bars = ax.bar(x + (k - (len(dS_list) - 1) / 2) * width, ys, width,
                       label=f"ΔS ≤ {int(dS)} cm",
                       color=colors.get(int(dS), "#666"))
        for bar, v in zip(bars, ys):
            if v > 0:
                ax.text(bar.get_x() + bar.get_width() / 2, v + 0.3,
                         f"{v:.1f}", ha="center", va="bottom",
                         fontsize=8, color=colors.get(int(dS), "#666"))
    ax.set_xticks(x); ax.set_xticklabels(names, fontsize=11)
    ax.set_ylabel("Lc (m)", fontsize=11)
    ax.set_title("So sánh Lc giữa các HK theo 4 ΔS", fontsize=12, weight="bold")
    ax.legend(fontsize=10, loc="upper left", ncol=4)
    ax.grid(axis="y", alpha=0.3)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=160, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_s_vs_lc(curves_data: dict) -> Optional[io.BytesIO]:
    if not curves_data or not curves_data.get("curves"):
        return None
    plt = _setup_mpl()
    fig, ax = plt.subplots(figsize=(11, 6))
    palette = {"ND-02": "#1d4ed8", "ND-03": "#dc2626", "ND-04": "#ea580c",
               "ND-05": "#16a34a", "ND-06": "#7c3aed", "ND-07": "#0891b2"}
    for hk_name, cd in curves_data["curves"].items():
        pts = cd.get("points") or []
        if not pts:
            continue
        xs = [pt["Lc_m"] for pt in pts]
        ys = [pt["S_total_cm"] for pt in pts]
        label = hk_name
        if cd.get("borrowed"):
            label += f" (Cc←{cd['cc_source']})"
        ax.plot(xs, ys, "-o", color=palette.get(hk_name, "#666"),
                 label=label, lw=2, ms=5)
        # label số tại Lc = 10, 15, 20, 25, 30 + cuối
        mark_lc = {10, 15, 20, 25, 30}
        for i, (xv, yv) in enumerate(zip(xs, ys)):
            if i == len(xs) - 1 or any(abs(xv - m) < 0.6 for m in mark_lc):
                ax.text(xv, yv + 1.5, f"{yv:.0f}",
                         fontsize=8, color=palette.get(hk_name, "#666"),
                         ha="center")
    # 4 ngưỡng ΔS
    for dS, col in [(10, "#991b1b"), (20, "#dc2626"),
                     (30, "#ea580c"), (40, "#f59e0b")]:
        ax.axhline(y=dS, ls="--", color=col, lw=1.4, alpha=0.8)
        ax.text(0.5, dS + 1, f"ΔS ≤ {dS} cm",
                 color=col, fontsize=9, weight="bold")
    ax.set_xlabel("Lc (m)", fontsize=11)
    ax.set_ylabel("S_total = S₁ + S₂ (cm)", fontsize=11)
    ax.set_title("Đường cong S(Lc) — quan hệ độ lún vs chiều dài cọc CDM",
                  fontsize=12, weight="bold")
    ax.set_ylim(bottom=0)
    ax.legend(fontsize=9, loc="upper right", ncol=2)
    ax.grid(alpha=0.3)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=160, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_smoothness_scatter(
    smoothness_pairs: list[dict], i_inv_eff: int,
) -> Optional[io.BytesIO]:
    if not smoothness_pairs:
        return None
    plt = _setup_mpl()
    fig, ax = plt.subplots(figsize=(10, 6))
    for p in smoothness_pairs:
        col = "#16a34a" if p["ok"] else "#dc2626"
        d = p["d_m"]; dS = p["dS_m"] * 100  # m → cm
        ax.scatter(d, dS, color=col, s=80, edgecolor="black", zorder=3)
        ax.annotate(f"{p['i']}↔{p['j']}\n1/{p['i_inv_actual']}"
                     if p['i_inv_actual'] else f"{p['i']}↔{p['j']}",
                     xy=(d, dS), xytext=(5, 5), textcoords="offset points",
                     fontsize=8)
    # Đường giới hạn i = 1/i_inv_eff (tức dS = d / i_inv_eff × 100)
    import numpy as np
    d_range = np.linspace(1, max(p["d_m"] for p in smoothness_pairs) * 1.1, 100)
    dS_lim = d_range / i_inv_eff * 100  # cm
    ax.plot(d_range, dS_lim, "--", color="#dc2626", lw=1.5,
             label=f"i = 1/{i_inv_eff} (giới hạn)")
    ax.fill_between(d_range, 0, dS_lim, alpha=0.1, color="#16a34a",
                     label="Vùng ĐẠT")
    ax.set_xlabel("Khoảng cách d giữa cặp HK (m)", fontsize=11)
    ax.set_ylabel("Chênh lún ΔS (cm)", fontsize=11)
    ax.set_title(f"Kiểm tra độ bằng phẳng — i ≤ 1/{i_inv_eff}",
                  fontsize=12, weight="bold")
    ax.legend(fontsize=10)
    ax.grid(alpha=0.3)
    ax.set_xlim(left=0); ax.set_ylim(bottom=0)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=160, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_grid_lc_map(
    grid_points: list[dict], hks_with_S: list[dict],
) -> Optional[io.BytesIO]:
    if not grid_points:
        return None
    plt = _setup_mpl()
    fig, ax = plt.subplots(figsize=(11, 8))
    import numpy as np
    # Build grid
    es = sorted({p["E"] for p in grid_points})
    ns = sorted({p["N"] for p in grid_points})
    e2i = {e: i for i, e in enumerate(es)}
    n2i = {n: i for i, n in enumerate(ns)}
    Z = np.full((len(ns), len(es)), np.nan)
    for gp in grid_points:
        if gp.get("Lc_m") is None or not gp.get("ok"):
            continue
        Z[n2i[gp["N"]], e2i[gp["E"]]] = gp["Lc_m"]
    masked = np.ma.masked_invalid(Z)
    im = ax.imshow(masked, aspect="equal", cmap="YlOrRd",
                    origin="lower", extent=[es[0], es[-1], ns[0], ns[-1]])
    cb = plt.colorbar(im, ax=ax, fraction=0.04)
    cb.set_label("Lc (m)", fontsize=11)
    # HK markers
    for hs in hks_with_S:
        ax.plot(hs["E"], hs["N"], "D", color="white",
                 markersize=12, markeredgecolor="black", markeredgewidth=2,
                 zorder=5)
        ax.annotate(hs["name"], xy=(hs["E"], hs["N"]),
                     xytext=(0, 12), textcoords="offset points",
                     ha="center", fontsize=10, weight="bold")
    ax.set_xlabel("Easting (m)", fontsize=11)
    ax.set_ylabel("Northing (m)", fontsize=11)
    ax.set_title("Bản đồ Lc grid 162 điểm 20m × 20m trong polygon QTT",
                  fontsize=12, weight="bold")
    ax.grid(alpha=0.3)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=160, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_stress_profile(
    bh_name: str, q_kPa: float,
    tip_depth_m: Optional[float] = None,
    db_path: Optional[Path] = None,
) -> Optional[io.BytesIO]:
    """Biểu đồ profile ứng suất σ'v0 vs Δσ vs vị trí Δσ/σ'v0 = 10%.

    - σ'v0(z) tính từ layers + γ_eff per symbol
    - Δσ = q kPa (constant — Terzaghi 1D giả định tải vô hạn)
    - Đường 10% σ'v0 — engine dừng tính S2 khi Δσ < 10% σ'v0
    - Hiển thị tên lớp đất (symbol) theo độ sâu
    - Đánh dấu mũi cọc CDM tại tip_depth_m
    """
    import sqlite3
    GAMMA_W = 9.81
    p = Path(db_path) if db_path else _ROOT / "data" / "TTHC.sqlite"

    with sqlite3.connect(p) as con:
        con.row_factory = sqlite3.Row
        layers = con.execute("""
            SELECT depth_top_m, depth_bot_m, symbol
            FROM layers
            WHERE borehole_id = (SELECT id FROM boreholes WHERE name=?)
            ORDER BY depth_top_m
        """, (bh_name,)).fetchall()
        gamma_rows = con.execute("""
            SELECT lt.symbol_tcvn AS sym, AVG(lt.gamma_kNm3) AS g
            FROM lab_tests lt
            JOIN boreholes b ON lt.borehole_id = b.id
            WHERE b.name=? AND lt.gamma_kNm3 > 0
            GROUP BY lt.symbol_tcvn
        """, (bh_name,)).fetchall()
    if not layers:
        return None

    gamma_by_sym = {r["sym"]: float(r["g"]) for r in gamma_rows}
    # Mặc định γ_sat theo symbol khi không có lab
    GAMMA_DEFAULT = {
        "F": 18.0, "1": 15.5, "1b": 16.0, "2": 16.5, "3": 18.5,
        "4": 19.0, "5": 19.5, "XMD": 15.0,
    }
    # Phân loại sét/cát theo SAND_SYMBOLS_S2 (settlement_calc.py)
    SAND_SYMBOLS_S2 = {"F", "2a", "2b", "2c", "3a", "3b", "3c",
                        "4", "5", "5a", "5b", "6", "7", "8"}

    def _method_for_sym(s: str) -> str:
        return "SPT (cát)" if s in SAND_SYMBOLS_S2 else "Terzaghi (sét)"

    # Build profile σ'v0 (z)
    depths_z = [0.0]
    sigmas = [0.0]
    layer_meta = []  # list of {top, bot, symbol, gamma_eff}
    for r in layers:
        top = float(r["depth_top_m"])
        bot = float(r["depth_bot_m"])
        sym = r["symbol"] or "?"
        g_tot = gamma_by_sym.get(sym, GAMMA_DEFAULT.get(sym, 16.0))
        g_eff = max(0.5, g_tot - GAMMA_W)  # bão hoà / dưới MNN
        layer_meta.append({
            "top": top, "bot": bot, "symbol": sym, "gamma_eff": g_eff,
        })
        z = depths_z[-1]
        sig = sigmas[-1]
        while z < bot - 1e-6:
            dz = min(0.5, bot - z)
            z += dz
            sig += g_eff * dz
            depths_z.append(z)
            sigmas.append(sig)

    # Tìm độ sâu z* tại Δσ/σ'v0 = 10% → σ'v0 = 10×q
    sigma_10pct_target = 10.0 * q_kPa
    z_cutoff = None
    for i in range(1, len(sigmas)):
        if sigmas[i] >= sigma_10pct_target:
            # linear interpolate
            f = (sigma_10pct_target - sigmas[i - 1]) / (sigmas[i] - sigmas[i - 1])
            z_cutoff = depths_z[i - 1] + f * (depths_z[i] - depths_z[i - 1])
            break

    plt = _setup_mpl()
    fig, ax = plt.subplots(figsize=(11, 8))

    # Layer background bands + symbol labels (sét warm / cát cool)
    LAYER_COLORS = {
        # Sét — warm tones
        "1": "#fca5a5", "1b": "#fb923c", "2": "#fbbf24",
        "XMD": "#fecaca", "CH": "#fca5a5", "MH": "#fb923c",
        "3": "#fde68a",
        # Cát — cool tones
        "F": "#bfdbfe", "2a": "#bfdbfe", "2b": "#93c5fd",
        "2c": "#60a5fa", "3a": "#a7f3d0", "3b": "#86efac",
        "3c": "#6ee7b7", "4": "#34d399", "5": "#10b981",
        "5a": "#10b981", "5b": "#059669", "6": "#047857",
        "7": "#065f46", "8": "#064e3b",
    }
    x_max = max(sigmas[-1] * 1.15, q_kPa * 12)
    for lm in layer_meta:
        sym = lm["symbol"]
        is_sand = sym in SAND_SYMBOLS_S2
        method = _method_for_sym(sym)
        ax.axhspan(
            lm["top"], lm["bot"],
            alpha=0.30,
            color=LAYER_COLORS.get(sym, "#e5e7eb"),
            zorder=1,
        )
        ax.text(
            x_max * 0.96, (lm["top"] + lm["bot"]) / 2,
            f"Lớp {sym} ({'cát' if is_sand else 'sét'})\n"
            f"γ'={lm['gamma_eff']:.1f}\n"
            f"S₂: {method}",
            ha="right", va="center",
            fontsize=8, color="#1f2937",
            bbox=dict(boxstyle="round,pad=0.25",
                      facecolor="white", alpha=0.85,
                      edgecolor="#3b82f6" if is_sand else "#dc2626",
                      linewidth=1.2),
        )
    # σ'v0 profile
    ax.plot(sigmas, depths_z, "-", color="#1d4ed8", lw=2.5,
             label="σ'v0 (ứng suất bản thân hữu hiệu)", zorder=3)
    # 10% σ'v0
    sigmas_10pct = [s * 0.10 for s in sigmas]
    ax.plot(sigmas_10pct, depths_z, "--", color="#ea580c", lw=2,
             label="10% × σ'v0 (ngưỡng dừng tính S₂)", zorder=3)
    # Δσ = q
    ax.axvline(x=q_kPa, color="#dc2626", lw=2.5,
                label=f"Δσ = q = {q_kPa:.1f} kPa (tải gây lún)", zorder=3)
    # Vị trí cutoff
    if z_cutoff:
        ax.axhline(y=z_cutoff, color="#16a34a", lw=2, ls="-.",
                    label=f"Δσ/σ'v0 = 10% tại z = {z_cutoff:.1f} m",
                    zorder=4)
        ax.scatter([q_kPa], [z_cutoff], color="#16a34a", s=120,
                    zorder=5, marker="X", edgecolor="black")
        ax.annotate(
            f"Engine dừng tính S₂\ntại z = {z_cutoff:.1f} m",
            xy=(q_kPa, z_cutoff),
            xytext=(q_kPa * 2.5, z_cutoff - 4),
            fontsize=10, color="#16a34a", weight="bold",
            arrowprops=dict(arrowstyle="->", color="#16a34a", lw=1.5),
        )
    # Tip CDM
    if tip_depth_m:
        ax.axhline(y=tip_depth_m, color="#7c3aed", lw=2,
                    label=f"Mũi cọc CDM tại z = {tip_depth_m:.1f} m",
                    zorder=4)

    ax.set_xlabel("Ứng suất (kPa)", fontsize=11)
    ax.set_ylabel("Độ sâu z từ mặt tự nhiên (m)", fontsize=11)
    ax.set_title(
        f"Profile ứng suất {bh_name} — so sánh σ'v0 vs Δσ + vị trí 10%",
        fontsize=12, weight="bold",
    )
    ax.invert_yaxis()
    ax.set_xlim(left=0, right=x_max)
    ax.set_ylim(top=0, bottom=max(depths_z) + 2)
    ax.legend(loc="lower right", fontsize=9, framealpha=0.95)
    ax.grid(alpha=0.3, zorder=0)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=160, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_zoning_map(
    grid_points: list[dict], zoning: dict, hks_with_S: list[dict],
) -> Optional[io.BytesIO]:
    if not grid_points or not zoning.get("stats"):
        return None
    plt = _setup_mpl()
    fig, ax = plt.subplots(figsize=(11, 8))
    import numpy as np
    es = sorted({p["E"] for p in grid_points})
    ns = sorted({p["N"] for p in grid_points})
    e2i = {e: i for i, e in enumerate(es)}
    n2i = {n: i for i, n in enumerate(ns)}
    Z = np.full((len(ns), len(es)), np.nan)
    for gp in grid_points:
        z = zoning["assignment"].get((gp["E"], gp["N"]))
        if z is None:
            continue
        Z[n2i[gp["N"]], e2i[gp["E"]]] = z
    masked = np.ma.masked_invalid(Z)
    n_zones = len(zoning["stats"])
    colors = ["#1565C0", "#D32F2F", "#2E7D32", "#F57F17", "#6A1B9A", "#00838F"]
    from matplotlib.colors import ListedColormap
    cmap = ListedColormap(colors[:max(n_zones, 1)])
    im = ax.imshow(masked, aspect="equal", cmap=cmap,
                    origin="lower", extent=[es[0], es[-1], ns[0], ns[-1]],
                    vmin=0, vmax=max(n_zones - 1, 0))
    cb = plt.colorbar(im, ax=ax, fraction=0.04, ticks=range(n_zones))
    cb.set_label("Vùng thiết kế", fontsize=11)
    cb.set_ticklabels([f"P{s['zone_id']+1}: Lc={s['Lc_design']:.1f}m"
                        for s in zoning["stats"]])
    for hs in hks_with_S:
        ax.plot(hs["E"], hs["N"], "D", color="white",
                 markersize=12, markeredgecolor="black", markeredgewidth=2)
        ax.annotate(hs["name"], xy=(hs["E"], hs["N"]),
                     xytext=(0, 12), textcoords="offset points",
                     ha="center", fontsize=10, weight="bold")
    ax.set_xlabel("Easting (m)", fontsize=11)
    ax.set_ylabel("Northing (m)", fontsize=11)
    ax.set_title(f"Phân vùng thiết kế CDM — {n_zones} vùng theo Lc",
                  fontsize=12, weight="bold")
    ax.grid(alpha=0.3)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=160, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


# ──────────────────────────────────────────────────────────────────────
# Public API
# ──────────────────────────────────────────────────────────────────────

def build_qtt_decision_docx(
    meta: dict,
    lc_matrix: list[dict],
    criteria: dict,
    smoothness_pairs: list[dict],
    zoning: dict,
    grid_meta: dict,
    curves_data: Optional[dict] = None,
    grid_points: Optional[list[dict]] = None,
    hks_with_S: Optional[list[dict]] = None,
) -> bytes:
    """Tạo bytes Word docx báo cáo quyết định thiết kế CDM QTT (đầy đủ)."""
    import sys as _sys
    _sys.path.insert(0, str(_ROOT / "scripts"))
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    _set_normal_style(doc)
    dS_list_default = [10.0, 15.0, 20.0, 25.0, 30.0, 40.0]
    # Tự suy ra dS_list từ matrix nếu có
    dS_list = (
        meta.get("delta_S_values_cm")
        or (sorted(lc_matrix[0]["by_dS"].keys()) if lc_matrix else dS_list_default)
    )

    # ── Trang bìa ─────────────────────────────────────────────────
    _add_heading(
        doc,
        "BÁO CÁO PHÂN TÍCH LỰA CHỌN CHIỀU DÀI CỌC CDM "
        "KHU VỰC QUẢNG TRƯỜNG TRUNG TÂM (QTT)",
        level=0,
    )
    p = doc.add_paragraph()
    r = p.add_run("Trung tâm Hành chính TP.HCM — Mã số 202605-TTHC")
    r.bold = True; r.font.size = Pt(13)
    _add_para_kv(doc, "Ngày lập", datetime.now().strftime("%d/%m/%Y"))
    _add_para_kv(
        doc, "Tiêu chuẩn áp dụng",
        "TCCS 41:2022 (Điều 6.2.3 Bảng 1, Phụ lục E.1) · "
        "TCVN 9403:2012 (Phụ lục C, Phụ lục B)",
    )

    # ── 1. Thông số đầu vào ────────────────────────────────────────
    _add_heading(doc, "1. Thông số đầu vào", level=1)
    _add_para_kv(doc, "Tải phân bố q",
                  f"{meta.get('q_kPa', 0):.1f} kPa (cấu tạo Áo đường + He + Hse)")
    _add_para_kv(doc, "Tổng chiều dày các lớp trên đỉnh CDM",
                  f"{meta.get('fill_thickness_m', 0):.1f} m")
    _add_para_kv(doc, "Đường kính cọc D", f"{meta.get('D_mm', 0):.0f} mm")
    _add_para_kv(doc, "Khoảng cách cọc s",
                  f"{meta.get('spacing_m', 0):.2f} m  ({meta.get('pattern', '')})")
    _add_para_kv(doc, "Tỷ lệ diện tích thay thế a",
                  f"{meta.get('a', 0):.4f}")
    _add_para_kv(doc, "Mô đun trụ Ec",
                  f"{meta.get('Ec_kPa', 0):,.0f} kPa  "
                  f"(k = {meta.get('Ec_factor', 0):.0f}, "
                  f"qu = {meta.get('qu_kPa', 0):.0f} kPa)")
    _add_para(doc,
        "Cao độ đỉnh CDM = cao độ thiết kế − tổng chiều dày các lớp đắp = "
        f"design − {meta.get('fill_thickness_m', 0):.1f} m. "
        "Tải trọng q không đổi giữa các HK — đảm bảo Δσ đồng nhất trên đỉnh CDM."
    )

    # ── 2. Cơ sở lý thuyết ────────────────────────────────────────
    _add_heading(doc, "2. Cơ sở lý thuyết", level=1)

    _add_heading(doc, "2.1 Lún khối gia cố S₁ (TCVN 9403 Phụ lục C)", level=2)
    _add_para(doc,
        "Lún đàn hồi tức thì của khối đất gia cố CDM dưới tải phân bố q:"
    )
    img = _formula_png([
        ("Công thức C.2:",
         r"$S_1 = \dfrac{q \cdot H_{gc}}{a \cdot E_c + (1-a) \cdot E_s} \times 100\quad[\text{cm}]$"),
        ("Mô đun trụ:",
         r"$E_c = k \cdot q_u / 2$,$\quad k = 50 \div 100$ (TCVN 9403 B.5.1)"),
        ("Mô đun đất:",
         r"$E_s = 250 \cdot c_u$, với$\quad c_u = \mu \cdot S_u$ (Bjerrum, TCCS 41 C.5)"),
        ("Tỷ lệ thay thế:",
         r"$a = \pi(D/2)^2 / s^2$ (bố trí ô vuông)"),
    ], title="Lún S₁ — khối gia cố CDM", height_cm=4.2)
    _add_picture(doc, img, width_cm=15.5)

    _add_heading(doc, "2.2 Lún cố kết S₂ dưới mũi cọc (Terzaghi 1D)", level=2)
    _add_para(doc,
        "Lún cố kết phân tử của lớp đất yếu CÒN LẠI bên dưới mũi cọc CDM. "
        "Mỗi phân tố h_i = 2 m, dừng khi Δσ/σ'v₀ < 10%."
    )
    img = _formula_png([
        ("OC (σ'_vf ≤ PC):",
         r"$S_{2,OC} = \dfrac{h_i \cdot C_s}{1+e_0} \log_{10}\dfrac{\sigma'_{vf}}{\sigma'_{v0}}$"),
        ("NC (σ'_v0 ≥ PC):",
         r"$S_{2,NC} = \dfrac{h_i \cdot C_c}{1+e_0} \log_{10}\dfrac{\sigma'_{vf}}{\sigma'_{v0}}$"),
        ("Cross-PC:",
         r"$S_{2,cross} = \dfrac{h_i}{1+e_0}\left[C_s \log_{10}\dfrac{P_C}{\sigma'_{v0}} + C_c \log_{10}\dfrac{\sigma'_{vf}}{P_C}\right]$"),
        ("Lún còn lại 15 năm:",
         r"$\Delta S = S_c (1 - U_t)$,$\quad U_t = $ độ cố kết tại t = 15 năm"),
    ], title="Lún cố kết S₂ — Terzaghi 1D phân tố", height_cm=4.5)
    _add_picture(doc, img, width_cm=15.5)

    _add_heading(doc, "2.3 Hiệu chỉnh Bjerrum cho Su VST", level=2)
    _add_para(doc,
        "Cường độ kháng cắt tính toán (TCCS 41 C.5):"
    )
    img = _formula_png([
        ("Công thức C.5:",
         r"$c_u = \mu \cdot S_u$, với$\quad \mu = f(I_p)$ tra Bảng C.1"),
        ("Bảng C.1:",
         r"$I_p$: 10 / 20 / 30 / 40 / 50 / 60 / 70%"),
        ("",
         r"$\mu$: 1.09 / 1.00 / 0.925 / 0.86 / 0.80 / 0.75 / 0.70"),
        ("Quy tắc:",
         r"$\mu$ chỉ áp cho Su VST. KHÔNG áp cho Cu_UU lab."),
    ], title="Hiệu chỉnh Bjerrum μ theo Ip", height_cm=4.5)
    _add_picture(doc, img, width_cm=15.5)

    _add_heading(doc, "2.4 Độ bằng phẳng theo Bảng E.1 (TCCS 41 Phụ lục E)", level=2)
    _add_para(doc,
        "Độ dốc dọc giữa các điểm sau khi cố kết. Yêu cầu i ≤ giới hạn tra theo "
        "(cấp đường × công trình × tốc độ thiết kế). Cho phép vồng trước i ≤ 1/125."
    )
    img = _formula_png([
        ("Định nghĩa:",
         r"$i = \dfrac{|\Delta S_i - \Delta S_j|}{d_{ij}}$  với cặp (i,j) cách d_{ij}"),
        ("Yêu cầu:",
         r"$i \leq i_{max}$, vd. $i_{max} = 1/200$ cho cao tốc cầu v=80 km/h"),
        ("Vồng:",
         r"$i_{eff} = \max(i_{table}, 1/125)$ khi cho phép tạo vồng trước"),
    ], title="Tiêu chí độ bằng phẳng — Bảng E.1", height_cm=3.8)
    _add_picture(doc, img, width_cm=15.5)

    # ── 3. Thuật toán chi tiết ────────────────────────────────────
    _add_heading(doc, "3. Thuật toán xác định Lc tối ưu", level=1)

    _add_heading(doc, "3.1 Tổng quan", level=2)
    _add_para(doc,
        "Thuật toán 5 bước: (1) Tính Lc cho mỗi HK × mỗi ΔS bằng find_cdm_length; "
        "(2) Tra giới hạn i từ Bảng E.1; (3) Kiểm tra độ bằng phẳng pair-wise; "
        "(4) Tính Lc grid 162 điểm 20m×20m; (5) Phân vùng quantile binning."
    )

    _add_heading(doc, "3.2 Pseudo-code", level=2)
    pseudo = (
        "for each HK i in 6 ND boreholes:\n"
        "    nat_i      = boreholes.elevation_m\n"
        "    design_i   = qtt_elevation_points.elev_des_m  (nearest grid)\n"
        "    cdm_top_i  = design_i − 1.9 m\n"
        "    clay_top_i, H_soft_i  = soft_profile_from_db(HK_i)\n"
        "    Su_i, source = nearest VST or lab UU\n"
        "    mu_i      = bjerrum_mu(Ip_avg)\n"
        "    Es_i      = 250 × mu_i × Su_i\n"
        "    if HK_i thiếu Cc → cc_source = nearest ND with Cc (§15)\n"
        "    for each ΔS in {10, 20, 30, 40}:\n"
        "        p* = find_cdm_length(cc_source, q=40.8, a, Ec, Su_i, ΔS, mu_i)\n"
        "        Lc_i(ΔS) = (clay_top_i + p*) − (nat_i − cdm_top_i)\n"
        "\n"
        "i_inv_eff = get_smoothness_i_inv(road_class, structure, speed)\n"
        "if allow_vong: i_inv_eff = max(i_inv_eff, 125)\n"
        "\n"
        "for each pair (i, j) in HK_list:\n"
        "    d_ij  = sqrt((E_i−E_j)² + (N_i−N_j)²)\n"
        "    ΔS_ij = |S_i − S_j| / 100  [m]\n"
        "    i_inv = d_ij / ΔS_ij\n"
        "    ok    = i_inv >= i_inv_eff\n"
        "\n"
        "for each grid point (E, N) in 162 qtt_elevation_points:\n"
        "    ref_hk   = nearest ND HK with Cc\n"
        "    fill_loc = max(0, elev_des − elev_nat)\n"
        "    Lc_grid  = find_cdm_length(ref_hk, ΔS_target) using local geometry\n"
        "\n"
        "zones = quantile_binning(Lc_grid, n_zones)\n"
        "Lc_design[zone] = max(Lc_grid in zone)   # an toàn\n"
    )
    p_pseudo = doc.add_paragraph()
    r_pseudo = p_pseudo.add_run(pseudo)
    r_pseudo.font.name = "Consolas"; r_pseudo.font.size = Pt(9)

    # ── 4. Tiêu chí thiết kế đã chọn ───────────────────────────────
    _add_heading(doc, "4. Tiêu chí thiết kế đã chọn", level=1)
    _add_para_kv(doc, "Cấp đường",
                  _ROAD_CLASS_DESC.get(criteria["road_class"], criteria["road_class"]))
    _add_para_kv(doc, "Công trình",
                  _STRUCT_DESC.get(criteria["structure"], criteria["structure"]))
    _add_para_kv(doc, "Tốc độ thiết kế", f"{criteria['speed_kmh']} km/h")
    _add_para_kv(doc, "Độ lún cho phép ΔS",
                  f"≤ {int(criteria['design_dS_cm'])} cm (TCCS 41 Bảng 1)")
    _add_para_kv(
        doc, "Độ bằng phẳng cho phép",
        f"i ≤ 1/{criteria['i_inv']}"
        + (" — cho phép vồng (nới về 1/125)" if criteria.get("allow_vong")
           else "")
    )
    if criteria.get("allow_vong") and criteria.get("i_inv_eff"):
        _add_para_kv(doc, "i áp dụng sau vồng",
                      f"1/{criteria['i_inv_eff']}")

    # ── 5. Lc tối ưu mỗi HK + 2 charts ────────────────────────────
    _add_heading(doc, "5. Chiều dài cọc CDM tối ưu", level=1)
    _add_para(doc,
        "Bảng dưới đây tổng hợp Lc tối ưu cho 6 HK ND × 4 mức ΔS. "
        "HK ND-07 thiếu dữ liệu lớp đất yếu — không tính được."
    )
    dS = criteria["design_dS_cm"]
    cols = ["HK", "Tự nhiên (m)", "Thiết kế (m)", "Đỉnh CDM (m)",
            "H_soft (m)", "Nguồn Cc", "Lc (m)", "p mũi (m)",
            "Trạng thái mũi", "S_total (cm)"]
    rows = []
    for h in lc_matrix:
        r = h.get("by_dS", {}).get(dS, {})
        _full = r.get("penetrates_full", False)
        if r.get("ok"):
            _trang_thai = "Hết lớp bùn" if _full else "Không hết lớp bùn"
        else:
            _trang_thai = "Không đạt"
        rows.append([
            h.get("name", "—"),
            f"{h.get('nat', 0):.2f}",
            f"{h.get('design', 0):.2f}",
            f"{h.get('cdm_top_elev', 0):.2f}",
            f"{h.get('H_soft_m', 0):.1f}" if h.get("H_soft_m") else "—",
            (h.get("cc_source") or "—")
            + (" (mượn)" if h.get("borrowed") else ""),
            f"{r['Lc_m']:.1f}" if r.get("ok") and r.get("Lc_m") else "—",
            f"{r['p_optimal_m']:.1f}" if r.get("p_optimal_m") else "—",
            _trang_thai,
            f"{r['S_total_cm']:.1f}" if r.get("S_total_cm") else "—",
        ])
    _add_data_table(doc, cols, rows)
    _add_para(doc,
        "Quy ước trạng thái mũi: Hết lớp bùn = mũi cọc qua HẾT lớp đất yếu "
        "theo ký hiệu lớp — S₂ VẪN TÍNH cho phần đất nén lún còn lại bên dưới "
        "(sét cứng / lớp xen có Cc nhỏ), thường nhỏ hơn nhiều so với trong lớp bùn. "
        "Không hết lớp bùn (cọc trong lớp bùn) = mũi cọc nằm TRONG lớp bùn — S₂ tính "
        "cho cả phần bùn còn lại bên dưới mũi → S₂ lớn hơn.",
        italic=True,
    )

    # ── 5b. Phân tích kỹ thuật Hết bùn vs Không hết bùn ──────────
    _add_heading(doc, "5b. Phân tích kỹ thuật — Cọc Hết lớp bùn vs Không hết lớp bùn",
                   level=2)

    _add_heading(doc, "5b.1 Định nghĩa kỹ thuật", level=3)
    _add_para(doc,
        "Hết lớp bùn (p_tip ≥ H_soft theo ký hiệu lớp): mũi cọc vượt qua đáy "
        "lớp bùn. Engine VẪN TÍNH S₂ cho phần đất nén lún còn lại bên dưới mũi "
        "(sét cứng, lớp xen) — S₂ thường nhỏ hơn nhiều so với trong lớp bùn nhưng "
        "KHÔNG bằng 0."
    )
    _add_para(doc,
        "Không hết lớp bùn — cọc trong lớp bùn (p_tip < H_soft): mũi cọc kết thúc "
        "trong lòng lớp bùn. Phần bùn còn lại bên dưới mũi (có Cc rất lớn) "
        "vẫn cố kết dưới tải q → S₂ lớn."
    )

    _add_heading(doc, "5b.2 Cơ sở vật lý + công thức S₂ — phân nhánh sét/cát", level=3)
    _add_para(doc,
        "Engine find_cdm_length tính S₂ bằng cách phân tố lớp đất dưới mũi cọc "
        "theo độ sâu (bước 2 m), tích lũy lún TỚI khi Δσ/σ'v0 < 10%. Mỗi phân "
        "tố áp dụng công thức KHÁC NHAU tùy loại đất:"
    )
    img_s2_clay = _formula_png([
        ("Lớp sét — Terzaghi 1D:",
         r"$S_{i,clay} = \dfrac{h_i \cdot C_c}{1+e_0} \log_{10}\dfrac{\sigma'_{vf}}{\sigma'_{v0}}$"),
        ("Phân nhánh:",
         r"OC / NC / cross-PC theo $\sigma'_{vf}$ vs $P_C$"),
        ("Fallback (thiếu Cc, có a₁₋₂):",
         r"$E_{oed} = \dfrac{1+e_0}{a_{1-2}} \times 98{,}0665$;$\quad S_i = \dfrac{\Delta\sigma \cdot h_i}{E_{oed}}$"),
    ], title="Công thức S₂ — lớp sét (Terzaghi / Eoed)", height_cm=4.0)
    _add_picture(doc, img_s2_clay, width_cm=15.5)

    img_s2_sand = _formula_png([
        ("Lớp cát — đàn hồi SPT:",
         r"$S_{i,sand} = \dfrac{\Delta\sigma \cdot h_i}{E_s}$,$\quad E_s = \alpha_{sand} \cdot N_{SPT}$"),
        ("Hằng số:",
         r"$\alpha_{sand} = 2000$ kPa mặc định"),
        ("SAND_SYMBOLS_S2:",
         r"{F, 2a, 2b, 2c, 3a, 3b, 3c, 4, 5, 5a, 5b, 6, 7, 8}"),
        ("Fallback:",
         r"$N \approx 10$ nếu không có SPT trong phân tố"),
    ], title="Công thức S₂ — lớp cát (SPT-based)", height_cm=4.0)
    _add_picture(doc, img_s2_sand, width_cm=15.5)

    _add_para(doc,
        "Điều kiện dừng tích lũy: Δσ/σ'v0 < 10% — áp dụng chung cho cả sét "
        "và cát."
    )
    _add_para(doc,
        "Cọc KHÔNG hết bùn: S₂ gồm cố kết PHẦN BÙN còn lại (Cc rất lớn, "
        "Terzaghi) + các lớp dưới (sét/cát) → S₂ lớn."
    )
    _add_para(doc,
        "Cọc HẾT bùn theo ký hiệu (qua hết 1/1b/2/XMD): S₂ chỉ còn lớp dưới — "
        "nếu LỚP CÁT (4, 5a...) thì dùng SPT (Es lớn → Si nhỏ); nếu sét cứng "
        "(3) vẫn Cc nhưng nhỏ → S₂ nhỏ hơn nhiều nhưng vẫn > 0."
    )

    # Biểu đồ ứng suất profile cho HK đại diện
    _representative_hk = None
    _representative_tip = None
    for h in lc_matrix:
        r = h.get("by_dS", {}).get(dS, {})
        if r.get("ok") and r.get("tip_depth_m"):
            _representative_hk = h["name"]
            _representative_tip = float(r["tip_depth_m"])
            break
    if _representative_hk:
        _add_para(doc,
            f"Biểu đồ dưới đây minh họa profile ứng suất tại {_representative_hk} "
            "— so sánh σ'v0 (ứng suất bản thân) với Δσ = q (tải gây lún) và đường "
            "10% σ'v0. Vị trí giao điểm là điểm engine dừng tính S₂. "
            "Tên lớp đất + γ' hiển thị bên phải.",
            italic=True,
        )
        img_stress = _chart_stress_profile(
            _representative_hk, meta.get("q_kPa", 40.8),
            tip_depth_m=_representative_tip,
        )
        if img_stress:
            _add_picture(doc, img_stress, width_cm=15.5)

    _add_heading(doc, "5b.3 So sánh ưu / nhược kinh tế - kỹ thuật", level=3)
    cols_cmp = ["Phương án", "Ưu điểm", "Nhược điểm", "Khi nên dùng"]
    rows_cmp = [
        [
            "Không hết lớp bùn (trong lớp bùn)",
            "Tiết kiệm 15-25% vật liệu cọc; thi công nhanh",
            "S₂ > 0; cần quan trắc; nhạy với sai số Cc",
            "Lớp bùn DÀY (>25m); S₂ chấp nhận theo TCCS 41",
        ],
        [
            "Hết lớp bùn",
            "S₂ ≈ 0; độ tin cậy cao; không cần quan trắc lâu",
            "Tốn vật liệu; thi công lâu hơn",
            "Lớp bùn MỎNG (<15m); tải lớn; gần cầu/cống nghiêm ngặt",
        ],
    ]
    _add_data_table(doc, cols_cmp, rows_cmp)

    _add_heading(doc, "5b.4 Áp dụng cho QTT", level=3)
    # Đếm trạng thái mũi cho từng ΔS
    _trang_thai_summary = []
    for d in dS_list:
        n_full = 0; n_floating = 0; n_fail = 0
        for h in lc_matrix:
            r = h.get("by_dS", {}).get(d, {})
            if r.get("ok"):
                if r.get("penetrates_full"):
                    n_full += 1
                else:
                    n_floating += 1
            else:
                n_fail += 1
        _trang_thai_summary.append([
            f"ΔS ≤ {int(d)} cm",
            f"{n_full}",
            f"{n_floating}",
            f"{n_fail}",
        ])
    _add_data_table(
        doc,
        ["Phương án", "Hết lớp bùn", "Không hết lớp bùn", "Không đạt"],
        _trang_thai_summary,
    )
    _add_para(doc,
        f"Với ΔS = {int(dS)} cm đã chọn: đa số HK ND có cọc KHÔNG hết lớp bùn "
        "(trong lớp bùn). Engine find_cdm_length tự chọn Lc ngắn nhất đạt S₁+S₂ ≤ "
        "ΔS — đây là tối ưu chi phí."
    )

    _add_heading(doc, "5b.5 Rủi ro cần quản lý khi cọc không hết lớp bùn", level=3)
    for line in [
        "Sai số Cc: nếu thực địa Cc > Cc_lab thì S₂ thực > S₂ tính → có thể "
        "vượt ΔS. Đặc biệt nguy hiểm với HK mượn Cc (ND-03/04/05 mượn từ "
        "ND-06 ở khoảng cách 91-135 m).",
        "Lún thứ cấp (creep) Cα chưa tính trong S₂ — bùn có thể tiếp tục lún "
        "sau cố kết sơ cấp.",
        "Mực nước ngầm biến động — σ'v0 thay đổi → profile lún khác → cần "
        "kiểm tra MNN min/max.",
        "Đề xuất: BỔ SUNG thí nghiệm Cc tại chính ND-03, ND-04, ND-05 trước "
        "khi chốt thiết kế cọc không hết lớp bùn.",
    ]:
        p_li = doc.add_paragraph(style="List Bullet")
        p_li.add_run(line).font.size = Pt(12)

    _add_para(doc, "Biểu đồ Heatmap Lc — HK × ΔS:", italic=True)
    _add_picture(doc, _chart_heatmap_lc(lc_matrix, dS_list), width_cm=15.5)

    _add_para(doc, "So sánh Lc giữa các HK theo 4 phương án ΔS:", italic=True)
    _add_picture(doc, _chart_bar_lc(lc_matrix, dS_list), width_cm=15.5)

    # ── 6. S(Lc) curves ───────────────────────────────────────────
    if curves_data:
        _add_heading(doc, "6. Đường cong S(Lc) — Quan hệ độ lún vs chiều dài cọc",
                       level=1)
        _add_para(doc,
            "Quét Lc từ 0.5m đến 35m cho mỗi HK ND có Cc. Giao điểm đường cong "
            "với ngưỡng ngang ΔS = Lc TỐI THIỂU cho HK đó. Đường nằm trên ngưỡng "
            "= không đạt được ΔS đó."
        )
        img_curve = _chart_s_vs_lc(curves_data)
        if img_curve:
            _add_picture(doc, img_curve, width_cm=15.5)

    # ── 7. Kiểm tra độ bằng phẳng + chart ─────────────────────────
    _add_heading(doc, "7. Kiểm tra độ bằng phẳng pair-wise", level=1)
    _add_para(doc,
        f"Kiểm tra độ dốc dọc i = ΔS/d cho tất cả {len(smoothness_pairs)} cặp HK. "
        f"Yêu cầu i ≤ 1/{criteria.get('i_inv_eff') or '—'} "
        f"({'có vồng' if criteria.get('allow_vong') else 'không vồng'})."
    )
    if smoothness_pairs:
        cols2 = ["Cặp HK", "Khoảng cách (m)", "ΔS (cm)",
                 "i thực tế", "i cho phép", "Đánh giá"]
        rows2 = []
        for p in smoothness_pairs:
            _i_act = p.get("i_inv_actual")
            rows2.append([
                f"{p['i']} ↔ {p['j']}",
                f"{p['d_m']:.1f}",
                f"{p['dS_m'] * 100:.2f}",
                f"1/{int(_i_act)}" if _i_act else "∞",
                f"1/{int(p['i_inv_max'])}",
                "Đạt" if p["ok"] else "Không đạt",
            ])
        _add_data_table(doc, cols2, rows2)
        _n_fail = sum(1 for p in smoothness_pairs if not p["ok"])
        _add_para(doc,
            f"Kết quả: {len(smoothness_pairs) - _n_fail}/{len(smoothness_pairs)} "
            f"cặp ĐẠT. {_n_fail} cặp KHÔNG đạt."
        )
        # Chart scatter
        img_smooth = _chart_smoothness_scatter(
            smoothness_pairs,
            int(criteria.get("i_inv_eff") or criteria.get("i_inv") or 200),
        )
        if img_smooth:
            _add_para(doc, "Đồ thị phân tán cặp HK vs đường giới hạn:",
                       italic=True)
            _add_picture(doc, img_smooth, width_cm=15.5)

    # ── 8. Phân vùng grid + 2 maps ────────────────────────────────
    _add_heading(doc, "8. Phân vùng thiết kế CDM theo Lc", level=1)
    _add_para(doc,
        f"Tính Lc tại {grid_meta.get('n_grid', 0)} điểm grid 20m × 20m trong "
        "polygon QTT, sau đó phân vùng quantile binning. Lc thiết kế của mỗi vùng "
        "= max Lc trong vùng (an toàn)."
    )
    stats = zoning.get("stats", [])
    if stats:
        cols3 = ["Vùng", "Lc (m) — khoảng", "Số điểm grid",
                 "Diện tích (m²)", "Lc thiết kế (m)"]
        rows3 = []
        for s in stats:
            rows3.append([
                f"P{s['zone_id'] + 1}",
                f"{s['Lc_min']:.1f} – {s['Lc_max']:.1f}",
                f"{s['n_points']}",
                f"{s['area_m2']:,.0f}",
                f"{s['Lc_design']:.1f}",
            ])
        _add_data_table(doc, cols3, rows3)
        _tot_area = sum(s["area_m2"] for s in stats)
        _Lc_max_design = max(s["Lc_design"] for s in stats)
        _add_para(doc,
            f"Tổng diện tích thiết kế CDM: {_tot_area:,.0f} m². "
            f"Lc thiết kế lớn nhất: {_Lc_max_design:.1f} m."
        )

    # Charts grid + zoning
    if grid_points and hks_with_S:
        img_grid = _chart_grid_lc_map(grid_points, hks_with_S)
        if img_grid:
            _add_para(doc, "Bản đồ Lc grid (162 điểm) — phân bố không gian:",
                       italic=True)
            _add_picture(doc, img_grid, width_cm=15.5)

        img_zone = _chart_zoning_map(grid_points, zoning, hks_with_S)
        if img_zone:
            _add_para(doc, f"Bản đồ phân vùng — {len(stats)} vùng thiết kế:",
                       italic=True)
            _add_picture(doc, img_zone, width_cm=15.5)

    # ── 9. Kết luận và Khuyến nghị ────────────────────────────────
    _add_heading(doc, "9. Kết luận và Khuyến nghị", level=1)
    _Lc_max_all = 0.0
    for h in lc_matrix:
        r = h.get("by_dS", {}).get(dS, {})
        if r.get("ok") and r.get("Lc_m"):
            _Lc_max_all = max(_Lc_max_all, r["Lc_m"])
    _add_para(doc,
        f"Với tiêu chí ΔS ≤ {int(dS)} cm và i ≤ 1/{criteria.get('i_inv_eff') or '—'}, "
        f"chiều dài cọc CDM tối đa cần thiết là {_Lc_max_all:.1f} m (đo từ đỉnh CDM). "
        f"Đỉnh cọc đặt tại cao độ design − {meta.get('fill_thickness_m', 0):.1f} m."
    )
    _add_para(doc, "Phương án thiết kế đề xuất:")
    for line in [
        f"Áp dụng phân vùng P1-P{len(stats)} với Lc thiết kế tương ứng (xem mục 8).",
        "Trong mỗi vùng, dùng Lc thiết kế bằng Lc tối đa của vùng — đảm bảo an toàn.",
        "Kiểm tra lại độ bằng phẳng sau khi xác định Lc thiết kế cuối cùng.",
        "Cho phép tạo vồng trước với i ≤ 1/125 (TCCS 41 Phụ lục E) khi cần.",
        "HK ND-03/04/05 thiếu Cc → bổ sung thí nghiệm để tăng độ tin cậy.",
    ]:
        p_li = doc.add_paragraph(style="List Bullet")
        p_li.add_run(line).font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────
# Generic multi-zone builder (BXN/NHC/KE_park/KE_levee) — kế thừa QTT layout
# ─────────────────────────────────────────────────────────────────────────

def build_zone_decision_docx(
    zone_code: str,
    meta: dict,
    lc_matrix: list[dict],
    criteria: dict,
    smoothness_pairs: list[dict],
    curves_data: Optional[dict] = None,
    hks_with_S: Optional[list[dict]] = None,
) -> bytes:
    """Tạo Word docx cho 1 zone (BXN/NHC/KE_park/KE_levee).

    Khác QTT: không có grid 162 điểm + không có zoning. Báo cáo gọn 7 mục:
      1. Thông số đầu vào
      2. Tiêu chí thiết kế
      3. Lc tối ưu mỗi HK
      4. Đường cong S(Lc)
      5. Độ bằng phẳng pair-wise
      6. Phân tích kỹ thuật mũi cọc
      7. Kết luận
    """
    import sys as _sys
    _sys.path.insert(0, str(_ROOT / "scripts"))
    from docx import Document
    from docx.shared import Pt

    zone_desc = meta.get("zone_desc", zone_code)
    is_levee = meta.get("force_full_penetration", False)

    doc = Document()
    _set_normal_style(doc)
    dS_list = meta.get("delta_S_values_cm") or [10.0, 15.0, 20.0, 25.0, 30.0, 40.0]
    dS = criteria.get("design_dS_cm", 30.0)

    # ── Trang bìa ─────────────────────────────────────────────────
    _add_heading(
        doc,
        f"BÁO CÁO PHÂN TÍCH LỰA CHỌN CHIỀU DÀI CỌC CDM "
        f"KHU VỰC {zone_desc.upper()} ({zone_code})",
        level=0,
    )
    p = doc.add_paragraph()
    r = p.add_run("Trung tâm Hành chính TP.HCM — Mã số 202605-TTHC")
    r.bold = True; r.font.size = Pt(13)
    _add_para_kv(doc, "Ngày lập", datetime.now().strftime("%d/%m/%Y"))
    _add_para_kv(
        doc, "Tiêu chuẩn áp dụng",
        "TCCS 41:2022 (Điều 6.2.3 Bảng 1, Phụ lục E.1) · "
        "TCVN 9403:2012 (Phụ lục C, Phụ lục B)",
    )
    if is_levee:
        _add_para(doc,
            "QUY TẮC ĐẶC BIỆT cho khu vực kè: cọc CDM LUÔN xuyên hết lớp bùn "
            f"(ngàm thêm {meta.get('L_ngam_m', 1.0):.1f}m vào lớp đất tốt). "
            "Lc không phụ thuộc ΔS — chỉ kiểm tra S_total có đạt ΔS hay không.",
            italic=True,
        )

    # ── 1. Thông số đầu vào ────────────────────────────────────────
    _add_heading(doc, "1. Thông số đầu vào", level=1)
    _add_para_kv(doc, "Khu vực", f"{zone_desc} ({zone_code})")
    _add_para_kv(doc, "Số HK selected", f"{len(lc_matrix)} hố khoan")
    _add_para_kv(doc, "Tải phân bố q",
                  f"{meta.get('q_kPa', 0):.1f} kPa (cấu tạo Áo đường + He + Hse)")
    _add_para_kv(doc, "Σh đắp trên đỉnh CDM",
                  f"{meta.get('fill_thickness_m', 0):.1f} m")
    _add_para_kv(doc, "Cao độ thiết kế",
                  f"{meta.get('design_elev_global_m', 2.7):.2f} m")
    _add_para_kv(doc, "Đường kính cọc D", f"{meta.get('D_mm', 800):.0f} mm")
    _add_para_kv(doc, "Khoảng cách cọc s",
                  f"{meta.get('spacing_m', 1.8):.2f} m  ({meta.get('pattern', 'square')})")
    _add_para_kv(doc, "Tỷ lệ thay thế a", f"{meta.get('a', 0):.4f}")
    _add_para_kv(doc, "Mô đun trụ Ec",
                  f"{meta.get('Ec_kPa', 0):,.0f} kPa")

    # ── 2. Tiêu chí thiết kế ───────────────────────────────────────
    _add_heading(doc, "2. Tiêu chí thiết kế đã chọn", level=1)
    _add_para_kv(doc, "Cấp đường",
                  _ROAD_CLASS_DESC.get(criteria.get("road_class", "cap_I_IV"),
                                        criteria.get("road_class", "—")))
    _add_para_kv(doc, "Công trình",
                  _STRUCT_DESC.get(criteria.get("structure", "cong"),
                                    criteria.get("structure", "—")))
    _add_para_kv(doc, "Tốc độ thiết kế", f"{criteria.get('speed_kmh', 60)} km/h")
    _add_para_kv(doc, "Độ lún cho phép ΔS", f"≤ {int(dS)} cm")
    if criteria.get("i_inv"):
        _add_para_kv(doc, "Độ bằng phẳng",
                      f"i ≤ 1/{criteria['i_inv']}"
                      + (" (vồng → 1/125)" if criteria.get("allow_vong") else ""))

    # ── 3. Lc tối ưu mỗi HK ────────────────────────────────────────
    _add_heading(doc, "3. Chiều dài cọc CDM tối ưu", level=1)
    cols = ["HK", "Tự nhiên (m)", "H_soft (m)", "Đỉnh CDM (m)",
            "Nguồn Cc", "p mũi (m)", "Lc (m)", "Trạng thái mũi",
            "S_total (cm)"]
    rows = []
    for h in lc_matrix:
        r = h.get("by_dS", {}).get(dS, {})
        full = r.get("penetrates_full", False)
        ok = r.get("ok", False)
        if ok:
            ts = "Hết lớp bùn" if full else "Không hết lớp bùn"
        else:
            ts = "Không đạt"
        rows.append([
            h.get("name", "—"),
            f"{h.get('nat', 0):.2f}",
            f"{h.get('H_soft_m', 0):.1f}" if h.get("H_soft_m") else "—",
            f"{h.get('cdm_top_elev', 0):.2f}",
            (h.get("cc_source") or "—")
            + (" (mượn)" if h.get("borrowed") else ""),
            f"{r['p_optimal_m']:.1f}" if r.get("p_optimal_m") else "—",
            f"{r['Lc_m']:.1f}" if r.get("Lc_m") is not None else "—",
            ts,
            f"{r['S_total_cm']:.1f}" if r.get("S_total_cm") else "—",
        ])
    _add_data_table(doc, cols, rows)
    if is_levee:
        _add_para(doc,
            "Khu vực kè bắt buộc xuyên hết lớp bùn — Lc cố định, "
            "chỉ kiểm tra S_total đạt ΔS hay không.",
            italic=True,
        )

    # Heatmap + bar chart
    _add_para(doc, "Heatmap Lc theo HK × ΔS:", italic=True)
    _add_picture(doc, _chart_heatmap_lc(lc_matrix, dS_list), width_cm=15.5)
    _add_para(doc, "So sánh Lc giữa các HK:", italic=True)
    _add_picture(doc, _chart_bar_lc(lc_matrix, dS_list), width_cm=15.5)

    # ── 4. S(Lc) curves ────────────────────────────────────────────
    if curves_data and curves_data.get("curves"):
        _add_heading(doc, "4. Đường cong S(Lc) — quan hệ lún vs chiều dài",
                       level=1)
        img_curve = _chart_s_vs_lc(curves_data)
        if img_curve:
            _add_picture(doc, img_curve, width_cm=15.5)

    # ── 5. Độ bằng phẳng pair-wise ─────────────────────────────────
    if smoothness_pairs:
        _add_heading(doc, "5. Kiểm tra độ bằng phẳng pair-wise", level=1)
        cols2 = ["Cặp HK", "Khoảng cách (m)", "ΔS (cm)",
                 "i thực tế", "i cho phép", "Đánh giá"]
        rows2 = []
        for p in smoothness_pairs:
            _i_act = p.get("i_inv_actual")
            rows2.append([
                f"{p['i']} ↔ {p['j']}",
                f"{p['d_m']:.1f}",
                f"{p['dS_m'] * 100:.2f}",
                f"1/{int(_i_act)}" if _i_act else "∞",
                f"1/{int(p['i_inv_max'])}",
                "Đạt" if p["ok"] else "Không đạt",
            ])
        _add_data_table(doc, cols2, rows2)
        img_smooth = _chart_smoothness_scatter(
            smoothness_pairs,
            int(criteria.get("i_inv_eff") or criteria.get("i_inv") or 200),
        )
        if img_smooth:
            _add_picture(doc, img_smooth, width_cm=15.5)

    # ── 6. Phân tích kỹ thuật mũi cọc ──────────────────────────────
    _add_heading(doc, "6. Phân tích kỹ thuật — Trạng thái mũi cọc", level=1)
    _add_para(doc,
        "S₂ phân nhánh theo loại đất: lớp sét dùng Cc Terzaghi 1D (OC/NC/cross-PC) "
        "hoặc Eoed fallback; lớp cát dùng SPT (Es = α·N, α = 2000 kPa). "
        "Engine dừng tích lũy khi Δσ/σ'v0 < 10%."
    )
    if is_levee:
        _add_para(doc,
            "Khu vực kè: cọc xuyên hết bùn → S₂ chỉ tính cho lớp dưới (sét cứng "
            "+ cát) → S nhỏ → đáp ứng nhanh ΔS thiết kế."
        )
    else:
        # Đếm trạng thái mũi
        n_full = sum(1 for h in lc_matrix
                     if h.get("by_dS", {}).get(dS, {}).get("penetrates_full"))
        n_floating = sum(1 for h in lc_matrix
                         if h.get("by_dS", {}).get(dS, {}).get("ok")
                         and not h.get("by_dS", {}).get(dS, {}).get("penetrates_full"))
        _add_para(doc,
            f"Tại ΔS = {int(dS)} cm: {n_full} HK hết lớp bùn, "
            f"{n_floating} HK không hết lớp bùn (cọc trong lớp bùn). "
            "Cọc trong lớp bùn tiết kiệm vật liệu nhưng nhạy với sai số Cc — "
            "cần quan trắc lâu dài."
        )

    # Profile ứng suất HK đại diện
    _rep_hk = None; _rep_tip = None
    for h in lc_matrix:
        r = h.get("by_dS", {}).get(dS, {})
        if r.get("ok") and r.get("tip_depth_m"):
            _rep_hk = h["name"]; _rep_tip = float(r["tip_depth_m"])
            break
    if _rep_hk:
        _add_para(doc,
            f"Biểu đồ profile ứng suất tại {_rep_hk} — so sánh σ'v0 vs Δσ + "
            "vị trí Δσ/σ'v0 = 10% + tên lớp đất:",
            italic=True,
        )
        img_stress = _chart_stress_profile(
            _rep_hk, meta.get("q_kPa", 40.8), tip_depth_m=_rep_tip,
        )
        if img_stress:
            _add_picture(doc, img_stress, width_cm=15.5)

    # ── 7. Kết luận ────────────────────────────────────────────────
    _add_heading(doc, "7. Kết luận và Khuyến nghị", level=1)
    _Lc_max = 0.0
    n_ok = 0
    for h in lc_matrix:
        r = h.get("by_dS", {}).get(dS, {})
        if r.get("ok") and r.get("Lc_m") is not None:
            _Lc_max = max(_Lc_max, r["Lc_m"])
            n_ok += 1
    _add_para(doc,
        f"Khu vực {zone_desc}: {n_ok}/{len(lc_matrix)} HK đạt tiêu chí "
        f"ΔS ≤ {int(dS)} cm với Lc tối đa = {_Lc_max:.1f} m."
    )
    if is_levee:
        _add_para(doc,
            f"Cọc kè luôn xuyên hết bùn (ngàm thêm {meta.get('L_ngam_m', 1.0):.1f}m) — "
            "Lc cố định theo H_soft của từng HK."
        )
    _add_para(doc, "Khuyến nghị:")
    for line in [
        f"Áp dụng Lc thiết kế = {_Lc_max:.1f} m (max) hoặc phân vùng theo HK "
        f"để tối ưu chi phí.",
        "Bổ sung thí nghiệm Cc tại HK mượn dữ liệu nếu khoảng cách > 100 m.",
        f"Theo dõi cố kết 15 năm cho HK không hết lớp bùn (TCCS 41 Bảng 1).",
    ]:
        p_li = doc.add_paragraph(style="List Bullet")
        p_li.add_run(line).font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
