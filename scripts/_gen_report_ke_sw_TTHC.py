"""Sinh báo cáo Word: Lựa chọn cọc ván SW — Dự án TTHC HCM.

Chạy: python scripts/_gen_report_ke_sw_TTHC.py
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────────
_ROOT      = Path(__file__).parent.parent
_KE_JSON   = _ROOT / "data" / "ke_sw_202605_TTHC.json"
_SOIL_JSON = _ROOT / "data" / "soil_profile_202605_TTHC.json"
_OUT_DIR   = Path(r"G:\My Drive\202605-TRUNG TAM HCM\KET CAU KE")
_OUT_FILE  = _OUT_DIR / "260514 BAO CAO LUA CHON COC SW-TTHC-HCM.docx"


# ── Helpers ────────────────────────────────────────────────────────────────────
def _set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def _heading(doc, text, level=1, size=14, bold=True, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, color=color)
    return p


def _para(doc, text="", size=12, bold=False, italic=False, indent=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        run = p.add_run(text)
        _set_font(run, size=size, bold=bold, italic=italic)
    return p


def _add_table(doc, headers, rows, col_widths=None, header_color=(31, 73, 125)):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(h)
        _set_font(run, size=10, bold=True, color=(255, 255, 255))
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        rgb = "".join(f"{c:02X}" for c in header_color)
        shading.set(qn("w:fill"), rgb)
        cell._tc.get_or_add_tcPr().append(shading)

    for r_idx, row in enumerate(rows):
        cells = tbl.rows[r_idx + 1].cells
        bg = "F2F2F2" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            txt = str(val) if val is not None else "—"
            bold_cell = txt in ("Đạt", "Không đạt", "SPECIAL") or "PASS" in txt
            run = p.add_run(txt)
            _set_font(run, size=10, bold=bold_cell)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), bg)
            cell._tc.get_or_add_tcPr().append(shading)

    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return tbl


# ── Main ───────────────────────────────────────────────────────────────────────
def build():
    ke   = json.loads(_KE_JSON.read_text(encoding="utf-8"))
    soil = json.loads(_SOIL_JSON.read_text(encoding="utf-8"))
    dc   = ke["design_conditions"]
    bhs  = ke["boreholes"]
    ml   = ke["NT2_multilayer_summary"]

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.0)

    # ── TIÊU ĐỀ ────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("BÁO CÁO LỰA CHỌN CỌC VÁN BÊ TÔNG SW")
    _set_font(run, size=16, bold=True, color=(31, 73, 125))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("KÈ BẢO VỆ BỜ — DỰ ÁN TRUNG TÂM HÀNH CHÍNH TP.HCM")
    _set_font(run, size=13, bold=True, color=(31, 73, 125))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Mã hồ sơ: 260512 CVTT-TTHC  |  Ngày lập: 2026-05-14  |  Cập nhật NT2 đa lớp")
    _set_font(run, size=11, italic=True, color=(89, 89, 89))

    doc.add_paragraph()

    # ── 1. ĐIỀU KIỆN THIẾT KẾ ──────────────────────────────────────────────────
    _heading(doc, "1. ĐIỀU KIỆN THIẾT KẾ", size=13, color=(31, 73, 125))

    su_layer = dc.get("su_per_layer_kNm2", {})
    su_str   = ", ".join(f"Lớp {k}: {v}" for k, v in su_layer.items() if not k.startswith("_"))

    rows_dc = [
        ("Cao độ đỉnh kè",            f"+{dc['top_ke_elevation_m']:.2f} m"),
        ("Cao độ tự nhiên",            f"{dc['natural_ground_elevation_m']:.2f} m"),
        ("Chiều cao kè trên mặt đất",  "2.70 m"),
        ("Lớp đất kiểm soát",          f"Lớp {dc['soft_clay_layer_symbol']} — {dc['soft_clay_description']}"),
        ("Xuyên qua lớp sét chảy",     f"≥ {dc['min_penetration_below_soft_clay_m']:.2f} m"),
        ("Tiêu chuẩn NT2",             "TCVN 11823-10:2017, Điều 7.3.8.6.2 — phương pháp α"),
        ("Hệ số sức kháng NT2",        f"φ_stat = {dc['phi_stat_alpha']} (Bảng 9)"),
        ("su theo từng lớp (kN/m²)",   su_str),
        ("Phương pháp alpha",           "Tomlinson (1980) — tự động tính theo su"),
        ("Catalog cọc",                 "BETON 6 Ver.2015"),
    ]
    _add_table(doc, ["Thông số", "Giá trị"], rows_dc, col_widths=[7.5, 9.0])
    doc.add_paragraph()

    # ── 2. HAI NGUYÊN TẮC ──────────────────────────────────────────────────────
    _heading(doc, "2. HAI NGUYÊN TẮC CHỌN CỌC", size=13, color=(31, 73, 125))

    _heading(doc, "2.1  Nguyên tắc 1 — Chiều dài tối thiểu", size=11, bold=True, color=(0, 0, 0))
    _para(doc, "Mũi cọc phải vượt qua đáy lớp sét chảy (Lớp 1) ít nhất 1.00 m:", size=11)
    _para(doc, "L_req  =  2.70  +  H(Lớp 1)  +  1.00  =  H(Lớp 1)  +  3.70  (m)", size=11, bold=True, indent=1)

    _heading(doc, "2.2  Nguyên tắc 2 — Sức chịu tải tính toán ≥ Trọng lượng bản thân cọc (TCVN 11823-10:2017)", size=11, bold=True, color=(0, 0, 0))
    _para(doc, "RR  =  φ_stat × (Rs + Rp)  ≥  W_cọc     (φ_stat = 0.35, Bảng 9)", size=11, bold=True, indent=1)
    _para(doc, "Trong đó:", size=11)
    _para(doc, "Rs  =  Σ_lớp [ α(su) × su × C × L_lớp ]   — tính từng lớp đất, bỏ qua đoạn đất đắp", size=10, indent=1)
    _para(doc, "Rp  =  9 × su_mũi × Ap   (Pt. 65 — sức kháng mũi, sét bão hòa; = 0 nếu mũi trong cát)", size=10, indent=1)
    _para(doc, "W_cọc  =  TL(T) × 9.81 / L_std × L   [kN/cọc]", size=10, indent=1)
    _para(doc,
          "Lưu ý đơn vị: w_per_pile [kN/m/cọc] = TL×9.81/L_std (dùng NT2). "
          "w_plaxis [kN/m/m] = w_per_pile / spacing (input PLAXIS Plate — hai giá trị khác nhau).",
          size=10, italic=True)
    doc.add_paragraph()

    # ── 3. CHIỀU DÀY LỚP SÉT CHẢY VÀ L_REQ ────────────────────────────────────
    _heading(doc, "3. CHIỀU DÀY LỚP SÉT CHẢY VÀ CHIỀU DÀI TỐI THIỂU", size=13, color=(31, 73, 125))

    bh_rows = []
    for bh in bhs:
        if bh["NT1"] == "SPECIAL":
            l_req = f"{bh['L_req_m']:.1f} [ĐB]"
            nt1   = "[ĐB] XMD"
        else:
            margin = bh["margin_NT1_m"]
            if bh["NT1"] == "PASS_CRITICAL":
                nt1 = f"Đạt (+{margin:.1f}m) CẢNH BÁO"
            else:
                nt1 = f"Đạt (+{margin:.1f}m)"
            l_req = f"{bh['L_req_m']:.1f}"
        bh_rows.append((
            bh["name"],
            f"{bh['Z_m']:+.3f}",
            f"{bh['H_layer1_m']:.1f}",
            l_req,
            bh["tip_layer"],
            bh["recommended_pile"],
            str(bh["recommended_L_m"]) if bh["recommended_L_m"] else "—",
            nt1,
        ))

    _add_table(doc,
        ["HK", "Z (m)", "H Lớp 1\n(m)", "L_req\n(m)", "Lớp tại\nmũi L=29m", "Cọc\nđề xuất", "L\n(m)", "NT1"],
        bh_rows,
        col_widths=[1.4, 1.5, 1.5, 1.5, 2.2, 2.0, 1.2, 5.2])
    _para(doc, "[ĐB] HK12: Lớp 1 là sét rất dẻo (không phải chảy); XMD hiện hữu 11.0–23.9m — cần thiết kế riêng.", size=10, italic=True)
    doc.add_paragraph()

    # ── 4. LỌC CỌC SW THEO L_MAX ───────────────────────────────────────────────
    _heading(doc, "4. LỌC CỌC SW THEO L_MAX ≥ L_REQ (L_req,max = 28.7 m tại HK10)", size=13, color=(31, 73, 125))

    pile_rows = [
        ("SW-600B", 600,  24, 5.34,  "3.724",          "Không đạt — L_max=24 < 28.7"),
        ("SW-740",  740,  28, 6.80,  "4.206",          "Không đạt — L_max=28 < 28.7"),
        ("SW-840",  840,  29, 7.29,  "4.595",          "Đạt — biên 0.3m tại HK10"),
        ("SW-940",  940,  30, 7.81,  "4.984 (nội suy)", "Đạt — an toàn HK10 (L_max=30m)"),
        ("SW-1100", 1100, 32, 10.54, "—",          "Đạt — dư thừa"),
        ("SW-1200", 1200, 34, 11.28, "—",          "Đạt — dư thừa"),
    ]
    _add_table(doc,
        ["Loại cọc", "H (mm)", "L_max (m)", "w_per_pile\n(kN/m/cọc)", "Chu vi\n(mm)", "Đánh giá NT1"],
        pile_rows,
        col_widths=[2.0, 1.7, 2.0, 2.5, 2.5, 5.8])
    _para(doc, "w_per_pile = TL(T) × 9.81 / L_std  [kN/m/cọc]  —  dùng để tính W_cọc cho NT2", size=10, italic=True)
    doc.add_paragraph()

    # ── 5. KIỂM TRA NT2 ĐA LỚP ────────────────────────────────────────────────
    _heading(doc, "5. KIỂM TRA NGUYÊN TẮC 2 — NT2 ĐA LỚP (Alpha-method, su theo từng lớp)", size=13, color=(31, 73, 125))

    _para(doc, "Tiêu chuẩn: TCVN 11823-10:2017, Điều 7.3.8.6.2 — RR = φ_stat × (Rs+Rp) ≥ W_cọc, φ_stat = 0.35", size=11, bold=True)

    # Bảng su và alpha
    _heading(doc, "5.1  Su và alpha theo từng lớp đất", size=11, bold=True, color=(0, 0, 0))
    su_rows = [
        ("Lớp 1 (sét chảy rất dẻo)", "10",  "1.000", "su ≤ 25 kN/m² — Tomlinson"),
        ("Lớp 1b (sét xám)",          "20",  "1.000", "su ≤ 25 kN/m² — Tomlinson"),
        ("Lớp 3 (sét dẻo)",           "35",  "0.968", "Nội suy tuyến tính Tomlinson"),
        ("Lớp 5 (sét cứng)",          "75",  "0.750", "Hình 18 TCVN 11823-10"),
        ("Lớp 5b (sét cứng nửa cứng)","100", "0.600", "Hình 18 TCVN 11823-10"),
        ("Cát / san lấp (F, 2a, 2b, 2c, 4, 5a, 6, 7, XMD)", "—", "—", "qs = 0"),
    ]
    _add_table(doc,
        ["Lớp đất", "su (kN/m²)", "α (Tomlinson)", "Ghi chú"],
        su_rows, col_widths=[5.5, 2.0, 2.5, 6.5])
    doc.add_paragraph()

    # Ví dụ chi tiết HK1
    _heading(doc, "5.2  Ví dụ tính toán chi tiết: SW-840, L=29m, HK1 (Z = −0.800 m)", size=11, bold=True, color=(0, 0, 0))
    nt2_ref = ke["NT2_check_SW840_L29"]
    ex_rows = [
        ("W_cọc = w × L",               f"{nt2_ref['W_pile_kN']} kN   (7.29 kN/m × 29 m)"),
        ("Đất đắp (bỏ qua)",            f"{nt2_ref['fill_m']} m   (đỉnh kè 2.70 − Z −0.800)"),
        ("L trong đất tự nhiên",         f"{nt2_ref['L_in_soil_m']} m   (29 − 3.50)"),
        ("Rs Lớp F (san lấp, qs=0)",     "0 kN"),
        ("Rs Lớp 1 (22.0m, su=10)",      "1.0 × 10 × 4.595 × 22.0 = 1.011 kN"),
        ("Rs Lớp 1b (2.5m, su=20)",      "1.0 × 20 × 4.595 × 2.5 = 230 kN"),
        ("Tổng Rs",                       "1.241 kN"),
        ("Rp (mũi dừng trong Lớp 1b)",  "9 × 20 × 310.700 × 10⁻⁶ = 56 kN"),
        ("Rn = Rs + Rp",                  "1.297 kN"),
        ("φ_stat (Bảng 9)",               "0.35"),
        ("RR = φ × Rn",                   "0.35 × 1.297 = 454 kN  ≥  W = 211.4 kN"),
        ("Kết quả NT2",                   "Đạt (tỷ số = 2.15)"),
    ]
    _add_table(doc, ["Thông số", "Giá trị / Tính toán"], ex_rows, col_widths=[5.5, 11.0])
    doc.add_paragraph()

    # Bảng kết quả đa lớp toàn bộ HK
    _heading(doc, "5.3  Kết quả NT2 đa lớp — Toàn bộ hố khoan (L = 29 m)", size=11, bold=True, color=(0, 0, 0))

    ml_rows = []
    for r in ml["results_by_borehole"]:
        if "note" in r:
            ml_rows.append((r["name"], "SPECIAL", "—", "—", "—", "—", "—", "—", r["note"]))
        else:
            worst = " (kiểm soát)" if r.get("worst_NT2") else ""
            nt2_str = "Đạt" if r["pass"] else "Không đạt"
            ml_rows.append((
                r["name"],
                r["pile"],
                str(r["tip_layer"]),
                str(r["Rs"]),
                str(r["Rp"]),
                str(r["RR"]),
                str(r["W"]),
                str(r["ratio"]) + worst,
                nt2_str,
            ))
    _add_table(doc,
        ["HK", "Cọc", "Lớp mũi", "Rs (kN)", "Rp (kN)", "RR (kN)", "W (kN)", "Tỷ số\nRR/W", "NT2"],
        ml_rows,
        col_widths=[1.2, 1.8, 1.6, 1.7, 1.7, 1.7, 1.7, 2.8, 2.3])
    _para(doc,
          "Hố kiểm soát NT2: HK8 (tỷ số 1.86) — Lớp 1 ngắn (19.5m), mũi vào cát 2b nên Rp=0. "
          "HK10 dùng SW-940, chu vi 4.984 mm (nội suy — cần xác nhận catalog BETON 6).",
          size=10, italic=True)
    doc.add_paragraph()

    # ── 6. BẢNG TỔNG HỢP ──────────────────────────────────────────────────────
    _heading(doc, "6. BẢNG TỔNG HỢP THIẾT KẾ — TOÀN BỘ HỐ KHOAN", size=13, color=(31, 73, 125))

    sum_rows = []
    for bh in bhs:
        w_str  = f"{bh['W_pile_kN']:.1f}" if bh["W_pile_kN"] else "—"
        ml_bh  = bh.get("NT2_multilayer")
        if ml_bh:
            rr_str    = str(ml_bh["RR_kN"])
            ratio_str = str(ml_bh["ratio"])
        else:
            rr_str = ratio_str = "—"
        nt1_map = {"PASS": "Đạt", "PASS_CRITICAL": "Đạt (Cảnh báo)", "SPECIAL": "Đặc biệt"}
        nt2_map = {"PASS": "Đạt", "SPECIAL": "Đặc biệt"}
        nt1_str = nt1_map.get(bh["NT1"], bh["NT1"])
        nt2_str = nt2_map.get(bh["NT2"], bh["NT2"])
        sum_rows.append((
            bh["name"],
            bh["recommended_pile"],
            str(bh["recommended_L_m"] or "—"),
            w_str,
            rr_str,
            ratio_str,
            nt1_str,
            nt2_str,
            bh["note"] if bh["note"] else "",
        ))
    _add_table(doc,
        ["HK", "Loại cọc", "L (m)", "W_cọc\n(kN)", "RR\n(kN)", "Tỷ số\nRR/W", "NT1", "NT2", "Ghi chú"],
        sum_rows,
        col_widths=[1.1, 1.8, 1.1, 1.6, 1.6, 1.5, 2.3, 1.4, 4.1])
    doc.add_paragraph()

    # ── 7. KHUYẾN NGHỊ ─────────────────────────────────────────────────────────
    _heading(doc, "7. KHUYẾN NGHỊ LỰA CHỌN CỌC", size=13, color=(31, 73, 125))

    rec_rows = [
        ("HK1–HK9, HK11 (10 HK)", "SW-840", "29",
         "NT1 Đạt  NT2 Đạt (tỷ số 1.86–2.15) — Toàn tuyến đồng nhất"),
        ("HK10", "SW-940", "29",
         "NT1 biên 0.3m với SW-840 — dùng SW-940 an toàn hơn (tỷ số NT2 = 2.17)"),
        ("HK12", "Thiết kế riêng", "—",
         "XMD hiện hữu — cần xác nhận chiều sâu, phương án thi công"),
    ]
    _add_table(doc,
        ["Khu vực", "Loại cọc", "L (m)", "Ghi chú"],
        rec_rows,
        col_widths=[3.5, 2.5, 1.5, 9.0])
    doc.add_paragraph()

    # ── 8. CHECKLIST ───────────────────────────────────────────────────────────
    _heading(doc, "8. CHECKLIST TRƯỚC THIẾT KẾ CHÍNH THỨC", size=13, color=(31, 73, 125))
    checks = [
        "Xác nhận su thực tế Lớp 1 từ thí nghiệm UU/CU — kiểm tra lại NT2 bằng GeoMCP",
        "Khảo sát bổ sung khu vực HK10 (Lớp 1 dày 25m — kiểm soát thiết kế NT1)",
        "Xác nhận chu vi SW-940 từ catalog gốc BETON 6 (hiện dùng 4.984 mm nội suy)",
        "Xác nhận ranh giới và chiều sâu XMD tại HK12 — điều chỉnh phương án thi công",
        "Kiểm tra Mcr: M_max (từ PLAXIS) < Mcr = 756.7 kN.m (SW-840)",
        "Kiểm tra ổn định đáy hố đào (basal heave) nếu có đào đất phía sau kè",
        "Xác nhận cao độ đỉnh kè +2.70m đồng nhất toàn tuyến hay thay đổi theo mặt bằng",
    ]
    for c in checks:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(c)
        _set_font(run, size=11)

    doc.add_paragraph()

    # ── Footer ──────────────────────────────────────────────────────────────────
    _para(doc,
          "Tài liệu này được lập tự động từ số liệu địa chất 12 hố khoan "
          "(260512 CVTT-TTHC. Tru DC.pdf) và catalog cọc SW BETON 6 Ver.2015. "
          "NT2 tính theo phương pháp đa lớp (alpha-method, su theo từng lớp thực tế, "
          "alpha Tomlinson 1980). Kết quả mang tính sơ bộ — cần xác nhận bởi "
          "kỹ sư địa kỹ thuật có thẩm quyền.",
          size=10, italic=True)

    # ── Save ────────────────────────────────────────────────────────────────────
    _OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(_OUT_FILE)
    print(f"Đã lưu: {_OUT_FILE}")


if __name__ == "__main__":
    build()
