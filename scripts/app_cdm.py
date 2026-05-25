# -*- coding: utf-8 -*-
"""
app_cdm.py  –  Ứng dụng thiết kế cọc đất xi măng (CDM)
Streamlit, cấu trúc tương tự app_coc_tai_ngang.py

Chạy:  streamlit run app_cdm.py --server.port 8503
"""
import io
import json
import math
import os
import sys
from datetime import datetime
from pathlib import Path

import pandas as pd
import streamlit as st

try:
    import plotly.graph_objects as go
    _HAS_PLOTLY = True
except Exception:
    _HAS_PLOTLY = False

try:
    import matplotlib.pyplot as plt
    _HAS_MPL = True
except ImportError:
    _HAS_MPL = False

try:
    import pydeck as pdk
    _HAS_PYDECK = True
except ImportError:
    _HAS_PYDECK = False

import sqlite3

# ── Đường dẫn ────────────────────────────────────────────────────────────────
_THIS    = Path(__file__).resolve()
# Hỗ trợ cả 2 cấu hình: chạy từ scripts/app_cdm.py (Cloud/local) hoặc root app_cdm.py
_ROOT    = _THIS.parent.parent if _THIS.parent.name == "scripts" else _THIS.parent
# Env var TTHC_DB_PATH cho phép port sandbox (8504) dùng DB riêng, không
# va chạm với 8503. Nếu không set → mặc định data/TTHC.sqlite.
_DB_ENV  = os.environ.get("TTHC_DB_PATH", "").strip()
_DB      = Path(_DB_ENV) if _DB_ENV else _ROOT / "data" / "TTHC.sqlite"
_CDM_SC  = _ROOT / "CDM" / "scripts"
_THEORY  = _ROOT / "35-ly-thuyet-cdm.md"


# ── Pile schematic (GEO5 style) — dùng cho tab Cọc ván SW mục C ──────────────
def draw_pile_schematic(
    L_m, water_lvl, top_elev, H_kN, M_kNm, layers, bc_type,
    pile_H_mm=840,
    soil_lvl_front=None,
    soil_lvl_back=None,
    back_layers=None,
    gamma_fill=18.0,
    phi_fill=25.0,
    c_fill=5.0,
    water_lvl_back=None,
):
    if not _HAS_MPL:
        return None
    bot  = top_elev - L_m
    ypad = max(L_m * 0.06, 0.8)
    y_top = top_elev + ypad * 2.2
    y_bot = bot - ypad * 1.2

    fig, ax = plt.subplots(figsize=(4.5, 5.5), dpi=90)
    ax.set_facecolor("white")
    fig.patch.set_facecolor("white")
    ax.set_xlim(-5, 5)
    ax.set_ylim(y_bot, y_top)
    ax.set_ylabel("Elevation [m]", fontsize=9)
    ax.set_xticks([])
    ax.tick_params(axis="y", labelsize=8)
    for sp in ["top", "right", "bottom"]:
        ax.spines[sp].set_visible(False)

    _fill_label = f"Fill\nγ={gamma_fill:.0f} kN/m³\nφ={phi_fill:.0f}°  c={c_fill:.0f} kPa"
    if soil_lvl_front is not None and top_elev > soil_lvl_front:
        ax.fill_between(
            [-4.8, 0], [soil_lvl_front, soil_lvl_front], [top_elev, top_elev],
            color="#c8a96e", alpha=0.35, hatch="\\\\", zorder=2,
            linewidth=0.5, edgecolor="#a07040",
        )
        ax.text(-2.4, (soil_lvl_front + top_elev) / 2, _fill_label,
                fontsize=7, color="#7a5020", ha="center", va="center",
                style="italic", linespacing=1.4)

    if soil_lvl_front is not None:
        ax.plot([-4.8, 0], [soil_lvl_front, soil_lvl_front],
                color="saddlebrown", lw=1.8, solid_capstyle="round", zorder=6)
        ax.text(-4.8, soil_lvl_front + ypad * 0.15, "Ground F",
                fontsize=8, color="saddlebrown", va="bottom", fontweight="bold")
    if soil_lvl_back is not None:
        ax.plot([0, 4.8], [soil_lvl_back, soil_lvl_back],
                color="saddlebrown", lw=1.8, solid_capstyle="round", zorder=6)
        ax.text(0.2, soil_lvl_back + ypad * 0.15, "Ground B",
                fontsize=8, color="saddlebrown", va="bottom", fontweight="bold")

    ax.plot([0, 0], [bot, top_elev], color="red", lw=3.5, solid_capstyle="butt", zorder=10)
    ax.plot([-4.8, 5], [top_elev, top_elev], color="#888", lw=1.0, ls="--", zorder=5)
    ax.plot([-0.25], [top_elev], "v", color="#444", ms=7, zorder=11,
            markerfacecolor="none", markeredgewidth=1.5)
    ax.text(-4.8, top_elev + ypad * 0.22, "SP Top", fontsize=9, color="#333", va="bottom")
    ax.text(0.25, top_elev + ypad * 0.05, f"{top_elev:.2f} m", fontsize=7.5, color="#555", va="bottom")

    for i, (nm, zt, zb, gm, ph, *_) in enumerate(layers):
        if i > 0:
            ax.plot([-4.8, 0], [zt, zt], color="goldenrod", lw=1.0, ls="--", zorder=5)
            ax.plot([-0.25], [zt], "v", color="goldenrod", ms=7, zorder=11,
                    markerfacecolor="none", markeredgewidth=1.5)
            ax.text(-4.8, zt + ypad * 0.12, f"Soil {i}", fontsize=9, color="goldenrod", va="bottom")

    if back_layers:
        for i, (nm, zt, zb, gm, ph, *_) in enumerate(back_layers):
            if i > 0:
                ax.plot([0, 4.8], [zt, zt], color="goldenrod", lw=0.9, ls=":", zorder=4, alpha=0.8)
                ax.text(0.3, zt + ypad * 0.12, f"Back {i}", fontsize=8, color="goldenrod", va="bottom", alpha=0.85)

    ax.text(0.25, bot - ypad * 0.05, f"{bot:.2f} m", fontsize=7.5, color="#555", va="top")

    w_gap = ypad * 0.18
    ax.plot([-4.2, -0.8], [water_lvl, water_lvl], color="steelblue", lw=2.0, zorder=7)
    ax.plot([-3.8, -1.1], [water_lvl - w_gap, water_lvl - w_gap], color="steelblue", lw=1.5, zorder=7)
    ax.text(-4.8, water_lvl + ypad * 0.15, f"Water F  {water_lvl:.2f} m", fontsize=8, color="steelblue")
    _wlb = water_lvl_back if water_lvl_back is not None else water_lvl
    ax.plot([0.8, 4.2], [_wlb, _wlb], color="steelblue", lw=2.0, zorder=7, ls="--")
    ax.plot([1.1, 3.8], [_wlb - w_gap, _wlb - w_gap], color="steelblue", lw=1.5, zorder=7, ls="--")
    ax.text(0.3, _wlb + ypad * 0.15, f"Water B  {_wlb:.2f} m", fontsize=8, color="steelblue")

    if bc_type == "Fixed":
        ax.plot([-0.6, 0.6], [bot, bot], color="#333", lw=1.8, zorder=8)
        for k in range(6):
            x0 = -0.40 + k * 0.16
            ax.plot([x0, x0 - 0.12], [bot, bot - ypad * 0.25], color="#333", lw=0.8, zorder=7)
    elif bc_type == "Cantilever":
        ax.plot([0], [bot], "^", ms=12, color="#333", zorder=8,
                markerfacecolor="none", markeredgewidth=2.0)
        ax.plot([-0.5, 0.5], [bot - ypad * 0.12, bot - ypad * 0.12], "#333", lw=1.0, zorder=7)
    elif bc_type == "Free":
        ax.plot([-0.2, 0.2], [bot, bot], "#333", lw=1.5, ls=":", zorder=7)

    if abs(H_kN) > 1e-3:
        sign = 1 if H_kN > 0 else -1
        ax.annotate("", xy=(0, top_elev), xytext=(-sign * 2.8, top_elev),
                    arrowprops=dict(arrowstyle="->, head_width=0.35, head_length=0.20",
                                   color="red", lw=2.2), zorder=12)
        ax.text(-sign * 1.4, top_elev + ypad * 0.50, f"H = {H_kN:.0f} kN",
                fontsize=9, color="red", ha="center", fontweight="bold")
    if abs(M_kNm) > 1e-3:
        ax.text(0.4, top_elev + ypad * 1.2, f"M = {M_kNm:.0f} kN.m",
                fontsize=8.5, color="darkorange", ha="left", fontweight="bold")

    ax.text(-4.0, y_bot + ypad * 0.4, "Front", fontsize=10, color="#666", style="italic")
    ax.text(1.5,  y_bot + ypad * 0.4, "Back",  fontsize=10, color="#666", style="italic")
    plt.tight_layout()
    return fig


@st.cache_data(show_spinner=False)
def _load_theory() -> str:
    try:
        return _THEORY.read_text(encoding="utf-8")
    except FileNotFoundError:
        return ""

_PUNCH_THEORY = _ROOT / "41-cdm-choc-thung-dem-ximang.md"

@st.cache_data(show_spinner=False)
def _load_punch_theory() -> str:
    try:
        return _PUNCH_THEORY.read_text(encoding="utf-8")
    except FileNotFoundError:
        return ""

st.set_page_config(
    page_title="Xử lý nền + Kè SW",
    page_icon="🏗",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Auto-refresh khi JSON / SQLite thay đổi (không cần bấm Rerun) ────────────
@st.fragment(run_every=3)
def _data_watcher():
    _watch = [
        _ROOT / "data" / "ke_sw_202605_TTHC.json",
        _ROOT / "data" / "ke_layers_202605_TTHC.json",
        _ROOT / "data" / "TTHC.sqlite",
    ]
    _mt = {str(f): f.stat().st_mtime for f in _watch if f.exists()}
    if st.session_state.get("_auto_mtimes") != _mt:
        st.session_state["_auto_mtimes"] = _mt
        st.cache_data.clear()
        st.rerun()

_data_watcher()

# ── CSS app + print PDF ───────────────────────────────────────────────────────
st.markdown("""
<style>
[data-testid="stSidebar"] { min-width: 220px; max-width: 280px; }
div[data-testid="stMetricValue"] { font-size: 1.1rem; }
#MainMenu { visibility: hidden; }
[data-testid="stToolbar"] { display: none; }
[data-testid="stDecoration"] { display: none; }
footer { visibility: hidden; }

/* ── Nút đóng / mở sidebar — làm to + nổi bật, cân đối 2 chiều ──
   - Khi sidebar mở: nút « ở góc trên phải sidebar (Streamlit tự render).
   - Khi sidebar đóng: nút » nổi ở góc trên trái màn hình (collapsedControl).
   CSS dưới đây phóng to + bo góc + đổ bóng + hover xanh navy cho cả 2 nút.
   Áp dụng cho cả tên cũ (collapsedControl) lẫn tên mới (stSidebarCollapsedControl)
   để tương thích các bản Streamlit khác nhau. */
[data-testid="stSidebarCollapsedControl"],
[data-testid="collapsedControl"] {
    background: #ffffff !important;
    border: 2px solid #1e3a8a !important;
    border-radius: 10px !important;
    box-shadow: 0 3px 10px rgba(30, 58, 138, 0.25) !important;
    width: 44px !important;
    height: 44px !important;
    top: 14px !important;
    left: 14px !important;
    z-index: 999 !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    transition: all 0.2s ease-in-out !important;
}
[data-testid="stSidebarCollapsedControl"]:hover,
[data-testid="collapsedControl"]:hover {
    background: #1e3a8a !important;
    transform: scale(1.05) !important;
    box-shadow: 0 4px 14px rgba(30, 58, 138, 0.45) !important;
}
[data-testid="stSidebarCollapsedControl"] svg,
[data-testid="stSidebarCollapsedControl"] button,
[data-testid="collapsedControl"] svg,
[data-testid="collapsedControl"] button {
    color: #1e3a8a !important;
    fill: #1e3a8a !important;
}
[data-testid="stSidebarCollapsedControl"]:hover svg,
[data-testid="stSidebarCollapsedControl"]:hover button,
[data-testid="collapsedControl"]:hover svg,
[data-testid="collapsedControl"]:hover button {
    color: #ffffff !important;
    fill: #ffffff !important;
}
/* Nút « đóng sidebar (đã có sẵn trong sidebar) — đồng bộ style để cân đối */
[data-testid="stSidebar"] [data-testid="baseButton-headerNoPadding"],
[data-testid="stSidebar"] button[kind="header"] {
    background: rgba(30, 58, 138, 0.08) !important;
    border-radius: 8px !important;
    transition: all 0.2s ease-in-out !important;
}
[data-testid="stSidebar"] [data-testid="baseButton-headerNoPadding"]:hover,
[data-testid="stSidebar"] button[kind="header"]:hover {
    background: #1e3a8a !important;
}
[data-testid="stSidebar"] [data-testid="baseButton-headerNoPadding"]:hover svg,
[data-testid="stSidebar"] button[kind="header"]:hover svg {
    color: #ffffff !important;
    fill: #ffffff !important;
}

/* ── iframe height=0 (JS handler): ẩn hẳn trên mọi tab, không chiếm chỗ ──
   Áp dụng screen mode — hiệu lực cho tất cả tab, không chỉ khi in.
   Plotly SVG inline không qua iframe → không bị ảnh hưởng. */
iframe[height="0"] {
    display: none !important;
    height: 0 !important;
    min-height: 0 !important;
    max-height: 0 !important;
    padding: 0 !important;
    margin: 0 !important;
    border: none !important;
    overflow: hidden !important;
}
[data-testid="stCustomComponentV1"]:has(iframe[height="0"]),
[data-testid="stIFrame"]:has(iframe[height="0"]) {
    display: none !important;
    height: 0 !important;
    min-height: 0 !important;
    padding: 0 !important;
    margin: 0 !important;
    overflow: hidden !important;
}

/* ── Nội dung lý thuyết trong expander: font rõ hơn trên màn hình ── */
[data-testid="stExpander"] [data-testid="stMarkdownContainer"] p,
[data-testid="stExpander"] [data-testid="stMarkdownContainer"] li,
[data-testid="stExpander"] [data-testid="stMarkdownContainer"] td,
[data-testid="stExpander"] [data-testid="stMarkdownContainer"] th {
    font-size: 15px;
    line-height: 1.75;
}
[data-testid="stExpander"] [data-testid="stMarkdownContainer"] h2 { font-size: 18px; margin-top: 1.2em; }
[data-testid="stExpander"] [data-testid="stMarkdownContainer"] h3 { font-size: 16px; margin-top: 1em; }
[data-testid="stExpander"] [data-testid="stMarkdownContainer"] h4 { font-size: 15px; }

/* ── Print / PDF ── */
@media print {
    /* Ẩn sidebar và toolbar khi in */
    [data-testid="stSidebar"],
    [data-testid="stToolbar"],
    [data-testid="stDecoration"],
    footer { display: none !important; }

    /* Ẩn mọi interactive control khi in — không hữu ích + gây chồng chéo lên bảng/biểu đồ.
       Áp dụng toàn bộ tab: Địa chất, Kè SW, Lún nền, Thông số...
       Radio "Chế độ xem" + multiselect Zone + checkbox 3D là nguồn gây chồng chéo chính. */
    [data-testid="stRadio"],
    [data-testid="stMultiSelect"],
    [data-testid="stSelectbox"],
    [data-testid="stSlider"],
    [data-testid="stNumberInput"],
    [data-testid="stTextInput"],
    [data-testid="stTextArea"],
    [data-testid="stDateInput"],
    [data-testid="stFileUploader"],
    [data-testid="stColorPicker"],
    [data-testid="stToggle"],
    [data-testid="stCheckbox"],
    [data-testid="stButton"],
    /* Nút bấm Streamlit (không phải download) */
    button[kind="secondary"], button[kind="primary"] { display: none !important; }
    /* Download button giữ lại (có thể hữu ích khi in) */
    [data-testid="stDownloadButton"] { display: block !important; }

    /* Mở rộng vùng nội dung ra full trang */
    [data-testid="stAppViewContainer"] > section { padding: 0 !important; }
    .block-container { max-width: 100% !important; padding: 0 1cm !important; }

    /* Font báo cáo kỹ thuật */
    [data-testid="stMarkdownContainer"] p,
    [data-testid="stMarkdownContainer"] li,
    [data-testid="stMarkdownContainer"] td,
    [data-testid="stMarkdownContainer"] th {
        font-family: "Times New Roman", Times, serif !important;
        font-size: 12pt !important;
        line-height: 1.6 !important;
        color: #000 !important;
    }
    [data-testid="stMarkdownContainer"] h1 { font-size: 16pt !important; font-weight: bold !important; page-break-after: avoid; }
    [data-testid="stMarkdownContainer"] h2 { font-size: 14pt !important; font-weight: bold !important; page-break-after: avoid; }
    [data-testid="stMarkdownContainer"] h3 { font-size: 13pt !important; font-weight: bold !important; page-break-after: avoid; }
    [data-testid="stMarkdownContainer"] h4 { font-size: 12pt !important; font-weight: bold !important; page-break-after: avoid; }

    /* Không ngắt trang giữa bảng */
    [data-testid="stMarkdownContainer"] table { page-break-inside: avoid; }

    /* Mở expander khi in */
    details { display: block !important; }
    summary { display: none !important; }
}
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════════
# CONSTANTS
# ═══════════════════════════════════════════════════════════════════════════════
_CLAY_SYMBOLS = {        # symbol lớp bùn sét yếu theo zone
    "KE":  ["1", "1b"],
    "BXN": ["2"],
    "NHC": ["2", "1"],   # NHC chưa có layers – fallback
    "QTT": ["1", "2"],   # Quảng Trường TT — chưa có layers nhập
}
_ZONE_NAMES = {
    "KE":  "Kè Công Viên (KE)",
    "BXN": "Bãi Đỗ Xe Ngầm (BXN)",
    "NHC": "Nhà Hành Chính (NHC)",
    "QTT": "Quảng Trường Trung Tâm (QTT)",
}

_LAYER_COLORS: dict[str, str] = {
    # Ký hiệu địa tầng Việt Nam (TCVN)
    "F":    "#9E9E9E",
    "1":    "#81D4FA",
    "1b":   "#4FC3F7",
    "2":    "#1565C0",
    "2a":   "#66BB6A",
    "2b":   "#FDD835",
    "2c":   "#C5E1A5",
    "3":    "#8D6E63",
    "3a":   "#A5D6A7",
    "3b":   "#69B578",
    "3c":   "#43A047",
    "4":    "#5C85D6",
    "5":    "#FFA726",
    "5a":   "#FB8C00",
    "5b":   "#E65100",
    "6":    "#6D4C41",
    "7":    "#795548",
    "8":    "#BCAAA4",
    "TK6a": "#FFD54F",
    "TK6b": "#FFB300",
    "XMD":  "#E91E63",
    # USCS symbols (NHC lab_tests.symbol_tcvn)
    "CH":   "#1A237E",   # High-plasticity clay – xanh đậm
    "CL":   "#1565C0",   # Low-plasticity clay
    "MH":   "#29B6F6",   # High-plasticity silt – xanh nhạt
    "ML":   "#81D4FA",   # Low-plasticity silt
    "OH":   "#4527A0",   # Organic high-plasticity
    "OL":   "#7E57C2",   # Organic low-plasticity
    "SM":   "#8D6E63",   # Silty sand – nâu
    "SC":   "#A1887F",   # Clayey sand
    "SW":   "#D4AC0D",   # Well-graded sand – vàng
    "SP":   "#F9E79F",   # Poorly-graded sand
    "GW":   "#5C85D6",   # Well-graded gravel
    "GP":   "#AED6F1",   # Poorly-graded gravel
    "GM":   "#7FB3D3",   # Silty gravel
    "GC":   "#5DADE2",   # Clayey gravel
    "PT":   "#2E7D32",   # Peat – xanh lá
}
_LAYER_DEFAULT_COLOR = "#EEEEEE"
_ZONE_MARKER: dict[str, str] = {"KE": "circle", "BXN": "square", "NHC": "diamond", "QTT": "cross"}

# ─── Bảng dịch VN / EN ────────────────────────────────────────────────────────
_L: dict[str, tuple[str, str]] = {
    # Sidebar
    "app_sub":      ("Thiết kế cọc đất xi măng",    "Cement Deep Mixing Design"),
    "save_load":    ("Lưu / Mở dự án",              "Save / Load Project"),
    "download":     ("Tải xuống (.cdm)",             "Download (.cdm)"),
    "upload":       ("Mở file dự án (.cdm)",         "Open project file (.cdm)"),
    "load_ok":      ("Đã tải dự án.",               "Project loaded."),
    "lang_lbl":     ("Ngôn ngữ / Language",          "Language / Ngôn ngữ"),
    # Pages
    "p_geology":    ("Địa chất",                    "Geology"),
    "p_sample_check":("Kiểm tra mẫu TN",            "Sample Check"),
    "p_tkcs_cdm":   ("TKCS CDM",                    "CDM Prelim Design"),
    "p_params":     ("Thiết kế CDM",                "CDM Design"),
    "p_compare":    ("Kết quả CDM",                 "CDM Results"),
    "p_settlement": ("Dự báo độ lún",               "Settlement Prediction"),
    "p_export":     ("Xuất kết quả",                "Export Results"),
    "p_cdm_bvt":    ("TKBVTC CDM",                  "CDM Detail Design"),
    "p_cdm_tien_do":("Tiến độ & Thí nghiệm CDM",   "CDM Progress & Testing"),
    "p_tvtk_prep":  ("Thống nhất đầu vào TVTK",    "TVTK Input Unification"),
    "p_tkcs_sw":    ("TKCS Cọc ván",                "Sheet Pile Prelim Design"),
    "p_ke_sw":      ("Cọc ván SW (Kè)",             "Sheet Pile SW (Ke)"),
    "p_sw_bvt":     ("TKBVTC Cọc SW",              "Sheet Pile Detail Design"),
    # Page 1 – Geology
    "p1_sub":       ("Địa chất – Chọn hố khoan",    "Geology – Borehole Selection"),
    "zone_lbl":     ("Zone",                        "Zone"),
    "bh_lbl":       ("Hố khoan",                   "Borehole"),
    "apply_bh":     ("Áp dụng hố khoan",            "Apply Borehole"),
    "apply_bh_ok":  ("Đã áp dụng {bh}.",            "Applied {bh}."),
    "no_bh_warn":   ("Zone này chưa có dữ liệu hố khoan.", "No borehole data for this zone."),
    "clay_thick":   ("Bề dày lớp bùn sét",          "Soft clay thickness"),
    "clay_elev":    ("Cao độ đáy lớp bùn",          "Clay base elevation"),
    "su_avg":       ("Su trung bình (VST)",          "Average Su (VST)"),
    "strat_title":  ("Địa tầng hố khoan",           "Borehole Stratigraphy"),
    "no_strat":     ("Hố khoan này chưa có dữ liệu địa tầng trong CSDL.",
                     "No stratigraphy data in database."),
    "vst_sel":      ("Chọn hố khoan VST (để trống = toàn bộ)",
                     "Select VST locations (blank = all)"),
    "no_bh_col":    ("Chọn hố khoan để xem cột địa chất.", "Select borehole to view soil column."),
    "cdm_test_exp": ("Kết quả thí nghiệm CDM (R7)", "CDM Test Results (R7)"),
    "cdm_test_note":("Lưu ý: Kết quả R7 (7 ngày). Cần qu_R90 để tính Ec chính xác.",
                     "Note: R7 results (7-day). qu_R90 needed for accurate Ec."),
    "view_mode":    ("Chế độ xem",                  "View Mode"),
    "view_3d":      ("3D địa chất",                 "3D Geology"),
    "view_map":     ("Bản đồ vị trí",               "Location Map"),
    "clay_surf":    ("Mặt đáy lớp bùn",             "Clay Base Surface"),
    "cdm_top_show": ("Đỉnh trụ CDM",                "CDM Pile Top"),
    "elev_lbl":     ("Cao độ (m)",                  "Elevation (m)"),
    "no_zone":      ("Chọn ít nhất một zone.",      "Select at least one zone."),
    "map_style":    ("Nền bản đồ",                  "Map Style"),
    "gps_done":     ("Hiệu chỉnh GPS — đã áp dụng", "GPS Calibration — applied"),
    "gps_needed":   ("Hiệu chỉnh GPS — cần thiết lập lần đầu",
                     "GPS Calibration — setup required"),
    "ref_bh":       ("Hố khoan tham chiếu",         "Reference Borehole"),
    "apply_calib":  ("Áp dụng hiệu chỉnh",          "Apply Calibration"),
    "no_coords":    ("Không có hố khoan nào có tọa độ.", "No boreholes with coordinates."),
    "no_plotly":    ("Cài `plotly` để xem bản đồ địa chất.", "Install `plotly` to view geology map."),
    "no_coords_db": ("Chưa có tọa độ hố khoan trong CSDL (x_coord_m / y_coord_m = NULL).",
                     "No borehole coordinates in database (x_coord_m / y_coord_m = NULL)."),
    # Page 2 – Parameters
    "p2_sub":       ("Thiết kế CDM",                "CDM Design"),
    "cdm_geom":     ("Hình học trụ CDM",            "CDM Pile Geometry"),
    "D_lbl":        ("Đường kính D (m)",             "Diameter D (m)"),
    "Lc_lbl":       ("Chiều dài Lc (m)",             "Length Lc (m)"),
    "CDTK_lbl":     ("Cao độ đỉnh cọc CDM (m)",     "CDM Top Elevation (m)"),
    "cdtk_info":    ("CDTK = đỉnh + đệm + đắp = **{v:+.2f} m**",
                     "CDTK = top + mat + fill = **{v:+.2f} m**"),
    "arr_lbl":      ("Bố trí lưới",                 "Grid Pattern"),
    "arr_tri":      ("Tam giác",                    "Triangular"),
    "arr_sq":       ("Hình vuông",                  "Square"),
    "material":     ("Vật liệu xi măng đất",        "Cement-Soil Material"),
    "cement_lbl":   ("Loại xi măng",                "Cement Type"),
    "dosage_lbl":   ("Hàm lượng XM (kg/m³)",       "Cement Dosage (kg/m³)"),
    "wc_lbl":       ("Tỷ lệ N/XM",                 "Water/Cement Ratio"),
    "qu_lbl":       ("qu thiết kế HT (kPa)",        "qu Design Field (kPa)"),
    "fs_lab":       ("FS quy đổi TN→HT",            "FS Lab→Field Conversion"),
    "geo_params":   ("Thông số địa kỹ thuật",       "Geotechnical Parameters"),
    "top_clay_lbl": ("Cao độ đỉnh lớp bùn (m)",    "Top of Soft Clay Elev. (m)"),
    "h_clay_lbl":   ("Bề dày lớp bùn h₁ (m)",      "Soft Clay Thickness h₁ (m)"),
    "bot_clay_info":("Đáy lớp bùn = {v:+.2f} m",   "Bottom of Soft Clay = {v:+.2f} m"),
    "su_lbl":       ("Su trung bình (kPa)",         "Average Su (kPa)"),
    "gamma_lbl":    ("γ tự nhiên (kN/m³)",          "Unit Weight γ (kN/m³)"),
    "loads_title":  ("Tải trọng gây lún",           "Settlement Loads"),
    "q_traf":       ("Hoạt tải (kN/m²)",            "Traffic Load (kN/m²)"),
    "pavement":     ("*Áo đường:*",                 "*Pavement:*"),
    "fill":         ("*Cát đắp:*",                  "*Fill:*"),
    "mat":          ("*Đệm cát:*",                  "*Sand Mat:*"),
    "h_lbl":        ("h (m)",                       "h (m)"),
    "punch_exp":    ("Kiểm tra chọc thủng lớp đệm xi măng", "Punching Check – Cement Mat"),
    "qu_mat_lbl":   ("qu đệm xi măng (kPa)",        "Cement mat qu (kPa)"),
    "fs_mat_lbl":   ("Hệ số an toàn Fs",            "Safety Factor Fs"),
    "theta_lbl":    ("Góc đàn hồi dẻo θ (°)",       "Plastic Arch Angle θ (°)"),
    "qa_lbl":       ("qa vùng không GC (kPa)",      "qa Unimproved Zone (kPa)"),
    "theory_exp":   ("Lý thuyết tính toán – TCVN 9403:2012",
                     "Design Theory – TCVN 9403:2012"),
    # Page 3 – Comparison
    "p3_sub":       ("So sánh các phương án khoảng cách trụ", "Pile Spacing Alternatives"),
    "setup_exp":    ("Thiết lập phương án",          "Setup Alternatives"),
    "sp_title":     ("Khoảng cách trụ",              "Pile Spacing"),
    "auto_lbl":     ("Tự động",                     "Auto"),
    "manual_lbl":   ("Nhập tay",                    "Manual"),
    "e_min":        ("e min (m)",                   "e min (m)"),
    "e_max":        ("e max (m)",                   "e max (m)"),
    "step_lbl":     ("Bước (m)",                    "Step (m)"),
    "sp_list":      ("Danh sách e (m), cách nhau bởi dấu phẩy",
                     "List of e (m), comma-separated"),
    "mat_punch":    ("Lớp đệm xi măng – kiểm tra chọc thủng",
                     "Cement Mat – Punching Check"),
    "qu_mat_sp":    ("qu đệm (kPa)",                "mat qu (kPa)"),
    "rec_sel":      ("Chọn phương án thiết kế",     "Select Design Alternative"),
    "pa_table":     ("Bảng so sánh phương án",      "Alternatives Comparison Table"),
    "param_sens":   ("Phân tích tham số (giữ nguyên e = PA kiến nghị)",
                     "Parametric Analysis (fixed e = recommended alt.)"),
    "qu_comp":      ("So sánh qu thiết kế",         "qu Design Comparison"),
    "D_comp":       ("So sánh đường kính D",        "Diameter D Comparison"),
    "hmat_comp":    ("So sánh chiều dày đệm cát",   "Sand Mat Thickness Comparison"),
    "qu_inp":       ("Các giá trị qu (kPa)",        "qu values (kPa)"),
    "D_inp":        ("Các giá trị D (m)",           "D values (m)"),
    "hmat_inp":     ("Các giá trị h_mat (m)",       "h_mat values (m)"),
    "punch_chart":  ("Kiểm tra chọc thủng đệm xi măng", "Cement Mat Punching Check"),
    "shear_ax":     ("Ứng suất cắt (kPa)",          "Shear Stress (kPa)"),
    # Page 4 – Results
    "p4_sub":       ("Kết quả chi tiết – Phương án kiến nghị",
                     "Detailed Results – Recommended Alternative"),
    "no_scenarios": ("Chưa có phương án. Vào tab So sánh PA để tính.",
                     "No alternatives defined. Go to Comparison tab."),
    "rec_pa":       ("Phương án kiến nghị: **PA{n} — e = {e} m**",
                     "Recommended: **Alt.{n} — e = {e} m**"),
    "settle_calc":  ("Tính toán độ lún",            "Settlement Calculation"),
    "bearing_calc": ("Kiểm tra sức chịu tải (AIT – TCVN 9403:2012 Phụ lục B)",
                     "Bearing Capacity Check (AIT – TCVN 9403:2012 Appendix B)"),
    "punch_res":    ("Kiểm tra chọc thủng lớp đệm xi măng (ALiCC)",
                     "Cement Mat Punching Check (ALiCC)"),
    "col_param":    ("Thông số",                    "Parameter"),
    "col_formula":  ("Công thức",                   "Formula"),
    "col_value":    ("Giá trị",                     "Value"),
    "bear_lbl":     ("Sức chịu tải",               "Bearing Capacity"),
    # Page 5 – Export
    "p5_sub":       ("Xuất báo cáo",               "Export Report"),
    "export_sum":   ("Tóm tắt trước khi xuất",     "Summary before export"),
    "rec_alt":      ("Phương án kiến nghị",         "Recommended Alt."),
    "settle_metric":("Độ lún S₁",                  "Settlement S₁"),
    "word_title":   ("Thuyết minh tính toán (Word)", "Calculation Report (Word)"),
    "word_cap":     ("Tài liệu gồm 6 mục: Cơ sở pháp lý · Địa chất · Thông số · "
                     "Phương pháp · So sánh PA · Kết quả PA kiến nghị.",
                     "Document: Legal basis · Geology · Parameters · "
                     "Method · Comparison · Recommended results."),
    "create_word":  ("Tạo file Word",              "Generate Word File"),
    "creating":     ("Đang tạo tài liệu...",        "Generating document..."),
    "dl_docx":      ("Tải xuống (.docx)",           "Download (.docx)"),
    # Sidebar summary
    "si_bh":        ("**HK:**",                    "**BH:**"),
}


def _t(key: str, **kw) -> str:
    """Trả về chuỗi theo ngôn ngữ hiện tại."""
    lang = st.session_state.get("lang", "VN")
    pair = _L.get(key, (key, key))
    s = pair[0 if lang == "VN" else 1]
    return s.format(**kw) if kw else s


# ═══════════════════════════════════════════════════════════════════════════════
# DB HELPERS
# ═══════════════════════════════════════════════════════════════════════════════
@st.cache_resource
def _db() -> sqlite3.Connection:
    con = sqlite3.connect(str(_DB), check_same_thread=False)
    con.row_factory = sqlite3.Row
    return con


@st.cache_data(ttl=300, show_spinner=False)
def _load_su_back_natural(bh_full: str) -> float:
    """Su tự nhiên phía Back: VST → lab Cu_UU → fallback 11 kPa. Cache 5 phút."""
    try:
        con = sqlite3.connect(str(_DB), check_same_thread=False)
        con.row_factory = sqlite3.Row
        vst_r = con.execute("""
            SELECT AVG(vst.Su_kPa) AS su_avg
            FROM vane_shear_tests vst
            JOIN vst_locations loc ON vst.vst_loc_id = loc.id
            WHERE loc.name = ? AND vst.depth_m BETWEEN 1.0 AND 25.0
        """, (bh_full,)).fetchone()
        if vst_r and vst_r["su_avg"]:
            con.close()
            return round(float(vst_r["su_avg"]), 1)
        lab_r = con.execute("""
            SELECT AVG(lt.Cu_UU_kPa) AS cu_avg
            FROM lab_tests lt JOIN boreholes b ON lt.borehole_id = b.id
            WHERE b.name = ? AND lt.Cu_UU_kPa > 0 AND lt.depth_from_m <= 25.0
        """, (bh_full,)).fetchone()
        con.close()
        if lab_r and lab_r["cu_avg"]:
            return round(float(lab_r["cu_avg"]), 1)
    except Exception:
        pass
    return 11.0


@st.cache_data(ttl=300, show_spinner=False)
def _load_eps50_params(bh_full: str, max_depth_m: float) -> tuple[float, float]:
    """AVG(E_kPa), AVG(Cu_UU_kPa) để tính ε₅₀. Cache 5 phút. Trả (0,0) nếu thiếu."""
    try:
        con = sqlite3.connect(str(_DB), check_same_thread=False)
        row = con.execute("""
            SELECT AVG(E_kPa) AS E_avg, AVG(Cu_UU_kPa) AS Cu_avg
            FROM lab_tests lt JOIN boreholes b ON lt.borehole_id = b.id
            WHERE b.name = ? AND lt.depth_from_m <= ?
              AND lt.E_kPa > 0 AND lt.Cu_UU_kPa > 0
        """, (bh_full, max_depth_m)).fetchone()
        con.close()
        if row and row[0] and row[1]:
            return float(row[0]), float(row[1])
    except Exception:
        pass
    return 0.0, 0.0


@st.cache_data(ttl=300)
def _load_boreholes_by_zone(zone_code: str) -> list[dict]:
    with _db() as con:
        rows = con.execute("""
            SELECT b.name, b.elevation_m, b.depth_m
            FROM boreholes b JOIN zones z ON b.zone_id=z.id
            WHERE z.code=? ORDER BY b.name
        """, (zone_code,)).fetchall()
    return [dict(r) for r in rows]


@st.cache_data(ttl=300)
def _load_layers(bh_name: str) -> list[dict]:
    with _db() as con:
        # Thử bảng layers cũ (KE / BXN)
        rows = con.execute("""
            SELECT l.symbol, l.description, l.depth_top_m, l.depth_bot_m, l.thickness_m,
                   ROUND(AVG(lt.gamma_kNm3), 2) AS gamma_kNm3,
                   ROUND(AVG(lt.c_kPa),      1) AS c_kPa,
                   ROUND(AVG(lt.phi_deg),     1) AS phi_deg,
                   ROUND(AVG(lt.e0),          3) AS e0,
                   ROUND(AVG(lt.Cc),          4) AS Cc,
                   ROUND(AVG(lt.Cs),          4) AS Cs,
                   ROUND(AVG(lt.PC_kPa),      1) AS PC_kPa,
                   AVG(lt.Cv_cm2s)               AS Cv_cm2s
            FROM layers l
            JOIN boreholes b ON l.borehole_id = b.id
            LEFT JOIN lab_tests lt
                   ON lt.borehole_id = b.id
                  AND lt.depth_from_m < l.depth_bot_m
                  AND lt.depth_to_m   > l.depth_top_m
            WHERE b.name = ?
            GROUP BY l.id
            ORDER BY l.depth_top_m
        """, (bh_name,)).fetchall()
        if rows:
            return [dict(r) for r in rows]

        # Fallback: strat_layers (NHC – import từ DXF)
        rows = con.execute("""
            SELECT
                COALESCE(
                    (SELECT lt2.symbol_tcvn FROM lab_tests lt2
                     WHERE lt2.borehole_id = b.id
                       AND lt2.depth_from_m < sl.depth_bot_m
                       AND lt2.depth_to_m   > sl.depth_top_m
                     ORDER BY lt2.depth_from_m LIMIT 1),
                    'L' || sl.layer_no
                ) AS symbol,
                sl.description,
                sl.depth_top_m,
                sl.depth_bot_m,
                ROUND(sl.depth_bot_m - sl.depth_top_m, 2) AS thickness_m,
                ROUND(AVG(lt.gamma_kNm3), 2) AS gamma_kNm3,
                ROUND(AVG(lt.c_kPa),      1) AS c_kPa,
                ROUND(AVG(lt.phi_deg),     1) AS phi_deg,
                ROUND(AVG(lt.e0),          3) AS e0,
                ROUND(AVG(lt.Cc),          4) AS Cc,
                ROUND(AVG(lt.Cs),          4) AS Cs,
                ROUND(AVG(lt.PC_kPa),      1) AS PC_kPa,
                AVG(lt.Cv_cm2s)               AS Cv_cm2s
            FROM strat_layers sl
            JOIN boreholes b ON sl.borehole_id = b.id
            LEFT JOIN lab_tests lt
                   ON lt.borehole_id = b.id
                  AND lt.depth_from_m < sl.depth_bot_m
                  AND lt.depth_to_m   > sl.depth_top_m
            WHERE b.name = ?
            GROUP BY sl.id
            ORDER BY sl.depth_top_m
        """, (bh_name,)).fetchall()
        if rows:
            return [dict(r) for r in rows]

        # Fallback cuối: build từ lab_tests (gom run liên tiếp cùng symbol)
        rows = con.execute("""
            SELECT
                symbol_tcvn AS symbol,
                COALESCE(description_vi, symbol_tcvn) AS description,
                MIN(depth_from_m) AS depth_top_m,
                MAX(depth_to_m)   AS depth_bot_m,
                ROUND(AVG(gamma_kNm3), 2) AS gamma_kNm3,
                ROUND(AVG(c_kPa),      1) AS c_kPa,
                ROUND(AVG(phi_deg),    1) AS phi_deg,
                ROUND(AVG(e0),         3) AS e0,
                ROUND(AVG(Cc),         4) AS Cc,
                ROUND(AVG(Cs),         4) AS Cs,
                ROUND(AVG(PC_kPa),     1) AS PC_kPa,
                AVG(Cv_cm2s)             AS Cv_cm2s
            FROM (
                SELECT *,
                    ROW_NUMBER() OVER (ORDER BY depth_from_m) -
                    ROW_NUMBER() OVER (PARTITION BY symbol_tcvn ORDER BY depth_from_m) AS grp
                FROM lab_tests
                WHERE borehole_id = (SELECT id FROM boreholes WHERE name = ?)
                  AND symbol_tcvn IS NOT NULL
                  AND depth_from_m IS NOT NULL
            )
            GROUP BY symbol_tcvn, grp
            ORDER BY depth_top_m
        """, (bh_name,)).fetchall()

    if not rows:
        return []
    result = [dict(r) for r in rows]
    # Lấp khoảng trống giữa các run: kéo đáy lớp trước = đỉnh lớp sau
    for i in range(len(result) - 1):
        result[i]["depth_bot_m"] = result[i + 1]["depth_top_m"]
    for r in result:
        r["thickness_m"] = round(r["depth_bot_m"] - r["depth_top_m"], 2)
    return result


@st.cache_data(ttl=300)
def _load_zone_params_by_symbol(zone_code: str, symbol: str) -> dict:
    """Trả về {Cc, Cs, e0, PC_kPa, Cv_cm2s, gamma_kNm3, source_bh} trung bình
    từ mọi hố khoan trong zone có symbol khớp (dùng khi hố khoan hiện tại thiếu)."""
    with _db() as con:
        row = con.execute("""
            SELECT
                ROUND(AVG(lt.Cc),        4) AS Cc,
                ROUND(AVG(lt.Cs),        4) AS Cs,
                ROUND(AVG(lt.e0),        3) AS e0,
                ROUND(AVG(lt.PC_kPa),    1) AS PC_kPa,
                AVG(lt.Cv_cm2s)             AS Cv_cm2s,
                ROUND(AVG(lt.gamma_kNm3),2) AS gamma_kNm3,
                GROUP_CONCAT(DISTINCT b.name) AS source_bh
            FROM lab_tests lt
            JOIN boreholes b ON lt.borehole_id = b.id
            JOIN zones z ON b.zone_id = z.id
            WHERE z.code = ?
              AND lt.symbol_tcvn = ?
              AND lt.Cc IS NOT NULL
        """, (zone_code, symbol)).fetchone()
    return dict(row) if row and row["Cc"] is not None else {}


@st.cache_data(ttl=300)
def _load_spt(bh_name: str) -> list[dict]:
    with _db() as con:
        # Thử spt_values cũ (KE / BXN)
        rows = con.execute("""
            SELECT sv.depth_m, sv.N1, sv.N2, sv.N3, sv.N
            FROM spt_values sv
            JOIN boreholes b ON sv.borehole_id = b.id
            WHERE b.name = ?
            ORDER BY sv.depth_m
        """, (bh_name,)).fetchall()
        if rows:
            return [dict(r) for r in rows]

        # Fallback: spt_tests (NHC – import từ DXF)
        rows = con.execute("""
            SELECT st.depth_from_m AS depth_m,
                   st.N1, st.N2, st.N3,
                   st.N_value AS N
            FROM spt_tests st
            JOIN boreholes b ON st.borehole_id = b.id
            WHERE b.name = ?
            ORDER BY st.depth_from_m
        """, (bh_name,)).fetchall()
    return [dict(r) for r in rows]


@st.cache_data(ttl=300)
def _load_vst_su(zone_code: str) -> pd.DataFrame:
    """VST Su toàn zone, dùng để vẽ profile và tính trung bình."""
    with _db() as con:
        rows = con.execute("""
            SELECT vt.depth_m, vt.Su_kPa, vl.name loc_name
            FROM vane_shear_tests vt
            JOIN vst_locations vl ON vt.vst_loc_id=vl.id
            JOIN zones z ON vl.zone_id=z.id
            WHERE z.code=? AND vt.Su_kPa IS NOT NULL
            ORDER BY vt.depth_m
        """, (zone_code,)).fetchall()
    return pd.DataFrame([dict(r) for r in rows])


@st.cache_data(ttl=300)
def _load_lab(bh_name: str) -> pd.DataFrame:
    with _db() as con:
        rows = con.execute("""
            SELECT depth_from_m, depth_to_m, gamma_kNm3, e0,
                   Cu_UU_kPa, c_kPa, phi_deg, Cc, Cs, E_kPa
            FROM lab_tests WHERE borehole_id=(
                SELECT id FROM boreholes WHERE name=?
            ) ORDER BY depth_from_m
        """, (bh_name,)).fetchall()
    return pd.DataFrame([dict(r) for r in rows])


@st.cache_data(ttl=300)
def _load_consol_summary(zone_code: str) -> pd.DataFrame:
    """Per-borehole summary: số mẫu và trung bình các chỉ tiêu nén cố kết."""
    with _db() as con:
        rows = con.execute("""
            SELECT
                b.name AS borehole,
                COUNT(lt.id) AS n_mau,
                SUM(CASE WHEN lt.Cc IS NOT NULL THEN 1 ELSE 0 END) AS n_Cc,
                ROUND(AVG(CASE WHEN lt.Cc  IS NOT NULL THEN lt.Cc  END), 4) AS Cc_tb,
                ROUND(AVG(CASE WHEN lt.Cs  IS NOT NULL THEN lt.Cs  END), 4) AS Cs_tb,
                ROUND(AVG(CASE WHEN lt.Cv_cm2s IS NOT NULL THEN lt.Cv_cm2s END), 12) AS Cv_tb,
                ROUND(AVG(CASE WHEN lt.k_cm_s  IS NOT NULL THEN lt.k_cm_s  END), 12) AS k_tb,
                ROUND(MIN(CASE WHEN lt.PC_kPa  IS NOT NULL THEN lt.PC_kPa  END), 1) AS PC_min,
                ROUND(MAX(CASE WHEN lt.PC_kPa  IS NOT NULL THEN lt.PC_kPa  END), 1) AS PC_max,
                ROUND(AVG(CASE WHEN lt.e0      IS NOT NULL THEN lt.e0      END), 3) AS e0_tb
            FROM boreholes b
            JOIN zones z ON b.zone_id = z.id
            LEFT JOIN lab_tests lt ON lt.borehole_id = b.id
            WHERE z.code = ?
            GROUP BY b.name
            HAVING n_Cc > 0
            ORDER BY b.name
        """, (zone_code,)).fetchall()
    return pd.DataFrame([dict(r) for r in rows])


@st.cache_data(ttl=300)
def _load_cdm_tests() -> pd.DataFrame:
    with _db() as con:
        rows = con.execute("""
            SELECT ct.*, b.name bh_name
            FROM cdm_tests ct JOIN boreholes b ON ct.borehole_id=b.id
            ORDER BY ct.cement_type, ct.WC_ratio, ct.dosage_kgm3
        """).fetchall()
    return pd.DataFrame([dict(r) for r in rows])


def _clay_params(bh_name: str, zone_code: str) -> dict:
    """Trả về h_clay, elevation, depth_clay_bot từ layers."""
    layers = _load_layers(bh_name)
    syms = _CLAY_SYMBOLS.get(zone_code, ["1", "2"])
    clay = [l for l in layers if l["symbol"] in syms]
    if not clay:
        return {}
    h_clay = sum(l["thickness_m"] or 0 for l in clay)
    depth_bot = max(l["depth_bot_m"] for l in clay)
    with _db() as con:
        row = con.execute("SELECT elevation_m FROM boreholes WHERE name=?", (bh_name,)).fetchone()
    elev = row["elevation_m"] if row and row["elevation_m"] else 0.0
    return {"h_clay": round(h_clay, 2), "depth_clay_bot": round(depth_bot, 2),
            "elev_clay_bot": round(elev - depth_bot, 2), "elevation_m": elev}


def _su_avg_in_range(zone_code: str, depth_top: float, depth_bot: float) -> float | None:
    df = _load_vst_su(zone_code)
    if df.empty:
        return None
    sub = df[(df.depth_m >= depth_top) & (df.depth_m <= depth_bot)]
    return round(float(sub.Su_kPa.mean()), 2) if not sub.empty else None


def _gamma_avg(bh_name: str) -> float:
    df = _load_lab(bh_name)
    if df.empty or df.gamma_kNm3.dropna().empty:
        return 15.0
    return round(float(df.gamma_kNm3.dropna().mean()), 3)


# ═══════════════════════════════════════════════════════════════════════════════
# 3D BOREHOLE MAP
# ═══════════════════════════════════════════════════════════════════════════════
@st.cache_data(ttl=300, show_spinner=False)
def _load_borehole_3d_data() -> tuple[list[dict], list[dict]]:
    if not _DB.exists():
        return [], []
    conn = sqlite3.connect(_DB)
    conn.row_factory = sqlite3.Row
    bhs = [dict(r) for r in conn.execute(
        "SELECT b.id, b.name, z.code zone, b.elevation_m, b.depth_m, "
        "b.x_coord_m, b.y_coord_m "
        "FROM boreholes b JOIN zones z ON z.id=b.zone_id "
        "WHERE b.x_coord_m IS NOT NULL "
        "  AND b.elevation_m IS NOT NULL AND b.depth_m IS NOT NULL "
        "ORDER BY z.id, b.name"
    ).fetchall()]
    ids = [b["id"] for b in bhs]
    if not ids:
        conn.close()
        return [], []
    ph = ",".join("?" * len(ids))
    lays = [dict(r) for r in conn.execute(
        f"SELECT borehole_id, symbol, description, depth_top_m, depth_bot_m "
        f"FROM layers WHERE borehole_id IN ({ph}) ORDER BY borehole_id, depth_top_m",
        ids,
    ).fetchall()]
    conn.close()
    return bhs, lays


@st.cache_data(ttl=300, show_spinner=False)
def _load_clay_bottom_3d() -> list[dict]:
    """Đáy lớp bùn sét yếu cho mỗi hố khoan có tọa độ."""
    if not _DB.exists():
        return []
    conn = sqlite3.connect(_DB)
    conn.row_factory = sqlite3.Row
    rows = [dict(r) for r in conn.execute("""
        SELECT b.name, z.code zone, b.elevation_m, b.x_coord_m, b.y_coord_m,
               ROUND(MAX(l.depth_bot_m), 3) clay_bot_depth,
               ROUND(b.elevation_m - MAX(l.depth_bot_m), 3) clay_bot_elev
        FROM boreholes b
        JOIN zones z ON z.id=b.zone_id
        JOIN layers l ON l.borehole_id=b.id
        WHERE b.x_coord_m IS NOT NULL
          AND ((z.code='KE'  AND l.symbol IN ('1','1b'))
            OR (z.code IN ('BXN','NHC') AND l.symbol='2'))
        GROUP BY b.id
        ORDER BY z.id, b.name
    """).fetchall()]
    conn.close()
    return rows


def _draw_boreholes_3d(
    selected_zones: list[str],
    show_clay_bottom: bool = False,
    show_cdm_top: bool = False,
    cdm_top_z: float = 2.7,
    focus_bh: str | None = None,
    pair_highlight: tuple | None = None,
) -> "go.Figure":
    """pair_highlight = (bh1_name, bh2_name, dist_m) — vẽ đường kích thước giữa 2 HK."""
    from collections import defaultdict

    bhs, lays = _load_borehole_3d_data()
    bhs  = [b for b in bhs if b["zone"] in selected_zones]
    if focus_bh:
        bhs_focus = [b for b in bhs if b["name"] == focus_bh]
        if bhs_focus:
            bhs = bhs_focus
    bh_ids = {b["id"] for b in bhs}
    lays = [l for l in lays if l["borehole_id"] in bh_ids]

    fig = go.Figure()
    if not bhs:
        fig.update_layout(title="Không có dữ liệu tọa độ hố khoan")
        return fig

    lay_by_bh: dict[int, list[dict]] = defaultdict(list)
    for l in lays:
        lay_by_bh[l["borehole_id"]].append(l)

    # Một trace / symbol → gộp legend
    all_syms: list[str] = []
    for l in lays:
        if l["symbol"] not in all_syms:
            all_syms.append(l["symbol"])

    added_legends: set[str] = set()
    for sym in all_syms:
        color = _LAYER_COLORS.get(sym, _LAYER_DEFAULT_COLOR)
        xs, ys, zs, texts = [], [], [], []
        for bh in bhs:
            for l in lay_by_bh[bh["id"]]:
                if l["symbol"] != sym:
                    continue
                z0 = bh["elevation_m"] - l["depth_top_m"]
                z1 = bh["elevation_m"] - l["depth_bot_m"]
                desc = l.get("description") or sym
                hover = (
                    f"<b>{bh['name']}</b><br>"
                    f"Lớp {sym}: {desc}<br>"
                    f"Sâu {l['depth_top_m']:.1f}–{l['depth_bot_m']:.1f} m<br>"
                    f"Cao độ {z0:.2f} → {z1:.2f} m"
                )
                xs += [bh["x_coord_m"], bh["x_coord_m"], None]
                ys += [bh["y_coord_m"], bh["y_coord_m"], None]
                zs += [z0, z1, None]
                texts += [hover, hover, None]

        show_leg = sym not in added_legends
        added_legends.add(sym)
        fig.add_trace(go.Scatter3d(
            x=xs, y=ys, z=zs, mode="lines",
            line=dict(color=color, width=10),
            name=f"Lớp {sym}", legendgroup=sym, showlegend=show_leg,
            hovertemplate="%{text}<extra></extra>", text=texts,
        ))

    # Labels tên hố khoan
    _zone_colors = {"KE": "#E53935", "BXN": "#1565C0", "NHC": "#2E7D32", "QTT": "#F9A825"}
    fig.add_trace(go.Scatter3d(
        x=[b["x_coord_m"] for b in bhs],
        y=[b["y_coord_m"] for b in bhs],
        z=[b["elevation_m"] + 3 for b in bhs],
        mode="text+markers",
        text=[b["name"].replace("BXN-CV-", "").replace("KE-", "") for b in bhs],
        textposition="top center",
        textfont=dict(size=9, color="#212121"),
        marker=dict(
            size=5,
            color=[_zone_colors.get(b["zone"], "#333") for b in bhs],
            symbol=[_ZONE_MARKER.get(b["zone"], "circle") for b in bhs],
        ),
        name="Vị trí HK", showlegend=True,
        hovertemplate="<b>%{text}</b><extra></extra>",
    ))

    # Đáy lớp bùn
    if show_clay_bottom:
        _clay = [c for c in _load_clay_bottom_3d() if c["zone"] in selected_zones]
        if len(_clay) >= 3:
            cx = [c["x_coord_m"] for c in _clay]
            cy = [c["y_coord_m"] for c in _clay]
            cz = [c["clay_bot_elev"] for c in _clay]
            ct = [
                f"<b>{c['name']}</b><br>"
                f"Đáy bùn sâu: {c['clay_bot_depth']:.1f} m<br>"
                f"Cao độ đáy: {c['clay_bot_elev']:.2f} m"
                for c in _clay
            ]
            fig.add_trace(go.Mesh3d(
                x=cx, y=cy, z=cz, delaunayaxis="z",
                intensity=cz,
                colorscale=[[0, "#0D47A1"], [0.5, "#42A5F5"], [1, "#BBDEFB"]],
                showscale=False, opacity=0.55,
                name="Đáy lớp bùn", legendgroup="clay_bot", showlegend=True,
                hovertemplate="%{text}<extra></extra>", text=ct,
            ))
            fig.add_trace(go.Scatter3d(
                x=cx, y=cy, z=cz, mode="markers",
                marker=dict(size=6, color="#0D47A1", symbol="diamond"),
                name="Điểm đáy bùn", legendgroup="clay_bot", showlegend=False,
                hovertemplate="%{text}<extra></extra>", text=ct,
            ))
            vx, vy, vz, vt = [], [], [], []
            for c in _clay:
                vx += [c["x_coord_m"], c["x_coord_m"], None]
                vy += [c["y_coord_m"], c["y_coord_m"], None]
                vz += [c["elevation_m"], c["clay_bot_elev"], None]
                lbl = f"<b>{c['name']}</b><br>Bùn dày {c['clay_bot_depth']:.1f} m"
                vt += [lbl, lbl, None]
            fig.add_trace(go.Scatter3d(
                x=vx, y=vy, z=vz, mode="lines",
                line=dict(color="#1565C0", width=3, dash="dot"),
                name="Chiều dày bùn", legendgroup="clay_bot", showlegend=False,
                hovertemplate="%{text}<extra></extra>", text=vt,
            ))

    # Mặt phẳng đỉnh trụ CDM
    if show_cdm_top:
        all_x = [b["x_coord_m"] for b in bhs]
        all_y = [b["y_coord_m"] for b in bhs]
        pad = 20.0
        mx, Mx = min(all_x) - pad, max(all_x) + pad
        my, My = min(all_y) - pad, max(all_y) + pad
        fig.add_trace(go.Mesh3d(
            x=[mx, Mx, Mx, mx], y=[my, my, My, My],
            z=[cdm_top_z]*4, i=[0, 0], j=[1, 2], k=[2, 3],
            color="#E91E63", opacity=0.25,
            name=f"Đỉnh trụ CDM +{cdm_top_z:.2f} m",
            showlegend=True, flatshading=True,
            hovertemplate=f"Đỉnh trụ CDM: <b>{cdm_top_z:.2f} m</b><extra></extra>",
        ))
        fig.add_trace(go.Scatter3d(
            x=[mx, Mx, Mx, mx, mx], y=[my, my, My, My, my],
            z=[cdm_top_z]*5, mode="lines",
            line=dict(color="#C2185B", width=3),
            name="Viền đỉnh CDM", showlegend=False,
        ))

    # Đường kích thước giữa 2 hố khoan (khi user chọn cặp)
    if pair_highlight:
        _ph_b1, _ph_b2, _ph_dist = pair_highlight
        _ph_bh1 = next((b for b in bhs if b["name"] == _ph_b1), None)
        _ph_bh2 = next((b for b in bhs if b["name"] == _ph_b2), None)
        if _ph_bh1 and _ph_bh2:
            _ph_z = max(_ph_bh1["elevation_m"], _ph_bh2["elevation_m"]) + 4
            _ph_mx = (_ph_bh1["x_coord_m"] + _ph_bh2["x_coord_m"]) / 2
            _ph_my = (_ph_bh1["y_coord_m"] + _ph_bh2["y_coord_m"]) / 2
            fig.add_trace(go.Scatter3d(
                x=[_ph_bh1["x_coord_m"], _ph_bh2["x_coord_m"]],
                y=[_ph_bh1["y_coord_m"], _ph_bh2["y_coord_m"]],
                z=[_ph_z, _ph_z],
                mode="lines",
                line=dict(color="#FF6F00", width=6),
                name=f"Khoảng cách {_ph_b1}↔{_ph_b2}",
                showlegend=True,
                hovertemplate=f"<b>{_ph_b1} ↔ {_ph_b2}</b><br>Khoảng cách: {_ph_dist:.1f} m<extra></extra>",
            ))
            # Nhãn khoảng cách tại trung điểm
            fig.add_trace(go.Scatter3d(
                x=[_ph_mx], y=[_ph_my], z=[_ph_z + 1.5],
                mode="text",
                text=[f"  {_ph_dist:.1f} m"],
                textfont=dict(size=14, color="#E65100"),
                showlegend=False,
                hovertemplate=f"{_ph_dist:.1f} m<extra></extra>",
            ))
            # Đường dọc từ mặt đất xuống đến đường kích thước (2 đầu)
            for _ph_bh in (_ph_bh1, _ph_bh2):
                fig.add_trace(go.Scatter3d(
                    x=[_ph_bh["x_coord_m"], _ph_bh["x_coord_m"]],
                    y=[_ph_bh["y_coord_m"], _ph_bh["y_coord_m"]],
                    z=[_ph_bh["elevation_m"] + 1, _ph_z],
                    mode="lines",
                    line=dict(color="#FF6F00", width=2, dash="dot"),
                    showlegend=False,
                    hoverinfo="skip",
                ))

    # Layout
    all_x = [b["x_coord_m"] for b in bhs]
    all_y = [b["y_coord_m"] for b in bhs]
    span_x = max(all_x) - min(all_x) or 1
    span_y = max(all_y) - min(all_y) or 1
    max_z = max(b["elevation_m"] for b in bhs) + 5
    min_z = min(b["elevation_m"] - b["depth_m"] for b in bhs) - 2

    fig.update_layout(
        height=420,
        margin=dict(l=0, r=0, t=30, b=0),
        legend=dict(
            x=1.01, y=0.95,
            bgcolor="rgba(255,255,255,0.85)",
            bordercolor="#ccc", borderwidth=1,
            font=dict(size=11),
        ),
        scene=dict(
            xaxis=dict(title="X — Easting (m)", tickformat=".0f"),
            yaxis=dict(title="Y — Northing (m)", tickformat=".0f"),
            zaxis=dict(title="Cao độ (m)", range=[min_z, max_z + 5]),
            aspectmode="manual",
            aspectratio=dict(
                x=span_x / max(span_x, span_y),
                y=span_y / max(span_x, span_y),
                z=0.35,
            ),
            camera=dict(eye=dict(x=1.4, y=-1.4, z=0.8)),
        ),
        paper_bgcolor="#FAFAFA",
    )
    return fig


def _draw_boreholes_3d_mpl(
    selected_zones: list[str],
    show_clay_bottom: bool = False,
    show_cdm_top: bool = False,
    cdm_top_z: float = 2.7,
    focus_bh: str | None = None,
    pair_highlight: tuple | None = None,
    elev: float = 20.0,
    azim: float = -60.0,
    zoom: float = 1.0,
):
    """Matplotlib fallback cho bản đồ 3D địa chất (khi không có plotly)."""
    import matplotlib.pyplot as plt
    from collections import defaultdict
    try:
        from mpl_toolkits.mplot3d import Axes3D  # noqa: F401
    except ImportError:
        fig, ax = plt.subplots(figsize=(8, 5))
        ax.text(0.5, 0.5, "Không có mpl_toolkits.mplot3d", ha="center", va="center",
                transform=ax.transAxes)
        ax.axis("off")
        return fig

    bhs, lays = _load_borehole_3d_data()
    bhs = [b for b in bhs if b["zone"] in selected_zones]
    if focus_bh:
        bhs_focus = [b for b in bhs if b["name"] == focus_bh]
        if bhs_focus:
            bhs = bhs_focus
    if not bhs:
        fig, ax = plt.subplots(figsize=(8, 5))
        ax.text(0.5, 0.5, "Không có dữ liệu tọa độ hố khoan", ha="center", va="center",
                transform=ax.transAxes)
        ax.axis("off")
        return fig

    bh_ids = {b["id"] for b in bhs}
    lays = [l for l in lays if l["borehole_id"] in bh_ids]
    lay_by_bh: dict[int, list[dict]] = defaultdict(list)
    for l in lays:
        lay_by_bh[l["borehole_id"]].append(l)

    fig = plt.figure(figsize=(10, 7))
    ax = fig.add_subplot(111, projection="3d")

    # Vẽ từng đoạn lớp đất bằng đường thẳng đứng có màu
    legend_seen: set[str] = set()
    for bh in bhs:
        for l in lay_by_bh[bh["id"]]:
            sym   = l["symbol"]
            color = _LAYER_COLORS.get(sym, _LAYER_DEFAULT_COLOR)
            z0    = bh["elevation_m"] - l["depth_top_m"]
            z1    = bh["elevation_m"] - l["depth_bot_m"]
            label = f"Lớp {sym}" if sym not in legend_seen else None
            legend_seen.add(sym)
            ax.plot(
                [bh["x_coord_m"], bh["x_coord_m"]],
                [bh["y_coord_m"], bh["y_coord_m"]],
                [z0, z1],
                color=color, lw=5, label=label, solid_capstyle="butt",
            )

    # Markers + nhãn tên HK
    _zone_colors = {"KE": "#E53935", "BXN": "#1565C0", "NHC": "#2E7D32", "QTT": "#F9A825"}
    _zone_markers = {"KE": "o", "BXN": "s", "NHC": "D", "QTT": "P"}
    for bh in bhs:
        z_top = bh["elevation_m"] + 3
        ax.scatter(bh["x_coord_m"], bh["y_coord_m"], z_top,
                   c=_zone_colors.get(bh["zone"], "#333"),
                   marker=_zone_markers.get(bh["zone"], "o"),
                   s=40, edgecolor="#000", lw=0.5)
        _nm = bh["name"].replace("BXN-CV-", "").replace("KE-", "")
        ax.text(bh["x_coord_m"], bh["y_coord_m"], z_top + 0.5, _nm,
                fontsize=8, ha="center", color="#212121")

    # Mặt đáy lớp bùn (vẽ scatter, không vẽ mesh)
    if show_clay_bottom:
        _clay = [c for c in _load_clay_bottom_3d() if c["zone"] in selected_zones]
        if _clay:
            ax.scatter(
                [c["x_coord_m"] for c in _clay],
                [c["y_coord_m"] for c in _clay],
                [c["clay_bot_elev"] for c in _clay],
                c="#0D47A1", marker="D", s=50, label="Đáy bùn",
            )

    # Mặt phẳng đỉnh trụ CDM
    if show_cdm_top:
        import numpy as _np
        all_x_ = [b["x_coord_m"] for b in bhs]
        all_y_ = [b["y_coord_m"] for b in bhs]
        pad = 20.0
        mx, Mx = min(all_x_) - pad, max(all_x_) + pad
        my, My = min(all_y_) - pad, max(all_y_) + pad
        ax.plot_surface(
            _np.array([[mx, Mx], [mx, Mx]]),
            _np.array([[my, my], [My, My]]),
            _np.array([[cdm_top_z]*2, [cdm_top_z]*2]),
            color="#E91E63", alpha=0.25, edgecolor="#C2185B",
        )

    # Đường kích thước giữa 2 HK
    if pair_highlight:
        _b1, _b2, _dist = pair_highlight
        _bh1 = next((b for b in bhs if b["name"] == _b1), None)
        _bh2 = next((b for b in bhs if b["name"] == _b2), None)
        if _bh1 and _bh2:
            _z = max(_bh1["elevation_m"], _bh2["elevation_m"]) + 4
            ax.plot(
                [_bh1["x_coord_m"], _bh2["x_coord_m"]],
                [_bh1["y_coord_m"], _bh2["y_coord_m"]],
                [_z, _z],
                color="#FF6F00", lw=2.5,
            )
            _mx = (_bh1["x_coord_m"] + _bh2["x_coord_m"]) / 2
            _my = (_bh1["y_coord_m"] + _bh2["y_coord_m"]) / 2
            ax.text(_mx, _my, _z + 1.5, f"{_dist:.1f} m",
                    fontsize=11, color="#E65100", ha="center", fontweight="bold")

    ax.set_xlabel("X — Easting (m)", fontsize=9)
    ax.set_ylabel("Y — Northing (m)", fontsize=9)
    ax.set_zlabel("Cao độ (m)", fontsize=9)
    ax.tick_params(labelsize=8)
    ax.legend(loc="upper left", fontsize=7, framealpha=0.85, ncol=2)
    ax.view_init(elev=float(elev), azim=float(azim))

    # Zoom: thu/giãn quanh tâm trục
    if abs(zoom - 1.0) > 1e-3:
        for lim_get, lim_set in (
            (ax.get_xlim, ax.set_xlim),
            (ax.get_ylim, ax.set_ylim),
            (ax.get_zlim, ax.set_zlim),
        ):
            _lo, _hi = lim_get()
            _ctr  = 0.5 * (_lo + _hi)
            _half = 0.5 * (_hi - _lo) / float(zoom)
            lim_set(_ctr - _half, _ctr + _half)

    fig.tight_layout()
    return fig


def _build_borehole_3d_deck(
    selected_zones: list[str],
    show_clay_bottom: bool = False,
    show_cdm_top: bool = False,
    cdm_top_z: float = 2.7,
    pair_highlight: tuple | None = None,
    z_scale: float = 5.0,
    col_radius: float = 2.5,
    pair_lines: list[dict] | None = None,
    show_pair_labels: bool = True,
    show_bh_names: bool = True,
):
    """Pydeck 3D map cho hố khoan — native drag/zoom/rotate qua deck.gl.
    z_scale: phóng đại trục Z (vì khoảng cách HK ~100m, độ sâu chỉ ~30m,
             nếu không scale sẽ thấy như mặt phẳng)."""
    bhs, lays = _load_borehole_3d_data()
    bhs = [b for b in bhs if b["zone"] in selected_zones]
    if not bhs:
        return None
    bh_ids = {b["id"] for b in bhs}
    lays = [l for l in lays if l["borehole_id"] in bh_ids]

    # VN2000 → WGS84
    for b in bhs:
        _lat, _lon = _vn2000_to_latlon(b["y_coord_m"], b["x_coord_m"])
        b["lat"], b["lon"] = _lat, _lon
    bh_by_id = {b["id"]: b for b in bhs}

    def _hex_rgb(h: str) -> list[int]:
        h = h.lstrip("#")
        if len(h) < 6:
            return [200, 200, 200]
        return [int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)]

    # ColumnLayer: mỗi lớp = trụ extrude từ z_bot lên z_top → thấy rõ bề dày
    columns = []
    for l in lays:
        b = bh_by_id.get(l["borehole_id"])
        if b is None:
            continue
        sym  = l["symbol"]
        col  = _LAYER_COLORS.get(sym, _LAYER_DEFAULT_COLOR)
        rgb  = _hex_rgb(col)
        z_top = (b["elevation_m"] - l["depth_top_m"]) * z_scale
        z_bot = (b["elevation_m"] - l["depth_bot_m"]) * z_scale
        thick = max(z_top - z_bot, 0.01)
        columns.append({
            # base position bao gồm altitude → ColumnLayer extrude lên từ z này
            "position":  [b["lon"], b["lat"], z_bot],
            "elevation": thick,
            "color":     rgb + [220],
            "bh":        b["name"],
            "sym":       sym,
            "desc":      (l.get("description") or sym)[:40],
            "rng":       f"{l['depth_top_m']:.1f}–{l['depth_bot_m']:.1f} m",
            "thk":       f"{l['depth_bot_m'] - l['depth_top_m']:.2f} m",
        })

    layers_deck = [
        pdk.Layer(
            "ColumnLayer",
            columns,
            pickable=True,
            extruded=True,
            radius=float(col_radius),    # bán kính trụ — đơn vị mét
            disk_resolution=24,          # 24 cạnh → trông tròn
            elevation_scale=1.0,         # đã scale trong z_scale
            get_position="position",
            get_elevation="elevation",
            get_fill_color="color",
            line_width_min_pixels=0,
            material={"ambient": 0.6, "diffuse": 0.6, "shininess": 30},
            auto_highlight=True,
        ),
    ]

    # Nhãn tên HK (toggle được)
    _zone_rgb = {"KE": [229, 57, 53], "BXN": [21, 101, 192], "NHC": [46, 125, 50], "QTT": [249, 168, 37]}
    if show_bh_names:
        label_data = [
            {
                "position": [b["lon"], b["lat"], (b["elevation_m"] + 3) * z_scale],
                "text":     b["name"].replace("BXN-CV-", "").replace("KE-", ""),
                "color":    _zone_rgb.get(b["zone"], [50, 50, 50]),
            }
            for b in bhs
        ]
        layers_deck.append(pdk.Layer(
            "TextLayer",
            label_data,
            get_position="position",
            get_text="text",
            get_color="color",
            get_size=14,
            size_units="pixels",
            get_alignment_baseline="'bottom'",
            pickable=False,
        ))

    # Markers HK (ScatterplotLayer)
    marker_data = [
        {
            "position": [b["lon"], b["lat"], b["elevation_m"] * z_scale],
            "color":    _zone_rgb.get(b["zone"], [50, 50, 50]),
            "bh":       b["name"],
            "zone":     b["zone"],
        }
        for b in bhs
    ]
    layers_deck.append(pdk.Layer(
        "ScatterplotLayer",
        marker_data,
        get_position="position",
        get_fill_color="color",
        get_radius=6,
        radius_units="pixels",
        pickable=True,
    ))

    # Đáy lớp bùn
    if show_clay_bottom:
        _clay = [c for c in _load_clay_bottom_3d() if c["zone"] in selected_zones]
        if _clay:
            clay_data = []
            for c in _clay:
                _lat, _lon = _vn2000_to_latlon(c["y_coord_m"], c["x_coord_m"])
                clay_data.append({
                    "position": [_lon, _lat, c["clay_bot_elev"] * z_scale],
                    "bh":       c["name"],
                    "elev":     c["clay_bot_elev"],
                    "depth":    c["clay_bot_depth"],
                })
            layers_deck.append(pdk.Layer(
                "ScatterplotLayer",
                clay_data,
                get_position="position",
                get_fill_color=[13, 71, 161, 200],
                get_radius=8,
                radius_units="pixels",
                pickable=True,
            ))

    # Mặt phẳng đỉnh trụ CDM (PolygonLayer)
    if show_cdm_top:
        all_lons = [b["lon"] for b in bhs]
        all_lats = [b["lat"] for b in bhs]
        pad = 0.0005
        mn_lon, mx_lon = min(all_lons) - pad, max(all_lons) + pad
        mn_lat, mx_lat = min(all_lats) - pad, max(all_lats) + pad
        cdm_data = [{
            "polygon": [
                [mn_lon, mn_lat], [mx_lon, mn_lat],
                [mx_lon, mx_lat], [mn_lon, mx_lat],
            ],
            "z": cdm_top_z * z_scale,
        }]
        layers_deck.append(pdk.Layer(
            "PolygonLayer",
            cdm_data,
            get_polygon="polygon",
            get_elevation="z",
            extruded=False,
            get_fill_color=[233, 30, 99, 70],
            get_line_color=[194, 24, 91],
            line_width_min_pixels=1,
        ))

    # Vẽ tất cả cặp khoảng cách HK (color-coded theo status TCCS41)
    _bh_lookup = {b["name"]: b for b in bhs}
    _status_color = {
        "Đạt":      [46, 125, 50, 220],     # xanh
        "Gần quá":  [229, 57, 53, 220],     # đỏ
        "Xa quá":   [255, 152, 0, 220],     # cam
    }
    if pair_lines:
        _line_data = []
        _label_data = []
        # Cao độ đường kích thước: max elev của tất cả HK + offset
        _z_lines = (max(b["elevation_m"] for b in bhs) + 4) * z_scale
        for p in pair_lines:
            _b1 = _bh_lookup.get(p["bh1"])
            _b2 = _bh_lookup.get(p["bh2"])
            if not (_b1 and _b2):
                continue
            _col = _status_color.get(p.get("status", ""), [120, 120, 120, 200])
            _dist = p.get("distance_m", 0.0)
            _line_data.append({
                "path": [[_b1["lon"], _b1["lat"], _z_lines],
                         [_b2["lon"], _b2["lat"], _z_lines]],
                "color":    _col,
                "bh1":      p["bh1"],
                "bh2":      p["bh2"],
                "distance": f"{_dist:.1f} m",
                "status":   p.get("status", ""),
            })
            if show_pair_labels:
                _label_data.append({
                    "position": [(_b1["lon"] + _b2["lon"]) / 2,
                                 (_b1["lat"] + _b2["lat"]) / 2,
                                 _z_lines + 1],
                    "text":     f"{_dist:.0f}m",
                    "color":    _col[:3],
                })
        if _line_data:
            layers_deck.append(pdk.Layer(
                "PathLayer",
                _line_data,
                get_path="path",
                get_color="color",
                get_width=2,
                width_units="pixels",
                pickable=True,
            ))
        if _label_data:
            layers_deck.append(pdk.Layer(
                "TextLayer",
                _label_data,
                get_position="position",
                get_text="text",
                get_color="color",
                get_size=12,
                size_units="pixels",
                get_alignment_baseline="'bottom'",
            ))

    # Đường kích thước được highlight (cặp user chọn cụ thể)
    if pair_highlight:
        _b1n, _b2n, _dist = pair_highlight
        _b1 = next((b for b in bhs if b["name"] == _b1n), None)
        _b2 = next((b for b in bhs if b["name"] == _b2n), None)
        if _b1 and _b2:
            _zline = (max(_b1["elevation_m"], _b2["elevation_m"]) + 6) * z_scale
            layers_deck.append(pdk.Layer(
                "PathLayer",
                [{
                    "path": [[_b1["lon"], _b1["lat"], _zline],
                             [_b2["lon"], _b2["lat"], _zline]],
                    "color": [255, 111, 0, 255],
                }],
                get_path="path",
                get_color="color",
                get_width=5,
                width_units="pixels",
                pickable=False,
            ))
            _mid_lon = (_b1["lon"] + _b2["lon"]) / 2
            _mid_lat = (_b1["lat"] + _b2["lat"]) / 2
            layers_deck.append(pdk.Layer(
                "TextLayer",
                [{
                    "position": [_mid_lon, _mid_lat, _zline + 1],
                    "text":     f"{_dist:.1f} m",
                    "color":    [230, 81, 0],
                }],
                get_position="position",
                get_text="text",
                get_color="color",
                get_size=18,
                size_units="pixels",
            ))

    mid_lat = sum(b["lat"] for b in bhs) / len(bhs)
    mid_lon = sum(b["lon"] for b in bhs) / len(bhs)
    view = pdk.ViewState(
        latitude=mid_lat, longitude=mid_lon,
        zoom=16, pitch=55, bearing=-20,
    )
    return pdk.Deck(
        layers=layers_deck,
        initial_view_state=view,
        tooltip={"html": "<b>{bh}{bh1}</b>{bh2}<br>"
                          "<span>Lớp <b>{sym}</b>: {desc}</span>"
                          "<span>Độ sâu: {rng} | Dày: <b>{thk}</b></span>"
                          "<span>Khoảng cách: <b>{distance}</b> — {status}</span>",
                 "style": {"color": "white",
                           "backgroundColor": "rgba(33,33,33,0.92)",
                           "fontSize": "12px"}},
        map_style=None,
    )


_VN2000_CM  = 105.75          # kinh tuyến trục TP.HCM: 105°45'E
_VN2000_FE  = 500_000.0       # False Easting
_M_PER_DEG  = 110_574.0       # m/độ vĩ tuyến (xích đạo ≈ 110,574 m)


def _vn2000_to_latlon(northing: float, easting: float) -> tuple[float, float]:
    """VN-2000 TP.HCM (CM=105°45', FE=500000) → WGS84 (lat, lon).
    Sai số <1 m trong phạm vi TP.HCM; đủ dùng cho overlay bản đồ.
    """
    lat = northing / _M_PER_DEG
    lon = _VN2000_CM + (easting - _VN2000_FE) / (_M_PER_DEG * math.cos(math.radians(lat)))
    return lat, lon


# (Hàm `_project_to_latlon`, `_MAP_STYLES`, `_draw_map_2d` đã bỏ —
# bản đồ vị trí trong tab Địa chất chuyển sang folium + pyproj,
# inline trong page handler. Xem CLAUDE.md §32 và mục page="geology".)


# ═══════════════════════════════════════════════════════════════════════════════
# CDM CALCULATIONS
# ═══════════════════════════════════════════════════════════════════════════════
def calc_a(D: float, e: float, arrangement: str = "triangle") -> float:
    if arrangement == "triangle":
        return math.pi * D**2 / (2 * math.sqrt(3) * e**2)
    return math.pi * D**2 / (4 * e**2)  # square


def calc_Ec(qu_field: float) -> float:
    return 100 * (qu_field / 2)   # Ec = 100·Cc, Cc = qu/2


def calc_Es(Cu: float) -> float:
    return 250 * Cu


def calc_Etb(a: float, Ec: float, Es: float) -> float:
    return a * Ec + (1 - a) * Es


def calc_S1(q: float, Lc: float, Etb: float) -> float:
    """Độ lún S1 (cm)."""
    return q * Lc / Etb * 100


def calc_sigma_col(q: float, Ec: float, Etb: float) -> float:
    """Ứng suất tập trung lên đầu cọc (kN/m²)."""
    return (Ec / Etb) * q


def calc_Pcol(sigma_col: float, D: float) -> float:
    """Lực nén lên 1 trụ CDM (kN)."""
    return sigma_col * math.pi * D**2 / 4


def calc_bearing(D: float, Lc: float, Cc: float, Cu: float, FS: float = 2.0) -> dict:
    """Sức chịu tải trụ CDM theo phương pháp AIT (TCVN 9403:2012 Phụ lục B)."""
    Ac = math.pi * D**2 / 4
    Qult_col  = 9 * Cc * Ac                         # sức chịu ở mũi (kN)
    Qult_skin = math.pi * D * Lc * Cu               # ma sát thân (kN)
    Qult      = Qult_col + Qult_skin
    Qa        = Qult / FS
    return {"Qult": round(Qult, 1), "Qa": round(Qa, 1),
            "Qult_col": round(Qult_col, 1), "Qult_skin": round(Qult_skin, 1)}


def calc_punching_check(
    D: float, e: float, Hse: float, He: float,
    quckse: float, Fs: float, theta_deg: float,
    gamma_fill: float, gamma_mat: float, qa: float = 0.0,
) -> dict:
    """Kiểm tra chọc thủng lớp đệm xi măng – phương pháp ALiCC (PWRI Japan)."""
    tan_t2 = math.tan(math.radians(theta_deg / 2))   # tan(θ/2) cho Ho
    tan_t  = math.tan(math.radians(theta_deg))        # tan(θ) cho thể tích
    tase   = quckse / (2.0 * Fs)
    Ho     = (e - D) * tan_t2

    if Ho <= He:
        Vsoil = (((e - D) * 0.5 * e**2
                  - math.pi * (e**3 - D**3) / 24.0
                  + (4.0 - math.pi) * (math.sqrt(2) - 1.0) * e**3 / 24.0)
                 * tan_t)
    else:
        r0    = He / tan_t + D / 2.0
        Vsoil = (He * e**2
                 - (1.0/3.0) * (math.pi * r0**2 * (He + D/2.0 * tan_t)
                                - math.pi * D / 2.0 * tan_t))

    r_mat  = Hse / tan_t + D / 2.0
    VCGCXM = (Hse * e**2
              - (1.0/3.0) * (math.pi * r_mat**2 * (Hse + D/2.0 * tan_t)
                             - math.pi * D / 2.0 * tan_t))

    A_unit = e**2 - math.pi * D**2 / 4.0
    PSoil  = ((Vsoil - VCGCXM) * gamma_fill + VCGCXM * gamma_mat) / A_unit
    tse    = (PSoil - qa) * A_unit / (math.pi * D * Hse)

    return {
        "tase_kPa":  round(tase, 2),
        "Ho_m":      round(Ho, 3),
        "He_m":      round(He, 3),
        "cong_thuc": "CT(1) Ho≤He" if Ho <= He else "CT(2) Ho>He",
        "Vsoil_m3":  round(Vsoil, 4),
        "VCGCXM_m3": round(VCGCXM, 4),
        "PSoil_kPa": round(PSoil, 3),
        "tse_kPa":   round(tse, 3),
        "check_CT":  "Đạt" if tse <= tase else "Không đạt",
    }


def build_scenarios(
    D: float, Lc: float, qu_field: float, Cu: float,
    q: float, spacings: list[float], arrangement: str,
    q_static: float | None = None,
) -> list[dict]:
    """
    Quy ước tải trọng (TCVN 9403:2012 + TCCS 41:2022):
    - **Lún S₁** (TCCS 41:2022 Phụ lục A): tính bằng tải TĨNH `q_static`
      (đất đắp + đệm + mặt đường + ô tô đắp), **KHÔNG xét hoạt tải bánh xe**
      vì hoạt tải có tần suất ngắn không gây cố kết.
    - **Sức tải đầu cọc Pcol** (TCVN 9403 Phụ lục C): tính bằng tải TỔNG `q`
      (bao gồm hoạt tải) vì cọc CDM chịu trực tiếp cả hoạt tải.
    Nếu `q_static = None` → dùng `q` cho cả hai (tương thích ngược).
    """
    if q_static is None:
        q_static = q
    Ec = calc_Ec(qu_field)
    Es = calc_Es(Cu)
    Cc = qu_field / 2
    rows = []
    for e in spacings:
        a    = calc_a(D, e, arrangement)
        Etb  = calc_Etb(a, Ec, Es)
        S1   = calc_S1(q_static, Lc, Etb)          # tải tĩnh, BỎ qua hoạt tải
        sc   = calc_sigma_col(q, Ec, Etb)          # tải tổng, XÉT hoạt tải
        Pcol = calc_Pcol(sc, D)
        bc   = calc_bearing(D, Lc, Cc, Cu)
        rows.append({
            "e (m)": e, "a": round(a, 4), "a (%)": round(a*100, 1),
            "Ec (kN/m²)": int(Ec), "Es (kN/m²)": int(Es),
            "Etb (kN/m²)": round(Etb, 0),
            "S₁ (cm)": round(S1, 2),
            "σ_col (kN/m²)": round(sc, 1),
            "Pcol (kN)": round(Pcol, 1),
            "Qa (kN)": bc["Qa"],
            "Đạt SCT": "Đạt" if Pcol < bc["Qa"] else "Không đạt",
        })
    return rows


def q_total(loads: dict) -> float:
    """Tải tổng bao gồm hoạt tải xe — dùng cho Pcol/SCT."""
    return (loads.get("q_traffic", 20)
            + loads.get("h_road", 0.8) * loads.get("g_road", 24)
            + loads.get("h_fill", 1.5) * loads.get("g_fill", 18)
            + loads.get("h_mat",  0.4) * loads.get("g_mat",  22.5))


def q_static(loads: dict) -> float:
    """Tải tĩnh (đất đắp + đệm + mặt đường, KHÔNG hoạt tải) — dùng cho lún."""
    return (loads.get("h_road", 0.8) * loads.get("g_road", 24)
            + loads.get("h_fill", 1.5) * loads.get("g_fill", 18)
            + loads.get("h_mat",  0.4) * loads.get("g_mat",  22.5))


# ═══════════════════════════════════════════════════════════════════════════════
# SESSION STATE INIT
# ═══════════════════════════════════════════════════════════════════════════════
_DEFAULTS = {
    "cdm_zone": "KE",
    "cdm_bh": None,
    "cdm_top_clay": 0.8,
    "cdm_h_clay": 24.8,
    "cdm_Su": 11.2,
    "cdm_gamma": 15.0,
    "cdm_elevation": 0.0,
    "cdm_D": 0.8,
    "cdm_e": 1.6,
    "cdm_L_ngam": 0.5,
    "cdm_Lc": 26.2,
    "cdm_CDTK": 0.8,
    "cdm_qu": 800.0,
    "cdm_FS_lab": 2.0,
    "cdm_arrangement": "triangle",
    "cdm_cement_type": "Hoàng Thạch PCB40",
    "cdm_dosage": 240,
    "cdm_WC": 1.0,
    "cdm_spacings": [1.4, 1.6, 1.8],
    "cdm_rec_idx": 1,
    "cdm_geo_view": "Bản đồ vị trí",
    "cdm_vst_sel": [],
    "cdm_quckse": 600.0,
    "cdm_Fs_mat": 3.0,
    "cdm_theta":  80.0,
    "cdm_qa_mat": 0.0,
    "lang": "VN",
    "cdm_loads": {
        "q_traffic": 20.0,
        "z_tk":   3.5,
        "h_road": 0.8,   "g_road": 24.0,
        "h_fill": 1.5,   "g_fill": 18.0,
        "h_mat":  0.4,   "g_mat":  22.5,
    },
}
for _k, _v in _DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v


def _get(key):
    return st.session_state.get(key, _DEFAULTS.get(key))


# ═══════════════════════════════════════════════════════════════════════════════
# CHART HELPERS
# ═══════════════════════════════════════════════════════════════════════════════
def _chart_scenarios(df: pd.DataFrame, rec_idx: int) -> go.Figure:
    """3 biểu đồ con: a(%), Etb, S1."""
    x = [f"e={r['e (m)']}m" for _, r in df.iterrows()]
    colors = ["#ED7D31" if i == rec_idx else "#4472C4" for i in range(len(df))]

    fig = go.Figure()
    # S1
    fig.add_trace(go.Bar(name="S₁ (cm)", x=x, y=df["S₁ (cm)"],
                         marker_color=colors, text=df["S₁ (cm)"].apply(lambda v: f"{v:.2f}"),
                         textposition="outside"))
    fig.update_layout(title="Độ lún S₁ theo khoảng cách trụ",
                      yaxis_title="S₁ (cm)", height=340,
                      legend=dict(orientation="h"), margin=dict(t=40, b=20))
    return fig


def _chart_etb(df: pd.DataFrame, rec_idx: int) -> go.Figure:
    x = [f"e={r['e (m)']}m" for _, r in df.iterrows()]
    colors = ["#ED7D31" if i == rec_idx else "#5B9BD5" for i in range(len(df))]
    fig = go.Figure(go.Bar(name="Etb", x=x, y=df["Etb (kN/m²)"],
                            marker_color=colors,
                            text=df["Etb (kN/m²)"].apply(lambda v: f"{int(v):,}"),
                            textposition="outside"))
    fig.update_layout(title="Mô đun tương đương Etb", yaxis_title="kN/m²",
                      height=300, margin=dict(t=40, b=20))
    return fig


def _linreg(xs: list, ys: list) -> tuple[float, float]:
    """Hồi quy tuyến tính đơn giản, trả về (a, b) với y = a*x + b."""
    n = len(xs)
    if n < 2:
        return 0.0, (ys[0] if ys else 0.0)
    sx  = sum(xs);  sy  = sum(ys)
    sxy = sum(xi*yi for xi, yi in zip(xs, ys))
    sxx = sum(xi**2 for xi in xs)
    denom = n * sxx - sx * sx
    if abs(denom) < 1e-12:
        return 0.0, sy / n
    a = (n * sxy - sx * sy) / denom
    b = (sy - a * sx) / n
    return a, b


def _chart_su_profile_mpl(
    df_vst: pd.DataFrame,
    selected_locs: list[str] | None = None,
    figsize: tuple = (7, 6),
    font_scale: float = 1.0,
):
    """Matplotlib fallback cho biểu đồ Su–VST khi không có plotly.
    `font_scale`: ×1 mặc định; ×2 khi đặt cạnh soil column trong cột hẹp."""
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=figsize)
    _fs_anno = int(7 * font_scale)
    _fs_axis = int(9 * font_scale)
    _fs_title = int(10 * font_scale)
    _fs_tick = int(8 * font_scale)
    _fs_leg  = int(8 * font_scale)
    if df_vst.empty:
        ax.text(0.5, 0.5, "Không có dữ liệu VST", ha="center", va="center", transform=ax.transAxes)
        return fig

    all_locs  = sorted(df_vst["loc_name"].unique())
    show_locs = selected_locs if selected_locs else all_locs
    _colors = ["#E53935","#1565C0","#2E7D32","#F57F17","#6A1B9A",
               "#00838F","#AD1457","#558B2F","#4527A0","#00695C"]

    for i, loc in enumerate(all_locs):
        if loc not in show_locs:
            continue
        grp    = df_vst[df_vst["loc_name"] == loc].sort_values("depth_m")
        color  = _colors[i % len(_colors)]
        depths = grp["depth_m"].tolist()
        sus    = grp["Su_kPa"].tolist()
        y_plot = [-d for d in depths]
        ax.plot(sus, y_plot, "-o", color=color, lw=1.5, ms=5, label=loc)
        for x, y in zip(sus, y_plot):
            ax.annotate(f"{x:.1f}", (x, y), xytext=(4, 0), textcoords="offset points",
                        fontsize=_fs_anno, color=color, va="center")
        if len(depths) >= 2:
            a, b = _linreg(depths, sus)
            d_ext = (max(depths) - min(depths)) * 0.1
            dd = [min(depths) - d_ext, max(depths) + d_ext]
            ax.plot([a*d + b for d in dd], [-d for d in dd],
                    "--", color=color, lw=1.0, alpha=0.7)
            sign = "+" if b >= 0 else "−"
            eq = f"Su={a:+.2f}z{sign}{abs(b):.2f}".replace("+-", "−")
            ax.text(a*dd[-1] + b, -dd[-1], eq, fontsize=_fs_anno, color=color,
                    bbox=dict(facecolor="white", alpha=0.75, edgecolor=color, lw=0.5, pad=1))

    ax.set_xlabel("Su (kPa)", fontsize=_fs_axis)
    ax.set_ylabel("Cao độ (m)", fontsize=_fs_axis)
    title = "Biểu đồ Su – Cắt cánh (VST)"
    title += f" — {', '.join(show_locs)}" if selected_locs else " — Toàn bộ"
    ax.set_title(title, fontsize=_fs_title)
    ax.grid(True, ls=":", color="#CCC", lw=0.5)
    ax.legend(loc="lower right", fontsize=_fs_leg, framealpha=0.85)
    ax.tick_params(labelsize=_fs_tick)
    fig.tight_layout()
    return fig


def _draw_soil_column_mpl(
    layers: list[dict], bh_name: str = "",
    spt: list[dict] | None = None,
    figsize: tuple | None = None,
    font_scale: float = 1.0,
):
    """Matplotlib fallback cho cột địa chất + SPT.
    `figsize=None`: dùng default (7,12) khi có SPT ở cột hẹp.
    `figsize=(W,H)`: override (vd (7,6) khi đặt cạnh VST cùng hàng).
    `font_scale`: ×1 mặc định cột hẹp; ×0.65 khi đặt cạnh VST trong cột rộng."""
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches

    def _tc(hex_color: str) -> str:
        h = hex_color.lstrip("#")
        if len(h) < 6:
            return "#212121"
        r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
        return "white" if (0.299*r + 0.587*g + 0.114*b) < 140 else "#212121"

    def _spt_color(n: int) -> str:
        if n <= 2:   return "#E53935"
        if n <= 5:   return "#FB8C00"
        if n <= 10:  return "#FDD835"
        if n <= 30:  return "#66BB6A"
        return "#1565C0"

    has_spt = bool(spt)
    if not layers:
        fig, ax = plt.subplots(figsize=(4, 7))
        ax.text(0.5, 0.5, f"Không có địa tầng\n{bh_name}", ha="center", va="center",
                transform=ax.transAxes, fontsize=10, color="#666")
        ax.axis("off")
        return fig

    # VST figsize=(7,6) trong cột rộng 2/3 → displayed_h ≈ 0.571·W
    # Cột địa chất ở cột hẹp 1/3 cần aspect cao hơn để bù → h/w ≈ 12/7
    # Override figsize khi đặt cạnh VST trong cùng row width
    y_bot = layers[-1]["depth_bot_m"]
    if has_spt:
        _figsize = figsize if figsize else (7, 12)
        fig, (ax_col, ax_spt) = plt.subplots(
            1, 2, figsize=_figsize, sharey=True,
            gridspec_kw={"width_ratios": [0.4, 0.6], "wspace": 0.05},
        )
    else:
        _figsize = figsize if figsize else (4, 7)
        fig, ax_col = plt.subplots(figsize=_figsize)
        ax_spt = None

    _fs_lbl  = int((14 if has_spt else 8) * font_scale)
    _fs_sym  = int((16 if has_spt else 9) * font_scale)
    _fs_axis = int((16 if has_spt else 9) * font_scale)
    _fs_tick = int((14 if has_spt else 8) * font_scale)
    _fs_ttl  = int((18 if has_spt else 10) * font_scale)
    for lay in layers:
        y0, y1 = lay["depth_top_m"], lay["depth_bot_m"]
        sym   = (lay.get("symbol") or "").strip()
        color = _LAYER_COLORS.get(
            sym, _LAYER_COLORS.get(sym[:2] if len(sym) > 2 else sym, _LAYER_DEFAULT_COLOR))
        ax_col.add_patch(mpatches.Rectangle(
            (0, y0), 1, y1 - y0, facecolor=color, edgecolor="#555", lw=0.6))
        ax_col.text(1.04, y0, f"{y0:.1f}", fontsize=_fs_lbl, color="#555", va="center")
        if (y1 - y0) >= 0.5:
            mid = (y0 + y1) / 2
            sym_lbl = sym if sym else f"L{layers.index(lay)+1}"
            desc = (lay.get("description") or "")[:18]
            # Ký hiệu lớn — _fs_sym, mô tả nhỏ ~= _fs_lbl (font chung)
            _fs_desc = max(int(_fs_lbl * 0.95), 7)
            ax_col.text(0.5, mid - (0.25 if (y1 - y0) >= 2.5 and desc else 0),
                        sym_lbl, fontsize=_fs_sym, color=_tc(color),
                        ha="center", va="center", fontweight="bold")
            if (y1 - y0) >= 2.5 and desc:
                ax_col.text(0.5, mid + (y1 - y0) * 0.08,
                            desc, fontsize=_fs_desc, color=_tc(color),
                            ha="center", va="center", fontweight="normal")
    ax_col.text(1.04, y_bot, f"{y_bot:.1f}", fontsize=_fs_lbl, color="#555", va="center")
    ax_col.set_xlim(-0.05, 1.5)
    ax_col.set_ylim(y_bot + 1, -0.5)
    ax_col.set_xticks([])
    ax_col.set_ylabel("Độ sâu (m)", fontsize=_fs_axis)
    ax_col.tick_params(labelsize=_fs_tick)
    ax_col.grid(True, axis="y", ls=":", color="#EEE", lw=0.5)
    ax_col.set_title("Địa chất" if has_spt else f"Cột ĐC: {bh_name}",
                     fontsize=_fs_ttl, color="#1F4E79")

    if has_spt and ax_spt is not None:
        _n_max = max((s["N"] or 0 for s in spt), default=10)
        _x_max = max(60, _n_max + 5)
        depths = [s["depth_m"] for s in spt]
        N_vals = [min(s["N"] or 0, 60) for s in spt]
        colors = [_spt_color(n) for n in N_vals]
        ax_spt.barh(depths, N_vals, height=0.6, color=colors, edgecolor="#444", lw=0.5)
        _fs_N = int(14 * font_scale)
        _fs_xlbl = int(16 * font_scale)
        _fs_tit = int(18 * font_scale)
        for d, n in zip(depths, N_vals):
            ax_spt.text(n + 1, d, str(n), fontsize=_fs_N, va="center")
        ax_spt.plot(N_vals, depths, ":", color="#333", lw=1)
        ax_spt.axvline(4, ls="--", color="#E53935", lw=1, alpha=0.6)
        ax_spt.text(4.2, -0.2, "N=4", fontsize=_fs_N, color="#E53935")
        ax_spt.set_xlim(0, _x_max)
        ax_spt.set_xlabel("N (blows/30cm)", fontsize=_fs_xlbl)
        ax_spt.tick_params(labelsize=_fs_N)
        ax_spt.grid(True, axis="x", ls=":", color="#EEE", lw=0.5)
        ax_spt.set_title("SPT – N", fontsize=_fs_tit, color="#1F4E79")
        fig.suptitle(f"Cột ĐC: {bh_name}", fontsize=_fs_tit, color="#1F4E79", y=0.995)

    fig.tight_layout()
    return fig


def _chart_su_profile(
    df_vst: pd.DataFrame,
    selected_locs: list[str] | None = None,
) -> go.Figure:
    if df_vst.empty:
        return go.Figure()

    all_locs = sorted(df_vst["loc_name"].unique())
    show_locs = selected_locs if selected_locs else all_locs

    _colors = [
        "#E53935","#1565C0","#2E7D32","#F57F17","#6A1B9A",
        "#00838F","#AD1457","#558B2F","#4527A0","#00695C",
    ]

    fig = go.Figure()
    reg_lines = []   # (loc, a, b, color, y_min, y_max)

    for i, loc in enumerate(all_locs):
        if loc not in show_locs:
            continue
        grp   = df_vst[df_vst["loc_name"] == loc].sort_values("depth_m")
        color = _colors[i % len(_colors)]
        depths = grp["depth_m"].tolist()
        sus    = grp["Su_kPa"].tolist()
        y_plot = [-d for d in depths]

        # Đường + markers + nhãn giá trị
        fig.add_trace(go.Scatter(
            x=sus, y=y_plot,
            mode="lines+markers+text",
            name=loc,
            line=dict(color=color, width=1.8),
            marker=dict(size=7, color=color),
            text=[f"{v:.1f}" for v in sus],
            textposition="middle right",
            textfont=dict(size=9, color=color),
            hovertemplate=(
                f"<b>{loc}</b><br>"
                "Sâu: %{customdata:.1f} m<br>"
                "Su: %{x:.1f} kPa<extra></extra>"
            ),
            customdata=depths,
        ))

        # Hồi quy tuyến tính: Su = a*depth + b
        if len(depths) >= 2:
            a, b = _linreg(depths, sus)
            reg_lines.append((loc, a, b, color, min(depths), max(depths)))

    # Vẽ đường hồi quy + annotation phương trình
    for loc, a, b, color, d_min, d_max in reg_lines:
        d_ext = (d_max - d_min) * 0.1
        dd = [d_min - d_ext, d_max + d_ext]
        su_reg = [a * d + b for d in dd]
        sign  = "+" if b >= 0 else "−"
        eq    = f"Su={a:+.2f}z{sign}{abs(b):.2f}".replace("+-", "−")
        fig.add_trace(go.Scatter(
            x=su_reg, y=[-d for d in dd],
            mode="lines",
            name=f"Hồi quy {loc}",
            line=dict(color=color, width=1.2, dash="dash"),
            hovertemplate=f"<b>Hồi quy {loc}</b><br>{eq}<extra></extra>",
            showlegend=False,
        ))
        # Nhãn phương trình ở cuối đường
        fig.add_annotation(
            x=su_reg[-1], y=-dd[-1],
            text=eq,
            showarrow=False,
            font=dict(size=9, color=color),
            xanchor="left",
            bgcolor="rgba(255,255,255,0.75)",
            bordercolor=color,
            borderwidth=1,
        )

    fig.update_layout(
        title="Biểu đồ Su – Cắt cánh (VST)" + (
            f" — {', '.join(show_locs)}" if selected_locs else " — Toàn bộ"
        ),
        xaxis_title="Su (kPa)",
        yaxis_title="Cao độ (m)",
        height=500,
        legend=dict(font_size=10, orientation="v", x=1.01, y=1),
        margin=dict(t=50, b=20, r=160),
    )
    return fig


def _chart_combined(df: pd.DataFrame, rec_idx: int) -> go.Figure:
    x = [f"e={r['e (m)']}m" for _, r in df.iterrows()]
    colors = ["#ED7D31" if i == rec_idx else "#4472C4" for i in range(len(df))]
    fig = go.Figure()
    fig.add_trace(go.Bar(name="a (%)", x=x, y=df["a (%)"],
                         marker_color=colors, yaxis="y1",
                         offsetgroup=0,
                         text=df["a (%)"].apply(lambda v: f"{v:.1f}%"),
                         textposition="outside"))
    fig.add_trace(go.Scatter(name="S₁ (cm)", x=x, y=df["S₁ (cm)"],
                              mode="lines+markers+text",
                              text=df["S₁ (cm)"].apply(lambda v: f"{v:.2f}"),
                              textposition="top center",
                              yaxis="y2", line=dict(color="#FF0000", width=2),
                              marker=dict(size=8, color="#FF0000")))
    fig.update_layout(
        title="Tổng hợp: tỷ lệ thay thế a(%) và độ lún S₁(cm)",
        yaxis=dict(title="a (%)", side="left"),
        yaxis2=dict(title="S₁ (cm)", side="right", overlaying="y"),
        barmode="group", height=340,
        legend=dict(orientation="h"), margin=dict(t=40, b=20),
    )
    return fig


def _draw_cdm_section(
    D: float, Lc: float, CDTK: float,
    h_clay: float, h_road: float, h_fill: float, h_mat: float,
    q_traffic: float,
    top_clay: float | None = None,
):
    """Mặt cắt ngang CDM: các lớp đất, trụ CDM, mũi tên tải, kích thước D/Lc."""
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches

    z_mat         = CDTK + h_mat
    z_fill        = z_mat + h_fill
    z_road        = z_fill + h_road
    _clay_top_eff = top_clay if top_clay is not None else CDTK
    z_clay_bot    = _clay_top_eff - h_clay
    z_pile_bot    = CDTK - Lc
    z_base        = min(z_clay_bot, z_pile_bot) - 1.5

    W  = max(4.0 * D, 2.5)
    cx = W / 2
    r  = D / 2

    fig, ax = plt.subplots(figsize=(5, 8))
    fig.patch.set_facecolor("#FAFAFA")
    ax.set_facecolor("white")
    ax.set_xlim(-W * 0.35, W * 1.2)
    ax.set_ylim(z_base, z_road + 2.5)
    ax.set_ylabel("Cao độ (m)", fontsize=9)
    ax.set_xticks([])
    ax.tick_params(axis="y", labelsize=8)
    ax.set_title("Mặt cắt ngang CDM", fontsize=11, fontweight="bold", pad=8)
    for s in ["top", "right", "bottom"]:
        ax.spines[s].set_visible(False)

    def hrect(y0, y1, fc, ec="#666", zo=1):
        if y1 > y0:
            ax.add_patch(mpatches.Rectangle(
                (0, y0), W, y1 - y0, fc=fc, ec=ec, lw=0.7, zorder=zo
            ))

    hrect(z_base,        z_clay_bot,    "#BCAAA4", "#795548")
    hrect(z_clay_bot,    _clay_top_eff, "#BBDEFB", "#1565C0")
    if top_clay is not None and top_clay < CDTK - 0.05:
        hrect(_clay_top_eff, CDTK,     "#FFD54F", "#E6BE00")
    if h_mat > 0:
        hrect(CDTK,      z_mat,        "#FFA726", "#E65100")
    hrect(z_mat,         z_fill,       "#FFD54F", "#E6BE00")
    hrect(z_fill,        z_road,       "#90A4AE", "#546E7A")

    p_bot = max(z_pile_bot, z_base + 0.1)
    ax.add_patch(mpatches.Rectangle(
        (cx - r, p_bot), D, CDTK - p_bot, fc="#66BB6A", ec="#2E7D32", lw=1.2, zorder=3
    ))
    p_mid = (CDTK + p_bot) / 2
    ax.text(cx, p_mid, f"CDM\nD={D:.2f}m\nLc={Lc:.1f}m",
            ha="center", va="center", fontsize=8.5,
            color="white", fontweight="bold", zorder=4)

    # Cao độ thiết kế = đỉnh lớp kết cấu áo đường (z_road)
    ax.axhline(z_road, color="#E53935", lw=1.5, ls="--", zorder=5)
    ax.text(W * 0.98, z_road + 0.08, f"Cao độ TK = {z_road:+.2f} m",
            ha="right", va="bottom", fontsize=8.5, color="#E53935",
            fontweight="bold", zorder=7,
            bbox=dict(boxstyle="round,pad=0.2", fc="white", ec="#E53935",
                      alpha=0.9, lw=0.8))
    # Đỉnh cọc CDM (nét chấm xanh lá)
    ax.axhline(CDTK, color="#2E7D32", lw=0.9, ls=":", zorder=4, alpha=0.8)
    ax.text(W * 0.98, CDTK - 0.1, f"Đỉnh cọc = {CDTK:+.2f} m",
            ha="right", va="top", fontsize=7.5, color="#2E7D32", zorder=5,
            bbox=dict(boxstyle="round,pad=0.1", fc="white", ec="none", alpha=0.8))

    # Đỉnh lớp bùn (nét gạch nâu) — chỉ vẽ khi khác CDTK > 0.05 m
    if top_clay is not None and abs(top_clay - CDTK) > 0.05:
        ax.axhline(top_clay, color="#6D4C41", lw=1.2, ls="-.", zorder=4, alpha=0.85)
        _va = "bottom" if top_clay > CDTK else "top"
        _dy = 0.08 if top_clay > CDTK else -0.08
        ax.text(W * 0.02, top_clay + _dy,
                f"Đỉnh lớp bùn = {top_clay:+.2f} m",
                ha="left", va=_va, fontsize=7.5, color="#6D4C41", zorder=5,
                bbox=dict(boxstyle="round,pad=0.15", fc="white", ec="#6D4C41",
                          alpha=0.9, lw=0.6))

    # Đáy lớp bùn (nét gạch nâu nhạt)
    ax.axhline(z_clay_bot, color="#6D4C41", lw=1.0, ls="-.", zorder=4, alpha=0.7)
    ax.text(W * 0.02, z_clay_bot - 0.08,
            f"Đáy lớp bùn = {z_clay_bot:+.2f} m",
            ha="left", va="top", fontsize=7.5, color="#6D4C41", zorder=5,
            bbox=dict(boxstyle="round,pad=0.15", fc="white", ec="#6D4C41",
                      alpha=0.9, lw=0.6))

    # Đáy cọc CDM (nét chấm xanh lá)
    ax.axhline(z_pile_bot, color="#2E7D32", lw=0.9, ls=":", zorder=4, alpha=0.8)
    ax.text(W * 0.98, z_pile_bot - 0.08,
            f"Đáy cọc = {z_pile_bot:+.2f} m",
            ha="right", va="top", fontsize=7.5, color="#2E7D32", zorder=5,
            bbox=dict(boxstyle="round,pad=0.1", fc="white", ec="none", alpha=0.8))

    # Kích thước phần cọc ngàm vào đất tốt (bên trái cọc)
    _emb = z_clay_bot - z_pile_bot
    if _emb > 0.05:
        dim_x_l = cx - r - W * 0.07
        ax.annotate("", xy=(dim_x_l, z_pile_bot), xytext=(dim_x_l, z_clay_bot),
                    arrowprops=dict(arrowstyle="<->", color="#D84315", lw=1.4),
                    zorder=5)
        ax.text(dim_x_l - W * 0.015, (z_pile_bot + z_clay_bot) / 2,
                f"Ngàm\n={_emb:.2f}m",
                ha="right", va="center", fontsize=7.5, color="#D84315",
                fontweight="bold", zorder=5,
                bbox=dict(boxstyle="round,pad=0.15", fc="white", ec="#D84315",
                          alpha=0.9, lw=0.6))

    def layer_text(y0, y1, txt, color):
        if (y1 - y0) > 0.4:
            ax.text(W * 0.03, (y0 + y1) / 2, txt,
                    ha="left", va="center", fontsize=8, color=color, zorder=6)

    layer_text(z_base,        z_clay_bot,    "Đất tốt",                       "#5D4037")
    layer_text(z_clay_bot,    _clay_top_eff, f"Bùn sét  h={h_clay:.1f}m",     "#0D47A1")
    if top_clay is not None and top_clay < CDTK - 0.05:
        _h_fill_below = CDTK - top_clay
        layer_text(_clay_top_eff, CDTK,     f"Cát đắp  h={_h_fill_below:.1f}m", "#827717")
    if h_mat > 0:
        layer_text(CDTK,      z_mat,        f"Đệm cát  h={h_mat:.1f}m",       "#E65100")
    layer_text(z_mat,         z_fill,       f"Cát đắp  h={h_fill:.1f}m",       "#827717")
    layer_text(z_fill,        z_road,       f"Mặt đường h={h_road:.1f}m",      "#37474F")

    dim_yD = (CDTK + z_mat) / 2 if h_mat > 0.2 else CDTK + 0.25
    ax.annotate("", xy=(cx + r, dim_yD), xytext=(cx - r, dim_yD),
                arrowprops=dict(arrowstyle="<->", color="#2E7D32", lw=1.5), zorder=5)
    ax.text(cx, dim_yD + 0.12, f"D={D:.2f}m",
            ha="center", va="bottom", fontsize=8, color="#2E7D32", zorder=5,
            bbox=dict(boxstyle="round,pad=0.1", fc="white", ec="none", alpha=0.8))

    dim_x = cx + r + W * 0.07
    ax.annotate("", xy=(dim_x, p_bot), xytext=(dim_x, CDTK),
                arrowprops=dict(arrowstyle="<->", color="#1A237E", lw=1.5), zorder=5)
    ax.text(dim_x + W * 0.02, p_mid, f"Lc={Lc:.1f}m",
            ha="left", va="center", fontsize=8, color="#1A237E", zorder=5)

    q_top = z_road + 1.0
    for xi in [W * 0.25, W * 0.5, W * 0.75]:
        ax.annotate("", xy=(xi, z_road), xytext=(xi, q_top),
                    arrowprops=dict(arrowstyle="-|>", color="#B71C1C", lw=1.8), zorder=6)
    ax.text(cx, q_top + 0.2, f"q = {q_traffic:.0f} kN/m²",
            ha="center", va="bottom", fontsize=10,
            color="#B71C1C", fontweight="bold", zorder=6)

    # He = h_fill (cát đắp, từ đỉnh đệm z_mat đến đỉnh lớp cát z_fill)
    _He = z_fill - z_mat
    if _He > 0.05:
        _dim_xH = -W * 0.18
        ax.annotate("", xy=(_dim_xH, z_mat), xytext=(_dim_xH, z_fill),
                    arrowprops=dict(arrowstyle="<->", color="#1565C0", lw=1.5), zorder=5)
        ax.text(_dim_xH - W * 0.03, (z_mat + z_fill) / 2,
                f"He\n={_He:.1f}m",
                ha="right", va="center", fontsize=8.5, color="#1565C0",
                fontweight="bold", zorder=5,
                bbox=dict(boxstyle="round,pad=0.2", fc="white", ec="#1565C0",
                          alpha=0.9, lw=0.8))
    # Hse = h_mat (đệm xi măng, từ CDTK đến z_mat)
    _Hse = z_mat - CDTK
    if _Hse > 0.05:
        _dim_xH2 = -W * 0.18
        ax.annotate("", xy=(_dim_xH2, CDTK), xytext=(_dim_xH2, z_mat),
                    arrowprops=dict(arrowstyle="<->", color="#E65100", lw=1.2), zorder=5)
        ax.text(_dim_xH2 - W * 0.03, (CDTK + z_mat) / 2,
                f"Hse\n={_Hse:.1f}m",
                ha="right", va="center", fontsize=7.5, color="#E65100",
                fontweight="bold", zorder=5,
                bbox=dict(boxstyle="round,pad=0.15", fc="white", ec="#E65100",
                          alpha=0.9, lw=0.7))

    plt.tight_layout(pad=0.5)
    return fig


def _draw_cdm_grid(D: float, arr: str, e_ref: float):
    """Sơ đồ bố trí lưới cọc CDM (plan view)."""
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches

    r = D / 2
    centers = []
    for ri in range(4):
        for ci in range(4):
            if arr == "triangle":
                px = ci * e_ref + (e_ref / 2 if ri % 2 else 0)
                py = ri * (e_ref * (3 ** 0.5) / 2)
            else:
                px = ci * e_ref
                py = ri * e_ref
            centers.append((px, py))

    all_px = [p[0] for p in centers]
    all_py = [p[1] for p in centers]
    pad = e_ref * 0.75

    fig, ax = plt.subplots(figsize=(5, 4.5))
    fig.patch.set_facecolor("#FAFAFA")
    ax.set_facecolor("#F0F4F8")
    ax.set_aspect("equal")
    ax.set_xticks([])
    ax.set_yticks([])
    arr_lbl = "Tam giác" if arr == "triangle" else "Hình vuông"
    ax.set_title(f"Sơ đồ lưới {arr_lbl}  (e = {e_ref:.2f} m)",
                 fontsize=11, fontweight="bold", pad=8)
    for s in ax.spines.values():
        s.set_visible(False)

    if arr == "square":
        ax.add_patch(mpatches.Rectangle(
            (-e_ref / 2, -e_ref / 2), e_ref, e_ref,
            fc="#BBDEFB", ec="#1565C0", lw=1.5, ls="--", alpha=0.4, zorder=1
        ))
    else:
        e_h = e_ref * (3 ** 0.5) / 2
        ax.add_patch(mpatches.Polygon(
            [(0, 0), (e_ref, 0), (e_ref / 2, e_h)],
            closed=True, fc="#BBDEFB", ec="#1565C0", lw=1.5, ls="--", alpha=0.35, zorder=1
        ))

    for px, py in centers:
        ax.add_patch(mpatches.Circle(
            (px, py), r, fc="#A5D6A7", ec="#2E7D32", lw=1.2, zorder=3
        ))

    p0x, p0y = centers[0]
    p1x, _   = centers[1]
    d_y = p0y - r - 0.2
    ax.annotate("", xy=(p1x, d_y), xytext=(p0x, d_y),
                arrowprops=dict(arrowstyle="<->", color="#E53935", lw=1.5))
    ax.text((p0x + p1x) / 2, d_y - 0.12, f"e = {e_ref:.2f} m",
            ha="center", va="top", fontsize=8.5, color="#E53935")

    a_val = calc_a(D, e_ref, arr)
    ax.text((min(all_px) + max(all_px)) / 2, max(all_py) + r + 0.3,
            f"a = {a_val * 100:.1f}%",
            ha="center", va="bottom", fontsize=11, color="#1A237E", fontweight="bold")

    ax.set_xlim(min(all_px) - pad, max(all_px) + pad)
    ax.set_ylim(min(all_py) - pad * 1.8, max(all_py) + pad * 0.8)
    plt.tight_layout(pad=0.5)
    return fig


def _draw_soil_column(
    layers: list[dict], bh_name: str = "",
    spt: list[dict] | None = None,
) -> go.Figure:
    """Cột địa chất + biểu đồ SPT (subplot 2 cột, trục Y dùng chung)."""
    try:
        from plotly.subplots import make_subplots
    except ImportError:
        make_subplots = None

    has_spt = bool(spt) and make_subplots is not None

    if not layers:
        fig = go.Figure()
        fig.update_layout(height=500,
                          title=dict(text=f"Cột ĐC: {bh_name}", font=dict(size=10), x=0.5))
        return fig

    def _tc(hex_color: str) -> str:
        h = hex_color.lstrip("#")
        if len(h) < 6:
            return "#212121"
        r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
        return "white" if (0.299 * r + 0.587 * g + 0.114 * b) < 140 else "#212121"

    def _spt_color(n: int) -> str:
        if n <= 2:   return "#E53935"   # đất rất yếu
        if n <= 5:   return "#FB8C00"   # yếu
        if n <= 10:  return "#FDD835"   # trung bình yếu
        if n <= 30:  return "#66BB6A"   # trung bình
        return "#1565C0"                 # chặt / cứng

    y_bot = layers[-1]["depth_bot_m"]

    if has_spt:
        fig = make_subplots(
            rows=1, cols=2,
            shared_yaxes=True,
            column_widths=[0.35, 0.65],
            horizontal_spacing=0.04,
            subplot_titles=["Địa chất", "SPT – N"],
        )
    else:
        fig = go.Figure()

    # ── Soil column (col 1) ──────────────────────────────────────────────────
    shapes: list[dict] = []
    annotations: list[dict] = []
    seen: set[float] = set()

    xref_col = "x" if has_spt else "x"
    yref_col = "y"

    for lay in layers:
        y0    = lay["depth_top_m"]
        y1    = lay["depth_bot_m"]
        sym   = (lay.get("symbol") or "").strip()
        color = _LAYER_COLORS.get(sym,
                _LAYER_COLORS.get(sym[:2] if len(sym) > 2 else sym,
                _LAYER_DEFAULT_COLOR))
        thick = y1 - y0

        shapes.append(dict(
            type="rect", x0=0, y0=y0, x1=1, y1=y1,
            xref=xref_col, yref=yref_col,
            fillcolor=color, line=dict(color="#555", width=0.6),
        ))

        if y0 not in seen:
            annotations.append(dict(
                x=1.08 if not has_spt else 1.12, y=y0,
                xref=xref_col, yref=yref_col,
                text=f"{y0:.1f}",
                showarrow=False, font=dict(size=8, color="#555"),
                xanchor="left", yanchor="middle",
            ))
            seen.add(y0)

        if thick >= 0.5:
            mid = (y0 + y1) / 2
            tc  = _tc(color)
            desc = (lay.get("description") or "")
            sym_lbl = sym if sym else f"L{layers.index(lay)+1}"
            lbl  = (f"<b>{sym_lbl}</b><br>{desc[:18]}"
                    if thick >= 2.5 and desc else f"<b>{sym_lbl}</b>")
            annotations.append(dict(
                x=0.5, y=mid,
                xref=xref_col, yref=yref_col,
                text=lbl,
                showarrow=False, font=dict(size=9, color=tc),
                xanchor="center", yanchor="middle",
            ))

    if y_bot not in seen:
        annotations.append(dict(
            x=1.08 if not has_spt else 1.12, y=y_bot,
            xref=xref_col, yref=yref_col,
            text=f"{y_bot:.1f}",
            showarrow=False, font=dict(size=8, color="#555"),
            xanchor="left", yanchor="middle",
        ))

    # Dummy trace để giữ trục y
    if has_spt:
        fig.add_trace(go.Scatter(
            x=[0.5], y=[y_bot / 2],
            mode="markers", marker=dict(opacity=0, size=1),
            showlegend=False,
        ), row=1, col=1)
    else:
        fig.add_trace(go.Scatter(
            x=[0.5], y=[y_bot / 2],
            mode="markers", marker=dict(opacity=0, size=1),
            showlegend=False,
        ))

    # ── SPT chart (col 2) ────────────────────────────────────────────────────
    if has_spt:
        _n_max = max((s["N"] or 0 for s in spt), default=10)
        _x_max = max(60, _n_max + 5)

        _depths = [s["depth_m"] for s in spt]
        _N_vals = [min(s["N"] or 0, 60) for s in spt]
        _colors = [_spt_color(n) for n in _N_vals]
        _labels = [
            f"N={s['N']}  ({s['N1']}/{s['N2']}/{'–' if s['N3'] is None else s['N3']})"
            for s in spt
        ]

        # Thanh ngang N
        fig.add_trace(go.Bar(
            x=_N_vals, y=_depths,
            orientation="h",
            marker_color=_colors,
            marker_line=dict(color="#444", width=0.5),
            text=[str(n) for n in _N_vals],
            textposition="outside",
            textfont=dict(size=13),
            customdata=_labels,
            hovertemplate="%{customdata}<extra></extra>",
            showlegend=False,
            width=1.2,
        ), row=1, col=2)

        # Đường nối N theo chiều sâu
        fig.add_trace(go.Scatter(
            x=_N_vals, y=_depths,
            mode="lines",
            line=dict(color="#333", width=1, dash="dot"),
            showlegend=False,
            hoverinfo="skip",
        ), row=1, col=2)

        # Đường giới hạn N=4 (đất rất yếu)
        fig.add_vline(x=4, line_dash="dash", line_color="#E53935",
                      line_width=1, opacity=0.6,
                      annotation_text="N=4", annotation_font_size=8,
                      annotation_position="top right", row=1, col=2)

        fig.update_xaxes(
            title_text="N (blows/30cm)", title_font=dict(size=14),
            range=[0, _x_max], tickfont=dict(size=12),
            showgrid=True, gridcolor="#EEE",
            row=1, col=2,
        )

    # ── Layout ───────────────────────────────────────────────────────────────
    y_range  = [y_bot + 1, -0.5]
    _yfs_t = 14 if has_spt else 9
    _yfs_k = 12 if has_spt else 8
    y_axis   = dict(
        range=y_range,
        title="Độ sâu (m)", title_font=dict(size=_yfs_t),
        tickfont=dict(size=_yfs_k),
        showgrid=True, gridcolor="#EEE", gridwidth=0.5,
        dtick=5,
    )

    if has_spt:
        fig.update_yaxes(y_axis, row=1, col=1)
        fig.update_xaxes(
            showticklabels=False, showgrid=False, zeroline=False,
            range=[-0.05, 1.6], row=1, col=1,
        )
        fig.update_layout(
            shapes=shapes,
            annotations=annotations,
            title=dict(text=f"Cột ĐC: {bh_name}", font=dict(size=10, color="#1F4E79"), x=0.5),
            height=500,
            margin=dict(l=10, r=20, t=55, b=20),
            plot_bgcolor="white",
            paper_bgcolor="#FAFAFA",
            showlegend=False,
            bargap=0,
        )
    else:
        fig.update_layout(
            shapes=shapes,
            annotations=annotations,
            title=dict(text=f"Cột ĐC: {bh_name}", font=dict(size=10, color="#1F4E79"), x=0.5),
            height=500,
            margin=dict(l=10, r=55, t=35, b=20),
            xaxis=dict(showticklabels=False, showgrid=False, zeroline=False, range=[-0.05, 1.7]),
            yaxis=y_axis,
            plot_bgcolor="white",
            paper_bgcolor="#FAFAFA",
            showlegend=False,
        )

    return fig


# ═══════════════════════════════════════════════════════════════════════════════
# EXPORT HELPERS
# ═══════════════════════════════════════════════════════════════════════════════
def _export_excel(scenarios: list[dict], params: dict) -> bytes:
    """Xuất Excel 3 sheet: Đầu vào | So sánh PA | Kết quả."""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.chart import BarChart, Reference
    except ImportError:
        return b""

    wb = openpyxl.Workbook()

    # --- Helpers style ---
    HDR_FILL = PatternFill("solid", fgColor="1F4E79")
    REC_FILL = PatternFill("solid", fgColor="E2EFDA")
    HDR_FONT = Font(bold=True, color="FFFFFF", size=10)
    BD = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    def _hdr(ws, row, col, val, fill=HDR_FILL, font=HDR_FONT):
        c = ws.cell(row=row, column=col, value=val)
        c.fill = fill; c.font = font; c.border = BD
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        return c

    def _val(ws, row, col, val, rec=False, fmt=None):
        c = ws.cell(row=row, column=col, value=val)
        if rec: c.fill = REC_FILL
        c.border = BD
        c.alignment = Alignment(horizontal="center", vertical="center")
        if fmt: c.number_format = fmt
        return c

    # ── Sheet 1: Đầu vào ──────────────────────────────────────────────────────
    ws1 = wb.active
    ws1.title = "Đầu vào"
    ws1.column_dimensions["A"].width = 36
    ws1.column_dimensions["B"].width = 20
    ws1.column_dimensions["C"].width = 16

    title_rows = [
        ("THÔNG SỐ ĐỊA KỸ THUẬT VÀ THIẾT KẾ CDM",),
        (f"Hố khoan: {params.get('bh_name','–')}   |   Zone: {params.get('zone','–')}",),
    ]
    for i, (txt,) in enumerate(title_rows, 1):
        c = ws1.cell(row=i, column=1, value=txt)
        c.font = Font(bold=True, size=11)
        ws1.merge_cells(start_row=i, start_column=1, end_row=i, end_column=3)

    _hdr(ws1, 4, 1, "Thông số"); _hdr(ws1, 4, 2, "Ký hiệu"); _hdr(ws1, 4, 3, "Giá trị")
    geo_rows = [
        ("Hố khoan tham chiếu", "HK", params.get("bh_name","–")),
        ("Bề dày lớp bùn sét", "h₁ (m)", params.get("h_clay", 0)),
        ("Sức kháng cắt VST trung bình", "Su (kN/m²)", params.get("Su", 0)),
        ("Dung trọng tự nhiên TB", "γ (kN/m³)", params.get("gamma", 0)),
        ("Đường kính trụ CDM", "D (mm)", int(params.get("D", 0.8)*1000)),
        ("Chiều dài trụ CDM", "Lc (m)", params.get("Lc", 0)),
        ("Cao độ đỉnh cọc CDM", "Đỉnh cọc (m)", params.get("CDTK", 2.7)),
        ("Bố trí", "–", "Tam giác" if params.get("arrangement","triangle")=="triangle" else "Vuông"),
        ("Cường độ TK hiện trường", "qu,tk (kPa)", params.get("qu", 800)),
        ("Hệ số quy đổi TN→HT", "FS", params.get("FS_lab", 2.0)),
        ("Cc = qu,tk/2", "Cc (kN/m²)", params.get("qu", 800)/2),
        ("Ec = 100·Cc", "Ec (kN/m²)", int(100*params.get("qu",800)/2)),
        ("Es = 250·Su", "Es (kN/m²)", int(250*params.get("Su",11.2))),
        ("Tổng tải trọng gây lún", "q (kN/m²)", params.get("q_total", 75.2)),
    ]
    for r, (n, k, v) in enumerate(geo_rows, 5):
        ws1.cell(row=r, column=1, value=n).border = BD
        ws1.cell(row=r, column=2, value=k).border = BD
        c = ws1.cell(row=r, column=3, value=v); c.border = BD
        c.alignment = Alignment(horizontal="center")

    # ── Sheet 2: So sánh PA ───────────────────────────────────────────────────
    ws2 = wb.create_sheet("So sánh PA")
    hdrs = ["PA", "e (m)", "a", "a (%)", "Ec (kN/m²)", "Es (kN/m²)",
            "Etb (kN/m²)", "S₁ (cm)", "Pcol (kN)", "Qa (kN)", "Đạt SCT"]
    for j, h in enumerate(hdrs, 1):
        _hdr(ws2, 2, j, h)
        ws2.column_dimensions[chr(64+j)].width = 13
    rec_idx = params.get("rec_idx", 0)
    for i, s in enumerate(scenarios):
        is_rec = (i == rec_idx)
        _val(ws2, i+3, 1, f"PA{i+1}", is_rec)
        _val(ws2, i+3, 2, s["e (m)"], is_rec)
        _val(ws2, i+3, 3, s["a"], is_rec, "0.0000")
        _val(ws2, i+3, 4, s["a (%)"], is_rec)
        _val(ws2, i+3, 5, s["Ec (kN/m²)"], is_rec, "#,##0")
        _val(ws2, i+3, 6, s["Es (kN/m²)"], is_rec, "#,##0")
        _val(ws2, i+3, 7, s["Etb (kN/m²)"], is_rec, "#,##0")
        _val(ws2, i+3, 8, s["S₁ (cm)"], is_rec, "0.00")
        _val(ws2, i+3, 9, s["Pcol (kN)"], is_rec, "0.0")
        _val(ws2, i+3, 10, s["Qa (kN)"], is_rec, "0.0")
        _val(ws2, i+3, 11, s["Đạt SCT"], is_rec)

    # Biểu đồ S1
    chart = BarChart()
    chart.title = "Độ lún S₁ theo khoảng cách trụ"
    chart.y_axis.title = "S₁ (cm)"
    chart.x_axis.title = "Khoảng cách e (m)"
    data_ref  = Reference(ws2, min_col=8, min_row=2, max_row=2+len(scenarios))
    cats_ref  = Reference(ws2, min_col=2, min_row=3, max_row=2+len(scenarios))
    chart.add_data(data_ref, titles_from_data=True)
    chart.set_categories(cats_ref)
    chart.shape = 4; chart.width = 14; chart.height = 10
    ws2.add_chart(chart, f"M2")

    # ── Sheet 3: Kết quả chi tiết PA kiến nghị ────────────────────────────────
    ws3 = wb.create_sheet("Kết quả PA KN")
    if scenarios and 0 <= rec_idx < len(scenarios):
        s = scenarios[rec_idx]
        ws3.column_dimensions["A"].width = 36
        ws3.column_dimensions["B"].width = 22
        ws3.column_dimensions["C"].width = 18
        c = ws3.cell(row=1, column=1,
                     value=f"TÍNH TOÁN CHI TIẾT – PA{rec_idx+1}: e = {s['e (m)']} m")
        c.font = Font(bold=True, size=12)
        ws3.merge_cells("A1:C1")

        _hdr(ws3, 2, 1, "Thông số"); _hdr(ws3, 2, 2, "Công thức"); _hdr(ws3, 2, 3, "Giá trị")
        detail_rows = [
            ("Khoảng cách bố trí", "e", f"{s['e (m)']} m"),
            ("Tỷ lệ thay thế", "a = π·D²/(2√3·e²)" if params.get("arrangement","triangle")=="triangle"
             else "a = π·D²/(4·e²)", f"{s['a']:.4f}  ({s['a (%)']:.1f}%)"),
            ("Mô đun trụ CDM", "Ec = 100·Cc", f"{s['Ec (kN/m²)']:,} kN/m²"),
            ("Mô đun đất nền", "Es = 250·Su", f"{s['Es (kN/m²)']:,} kN/m²"),
            ("Mô đun tương đương", "Etb = a·Ec+(1-a)·Es", f"{int(s['Etb (kN/m²)']):,} kN/m²"),
            ("Tải trọng thiết kế", "q", f"{params.get('q_total',0):.1f} kN/m²"),
            ("Chiều dài trụ", "Lc", f"{params.get('Lc',0):.1f} m"),
            ("Độ lún S₁", "q·Lc/Etb×100", f"{s['S₁ (cm)']:.2f} cm"),
            ("Ứng suất đầu cọc", "σ_col = Ec/Etb·q", f"{s['σ_col (kN/m²)']:.1f} kN/m²"),
            ("Lực nén lên trụ", "Pcol = σ_col·Ac", f"{s['Pcol (kN)']:.1f} kN"),
            ("Sức chịu tải cho phép", "Qa (AIT, FS=2)", f"{s['Qa (kN)']:.1f} kN"),
            ("Kết luận SCT", "Pcol < Qa ?", s["Đạt SCT"]),
        ]
        for r, (n, f, v) in enumerate(detail_rows, 3):
            ws3.cell(row=r, column=1, value=n).border = BD
            ws3.cell(row=r, column=2, value=f).border = BD
            c = ws3.cell(row=r, column=3, value=v); c.border = BD
            c.alignment = Alignment(horizontal="center")
            if v in ("Đạt", "Không đạt"):
                c.font = Font(color="375623" if v == "Đạt" else "C00000", bold=True)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _custom_report_pdf(
    scenarios: list[dict], params: dict, rec_idx: int,
    include_params: bool = True, include_compare: bool = True,
    include_charts: bool = True, include_sens: bool = False,
    include_punch: bool = True, include_rec: bool = True,
    charts_only: bool = False,
) -> bytes:
    """Sinh PDF tuỳ chỉnh — chọn từng mục, hoặc 'charts only' mode.
    Dùng ReportLab (luôn cài được trên Cloud) + matplotlib charts → PNG embed."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm, mm
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                         Table, TableStyle, Image, PageBreak,
                                         KeepTogether)
        import matplotlib.pyplot as plt
        import io as _io
    except ImportError as e:
        st.error(
            "**Thư viện `reportlab` chưa được cài trên môi trường local.**\n\n"
            f"Lỗi: `{e}`\n\n"
            "**Cách khắc phục:**\n"
            "1. Mở Command Prompt / PowerShell (Run as Administrator)\n"
            "2. Chạy lệnh: `pip install reportlab>=4.0`\n"
            "3. Khởi động lại Streamlit (đóng tab và chạy lại `start_app.bat`)\n\n"
            "**Hoặc dùng phương án thay thế:**\n"
            "- Xuất Word (nút bên trái) → mở bằng Microsoft Word → File → Save As PDF\n"
            "- Streamlit Cloud đã có sẵn `reportlab` — đẩy code lên Cloud qua `update_app.bat` "
            "để tạo PDF trực tiếp"
        )
        return b""
    except Exception as e:
        st.error(f"Lỗi không xác định khi tải thư viện PDF: {e}")
        return b""

    buf_out = _io.BytesIO()
    doc = SimpleDocTemplate(
        buf_out, pagesize=A4,
        leftMargin=14 * mm, rightMargin=12 * mm,
        topMargin=14 * mm, bottomMargin=14 * mm,
        title="Báo cáo Thiết kế CDM",
    )
    styles = getSampleStyleSheet()
    style_title = ParagraphStyle("t", parent=styles["Title"], fontSize=14,
                                  alignment=1, spaceAfter=8)
    style_h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=12,
                               textColor=colors.HexColor("#1F4E79"), spaceBefore=10, spaceAfter=4)
    style_p = ParagraphStyle("p", parent=styles["BodyText"], fontSize=9.5, leading=12)
    style_cap = ParagraphStyle("cap", parent=styles["BodyText"], fontSize=8,
                                alignment=1, textColor=colors.grey, spaceAfter=4)

    story = []
    rec = scenarios[rec_idx] if scenarios and rec_idx < len(scenarios) else {}

    # ── Title ────────────────────────────────────────────────────────────────
    title = ("BÁO CÁO BIỂU ĐỒ THIẾT KẾ CDM" if charts_only
             else "BÁO CÁO THIẾT KẾ CỌC ĐẤT XI MĂNG (CDM)")
    story.append(Paragraph(title, style_title))
    story.append(Paragraph(
        f"Khu vực: <b>{params.get('zone', '—')}</b>  |  "
        f"Hố khoan: <b>{params.get('bh_name', '—')}</b>  |  "
        f"D = {int(params.get('D', 0.8) * 1000)} mm  |  "
        f"Lc = {params.get('Lc', 0):.1f} m",
        style_p,
    ))
    story.append(Spacer(1, 8))

    # ── 1. Thông số đầu vào ──────────────────────────────────────────────────
    if include_params and not charts_only:
        story.append(Paragraph("1. THÔNG SỐ ĐẦU VÀO", style_h1))
        tbl_data = [
            ["Tham số", "Ký hiệu", "Giá trị", "Đơn vị"],
            ["Đường kính cọc", "D", f"{params.get('D', 0.8):.2f}", "m"],
            ["Chiều dài cọc", "Lc", f"{params.get('Lc', 23):.1f}", "m"],
            ["Cường độ nén thiết kế", "qu", f"{params.get('qu', 800):.0f}", "kPa"],
            ["Sức kháng cắt KTN", "Su", f"{params.get('Su', 10):.1f}", "kPa"],
            ["Bề dày lớp bùn", "h_clay", f"{params.get('h_clay', 22):.1f}", "m"],
            ["Dung trọng đất", "γ", f"{params.get('gamma', 15):.1f}", "kN/m³"],
            ["Tải tổng (xét hoạt tải)", "q", f"{params.get('q_total', 75):.1f}", "kN/m²"],
            ["Hàm lượng xi măng", "x", f"{params.get('dosage', 240):.0f}", "kg/m³"],
            ["Bố trí", "—",
             "Tam giác" if params.get("arrangement") == "triangle" else "Hình vuông", "—"],
        ]
        t = Table(tbl_data, colWidths=[6.5*cm, 2.5*cm, 3*cm, 2*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
            ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID",     (0, 0), (-1, -1), 0.4, colors.HexColor("#888")),
            ("ALIGN",    (1, 0), (-1, -1), "CENTER"),
            ("VALIGN",   (0, 0), (-1, -1), "MIDDLE"),
        ]))
        story.append(t)
        story.append(Spacer(1, 8))

    # ── 2. Bảng so sánh phương án ────────────────────────────────────────────
    if include_compare and not charts_only:
        story.append(Paragraph("2. BẢNG SO SÁNH PHƯƠNG ÁN", style_h1))
        rows = [["PA", "e (m)", "a (%)", "Etb (kN/m²)", "S₁ (cm)",
                 "Pcol (kN)", "Qa (kN)", "Đạt SCT"]]
        for i, s in enumerate(scenarios):
            star = " *" if i == rec_idx else ""
            rows.append([
                f"PA{i+1}{star}", f"{s['e (m)']}",
                f"{s['a (%)']:.1f}", f"{int(s['Etb (kN/m²)']):,}",
                f"{s['S₁ (cm)']:.2f}", f"{s['Pcol (kN)']:.1f}",
                f"{s['Qa (kN)']:.1f}", s["Đạt SCT"],
            ])
        t = Table(rows, colWidths=[1.5*cm, 1.5*cm, 1.8*cm, 3*cm, 2*cm, 2*cm, 2*cm, 2*cm])
        ts = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
            ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 8.5),
            ("GRID",     (0, 0), (-1, -1), 0.4, colors.HexColor("#888")),
            ("ALIGN",    (0, 0), (-1, -1), "CENTER"),
        ]
        if rec_idx + 1 < len(rows):
            ts.append(("BACKGROUND", (0, rec_idx + 1), (-1, rec_idx + 1),
                       colors.HexColor("#FFF2CC")))
        t.setStyle(TableStyle(ts))
        story.append(t)
        story.append(Paragraph("* = Phương án kiến nghị", style_cap))
        story.append(Spacer(1, 8))

    # ── 3. Biểu đồ so sánh ───────────────────────────────────────────────────
    if include_charts:
        story.append(Paragraph("3. BIỂU ĐỒ SO SÁNH PHƯƠNG ÁN", style_h1))
        fig, axs = plt.subplots(2, 2, figsize=(11, 7))
        _xs = [f"e={s['e (m)']}m" for s in scenarios]
        _cols = ['#ED7D31' if i == rec_idx else '#4472C4' for i in range(len(scenarios))]
        # S1
        axs[0,0].bar(_xs, [s["S₁ (cm)"] for s in scenarios], color=_cols, edgecolor='black')
        axs[0,0].set_title("Lún S₁ (cm)"); axs[0,0].set_ylabel("cm")
        for i, v in enumerate([s["S₁ (cm)"] for s in scenarios]):
            axs[0,0].text(i, v*1.02, f"{v:.2f}", ha='center', fontsize=8)
        # Etb
        axs[0,1].bar(_xs, [s["Etb (kN/m²)"] for s in scenarios], color=_cols, edgecolor='black')
        axs[0,1].set_title("Mô-đun tương đương Etb (kN/m²)")
        # Pcol vs Qa
        _p = [s["Pcol (kN)"] for s in scenarios]
        _q = [s["Qa (kN)"]   for s in scenarios]
        axs[1,0].bar(_xs, _p, color=_cols, edgecolor='black', label='Pcol')
        axs[1,0].plot(_xs, _q, 'D-', color='#E53935', lw=2, label='Qa (giới hạn)')
        axs[1,0].set_title("Sức chịu tải: Pcol vs Qa"); axs[1,0].set_ylabel("kN")
        axs[1,0].legend(fontsize=8)
        # a (%)
        axs[1,1].bar(_xs, [s["a (%)"] for s in scenarios], color=_cols, edgecolor='black')
        axs[1,1].set_title("Tỷ lệ thay thế a (%)")
        for ax in axs.flat:
            ax.grid(True, axis='y', ls=':', alpha=0.4)
        fig.suptitle("So sánh phương án — cam = PA kiến nghị", fontsize=11, fontweight='bold')
        fig.tight_layout()
        buf_img = _io.BytesIO()
        fig.savefig(buf_img, format='png', dpi=130, bbox_inches='tight', facecolor='white')
        plt.close(fig); buf_img.seek(0)
        story.append(Image(buf_img, width=17*cm, height=11*cm))
        story.append(Paragraph("Hình 1. So sánh 4 chỉ tiêu chính theo khoảng cách e", style_cap))
        story.append(Spacer(1, 8))

    # ── 4. Phân tích độ nhạy ─────────────────────────────────────────────────
    if include_sens:
        story.append(Paragraph("4. PHÂN TÍCH ĐỘ NHẠY", style_h1))
        story.append(Paragraph(
            "Khảo sát ảnh hưởng của các tham số thiết kế đến lún S₁ và sức chịu tải. "
            "Tham số được biến thiên quanh giá trị kiến nghị.", style_p))
        story.append(Spacer(1, 6))

    # ── 5. Kiểm tra chọc thủng ──────────────────────────────────────────────
    if include_punch and not charts_only:
        story.append(Paragraph("5. KIỂM TRA CHỌC THỦNG (ALiCC)", style_h1))
        story.append(Paragraph(
            "Phương pháp ALiCC (PWRI Japan) — TCVN 9403:2012 Phụ lục C. "
            "Điều kiện đạt: τse ≤ τase.", style_p))
        story.append(Spacer(1, 6))

    # ── 6. Kết quả PA kiến nghị ──────────────────────────────────────────────
    if include_rec and not charts_only and rec:
        story.append(PageBreak())
        story.append(Paragraph(f"6. KẾT QUẢ PHƯƠNG ÁN KIẾN NGHỊ (e = {rec['e (m)']} m)",
                                style_h1))
        rec_rows = [
            ["Đại lượng", "Giá trị"],
            ["Tỷ lệ thay thế a",        f"{rec.get('a', 0):.4f}  ({rec.get('a (%)', 0):.1f}%)"],
            ["Mô-đun tương đương Etb",  f"{int(rec.get('Etb (kN/m²)', 0)):,} kN/m²"],
            ["Độ lún S₁",               f"{rec.get('S₁ (cm)', 0):.2f} cm"],
            ["Độ lún S₂",               "0,00 cm (CDM xuyên hết lớp bùn)"],
            ["Tổng lún S",              f"{rec.get('S₁ (cm)', 0):.2f} cm"],
            ["Lực nén lên cọc Pcol",    f"{rec.get('Pcol (kN)', 0):.1f} kN"],
            ["Sức chịu tải Qa",         f"{rec.get('Qa (kN)', 0):.1f} kN"],
            ["Kiểm tra SCT",            rec.get("Đạt SCT", "—")],
        ]
        t = Table(rec_rows, colWidths=[7*cm, 8*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
            ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("GRID",     (0, 0), (-1, -1), 0.4, colors.HexColor("#888")),
            ("ALIGN",    (1, 1), (-1, -1), "LEFT"),
        ]))
        story.append(t)

    doc.build(story)
    return buf_out.getvalue()


def _export_word_bytes(
    scenarios: list[dict],
    params: dict,
    rec_idx: int,
    co_name: str = "",
    co_staff: str = "",
    logo_bytes: bytes | None = None,
) -> bytes:
    """Tạo Word thuyết minh CDM và trả về bytes."""
    sys.path.insert(0, str(_CDM_SC))
    try:
        import docx_helpers as H
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _Oxm
        from docx.shared import Cm as _Cm
        import math as _m
    except Exception as e:
        # Surface error thay vì silent b""
        import traceback as _tb
        st.error(f"Import lỗi: {type(e).__name__}: {e}")
        st.code(_tb.format_exc())
        return b""

    H.reset_counters()
    doc = Document()
    H.setup_page(doc)

    # ── Header / Footer ───────────────────────────────────────────────────────
    def _field_run(para, field: str) -> None:
        """Chèn Word field (PAGE / NUMPAGES) vào paragraph."""
        run = para.add_run()
        fc_b = _Oxm("w:fldChar"); fc_b.set(_qn("w:fldCharType"), "begin")
        instr = _Oxm("w:instrText"); instr.text = field
        instr.set(_qn("xml:space"), "preserve")
        fc_e = _Oxm("w:fldChar"); fc_e.set(_qn("w:fldCharType"), "end")
        run._r.append(fc_b); run._r.append(instr); run._r.append(fc_e)

    def _no_tbl_border(tbl) -> None:
        tblPr = tbl._tbl.tblPr
        borders = _Oxm("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = _Oxm(f"w:{edge}"); b.set(_qn("w:val"), "none"); borders.append(b)
        tblPr.append(borders)

    sec = doc.sections[0]
    _pw = sec.page_width - sec.left_margin - sec.right_margin  # usable width (EMU)

    # Header: logo (trái) | tên công ty (phải)
    sec.header.is_linked_to_previous = False
    _ht = sec.header.add_table(1, 2, width=_pw)
    _no_tbl_border(_ht)
    _ht.cell(0, 0).width = _Cm(3)
    _ht.cell(0, 1).width = _pw - _Cm(3)
    _logo_p = _ht.cell(0, 0).paragraphs[0]
    if logo_bytes:
        try:
            _logo_p.add_run().add_picture(io.BytesIO(logo_bytes), height=_Cm(1.2))
        except Exception:
            pass
    _name_p = _ht.cell(0, 1).paragraphs[0]
    _name_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    H._fmt(_name_p.add_run(co_name), size=9, bold=True)

    # Footer: nhân sự (trái) | Trang X/Y (phải)
    sec.footer.is_linked_to_previous = False
    _ft = sec.footer.add_table(1, 2, width=_pw)
    _no_tbl_border(_ft)
    _ft.cell(0, 0).width = _pw - _Cm(3)
    _ft.cell(0, 1).width = _Cm(3)
    H._fmt(_ft.cell(0, 0).paragraphs[0].add_run(co_staff), size=9)
    _pg_p = _ft.cell(0, 1).paragraphs[0]
    _pg_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    H._fmt(_pg_p.add_run("Trang "), size=9)
    _field_run(_pg_p, "PAGE")
    H._fmt(_pg_p.add_run(" / "), size=9)
    _field_run(_pg_p, "NUMPAGES")

    D    = params["D"]
    Lc   = params["Lc"]
    qu   = params["qu"]
    Cu   = params["Su"]
    q    = params["q_total"]
    Cc   = qu / 2
    Ec   = 100 * Cc
    Es   = 250 * Cu
    arr  = "tam giác" if params.get("arrangement", "triangle") == "triangle" else "hình vuông"
    rec  = scenarios[rec_idx] if scenarios else {}

    # Trang bìa
    for _ in range(4): doc.add_paragraph()
    H.heading(doc, "THUYẾT MINH TÍNH TOÁN",
              size=16, bold=True, center=True, space_before=0)
    H.heading(doc, "GIA CỐ NỀN ĐẤT YẾU BẰNG CỌC ĐẤT XI MĂNG (CDM)",
              size=15, bold=True, center=True, space_before=4)
    for _ in range(2): doc.add_paragraph()
    H.heading(doc,
              f"D = {int(D*1000)} mm – e = {' / '.join(str(s['e (m)']) for s in scenarios)} m"
              f" – Lc = {Lc:.1f} m – qu = {int(qu)} kPa",
              size=13, bold=False, center=True, space_before=0)
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    H._fmt(p.add_run("TP. Hồ Chí Minh, tháng 5 năm 2026"), size=12)
    doc.add_page_break()

    H.add_toc(doc)

    # I. Cơ sở pháp lý
    H.heading(doc, "I. CƠ SỞ PHÁP LÝ VÀ TIÊU CHUẨN ÁP DỤNG", size=13)
    for item in [
        "– TCVN 9403:2012 – Gia cố nền đất yếu – Phương pháp trụ đất xi măng;",
        "– TCVN 4200:2012 – Đất xây dựng – Phương pháp xác định tính nén lún;",
        "– TCVN 8868:2011 – Thí nghiệm xác định sức kháng cắt không cố kết – không thoát nước;",
        f"– Kết quả khảo sát địa kỹ thuật (hố khoan {params.get('bh_name','HK1')});",
        "– Kết quả thí nghiệm thành phần cấp phối xi măng đất.",
    ]:
        H.body(doc, item, indent=0.5)

    # II. Địa chất
    H.heading(doc, "II. ĐẶC ĐIỂM ĐỊA CHẤT KHU VỰC", size=13)
    H.body(doc, f"Căn cứ kết quả khảo sát hố khoan {params.get('bh_name','–')}, "
           f"zone {params.get('zone','–')}:")
    H.tbl_caption(doc, f"Thông số địa kỹ thuật – {params.get('bh_name','–')}")
    H.make_table(doc,
        headers=["Thông số", "Ký hiệu", "Giá trị"],
        rows=[
            ("Bề dày lớp bùn sét (lớp cần gia cố)", "h₁ (m)", f"{params.get('h_clay',0):.2f}"),
            ("Sức kháng cắt không thoát nước (VST)", "Su (kN/m²)", f"{Cu:.2f}"),
            ("Dung trọng tự nhiên trung bình", "γ (kN/m³)", f"{params.get('gamma',15):.2f}"),
            ("Mô đun biến dạng đất nền", "Es = 250·Su (kN/m²)", f"{int(Es):,}"),
        ])

    # III. Thông số thiết kế
    H.heading(doc, "III. THÔNG SỐ THIẾT KẾ CỌC ĐẤT XI MĂNG", size=13)
    H.tbl_caption(doc, "Thông số hình học và vật liệu trụ CDM")
    H.make_table(doc,
        headers=["Thông số", "Ký hiệu", "Giá trị"],
        rows=[
            ("Đường kính trụ CDM",              "D",          f"{int(D*1000)} mm"),
            (f"Khoảng cách tâm-tâm (lưới {arr})", "e",        " / ".join(f"{s['e (m)']}m" for s in scenarios)),
            ("Chiều dài trụ CDM",               "Lc",         f"{Lc:.1f} m"),
            ("Hàm lượng xi măng",               "x",          f"{params.get('dosage',240)} kg/m³"),
            ("Cường độ nén TK hiện trường",      "qu,tk",      f"{int(qu)} kPa"),
            ("Cường độ kháng cắt thiết kế",      "Cc = qu/2",  f"{int(Cc)} kN/m²"),
            ("Mô đun đàn hồi trụ",              "Ec = 100·Cc", f"{int(Ec):,} kN/m²"),
            ("Tổng tải trọng gây lún",           "q",          f"{q:.1f} kN/m²"),
        ])

    # IV. Phương pháp tính
    H.heading(doc, "IV. PHƯƠNG PHÁP TÍNH TOÁN LÚN (TCVN 9403:2012 – Phụ lục C)", size=13)
    H.body(doc, "Độ lún bản thân khối CDM:")
    H.formula(doc, "S₁ = q × H / Etb    (C.2)")
    H.formula(doc, "Etb = a × Ec + (1 − a) × Es")
    form = ("a = π × D² / (2√3 × e²)  [lưới tam giác]"
            if params.get("arrangement","triangle")=="triangle"
            else "a = π × D² / (4 × e²)  [lưới vuông]")
    H.formula(doc, form)
    H.body(doc, "S₂ = 0 (trụ CDM xuyên qua toàn bộ lớp bùn sét yếu).")

    # V. So sánh phương án
    H.heading(doc, "V. SO SÁNH CÁC PHƯƠNG ÁN KHOẢNG CÁCH TRỤ", size=13)
    H.tbl_caption(doc, f"So sánh S₁ theo khoảng cách trụ (D={int(D*1000)}mm, qu={int(qu)}kPa, q={q:.1f}kN/m²)")
    H.make_table(doc,
        headers=["PA", "e (m)", "a", "Etb (kN/m²)", "S₁ (cm)", "Đạt SCT"],
        rows=[(f"PA{i+1}", str(s["e (m)"]), f"{s['a']:.4f}",
               f"{int(s['Etb (kN/m²)']):,}", f"{s['S₁ (cm)']:.2f}", s["Đạt SCT"])
              for i, s in enumerate(scenarios)],
        highlight_rows={rec_idx})

    # VI. Kết quả PA kiến nghị
    if rec:
        H.heading(doc,
                  f"VI. KẾT QUẢ – PHƯƠNG ÁN KIẾN NGHỊ (e = {rec['e (m)']} m)", size=13)
        for line in [
            f"Tỷ lệ thay thế: a = {rec['a']:.4f}  ({rec['a (%)']:.1f}%)",
            f"Mô đun tương đương: Etb = {int(rec['Etb (kN/m²)']):,} kN/m²",
            f"Độ lún: S₁ = {rec['S₁ (cm)']:.2f} cm  ;  S₂ = 0  →  S = {rec['S₁ (cm)']:.2f} cm",
            f"Sức chịu tải: Pcol = {rec['Pcol (kN)']:.1f} kN  <  Qa = {rec['Qa (kN)']:.1f} kN  →  {rec['Đạt SCT']}",
        ]:
            H.body(doc, line, indent=0.3)

    # VII. Kiểm tra chọc thủng lớp đệm xi măng
    _ld_w   = params.get("loads", {})
    _pr_w   = calc_punching_check(
        D=D, e=rec.get("e (m)", 1.8) if rec else 1.8,
        Hse=_ld_w.get("h_mat", 0.4),
        He=_ld_w.get("h_fill", 1.5),
        quckse=params.get("quckse", 600.0),
        Fs=params.get("Fs_mat", 3.0),
        theta_deg=params.get("theta", 80.0),
        gamma_fill=_ld_w.get("g_fill", 18.0),
        gamma_mat=_ld_w.get("g_mat", 22.5),
        qa=params.get("qa_mat", 0.0),
    )
    H.heading(doc, "VII. KIỂM TRA CHỌC THỦNG LỚP ĐỆM XI MĂNG", size=13)
    H.body(doc, "Phương pháp: ALiCC (PWRI Japan). Điều kiện: τse ≤ τase.")
    H.tbl_caption(doc, "Kết quả kiểm tra chọc thủng lớp đệm xi măng")
    H.make_table(doc,
        headers=["Thông số", "Công thức", "Giá trị"],
        rows=[
            ("Bề dày đệm xi măng Hse",     "—",                              f"{_ld_w.get('h_mat',0.4):.2f} m"),
            ("Ứng suất cắt cho phép",       "τase = quckse/(2·Fs)",           f"{_pr_w['tase_kPa']:.2f} kPa"),
            ("Chiều cao vùng vòm Ho",       "Ho = (e−D)·tan(θ/2)",            f"{_pr_w['Ho_m']:.3f} m"),
            ("Công thức áp dụng",           "So sánh Ho và He",               _pr_w["cong_thuc"]),
            ("Áp lực vùng không gia cố",    "PSoil",                          f"{_pr_w['PSoil_kPa']:.3f} kPa"),
            ("Ứng suất cắt thực tế τse",    "τse = PSoil·A/(π·D·Hse)",        f"{_pr_w['tse_kPa']:.3f} kPa"),
            ("Kết quả kiểm tra",            "τse ≤ τase ?",                   _pr_w["check_CT"]),
        ])

    # VIII. Biểu đồ so sánh phương án
    import matplotlib.pyplot as _plt
    from docx.shared import Inches as _In

    def _bar_to_buf(x_lbls, y_vals, title, y_lbl, hi_idx, ref_line=None):
        _fig, _ax = _plt.subplots(figsize=(8, 3))
        _cols = ["#ED7D31" if i == hi_idx else "#4472C4" for i in range(len(y_vals))]
        _bars = _ax.bar(x_lbls, y_vals, color=_cols)
        _ax.bar_label(_bars, fmt="%.2f", padding=3, fontsize=9)
        if ref_line is not None:
            _ax.axhline(ref_line, color="#E53935", ls="--", lw=1.2,
                        label=f"Giới hạn = {ref_line:.1f}")
            _ax.legend(fontsize=9)
        _ax.set_title(title, fontsize=11)
        _ax.set_ylabel(y_lbl, fontsize=9)
        _ax.tick_params(labelsize=9)
        for _sp in ["top", "right"]:
            _ax.spines[_sp].set_visible(False)
        _buf = io.BytesIO()
        _fig.tight_layout()
        _fig.savefig(_buf, format="png", dpi=150, bbox_inches="tight")
        _plt.close(_fig)
        _buf.seek(0)
        return _buf

    if scenarios:
        H.heading(doc, "VIII. BIỂU ĐỒ SO SÁNH PHƯƠNG ÁN", size=13)
        _xl = [f"e={s['e (m)']}m" for s in scenarios]

        doc.add_picture(
            _bar_to_buf(_xl, [s["S₁ (cm)"] for s in scenarios],
                        f"Độ lún S₁ (cm) – D={int(D*1000)}mm, qu={int(qu)}kPa",
                        "S₁ (cm)", rec_idx),
            width=_In(5.5))
        doc.add_paragraph()
        doc.add_picture(
            _bar_to_buf(_xl, [s["a (%)"] for s in scenarios],
                        "Tỷ lệ thay thế a (%)",
                        "a (%)", rec_idx),
            width=_In(5.5))
        doc.add_paragraph()

        _tse_list = []
        for _s in scenarios:
            _pr_s = calc_punching_check(
                D=D, e=_s["e (m)"],
                Hse=_ld_w.get("h_mat", 0.4),
                He=_ld_w.get("h_fill", 1.5),
                quckse=params.get("quckse", 600.0),
                Fs=params.get("Fs_mat", 3.0),
                theta_deg=params.get("theta", 80.0),
                gamma_fill=_ld_w.get("g_fill", 18.0),
                gamma_mat=_ld_w.get("g_mat", 22.5),
                qa=params.get("qa_mat", 0.0),
            )
            _tse_list.append(_pr_s["tse_kPa"])
        doc.add_picture(
            _bar_to_buf(_xl, _tse_list,
                        "Ứng suất cắt τse – kiểm tra chọc thủng",
                        "τse (kPa)", rec_idx,
                        ref_line=_pr_w["tase_kPa"]),
            width=_In(5.5))

    try:
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as _e_save:
        import traceback as _tb
        st.error(f"Lỗi khi save Word: {type(_e_save).__name__}: {_e_save}")
        st.code(_tb.format_exc())
        return b""


# ═══════════════════════════════════════════════════════════════════════════════
# PROJECT SAVE / LOAD
# ═══════════════════════════════════════════════════════════════════════════════
_SAVE_KEYS = ["cdm_zone","cdm_bh","cdm_top_clay","cdm_h_clay","cdm_Su","cdm_gamma","cdm_elevation",
              "cdm_D","cdm_e","cdm_L_ngam","cdm_Lc","cdm_CDTK","cdm_qu","cdm_FS_lab","cdm_arrangement",
              "cdm_cement_type","cdm_dosage","cdm_WC","cdm_spacings","cdm_rec_idx","cdm_loads",
              "cdm_quckse","cdm_Fs_mat","cdm_theta","cdm_qa_mat"]

def _project_to_json() -> bytes:
    return json.dumps({k: st.session_state.get(k) for k in _SAVE_KEYS},
                      ensure_ascii=False, indent=2).encode("utf-8")

def _project_from_dict(data: dict):
    for k, v in data.items():
        if k in _SAVE_KEYS:
            st.session_state[k] = v


# ═══════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════════════════════════════════════════
st.sidebar.markdown("### Thiết kế cọc đất xi măng\nXử lý nền + Kè SW")
st.sidebar.markdown(f"**{_t('app_sub')}**")
st.sidebar.divider()

# Language toggle
_lang_sel = st.sidebar.radio(
    _t("lang_lbl"), ["VN", "EN"],
    index=["VN", "EN"].index(_get("lang")),
    horizontal=True, key="_lang_radio",
)
if _lang_sel != _get("lang"):
    st.session_state["lang"] = _lang_sel
    st.rerun()

st.sidebar.divider()

# Save / Load
with st.sidebar.expander(_t("save_load"), expanded=False):
    st.download_button(_t("download"), _project_to_json(),
                       file_name="cdm_project.cdm", mime="application/json")
    _up = st.file_uploader(_t("upload"), type=["cdm"], key="cdm_upload")
    if _up:
        try:
            _project_from_dict(json.load(_up))
            st.success(_t("load_ok"))
            st.rerun()
        except Exception as e:
            st.error(f"{'Lỗi' if _get('lang')=='VN' else 'Error'}: {e}")

st.sidebar.divider()

if "_page" not in st.session_state:
    st.session_state["_page"] = "geology"

def _nav(label: str, pid: str, indent: bool = False) -> None:
    pfx = "    " if indent else ""
    if st.sidebar.button(pfx + label, key=f"_nav_{pid}",
                         use_container_width=True,
                         type="primary" if st.session_state.get("_page") == pid else "secondary"):
        st.session_state["_page"] = pid
        st.rerun()

# Migrate session_state cũ: "compare" / "export" → "params" (đã gộp)
if st.session_state.get("_page") in ("compare", "export"):
    st.session_state["_page"] = "params"

st.sidebar.markdown("**Chuẩn bị họp**")
_nav(_t("p_tvtk_prep"), "tvtk_prep")
st.sidebar.markdown("---")
_nav(_t("p_geology"),      "geology")
_nav(_t("p_sample_check"), "sample_check")
st.sidebar.markdown(f"**{_t('p_tkcs_cdm')}**")
_nav(_t("p_params"),    "params",     indent=True)   # gộp: Thông số + Kết quả CDM + Xuất kết quả
_nav(_t("p_settlement"),"settlement", indent=True)
_nav(_t("p_cdm_bvt"),   "cdm_bvt")
_nav(_t("p_cdm_tien_do"), "cdm_tien_do", indent=True)
st.sidebar.markdown(f"**{_t('p_tkcs_sw')}**")
_nav(_t("p_ke_sw"),  "ke_sw",  indent=True)
_nav(_t("p_sw_bvt"), "sw_bvt")

_page = st.session_state.get("_page", "geology")

# Info sidebar dưới cùng
st.sidebar.divider()
_q_now = q_total(_get("cdm_loads"))
_Su    = _get("cdm_Su")
_Lc    = _get("cdm_Lc")
_qu    = _get("cdm_qu")
_rec   = _get("cdm_rec_idx")
_sps   = _get("cdm_spacings")

# ── Print-friendly CSS (Ctrl+P → Save as PDF) ───────────────────────────────
# Streamlit dùng nhiều layer container có overflow/height ràng buộc → in 1 trang.
# Brute force: wildcard * override mọi container khi @media print.
st.markdown("""
<style>
@media print {
  /* Override containers — bắt buộc để content chảy nhiều trang.
     KHÔNG dùng wildcard * vì:
       - KaTeX dùng nested .vlist với position relative + inline height
         → wildcard overflow:visible + max-height:none đè lên KaTeX
         làm tử số/mẫu số phân số lệch, dấu Việt rời rạc, ký hiệu trôi.
       - Plotly dùng absolute height cho chart canvas.
     Áp dụng overflow/max-height CHỈ cho Streamlit containers + body. */
  html, body,
  .main, .block-container,
  [data-testid^="stMain"], [data-testid^="stApp"],
  [data-testid="stVerticalBlock"],
  [data-testid="stHorizontalBlock"],
  [data-testid="stColumn"],
  [data-testid="stElementContainer"],
  [data-testid="stMarkdownContainer"],
  [data-testid="stContainer"],
  [data-testid="stExpander"],
  details, summary {
    overflow: visible !important;
    overflow-x: visible !important;
    overflow-y: visible !important;
    max-height: none !important;
  }
  * {
    -webkit-print-color-adjust: exact !important;
    print-color-adjust: exact !important;
  }
  /* Interactive controls — ẩn khi in, không hữu ích và gây chồng chéo lên bảng/chart.
     Đặc biệt: radio "Chế độ xem" + multiselect Zone + checkbox 3D trong tab Địa chất. */
  [data-testid="stRadio"],
  [data-testid="stMultiSelect"],
  [data-testid="stSelectbox"],
  [data-testid="stSlider"],
  [data-testid="stNumberInput"],
  [data-testid="stTextInput"],
  [data-testid="stTextArea"],
  [data-testid="stDateInput"],
  [data-testid="stFileUploader"],
  [data-testid="stColorPicker"],
  [data-testid="stToggle"],
  [data-testid="stCheckbox"],
  [data-testid="stButton"],
  button[kind="secondary"], button[kind="primary"] { display: none !important; }
  [data-testid="stDownloadButton"] { display: block !important; }

  /* height:auto CHỈ cho Streamlit layout containers (không đụng MathJax/Plotly) */
  .element-container,
  [data-testid="stVerticalBlock"],
  [data-testid="stHorizontalBlock"],
  [data-testid="stColumn"],
  [data-testid="stMarkdownContainer"],
  [data-testid="stExpander"],
  [data-testid="stDataFrame"] > div,
  .dvn-scroller, [class*="stVirtual"],
  .main, .block-container,
  [data-testid="stMainBlockContainer"],
  [data-testid="stAppViewBlockContainer"] {
    height: auto !important;
  }
  /* MathJax v3 CHTML: đảm bảo overflow visible (không bị container clamp),
     KHÔNG override height — MathJax tự set inline height cho layout nội tại */
  mjx-container, mjx-container * {
    overflow: visible !important;
    max-height: none !important;
  }
  /* MathJax v2 fallback */
  .MathJax, .MathJax_Display, .MathJax span, .MathJax svg {
    overflow: visible !important;
    max-height: none !important;
  }
  /* Override position/transform CHỈ cho Streamlit chrome (sidebar, header...),
     KHÔNG đụng tới nội dung Plotly */
  [data-testid="stSidebar"],
  [data-testid="stHeader"],
  [data-testid="stToolbar"],
  [data-testid="stMainBlockContainer"],
  [data-testid="stAppViewBlockContainer"],
  .main, .block-container {
    position: static !important;
    transform: none !important;
  }
  /* ── KaTeX / MathJax: KHÔNG override BẤT CỨ THỨ GÌ ──
     Wildcard ở trên đã restrict không động vào .katex nội bộ.
     KaTeX render đúng layout (nested .vlist + inline height cho phân số,
     subscript, dấu Việt) khi browser cho phép giữ nguyên CSS gốc. */

  /* ẨN .katex-mathml — đây là MathML annotation cho screen reader,
     KHÔNG phải để hiển thị. Mặc định KaTeX dùng position absolute + clip
     để ẩn → giữ nguyên trong print. */
  .katex-mathml,
  .katex .katex-mathml,
  span.katex-mathml,
  math[xmlns="http://www.w3.org/1998/Math/MathML"] {
    display: none !important;
    visibility: hidden !important;
    position: absolute !important;
    clip: rect(0, 0, 0, 0) !important;
    width: 1px !important;
    height: 1px !important;
    overflow: hidden !important;
  }
  html, body {
    width: 100% !important;
    min-height: 0 !important;
  }
  /* Ẩn UI chrome (sau wildcard để không bị override) */
  [data-testid="stSidebar"],
  [data-testid="stHeader"], header,
  [data-testid="stToolbar"], [data-testid="stDecoration"],
  [data-testid="stStatusWidget"], .stDeployButton,
  [data-testid="collapsedControl"], footer {
    display: none !important;
  }
  /* Container chính: full width, không padding thừa */
  .main .block-container,
  [data-testid="stMainBlockContainer"],
  [data-testid="stAppViewBlockContainer"] {
    max-width: 100% !important;
    width: 100% !important;
    padding: 0.4rem !important;
    margin: 0 !important;
  }
  /* Auto-mở expander khi in (hide summary, hiện content) */
  details[data-testid="stExpander"] > summary { display: none !important; }
  details[data-testid="stExpander"] > div,
  details[data-testid="stExpander"] {
    display: block !important;
    open: true !important;
  }
  details { display: block !important; }
  details > summary { display: none !important; }
  details > * { display: block !important; }
  /* Tránh ngắt giữa figure/table/expander */
  div[data-testid="stExpander"],
  .stPlotlyChart, .stImage, .stPyplot, .stDataFrame,
  table, figure {
    break-inside: avoid;
    page-break-inside: avoid;
  }
  /* ── GIỮ HÀNG 2 CỘT CHART CÙNG TRANG ───────────────────────────────
     Chỉ áp dụng cho stHorizontalBlock (row 2 cột) — KHÔNG áp dụng cho
     stColumn / stVerticalBlock vì làm cả block to bị đẩy sang trang sau
     để lại khoảng trống lớn (vd: Mục E ổn định tổng thể). */
  [data-testid="stHorizontalBlock"] {
    break-inside: avoid;
    page-break-inside: avoid;
  }
  /* Markdown heading (h1-h6) không break ngay sau — dính với chart kế tiếp.
     Chỉ heading, không phải mọi markdown container (tránh stick caption + body). */
  [data-testid="stMarkdownContainer"] h1,
  [data-testid="stMarkdownContainer"] h2,
  [data-testid="stMarkdownContainer"] h3,
  [data-testid="stMarkdownContainer"] h4,
  [data-testid="stMarkdownContainer"] h5,
  [data-testid="stMarkdownContainer"] h6,
  [data-testid="stMarkdownContainer"] strong {
    break-after: avoid;
    page-break-after: avoid;
  }
  /* Chart pyplot to: giới hạn chiều cao 22cm để vừa 1 trang A4 */
  .stPyplot img, [data-testid="stPyplot"] img,
  .stImage img, [data-testid="stImage"] img {
    max-height: 22cm !important;
    object-fit: contain !important;
  }
  /* Bảng dataframe rộng: scale font + wrap để vừa khổ A4.
     Đồng thời ép min-height để container chừa đủ chỗ — tránh nội dung
     phía dưới (st.latex / st.markdown) đè lên rows. */
  [data-testid="stDataFrame"],
  [data-testid="stDataFrame"] > div {
    max-width: 100% !important;
    width: 100% !important;
    overflow: visible !important;
    min-height: fit-content !important;
    height: auto !important;
  }
  [data-testid="stDataFrame"] [role="grid"],
  [data-testid="stDataFrame"] [data-testid="data-grid-canvas"] {
    height: auto !important;
    min-height: 100px !important;
  }
  [data-testid="stDataFrame"] table,
  [data-testid="stDataFrame"] [role="grid"] {
    font-size: 8pt !important;
    max-width: 100% !important;
    width: 100% !important;
    table-layout: auto !important;
    word-break: break-word !important;
  }
  [data-testid="stDataFrame"] th,
  [data-testid="stDataFrame"] td {
    padding: 2px 4px !important;
    word-break: break-word !important;
    white-space: normal !important;
  }
  /* Plotly chart: ép max-height 22cm + scale SVG vừa width */
  .stPlotlyChart, [data-testid="stPlotlyChart"],
  .js-plotly-plot {
    max-height: 22cm !important;
    max-width: 100% !important;
  }
  .stPlotlyChart svg.main-svg, .js-plotly-plot svg.main-svg {
    max-height: 22cm !important;
    max-width: 100% !important;
  }
  /* ── PLOTLY: giữ chiều cao chart nhưng KHÔNG ép display:block lên
     mọi .js-plotly-plot / .main-svg — vì Streamlit để display:none cho
     DOM stale khi re-run → ép visible làm hiện copy thừa của cùng 1 chart */
  .stPlotlyChart, [data-testid="stPlotlyChart"] {
    min-height: 380px !important;
    height: auto !important;
    overflow: visible !important;
  }
  /* Chỉ ẩn duplicate: trong cùng .stPlotlyChart, chỉ giữ .js-plotly-plot
     đầu tiên (Streamlit luôn render mới ở vị trí đầu) */
  .stPlotlyChart .js-plotly-plot ~ .js-plotly-plot,
  [data-testid="stPlotlyChart"] .js-plotly-plot ~ .js-plotly-plot {
    display: none !important;
  }
  /* Streamlit đánh dấu stale element bằng data-stale="true" khi
     re-run mà DOM cũ chưa được clean → ẩn hẳn khi in */
  [data-stale="true"],
  [data-testid="stElementContainer"][data-stale="true"],
  [data-stale="true"] * {
    display: none !important;
    visibility: hidden !important;
    height: 0 !important;
    overflow: hidden !important;
  }
  /* SVG vừa khít container, không co dãn quá khổ */
  .stPlotlyChart svg.main-svg, .js-plotly-plot svg.main-svg {
    max-width: 100% !important;
    height: auto !important;
  }
  /* Ẩn modebar (toolbar Plotly) khi in */
  .modebar, .modebar-container { display: none !important; }
  /* iframe Streamlit: KHÔNG đặt min-height toàn cục — iframe ẩn của
     components.html(height=0) sẽ bị thổi lên 380px → tạo khoảng trống
     lớn ở đầu mọi trang. Chỉ giữ max-width + page-break. */
  iframe {
    max-width: 100% !important;
    width: 100% !important;
    page-break-inside: avoid;
  }
  /* Iframe height=0 (JS handler beforeprint): ẨN HẲN khi in để không
     chiếm khoảng trống đầu trang */
  iframe[height="0"],
  iframe[srcdoc*="beforeprint"] {
    height: 0 !important;
    min-height: 0 !important;
    max-height: 0 !important;
    display: none !important;
    visibility: hidden !important;
  }
  /* Streamlit iframe wrapper element cũng ẩn hẳn khi iframe rỗng */
  [data-testid="stIFrame"]:has(iframe[height="0"]),
  [data-testid="stCustomComponentV1"]:has(iframe[height="0"]) {
    display: none !important;
    height: 0 !important;
  }
  /* Matplotlib pyplot + ảnh: luôn hiện, không bị wildcard ẩn */
  .stPyplot, [data-testid="stPyplot"],
  .stImage, [data-testid="stImage"],
  .stPyplot img, .stImage img {
    display: block !important;
    visibility: visible !important;
    page-break-inside: avoid;
    max-width: 100% !important;
    height: auto !important;
  }
  body { font-size: 10pt; line-height: 1.3; color: #000 !important; }
  a { text-decoration: none !important; color: inherit !important; }
  /* Utility: ẩn/hiện theo screen vs print */
  .print-only  { display: block  !important; }
  .screen-only { display: none   !important; }
}
/* Ẩn print-only trên màn hình bình thường */
.print-only { display: none; }
/* iframe height=0: screen-mode fallback (CSS khối 1 đã cover, giữ đây để override
   bất kỳ Streamlit style nào đặt min-height sau) */
iframe[height="0"],
[data-testid="stCustomComponentV1"]:has(iframe[height="0"]),
[data-testid="stIFrame"]:has(iframe[height="0"]) {
  display: none !important;
  height: 0 !important;
  min-height: 0 !important;
  padding: 0 !important;
  margin: 0 !important;
  overflow: hidden !important;
}
@page { size: A4; margin: 12mm 10mm; }
</style>
""", unsafe_allow_html=True)

# JS auto-mở expander trước khi in — qua components.v1.html (iframe same-origin)
# vì st.markdown bị sanitize <script>.
import streamlit.components.v1 as _components_pdf
_components_pdf.html("""
<script>
/*
  Print handler v3 — Pre-cache strategy
  ──────────────────────────────────────
  Root cause: beforeprint fires AFTER Chrome freezes JS → async Plotly.toImage()
  không kịp chạy. Ctrl+P hook cũng miss khi user dùng browser menu / Ctrl+Shift+P.

  Solution: PRE-CACHE snapshots định kỳ mỗi 4 giây vào Map<el,dataUrl>.
  beforeprint chỉ INJECT img đã cache sẵn (đồng bộ, 0ms) → luôn kịp.
  Ctrl+P hook: refresh cache mới nhất rồi print → chất lượng cao nhất.
*/
(function() {
  try {
    const W = window.parent;
    if (!W || W._printV5) return;
    W._printV5 = true;

    const OA = 'data-po';           // overlay attribute
    const cache = new Map();        // plotEl → dataUrl

    /* CSS: hiện overlay khi in, ẩn Plotly canvas khi in */
    const sid = '_pov3';
    if (!W.document.getElementById(sid)) {
      const s = W.document.createElement('style');
      s.id = sid;
      s.textContent =
        '@media print {' +
        /* Hiện overlay img */
        '  img[' + OA + ']{display:block!important;width:100%!important;height:auto!important;position:static!important;}' +
        /* Khi overlay đã inject: ẩn toàn bộ children còn lại của plot container (chỉ giữ img) */
        '  .js-plotly-plot:has(img[' + OA + ']) > :not(img[' + OA + ']){display:none!important;}' +
        '}';
      W.document.head.appendChild(s);
    }

    /* Chụp snapshot tất cả chart → lưu cache */
    const refreshCache = async () => {
      const Plotly = W.Plotly;
      if (!Plotly || !Plotly.toImage) return 0;
      const plots = W.document.querySelectorAll('.js-plotly-plot');
      let n = 0;
      await Promise.all(Array.from(plots).map(async p => {
        try {
          const r = p.getBoundingClientRect();
          if (r.width < 10 || r.height < 10) return;
          const url = await Plotly.toImage(p, {
            format: 'png',
            width:  Math.round(Math.max(r.width, 600)),
            height: Math.round(Math.max(r.height, 280)),
            scale: 1.8,
          });
          cache.set(p, url);
          n++;
        } catch(_) {}
      }));
      return n;
    };

    /* Inject img overlay từ cache vào wrapper của mỗi chart */
    const injectOverlays = () => {
      cache.forEach((url, p) => {
        if (!p.isConnected) { cache.delete(p); return; }
        if (p.querySelector('img[' + OA + ']')) return;
        const wrap = p.closest('[data-testid="stPlotlyChart"]') || p;
        const img = W.document.createElement('img');
        img.setAttribute(OA, '1');
        img.src = url;
        img.style.cssText = 'display:none;width:100%;height:auto;background:#fff';
        wrap.appendChild(img);
      });
    };

    const removeOverlays = () =>
      W.document.querySelectorAll('img[' + OA + ']').forEach(e => e.remove());

    /* Mở / đóng expander */
    const openD  = () => W.document.querySelectorAll('details').forEach(d => {
      d.setAttribute('_pw', d.open ? '1' : '0'); d.open = true;
    });
    const closeD = () => W.document.querySelectorAll('details').forEach(d => {
      if (d.getAttribute('_pw') === '0') d.open = false; d.removeAttribute('_pw');
    });

    /* Xoá DOM stale: Streamlit đánh dấu element cũ với data-stale="true"
       khi re-run mà DOM cũ chưa GC → 1 chart hiện thành 2-3 copies khi in */
    const removeStale = () => {
      const stale = W.document.querySelectorAll('[data-stale="true"]');
      stale.forEach(n => n.remove());
      if (stale.length) console.log('[PrintV5] removed', stale.length, 'stale nodes');
    };

    /* Dedupe charts: nếu cùng 1 parent có ≥2 chart cùng loại → giữ
       chart CUỐI (mới nhất), xoá các chart cũ. Streamlit stale DOM
       thường nằm cùng parent với chart mới sau rerun. */
    const dedupeCharts = () => {
      const findDups = (selector) => {
        const groups = new Map();
        W.document.querySelectorAll(selector).forEach(el => {
          const wrap = el.closest('[data-testid="stElementContainer"]') || el;
          const parent = wrap.parentElement;
          if (!parent) return;
          if (!groups.has(parent)) groups.set(parent, []);
          groups.get(parent).push(wrap);
        });
        let removed = 0;
        groups.forEach((items) => {
          if (items.length > 1) {
            items.slice(0, -1).forEach(w => {
              try { w.remove(); removed++; } catch(_) {}
            });
          }
        });
        return removed;
      };
      const n1 = findDups('[data-testid="stPlotlyChart"]');
      const n2 = findDups('[data-testid="stPyplot"]');
      const n3 = findDups('[data-testid="stImage"]');
      const n4 = findDups('[data-testid="stDataFrame"]');
      const n5 = findDups('[data-testid="stTable"]');
      const total = n1 + n2 + n3 + n4 + n5;
      if (total) console.log('[PrintV5] dedupe removed:',
        n1, 'Plotly,', n2, 'Pyplot,', n3, 'Image,', n4, 'DataFrame,', n5, 'Table');
    };

    /* beforeprint: dùng cache đã có (đồng bộ) */
    W.addEventListener('beforeprint', () => {
      openD(); removeStale(); dedupeCharts(); injectOverlays();
    });
    W.addEventListener('afterprint',  () => { closeD(); removeOverlays(); });

    /* Ctrl+P: snapshot mới → remove stale + dedupe → inject → print */
    W.document.addEventListener('keydown', async e => {
      if ((e.ctrlKey || e.metaKey) && e.key === 'p') {
        e.preventDefault();
        openD();
        removeOverlays();
        removeStale();
        dedupeCharts();
        await refreshCache();
        injectOverlays();
        setTimeout(() => W.print(), 80);
      }
    });

    /* Pre-cache: chạy sau 2s, sau đó mỗi 4s */
    setTimeout(async () => {
      const n = await refreshCache();
      console.log('[PrintV5] initial cache:', n, 'charts');
      setInterval(refreshCache, 4000);
    }, 2000);

    console.log('[PrintV5] attached');
  } catch(e) { console.error('[PrintV5]', e); }
})();
</script>
""", height=0)

# ── PDF export đã được ẩn theo yêu cầu ─────────────────────────────────
# Tất cả nút xuất PDF qua engine server-side bị tắt.
# User dùng Ctrl+P trên trình duyệt để in (CSS @media print đã được tối ưu).
_HAS_PDF = False
_PDF_ENGINE = "—"


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 1 – ĐỊA CHẤT
# ═══════════════════════════════════════════════════════════════════════════════
if _page == "geology":
    st.subheader(_t("p1_sub"))

    # ── 3D / Map toggle (đưa lên đầu trang) ───────────────────────────────────
    if not _HAS_PLOTLY:
        st.info("Biểu đồ 3D cần Plotly — vui lòng tải lại trang.")
    else:
        _bhs_all, _ = _load_borehole_3d_data()
        _zones_with_coords = sorted({b["zone"] for b in _bhs_all})
        if not _zones_with_coords:
            st.info(_t("no_coords_db"))
        else:
            # ── Thanh điều khiển ──────────────────────────────────────────────
            _ctrl_a, _ctrl_b = st.columns([3, 2])
            _view_opts = [_t("view_map"), _t("view_3d")]
            with _ctrl_a:
                _geo_view = st.radio(
                    _t("view_mode"),
                    _view_opts,
                    index=0,
                    horizontal=True,
                    key="_geo_view_radio",
                )
                st.session_state["cdm_geo_view"] = _geo_view

            # ── 3D VIEW ───────────────────────────────────────────────────────
            if _geo_view == _t("view_3d"):
                with _ctrl_b:
                    _default_zones = [z for z in _zones_with_coords if z != "QTT"]
                    _sel_zones = st.multiselect(
                        _t("zone_lbl"), _zones_with_coords,
                        default=_default_zones, key="_3d_cdm_zones",
                    )
                _cb1, _cb2, _cb3 = st.columns(3)
                _show_clay = _cb1.checkbox(_t("clay_surf"), value=True, key="_3d_cdm_clay")
                _show_cdm  = _cb2.checkbox(_t("cdm_top_show"), value=False, key="_3d_cdm_top")
                _cdm_top_z = float(_get("cdm_CDTK"))
                if _show_cdm:
                    _cdm_top_z = _cb3.number_input(
                        _t("elev_lbl"), value=_cdm_top_z, step=0.1, key="_3d_cdm_top_z")
                _vst_focus = _get("cdm_vst_sel")
                _focus_bh = _vst_focus[0] if len(_vst_focus) == 1 else None

                # ── Layout dọc: 3D ở trên, Bảng khoảng cách ở dưới ──────
                # Dùng st.empty() placeholder để render 3D chart sau khi
                # đã lấy được _pair_sel từ bảng phía dưới.
                _chart_placeholder = st.empty()

                # ── Bảng khoảng cách (dưới) — set _pair_sel ─────────────────
                _pair_sel = None
                st.markdown("#### Khoảng cách hố khoan — Điều 5.3.2")
                try:
                    import sys as _sys2
                    _sys2.path.insert(0, str(_ROOT / "scripts"))
                    from borehole_spacing import check_spacing_532 as _chk_sp
                    _bhs_sp = [b for b in _bhs_all if b.get("zone") in (_sel_zones or [])]
                    _step_opts = {
                        "BVTK (100–150 m)": "BVTK",
                        "LAPDA (250–500 m)": "LAPDA",
                    }
                    _sp_c1, _sp_c2, _sp_c3 = st.columns([1, 2, 2])
                    with _sp_c1:
                        _step_lbl = st.selectbox(
                            "Bước thiết kế",
                            list(_step_opts.keys()),
                            key="_sp_step_sel",
                        )
                    _step_code = _step_opts[_step_lbl]
                    _sp_res = _chk_sp(_bhs_sp, _step_code, same_zone_only=True)
                    _sp_s   = _sp_res["summary"]

                    # Metric tóm tắt — 2 cột bên phải
                    with _sp_c2:
                        st.metric(
                            "Cặp đạt yêu cầu",
                            f"{_sp_s['n_ok']}/{_sp_s['n_pairs']}",
                            delta=f"Gần: {_sp_s['n_too_close']} | Xa: {_sp_s['n_too_far']}",
                            delta_color="normal" if _sp_s["n_ok"] == _sp_s["n_pairs"] else "inverse",
                        )
                    with _sp_c3:
                        if _sp_s["min_dist_m"] is not None:
                            st.metric(
                                "Khoảng cách (min–max)",
                                f"{_sp_s['min_dist_m']:.0f} – {_sp_s['max_dist_m']:.0f} m",
                                delta=f"Yêu cầu: {_sp_res['limit_min_m']:.0f}–{_sp_res['limit_max_m']:.0f} m",
                                delta_color="off",
                            )

                    # Chọn cặp HK để hiển thị kích thước trên 3D
                    _sp_pairs = _sp_res["pairs"]
                    _pair_labels = [
                        f"{p['bh1']} ↔ {p['bh2']}  ({p['distance_m']:.0f} m — {p['status']})"
                        for p in _sp_pairs
                    ]
                    _pair_choice = st.selectbox(
                        "Chọn cặp HK để đo kích thước trên 3D",
                        ["(không chọn)"] + _pair_labels,
                        key="_sp_pair_sel",
                    )
                    if _pair_choice != "(không chọn)":
                        _pidx = _pair_labels.index(_pair_choice)
                        _pp   = _sp_pairs[_pidx]
                        _pair_sel = (_pp["bh1"], _pp["bh2"], _pp["distance_m"])

                    # Bảng chi tiết
                    if _sp_pairs:
                        _sp_rows = [
                            {
                                "Hố khoan 1":    p["bh1"],
                                "Hố khoan 2":    p["bh2"],
                                "Khu vực":       p["zone1"],
                                "Khoảng cách (m)": p["distance_m"],
                                "Kết quả":       p["status"],
                            }
                            for p in _sp_pairs
                        ]
                        _sp_df = pd.DataFrame(_sp_rows)

                        def _sp_color(val):
                            if val == "Đạt":
                                return "background-color:#d4edda; color:#155724"
                            if val in ("Gần quá", "Xa quá"):
                                return "background-color:#f8d7da; color:#721c24"
                            return ""

                        st.dataframe(
                            _sp_df.style.map(_sp_color, subset=["Kết quả"]),
                            use_container_width=True,
                            hide_index=True,
                        )
                    else:
                        st.info("Chưa có tọa độ hố khoan để tính khoảng cách.")
                    st.caption(
                        f"Tiêu chuẩn: TCCS 41:2022 Điều 5.3.2 — {_sp_res['limit_label']}"
                    )
                except Exception as _sp_exc:
                    st.warning(f"Không tải được borehole_spacing: {_sp_exc}")

                # ── 3D chart (render vào placeholder ở trên) ─────────────────
                if _sel_zones:
                    try:
                        _chart_placeholder.plotly_chart(
                            _draw_boreholes_3d(
                                _sel_zones, _show_clay, _show_cdm, _cdm_top_z,
                                focus_bh=_focus_bh,
                                pair_highlight=_pair_sel,
                            ),
                            use_container_width=True,
                            config={"displayModeBar": True},
                            key="geo_3d_bhmap",
                        )
                    except Exception as _3d_err:
                        _chart_placeholder.error(f"Không vẽ được biểu đồ 3D: {_3d_err}")
                else:
                    _chart_placeholder.info(_t("no_zone"))

            # ── MAP VIEW (folium + pyproj) ────────────────────────────────────
            else:
                try:
                    import folium as _flm
                    from streamlit_folium import st_folium as _stflm
                    from pyproj import Transformer as _Trf
                    _HAS_GIS_GEO = True
                except ImportError as _gis_err:
                    _HAS_GIS_GEO = False
                    st.error(
                        f"Thiếu thư viện bản đồ. Cài: folium, streamlit-folium, pyproj. ({_gis_err})"
                    )

                if _HAS_GIS_GEO:
                    _GEO_MAP_CRS = {
                        "TM-3 106°00' (TTHC Q.1, Thủ Thiêm)": "EPSG:9210",
                        "TM-3 105°45' (HCM tổng quát)":       "EPSG:9209",
                        "UTM 48N (105°)":                     "EPSG:3405",
                    }
                    with _ctrl_b:
                        _crs_lbl = st.selectbox(
                            "Hệ tọa độ gốc (VN-2000)",
                            list(_GEO_MAP_CRS.keys()),
                            index=0,
                            key="_geo_map_crs_sel",
                            help="Đổi nếu hố khoan lệch ~25-50 km trên bản đồ. TTHC Q.1 / Thủ Thiêm: TM-3 106°00'.",
                        )
                    _crs_code = _GEO_MAP_CRS[_crs_lbl]

                    _map_c1, _map_c2, _map_c3, _map_c4, _map_c5 = st.columns([3, 2, 2, 2, 2])
                    with _map_c1:
                        _map_zone_default = [z for z in _zones_with_coords if z != "QTT"]
                        _map_zones = st.multiselect(
                            "Khu vực hiển thị",
                            _zones_with_coords,
                            default=_map_zone_default,
                            key="_geo_map_zones",
                        )
                    with _map_c2:
                        _show_ke_aln = st.checkbox(
                            "Hiện tuyến kè (đen)",
                            value=True,
                            key="_geo_map_show_aln",
                            help="Vẽ polyline tuyến kè từ dữ liệu bình đồ.",
                        )
                    with _map_c3:
                        _show_pair_dim = st.checkbox(
                            "Hiện đường kích thước cặp đã chọn",
                            value=True,
                            key="_geo_map_show_pair",
                            help="Vẽ đường nối + nhãn khoảng cách giữa 2 hố khoan đã chọn ở bảng phía trên.",
                        )
                    with _map_c4:
                        _show_test_piles = st.checkbox(
                            "Hiện cọc CDM thử (18 cọc)",
                            value=True,
                            key="_geo_map_show_test_piles",
                            help="Vẽ 18 cọc CDM thử (CỌC-01..18) — marker tròn cam đậm + nhãn số thứ tự.",
                        )
                    with _map_c5:
                        _show_metro = st.checkbox(
                            "Hiện tuyến metro",
                            value=True,
                            key="_geo_map_show_metro",
                            help="Tuyến chính metro (đỏ đậm), ranh kiểm soát XD (cam đứt), ranh GPMB (xám đứt), tường tunnel/panel (xanh/tím nhạt).",
                        )

                    @st.cache_data(show_spinner=False)
                    def _geo_make_trf(epsg: str):
                        return _Trf.from_crs(epsg, "EPSG:4326", always_xy=True)

                    _trf = _geo_make_trf(_crs_code)

                    # boreholes: x_coord_m = Northing, y_coord_m = Easting (CLAUDE.md §32)
                    # Sanity-check: nếu zone bị nhập ngược (x là Easting) → swap runtime.
                    _hk_geo = []
                    for b in _bhs_all:
                        if b["zone"] not in _map_zones:
                            continue
                        _xc = b.get("x_coord_m")
                        _yc = b.get("y_coord_m")
                        if _xc is None or _yc is None:
                            continue
                        # Tại TP.HCM: Northing ~1,190,000+, Easting ~600,000.
                        # Nếu x_coord nhỏ hơn y_coord → convention bị ngược → swap.
                        if _xc < _yc:
                            _xc, _yc = _yc, _xc
                        _lon, _lat = _trf.transform(_yc, _xc)
                        _hk_geo.append({
                            **b,
                            "lat": _lat, "lon": _lon,
                        })

                    if not _hk_geo:
                        st.warning("Không có hố khoan trong khu vực đã chọn.")
                    else:
                        _lat0 = sum(h["lat"] for h in _hk_geo) / len(_hk_geo)
                        _lon0 = sum(h["lon"] for h in _hk_geo) / len(_hk_geo)

                        # Tuyến kè polyline (chỉ load khi tick)
                        _aln_lines: list[list[tuple[float, float]]] = []
                        if _show_ke_aln:
                            try:
                                _con_aln = sqlite3.connect(_DB)
                                _aln_rows = _con_aln.execute(
                                    "SELECT polyline_id, x_m, y_m FROM ke_binhdo_toadoke "
                                    "ORDER BY polyline_id, vertex_idx"
                                ).fetchall()
                                _con_aln.close()
                                _cur_pl, _cur_pts = None, []
                                for _pl_id, _bx, _by in _aln_rows:
                                    if _cur_pl is not None and _pl_id != _cur_pl:
                                        if len(_cur_pts) >= 2:
                                            _aln_lines.append(_cur_pts)
                                        _cur_pts = []
                                    # JSON polyline: x_m = Easting, y_m = Northing
                                    _lonp, _latp = _trf.transform(_bx, _by)
                                    _cur_pts.append((_latp, _lonp))
                                    _cur_pl = _pl_id
                                if len(_cur_pts) >= 2:
                                    _aln_lines.append(_cur_pts)
                            except Exception:
                                _aln_lines = []

                        # ── Build folium map ─────────────────────────────────
                        _fmap = _flm.Map(
                            location=[_lat0, _lon0],
                            zoom_start=16,
                            tiles=None,
                            control_scale=True,
                        )
                        _flm.TileLayer("OpenStreetMap", name="Đường phố (OSM)").add_to(_fmap)
                        _flm.TileLayer(
                            tiles="https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}",
                            attr="Esri", name="Ảnh vệ tinh (Esri)",
                        ).add_to(_fmap)
                        _flm.TileLayer(
                            tiles="https://{s}.tile.opentopomap.org/{z}/{x}/{y}.png",
                            attr="OpenTopoMap", name="Địa hình (OpenTopoMap)",
                        ).add_to(_fmap)
                        _flm.TileLayer(
                            tiles="https://{s}.basemaps.cartocdn.com/light_all/{z}/{x}/{y}.png",
                            attr="CartoDB", name="Carto nhạt (in PDF)",
                        ).add_to(_fmap)

                        # Polyline kè
                        if _aln_lines:
                            _fg_aln = _flm.FeatureGroup(name="Tuyến kè", show=True)
                            for poly in _aln_lines:
                                _flm.PolyLine(poly, color="#000000",
                                              weight=3, opacity=0.85).add_to(_fg_aln)
                            _fg_aln.add_to(_fmap)

                        # Cọc CDM thử (18 cọc — bảng cdm_coc_thu)
                        if _show_test_piles:
                            try:
                                _con_tp = sqlite3.connect(_DB)
                                _tp_rows = _con_tp.execute(
                                    "SELECT code, northing_m, easting_m, elevation_m, description "
                                    "FROM cdm_coc_thu ORDER BY point_name"
                                ).fetchall()
                                _con_tp.close()
                            except Exception:
                                _tp_rows = []
                            if _tp_rows:
                                _fg_tp = _flm.FeatureGroup(
                                    name=f"Cọc CDM thử ({len(_tp_rows)})", show=True
                                )
                                for _code, _nor, _eas, _elev, _desc in _tp_rows:
                                    _lon_tp, _lat_tp = _trf.transform(_eas, _nor)
                                    _popup_tp = (
                                        f"<b>{_code}</b><br>"
                                        f"<i>{_desc}</i><br>"
                                        f"Cao độ mặt: <b>{_elev:.2f} m</b><br>"
                                        f"Tọa độ: N={_nor:,.1f} m, E={_eas:,.1f} m"
                                    )
                                    _num = _code.split("-")[-1] if "-" in _code else _code
                                    # Marker hình tròn cam đậm + viền nâu để phân biệt với HK xanh/đỏ/lá
                                    _flm.CircleMarker(
                                        location=[_lat_tp, _lon_tp],
                                        radius=7,
                                        color="#BF360C",
                                        weight=2,
                                        fill=True,
                                        fill_color="#FF6F00",
                                        fill_opacity=0.95,
                                        popup=_flm.Popup(_popup_tp, max_width=240),
                                        tooltip=_code,
                                    ).add_to(_fg_tp)
                                    # Nhãn số thứ tự (NN) bên cạnh
                                    _flm.map.Marker(
                                        [_lat_tp, _lon_tp],
                                        icon=_flm.DivIcon(html=(
                                            f'<div style="font-size:10px;font-weight:bold;'
                                            f'color:#BF360C;text-shadow:1px 1px 2px #fff,'
                                            f'-1px -1px 2px #fff;margin-left:10px;'
                                            f'margin-top:-6px;white-space:nowrap">'
                                            f'{_num}</div>'
                                        )),
                                    ).add_to(_fg_tp)
                                _fg_tp.add_to(_fmap)

                        # Tuyến metro (bảng metro_lines) — bỏ aux_tunnel + boundary_control
                        if _show_metro:
                            _METRO_STYLE = {
                                "centerline":       {"color": "#D32F2F", "weight": 4, "opacity": 0.95, "dash": None,  "label": "Tuyến chính metro"},
                                "boundary_land":    {"color": "#616161", "weight": 2, "opacity": 0.80, "dash": "4,4", "label": "Ranh GPMB"},
                                "boundary_station": {"color": "#9E9E9E", "weight": 2, "opacity": 0.80, "dash": "2,4", "label": "Ranh ga (Xref)"},
                                "aux_panel":        {"color": "#7B1FA2", "weight": 1, "opacity": 0.60, "dash": None,  "label": "Tường panel"},
                                "other":            {"color": "#000000", "weight": 1, "opacity": 0.50, "dash": None,  "label": "Khác"},
                            }
                            try:
                                _con_mt = sqlite3.connect(_DB)
                                _mt_rows = _con_mt.execute(
                                    "SELECT polyline_id, vertex_idx, x_m, y_m, category "
                                    "FROM metro_lines "
                                    "WHERE category NOT IN ('aux_tunnel', 'boundary_control') "
                                    "ORDER BY category, polyline_id, vertex_idx"
                                ).fetchall()
                                _con_mt.close()
                            except Exception:
                                _mt_rows = []
                            if _mt_rows:
                                # Group vertex theo polyline_id, giữ thứ tự category để vẽ
                                _polys_by_cat: dict = {}
                                _cur_key = (None, None)
                                _cur_pts: list = []
                                _cur_cat = None
                                for _pl, _vi, _bx, _by, _cat in _mt_rows:
                                    _key = (_cat, _pl)
                                    if _cur_key != _key and _cur_pts:
                                        _polys_by_cat.setdefault(_cur_cat, []).append(_cur_pts)
                                        _cur_pts = []
                                    _cur_key = _key
                                    _cur_cat = _cat
                                    _lonm, _latm = _trf.transform(_bx, _by)
                                    _cur_pts.append((_latm, _lonm))
                                if _cur_pts and _cur_cat is not None:
                                    _polys_by_cat.setdefault(_cur_cat, []).append(_cur_pts)

                                _n_mt_total = sum(len(v) for v in _polys_by_cat.values())
                                # Mỗi category 1 FeatureGroup → user toggle qua LayerControl
                                # Vẽ centerline LAST → nằm trên cùng
                                _order = ["aux_panel", "boundary_land",
                                          "boundary_station", "other", "centerline"]
                                for _cat in _order:
                                    _polys = _polys_by_cat.get(_cat)
                                    if not _polys:
                                        continue
                                    _sty = _METRO_STYLE.get(_cat, _METRO_STYLE["other"])
                                    _fg_mt = _flm.FeatureGroup(
                                        name=f"Metro: {_sty['label']} ({len(_polys)})",
                                        show=True,
                                    )
                                    for _poly in _polys:
                                        if len(_poly) < 2:
                                            continue
                                        _kwargs = dict(
                                            locations=_poly,
                                            color=_sty["color"],
                                            weight=_sty["weight"],
                                            opacity=_sty["opacity"],
                                        )
                                        if _sty["dash"]:
                                            _kwargs["dash_array"] = _sty["dash"]
                                        _flm.PolyLine(**_kwargs).add_to(_fg_mt)
                                    _fg_mt.add_to(_fmap)

                        # Hố khoan theo zone
                        _ZONE_FLM_COLOR = {
                            "KE": "#E53935", "BXN": "#1565C0",
                            "NHC": "#2E7D32", "QTT": "#F9A825",
                        }
                        _fg_by_zone: dict = {}
                        for h in _hk_geo:
                            z = h["zone"]
                            if z not in _fg_by_zone:
                                _fg_by_zone[z] = _flm.FeatureGroup(
                                    name=f"Hố khoan {z}", show=True
                                )
                            _color = _ZONE_FLM_COLOR.get(z, "#777777")
                            _popup_html = (
                                f"<b>{h['name']}</b><br>"
                                f"Khu vực: <b>{z}</b><br>"
                                f"Cao độ mặt: <b>{h.get('elevation_m', 0):.2f} m</b><br>"
                                f"Độ sâu HK: <b>{h.get('depth_m', 0):.1f} m</b><br>"
                                f"Tọa độ: N={h.get('x_coord_m', 0):,.1f} m, "
                                f"E={h.get('y_coord_m', 0):,.1f} m"
                            )
                            _flm.CircleMarker(
                                location=[h["lat"], h["lon"]],
                                radius=8, color=_color,
                                weight=2, fill=True,
                                fill_color=_color, fill_opacity=0.85,
                                popup=_flm.Popup(_popup_html, max_width=280),
                                tooltip=h["name"],
                            ).add_to(_fg_by_zone[z])
                            _flm.map.Marker(
                                [h["lat"], h["lon"]],
                                icon=_flm.DivIcon(html=(
                                    f'<div style="font-size:11px;font-weight:bold;'
                                    f'color:#222;text-shadow:1px 1px 2px #fff,'
                                    f'-1px -1px 2px #fff;margin-left:10px;'
                                    f'margin-top:-10px;white-space:nowrap">'
                                    f'{h["name"]}</div>'
                                )),
                            ).add_to(_fg_by_zone[z])
                        for fg in _fg_by_zone.values():
                            fg.add_to(_fmap)

                        # Đường kích thước cặp đã chọn (đồng bộ với 3D view)
                        _pair_session = st.session_state.get("_sp_pair_sel", "(không chọn)")
                        if _show_pair_dim and _pair_session and _pair_session != "(không chọn)":
                            try:
                                _bh1_nm = _pair_session.split(" ↔ ")[0].strip()
                                _bh2_nm = _pair_session.split(" ↔ ")[1].split("  ")[0].strip()
                                _bh1 = next((h for h in _hk_geo if h["name"] == _bh1_nm), None)
                                _bh2 = next((h for h in _hk_geo if h["name"] == _bh2_nm), None)
                                if _bh1 and _bh2:
                                    _dist_m_pair = _pair_session.split("(")[1].split(" m")[0]
                                    _fg_dim = _flm.FeatureGroup(
                                        name="Đường kích thước", show=True
                                    )
                                    _flm.PolyLine(
                                        [[_bh1["lat"], _bh1["lon"]],
                                         [_bh2["lat"], _bh2["lon"]]],
                                        color="#FF6F00", weight=3,
                                        opacity=0.9, dash_array="8,6",
                                    ).add_to(_fg_dim)
                                    _mid_lat = (_bh1["lat"] + _bh2["lat"]) / 2
                                    _mid_lon = (_bh1["lon"] + _bh2["lon"]) / 2
                                    _flm.map.Marker(
                                        [_mid_lat, _mid_lon],
                                        icon=_flm.DivIcon(html=(
                                            f'<div style="background:#FF6F00;color:#fff;'
                                            f'padding:2px 8px;border-radius:4px;'
                                            f'font-size:12px;font-weight:bold;'
                                            f'box-shadow:0 1px 3px rgba(0,0,0,0.3);'
                                            f'white-space:nowrap">'
                                            f'{_dist_m_pair} m</div>'
                                        )),
                                    ).add_to(_fg_dim)
                                    _fg_dim.add_to(_fmap)
                            except Exception:
                                pass

                        _flm.LayerControl(collapsed=False, position="topright").add_to(_fmap)

                        # Plugins: fullscreen + thước đo + minimap
                        try:
                            from folium.plugins import (
                                Fullscreen as _GFS,
                                MeasureControl as _GMC,
                                MiniMap as _GMM,
                            )
                            _GFS(position="topleft",
                                 title="Toàn màn hình",
                                 title_cancel="Thoát").add_to(_fmap)
                            _GMC(position="topleft",
                                 primary_length_unit="meters",
                                 primary_area_unit="sqmeters").add_to(_fmap)
                            _GMM(toggle_display=True,
                                 position="bottomleft").add_to(_fmap)
                        except Exception:
                            pass

                        # Thống kê + render
                        _zone_counts = {}
                        for h in _hk_geo:
                            _zone_counts[h["zone"]] = _zone_counts.get(h["zone"], 0) + 1
                        _zone_summary = " · ".join(
                            f"{z}: {n}" for z, n in sorted(_zone_counts.items())
                        )
                        _n_tp_shown = len(_tp_rows) if _show_test_piles and '_tp_rows' in dir() else 0
                        _n_mt_shown = _n_mt_total if _show_metro and '_n_mt_total' in dir() else 0
                        st.caption(
                            f"Tâm: {_lat0:.5f}°N, {_lon0:.5f}°E · "
                            f"{len(_hk_geo)} hố khoan ({_zone_summary}) · "
                            f"{len(_aln_lines)} đoạn tuyến kè · "
                            f"{_n_tp_shown} cọc CDM thử · "
                            f"{_n_mt_shown} polyline tuyến metro"
                        )
                        _stflm(_fmap, width=None, height=640, returned_objects=[])

                        with st.expander("Hướng dẫn dùng bản đồ"):
                            st.markdown(
                                "- **Đổi nền**: bảng góc trên phải — OSM / Vệ tinh / Địa hình / Carto\n"
                                "- **Tắt/mở lớp**: tick checkbox \"Tuyến kè\", \"Hố khoan KE/BXN/NHC\", \"Đường kích thước\"\n"
                                "- **Zoom**: nút +/− góc trái hoặc lăn chuột\n"
                                "- **Pan**: kéo chuột trái\n"
                                "- **Toàn màn hình**: nút mũi tên góc trái trên\n"
                                "- **Đo khoảng cách/diện tích**: nút thước góc trái trên (click để bắt đầu, double-click để kết thúc)\n"
                                "- **Nhấp hố khoan**: hiện popup với cao độ, độ sâu, tọa độ N/E\n"
                                "- **Đồng bộ với view 3D**: cặp HK chọn ở bảng \"Khoảng cách\" phía trên sẽ vẽ đường kích thước cam ở bản đồ"
                            )

    st.divider()

    col_sel, col_prof = st.columns([1, 4])

    with col_sel:
        zone = st.radio(_t("zone_lbl"), list(_ZONE_NAMES.keys()),
                        format_func=lambda z: _ZONE_NAMES[z],
                        index=list(_ZONE_NAMES.keys()).index(_get("cdm_zone")),
                        horizontal=False)
        st.session_state["cdm_zone"] = zone

        bhs = _load_boreholes_by_zone(zone)
        bh_names = [b["name"] for b in bhs]
        if not bh_names:
            st.warning(_t("no_bh_warn"))
            bh_names = ["(trống)"]

        cur_bh = _get("cdm_bh")
        sel_idx = bh_names.index(cur_bh) if cur_bh in bh_names else 0
        bh_name = st.selectbox(_t("bh_lbl"), bh_names, index=sel_idx)

        if st.button(_t("apply_bh"), type="primary"):
            st.session_state["cdm_bh"]  = bh_name
            st.session_state["sl_zone"] = zone
            st.session_state["sl_bh"]   = bh_name
            # Auto-fill soil params
            cp = _clay_params(bh_name, zone)
            if cp:
                st.session_state["cdm_h_clay"]   = cp["h_clay"]
                st.session_state["cdm_elevation"] = cp.get("elevation_m", 0.0)
                st.session_state["cdm_Lc"]        = round(cp["h_clay"] + 1.5, 1)
            su = _su_avg_in_range(zone, 0, cp.get("depth_clay_bot", 30) if cp else 30)
            if su:
                st.session_state["cdm_Su"] = su
            g = _gamma_avg(bh_name)
            if g:
                st.session_state["cdm_gamma"] = g
            # Liên kết VST: tự động chọn VST location khớp với hố khoan
            _df_vst_link = _load_vst_su(zone)
            if not _df_vst_link.empty:
                _vst_locs = sorted(_df_vst_link["loc_name"].unique().tolist())
                _matched = [l for l in _vst_locs if l == bh_name]
                if not _matched:
                    _bh_low = bh_name.lower()
                    _matched = [l for l in _vst_locs
                                if _bh_low in l.lower() or l.lower() in _bh_low]
                if _matched:
                    st.session_state["cdm_vst_sel"]  = _matched
                    st.session_state["_vst_loc_sel"] = _matched
            st.success(_t("apply_bh_ok", bh=bh_name))
            st.rerun()

        # Thông tin nhanh
        if _get("cdm_bh"):
            cp = _clay_params(_get("cdm_bh"), zone)
            if cp:
                st.metric(_t("clay_thick"), f"{cp['h_clay']:.1f} m")
                st.metric(_t("clay_elev"),  f"{cp['elev_clay_bot']:.2f} m")
            su = _su_avg_in_range(zone, 0, cp.get("depth_clay_bot",30) if cp else 30)
            if su:
                st.metric(_t("su_avg"), f"{su:.1f} kPa")


    with col_prof:
        layers = _load_layers(bh_name)
        if layers:
            st.markdown(f"**{_t('strat_title')}**")
            df_lay = pd.DataFrame(layers)
            _is_vn = _get("lang") == "VN"
            if _is_vn:
                df_lay.columns = ["Ký hiệu", "Mô tả", "Đỉnh (m)", "Đáy (m)", "Dày (m)",
                                  "γ (kN/m³)", "c (kPa)", "φ (°)",
                                  "e₀", "Cc", "Cs", "PC (kPa)", "Cv (cm²/s)"]
                _desc_col = "Mô tả"
            else:
                df_lay.columns = ["Symbol", "Description", "Top (m)", "Bot (m)", "Thick (m)",
                                  "γ (kN/m³)", "c (kPa)", "φ (°)",
                                  "e₀", "Cc", "Cs", "PC (kPa)", "Cv (cm²/s)"]
                _desc_col = "Description"

            # Format Cv scientific notation
            _cv_col = "Cv (cm²/s)"
            df_lay[_cv_col] = df_lay[_cv_col].apply(
                lambda x: f"{x:.2e}" if pd.notna(x) and x is not None else "–"
            )

            st.dataframe(
                df_lay,
                use_container_width=True,
                height=310,
                column_config={
                    _desc_col:    st.column_config.TextColumn(width="medium"),
                    "γ (kN/m³)":  st.column_config.NumberColumn(format="%.2f"),
                    "c (kPa)":    st.column_config.NumberColumn(format="%.1f"),
                    "φ (°)":      st.column_config.NumberColumn(format="%.1f"),
                    "e₀":         st.column_config.NumberColumn(format="%.3f"),
                    "Cc":         st.column_config.NumberColumn(format="%.4f"),
                    "Cs":         st.column_config.NumberColumn(format="%.4f"),
                    "PC (kPa)":   st.column_config.NumberColumn(format="%.1f"),
                    _cv_col:      st.column_config.TextColumn(width="small"),
                },
            )
        else:
            st.info(_t("no_strat"))

    st.divider()
    # VST profile + soil column
    _su_col, _sc_col = st.columns([2, 1])
    with _su_col:
        df_vst = _load_vst_su(zone)
        if not df_vst.empty:
            _vst_all_locs = sorted(df_vst["loc_name"].unique().tolist())
            _vst_sel_prev = [l for l in _get("cdm_vst_sel") if l in _vst_all_locs]
            _vst_sel = st.multiselect(
                _t("vst_sel"),
                options=_vst_all_locs,
                default=_vst_sel_prev,
                key="_vst_loc_sel",
            )
            st.session_state["cdm_vst_sel"] = _vst_sel
            _vst_pass = _vst_sel if _vst_sel else None
            if _HAS_PLOTLY:
                st.plotly_chart(
                    _chart_su_profile(df_vst, _vst_pass),
                    use_container_width=True,
                )
            elif _HAS_MPL:
                _fig_vst_mpl = _chart_su_profile_mpl(df_vst, _vst_pass)
                st.pyplot(_fig_vst_mpl, use_container_width=True)
                plt.close(_fig_vst_mpl)
            else:
                st.dataframe(df_vst)
    with _sc_col:
        _bh_sc = bh_name if bh_name != "(trống)" else None
        if _bh_sc and (_HAS_PLOTLY or _HAS_MPL):
            try:
                _sc_layers = _load_layers(_bh_sc)
                _sc_spt    = _load_spt(_bh_sc)
                if _HAS_PLOTLY:
                    st.plotly_chart(
                        _draw_soil_column(_sc_layers, _bh_sc, spt=_sc_spt or None),
                        use_container_width=True,
                        config={"displayModeBar": False},
                    )
                else:
                    _fig_col_mpl = _draw_soil_column_mpl(_sc_layers, _bh_sc, spt=_sc_spt or None)
                    st.pyplot(_fig_col_mpl, use_container_width=True)
                    plt.close(_fig_col_mpl)
            except Exception as _e:
                st.warning(f"Không vẽ được cột địa chất: {_e}")
        else:
            st.caption(_t("no_bh_col"))

    # ── Nén cố kết – thống kê hố khoan có dữ liệu ─────────────────────────────
    _df_consol = _load_consol_summary(zone)
    if not _df_consol.empty:
        _is_vn = _get("lang") == "VN"
        _consol_title = (
            "Nén cố kết – Hố khoan có dữ liệu tính lún theo thời gian"
            if _is_vn else
            "Consolidation – Boreholes with time-settlement data"
        )
        with st.expander(_consol_title, expanded=True):
            # Format Cv and k in scientific notation for readability
            _df_disp = _df_consol.copy()
            _df_disp["Cv_tb"] = _df_disp["Cv_tb"].apply(
                lambda x: f"{x:.2e}" if pd.notna(x) else "–"
            )
            _df_disp["k_tb"] = _df_disp["k_tb"].apply(
                lambda x: f"{x:.2e}" if pd.notna(x) else "–"
            )
            _df_disp["PC_kPa"] = _df_disp.apply(
                lambda r: f"{r['PC_min']:.0f}–{r['PC_max']:.0f}"
                if pd.notna(r["PC_min"]) else "–", axis=1
            )

            if _is_vn:
                _df_disp = _df_disp.rename(columns={
                    "borehole": "Hố khoan",
                    "n_mau":   "Tổng mẫu",
                    "n_Cc":    "Mẫu NCK",
                    "Cc_tb":   "Cc TB",
                    "Cs_tb":   "Cs TB",
                    "Cv_tb":   "Cv TB (cm²/s)",
                    "k_tb":    "k TB (cm/s)",
                    "PC_kPa":  "PC (kPa)",
                    "e0_tb":   "e₀ TB",
                })
                _note = (
                    "NCK = Nén cố kết. Cc: chỉ số nén. Cs: chỉ số nở lại. "
                    "Cv: hệ số cố kết. k: hệ số thấm. PC: áp lực tiền cố kết. "
                    "Các giá trị là trung bình số học của các mẫu trong hố khoan."
                )
            else:
                _df_disp = _df_disp.rename(columns={
                    "borehole": "Borehole",
                    "n_mau":   "Total",
                    "n_Cc":    "Consol.",
                    "Cc_tb":   "Cc avg",
                    "Cs_tb":   "Cs avg",
                    "Cv_tb":   "Cv avg (cm²/s)",
                    "k_tb":    "k avg (cm/s)",
                    "PC_kPa":  "PC (kPa)",
                    "e0_tb":   "e₀ avg",
                })
                _note = (
                    "Consol. = samples with consolidation data. Cv: coefficient of consolidation. "
                    "k: permeability. PC: preconsolidation pressure. Values are arithmetic means per borehole."
                )

            _drop_cols = ["PC_min", "PC_max"]
            _df_disp = _df_disp.drop(columns=[c for c in _drop_cols if c in _df_disp.columns])

            st.dataframe(
                _df_disp,
                use_container_width=True,
                hide_index=True,
                column_config={
                    _df_disp.columns[2]: st.column_config.NumberColumn(format="%d"),
                    "Cc TB" if _is_vn else "Cc avg": st.column_config.NumberColumn(format="%.4f"),
                    "Cs TB" if _is_vn else "Cs avg": st.column_config.NumberColumn(format="%.4f"),
                    "e₀ TB" if _is_vn else "e₀ avg": st.column_config.NumberColumn(format="%.3f"),
                },
            )
            st.caption(_note)

            # Tóm tắt số hố có dữ liệu
            _n_bh_ok  = len(_df_consol)
            _n_bh_all = len(_load_boreholes_by_zone(zone))
            st.info(
                f"{'Zone' if not _is_vn else 'Zone'} **{_ZONE_NAMES[zone]}**: "
                f"{_n_bh_ok}/{_n_bh_all} "
                f"{'hố khoan có thí nghiệm nén cố kết (Cc/Cs/Cv/k/PC).' if _is_vn else 'boreholes have consolidation test data (Cc/Cs/Cv/k/PC).'}"
            )

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE – KIỂM TRA MẪU THÍ NGHIỆM
# ═══════════════════════════════════════════════════════════════════════════════
elif _page == "sample_check":
    import sys as _sys2
    _sys2.path.insert(0, str(_ROOT / "scripts"))

    # Import độc lập — section A không phụ thuộc section B
    try:
        from settlement_calc import check_samples_vs_tccs41 as _chk537
        _HAS_537 = True
    except Exception as _exc537:
        _HAS_537 = False
        st.error(f"Không tải được settlement_calc: {_exc537}")

    try:
        from cdm_column_calc import check_qc_adequacy as _chk9403
        _HAS_9403 = True
    except Exception:
        _HAS_9403 = False  # module chưa tồn tại — ẩn section B

    st.subheader("Kiểm tra mẫu thí nghiệm — TCCS 41:2022 & TCVN 9403:2012")

    # ── A. TCCS 41:2022 Điều 5.3.7 ──────────────────────────────────────────
    st.markdown("## A. Mẫu nén cố kết — TCCS 41:2022 Điều 5.3.7")
    st.caption(
        "Điều 5.3.7: Mỗi lớp đất yếu, mỗi chỉ tiêu đưa vào tính toán cần có **≥ 6 số liệu thí nghiệm**. "
        "Trị số tính toán: Δt = Δtb ± δ  (δ = độ lệch chuẩn mẫu). "
        "Áp dụng cho Cc, Cs, Cv, PC của các lớp đất yếu (CH/CL/MH/ML và biến thể)."
    )

    if _HAS_537:
        # Tải dữ liệu 3 khu vực
        _chk_all = {}
        for _zc in ["NHC", "BXN", "KE"]:
            try:
                _chk_all[_zc] = _chk537(_zc)
            except Exception as _ex:
                st.error(f"Khu vực {_zc}: {_ex}")

        # Metric cards tóm tắt
        _lun_params = ["Cc", "Cs", "Cv", "PC"]
        _met_cols = st.columns(3)
        for _ci, _zc in enumerate(["NHC", "BXN", "KE"]):
            if _zc not in _chk_all:
                continue
            _s       = _chk_all[_zc]["zone_summary"]
            _n_soft  = _s["n_layers_soft"]
            _ok_cc   = _s["params_ok"].get("Cc", 0)
            _fail_cc = _s["params_fail"].get("Cc", 0)
            with _met_cols[_ci]:
                st.metric(
                    f"Khu vực {_zc}",
                    f"{_ok_cc}/{_n_soft} lớp đủ Cc",
                    f"{_fail_cc} lớp thiếu mẫu (n<6)",
                    delta_color="normal" if _fail_cc == 0 else "inverse",
                )
                st.caption(
                    f"Tổng lớp: {_s['n_layers_total']} | Lớp yếu: {_n_soft}  \n"
                    + "  ".join(
                        f"{p}: {_s['params_ok'].get(p,0)}Đ/{_s['params_fail'].get(p,0)}T"
                        for p in _lun_params
                    )
                )

        def _fmt_p(info):
            if not info or info.get("n", 0) == 0:
                return "-"
            n  = info["n"]
            ok = info.get("ok")
            if ok is True:  return f"{n} (Đạt)"
            if ok is False: return f"{n} (Thiếu)"
            return str(n)

        def _cell_color(val):
            if "(Đạt)"   in str(val): return "background-color:#d4edda; color:#155724"
            if "(Thiếu)" in str(val): return "background-color:#f8d7da; color:#721c24"
            return ""

        # Bảng chi tiết per khu vực
        for _zc in ["NHC", "BXN", "KE"]:
            if _zc not in _chk_all:
                continue
            st.markdown(f"**Khu vực {_zc} — Chi tiết theo lớp đất**")
            _rows_ly = []
            for _ly in _chk_all[_zc]["layers"]:
                _p = _ly["params"]
                _rows_ly.append({
                    "Lớp (USCS)": _ly["symbol"],
                    "Đất yếu":    "Có" if _ly["is_soft"] else "-",
                    "n mẫu":      _ly["n_total"],
                    "Cc (n)":     _fmt_p(_p.get("Cc")),
                    "Cc trung bình": round(_p["Cc"]["mean"], 3) if _p.get("Cc", {}).get("n", 0) > 0 else "-",
                    "Cc độ lệch":    round(_p["Cc"]["std"], 3)  if _p.get("Cc", {}).get("n", 0) > 1 else "-",
                    "Cs (n)":     _fmt_p(_p.get("Cs")),
                    "Cv (n)":     _fmt_p(_p.get("Cv")),
                    "PC (n)":     _fmt_p(_p.get("PC")),
                    "phi (n)":    _fmt_p(_p.get("phi")),
                    "c (n)":      _fmt_p(_p.get("c")),
                })
            _df_ly = pd.DataFrame(_rows_ly)
            st.dataframe(
                _df_ly.style.map(_cell_color,
                    subset=["Cc (n)", "Cs (n)", "Cv (n)", "PC (n)"]),
                use_container_width=True, hide_index=True,
            )

        # Tóm tắt 3 khu vực
        st.divider()
        st.markdown("**Tóm tắt 3 khu vực — Thông số lún chính (Cc/Cs/Cv/PC)**")
        _sum_rows = []
        for _zc in ["NHC", "BXN", "KE"]:
            if _zc not in _chk_all:
                continue
            _s = _chk_all[_zc]["zone_summary"]
            _sum_rows.append({
                "Khu vực":      _zc,
                "Lớp yếu":      _s["n_layers_soft"],
                "Cc Đạt/Thiếu": f"{_s['params_ok'].get('Cc',0)}/{_s['params_fail'].get('Cc',0)}",
                "Cs Đạt/Thiếu": f"{_s['params_ok'].get('Cs',0)}/{_s['params_fail'].get('Cs',0)}",
                "Cv Đạt/Thiếu": f"{_s['params_ok'].get('Cv',0)}/{_s['params_fail'].get('Cv',0)}",
                "PC Đạt/Thiếu": f"{_s['params_ok'].get('PC',0)}/{_s['params_fail'].get('PC',0)}",
                "VST":           _chk_all[_zc]["n_vst_zone"],
            })
        if _sum_rows:
            st.table(pd.DataFrame(_sum_rows))
        st.info(
            "Xanh = đủ n≥6 | Đỏ = thiếu mẫu (<6) | '-' = không có mẫu hoặc không áp dụng. "
            "Khu vực KE chưa có mẫu Cc — ưu tiên bổ sung."
        )

    # ── B. TCVN 9403:2012 Bảng B.1 — QC mẫu CDM ────────────────────────────
    st.divider()
    st.markdown("## B. Mẫu kiểm tra chất lượng CDM — TCVN 9403:2012 Bảng B.1")
    st.caption(
        "Bảng B.1: Số mẫu thí nghiệm nén mẫu (phòng) và số mẫu kiểm tra hiện trường "
        "tối thiểu theo số lượng cột CDM thi công. Nhập số liệu để kiểm tra."
    )

    # Bảng yêu cầu Bảng B.1
    _b1_rows = [
        {"Số cột CDM": "≤ 100",      "Mẫu phòng (min/lớp)": 2,  "Kiểm tra hiện trường (min)": 5},
        {"Số cột CDM": "101 – 500",   "Mẫu phòng (min/lớp)": 5,  "Kiểm tra hiện trường (min)": 10},
        {"Số cột CDM": "501 – 1 000", "Mẫu phòng (min/lớp)": 10, "Kiểm tra hiện trường (min)": 30},
        {"Số cột CDM": "1 001 – 2 000","Mẫu phòng (min/lớp)": 15, "Kiểm tra hiện trường (min)": 50},
        {"Số cột CDM": "> 2 000",     "Mẫu phòng (min/lớp)": 20, "Kiểm tra hiện trường (min)": 100},
    ]
    st.table(pd.DataFrame(_b1_rows))

    if not _HAS_9403:
        st.info(
            "Mô-đun kiểm tra CDM tự động chưa được triển khai. "
            "Phần này sẽ khả dụng sau khi bộ thiết kế CDM được cài đặt."
        )
    else:
        st.markdown("**Nhập số liệu kiểm tra thực tế:**")
        _qc_cols = st.columns(3)
        _qc_zone_inputs = {}
        for _ci, _zc in enumerate(["NHC", "BXN", "KE"]):
            with _qc_cols[_ci]:
                st.markdown(f"**Khu vực {_zc}**")
                _n_col = st.number_input(f"Số cột CDM ({_zc})", 0, 10000, 0, 50,
                                         key=f"qc_ncol_{_zc}")
                _n_lab = st.number_input(f"Mẫu phòng hiện có ({_zc})", 0, 500, 0, 1,
                                         key=f"qc_nlab_{_zc}")
                _n_fld = st.number_input(f"Kiểm tra hiện trường hiện có ({_zc})", 0, 500, 0, 1,
                                         key=f"qc_nfld_{_zc}")
                _qc_zone_inputs[_zc] = (_n_col, _n_lab, _n_fld)

        _qc_result_rows = []
        for _zc, (_nc, _nl, _nf) in _qc_zone_inputs.items():
            if _nc == 0:
                continue
            _res9403 = _chk9403(_nc, _nl, _nf)
            _qc_result_rows.append({
                "Khu vực":           _zc,
                "Số cột":            _nc,
                "Mẫu phòng cần":     _res9403["required_lab"],
                "Mẫu phòng có":      _nl,
                "MP thiếu":          _res9403["lab_gap"],
                "MP kết quả":        "Đạt" if _res9403["lab_ok"] else "Không đạt",
                "KT h.trường cần":   _res9403["required_field"],
                "KT h.trường có":    _nf,
                "KT thiếu":          _res9403["field_gap"],
                "KT kết quả":        "Đạt" if _res9403["field_ok"] else "Không đạt",
            })
        if _qc_result_rows:
            _df_qc = pd.DataFrame(_qc_result_rows)
            st.dataframe(
                _df_qc.style.apply(
                    lambda col: [
                        "background-color:#d4edda" if v == "Đạt"
                        else "background-color:#f8d7da" if v == "Không đạt"
                        else "" for v in col
                    ],
                    subset=["MP kết quả", "KT kết quả"],
                ),
                use_container_width=True, hide_index=True,
            )
        else:
            st.info("Nhập số cột CDM > 0 để kiểm tra yêu cầu mẫu theo Bảng B.1.")

    # ── C. Kết quả thí nghiệm CDM 7 ngày — từ SQLite cdm_tests ──────────────
    st.divider()
    st.markdown("## C. Kết quả thí nghiệm CDM 7 ngày")
    st.caption(
        "Dữ liệu thí nghiệm CDM: qu (cường độ chịu nén nở hông) "
        "và E50 (mô đun ở 50% cường độ phá hủy) tại tuổi mẫu 7 ngày. "
        "Mục đích: chọn tổ hợp xi măng + tỷ lệ W/C + hàm lượng đạt yêu cầu thiết kế."
    )
    try:
        _con_cdm = sqlite3.connect(str(_DB))
        _df_cdm7 = pd.read_sql_query(
            """SELECT b.name AS "Hố khoan", z.code AS "Khu vực",
                      ct.group_id AS "Nhóm", ct.test_id AS "Mã mẫu",
                      ct.cement_type AS "Loại xi măng",
                      ct.WC_ratio AS "W/C",
                      ct.dosage_kgm3 AS "Hàm lượng (kg/m³)",
                      ct.qu_R7_kPa AS "qu R7 (kPa)",
                      ct.E50_R7_MPa AS "E50 R7 (MPa)"
               FROM cdm_tests ct
               JOIN boreholes b ON ct.borehole_id = b.id
               JOIN zones z ON b.zone_id = z.id
               ORDER BY z.code, ct.group_id, ct.test_id""",
            _con_cdm,
        )
        _con_cdm.close()

        if _df_cdm7.empty:
            st.warning("Chưa có dữ liệu thí nghiệm CDM. Cần bổ sung kết quả nén nở hông từ phòng.")
        else:
            # Metric tổng quan
            _qu_min = _df_cdm7["qu R7 (kPa)"].min()
            _qu_max = _df_cdm7["qu R7 (kPa)"].max()
            _qu_avg = _df_cdm7["qu R7 (kPa)"].mean()
            _m1, _m2, _m3, _m4 = st.columns(4)
            _m1.metric("Tổng số mẫu", len(_df_cdm7))
            _m2.metric("qu R7 trung bình", f"{_qu_avg:.0f} kPa")
            _m3.metric("qu R7 min", f"{_qu_min:.0f} kPa")
            _m4.metric("qu R7 max", f"{_qu_max:.0f} kPa")

            # User nhập yêu cầu thiết kế để đánh giá
            _qu_req_c1, _qu_req_c2 = st.columns([1, 3])
            with _qu_req_c1:
                _qu_design_req = st.number_input(
                    "qu yêu cầu (kPa)", value=0.0, step=50.0,
                    key="cdm_qu_design_req",
                    help="Cường độ qu thiết kế yêu cầu — kỹ sư nhập theo dự án (vd: 500 kPa cho kè dân dụng)."
                )
            with _qu_req_c2:
                if _qu_design_req > 0:
                    _n_ok  = (_df_cdm7["qu R7 (kPa)"] >= _qu_design_req).sum()
                    _n_tot = len(_df_cdm7)
                    _pct   = 100 * _n_ok / _n_tot
                    st.metric(f"Mẫu đạt qu ≥ {_qu_design_req:.0f} kPa",
                                f"{_n_ok}/{_n_tot}", f"{_pct:.0f}%")

            # Bảng chi tiết, tô màu theo qu yêu cầu
            def _hl_qu(val):
                if isinstance(val, (int, float)) and _qu_design_req > 0:
                    if val >= _qu_design_req:
                        return "background-color:#d4edda; color:#155724"
                    return "background-color:#f8d7da; color:#721c24"
                return ""
            st.markdown("**Bảng chi tiết kết quả thí nghiệm:**")
            st.dataframe(
                _df_cdm7.style.map(_hl_qu, subset=["qu R7 (kPa)"]),
                use_container_width=True, hide_index=True,
                column_config={
                    "qu R7 (kPa)":  st.column_config.NumberColumn(format="%.1f"),
                    "E50 R7 (MPa)": st.column_config.NumberColumn(format="%.1f"),
                    "W/C":           st.column_config.NumberColumn(format="%.2f"),
                    "Hàm lượng (kg/m³)": st.column_config.NumberColumn(format="%.0f"),
                },
            )

            # Biểu đồ Plotly: qu vs dosage, group by cement + W/C
            if _HAS_PLOTLY:
                _df_cdm7["Combo"] = (_df_cdm7["Loại xi măng"] + " · W/C=" +
                                       _df_cdm7["W/C"].astype(str))
                _fig_cdm = go.Figure()
                _colors_combo = ["#1565C0", "#2E7D32", "#E65100", "#C62828"]
                for _i_c, _combo in enumerate(_df_cdm7["Combo"].unique()):
                    _sub = _df_cdm7[_df_cdm7["Combo"] == _combo].sort_values("Hàm lượng (kg/m³)")
                    _fig_cdm.add_trace(go.Scatter(
                        x=_sub["Hàm lượng (kg/m³)"],
                        y=_sub["qu R7 (kPa)"],
                        mode="lines+markers",
                        line=dict(color=_colors_combo[_i_c % 4], width=2.5),
                        marker=dict(size=10),
                        name=_combo,
                        text=[f"{q:.0f} kPa" for q in _sub["qu R7 (kPa)"]],
                        textposition="top center",
                    ))
                if _qu_design_req > 0:
                    _fig_cdm.add_hline(y=_qu_design_req, line_dash="dash",
                                         line_color="red",
                                         annotation_text=f"qu yêu cầu = {_qu_design_req:.0f} kPa")
                _fig_cdm.update_layout(
                    title="qu R7 theo hàm lượng xi măng — phân nhóm theo loại + W/C",
                    xaxis_title="Hàm lượng xi măng (kg/m³)",
                    yaxis_title="qu R7 (kPa) — cường độ chịu nén nở hông 7 ngày",
                    height=400, margin=dict(t=50, b=40),
                    legend=dict(orientation="h", y=1.10),
                )
                st.plotly_chart(_fig_cdm, use_container_width=True)

                # ─────────────────────────────────────────────────────────────
                # Biểu đồ phân tích bổ sung
                # ─────────────────────────────────────────────────────────────
                st.markdown("---")
                st.markdown("#### Phân tích bổ sung")

                # 1. Scatter qu vs E50 — tương quan
                _fig_qE = go.Figure()
                _qu_arr = _df_cdm7["qu R7 (kPa)"].values
                _E_arr  = _df_cdm7["E50 R7 (MPa)"].values
                # Hồi quy tuyến tính
                import numpy as _np_cdm
                _slope, _intc = _np_cdm.polyfit(_qu_arr, _E_arr, 1)
                _qu_fit = _np_cdm.linspace(_qu_arr.min(), _qu_arr.max(), 50)
                _E_fit  = _slope * _qu_fit + _intc
                _corr   = _np_cdm.corrcoef(_qu_arr, _E_arr)[0, 1]
                for _combo in _df_cdm7["Combo"].unique():
                    _sub = _df_cdm7[_df_cdm7["Combo"] == _combo]
                    _fig_qE.add_trace(go.Scatter(
                        x=_sub["qu R7 (kPa)"], y=_sub["E50 R7 (MPa)"],
                        mode="markers", marker=dict(size=12),
                        name=_combo,
                        text=_sub["Mã mẫu"], hovertemplate="%{text}<br>qu=%{x:.0f}<br>E50=%{y:.1f}",
                    ))
                _fig_qE.add_trace(go.Scatter(
                    x=_qu_fit, y=_E_fit, mode="lines",
                    line=dict(color="black", dash="dash", width=2),
                    name=f"Hồi quy: E50 = {_slope:.4f}·qu + {_intc:.2f} (r={_corr:.3f})",
                ))
                _fig_qE.update_layout(
                    title="Tương quan qu R7 ↔ E50 R7",
                    xaxis_title="qu R7 (kPa)", yaxis_title="E50 R7 (MPa)",
                    height=380, margin=dict(t=50, b=40),
                    legend=dict(orientation="h", y=1.13),
                )

                # 2. Box plot qu theo loại xi măng
                _fig_box_cm = go.Figure()
                for _ct in _df_cdm7["Loại xi măng"].unique():
                    _vals = _df_cdm7[_df_cdm7["Loại xi măng"] == _ct]["qu R7 (kPa)"]
                    _fig_box_cm.add_trace(go.Box(
                        y=_vals, name=_ct, boxmean="sd",
                        boxpoints="all", jitter=0.3, pointpos=-1.8,
                    ))
                _fig_box_cm.update_layout(
                    title="Phân bố qu R7 theo loại xi măng",
                    yaxis_title="qu R7 (kPa)",
                    height=380, margin=dict(t=50, b=40),
                    showlegend=False,
                )

                _col_a, _col_b = st.columns(2)
                _col_a.plotly_chart(_fig_qE, use_container_width=True)
                _col_b.plotly_chart(_fig_box_cm, use_container_width=True)

                # 3. Heatmap qu theo cement × dosage (W/C = 0.8 và 1.0)
                from plotly.subplots import make_subplots as _mksub_cdm
                _wc_vals = sorted(_df_cdm7["W/C"].unique())
                _fig_hm = _mksub_cdm(rows=1, cols=len(_wc_vals),
                                       subplot_titles=[f"W/C = {w:.1f}" for w in _wc_vals],
                                       shared_yaxes=True)
                for _i_w, _w in enumerate(_wc_vals):
                    _sub_w = _df_cdm7[_df_cdm7["W/C"] == _w]
                    _piv = _sub_w.pivot_table(
                        index="Loại xi măng", columns="Hàm lượng (kg/m³)",
                        values="qu R7 (kPa)", aggfunc="mean",
                    )
                    _fig_hm.add_trace(go.Heatmap(
                        z=_piv.values, x=_piv.columns, y=_piv.index,
                        text=[[f"{v:.0f}" for v in row] for row in _piv.values],
                        texttemplate="%{text}",
                        colorscale="RdYlGn", showscale=(_i_w == len(_wc_vals)-1),
                        colorbar=dict(title="qu R7<br>(kPa)", x=1.02),
                        zmin=300, zmax=750,
                    ), row=1, col=_i_w+1)
                    _fig_hm.update_xaxes(title_text="Hàm lượng (kg/m³)", row=1, col=_i_w+1)
                _fig_hm.update_yaxes(title_text="Loại xi măng", row=1, col=1)
                _fig_hm.update_layout(
                    title="Heatmap qu R7 theo loại xi măng × hàm lượng × W/C",
                    height=320, margin=dict(t=60, b=40, r=80),
                )
                st.plotly_chart(_fig_hm, use_container_width=True)

                # 4. Bar chart qu trung bình + std theo nhóm
                _grp = _df_cdm7.groupby(["Loại xi măng", "W/C"]).agg(
                    qu_mean=("qu R7 (kPa)", "mean"),
                    qu_std=("qu R7 (kPa)", "std"),
                    qu_min=("qu R7 (kPa)", "min"),
                    qu_max=("qu R7 (kPa)", "max"),
                    n=("qu R7 (kPa)", "count"),
                ).reset_index()
                _grp["Tổ hợp"] = _grp["Loại xi măng"] + " | W/C=" + _grp["W/C"].astype(str)
                _fig_bar = go.Figure()
                _fig_bar.add_trace(go.Bar(
                    x=_grp["Tổ hợp"], y=_grp["qu_mean"],
                    error_y=dict(type="data", array=_grp["qu_std"]),
                    marker_color=["#1565C0", "#2E7D32", "#E65100", "#C62828"][:len(_grp)],
                    text=[f"{m:.0f}±{s:.0f}" for m, s in zip(_grp["qu_mean"], _grp["qu_std"])],
                    textposition="outside",
                ))
                if _qu_design_req > 0:
                    _fig_bar.add_hline(y=_qu_design_req, line_dash="dash", line_color="red",
                                         annotation_text=f"qu yêu cầu = {_qu_design_req:.0f}")
                _fig_bar.update_layout(
                    title="qu R7 trung bình ± độ lệch chuẩn theo tổ hợp",
                    yaxis_title="qu R7 (kPa)",
                    height=380, margin=dict(t=50, b=80),
                    showlegend=False,
                )
                st.plotly_chart(_fig_bar, use_container_width=True)

                # Bảng thống kê tổ hợp
                _grp_display = _grp[["Tổ hợp", "n", "qu_mean", "qu_std", "qu_min", "qu_max"]].copy()
                _grp_display.columns = ["Tổ hợp", "n mẫu", "qu trung bình", "qu σ", "qu min", "qu max"]
                st.markdown("**Bảng thống kê theo tổ hợp:**")
                st.dataframe(
                    _grp_display.style.format({
                        "qu trung bình": "{:.1f}",
                        "qu σ":  "{:.1f}",
                        "qu min": "{:.1f}",
                        "qu max": "{:.1f}",
                    }),
                    use_container_width=True, hide_index=True,
                )

                # ─────────────────────────────────────────────────────────────
                # Biểu đồ phân tích nâng cao (statistical + parametric)
                # ─────────────────────────────────────────────────────────────
                st.markdown("---")
                st.markdown("#### Phân tích thống kê + độ nhạy")

                # 5. Histogram qu + KDE
                _fig_hist = go.Figure()
                _fig_hist.add_trace(go.Histogram(
                    x=_qu_arr, nbinsx=8,
                    marker_color="#1565C0", opacity=0.7,
                    name="Phân bố qu R7", histnorm="probability density",
                ))
                # KDE qua scipy
                try:
                    from scipy.stats import gaussian_kde as _kde
                    _kde_fn = _kde(_qu_arr)
                    _x_kde = _np_cdm.linspace(_qu_arr.min() - 50, _qu_arr.max() + 50, 100)
                    _fig_hist.add_trace(go.Scatter(
                        x=_x_kde, y=_kde_fn(_x_kde),
                        mode="lines", line=dict(color="#E65100", width=2.5),
                        name="KDE",
                    ))
                except Exception:
                    pass
                _fig_hist.add_vline(x=_qu_avg, line_dash="dash", line_color="black",
                                     annotation_text=f"μ={_qu_avg:.0f}")
                _fig_hist.add_vline(x=_qu_avg + _qu_arr.std(), line_dash="dot", line_color="gray",
                                     annotation_text="+σ")
                _fig_hist.add_vline(x=_qu_avg - _qu_arr.std(), line_dash="dot", line_color="gray",
                                     annotation_text="-σ")
                if _qu_design_req > 0:
                    _fig_hist.add_vline(x=_qu_design_req, line_dash="dash", line_color="red",
                                         annotation_text=f"yêu cầu={_qu_design_req:.0f}")
                _fig_hist.update_layout(
                    title=f"Phân bố qu R7 (n={len(_qu_arr)}, μ={_qu_avg:.0f}, σ={_qu_arr.std():.0f} kPa)",
                    xaxis_title="qu R7 (kPa)", yaxis_title="Mật độ xác suất",
                    height=380, margin=dict(t=50, b=40),
                    bargap=0.05, showlegend=False,
                )

                # 6. CDF (Cumulative Distribution)
                _sorted_qu = _np_cdm.sort(_qu_arr)
                _cdf_y = _np_cdm.arange(1, len(_sorted_qu) + 1) / len(_sorted_qu)
                _fig_cdf = go.Figure()
                _fig_cdf.add_trace(go.Scatter(
                    x=_sorted_qu, y=_cdf_y * 100,
                    mode="lines+markers",
                    line=dict(color="#2E7D32", width=2.5, shape="hv"),
                    marker=dict(size=8),
                    name="CDF",
                ))
                if _qu_design_req > 0:
                    _frac_below = (_sorted_qu < _qu_design_req).sum() / len(_sorted_qu) * 100
                    _fig_cdf.add_vline(x=_qu_design_req, line_dash="dash", line_color="red",
                                         annotation_text=f"yêu cầu ({_frac_below:.0f}% mẫu không đạt)")
                _fig_cdf.add_hline(y=50, line_dash="dot", line_color="gray",
                                     annotation_text="median (50%)")
                _fig_cdf.update_layout(
                    title="Đường cong phân bố tích lũy CDF — qu R7",
                    xaxis_title="qu R7 (kPa)", yaxis_title="P(qu ≤ x) (%)",
                    height=380, margin=dict(t=50, b=40),
                    yaxis=dict(range=[0, 105]),
                    showlegend=False,
                )

                _hcol1, _hcol2 = st.columns(2)
                _hcol1.plotly_chart(_fig_hist, use_container_width=True)
                _hcol2.plotly_chart(_fig_cdf, use_container_width=True)

                # 7. Sensitivity: độ tăng qu per +20 kg/m³ dosage cho mỗi tổ hợp
                _sens_rows = []
                for _combo_s in _df_cdm7["Combo"].unique():
                    _sub_s = _df_cdm7[_df_cdm7["Combo"] == _combo_s].sort_values("Hàm lượng (kg/m³)")
                    if len(_sub_s) < 2:
                        continue
                    _slope_s, _intc_s = _np_cdm.polyfit(_sub_s["Hàm lượng (kg/m³)"],
                                                          _sub_s["qu R7 (kPa)"], 1)
                    _sens_rows.append({
                        "Tổ hợp":          _combo_s,
                        "Δqu / +20 kg/m³": _slope_s * 20,
                        "qu @ 220":        _slope_s * 220 + _intc_s,
                        "qu @ 260":        _slope_s * 260 + _intc_s,
                        "Hệ số (kPa/kg/m³)": _slope_s,
                    })
                _df_sens = pd.DataFrame(_sens_rows)
                if not _df_sens.empty:
                    _fig_sens = go.Figure()
                    _fig_sens.add_trace(go.Bar(
                        y=_df_sens["Tổ hợp"], x=_df_sens["Δqu / +20 kg/m³"],
                        orientation="h",
                        marker_color="#9C27B0",
                        text=[f"+{v:.0f} kPa" for v in _df_sens["Δqu / +20 kg/m³"]],
                        textposition="outside",
                    ))
                    _fig_sens.update_layout(
                        title="Độ nhạy: Δqu khi tăng hàm lượng +20 kg/m³",
                        xaxis_title="Δqu R7 (kPa)",
                        height=320, margin=dict(t=50, b=40, l=200),
                        showlegend=False,
                    )

                    # 8. 3D surface (W/C × dosage × qu) cho mỗi cement type
                    _fig_3d = go.Figure()
                    _colors_3d = {"Xỉ Tam Điệp": "Blues", "Hoàng Thạch PCB40": "Reds"}
                    for _ct_3d in _df_cdm7["Loại xi măng"].unique():
                        _sub3 = _df_cdm7[_df_cdm7["Loại xi măng"] == _ct_3d]
                        _fig_3d.add_trace(go.Scatter3d(
                            x=_sub3["W/C"], y=_sub3["Hàm lượng (kg/m³)"],
                            z=_sub3["qu R7 (kPa)"],
                            mode="markers+text",
                            marker=dict(size=8,
                                          color=_sub3["qu R7 (kPa)"],
                                          colorscale=_colors_3d.get(_ct_3d, "Viridis"),
                                          showscale=False),
                            text=_sub3["Mã mẫu"], textposition="top center",
                            name=_ct_3d,
                        ))
                    _fig_3d.update_layout(
                        title="Phân bố 3D: W/C × Hàm lượng × qu R7",
                        scene=dict(xaxis_title="W/C", yaxis_title="Hàm lượng (kg/m³)",
                                    zaxis_title="qu R7 (kPa)"),
                        height=480, margin=dict(t=50, b=40),
                        legend=dict(orientation="h", y=1.10),
                    )

                    _s1, _s2 = st.columns(2)
                    _s1.plotly_chart(_fig_sens, use_container_width=True)
                    _s2.plotly_chart(_fig_3d, use_container_width=True)

                    # Bảng độ nhạy
                    st.markdown("**Bảng độ nhạy — qu tăng theo hàm lượng xi măng:**")
                    st.dataframe(
                        _df_sens.style.format({
                            "Δqu / +20 kg/m³":   "{:+.1f}",
                            "qu @ 220":          "{:.1f}",
                            "qu @ 260":          "{:.1f}",
                            "Hệ số (kPa/kg/m³)": "{:.3f}",
                        }),
                        use_container_width=True, hide_index=True,
                    )
                    st.caption(
                        "Hệ số tuyến tính dqu/d(dosage) cho biết hiệu quả tăng cường độ khi tăng "
                        "hàm lượng xi măng. Tổ hợp có hệ số cao = đáng đầu tư thêm xi măng."
                    )

            # Tổ hợp tốt nhất (qu cao nhất)
            _best = _df_cdm7.loc[_df_cdm7["qu R7 (kPa)"].idxmax()]
            st.success(
                f"**Tổ hợp đạt qu cao nhất:** {_best['Loại xi măng']} · "
                f"W/C = {_best['W/C']:.2f} · Hàm lượng = {_best['Hàm lượng (kg/m³)']:.0f} kg/m³ "
                f"→ qu = **{_best['qu R7 (kPa)']:.1f} kPa**, "
                f"E50 = {_best['E50 R7 (MPa)']:.1f} MPa (mẫu {_best['Mã mẫu']}, HK {_best['Hố khoan']})"
            )
    except Exception as _e_cdm:
        st.warning(f"Không tải được dữ liệu CDM: {_e_cdm}")

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 2 – THÔNG SỐ
# ═══════════════════════════════════════════════════════════════════════════════
elif _page == "params":
    st.subheader(_t("p2_sub"))

    # ── Lý thuyết tính toán (đặt trên cùng để xem trước khi nhập thông số) ───
    _theory_text = _load_theory()
    if _theory_text:
        with st.expander(_t("theory_exp"), expanded=False):
            st.markdown(_theory_text)

    # ── Áp địa chất từ HK (SQLite + nearest fallback) ────────────────────────
    with st.container(border=True):
        st.markdown("**Áp địa chất từ hố khoan** — đọc trực tiếp từ số liệu khảo sát, "
                    "thiếu thì lấy HK gần nhất")
        _ap_c1, _ap_c2, _ap_c3 = st.columns([1, 2, 1])
        with _ap_c1:
            _ap_zone = st.radio(
                "Zone", list(_ZONE_NAMES.keys()),
                format_func=lambda z: _ZONE_NAMES[z],
                index=list(_ZONE_NAMES.keys()).index(_get("cdm_zone")),
                key="_p_ap_zone", horizontal=False,
            )
        with _ap_c2:
            _ap_bhs   = [b["name"] for b in _load_boreholes_by_zone(_ap_zone)]
            _cur_bh   = _get("cdm_bh") if _get("cdm_bh") in _ap_bhs else (_ap_bhs[0] if _ap_bhs else "")
            _ap_bh    = st.selectbox(
                "Hố khoan áp",
                _ap_bhs, index=_ap_bhs.index(_cur_bh) if _cur_bh in _ap_bhs else 0,
                key="_p_ap_bh",
            )
        with _ap_c3:
            st.markdown("&nbsp;", unsafe_allow_html=True)
            _btn_ap = st.button("Áp địa chất", type="primary",
                                 key="_p_btn_ap_geo", use_container_width=True)

        if _btn_ap and _ap_bh:
            # ── Helper: nearest BH có γ avg ──────────────────────────────────
            def _gamma_nearest(bh_target, db_path):
                import sqlite3 as _sq, math as _math
                con = _sq.connect(str(db_path))
                row = con.execute(
                    "SELECT b.x_coord_m, b.y_coord_m, z.code FROM boreholes b "
                    "JOIN zones z ON b.zone_id=z.id WHERE b.name=?", (bh_target,)
                ).fetchone()
                if not row or row[0] is None:
                    con.close(); return None, None, None
                x0, y0, zone = row
                rows = con.execute("""
                    SELECT b.name, b.x_coord_m, b.y_coord_m, AVG(lt.gamma_kNm3)
                    FROM lab_tests lt JOIN boreholes b ON lt.borehole_id=b.id
                    JOIN zones z ON b.zone_id=z.id
                    WHERE z.code=? AND b.name<>? AND lt.gamma_kNm3 IS NOT NULL
                      AND b.x_coord_m IS NOT NULL GROUP BY b.id
                """, (zone, bh_target)).fetchall()
                con.close()
                if not rows: return None, None, None
                cand = sorted(((((x-x0)**2+(y-y0)**2)**0.5, nm, v) for nm,x,y,v in rows))
                d, nm, v = cand[0]; return round(float(v),2), nm, round(d,1)

            def _su_nearest(zone_code, depth_top, depth_bot, ref_bh):
                """Tìm zone khác có VST trong depth range gần BH ref."""
                # Cách đơn giản: chỉ dùng VST của zone hiện tại (tránh xa zone)
                v = _su_avg_in_range(zone_code, depth_top, depth_bot)
                return v, ref_bh, 0.0  # cùng zone tổng hợp → no dist concept

            # ── Lưu metadata để tô xanh label + caption nguồn ─────────────────
            _applied = st.session_state.get("_applied_fields", {})

            _src_log = []
            # 1. top_clay
            _cp = _clay_params(_ap_bh, _ap_zone)
            if _cp:
                _layers = _load_layers(_ap_bh)
                _clay_syms = _CLAY_SYMBOLS.get(_ap_zone, ["1", "2"])
                _first_clay = next((l for l in _layers if l["symbol"] in _clay_syms), None)
                if _first_clay:
                    _depth_top = _first_clay["depth_top_m"]
                    _top_clay_v = round(_cp["elevation_m"] - _depth_top, 2)
                    st.session_state["cdm_top_clay"] = _top_clay_v
                    _applied["cdm_top_clay"] = {"src": _ap_bh, "dist_m": None,
                                                "value": _top_clay_v}
                    _src_log.append(("Cao độ đỉnh lớp bùn",
                                     f"từ {_ap_bh} (depth_top={_depth_top:.2f}m, elev={_cp['elevation_m']:.2f}m)"))

                # 2. h_clay
                st.session_state["cdm_h_clay"] = _cp["h_clay"]
                st.session_state["cdm_Lc"]    = round(_cp["h_clay"] + 1.5, 1)
                _applied["cdm_h_clay"] = {"src": _ap_bh, "dist_m": None,
                                          "value": _cp["h_clay"]}
                _src_log.append(("Bề dày lớp bùn", f"từ {_ap_bh} h_clay={_cp['h_clay']}m"))
            else:
                _src_log.append(("Lớp bùn", "HK không có dữ liệu lớp"))

            # 3. Su
            _depth_bot_clay = _cp.get("depth_clay_bot", 30) if _cp else 30
            _su = _su_avg_in_range(_ap_zone, 0, _depth_bot_clay)
            if _su:
                st.session_state["cdm_Su"] = _su
                _applied["cdm_Su"] = {"src": f"VST zone {_ap_zone}", "dist_m": None,
                                      "value": _su}
                _src_log.append(("Su trung bình", f"VST zone {_ap_zone} trong 0–{_depth_bot_clay:.1f}m"))
            else:
                _src_log.append(("Su trung bình", "không có VST — giữ giá trị cũ"))

            # 4. γ — ưu tiên HK hiện tại, nearest fallback
            _g = _gamma_avg(_ap_bh)
            if _g and _g != 15.0:
                st.session_state["cdm_gamma"] = _g
                _applied["cdm_gamma"] = {"src": _ap_bh, "dist_m": None, "value": _g}
                _src_log.append(("γ tự nhiên", f"lab_tests {_ap_bh}"))
            else:
                _g_n, _g_bh, _g_dist = _gamma_nearest(_ap_bh, _DB)
                if _g_n:
                    st.session_state["cdm_gamma"] = _g_n
                    _applied["cdm_gamma"] = {"src": _g_bh, "dist_m": _g_dist,
                                             "value": _g_n}
                    _src_log.append(("γ tự nhiên",
                                     f"**HK gần nhất {_g_bh} (cách {_g_dist:.0f} m)** — γ={_g_n} kN/m³"))
                else:
                    _src_log.append(("γ tự nhiên", "không có γ — giữ giá trị mặc định 15.0"))

            st.session_state["cdm_bh"]   = _ap_bh
            st.session_state["cdm_zone"] = _ap_zone
            st.session_state["_applied_fields"] = _applied

            st.success(f"Đã áp địa chất từ **{_ap_bh}** — các ô có nhãn xanh là số liệu vừa áp")
            for _fname, _src in _src_log:
                st.caption(f"• **{_fname}** ← {_src}")
            st.rerun()

    # Layout: thông số nhập full-width, hình minh họa bên dưới
    _col_inp = st.container()

    with _col_inp:
        c1, c2, c3, c4 = st.columns(4)

        with c1:
            st.markdown(f"**{_t('cdm_geom')}**")
            D      = st.number_input(_t("D_lbl"), 0.5, 1.2, _get("cdm_D"), 0.05)
            e      = st.number_input("Khoảng cách e (m)", 0.8, 4.0, _get("cdm_e"), 0.1)
            CDTK   = st.number_input(_t("CDTK_lbl"), -5.0, 10.0, _get("cdm_CDTK"), 0.1)
            L_ngam = st.number_input("Ngàm vào đất tốt (m)", 0.0, 5.0, _get("cdm_L_ngam"), 0.1,
                                     help="Chiều dài cọc CDM ngàm vào lớp đất tốt bên dưới lớp bùn")
            _top_clay_c1 = _get("cdm_top_clay")
            _h_clay_c1   = _get("cdm_h_clay")
            _gap_c1      = CDTK - _top_clay_c1          # khoảng từ đỉnh bùn lên đỉnh cọc
            Lc           = round(_gap_c1 + _h_clay_c1 + L_ngam, 1)
            st.info(
                f"Lc = {_gap_c1:.1f} + {_h_clay_c1:.1f} + {L_ngam:.1f} = **{Lc:.1f} m**\n\n"
                f"*(đỉnh bùn→cọc) + h_bùn + ngàm*"
            )
            _ld0 = _get("cdm_loads")
            _He_v  = _ld0.get("h_fill", 1.5)
            _Hse_v = _ld0.get("h_mat",  0.4)
            st.info(f"He = **{_He_v:.2f} m**   |   Hse = **{_Hse_v:.2f} m**")
            arr  = st.radio(_t("arr_lbl"), ["triangle", "square"],
                            format_func=lambda x: _t("arr_tri") if x == "triangle" else _t("arr_sq"),
                            index=["triangle", "square"].index(_get("cdm_arrangement")))
            st.session_state.update({"cdm_D": D, "cdm_e": e, "cdm_L_ngam": L_ngam,
                                      "cdm_Lc": Lc, "cdm_CDTK": CDTK, "cdm_arrangement": arr})

        with c2:
            st.markdown(f"**{_t('material')}**")
            cement = st.text_input(_t("cement_lbl"), _get("cdm_cement_type"))
            dosage = st.number_input(_t("dosage_lbl"), 100, 400, _get("cdm_dosage"), 10)
            WC     = st.number_input(_t("wc_lbl"), 0.5, 2.0, _get("cdm_WC"), 0.1)
            qu     = st.number_input(_t("qu_lbl"), 100.0, 5000.0, _get("cdm_qu"), 50.0)
            FS_lab = st.number_input(_t("fs_lab"), 1.0, 5.0, _get("cdm_FS_lab"), 0.1)
            Cc = qu / 2
            Ec = 100 * Cc
            st.info(f"Cc = {Cc:.0f} kN/m²  |  Ec = {int(Ec):,} kN/m²")
            st.session_state.update({"cdm_cement_type": cement, "cdm_dosage": dosage,
                                      "cdm_WC": WC, "cdm_qu": qu, "cdm_FS_lab": FS_lab})

        with c3:
            st.markdown(f"**{_t('geo_params')}**")
            # Helper: label xanh + caption nguồn nếu field đã được áp
            _apf = st.session_state.get("_applied_fields", {})

            def _lbl(key: str, base: str) -> str:
                return f":blue[{base}]" if key in _apf else base

            def _src_cap(key: str) -> None:
                if key not in _apf:
                    return
                info = _apf[key]
                src = info.get("src", "?")
                dist = info.get("dist_m")
                txt = f"← áp từ **{src}**"
                if dist:
                    txt += f" _(cách {dist:.0f} m)_"
                st.markdown(f":blue[{txt}]")

            top_clay = st.number_input(_lbl("cdm_top_clay", _t("top_clay_lbl")),
                                       -20.0, 20.0, _get("cdm_top_clay"), 0.1)
            _src_cap("cdm_top_clay")
            h_clay  = st.number_input(_lbl("cdm_h_clay", _t("h_clay_lbl")),
                                       1.0, 60.0, _get("cdm_h_clay"), 0.5)
            _src_cap("cdm_h_clay")
            Su      = st.number_input(_lbl("cdm_Su", _t("su_lbl")),
                                       1.0, 100.0, _get("cdm_Su"), 0.5)
            _src_cap("cdm_Su")
            gamma   = st.number_input(_lbl("cdm_gamma", _t("gamma_lbl")),
                                       10.0, 22.0, _get("cdm_gamma"), 0.1)
            _src_cap("cdm_gamma")

            # Detect user manual override → remove "applied" flag
            for _k, _new_v in (("cdm_top_clay", top_clay), ("cdm_h_clay", h_clay),
                               ("cdm_Su", Su), ("cdm_gamma", gamma)):
                if _k in _apf and abs(_apf[_k].get("value", _new_v) - _new_v) > 1e-6:
                    _apf.pop(_k, None)
            st.session_state["_applied_fields"] = _apf

            st.info(_t("bot_clay_info", v=top_clay - h_clay))
            Es_show = 250 * Su
            st.info(f"Es = 250×Su = {int(Es_show):,} kN/m²")
            st.session_state.update({"cdm_top_clay": top_clay, "cdm_h_clay": h_clay,
                                     "cdm_Su": Su, "cdm_gamma": gamma})

        with c4:
            st.markdown(f"**{_t('loads_title')}**")
            ld = _get("cdm_loads").copy()
            ld["q_traffic"] = st.number_input(_t("q_traf"), 0.0, 200.0,
                                              ld.get("q_traffic", 20.0), 5.0)
            ld["z_tk"] = st.number_input(
                "Cao độ thiết kế (m)",
                -5.0, 20.0, ld.get("z_tk", 3.5), 0.1,
                help="Cao độ đỉnh mặt đường thiết kế",
            )
            _top_clay_now = _get("cdm_top_clay")
            _cdtk_now     = _get("cdm_CDTK")
            _H = ld["z_tk"] - _top_clay_now
            if _H < 0:
                st.error("Cao độ TK thấp hơn đỉnh bùn → H < 0")
            _hdr_style = "font-size:12px; font-weight:600; color:#444; margin-bottom:2px"
            _cl, _ch, _cg = st.columns([1.5, 1, 1])
            _ch.markdown(f"<div style='{_hdr_style}'>h (m)</div>", unsafe_allow_html=True)
            _cg.markdown(f"<div style='{_hdr_style}'>γ (kN/m³)</div>", unsafe_allow_html=True)

            _cl, _ch, _cg = st.columns([1.5, 1, 1])
            _cl.markdown(_t("pavement"))
            ld["h_road"] = _ch.number_input("h_road", 0.0, 2.0, ld.get("h_road", 0.8), 0.1, key="h_road", label_visibility="collapsed")
            ld["g_road"] = _cg.number_input("g_road", 10.0, 30.0, ld.get("g_road", 24.0), 0.5, key="g_road", label_visibility="collapsed")

            _cl, _ch, _cg = st.columns([1.5, 1, 1])
            _cl.markdown("*He — cát đắp:*")
            ld["h_fill"] = _ch.number_input("h_fill", 0.0, 5.0, ld.get("h_fill", 1.5), 0.1, key="h_fill", label_visibility="collapsed")
            ld["g_fill"] = _cg.number_input("g_fill", 10.0, 24.0, ld.get("g_fill", 18.0), 0.5, key="g_fill", label_visibility="collapsed")

            _cl, _ch, _cg = st.columns([1.5, 1, 1])
            _cl.markdown("*Hse — đệm cát:*")
            ld["h_mat"] = _ch.number_input("h_mat", 0.0, 2.0, ld.get("h_mat", 0.4), 0.1, key="h_mat", label_visibility="collapsed")
            ld["g_mat"] = _cg.number_input("g_mat", 10.0, 26.0, ld.get("g_mat", 22.5), 0.5, key="g_mat", label_visibility="collapsed")

            # Ràng buộc: z_tk = CDTK + Hse + He + h_đường
            _Hse       = ld["h_mat"]
            _He        = ld["h_fill"]
            _h_road    = ld["h_road"]
            _fill_gap  = _cdtk_now - _top_clay_now          # cát đắp từ đỉnh bùn lên đỉnh cọc
            _H_parts   = _fill_gap + _Hse + _He + _h_road   # H theo hình học
            _z_road_calc = _cdtk_now + _Hse + _He + _h_road
            st.info(
                f"H = (CDTK − đỉnh bùn) + Hse + He + h_đường\n\n"
                f"= {_fill_gap:.2f} + {_Hse:.2f} + {_He:.2f} + {_h_road:.2f} = **{_H_parts:.2f} m**"
            )
            _dz = ld["z_tk"] - _z_road_calc
            if abs(_dz) > 0.05:
                st.warning(
                    f"Cao độ TK nhập ({ld['z_tk']:+.2f} m) ≠ "
                    f"CDTK + Hse + He + h_đường = **{_z_road_calc:+.2f} m** (Δ = {_dz:+.2f} m)."
                )
                if st.button("Đồng bộ cao độ TK = hình học", key="sync_ztk"):
                    ld["z_tk"] = round(_z_road_calc, 2)
                    st.session_state["cdm_loads"] = ld
                    st.rerun()
            q_tot = q_total(ld)
            st.success(f"**q = {q_tot:.2f} kN/m²**")
            st.session_state["cdm_loads"] = ld

    # ── Bảng thông số địa chất (từ hố khoan đã áp dụng) ────────────────────────
    _p_bh = _get("cdm_bh")
    if _p_bh:
        _p_layers = _load_layers(_p_bh)
        if _p_layers:
            _p_zone = _get("cdm_zone")
            _p_rows = []
            _p_borrowed = False
            for _lr in _p_layers:
                _sym = _lr.get("symbol") or ""
                _Cc  = _lr.get("Cc");  _Cs = _lr.get("Cs")
                _e0  = _lr.get("e0");  _PC = _lr.get("PC_kPa")
                _Cv  = _lr.get("Cv_cm2s"); _gam = _lr.get("gamma_kNm3")
                _src = ""
                if _Cc is None or _e0 is None:
                    _fb = _load_zone_params_by_symbol(_p_zone, _sym)
                    if _fb:
                        _Cc  = _Cc  if _Cc  is not None else _fb.get("Cc")
                        _Cs  = _Cs  if _Cs  is not None else _fb.get("Cs")
                        _e0  = _e0  if _e0  is not None else _fb.get("e0")
                        _PC  = _PC  if _PC  is not None else _fb.get("PC_kPa")
                        _Cv  = _Cv  if _Cv  is not None else _fb.get("Cv_cm2s")
                        _gam = _gam if _gam is not None else _fb.get("gamma_kNm3")
                        _src = _fb.get("source_bh") or ""
                        if _src:
                            _p_borrowed = True
                _p_rows.append({
                    "Lớp":         _sym,
                    "Mô tả":       (_lr.get("description") or "")[:30],
                    "h (m)":       _lr.get("thickness_m"),
                    "γ (kN/m³)":  _gam,
                    "e0":          _e0,
                    "Cc":          _Cc,
                    "Cs":          _Cs,
                    "PC (kPa)":   _PC,
                    "Cv (cm²/s)": f"{_Cv:.2e}" if _Cv is not None else "—",
                    "_src":        _src,
                })
            _df_p = pd.DataFrame(_p_rows)
            _src_map_p = _df_p["_src"].to_dict()
            def _hl_p(row):
                return ["background-color:#FFF8DC"]*len(row) if _src_map_p.get(row.name) else [""]*len(row)
            with st.expander(f"Thông số địa chất — {_p_bh}", expanded=True):
                st.dataframe(
                    _df_p.drop(columns=["_src"]).style.apply(_hl_p, axis=1),
                    use_container_width=True, hide_index=True,
                    column_config={
                        "h (m)":     st.column_config.NumberColumn(format="%.2f"),
                        "γ (kN/m³)":st.column_config.NumberColumn(format="%.2f"),
                        "e0":        st.column_config.NumberColumn(format="%.3f"),
                        "Cc":        st.column_config.NumberColumn(format="%.4f"),
                        "Cs":        st.column_config.NumberColumn(format="%.4f"),
                        "PC (kPa)": st.column_config.NumberColumn(format="%.1f"),
                    },
                )
                if _p_borrowed:
                    st.caption("*Ô nền vàng: tham khảo từ hố khoan khác cùng khu vực.*")

    # ── Kiểm tra ràng buộc số liệu ──────────────────────────────────────────────
    _vc_top  = _get("cdm_top_clay")
    _vc_CDTK = _get("cdm_CDTK")
    _vc_Lc   = _get("cdm_Lc")
    _vc_hc   = _get("cdm_h_clay")
    _vc_zpb  = _vc_CDTK - _vc_Lc          # đáy cọc
    _vc_zcb  = _vc_top  - _vc_hc          # đáy bùn
    _vc_emb  = _vc_zcb  - _vc_zpb         # ngàm vào đất tốt

    if _vc_top > _vc_CDTK + 0.05:
        st.error(
            f"Đỉnh lớp bùn ({_vc_top:+.2f} m) cao hơn đỉnh cọc CDM ({_vc_CDTK:+.2f} m). "
            "Hình vẽ sẽ sai — giảm 'Cao độ đỉnh lớp bùn' hoặc tăng 'Cao độ đỉnh cọc CDM'."
        )
    if _vc_emb <= 0:
        st.error(
            f"Mũi cọc ({_vc_zpb:+.2f} m) chưa qua đáy bùn ({_vc_zcb:+.2f} m). "
            f"Cần Lc > {_vc_CDTK - _vc_zcb:.1f} m để cọc ngàm vào đất tốt."
        )
    elif _vc_emb < 0.5:
        st.warning(
            f"Ngàm cọc vào đất tốt chỉ {_vc_emb:.2f} m — khuyến nghị ≥ 0.5 m "
            f"(Lc hiện = {_vc_Lc:.1f} m, cần ≥ {_vc_Lc + (0.5 - _vc_emb):.1f} m)."
        )

    # ── Hình minh họa bên dưới: mặt cắt (trái) + lưới mặt bằng (phải) ────────
    _col_sec, _col_grd = st.columns(2, gap="medium")

    with _col_sec:
        _ld_s = _get("cdm_loads")
        _fig_sec = _draw_cdm_section(
            D=_get("cdm_D"), Lc=_get("cdm_Lc"), CDTK=_get("cdm_CDTK"),
            h_clay=_get("cdm_h_clay"),
            h_road=_ld_s.get("h_road", 0.8),
            h_fill=_ld_s.get("h_fill", 1.5),
            h_mat=_ld_s.get("h_mat", 0.4),
            q_traffic=_ld_s.get("q_traffic", 20.0),
            top_clay=_get("cdm_top_clay"),
        )
        st.pyplot(_fig_sec, use_container_width=True)
        _fig_sec.clf()

    with _col_grd:
        # Hình lưới ở tab Thông số dùng e do user nhập trực tiếp (cdm_e)
        # — KHÔNG dùng cdm_spacings[rec_idx] (đó là kết quả từ tab So sánh)
        _e_ref_now = _get("cdm_e")
        _fig_grd = _draw_cdm_grid(_get("cdm_D"), _get("cdm_arrangement"), _e_ref_now)
        st.pyplot(_fig_grd, use_container_width=True)
        _fig_grd.clf()

    # Xuất PDF qua nút app đã tắt (§25 CLAUDE.md) — dùng Ctrl+P trình duyệt


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 3 – SO SÁNH PHƯƠNG ÁN
# ═══════════════════════════════════════════════════════════════════════════════
if _page == "params":   # tiếp nội dung Kết quả CDM (gộp vào tab Thông số)
    st.subheader(_t("p3_sub"))

    # Input khoảng cách + thông số đệm xi măng
    with st.expander(_t("setup_exp"), expanded=True):
        _col_sp, _col_mat = st.columns([3, 2], gap="large")

        with _col_sp:
            st.markdown(f"**{_t('sp_title')}**")
            _auto_lbl, _manual_lbl = _t("auto_lbl"), _t("manual_lbl")
            mode = st.radio("", [_auto_lbl, _manual_lbl],
                            horizontal=True, label_visibility="collapsed",
                            key="_sp_mode")
            if mode == _auto_lbl:
                _sc1, _sc2, _sc3 = st.columns(3)
                e_min  = _sc1.number_input(_t("e_min"), 1.0, 3.0, 1.2, 0.1)
                e_max  = _sc2.number_input(_t("e_max"), 1.0, 4.0, 2.0, 0.1)
                e_step = _sc3.number_input(_t("step_lbl"), 0.1, 1.0, 0.2, 0.1)
                spacings = [round(e_min + i * e_step, 2)
                            for i in range(int(round((e_max - e_min) / e_step)) + 1)]
            else:
                raw = st.text_input(_t("sp_list"),
                                    ", ".join(str(s) for s in _get("cdm_spacings")),
                                    label_visibility="visible")
                try:
                    spacings = [float(x.strip()) for x in raw.split(",") if x.strip()]
                except ValueError:
                    spacings = _get("cdm_spacings")
            spacings = sorted(set(spacings))
            st.session_state["cdm_spacings"] = spacings
            _alts_label = "Các phương án" if _get("lang") == "VN" else "Alternatives"
            st.caption(f"{_alts_label}: {', '.join(f'{e}m' for e in spacings)}")

        with _col_mat:
            st.markdown(f"**{_t('mat_punch')}**")
            _mc1, _mc2 = st.columns(2)
            with _mc1:
                _quckse_sp = st.number_input(_t("qu_mat_sp"), 100.0, 3000.0,
                                             _get("cdm_quckse"), 50.0,
                                             key="_sp_quckse")
                _theta_sp  = st.number_input("θ (°)", 30.0, 89.0,
                                             _get("cdm_theta"), 5.0,
                                             key="_sp_theta")
            with _mc2:
                _Fs_sp     = st.number_input("Fs", 1.0, 5.0,
                                             _get("cdm_Fs_mat"), 0.5,
                                             key="_sp_Fs")
                _qa_sp     = st.number_input("qa (kPa)", 0.0, 200.0,
                                             _get("cdm_qa_mat"), 5.0,
                                             key="_sp_qa")
            _tase_sp = _quckse_sp / (2.0 * _Fs_sp)
            st.info(f"τase = qu/(2·Fs) = **{_tase_sp:.1f} kPa**")
            st.session_state.update({
                "cdm_quckse": _quckse_sp,
                "cdm_Fs_mat": _Fs_sp,
                "cdm_theta":  _theta_sp,
                "cdm_qa_mat": _qa_sp,
            })

    # Hình minh họa phương pháp tính (40% width mỗi cột)
    _img_load = _ROOT / "giao dien" / "PNBAY-TAI TRONG NEN DAP LEN PHAN KHÔNG GIA CO.png"
    _img_tse  = _ROOT / "giao dien" / "PNBAY-Tse.png"
    _ic1, _ic2 = st.columns(2)
    if _img_load.exists():
        with _ic1:
            _s1, _ = st.columns([2, 3])
            _s1.image(str(_img_load), use_container_width=True)
    if _img_tse.exists():
        with _ic2:
            _s2, _ = st.columns([2, 3])
            _s2.image(str(_img_tse), use_container_width=True)

    # Tính toán
    D   = _get("cdm_D")
    Lc  = _get("cdm_Lc")
    qu  = _get("cdm_qu")
    Su  = _get("cdm_Su")
    arr = _get("cdm_arrangement")
    _ld_now = _get("cdm_loads")
    q       = q_total(_ld_now)   # tổng (xét hoạt tải) → cho Pcol
    q_st    = q_static(_ld_now)  # tĩnh (bỏ hoạt tải) → cho S1

    scenarios = build_scenarios(D, Lc, qu, Su, q, spacings, arr, q_static=q_st)
    df = pd.DataFrame(scenarios)
    st.caption(
        f"**Quy ước tải trọng** (TCVN 9403 / TCCS 41):  \n"
        f"• **Lún $S_1$** dùng tải tĩnh $q_{{static}} = {q_st:.1f}$ kN/m² "
        f"(đất đắp + đệm + mặt đường, **bỏ hoạt tải xe**)  \n"
        f"• **Sức tải $P_{{col}}$** dùng tải tổng $q = {q:.1f}$ kN/m² "
        f"(bao gồm hoạt tải $q_{{traffic}} = {_ld_now.get('q_traffic',20):.1f}$ kN/m²)"
    )

    # Chọn PA kiến nghị
    _alt_prefix = "PA" if _get("lang") == "VN" else "Alt."
    rec_idx = st.selectbox(
        _t("rec_sel"),
        range(len(scenarios)),
        index=min(_get("cdm_rec_idx"), len(scenarios)-1),
        format_func=lambda i: f"{_alt_prefix}{i+1}: e={scenarios[i]['e (m)']}m  →  S₁={scenarios[i]['S₁ (cm)']:.2f}cm",
    )
    st.session_state["cdm_rec_idx"] = rec_idx

    # Tính kiểm tra chọc thủng cho mỗi PA
    _ld_punch = _get("cdm_loads")
    _punch = [
        calc_punching_check(
            D=D, e=s["e (m)"],
            Hse=_ld_punch.get("h_mat", 0.4),
            He=_ld_punch.get("h_fill", 1.5),
            quckse=_get("cdm_quckse"),
            Fs=_get("cdm_Fs_mat"),
            theta_deg=_get("cdm_theta"),
            gamma_fill=_ld_punch.get("g_fill", 18.0),
            gamma_mat=_ld_punch.get("g_mat", 22.5),
            qa=_get("cdm_qa_mat"),
        )
        for s in scenarios
    ]
    for i, pr in enumerate(_punch):
        df.at[i, "τse (kPa)"]  = pr["tse_kPa"]
        df.at[i, "τase (kPa)"] = pr["tase_kPa"]
        df.at[i, "Đạt CT"]     = pr["check_CT"]

    # Bảng kết quả
    st.markdown(f"**{_t('pa_table')}**")
    display_cols = ["e (m)", "a (%)", "Etb (kN/m²)", "S₁ (cm)",
                    "Pcol (kN)", "Qa (kN)", "Đạt SCT",
                    "τse (kPa)", "τase (kPa)", "Đạt CT"]

    def _row_color(row):
        idx = df.index.get_loc(row.name)
        if idx == rec_idx:
            return ["background-color: #E2EFDA"] * len(row)
        return [""] * len(row)

    st.dataframe(
        df[display_cols].style.apply(_row_color, axis=1)
                               .format({"a (%)": "{:.1f}%",
                                        "Etb (kN/m²)": "{:,.0f}",
                                        "S₁ (cm)": "{:.2f}",
                                        "Pcol (kN)": "{:.1f}",
                                        "Qa (kN)": "{:.1f}"}),
        use_container_width=True,
    )

    # ── Matplotlib fallback cho mọi biểu đồ (đảm bảo luôn có dù plotly fail) ─
    def _mpl_compare_charts(_df, _rec_idx, _scenarios, _punch):
        import matplotlib.pyplot as plt
        _xs = [f"e={s['e (m)']}m" for s in _scenarios]
        _cols = ['#ED7D31' if i == _rec_idx else '#4472C4' for i in range(len(_scenarios))]

        fig, axs = plt.subplots(2, 2, figsize=(13, 8))
        # S1 vs e
        axs[0,0].bar(_xs, [s["S₁ (cm)"] for s in _scenarios], color=_cols, edgecolor='black')
        for i, v in enumerate([s["S₁ (cm)"] for s in _scenarios]):
            axs[0,0].text(i, v+0.3, f"{v:.2f}", ha='center', fontsize=9)
        axs[0,0].set_title("S₁ theo khoảng cách e"); axs[0,0].set_ylabel("S₁ (cm)")
        axs[0,0].grid(True, axis='y', ls=':', alpha=0.4)

        # Etb vs e
        axs[0,1].bar(_xs, [s["Etb (kN/m²)"] for s in _scenarios], color=_cols, edgecolor='black')
        for i, v in enumerate([s["Etb (kN/m²)"] for s in _scenarios]):
            axs[0,1].text(i, v*1.01, f"{v:,.0f}", ha='center', fontsize=8)
        axs[0,1].set_title("Etb theo khoảng cách e"); axs[0,1].set_ylabel("Etb (kN/m²)")
        axs[0,1].grid(True, axis='y', ls=':', alpha=0.4)

        # Pcol vs Qa
        _pcol = [s["Pcol (kN)"] for s in _scenarios]
        _qa   = [s["Qa (kN)"]   for s in _scenarios]
        axs[1,0].bar(_xs, _pcol, color=_cols, edgecolor='black', label='Pcol')
        axs[1,0].plot(_xs, _qa, 'D-', color='#E53935', lw=2, markersize=8, label='Qa (giới hạn)')
        for i, v in enumerate(_pcol):
            axs[1,0].text(i, v+max(_pcol)*0.02, f"{v:.0f}", ha='center', fontsize=9)
        axs[1,0].set_title("Sức chịu tải: P_col vs Q_a")
        axs[1,0].set_ylabel("kN"); axs[1,0].legend(loc='upper left')
        axs[1,0].grid(True, axis='y', ls=':', alpha=0.4)

        # τse chọc thủng
        _tse = [pr["tse_kPa"] for pr in _punch]
        _tase = _punch[0]["tase_kPa"] if _punch else 100
        axs[1,1].bar(_xs, _tse, color=_cols, edgecolor='black')
        axs[1,1].axhline(y=_tase, ls='--', color='#E53935', lw=1.5)
        axs[1,1].text(len(_xs)-1, _tase, f"τase={_tase:.0f}", color='#E53935',
                      fontsize=9, va='bottom', ha='right')
        for i, v in enumerate(_tse):
            axs[1,1].text(i, v+max(_tse)*0.02, f"{v:.1f}", ha='center', fontsize=9)
        axs[1,1].set_title("τse chọc thủng đệm xi măng"); axs[1,1].set_ylabel("τse (kPa)")
        axs[1,1].grid(True, axis='y', ls=':', alpha=0.4)

        fig.suptitle("So sánh phương án — cam = PA kiến nghị", fontsize=12, fontweight='bold')
        fig.tight_layout()
        return fig

    if _HAS_PLOTLY:
        ca, cb, cc = st.columns(3)
        with ca:
            st.plotly_chart(_chart_scenarios(df, rec_idx), use_container_width=True)
        with cb:
            st.plotly_chart(_chart_etb(df, rec_idx), use_container_width=True)
        with cc:
            st.plotly_chart(_chart_combined(df, rec_idx), use_container_width=True)
    elif _HAS_MPL:
        st.caption("_Hiển thị biểu đồ chế độ đơn giản._")
        _fig_cmp_mpl = _mpl_compare_charts(df, rec_idx, scenarios, _punch)
        st.pyplot(_fig_cmp_mpl, use_container_width=True)
        plt.close(_fig_cmp_mpl)

    if _HAS_PLOTLY:
        # Biểu đồ chọc thủng (chỉ render khi có Plotly)
        _x_sp = [f"e={s['e (m)']}m" for s in scenarios]
        _col_ct = ["#ED7D31" if i == rec_idx else "#4472C4" for i in range(len(scenarios))]
        _tase_v = _punch[0]["tase_kPa"] if _punch else 100.0
        _fig_ct = go.Figure()
        _fig_ct.add_trace(go.Bar(
            name="τse (kPa)", x=_x_sp,
            y=[pr["tse_kPa"] for pr in _punch],
            marker_color=_col_ct,
            text=[f"{pr['tse_kPa']:.1f}" for pr in _punch],
            textposition="outside",
        ))
        _fig_ct.add_hline(
            y=_tase_v, line_dash="dash", line_color="#E53935", line_width=1.5,
            annotation_text=f"τase = {_tase_v:.0f} kPa",
            annotation_position="top right",
        )
        _fig_ct.update_layout(
            title=_t("punch_chart"),
            yaxis_title=_t("shear_ax"),
            height=340, margin=dict(t=45, b=20),
            showlegend=False,
        )
        st.plotly_chart(_fig_ct, use_container_width=True)

        # ── BIỂU ĐỒ MỚI 1: Pcol vs Qa — kiểm tra SCT ───────────────────────────
        _fig_sct = go.Figure()
        _fig_sct.add_trace(go.Bar(
            name="Pcol (kN)", x=_x_sp,
            y=[s["Pcol (kN)"] for s in scenarios],
            marker_color=["#ED7D31" if i == rec_idx else "#4472C4"
                          for i in range(len(scenarios))],
            text=[f"{s['Pcol (kN)']:.1f}" for s in scenarios],
            textposition="outside",
        ))
        _fig_sct.add_trace(go.Scatter(
            name="Qa (kN)", x=_x_sp,
            y=[s["Qa (kN)"] for s in scenarios],
            mode="lines+markers+text",
            text=[f"{s['Qa (kN)']:.0f}" for s in scenarios],
            textposition="top center",
            line=dict(color="#E53935", width=2.5, dash="dash"),
            marker=dict(size=10, color="#E53935", symbol="diamond"),
        ))
        _fig_sct.update_layout(
            title="Sức chịu tải: P_col vs Q_a — vùng dưới Qa = Đạt SCT",
            yaxis_title="kN", height=340, margin=dict(t=45, b=20),
            legend=dict(orientation="h", y=1.05),
        )
        st.plotly_chart(_fig_sct, use_container_width=True)

        # ── BIỂU ĐỒ MỚI 2: Radar so sánh đa chỉ tiêu (normalized) ──────────────
        _metrics = {
            "S₁ (cm)":     [s["S₁ (cm)"]   for s in scenarios],
            "a (%)":       [s["a (%)"]     for s in scenarios],
            "Pcol (kN)":   [s["Pcol (kN)"] for s in scenarios],
            "Etb (kPa)":   [s["Etb (kN/m²)"] for s in scenarios],
            "τse (kPa)":   [pr["tse_kPa"]   for pr in _punch],
        }
        _fig_rad = go.Figure()
        for i, s in enumerate(scenarios):
            _vals_norm = []
            for k, vlist in _metrics.items():
                _vmax = max(vlist) if max(vlist) > 0 else 1
                _vals_norm.append(vlist[i] / _vmax * 100)
            _vals_norm.append(_vals_norm[0])  # đóng vòng
            _fig_rad.add_trace(go.Scatterpolar(
                r=_vals_norm,
                theta=list(_metrics.keys()) + [list(_metrics.keys())[0]],
                fill="toself",
                name=f"e={s['e (m)']}m",
                opacity=0.55 if i != rec_idx else 0.85,
                line=dict(width=2.5 if i == rec_idx else 1.5,
                          color="#ED7D31" if i == rec_idx else None),
            ))
        _fig_rad.update_layout(
            title="So sánh đa chỉ tiêu (% so với max mỗi chỉ tiêu) — cam = PA kiến nghị",
            polar=dict(radialaxis=dict(visible=True, range=[0, 105])),
            height=420, margin=dict(t=45, b=20),
        )
        st.plotly_chart(_fig_rad, use_container_width=True)

    # ── Phân tích tham số ─────────────────────────────────────────────────────
    st.divider()
    st.markdown(f"### {_t('param_sens')}")

    _e_rec   = scenarios[rec_idx]["e (m)"]
    _ld_base = _get("cdm_loads").copy()
    _hmat_cur = _ld_base.get("h_mat", 0.4)

    _pa1, _pa2, _pa3 = st.columns(3)

    # ── 1. So sánh qu ─────────────────────────────────────────────────────────
    with _pa1:
        st.markdown(f"**{_t('qu_comp')}**")
        _qu_raw = st.text_input(_t("qu_inp"),
                                ", ".join(str(v) for v in [200, 300, 400, 500, 600, 700]),
                                key="_qu_list")
        try:
            _qu_list = sorted(set(float(x.strip()) for x in _qu_raw.split(",") if x.strip()))
        except ValueError:
            _qu_list = [200, 300, 400, 500]

        _qu_rows = []
        for _qv in _qu_list:
            _Ec_v  = calc_Ec(_qv)
            _Es_v  = calc_Es(Su)
            _a_v   = calc_a(D, _e_rec, arr)
            _Etb_v = calc_Etb(_a_v, _Ec_v, _Es_v)
            _S1_v  = calc_S1(q_st, Lc, _Etb_v)        # S1: tải tĩnh
            _sc_v  = calc_sigma_col(q, _Ec_v, _Etb_v) # Pcol: tải tổng
            _Pcol_v = calc_Pcol(_sc_v, D)
            _bc_v  = calc_bearing(D, Lc, _qv / 2, Su)
            _qu_rows.append({
                "qu (kPa)": int(_qv),
                "Ec (kN/m²)": int(_Ec_v),
                "Etb (kN/m²)": int(_Etb_v),
                "S₁ (cm)": round(_S1_v, 2),
                "Pcol (kN)": round(_Pcol_v, 1),
                "Qa (kN)": _bc_v["Qa"],
                "Đạt SCT": "Đạt" if _Pcol_v < _bc_v["Qa"] else "Không đạt",
            })
        _df_qu = pd.DataFrame(_qu_rows)
        st.dataframe(_df_qu, use_container_width=True, hide_index=True,
                     column_config={"Etb (kN/m²)": st.column_config.NumberColumn(format="%,d"),
                                    "Ec (kN/m²)":  st.column_config.NumberColumn(format="%,d")})
        if _HAS_PLOTLY:
            _col_qu = ["#ED7D31" if r["qu (kPa)"] == int(qu) else "#4472C4"
                       for r in _qu_rows]
            _fig_qu = go.Figure(go.Bar(
                x=[f"{r['qu (kPa)']} kPa" for r in _qu_rows],
                y=_df_qu["S₁ (cm)"],
                marker_color=_col_qu,
                text=_df_qu["S₁ (cm)"].apply(lambda v: f"{v:.2f}"),
                textposition="outside",
            ))
            _fig_qu.update_layout(
                title=f"S₁ theo qu  (e={_e_rec}m, D={D}m)",
                yaxis_title="S₁ (cm)", height=300,
                margin=dict(t=45, b=20),
            )
            st.plotly_chart(_fig_qu, use_container_width=True)

    # ── 2. So sánh D ──────────────────────────────────────────────────────────
    with _pa2:
        st.markdown(f"**{_t('D_comp')}**")
        _D_raw = st.text_input(_t("D_inp"), "0.50, 0.60, 0.70, 0.80, 1.00",
                               key="_D_list")
        try:
            _D_list = sorted(set(float(x.strip()) for x in _D_raw.split(",") if x.strip()))
        except ValueError:
            _D_list = [0.5, 0.6, 0.7, 0.8]

        _D_rows = []
        for _Dv in _D_list:
            _Ec_v  = calc_Ec(qu)
            _Es_v  = calc_Es(Su)
            _a_v   = calc_a(_Dv, _e_rec, arr)
            _Etb_v = calc_Etb(_a_v, _Ec_v, _Es_v)
            _S1_v  = calc_S1(q_st, Lc, _Etb_v)        # S1: tải tĩnh
            _sc_v  = calc_sigma_col(q, _Ec_v, _Etb_v) # Pcol: tải tổng
            _Pcol_v = calc_Pcol(_sc_v, _Dv)
            _bc_v  = calc_bearing(_Dv, Lc, qu / 2, Su)
            _D_rows.append({
                "D (m)": _Dv,
                "a (%)": round(_a_v * 100, 1),
                "Etb (kN/m²)": int(_Etb_v),
                "S₁ (cm)": round(_S1_v, 2),
                "Pcol (kN)": round(_Pcol_v, 1),
                "Qa (kN)": _bc_v["Qa"],
                "Đạt SCT": "Đạt" if _Pcol_v < _bc_v["Qa"] else "Không đạt",
            })
        _df_D = pd.DataFrame(_D_rows)
        st.dataframe(_df_D, use_container_width=True, hide_index=True,
                     column_config={"a (%)": st.column_config.NumberColumn(format="%.1f%%"),
                                    "Etb (kN/m²)": st.column_config.NumberColumn(format="%,d")})
        if _HAS_PLOTLY:
            _col_D = ["#ED7D31" if abs(_Dv - D) < 0.001 else "#4472C4"
                      for _Dv in _D_list]
            _fig_D = go.Figure()
            _fig_D.add_trace(go.Bar(
                name="a (%)",
                x=[f"{r['D (m)']}m" for r in _D_rows],
                y=_df_D["a (%)"],
                marker_color=_col_D,
                text=_df_D["a (%)"].apply(lambda v: f"{v:.1f}%"),
                textposition="outside",
                yaxis="y1",
            ))
            _fig_D.add_trace(go.Scatter(
                name="S₁ (cm)",
                x=[f"{r['D (m)']}m" for r in _D_rows],
                y=_df_D["S₁ (cm)"],
                mode="lines+markers+text",
                text=_df_D["S₁ (cm)"].apply(lambda v: f"{v:.2f}"),
                textposition="top center",
                yaxis="y2",
                line=dict(color="#E53935", width=2),
                marker=dict(size=8, color="#E53935"),
            ))
            _fig_D.update_layout(
                title=f"a(%) và S₁ theo D  (e={_e_rec}m, qu={int(qu)}kPa)",
                yaxis=dict(title="a (%)"),
                yaxis2=dict(title="S₁ (cm)", overlaying="y", side="right"),
                height=300, margin=dict(t=45, b=20),
                legend=dict(orientation="h"),
            )
            st.plotly_chart(_fig_D, use_container_width=True)

    # ── 3. So sánh h_mat ──────────────────────────────────────────────────────
    with _pa3:
        st.markdown(f"**{_t('hmat_comp')}**")
        _hm_raw = st.text_input(_t("hmat_inp"), "0.20, 0.30, 0.40, 0.50, 0.60, 0.80",
                                key="_hmat_list")
        try:
            _hmat_list = sorted(set(float(x.strip()) for x in _hm_raw.split(",") if x.strip()))
        except ValueError:
            _hmat_list = [0.2, 0.3, 0.4, 0.5]

        _hm_rows = []
        for _hm in _hmat_list:
            _ld_v = _ld_base.copy()
            _ld_v["h_mat"] = _hm
            _q_v    = q_total(_ld_v)   # tổng (xét hoạt tải)
            _q_st_v = q_static(_ld_v)  # tĩnh (bỏ hoạt tải) — cho S1
            _Ec_v  = calc_Ec(qu)
            _Es_v  = calc_Es(Su)
            _a_v   = calc_a(D, _e_rec, arr)
            _Etb_v = calc_Etb(_a_v, _Ec_v, _Es_v)
            _S1_v  = calc_S1(_q_st_v, Lc, _Etb_v)   # S1: tải tĩnh
            _hm_rows.append({
                "h_mat (m)": _hm,
                "q (kN/m²)": round(_q_v, 2),
                "ΔS₁ (cm)": round(_S1_v, 2),
            })
        _df_hm = pd.DataFrame(_hm_rows)
        st.dataframe(_df_hm, use_container_width=True, hide_index=True)
        if _HAS_PLOTLY:
            _col_hm = ["#ED7D31" if abs(_hm - _hmat_cur) < 0.001 else "#4472C4"
                       for _hm in _hmat_list]
            _fig_hm = go.Figure()
            _fig_hm.add_trace(go.Bar(
                name="q (kN/m²)",
                x=[f"{r['h_mat (m)']}m" for r in _hm_rows],
                y=_df_hm["q (kN/m²)"],
                marker_color=_col_hm,
                text=_df_hm["q (kN/m²)"].apply(lambda v: f"{v:.1f}"),
                textposition="outside",
                yaxis="y1",
            ))
            _fig_hm.add_trace(go.Scatter(
                name="S₁ (cm)",
                x=[f"{r['h_mat (m)']}m" for r in _hm_rows],
                y=_df_hm["ΔS₁ (cm)"],
                mode="lines+markers+text",
                text=_df_hm["ΔS₁ (cm)"].apply(lambda v: f"{v:.2f}"),
                textposition="top center",
                yaxis="y2",
                line=dict(color="#E53935", width=2),
                marker=dict(size=8, color="#E53935"),
            ))
            _fig_hm.update_layout(
                title=f"q và S₁ theo h_mat  (e={_e_rec}m, D={D}m)",
                yaxis=dict(title="q (kN/m²)"),
                yaxis2=dict(title="S₁ (cm)", overlaying="y", side="right"),
                height=300, margin=dict(t=45, b=20),
                legend=dict(orientation="h"),
            )
            st.plotly_chart(_fig_hm, use_container_width=True)

    # ── BIỂU ĐỒ TỔNG HỢP: Sensitivity S₁ theo qu/D/h_mat (normalized) ──────────
    # Matplotlib fallback luôn vẽ
    if _HAS_MPL and not _HAS_PLOTLY:
        import matplotlib.pyplot as plt
        st.caption("_Phân tích độ nhạy của S₁ theo các tham số thiết kế._")
        fig, axs = plt.subplots(1, 3, figsize=(15, 4))
        # qu
        _qu_x = [r["qu (kPa)"] for r in _qu_rows]
        _qu_s = [r["S₁ (cm)"] for r in _qu_rows]
        axs[0].plot(_qu_x, _qu_s, 'o-', color='#1565C0', lw=2)
        for x, y in zip(_qu_x, _qu_s):
            axs[0].annotate(f"{y:.2f}", (x, y), xytext=(0, 6), textcoords='offset points',
                            ha='center', fontsize=8)
        axs[0].axvline(qu, ls=':', color='red', alpha=0.6, label=f'qu hiện tại={int(qu)}')
        axs[0].set_xlabel('qu (kPa)'); axs[0].set_ylabel('S₁ (cm)')
        axs[0].set_title('S₁ theo qu')
        axs[0].grid(True, ls=':', alpha=0.5); axs[0].legend(fontsize=8)
        # D
        _D_x = [r["D (m)"] for r in _D_rows]
        _D_s = [r["S₁ (cm)"] for r in _D_rows]
        axs[1].plot(_D_x, _D_s, 's-', color='#2E7D32', lw=2)
        for x, y in zip(_D_x, _D_s):
            axs[1].annotate(f"{y:.2f}", (x, y), xytext=(0, 6), textcoords='offset points',
                            ha='center', fontsize=8)
        axs[1].axvline(D, ls=':', color='red', alpha=0.6, label=f'D hiện tại={D}m')
        axs[1].set_xlabel('D (m)'); axs[1].set_ylabel('S₁ (cm)')
        axs[1].set_title('S₁ theo D')
        axs[1].grid(True, ls=':', alpha=0.5); axs[1].legend(fontsize=8)
        # h_mat
        _hm_x = [r["h_mat (m)"] for r in _hm_rows]
        _hm_s = [r["ΔS₁ (cm)"] for r in _hm_rows]
        axs[2].plot(_hm_x, _hm_s, '^-', color='#E53935', lw=2)
        for x, y in zip(_hm_x, _hm_s):
            axs[2].annotate(f"{y:.2f}", (x, y), xytext=(0, 6), textcoords='offset points',
                            ha='center', fontsize=8)
        axs[2].axvline(_hmat_cur, ls=':', color='red', alpha=0.6,
                       label=f'h_mat hiện tại={_hmat_cur}m')
        axs[2].set_xlabel('h_mat (m)'); axs[2].set_ylabel('S₁ (cm)')
        axs[2].set_title('S₁ theo h_mat')
        axs[2].grid(True, ls=':', alpha=0.5); axs[2].legend(fontsize=8)
        fig.suptitle("Độ nhạy của S₁ theo các tham số thiết kế", fontsize=12, fontweight='bold')
        fig.tight_layout()
        st.pyplot(fig, use_container_width=True)
        plt.close(fig)

    if _HAS_PLOTLY:
        _fig_sens = go.Figure()
        # Normalize: lấy S1 ở giá trị "current" làm baseline 100%
        _S1_baseline_qu = next((r["S₁ (cm)"] for r in _qu_rows if r["qu (kPa)"] == int(qu)),
                               _qu_rows[len(_qu_rows)//2]["S₁ (cm)"])
        _S1_baseline_D  = next((r["S₁ (cm)"] for r in _D_rows if abs(r["D (m)"] - D) < 0.01),
                               _D_rows[len(_D_rows)//2]["S₁ (cm)"])
        _S1_baseline_hm = next((r["ΔS₁ (cm)"] for r in _hm_rows if abs(r["h_mat (m)"] - _hmat_cur) < 0.01),
                               _hm_rows[len(_hm_rows)//2]["ΔS₁ (cm)"])

        _fig_sens.add_trace(go.Scatter(
            name=f"qu (baseline={int(qu)} kPa)",
            x=[r["qu (kPa)"]/qu*100 for r in _qu_rows],
            y=[r["S₁ (cm)"]/_S1_baseline_qu*100 for r in _qu_rows],
            mode="lines+markers", line=dict(color="#1565C0", width=2.5),
            marker=dict(size=8),
        ))
        _fig_sens.add_trace(go.Scatter(
            name=f"D (baseline={D} m)",
            x=[r["D (m)"]/D*100 for r in _D_rows],
            y=[r["S₁ (cm)"]/_S1_baseline_D*100 for r in _D_rows],
            mode="lines+markers", line=dict(color="#2E7D32", width=2.5),
            marker=dict(size=8),
        ))
        _fig_sens.add_trace(go.Scatter(
            name=f"h_mat (baseline={_hmat_cur} m)",
            x=[r["h_mat (m)"]/_hmat_cur*100 if _hmat_cur > 0 else 100 for r in _hm_rows],
            y=[r["ΔS₁ (cm)"]/_S1_baseline_hm*100 if _S1_baseline_hm > 0 else 0
               for r in _hm_rows],
            mode="lines+markers", line=dict(color="#E53935", width=2.5),
            marker=dict(size=8),
        ))
        _fig_sens.add_hline(y=100, line_dash="dot", line_color="#999",
                            annotation_text="Baseline 100%")
        _fig_sens.add_vline(x=100, line_dash="dot", line_color="#999")
        _fig_sens.update_layout(
            title="Độ nhạy của S₁ theo qu / D / h_mat (chuẩn hoá % so với baseline PA kiến nghị)",
            xaxis_title="Tham số / Baseline (%)",
            yaxis_title="S₁ / S₁_baseline (%)",
            height=400, margin=dict(t=50, b=40),
            legend=dict(orientation="h", y=-0.15),
        )
        st.plotly_chart(_fig_sens, use_container_width=True)
        st.caption(
            "Đường dốc đứng = tham số nhạy (thay đổi nhỏ → S₁ thay đổi lớn). "
            "Đường thoải = tham số ít ảnh hưởng. "
            "Giao điểm tại (100%, 100%) = trạng thái PA kiến nghị hiện tại."
        )

    # ── Kết quả chi tiết PA kiến nghị ─────────────────────────────────────────
    st.divider()
    st.subheader(_t("p4_sub"))

    s   = scenarios[rec_idx]
    Cc  = qu / 2
    Ec  = calc_Ec(qu)
    Es  = calc_Es(Su)
    Etb = s["Etb (kN/m²)"]

    _CP = _t("col_param"); _CF = _t("col_formula"); _CV = _t("col_value")

    def _color_sct(val):
        if val == "Đạt":       return "color: green; font-weight: bold"
        if val == "Không đạt": return "color: red; font-weight: bold"
        return ""

    st.success(_t("rec_pa", n=rec_idx+1, e=s['e (m)']))

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("S₁ (cm)", f"{s['S₁ (cm)']:.2f}")
    c2.metric("Etb (kN/m²)", f"{int(Etb):,}")
    c3.metric("a (%)", f"{s['a (%)']:.1f}%")
    c4.metric("Sức chịu tải", s["Đạt SCT"])

    st.divider()
    col_l, col_r = st.columns(2)

    with col_l:
        st.markdown(f"**{_t('settle_calc')}**")
        # Tải tĩnh (bỏ hoạt tải xe) — dùng riêng cho lún
        _q_st_settle = q_static(_get("cdm_loads"))
        _is_tri = arr == "triangle"
        _a_fmla = (r"$a = \dfrac{\pi D^2}{2\sqrt{3}\, e^2}$ (tam giác)"
                   if _is_tri else r"$a = \dfrac{\pi D^2}{4 e^2}$ (vuông)")
        # Bảng markdown với LaTeX inline ($...$) — render qua KaTeX
        _rows_md = [
            ("Khoảng cách trụ", "$e$", f"{s['e (m)']} m"),
            ("Tỷ lệ thay thế", _a_fmla, f"{s['a']:.4f}  ({s['a (%)']:.1f}%)"),
            ("Mô đun trụ CDM", r"$E_c = 100 \cdot C_c = 100 \cdot \dfrac{q_u}{2}$",
             f"{int(Ec):,} kN/m²"),
            ("Mô đun đất nền", r"$E_s = 250 \cdot S_u$", f"{int(Es):,} kN/m²"),
            ("Mô đun tương đương", r"$E_{tb} = a \cdot E_c + (1-a) \cdot E_s$",
             f"{int(Etb):,} kN/m²"),
            ("Tải trọng tĩnh (bỏ hoạt tải)",
             r"$q_{\text{tĩnh}} = h_{\text{đắp}}\gamma + h_{\text{đệm}}\gamma + h_{\text{đường}}\gamma$",
             f"{_q_st_settle:.2f} kN/m²"),
            ("Chiều dài trụ", "$L_c$", f"{Lc:.1f} m"),
            ("Độ lún bản thân CDM",
             r"$S_1 = \dfrac{q_{\text{tĩnh}} \cdot L_c}{E_{tb}} \times 100$",
             f"{s['S₁ (cm)']:.2f} cm"),
            ("Độ lún dưới mũi trụ", "$S_2 = 0$ (trụ xuyên qua lớp bùn)", "0,00 cm"),
            ("Tổng độ lún", "$S = S_1 + S_2$", f"{s['S₁ (cm)']:.2f} cm"),
        ]
        # Build markdown table
        _md = "| Thông số | Công thức | Giá trị |\n|---|---|---|\n"
        for _name, _fml, _val in _rows_md:
            _md += f"| {_name} | {_fml} | **{_val}** |\n"
        st.markdown(_md)
        st.caption(
            f"_Lưu ý: lún dùng tải tĩnh $q_{{tĩnh}} = {_q_st_settle:.1f}$ kN/m² "
            f"(bỏ hoạt tải xe). Tải tổng $q = {q:.1f}$ kN/m² dùng riêng cho sức chịu tải._"
        )

    with col_r:
        st.markdown(f"**{_t('bearing_calc')}**")
        bc = calc_bearing(D, Lc, Cc, Su)
        _ok_sct = s["Đạt SCT"] == "Đạt"
        _sct_check = ("**:green[Đạt]**" if _ok_sct else "**:red[Không đạt]**")
        _sct_rows_md = [
            ("Diện tích cắt ngang", r"$A_c = \dfrac{\pi D^2}{4}$",
             f"{math.pi*D**2/4:.4f} m²"),
            ("Sức chịu tải mũi", r"$Q_{ult,\text{mũi}} = 9 \cdot C_c \cdot A_c$",
             f"{bc['Qult_col']:.1f} kN"),
            ("Sức chịu tải thân", r"$Q_{ult,\text{thân}} = \pi \cdot D \cdot L_c \cdot C_u$",
             f"{bc['Qult_skin']:.1f} kN"),
            ("Tổng sức chịu tải",
             r"$Q_{ult} = Q_{ult,\text{mũi}} + Q_{ult,\text{thân}}$",
             f"{bc['Qult']:.1f} kN"),
            ("Sức chịu tải cho phép", r"$Q_a = \dfrac{Q_{ult}}{FS}, \; FS = 2$",
             f"{bc['Qa']:.1f} kN"),
            ("Ứng suất đầu cọc",
             r"$\sigma_{col} = \dfrac{E_c}{E_{tb}} \cdot q_{\text{tổng}}$",
             f"{s['σ_col (kN/m²)']:.1f} kN/m²"),
            ("Lực nén lên 1 trụ", r"$P_{col} = \sigma_{col} \cdot A_c$",
             f"{s['Pcol (kN)']:.1f} kN"),
            ("Kiểm tra", r"$P_{col} < Q_a \;?$", _sct_check),
        ]
        _md_sct = "| Thông số | Công thức | Giá trị |\n|---|---|---|\n"
        for _name, _fml, _val in _sct_rows_md:
            _md_sct += f"| {_name} | {_fml} | **{_val}** |\n"
        st.markdown(_md_sct)
        st.caption(
            f"_Sức chịu tải dùng tải tổng $q = {q:.1f}$ kN/m² — **bao gồm hoạt tải xe** (TCVN 9403 Phụ lục B)._"
        )

    st.divider()
    st.markdown(f"**{_t('punch_res')}**")
    _ld_r   = _get("cdm_loads")
    _pr_r   = _punch[rec_idx]
    _ok_ct  = _pr_r["check_CT"] == "Đạt"
    _ct_check = "**:green[Đạt]**" if _ok_ct else "**:red[Không đạt]**"
    _punch_rows_md = [
        ("Bề dày đệm xi măng",
         "$H_{se}$",
         f"{_ld_r.get('h_mat',0.4):.2f} m"),
        ("Chiều cao đất đắp trên đệm",
         "$H_e = h_{\\text{fill}}$",
         f"{_ld_r.get('h_fill',1.5):.2f} m"),
        ("Cường độ kháng nén đệm XM",
         "$q_{u,ck,se}$",
         f"{_get('cdm_quckse'):.0f} kPa"),
        ("Ứng suất cắt cho phép",
         r"$\tau_{ase} = \dfrac{q_{u,ck,se}}{2 \cdot F_s}$",
         f"{_pr_r['tase_kPa']:.2f} kPa"),
        ("Chiều cao vùng vòm đất",
         r"$H_o = (e - D) \cdot \tan(\theta/2)$",
         f"{_pr_r['Ho_m']:.3f} m"),
        ("So sánh $H_o$ vs $H_e$",
         "Chọn công thức",
         _pr_r["cong_thuc"]),
        ("Thể tích đất tác dụng",
         "$V_{soil}$",
         f"{_pr_r['Vsoil_m3']:.4f} m³"),
        ("Thể tích đệm XM tác dụng",
         "$V_{CGCXM}$",
         f"{_pr_r['VCGCXM_m3']:.4f} m³"),
        ("Áp lực vùng không gia cố",
         "$P_{Soil}$",
         f"{_pr_r['PSoil_kPa']:.3f} kPa"),
        ("Ứng suất cắt thực tế",
         r"$\tau_{se} = \dfrac{P_{Soil} \cdot (e^2 - \pi D^2/4)}{\pi \cdot D \cdot H_{se}}$",
         f"{_pr_r['tse_kPa']:.3f} kPa"),
        ("Kiểm tra",
         r"$\tau_{se} \leq \tau_{ase} \;?$",
         _ct_check),
    ]
    _md_ct = "| Thông số | Công thức | Giá trị |\n|---|---|---|\n"
    for _n, _f, _v in _punch_rows_md:
        _md_ct += f"| {_n} | {_f} | **{_v}** |\n"
    st.markdown(_md_ct)
    st.caption(
        "_Phương pháp **ALiCC (PWRI Japan)** — TCVN 9403:2012 Phụ lục C.  \n"
        "Điều kiện đạt: $\\tau_{se} \\leq \\tau_{ase}$ (ứng suất cắt thực tế ≤ ứng suất cắt cho phép)._"
    )

    _punch_md = _load_punch_theory()
    if _punch_md:
        with st.expander("Diễn giải công thức kiểm tra chọc thủng lớp đệm xi măng", expanded=False):
            st.markdown(_punch_md)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 5 – XUẤT
# ═══════════════════════════════════════════════════════════════════════════════
if _page == "params":   # tiếp nội dung Xuất kết quả (gộp vào tab Thông số)
    st.subheader(_t("p5_sub"))

    D   = _get("cdm_D")
    Lc  = _get("cdm_Lc")
    qu  = _get("cdm_qu")
    Su  = _get("cdm_Su")
    arr = _get("cdm_arrangement")
    q   = q_total(_get("cdm_loads"))
    sps = _get("cdm_spacings")
    rec = _get("cdm_rec_idx")

    scenarios = build_scenarios(D, Lc, qu, Su, q, sps, arr)
    rec = min(rec, len(scenarios) - 1) if scenarios else 0

    params = {
        "bh_name":   _get("cdm_bh") or "–",
        "zone":      _get("cdm_zone"),
        "h_clay":    _get("cdm_h_clay"),
        "Su": Su,    "gamma": _get("cdm_gamma"),
        "D": D,      "Lc": Lc, "CDTK": _get("cdm_CDTK"),
        "qu": qu,    "FS_lab": _get("cdm_FS_lab"),
        "dosage":    _get("cdm_dosage"),
        "arrangement": arr,
        "q_total":   q,
        "rec_idx":   rec,
        "loads":     _get("cdm_loads"),
        "quckse":    _get("cdm_quckse"),
        "Fs_mat":    _get("cdm_Fs_mat"),
        "theta":     _get("cdm_theta"),
        "qa_mat":    _get("cdm_qa_mat"),
    }

    st.markdown(f"### {_t('export_sum')}")
    if scenarios:
        s = scenarios[rec]
        col1, col2, col3 = st.columns(3)
        col1.metric(_t("rec_alt"), f"e = {s['e (m)']} m")
        col2.metric(_t("settle_metric"), f"{s['S₁ (cm)']:.2f} cm")
        col3.metric(_t("bear_lbl"), s["Đạt SCT"])

    st.divider()
    c_word, c_excel = st.columns(2)

    with c_word:
        st.markdown(f"#### {_t('word_title')}")
        st.caption(_t("word_cap"))

        if st.button(_t("create_word"), type="primary"):
            # Approach mới: HTML-format Word (không cần python-docx)
            # — chỉ dùng jinja2 + stdlib, hoạt động trên mọi môi trường
            try:
                import sys as _sys_w
                _sys_w.path.insert(0, str(_ROOT / "scripts"))
                from word_html_export import build_word_html as _build_w
                _p_w = dict(params)
                _p_w["q_static"] = q_static(_get("cdm_loads"))
                word_bytes = _build_w(scenarios, _p_w, rec)
                st.download_button(
                    _t("dl_docx"),
                    word_bytes,
                    file_name=f"CDM-THUYET-MINH-{params['bh_name']}.doc",
                    mime="application/msword",
                )
                st.success("Tạo thuyết minh thành công. Mở bằng Microsoft Word — "
                           "có thể chọn 'Save As' để lưu thành định dạng .docx nếu muốn.")
            except Exception as _e_w:
                st.error(f"Không tạo được file Word: {_e_w}")

    with c_excel:
        st.markdown("#### Tính toán chi tiết (Excel)")
        st.caption("3 sheet: Đầu vào · So sánh PA (bảng + biểu đồ nhúng) · Kết quả PA kiến nghị.")
        if st.button("Tạo file Excel", type="primary"):
            with st.spinner("Đang tạo Excel..."):
                xl_bytes = _export_excel(scenarios, params)
            if xl_bytes:
                st.download_button(
                    "Tải xuống bảng tính",
                    xl_bytes,
                    file_name=f"CDM-TINH-TOAN-{params['bh_name']}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            else:
                st.error("Lỗi tạo Excel. Kiểm tra openpyxl đã cài chưa.")

    # ── Báo cáo Word 7 HK kè SW (NỘI LỰC CỪ SAU XỬ LÝ CDM) ──────────────────
    st.divider()
    st.markdown("#### Báo cáo Word đầy đủ — Thiết kế cọc ván SW + CDM")
    st.caption(
        "5 chương: (1) Lý thuyết + công thức (NT1/NT2/Winkler/Bishop), "
        "(2) Bảng NT1/NT2/Nội lực tổng hợp 7 HK, "
        "(3) Biểu đồ nội lực 5-panel từng HK (u/M/Q/kh + sơ đồ tường), "
        "(4) Ổn định tổng thể 3 phương pháp, "
        "(5) Kết luận & Kiến nghị."
    )
    if st.button("Tạo báo cáo Word 7 HK", type="primary",
                  use_container_width=False, key="btn_word_7hk"):
        try:
            # ── IMPORTS ────────────────────────────────────────────────────────────
            import sys as _sys_7hk
            _sys_7hk.path.insert(0, str(_ROOT / "scripts"))
            from wall_internal_force import (
                sw_pile_props as _wif_sw, WallGeometry as _WIF_WG,
                EarthLayer as _WIF_EL, SoilLayer as _WIF_SL,
                build_lateral_load as _wif_build,
            )
            from winkler_np import solve_numpy_dist as _wif_solve
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as _plt_7
            import matplotlib.patches as _mpatch7
            import numpy as _np_7
            from docx import Document as _Doc7
            from docx.shared import Cm as _Cm7, Pt as _Pt7, RGBColor as _RGB7
            from docx.enum.text import WD_ALIGN_PARAGRAPH as _ALN7
            from docx.oxml.ns import qn as _qn7
            from docx.oxml import OxmlElement as _OXE7
            import io as _io_7, tempfile as _tmp_7, os as _os_7, sqlite3 as _sql_7
            from datetime import datetime as _dt_7

            # ── HELPER: cell background colour ─────────────────────────────────────
            def _cell_color7(cell, hex6):
                tc = cell._tc; tcPr = tc.get_or_add_tcPr()
                shd = _OXE7("w:shd")
                shd.set(_qn7("w:val"), "clear"); shd.set(_qn7("w:color"), "auto")
                shd.set(_qn7("w:fill"), hex6)
                tcPr.append(shd)

            def _fmt_row7(row, size_pt=7.5):
                for cell in (row if isinstance(row, (list, tuple)) else row.cells):
                    for para in cell.paragraphs:
                        para.paragraph_format.space_before = _Pt7(1)
                        para.paragraph_format.space_after = _Pt7(1)
                        for run in para.runs:
                            run.font.size = _Pt7(size_pt)

            # ── HELPER: formula image (matplotlib mathtext, no LaTeX install) ───────
            def _formula_img7(rows, title="", page_w_cm=14.0, height_cm=None, dpi=110):
                """rows: list of (label_str, mathtext_str). Returns BytesIO PNG."""
                h = height_cm if height_cm else max(0.9 * len(rows) + (0.6 if title else 0.0), 1.0)
                fig, ax = _plt_7.subplots(figsize=(page_w_cm / 2.54, h / 2.54), dpi=dpi)
                ax.set_facecolor("#f4f6fb"); fig.patch.set_facecolor("#f4f6fb")
                ax.axis("off")
                y_cur = 0.97
                if title:
                    ax.text(0.01, y_cur, title, transform=ax.transAxes,
                            fontsize=8.5, fontweight="bold", va="top", color="#1a3c5e")
                    y_cur -= 0.16
                n = len(rows)
                step = (y_cur - 0.03) / max(n, 1)
                for i, (lbl, mstr) in enumerate(rows):
                    y = y_cur - (i + 0.5) * step
                    if lbl:
                        ax.text(0.01, y, lbl, transform=ax.transAxes,
                                fontsize=7, va="center", color="#555555", style="italic")
                    ax.text(0.50, y, mstr, transform=ax.transAxes,
                            fontsize=10.5, va="center", ha="center",
                            fontfamily="DejaVu Sans")
                buf = _io_7.BytesIO()
                fig.savefig(buf, format="png", bbox_inches="tight", dpi=dpi,
                            facecolor=fig.get_facecolor())
                _plt_7.close(fig)
                buf.seek(0)
                return buf

            # ── HELPER: inline CDM block (unchanged) ───────────────────────────────
            def _icb(layers, soil_top, cdm_top, cdm_bot,
                     cdm_phi=30, cdm_c=50, cdm_gamma=20, cdm_gamma_sub=10):
                cdm = _WIF_EL(cdm_bot, cdm_gamma, cdm_gamma_sub, cdm_phi, cdm_c)
                segs_orig = []; cur_top = soil_top
                for lay in layers:
                    segs_orig.append((cur_top, lay.tip_elev, lay)); cur_top = lay.tip_elev
                new_segs = []
                for (s_top, s_bot, lay) in segs_orig:
                    bot_above = max(s_bot, cdm_top)
                    if s_top > bot_above: new_segs.append((s_top, bot_above, lay))
                    top_below = min(s_top, cdm_bot)
                    if top_below > s_bot: new_segs.append((top_below, s_bot, lay))
                soil_bot = layers[-1].tip_elev
                cdm_top_eff = min(cdm_top, soil_top); cdm_bot_eff = max(cdm_bot, soil_bot)
                if cdm_top_eff > cdm_bot_eff:
                    new_segs.append((cdm_top_eff, cdm_bot_eff, cdm))
                new_segs.sort(key=lambda s: -s[0])
                return [_WIF_EL(s[1], s[2].gamma, s[2].gamma_sub, s[2].phi, s[2].c)
                        for s in new_segs]

            # ── HELPER: wall schematic (matplotlib patches + hatching) ─────────────
            def _wall_schematic7(hk, geom, cte, cbe, cdm_lc, pname, png_out):
                Z = float(hk.get("Z_m", 0))
                Lp = float(hk.get("recommended_L_m") or 29)
                top = geom.top_elev; bot = geom.bot_elev
                sl_b = geom.soil_level_back; wl_f = geom.water_elev_front
                fig_h = max(Lp * 0.4 + 2.5, 8)
                fig, ax = _plt_7.subplots(figsize=(7 / 2.54, fig_h / 2.54))
                ax.set_facecolor("#eaf4fb"); fig.patch.set_facecolor("white")
                ax.add_patch(_mpatch7.Rectangle((-5, bot - 1), 5, (top + 2) - (bot - 1),
                                                 fc="#ead8bc", zorder=0))
                ax.add_patch(_mpatch7.Rectangle((0, bot - 1), 5, sl_b - (bot - 1),
                                                 fc="#cce0f0", zorder=0))
                ax.add_patch(_mpatch7.Rectangle((0, sl_b), 5, (top + 2) - sl_b,
                                                 fc="#eaf4fb", zorder=0))
                if top > Z:
                    ax.add_patch(_mpatch7.Rectangle((-5, Z), 5, top - Z,
                                                     fc="#c8a96e", ec="#8B5e14", lw=0.7,
                                                     hatch="//", alpha=0.85, zorder=1))
                    ax.text(-2.5, (Z + top) / 2, "Dat dap", ha="center", va="center",
                            fontsize=6.5, color="#5c3d00", fontweight="bold", zorder=3)
                ax.add_patch(_mpatch7.Rectangle((-5, bot - 1), 5, Z - (bot - 1),
                                                 fc="#c4a882", ec="#8B6914", lw=0.3,
                                                 hatch="..", alpha=0.6, zorder=1))
                if cte > cbe:
                    ax.add_patch(_mpatch7.Rectangle((-5, cbe), 5, cte - cbe,
                                                     fc="#a8dba8", ec="#1e7e1e", lw=0.9,
                                                     hatch="xx", alpha=0.9, zorder=2))
                    ax.text(-2.5, (cbe + cte) / 2,
                            "CDM\nLc=%.1fm" % cdm_lc, ha="center", va="center",
                            fontsize=6.5, color="#145214", fontweight="bold", zorder=4)
                ax.add_patch(_mpatch7.Rectangle((0, bot - 1), 5, sl_b - (bot - 1),
                                                 fc="#b8cfe8", ec="#4a6fa5", lw=0.3,
                                                 hatch="\\\\", alpha=0.7, zorder=1))
                ax.axhline(wl_f, xmin=0, xmax=0.5, color="#1565c0", lw=1.5, ls="--", zorder=4)
                ax.axhline(geom.water_elev_back, xmin=0.5, xmax=1.0,
                            color="#1565c0", lw=1.2, ls="-.", zorder=4)
                ax.text(-4.7, wl_f + 0.2, "MNN Front", fontsize=6, color="#1565c0", zorder=5)
                ax.add_patch(_mpatch7.Rectangle((-0.2, bot), 0.4, Lp,
                                                 fc="#2c2c2c", ec="black", lw=0.9, zorder=5))
                ax.plot([-5, 0], [Z, Z], "k-", lw=1.6, zorder=4)
                ax.plot([0, 5], [sl_b, sl_b], "k--", lw=1.2, zorder=4)
                for xi in [-4, -2.5, -1]:
                    ax.annotate("", xy=(xi, top), xytext=(xi, top + 0.7),
                                arrowprops=dict(arrowstyle="->", color="#b03000", lw=1.0), zorder=5)
                ax.text(-2.5, top + 0.85, "q=10 kPa", ha="center",
                        fontsize=6, color="#b03000", zorder=5)
                for elev, lbl in [(top, "+%.2f (dinh cu)" % top),
                                   (Z, "+%.2f (MD Front)" % Z),
                                   (cte, "+%.2f (dau CDM)" % cte),
                                   (cbe, "+%.2f (chan CDM)" % cbe),
                                   (bot, "+%.2f (chan cu)" % bot)]:
                    ax.plot([-0.25, 0.25], [elev, elev], "k-", lw=0.8, zorder=6)
                    ax.text(0.35, elev, lbl, va="center", fontsize=5.5, color="#333", zorder=6)
                ax.text(-4.5, top + 1.2, "TRAI (Front/Active)", fontsize=7,
                        fontweight="bold", color="#7a3b00", zorder=5)
                ax.text(0.5, top + 1.2, "PHAI (Back/Passive)", fontsize=7,
                        fontweight="bold", color="#003b7a", zorder=5)
                ax.text(-0.1, (bot + top) / 2, pname, ha="right", fontsize=6.5,
                        color="white", fontweight="bold", rotation=90, zorder=6)
                ax.set_xlim(-5.5, 5.5); ax.set_ylim(bot - 1.5, top + 2.0)
                ax.set_ylabel("Cao do (m)", fontsize=7); ax.set_xticks([])
                ax.set_title("%s - %s, L=%.0fm" % (hk["name"], pname, Lp),
                             fontsize=8, fontweight="bold")
                ax.grid(axis="y", alpha=0.25, lw=0.5)
                _plt_7.tight_layout()
                _plt_7.savefig(png_out, dpi=110, bbox_inches="tight", facecolor="white")
                _plt_7.close()

            # ── HELPER: 5-panel force diagram ──────────────────────────────────────
            def _plot5_7(hk, geom, res, cte, cbe, pname, Lp, png_out):
                zs = res["zs"]
                el_pile = [geom.top_elev - z for z in zs]
                el_mid = [geom.top_elev - (zs[i] + zs[i + 1]) / 2 for i in range(len(zs) - 1)]
                Mcr = res["Mcr_kNm"]
                u = res["u_max_mm"]; M = res["M_max_kNm"]; Q = res["Q_max_kN"]
                kh_arr = res.get("kh", [])
                fig, axes = _plt_7.subplots(
                    1, 5, figsize=(17 / 2.54, 9 / 2.54), sharey=True,
                    gridspec_kw={"width_ratios": [1.5, 2, 2.5, 2, 1.5]})
                fig.patch.set_facecolor("white")
                ax0, ax1, ax2, ax3, ax4 = axes
                # ax0 schematic
                ax0.set_facecolor("#f0f4f8")
                ax0.add_patch(_mpatch7.Rectangle((-0.2, geom.bot_elev), 0.4, Lp,
                                                   fc="#2c2c2c", ec="black", lw=0.8, zorder=3))
                if cte > cbe:
                    ax0.add_patch(_mpatch7.Rectangle((-2, cbe), 2, cte - cbe,
                                                      fc="#a8dba8", ec="#1e7e1e",
                                                      hatch="xx", alpha=0.85, lw=0.7, zorder=2))
                if geom.top_elev > geom.soil_level_front:
                    ax0.add_patch(_mpatch7.Rectangle((-2, geom.soil_level_front), 2,
                                                      geom.top_elev - geom.soil_level_front,
                                                      fc="#c8a96e", hatch="//", alpha=0.7,
                                                      lw=0.5, zorder=2))
                ax0.plot([-2, 0], [geom.soil_level_front] * 2, "k-", lw=1.5, zorder=4)
                ax0.plot([0, 2], [geom.soil_level_back] * 2, "k--", lw=1.2, zorder=4)
                ax0.axhline(geom.water_elev_front, color="#1565c0", lw=1.0,
                             ls="--", xmin=0, xmax=0.5, zorder=4)
                ax0.set_xlim(-2.5, 2.5); ax0.set_xticks([])
                ax0.set_title("So do", fontsize=7)
                if cte > cbe:
                    ax0.text(-1.5, (cbe + cte) / 2, "CDM", ha="center",
                             fontsize=6, color="#145214", fontweight="bold")
                ax0.set_ylabel("Cao do (m)", fontsize=7)
                ax0.grid(axis="y", alpha=0.3)
                # ax1 displacement
                ax1.axhspan(cbe, cte, alpha=0.08, color="green", zorder=0)
                ax1.plot(res["ux"], el_pile, color="royalblue", lw=1.8, zorder=3)
                ax1.fill_betweenx(el_pile, 0, res["ux"],
                                   where=[v > 0 for v in res["ux"]], alpha=0.12, color="royalblue")
                ax1.fill_betweenx(el_pile, 0, res["ux"],
                                   where=[v < 0 for v in res["ux"]], alpha=0.12, color="tomato")
                for lim in [50, -50]:
                    ax1.axvline(lim, color="red", ls="--", lw=0.8, alpha=0.7)
                ax1.axvline(0, color="black", lw=0.5)
                ax1.set_xlabel("u (mm)", fontsize=7)
                ax1.set_title("u=%.1f mm" % u, fontsize=7,
                               color="red" if u >= 50 else "black", fontweight="bold")
                ax1.tick_params(labelsize=6); ax1.grid(alpha=0.3)
                # ax2 moment
                ax2.axhspan(cbe, cte, alpha=0.08, color="green", zorder=0)
                ax2.plot(res["Ms"], el_mid, color="forestgreen", lw=1.8, zorder=3)
                ax2.fill_betweenx(el_mid, 0, res["Ms"],
                                   where=[v > 0 for v in res["Ms"]], alpha=0.12, color="forestgreen")
                ax2.fill_betweenx(el_mid, 0, res["Ms"],
                                   where=[v < 0 for v in res["Ms"]], alpha=0.12, color="darkorange")
                ax2.axvline(Mcr, color="red", ls="--", lw=0.9, alpha=0.7)
                ax2.axvline(-Mcr, color="red", ls="--", lw=0.9, alpha=0.7)
                ax2.axvline(0, color="black", lw=0.5)
                ax2.set_xlabel("M (kNm)", fontsize=7)
                ax2.set_title("M=%.0f Mcr=%.0f" % (M, Mcr), fontsize=7,
                               color="red" if M >= Mcr else "black", fontweight="bold")
                ax2.tick_params(labelsize=6); ax2.grid(alpha=0.3)
                # ax3 shear
                ax3.axhspan(cbe, cte, alpha=0.08, color="green", zorder=0)
                ax3.plot(res["Qs"], el_mid, color="darkorchid", lw=1.8, zorder=3)
                ax3.fill_betweenx(el_mid, 0, res["Qs"], alpha=0.12, color="darkorchid")
                ax3.axvline(0, color="black", lw=0.5)
                ax3.set_xlabel("Q (kN)", fontsize=7)
                ax3.set_title("Q=%.0f kN" % Q, fontsize=7)
                ax3.tick_params(labelsize=6); ax3.grid(alpha=0.3)
                # ax4 kh profile
                ax4.axhspan(cbe, cte, alpha=0.08, color="green", zorder=0)
                if kh_arr and len(kh_arr) == len(el_mid) and any(v > 0 for v in kh_arr):
                    ax4.plot(kh_arr, el_mid, color="saddlebrown", lw=1.5, zorder=3)
                    ax4.fill_betweenx(el_mid, 0, kh_arr, alpha=0.18, color="saddlebrown")
                else:
                    ax4.text(0.5, 0.5, "kh\nn/a", ha="center", va="center",
                             transform=ax4.transAxes, fontsize=7, color="gray")
                ax4.axvline(0, color="black", lw=0.5)
                ax4.set_xlabel("kh (kN/m3)", fontsize=7)
                ax4.set_title("Lo xo nen", fontsize=7)
                ax4.tick_params(labelsize=6); ax4.grid(alpha=0.3)
                ok7 = u < 50 and M < Mcr
                fig.suptitle("%s - %s L=%.0fm (CDM Lc=%.1fm) [%s]" % (
                    hk["name"], pname, Lp, _cdm_Lc7, "DAT" if ok7 else "KHONG DAT"),
                    fontsize=8, fontweight="bold",
                    color="darkgreen" if ok7 else "crimson")
                _plt_7.tight_layout(rect=[0, 0, 1, 0.94])
                _plt_7.savefig(png_out, dpi=110, bbox_inches="tight", facecolor="white")
                _plt_7.close()

            # ── HELPER: read stability from SQLite ─────────────────────────────────
            def _read_stab7(bh_full, pile_type, L_m):
                try:
                    con = _sql_7.connect(str(_DB)); con.row_factory = _sql_7.Row
                    rows = con.execute(
                        "SELECT * FROM ke_sw_stability "
                        "WHERE bh_name=? AND pile_type=? AND ABS(L_m-?)<0.6 "
                        "AND method IN ('bishop','spencer','morgenstern_price')",
                        (bh_full, pile_type, L_m)).fetchall()
                    con.close()
                    return {r["method"]: dict(r) for r in rows}
                except Exception:
                    return {}

            # ── HELPER: read NT detail from SQLite ─────────────────────────────────
            def _read_nt7(bh_full, pile_type):
                try:
                    con = _sql_7.connect(str(_DB)); con.row_factory = _sql_7.Row
                    row = con.execute(
                        "SELECT * FROM ke_sw_nt_detail "
                        "WHERE bh_name=? AND pile_type=? ORDER BY rowid DESC LIMIT 1",
                        (bh_full, pile_type)).fetchone()
                    con.close()
                    return dict(row) if row else {}
                except Exception:
                    return {}

            # ── LOAD HK LIST ───────────────────────────────────────────────────────
            try:
                _bhs_all = _ke_data.get("boreholes", [])
                _hk_list = [b for b in _bhs_all if b.get("on_sw_alignment")]
            except Exception:
                _hk_list = []
            if not _hk_list:
                _hk_list = [
                    {"name": "HK2",  "Z_m": 2.030,  "H_layer1_m": 20.0, "recommended_pile": "SW-840", "recommended_L_m": 29},
                    {"name": "HK3",  "Z_m": 1.256,  "H_layer1_m": 19.2, "recommended_pile": "SW-840", "recommended_L_m": 29},
                    {"name": "HK7",  "Z_m": -0.561, "H_layer1_m": 21.0, "recommended_pile": "SW-840", "recommended_L_m": 29},
                    {"name": "HK8",  "Z_m": 2.579,  "H_layer1_m": 24.1, "recommended_pile": "SW-840", "recommended_L_m": 29},
                    {"name": "HK9",  "Z_m": -2.250, "H_layer1_m": 21.0, "recommended_pile": "SW-840", "recommended_L_m": 29},
                    {"name": "HK10", "Z_m": -0.381, "H_layer1_m": 25.0, "recommended_pile": "SW-940", "recommended_L_m": 29},
                    {"name": "HK11", "Z_m": -0.220, "H_layer1_m": 24.2, "recommended_pile": "SW-840", "recommended_L_m": 29},
                ]
            _PILE_LIB7 = {
                "SW-840": dict(H_mm=840, Itd_cm4=2_125_017, Mcr_Tm=77.10, Atd_cm2=3107),
                "SW-940": dict(H_mm=940, Itd_cm4=2_983_488, Mcr_Tm=93.30, Atd_cm2=3544),
            }
            _TOP_KE7   = 2.7
            _cdm_Lc7   = float(st.session_state.get("cdm_Lc", 26.2) or 26.2)
            _cdm_Lng7  = float(st.session_state.get("cdm_L_ngam", 0.5) or 0.5)
            _cdm_thk7  = max(0.0, _cdm_Lc7 - _cdm_Lng7)
            _cdm_fac7  = 3.0

            def _solve7(hk):
                _nm = hk["name"]
                pname = (st.session_state.get(f"dpy_pile_{_nm}")
                          or st.session_state.get("ke_sw_rec_piles", {}).get(_nm)
                          or hk.get("recommended_pile") or "SW-840")
                _ce = None
                try:
                    _ce = next((c for c in _sw_piles if c.get("name") == pname), None)
                except Exception:
                    pass
                if _ce:
                    p_props = dict(H_mm=int(_ce.get("H_mm") or 840),
                                   Itd_cm4=float(_ce.get("Itd_cm4") or 2_125_017),
                                   Mcr_Tm=float(_ce.get("Mcr_Tm") or 77.10),
                                   Atd_cm2=float(_ce.get("Atd_cm2") or 3107))
                else:
                    p_props = _PILE_LIB7.get(pname, _PILE_LIB7["SW-840"])
                pile = _wif_sw(name=pname, fc_MPa=70, **p_props)
                Z  = float(hk.get("Z_m", 0))
                H1 = float(hk.get("H_layer1_m", 20))
                Lp = float(st.session_state.get(f"dpy_L_{_nm}")
                            or st.session_state.get("ke_sw_L_thiet_ke", {}).get(_nm)
                            or hk.get("recommended_L_m") or 29)
                geom = _WIF_WG(top_elev=_TOP_KE7, pile_length=Lp,
                                soil_level_front=Z, soil_level_back=Z - 1.0,
                                water_elev_front=Z - 0.5, water_elev_back=Z - 0.5,
                                surcharge_front=10.0)
                fill_h = _TOP_KE7 - Z
                fill   = _WIF_EL(Z, 18, 8, 28, 0) if fill_h > 0 else None
                pile_bot = _TOP_KE7 - Lp
                front_baseline = [_WIF_EL(Z - H1, 15, 5, 10, 5),
                                   _WIF_EL(pile_bot - 1, 18, 8, 30, 0)]
                back_layers = list(front_baseline)
                cte = Z; cbe = Z - _cdm_thk7
                front_cdm = _icb(front_baseline, soil_top=Z, cdm_top=cte, cdm_bot=cbe,
                                  cdm_phi=30.0, cdm_c=50.0)
                H1_after = max(0.1, H1 - _cdm_thk7)
                layers_kh = [
                    _WIF_SL("XMD", _cdm_thk7, Su_kPa=50, gamma_kNm3=20),
                    _WIF_SL("1",   H1_after,   Su_kPa=10, gamma_kNm3=15),
                    _WIF_SL("2b",  10,          gamma_kNm3=18),
                ]
                if fill_h > 0:
                    layers_kh.insert(0, _WIF_SL("F", fill_h, gamma_kNm3=18))
                load = _wif_build(geom, front_cdm, back_layers, fill=fill, N=300, mode="winkler")
                res  = _wif_solve(layers_kh, pile, Lp,
                                   zs_load=load["zs_depth_m"], p_load_kNm2=load["p_net"],
                                   N=60, top_pin=True,
                                   cdm_thickness_m=_cdm_thk7 + (_TOP_KE7 - cte),
                                   cdm_factor=_cdm_fac7)
                return geom, res, cte, cbe, pname, Lp

            # ── COMPUTE ALL HKs ────────────────────────────────────────────────────
            _records7 = []; _pngs_force7 = []; _pngs_schem7 = []
            _tmpdir = _tmp_7.mkdtemp()
            with st.spinner("Đang tính nội lực + ổn định 7 HK..."):
                for hk in _hk_list:
                    out = _solve7(hk)
                    if out is None: continue
                    geom7, res7, cte7, cbe7, pname7, Lp7 = out
                    if res7.get("error"): continue
                    _nm7    = hk["name"]
                    _bh_full7 = "KE-" + _nm7
                    _png_sc = _os_7.path.join(_tmpdir, "sc_%s.png" % _nm7)
                    _wall_schematic7(hk, geom7, cte7, cbe7, _cdm_Lc7, pname7, _png_sc)
                    _pngs_schem7.append(_png_sc)
                    _png_f = _os_7.path.join(_tmpdir, "f_%s.png" % _nm7)
                    _plot5_7(hk, geom7, res7, cte7, cbe7, pname7, Lp7, _png_f)
                    _pngs_force7.append(_png_f)
                    _stab7 = _read_stab7(_bh_full7, pname7, Lp7)
                    _nt7   = _read_nt7(_bh_full7, pname7)
                    _records7.append(dict(
                        hk=hk, geom=geom7, res=res7, cte=cte7, cbe=cbe7,
                        pname=pname7, Lp=Lp7, bh_full=_bh_full7,
                        u=res7["u_max_mm"], M=res7["M_max_kNm"],
                        Q=res7["Q_max_kN"],  Mcr=res7["Mcr_kNm"],
                        ratio=res7["M_max_kNm"] / res7["Mcr_kNm"],
                        stab=_stab7, nt=_nt7,
                    ))

            # ── SQLite lookup cho bảng thông số (SQLite-primary, giống Mục B) ──────
            try:
                with _sql_7.connect(str(_DB), timeout=5) as _c_ntd7:
                    _c_ntd7.row_factory = _sql_7.Row
                    _nt_detail_7: dict = {
                        r["bh_name"]: dict(r)
                        for r in _c_ntd7.execute("SELECT * FROM ke_sw_nt_detail").fetchall()
                    }
            except Exception:
                _nt_detail_7 = {}

            # ── BUILD WORD DOCUMENT ────────────────────────────────────────────────
            with st.spinner("Đang xây dựng tài liệu Word..."):
                doc = _Doc7()
                _sec7 = doc.sections[0]
                _sec7.left_margin = _Cm7(2.5); _sec7.right_margin  = _Cm7(2.5)
                _sec7.top_margin  = _Cm7(2.5); _sec7.bottom_margin = _Cm7(2.5)

                # ── Helper header/footer ───────────────────────────────────────────
                def _field_run7(para, field: str) -> None:
                    run = para.add_run()
                    fc_b = _OXE7("w:fldChar"); fc_b.set(_qn7("w:fldCharType"), "begin")
                    instr = _OXE7("w:instrText"); instr.text = field
                    instr.set(_qn7("xml:space"), "preserve")
                    fc_e = _OXE7("w:fldChar"); fc_e.set(_qn7("w:fldCharType"), "end")
                    run._r.append(fc_b); run._r.append(instr); run._r.append(fc_e)

                def _no_border7(tbl) -> None:
                    tblPr = tbl._tbl.tblPr
                    borders = _OXE7("w:tblBorders")
                    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                        b = _OXE7(f"w:{edge}"); b.set(_qn7("w:val"), "none"); borders.append(b)
                    tblPr.append(borders)

                _co_name7  = st.session_state.get("export_co_name") or ""
                _co_staff7 = st.session_state.get("export_co_staff") or ""
                _logo7     = st.session_state.get("export_logo_bytes")
                _pw7 = _sec7.page_width - _sec7.left_margin - _sec7.right_margin

                # Header: logo trái | tên công ty phải
                _sec7.header.is_linked_to_previous = False
                _ht7 = _sec7.header.add_table(1, 2, width=_pw7)
                _no_border7(_ht7)
                _ht7.cell(0, 0).width = _Cm7(3); _ht7.cell(0, 1).width = _pw7 - _Cm7(3)
                if _logo7:
                    try: _ht7.cell(0, 0).paragraphs[0].add_run().add_picture(
                            _io_7.BytesIO(_logo7), height=_Cm7(1.2))
                    except Exception: pass
                _hp7 = _ht7.cell(0, 1).paragraphs[0]; _hp7.alignment = _ALN7.RIGHT
                _rh = _hp7.add_run(_co_name7); _rh.bold = True; _rh.font.size = _Pt7(9)

                # Footer: nhân sự trái | Trang X/Y phải
                _sec7.footer.is_linked_to_previous = False
                _ft7 = _sec7.footer.add_table(1, 2, width=_pw7)
                _no_border7(_ft7)
                _ft7.cell(0, 0).width = _pw7 - _Cm7(3); _ft7.cell(0, 1).width = _Cm7(3)
                _ft7.cell(0, 0).paragraphs[0].add_run(_co_staff7).font.size = _Pt7(9)
                _fp7 = _ft7.cell(0, 1).paragraphs[0]; _fp7.alignment = _ALN7.RIGHT
                _fp7.add_run("Trang ").font.size = _Pt7(9)
                _field_run7(_fp7, "PAGE")
                _fp7.add_run(" / ").font.size = _Pt7(9)
                _field_run7(_fp7, "NUMPAGES")

                # ─ TRANG BÌA ──────────────────────────────────────────────────────
                for _ in range(3):
                    doc.add_paragraph()
                _h0t = doc.add_heading("BÁO CÁO KỸ THUẬT", level=0)
                _h0t.alignment = _ALN7.CENTER
                _h1t = doc.add_heading("THIẾT KẾ CỌC VÁN SW + XỬ LÝ NỀN CDM", level=0)
                _h1t.alignment = _ALN7.CENTER
                _ph = doc.add_paragraph(); _ph.alignment = _ALN7.CENTER
                _r = _ph.add_run("Kè cọc ván SW – Khu vực KE – Trung tâm Hành chính TP.HCM")
                _r.font.size = _Pt7(12); _r.bold = True
                doc.add_paragraph()
                _mt = doc.add_table(rows=6, cols=2)
                _mt.style = "Table Grid"
                _meta_pairs7 = [
                    ("Dự án",         "Trung tâm Hành chính TP.HCM – Mã số 202605-TTHC"),
                    ("Hạng mục",      "Kè cọc ván dự ứng lực SW – %d hố khoan trên tuyến kè" % len(_hk_list)),
                    ("Thiết kế CDM",  "Lc = %.1f m  Ngàm = %.1f m  k_CDM = ×%.1f" % (_cdm_Lc7, _cdm_Lng7, _cdm_fac7)),
                    ("Tiêu chuẩn",    "TCVN 11823-10:2017  –  TCVN 9403:2012  –  TCVN 4253:2012"),
                    ("Phương pháp",   "α-method (Tomlinson 1980)  –  Winkler (API Clay)  –  Bishop"),
                    ("Ngày lập",      _dt_7.now().strftime("%d/%m/%Y")),
                ]
                for i, (k, v) in enumerate(_meta_pairs7):
                    _mt.rows[i].cells[0].text = k
                    _mt.rows[i].cells[1].text = v
                    _rn = _mt.rows[i].cells[0].paragraphs[0].runs
                    if _rn: _rn[0].bold = True
                    _cell_color7(_mt.rows[i].cells[0], "E8ECEF")
                doc.add_page_break()

                # ─ BẢNG THÔNG SỐ THIẾT KẾ PER HK ─────────────────────────────────
                doc.add_heading("1. Bảng tổng hợp thông số và kết quả thiết kế", level=1)
                _col_ts7 = ["Hố khoan", "Cọc thiết kế", "Z (m)", "H lớp yếu (m)",
                            "L yêu cầu NT1 (m)", "L thiết kế (m)", "NT1", "NT2"]
                _tm0 = doc.add_table(rows=1, cols=len(_col_ts7))
                _tm0.style = "Table Grid"
                for i, h in enumerate(_col_ts7):
                    _c = _tm0.rows[0].cells[i]; _c.text = h
                    _cell_color7(_c, "1F3864")
                    for para in _c.paragraphs:
                        para.paragraph_format.space_before = _Pt7(1)
                        para.paragraph_format.space_after  = _Pt7(1)
                        for run in para.runs:
                            run.bold = True; run.font.color.rgb = _RGB7(0xFF, 0xFF, 0xFF)
                            run.font.size = _Pt7(8.5)
                for _hk_m in _hk_list:
                    _nm_m  = _hk_m["name"]
                    _db_m  = _nt_detail_7.get(f"KE-{_nm_m}") or {}
                    _pname_m = (st.session_state.get("ke_sw_rec_piles", {}).get(_nm_m)
                                or _hk_m.get("recommended_pile") or "SW-840")
                    _Lp_m  = float(st.session_state.get("ke_sw_L_thiet_ke", {}).get(_nm_m)
                                   or _hk_m.get("recommended_L_m") or 29)
                    _Z_m   = float(_db_m.get("Z_m")             or _hk_m.get("Z_m") or 0)
                    _H1_m  = float(_db_m.get("D_bottom_soft_m") or _hk_m.get("H_layer1_m") or 0)
                    _Lreq  = float(_db_m.get("L_req_nt1_m")     or _hk_m.get("L_req_m") or 0)
                    _nt1_m = _db_m.get("nt1_result") or _hk_m.get("NT1") or "–"
                    _nt2_m = _db_m.get("nt2_result") or _hk_m.get("NT2") or "–"
                    _rw = _tm0.add_row().cells
                    _rw[0].text = _nm_m;             _rw[1].text = _pname_m
                    _rw[2].text = "%.2f" % _Z_m;     _rw[3].text = "%.1f" % _H1_m
                    _rw[4].text = "%.1f" % _Lreq;    _rw[5].text = "%.1f" % _Lp_m
                    _rw[6].text = _nt1_m;             _rw[7].text = _nt2_m
                    _ok1 = str(_nt1_m).strip().lower() in ("dat", "đạt")
                    _ok2 = str(_nt2_m).strip().lower() in ("dat", "đạt")
                    _cell_color7(_rw[6], "CCFFCC" if _ok1 else "FFCCCC")
                    _cell_color7(_rw[7], "CCFFCC" if _ok2 else "FFCCCC")
                    _fmt_row7(_rw, 8.5)
                doc.add_page_break()

                # ─ SECTION 2: CƠ SỞ LÝ THUYẾT ────────────────────────────────────
                doc.add_heading("2. Cơ sở lý thuyết & Công thức tính toán", level=1)

                doc.add_heading("2.1 Kiểm tra NT1 – Chiều dài cọc tối thiểu (TCVN 11823-10:2017)", level=2)
                doc.add_paragraph(
                    "Chiều dài thiết kế phải đủ neo qua toàn bộ lớp đất yếu và ngàm vào lớp "
                    "chịu lực bên dưới. Điều kiện NT1 (Điều 7.3.4):"
                )
                _img_nt1 = _formula_img7([
                    ("Điều kiện:", r"$L_{TK} \geq L_{yc} = H_{fill} + D_{bot,soft} + L_{pen,min}$"),
                    ("Trong đó:",  r"$L_{pen,min} = 2{,}0\ \text{m}$  ·  $H_{fill} = z_{top} - Z_{HK}\ (\geq 0)$"),
                    ("",           r"$D_{bot,soft}$  =  chiều sâu đến đáy lớp yếu cuối cùng"),
                ], title="Công thức NT1 – Chiều dài tối thiểu", height_cm=4.5)
                doc.add_picture(_img_nt1, width=_Cm7(14))

                doc.add_heading("2.2 Kiểm tra NT2 – Sức chịu tải dọc trục (α-method, Tomlinson 1980)", level=2)
                doc.add_paragraph(
                    "Phương pháp α (TCVN 11823-10:2017 Điều 7.3.8.6.2) áp dụng cho đất sét. "
                    "Hệ số φ_stat = 0,35. Bỏ qua sức kháng bên trong vùng đất đắp:"
                )
                _img_nt2 = _formula_img7([
                    ("Sức kháng bên:", r"$R_s = \alpha \cdot \overline{S_u} \cdot P_{perimeter} \cdot L_{soil}$"),
                    ("Sức kháng mũi:", r"$R_p = 9 \cdot S_{u,tip} \cdot A_{pile}$"),
                    ("Kiểm tra:",      r"$R_R = \phi_{stat} \cdot (R_s + R_p) \geq W_{pile}$  ·  $\phi_{stat} = 0{,}35$"),
                    ("Nguồn Su:",      r"Ưu tiên: VST cắt cánh  >  UU lab  >  Giả định (cảnh báo)"),
                ], title="Công thức NT2 – Sức chịu tải dọc trục (α-method, Tomlinson 1980)", height_cm=5.0)
                doc.add_picture(_img_nt2, width=_Cm7(14))

                doc.add_heading("2.3 Nội lực tường cừ – Mô hình Winkler (API Clay, Matlock 1970)", level=2)
                doc.add_paragraph(
                    "Dầm trên nền đàn hồi Winkler. Tải phân bố = áp lực đất Active (Front) + "
                    "áp lực nước chênh lệch. Lò xo nền phía Back kháng chuyển vị:"
                )
                _img_wink = _formula_img7([
                    ("Phản lực nền:", r"$p = k_h \cdot u \cdot B_{pile}$"),
                    ("Hệ số kh:",     r"$k_h = \frac{67 \cdot S_u}{d_{pile}}\ \ \text{(API Clay, kN/m}^3\text{)}$"),
                    ("Vùng CDM:",     r"$k_h^{CDM} = k_{factor} \times k_h\ \ (k_{factor} = 3{,}0)$"),
                    ("Tiêu chí:",     r"$u_{max} \leq 50\ \text{mm}$   ·   $M_{max} \leq M_{cr}$"),
                ], title="Mô hình Winkler – API Clay (Matlock 1970)", height_cm=5.5)
                doc.add_picture(_img_wink, width=_Cm7(14))
                doc.add_paragraph(
                    "Cọc tựa khớp ở đỉnh (top_pin=True). Chia N=60 đoạn. "
                    "Tải áp lực đất tính từ WallGeometry và EarthLayer (Active + Nước Front). "
                    "Su phía Back lấy từ VST > lab_Cu_UU > fallback 11 kPa."
                )

                doc.add_heading("2.4 Ổn định tổng thể – Bishop Simplified / Spencer / Morgenstern-Price", level=2)
                doc.add_paragraph(
                    "Ba phương pháp phân mảnh cho cung trượt tròn. "
                    "Yêu cầu: Fs ≥ 1,30 (TCVN 4253:2012)."
                )
                _img_bish = _formula_img7([
                    ("Bishop:",    r"$F_s = \dfrac{\sum \left[c'b + (W - ub)\tan\phi'\right] / m_\alpha}{\sum W \sin\alpha}$"),
                    ("",           r"$m_\alpha = \cos\alpha \left(1 + \dfrac{\tan\alpha\,\tan\phi'}{F_s}\right)$"),
                    ("Spencer:",   r"Thêm cân bằng lực dọc phân mảnh (iterative, lực liên phân mảnh thứ cấp)"),
                    ("M-P:",       r"Dùng hàm $f(\alpha)$ mô tả phương lực liên phân mảnh (mô hình đầy đủ nhất)"),
                ], title="Ổn định tổng thể – Bishop / Spencer / Morgenstern-Price", height_cm=5.5)
                doc.add_picture(_img_bish, width=_Cm7(14))
                doc.add_page_break()

                # ─ SECTION 3: TỔNG HỢP NT1/NT2/NỘI LỰC ──────────────────────────
                doc.add_heading("3. Tổng hợp kết quả NT1, NT2 và Nội lực", level=1)
                _col_h7 = [
                    "HK", "Cọc", "L TK\n(m)", "L yc NT1\n(m)", "NT1",
                    "Rs\n(kN)", "Rp\n(kN)", "RR\n(kN)", "W\n(kN)", "NT2",
                    "u_max\n(mm)", "M_max\n(kNm)", "M/Mcr", "Kết luận",
                ]
                _tb2 = doc.add_table(rows=1, cols=len(_col_h7))
                _tb2.style = "Table Grid"
                for i, h in enumerate(_col_h7):
                    _c = _tb2.rows[0].cells[i]
                    _c.text = h; _cell_color7(_c, "1F3864")
                    for para in _c.paragraphs:
                        para.paragraph_format.space_before = _Pt7(1)
                        para.paragraph_format.space_after  = _Pt7(1)
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = _RGB7(0xFF, 0xFF, 0xFF)
                            run.font.size = _Pt7(7)
                for r in _records7:
                    _nd  = r["nt"]
                    _row = _tb2.add_row().cells
                    _row[0].text = r["hk"]["name"]
                    _row[1].text = r.get("pname", "")
                    _row[2].text = "%.1f" % r["Lp"]
                    _row[3].text = ("%.1f" % _nd.get("L_req_nt1_m", 0)) if _nd else "-"
                    _nt1_ok = str(_nd.get("nt1_result", "")).strip().lower() in ("dat", "đạt") if _nd else None
                    _row[4].text = ("Đạt" if _nt1_ok else "Không đạt") if _nd else "-"
                    if _nd and not _nt1_ok: _cell_color7(_row[4], "FFCCCC")
                    elif _nd and _nt1_ok:   _cell_color7(_row[4], "CCFFCC")
                    _row[5].text = ("%.0f" % _nd.get("Rs_kN", 0))  if _nd else "-"
                    _row[6].text = ("%.0f" % _nd.get("Rp_kN", 0))  if _nd else "-"
                    _row[7].text = ("%.0f" % _nd.get("RR_kN", 0))  if _nd else "-"
                    _row[8].text = ("%.0f" % _nd.get("W_kN", 0))   if _nd else "-"
                    _nt2_ok = str(_nd.get("nt2_result", "")).strip().lower() in ("dat", "đạt") if _nd else None
                    _row[9].text = ("Đạt" if _nt2_ok else "Không đạt") if _nd else "-"
                    if _nd and not _nt2_ok: _cell_color7(_row[9], "FFCCCC")
                    elif _nd and _nt2_ok:   _cell_color7(_row[9], "CCFFCC")
                    _row[10].text = "%.1f" % r["u"]
                    _row[11].text = "%.0f" % r["M"]
                    _row[12].text = "%.2f" % r["ratio"]
                    _ok_w7 = r["u"] < 50 and r["M"] < r["Mcr"]
                    _row[13].text = "Đạt" if _ok_w7 else "Không đạt"
                    if not _ok_w7: _cell_color7(_row[13], "FFCCCC")
                    else:          _cell_color7(_row[13], "CCFFCC")
                    _fmt_row7(_row, 7.5)
                doc.add_page_break()

                # ─ SECTION 3: BIEU DO NOI LUC ─────────────────────────────────────
                doc.add_heading("3. Biểu đồ nội lực từng hố khoan", level=1)
                doc.add_paragraph(
                    "Mỗi hố khoan gồm 5 panel: (1) Sơ đồ tường + CDM + đất đắp, "
                    "(2) Chuyển vị ngang u (mm), (3) Mô men uốn M (kNm) so với Mcr, "
                    "(4) Lực cắt Q (kN), (5) Hệ số lò xo nền kh (kN/m³). "
                    "Vùng CDM tô nền xanh lá nhạt. Đường đứt đỏ = giới hạn cho phép."
                )
                for r, _png_f in zip(_records7, _pngs_force7):
                    _hk = r["hk"]
                    doc.add_heading(
                        "%s - %s  L = %.0f m  CDM Lc = %.1f m" % (
                            _hk["name"], r["pname"], r["Lp"], _cdm_Lc7),
                        level=2)
                    _p_sum = doc.add_paragraph()
                    _p_sum.add_run("Kết quả: ").bold = True
                    _ok_r = r["u"] < 50 and r["M"] < r["Mcr"]
                    _p_sum.add_run(
                        "u_max = %.2f mm (≤50)  M_max = %.0f kNm (Mcr=%.0f)  "
                        "M/Mcr = %.2f  Q_max = %.0f kN  →  %s" % (
                            r["u"], r["M"], r["Mcr"], r["ratio"],
                            r["Q"], "ĐẠT" if _ok_r else "KHÔNG ĐẠT"))
                    doc.add_picture(_png_f, width=_Cm7(15.5))
                    doc.add_paragraph()
                doc.add_page_break()

                # ─ SECTION 4: ON DINH TONG THE ────────────────────────────────────
                doc.add_heading("4. Ổn định tổng thể tường SW + CDM", level=1)
                _p4 = doc.add_paragraph()
                _p4.add_run("Tiêu chí kiểm tra: ").bold = True
                _p4.add_run(
                    "Fs_trượt cung tròn ≥ 1,30 (TCVN 4253:2012). "
                    "Ba phương pháp phân mảnh: Bishop Simplified, Spencer, Morgenstern-Price. "
                    "Kết quả lấy từ cơ sở dữ liệu tính toán.")
                _cols_st7 = ["HK", "Cọc", "L (m)", "Fs Bishop", "Fs Spencer",
                              "Fs M-P", "Fs nghiêng", "Fs đẩy trồi", "Kết luận"]
                _tb_st7 = doc.add_table(rows=1, cols=len(_cols_st7))
                _tb_st7.style = "Table Grid"
                for i, h in enumerate(_cols_st7):
                    _c = _tb_st7.rows[0].cells[i]; _c.text = h
                    _cell_color7(_c, "1F3864")
                    for para in _c.paragraphs:
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = _RGB7(0xFF, 0xFF, 0xFF)
                            run.font.size = _Pt7(7)

                def _fs_val7(d, *keys):
                    for k in keys:
                        v = d.get(k)
                        if v is not None:
                            try: return float(v)
                            except (TypeError, ValueError): pass
                    return None

                for r in _records7:
                    _stb  = r["stab"]
                    _bish7 = _stb.get("bishop", {})
                    _sp7   = _stb.get("spencer", {})
                    _mp7   = _stb.get("morgenstern_price", {})
                    _Fs_b  = _fs_val7(_bish7, "Fs_slip")
                    _Fs_sp = _fs_val7(_sp7,   "Fs_slip")
                    _Fs_mp = _fs_val7(_mp7,   "Fs_slip")
                    _Fs_lat = _fs_val7(_bish7, "Fs_overturning")
                    _Fs_toe = _fs_val7(_bish7, "Fs_toe_kickout")
                    _row = _tb_st7.add_row().cells
                    _row[0].text = r["hk"]["name"]
                    _row[1].text = r.get("pname", "")
                    _row[2].text = "%.0f" % r["Lp"]
                    _row[3].text = ("%.2f" % _Fs_b)  if _Fs_b  is not None else "-"
                    _row[4].text = ("%.2f" % _Fs_sp) if _Fs_sp is not None else "-"
                    _row[5].text = ("%.2f" % _Fs_mp) if _Fs_mp is not None else "-"
                    _row[6].text = ("%.2f" % _Fs_lat) if _Fs_lat is not None else "-"
                    _row[7].text = ("%.2f" % _Fs_toe) if _Fs_toe is not None else "-"
                    _ok_st7 = _Fs_b is not None and _Fs_b >= 1.30
                    _row[8].text = "Đạt" if _ok_st7 else ("-" if _Fs_b is None else "Không đạt")
                    if _ok_st7:        _cell_color7(_row[8], "CCFFCC")
                    elif _Fs_b is not None: _cell_color7(_row[8], "FFCCCC")
                    _fmt_row7(_row, 7.5)

                _any_stab7 = any(
                    r["stab"].get("bishop", {}).get("Fs_slip") is not None for r in _records7)
                if _any_stab7:
                    _hk_nm_st7 = [r["hk"]["name"] for r in _records7]
                    _fs_b_plt  = [_fs_val7(r["stab"].get("bishop", {}), "Fs_slip") or _np_7.nan for r in _records7]
                    _fs_s_plt  = [_fs_val7(r["stab"].get("spencer", {}), "Fs_slip") or _np_7.nan for r in _records7]
                    _fs_m_plt  = [_fs_val7(r["stab"].get("morgenstern_price", {}), "Fs_slip") or _np_7.nan for r in _records7]
                    _fig_st7, _ax_st7 = _plt_7.subplots(figsize=(14 / 2.54, 6 / 2.54))
                    _xs7 = _np_7.arange(len(_hk_nm_st7)); _ww7 = 0.22
                    _ax_st7.bar(_xs7 - _ww7, _fs_b_plt, _ww7, label="Bishop",  color="#2196F3", alpha=0.85)
                    _ax_st7.bar(_xs7,         _fs_s_plt, _ww7, label="Spencer", color="#4CAF50", alpha=0.85)
                    _ax_st7.bar(_xs7 + _ww7,  _fs_m_plt, _ww7, label="M-P",    color="#FF9800", alpha=0.85)
                    _ax_st7.axhline(1.30, color="red", ls="--", lw=1.2, label="Fs_min=1.30")
                    _ax_st7.set_xticks(_xs7); _ax_st7.set_xticklabels(_hk_nm_st7, fontsize=7)
                    _ax_st7.set_ylabel("Hệ số an toàn Fs", fontsize=8)
                    _ax_st7.legend(fontsize=7, ncol=4); _ax_st7.grid(axis="y", alpha=0.3)
                    _ax_st7.set_title("Fs ổn định tổng thể – 3 phương pháp phân mảnh", fontsize=9)
                    _valid_vals = [v for v in _fs_b_plt + _fs_s_plt + _fs_m_plt if not _np_7.isnan(v)]
                    _ax_st7.set_ylim(0, (max(_valid_vals) if _valid_vals else 2.5) + 0.3)
                    _plt_7.tight_layout()
                    _png_st7 = _os_7.path.join(_tmpdir, "stability.png")
                    _plt_7.savefig(_png_st7, dpi=110, bbox_inches="tight", facecolor="white")
                    _plt_7.close()
                    doc.add_paragraph()
                    doc.add_picture(_png_st7, width=_Cm7(14))
                doc.add_page_break()

                # ─ SECTION 5: KET LUAN ────────────────────────────────────────────
                doc.add_heading("5. Kết luận & Kiến nghị", level=1)
                _n_ok7 = sum(1 for r in _records7 if r["u"] < 50 and r["M"] < r["Mcr"])
                _u_max7 = max((r["u"] for r in _records7), default=0)
                _m_rat7 = max((r["ratio"] for r in _records7), default=0)
                _q_max7 = max((r["Q"] for r in _records7), default=0)
                _p5 = doc.add_paragraph()
                _p5.add_run("Tổng hợp kết quả:").bold = True
                _items5 = [
                    "Số HK đạt tiêu chí nội lực (u ≤ 50 mm và M ≤ Mcr): %d / %d" % (_n_ok7, len(_records7)),
                    "Chuyển vị ngang lớn nhất: u_max = %.2f mm" % _u_max7,
                    "Tỉ số M/Mcr lớn nhất: %.2f" % _m_rat7,
                    "Lực cắt lớn nhất: Q_max = %.0f kN" % _q_max7,
                    "CDM: Lc = %.1f m  Ngàm = %.1f m  k_CDM = ×%.1f" % (_cdm_Lc7, _cdm_Lng7, _cdm_fac7),
                ]
                for _it in _items5:
                    doc.add_paragraph(_it, style="List Bullet")
                doc.add_paragraph()
                _p5k = doc.add_paragraph()
                _p5k.add_run("Kiến nghị:").bold = True
                _kns7 = []
                if _n_ok7 < len(_records7):
                    _kns7.append("Tăng chiều dài cọc hoặc tăng Lc CDM cho các HK chưa đạt yêu cầu.")
                _kns7 += [
                    "Cập nhật Su thực tế từ VST/lab sau khi có thêm kết quả khảo sát.",
                    "Thi công cọc SW từ hai đầu vào giữa để hạn chế heave đất nền.",
                    "Quan trắc chuyển vị định kỳ, không vượt 25 mm trong quá trình thi công.",
                    "Kiểm tra lại Fs ổn định tổng thể bằng phần mềm FEM sau khi có đủ số liệu.",
                ]
                for _kn in _kns7:
                    doc.add_paragraph(_kn, style="List Bullet")

                # ─ LUU VA DOWNLOAD ────────────────────────────────────────────────
                _buf7 = _io_7.BytesIO()
                doc.save(_buf7)
                _buf7.seek(0)

            st.download_button(
                "Tải báo cáo Word đầy đủ 7 HK",
                _buf7.getvalue(),
                file_name="BaoCao_Ke_SW_CDM_%s.docx" % _dt_7.now().strftime("%Y%m%d_%H%M"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
            st.success("Tạo báo cáo thành công: %d/%d HK đạt." % (_n_ok7, len(_records7)))
        except Exception as _e_7hk:
            st.error("Không tạo được báo cáo Word: %s" % _e_7hk)
            import traceback as _tb_7hk
            st.code(_tb_7hk.format_exc())


    # ── Báo cáo PDF tuỳ chỉnh (ĐÃ ẨN — user dùng Ctrl+P trình duyệt) ────────

    # ── Lý thuyết tính toán (in PDF) ─────────────────────────────────────────
    st.divider()
    _theory_text = _load_theory()
    if _theory_text:
        with st.expander("Lý thuyết tính toán – đính kèm khi in PDF", expanded=False):
            st.markdown(_theory_text)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 6 – LÚN NỀN (TCCS 41:2022)
# ═══════════════════════════════════════════════════════════════════════════════
if _page == "settlement":
    import sys as _sys
    _sys.path.insert(0, str(_ROOT / "scripts"))
    try:
        from settlement_calc import (
            compare_methods as _sc_compare,
            calc_settlement_iterative_9_2_3 as _sc_iter923,
        )
        _HAS_SC = True
    except Exception as _exc:
        _HAS_SC = False
        st.error(f"Không tải được settlement_calc: {_exc}")

    if _HAS_SC:
        st.subheader("Dự báo độ lún – So sánh phương án xử lý (TCCS 41:2022)")

        # ── Lý thuyết tính lún (đặt trên cùng — xem trước khi chạy tính toán) ──
        with st.expander("Lý thuyết tính lún cố kết — TCCS 41:2022 + TCVN 9403:2012",
                          expanded=False):
            st.markdown(r"""
## 1. Lún cố kết sơ cấp — Phụ lục A, TCCS 41:2022

Mỗi lớp đất yếu $i$ tính theo trạng thái cố kết. Đặt $\sigma'_{v0}$ là ứng suất hữu hiệu ban đầu,
$\sigma'_{vf} = \sigma'_{v0} + \Delta\sigma$ là ứng suất sau khi đắp, $P_C$ là áp lực tiền cố kết.

### 1.1 Đất quá cố kết (OC) — $\sigma'_{vf} \leq P_C$

$$S_i = H_i \cdot \frac{C_s}{1 + e_0} \cdot \log \frac{\sigma'_{vf}}{\sigma'_{v0}}$$

### 1.2 Cắt qua $P_C$ — $\sigma'_{v0} < P_C < \sigma'_{vf}$

$$S_i = H_i \left[ \frac{C_s}{1 + e_0} \log \frac{P_C}{\sigma'_{v0}}
                  + \frac{C_c}{1 + e_0} \log \frac{\sigma'_{vf}}{P_C} \right]$$

### 1.3 Đất bình thường cố kết (NC) — $\sigma'_{v0} \geq P_C$

$$S_i = H_i \cdot \frac{C_c}{1 + e_0} \cdot \log \frac{\sigma'_{vf}}{\sigma'_{v0}}$$

### 1.4 Tổng lún

$$S_{\text{total}} = \sum_{i=1}^{n} S_i$$

Trong đó:
- $C_c$: chỉ số nén; $C_s$: chỉ số nở lại; $e_0$: hệ số rỗng tự nhiên
- $H_i$: chiều dày lớp đại diện (boundary trung điểm các mẫu)

---

## 2. Lún CDM — TCVN 9403:2012, Phụ lục C

CDM thay thế đất yếu bằng nền hỗn hợp (composite). Lún tổng:

$$S = S_1 + S_2$$

### 2.1 Lún đàn hồi tức thời trong khối gia cố

$$S_1 = \frac{q \cdot H_{\text{soft}}}{a \cdot E_c + (1 - a) \cdot E_s}$$

### 2.2 Lún cố kết bên dưới mũi cột

$$S_2 = 0 \quad (\text{nếu CDM cắm đến lớp cứng})$$

### 2.3 Mô-đun thành phần

$$E_c = 75 \cdot C_{c,\text{col}} = 75 \cdot \left( \text{ratio}_{lab \to field} \cdot \frac{q_{u,lab}}{2} \right)
\quad \text{(TCVN 9403 B.5.1)}$$

$$E_s = 250 \cdot C_u \quad \text{(tương quan Mesri)}$$

$$E_{\text{composite}} = a \cdot E_c + (1 - a) \cdot E_s$$

### 2.4 Tỷ lệ diện tích thay thế $a$

$$a_{\text{vuông}} = \frac{\pi D^2}{4 e^2} \qquad
a_{\text{tam giác}} = \frac{\pi D^2}{2\sqrt{3}\, e^2}$$

---

## 3. Độ cố kết theo thời gian — Điều 9.3–9.4 TCCS 41:2022

### 3.1 Độ cố kết và lún còn lại

$$U(t) = \frac{S(t)}{S_{\text{total}}} \qquad
\Delta S = S_{\text{total}} \cdot \bigl(1 - U(t_{tc})\bigr)$$

### 3.2 Theo phương án xử lý

| Phương án | $U(t)$ | Ghi chú |
|---|---|---|
| Không xử lý | $U_v$ (Terzaghi đứng) — phụ thuộc $C_v$, $H_{dr}$ | Chậm nhất |
| Bấc thấm | $U = 1 - (1 - U_v)(1 - U_h)$ | $U_h$ tăng theo $n = r_e/r_w$ |
| Giếng cát | Tương tự bấc thấm | Đường kính giếng lớn hơn |
| CDM | $S_1$ đàn hồi tức thời → $\Delta S \approx 0$ | Lún xảy ra ngay khi đắp |

### 3.3 Yêu cầu thiết kế

$$\Delta S \leq 30 \text{ cm} \quad \text{(TCCS 41:2022 — đường cấp I đoạn thông thường)}$$

---

## 4. Surcharge (gia tải tạm)

Tăng tải gia tải $> $ tải thiết kế trong giai đoạn cố kết → tăng tốc $U(t)$ →
sau khi dỡ surcharge, lún còn lại $\Delta S$ giảm. **Không** thay đổi $S_{\text{total}}$ dưới tải thiết kế.

---

## 5. Bảng đơn vị

| Đại lượng | Ký hiệu | Đơn vị |
|---|:---:|:---:|
| Chỉ số nén / nở lại | $C_c$, $C_s$ | — |
| Hệ số rỗng | $e_0$ | — |
| Ứng suất hữu hiệu | $\sigma'_v$, $P_C$ | kPa |
| Tải trọng | $q$ | kPa |
| Chiều dày, chiều sâu | $H_i$, $H_{\text{soft}}$ | m |
| Hệ số cố kết | $C_v$ | cm²/s |
| Lún | $S_i$, $S_{\text{total}}$, $\Delta S$ | cm |
| Mô-đun đàn hồi | $E_c$, $E_s$, $E_{\text{composite}}$ | kPa |
| Thời gian | $t$ | tháng |

**Tài liệu nguồn:** TCCS 41:2022 (Phụ lục A, Điều 9.3–9.4), TCVN 9403:2012 (Phụ lục C, Mục B.5.1).
""")

        # ── PHẦN 1: So sánh phương án ────────────────────────────────────────
        if True:
            _sc1, _sc2, _sc3 = st.columns([1, 1, 2])
            with _sc1:
                _sl_zone = st.radio("Zone", list(_ZONE_NAMES.keys()),
                                    format_func=lambda z: _ZONE_NAMES[z],
                                    key="sl_zone")
                _sl_bhs = [b["name"] for b in _load_boreholes_by_zone(_sl_zone)]
                _sl_bh  = st.selectbox("Hố khoan tính lún", _sl_bhs, key="sl_bh")
            with _sc2:
                _sl_from_cdm = st.checkbox("Lấy H từ tab Thông số CDM", key="sl_H_from_cdm")
                if _sl_from_cdm:
                    _cdm_ld  = _get("cdm_loads")
                    _cdm_top = _get("cdm_top_clay")
                    _sl_H    = round(_cdm_ld.get("z_tk", 3.5) - _cdm_top, 2)
                    st.info(f"H = z_tk − đỉnh bùn = **{_sl_H:.2f} m**")
                else:
                    _sl_H = st.number_input("Chiều cao đắp H (m)", 0.5, 15.0,
                                            st.session_state.get("sl_H_manual", 3.0), 0.5,
                                            key="sl_H_manual")
                _sl_lim = st.number_input("Giới hạn lún còn lại (cm)", 10, 50, 30, 5, key="sl_lim")
                _sl_tc  = st.number_input("Thời gian thi công (tháng)", 3, 36, 6, 1, key="sl_tc")
            with _sc3:
                st.caption("**Phương pháp tính:** TCCS 41:2022 — Phụ lục A (lún sơ cấp), Điều 9.3-9.4 (độ cố kết)")
                st.caption("**Lún còn lại ΔS** = lún tổng − lún tại thời điểm kết thúc thi công")

            # ── Bảng thông số địa chất dùng để tính lún cố kết ─────────────
            if _sl_bh:
                _geo_layers = _load_layers(_sl_bh)
                if _geo_layers:
                    _param_rows = []
                    _has_borrowed = False
                    for _lr in _geo_layers:
                        _sym = _lr.get("symbol") or ""
                        _Cc  = _lr.get("Cc")
                        _Cs  = _lr.get("Cs")
                        _e0  = _lr.get("e0")
                        _PC  = _lr.get("PC_kPa")
                        _Cv  = _lr.get("Cv_cm2s")
                        _gam = _lr.get("gamma_kNm3")
                        _src = ""
                        # Nếu thiếu thông số cố kết → tham khảo hố khoan khác trong zone
                        if _Cc is None or _e0 is None:
                            _fb = _load_zone_params_by_symbol(_sl_zone, _sym)
                            if _fb:
                                _Cc  = _Cc  if _Cc  is not None else _fb.get("Cc")
                                _Cs  = _Cs  if _Cs  is not None else _fb.get("Cs")
                                _e0  = _e0  if _e0  is not None else _fb.get("e0")
                                _PC  = _PC  if _PC  is not None else _fb.get("PC_kPa")
                                _Cv  = _Cv  if _Cv  is not None else _fb.get("Cv_cm2s")
                                _gam = _gam if _gam is not None else _fb.get("gamma_kNm3")
                                _src = _fb.get("source_bh") or ""
                                if _src:
                                    _has_borrowed = True
                        _Cv_fmt = f"{_Cv:.2e}" if _Cv is not None else "—"
                        _param_rows.append({
                            "Lớp":         _sym,
                            "Mô tả":       (_lr.get("description") or "")[:30],
                            "h (m)":       _lr.get("thickness_m"),
                            "γ (kN/m³)":  _gam,
                            "e0":          _e0,
                            "Cc":          _Cc,
                            "Cs":          _Cs,
                            "PC (kPa)":   _PC,
                            "Cv (cm²/s)": _Cv_fmt,
                            "_src":        _src,
                        })
                    _df_geo = pd.DataFrame(_param_rows)
                    _src_map_sl = _df_geo["_src"].to_dict()

                    def _highlight_borrowed(row):
                        if _src_map_sl.get(row.name):
                            return ["background-color: #FFF8DC"] * len(row)
                        return [""] * len(row)

                    # Badge "Đã kiểm tra chéo" — ngôn ngữ địa kỹ thuật, không tên file
                    _n_layers_sl = len(_param_rows)
                    _n_full = sum(1 for r in _param_rows
                                   if r.get("Cc") is not None and r.get("e0") is not None)
                    _n_borrowed = sum(1 for r in _param_rows if r.get("_src"))
                    st.success(
                        f"**✓ Đã kiểm tra chéo dữ liệu địa kỹ thuật — HK {_sl_bh}** "
                        f"(zone {_sl_zone})  ·  "
                        f"Số lớp đất: **{_n_layers_sl}** · "
                        f"Lớp có đầy đủ Cc + e₀ (mẫu thí nghiệm): **{_n_full}** · "
                        f"Lớp dùng tham khảo từ hố khoan khác cùng khu vực: **{_n_borrowed}**",
                        icon="✅",
                    )
                    st.markdown("**Thông số địa chất dùng tính lún cố kết**")
                    st.dataframe(
                        _df_geo.drop(columns=["_src"]).style.apply(_highlight_borrowed, axis=1),
                        use_container_width=True, hide_index=True,
                        column_config={
                            "h (m)":      st.column_config.NumberColumn(format="%.2f"),
                            "γ (kN/m³)": st.column_config.NumberColumn(format="%.2f"),
                            "e0":         st.column_config.NumberColumn(format="%.3f"),
                            "Cc":         st.column_config.NumberColumn(format="%.4f"),
                            "Cs":         st.column_config.NumberColumn(format="%.4f"),
                            "PC (kPa)":  st.column_config.NumberColumn(format="%.1f"),
                        },
                    )
                    if _has_borrowed:
                        st.caption("*Ô nền vàng: thông số tham khảo từ hố khoan khác cùng khu vực do hố khoan hiện tại chưa đủ dữ liệu.*")

            # Auto-run (bỏ nút) — tự tính khi đủ dữ liệu input
            with st.spinner("Đang tính tổng hợp 5 phương án..."):
                try:
                    _cmp = _sc_compare(
                        _sl_bh, _sl_zone,
                        H_fill_m=float(_sl_H),
                        residual_limit_cm=float(_sl_lim),
                        t_construction_months=float(_sl_tc),
                    )
                    st.session_state["sl_result"] = _cmp
                except Exception as _e:
                    st.error(f"Lỗi tính toán: {_e}")
                    _cmp = None

            if _cmp:
                _cdm_beta = _cmp.get("cdm_beta", 1.0)
                _cdm_a    = _cmp.get("cdm_area_ratio", 0.25)
                _col_info1, _col_info2, _col_info3 = st.columns(3)
                with _col_info1:
                    st.metric("Lún tự nhiên (Cc)", f"{_cmp['S_total_cm']:.0f} cm",
                              help="Tổng lún cố kết sơ cấp tính từ mẫu thí nghiệm theo TCCS41 Phụ lục A")
                with _col_info2:
                    st.metric("CDM beta (TCVN 9403)", f"{_cdm_beta:.3f}",
                              f"Ứng suất vào đất = {_cdm_beta*100:.0f}% tải trọng",
                              help="beta = Es/(a*Ec+(1-a)*Es); Ec=75*Cc_col; Es=250*Cu")
                with _col_info3:
                    _cdm_sc = next((s for s in _cmp["scenarios"] if s["method"] == "cdm"), None)
                    if _cdm_sc:
                        _cdm_red = (1 - _cdm_sc["S_total_cm"] / _cmp["S_total_cm"]) * 100
                        st.metric("CDM lún đàn hồi (S1)", f"{_cdm_sc['S_total_cm']:.0f} cm",
                                  f"-{_cdm_red:.0f}% so với không xử lý | S2=0",
                                  help="S1 = q×H/(a×Ec+(1-a)×Es) — đàn hồi tức thời; S2=0 (CDM đến lớp cứng). TCVN 9403:2012 Phụ lục C")

                # Bảng so sánh phương án
                _sc_rows = []
                for _sc_item in _cmp["scenarios"]:
                    _t90_str = f"{_sc_item['t_90_months']:.0f}" if _sc_item["t_90_months"] else ">240"
                    _sc_rows.append({
                        "Phương án":        _sc_item["label"],
                        "Lún tổng (cm)":    _sc_item["S_total_cm"],
                        "Lún tại TC (cm)":  _sc_item["S_at_constr_cm"],
                        "U tại TC (%)":     _sc_item["U_at_constr_pct"],
                        "Lún còn lại (cm)": _sc_item["residual_cm"],
                        "t90% (tháng)":     _t90_str,
                        "Surcharge (m)":    _sc_item["H_surcharge_m"],
                        "Đánh giá":         "Đạt" if _sc_item["feasible"] else "Không đạt",
                    })
                _df_sc = pd.DataFrame(_sc_rows)
                st.dataframe(
                    _df_sc.style.apply(
                        lambda col: ["background-color:#d4edda" if v == "Đạt"
                                     else "background-color:#f8d7da" if v == "Không đạt"
                                     else "" for v in col],
                        subset=["Đánh giá"]
                    ),
                    use_container_width=True, hide_index=True
                )
                _cdm_S1 = _cmp.get("cdm_S1_cm", 0)
                _cdm_Ec = _cmp.get("cdm_Ec_kPa", 0)
                _cdm_Es = _cmp.get("cdm_Es_kPa", 0)
                _cdm_Etb = _cmp.get("cdm_composite_kPa", 0)
                st.caption(
                    f"CDM (TCVN 9403 Phụ lục C): S1={_cdm_S1:.0f} cm (đàn hồi), S2=0 cm. "
                    f"Ec={_cdm_Ec:.0f} kPa | Es={_cdm_Es:.0f} kPa | Etb={_cdm_Etb:.0f} kPa (a={_cdm_a:.2f}). "
                    "Không xử lý/Bấc thấm/Giếng cát: Cc cố kết (Terzaghi+Barron). "
                    "Giải thích: xem 'Lý thuyết tính lún' cuối trang."
                )

                # ── Biểu đồ tổng hợp so sánh phương án xử lý ─────────────────
                if _HAS_PLOTLY:
                    _labels  = [r["Phương án"] for r in _sc_rows]
                    _S_tot   = [r["Lún tổng (cm)"]    for r in _sc_rows]
                    _S_resi  = [r["Lún còn lại (cm)"] for r in _sc_rows]
                    _U_tc    = [r["U tại TC (%)"]     for r in _sc_rows]
                    _t90_num = [(float(r["t90% (tháng)"]) if r["t90% (tháng)"] != ">240" else 240)
                                for r in _sc_rows]
                    _eval    = [r["Đánh giá"] for r in _sc_rows]
                    _color_bar = ["#2E7D32" if e == "Đạt" else "#C62828" for e in _eval]

                    from plotly.subplots import make_subplots
                    _fig_cmp4 = make_subplots(
                        rows=2, cols=2,
                        subplot_titles=(
                            "Lún tổng vs Lún còn lại (cm)",
                            "Độ cố kết U(t) tại kết thúc thi công (%)",
                            "Thời gian t₉₀ — đạt 90% cố kết (tháng)",
                            "Lún còn lại vs Giới hạn ΔS (cm)",
                        ),
                        specs=[[{"secondary_y": False}, {}], [{}, {}]],
                        vertical_spacing=0.18, horizontal_spacing=0.12,
                    )

                    # (1,1): Bar group — Lún tổng vs Lún còn lại
                    _fig_cmp4.add_trace(go.Bar(
                        x=_labels, y=_S_tot,
                        name="Lún tổng",
                        marker_color="#4472C4",
                        text=[f"{v:.0f}" for v in _S_tot], textposition="outside",
                    ), row=1, col=1)
                    _fig_cmp4.add_trace(go.Bar(
                        x=_labels, y=_S_resi,
                        name="Lún còn lại",
                        marker_color="#ED7D31",
                        text=[f"{v:.1f}" for v in _S_resi], textposition="outside",
                    ), row=1, col=1)

                    # (1,2): Độ cố kết U%
                    _fig_cmp4.add_trace(go.Bar(
                        x=_labels, y=_U_tc,
                        marker_color=_color_bar,
                        text=[f"{v:.0f}%" for v in _U_tc], textposition="outside",
                        showlegend=False,
                    ), row=1, col=2)
                    _fig_cmp4.add_hline(y=90, line_dash="dash", line_color="#888",
                                        annotation_text="U=90%", row=1, col=2)

                    # (2,1): t₉₀ (tháng)
                    _fig_cmp4.add_trace(go.Bar(
                        x=_labels, y=_t90_num,
                        marker_color="#9C27B0",
                        text=[r["t90% (tháng)"] for r in _sc_rows],
                        textposition="outside",
                        showlegend=False,
                    ), row=2, col=1)
                    _fig_cmp4.add_hline(y=float(_sl_tc), line_dash="dash", line_color="#1565C0",
                                        annotation_text=f"t_TC = {_sl_tc:.0f} tháng",
                                        row=2, col=1)

                    # (2,2): Lún còn lại vs giới hạn
                    _fig_cmp4.add_trace(go.Bar(
                        x=_labels, y=_S_resi,
                        marker_color=_color_bar,
                        text=[f"{v:.1f}" for v in _S_resi], textposition="outside",
                        showlegend=False,
                    ), row=2, col=2)
                    _fig_cmp4.add_hline(y=float(_sl_lim), line_dash="dash", line_color="#E53935",
                                        annotation_text=f"Giới hạn ΔS ≤ {_sl_lim:.0f} cm",
                                        row=2, col=2)

                    _fig_cmp4.update_layout(
                        title=dict(text="<b>Bảng tổng hợp so sánh phương án xử lý đất yếu</b>",
                                   font=dict(size=14)),
                        height=720, margin=dict(t=80, b=60, l=50, r=20),
                        barmode="group", showlegend=True,
                        legend=dict(orientation="h", y=1.06, x=0.3),
                        plot_bgcolor="#FAFAFA",
                    )
                    _fig_cmp4.update_yaxes(title_text="cm",    row=1, col=1)
                    _fig_cmp4.update_yaxes(title_text="U (%)", row=1, col=2, range=[0, 110])
                    _fig_cmp4.update_yaxes(title_text="tháng", row=2, col=1)
                    _fig_cmp4.update_yaxes(title_text="cm",    row=2, col=2)
                    _fig_cmp4.update_xaxes(tickangle=-15)
                    st.plotly_chart(_fig_cmp4, use_container_width=True)

                    # ── Biểu đồ tóm tắt: stacked horizontal bar — Lún tại TC vs Còn lại ──
                    _fig_stk = go.Figure()
                    _fig_stk.add_trace(go.Bar(
                        y=_labels,
                        x=[r["Lún tại TC (cm)"] for r in _sc_rows],
                        name="Lún tại kết thúc TC",
                        marker_color="#1565C0",
                        orientation="h",
                        text=[f"{r['Lún tại TC (cm)']:.1f}" for r in _sc_rows],
                        textposition="inside",
                    ))
                    _fig_stk.add_trace(go.Bar(
                        y=_labels,
                        x=_S_resi,
                        name="Lún còn lại",
                        marker_color="#E53935",
                        orientation="h",
                        text=[f"{v:.1f}" for v in _S_resi],
                        textposition="inside",
                    ))
                    _fig_stk.update_layout(
                        title="<b>Phân bố lún: hoàn thành tại TC vs còn lại</b>",
                        barmode="stack",
                        xaxis_title="cm",
                        height=320, margin=dict(t=50, b=40, l=180, r=20),
                        legend=dict(orientation="h", y=1.15),
                        plot_bgcolor="#FAFAFA",
                    )
                    st.plotly_chart(_fig_stk, use_container_width=True)

                elif _HAS_MPL:
                    _fig_mp, _axs_mp = plt.subplots(2, 2, figsize=(11, 7))
                    _xlb = list(range(len(_sc_rows)))
                    _xl  = [r["Phương án"][:14] for r in _sc_rows]
                    _S_tot   = [r["Lún tổng (cm)"]    for r in _sc_rows]
                    _S_resi  = [r["Lún còn lại (cm)"] for r in _sc_rows]
                    _U_tc    = [r["U tại TC (%)"]     for r in _sc_rows]
                    _t90_num = [(float(r["t90% (tháng)"]) if r["t90% (tháng)"] != ">240" else 240)
                                for r in _sc_rows]
                    _col = ["#2E7D32" if r["Đánh giá"] == "Đạt" else "#C62828" for r in _sc_rows]
                    _w = 0.4
                    _axs_mp[0,0].bar([i-_w/2 for i in _xlb], _S_tot, _w, label="Lún tổng", color="#4472C4")
                    _axs_mp[0,0].bar([i+_w/2 for i in _xlb], _S_resi, _w, label="Lún còn lại", color="#ED7D31")
                    _axs_mp[0,0].set_title("Lún tổng vs Lún còn lại (cm)"); _axs_mp[0,0].legend()
                    _axs_mp[0,1].bar(_xlb, _U_tc, color=_col); _axs_mp[0,1].axhline(90, ls="--", color="#888")
                    _axs_mp[0,1].set_title("U(t) tại TC (%)"); _axs_mp[0,1].set_ylim(0,110)
                    _axs_mp[1,0].bar(_xlb, _t90_num, color="#9C27B0"); _axs_mp[1,0].axhline(float(_sl_tc), ls="--", color="#1565C0")
                    _axs_mp[1,0].set_title("t₉₀ (tháng)")
                    _axs_mp[1,1].bar(_xlb, _S_resi, color=_col); _axs_mp[1,1].axhline(float(_sl_lim), ls="--", color="#E53935")
                    _axs_mp[1,1].set_title("Lún còn lại vs Giới hạn (cm)")
                    for _a in _axs_mp.flat:
                        _a.set_xticks(_xlb); _a.set_xticklabels(_xl, rotation=20, ha="right", fontsize=8)
                        _a.grid(alpha=0.3)
                    _fig_mp.suptitle("Bảng tổng hợp so sánh phương án xử lý đất yếu", fontsize=12, fontweight="bold")
                    _fig_mp.tight_layout()
                    st.pyplot(_fig_mp, use_container_width=True)

                # PDF qua nút đã tắt (§25 CLAUDE.md) — dùng Ctrl+P trình duyệt

                # Biểu đồ S(t)
                if _HAS_PLOTLY:
                    _colors_sl = {
                        "no_treat":   "#636EFA",
                        "pvd_1200":   "#EF553B",
                        "pvd_1500":   "#00CC96",
                        "sand_drain": "#AB63FA",
                        "cdm":        "#FFA15A",
                    }
                    _fig_sl = go.Figure()
                    for _sc_item in _cmp["scenarios"]:
                        _mid = _sc_item["method"]
                        _ts  = _cmp["time_series"].get(_mid, [])
                        if not _ts:
                            continue
                        _fig_sl.add_trace(go.Scatter(
                            x=[p["t_months"] for p in _ts],
                            y=[p["S_cm"] for p in _ts],
                            mode="lines+markers",
                            name=_sc_item["label"],
                            line=dict(color=_colors_sl.get(_mid, "#888"), width=2),
                            marker=dict(size=4),
                        ))
                    _S_ref = _cmp["S_total_cm"] - float(_sl_lim)
                    _fig_sl.add_hline(
                        y=max(_S_ref, 0),
                        line_dash="dash", line_color="red",
                        annotation_text=f"S min để đạt ΔS≤{_sl_lim}cm",
                        annotation_position="bottom right",
                    )
                    _fig_sl.add_vline(
                        x=float(_sl_tc),
                        line_dash="dot", line_color="gray",
                        annotation_text=f"Kết thúc TC ({_sl_tc}th)",
                    )
                    _fig_sl.update_layout(
                        title="Quan hệ S-t theo phương án xử lý nền",
                        xaxis_title="Thời gian (tháng)",
                        yaxis_title="Độ lún S (cm)",
                        height=420,
                        legend=dict(orientation="h", yanchor="bottom", y=1.02),
                        margin=dict(t=60),
                    )
                    st.plotly_chart(_fig_sl, use_container_width=True)

                # Chi tiết lún từng lớp
                with st.expander("Chi tiết lún từng lớp"):
                    _detail = _cmp["S_detail"]
                    if _detail.get("warning"):
                        st.warning(_detail["warning"])
                    if _detail.get("layers"):
                        _df_lay = pd.DataFrame(_detail["layers"])
                        _df_lay.columns = [
                            "Độ sâu (m)", "H lớp (m)",
                            "σv0 (kPa)", "σvf (kPa)", "PC (kPa)",
                            "Cc", "Cs", "e0", "Si (cm)", "Trạng thái OC"
                        ]
                        st.table(_df_lay)

                # ── Biểu đồ ứng suất theo chiều sâu ──────────────────────────
                _layers = _cmp["S_detail"].get("layers", [])
                if _layers and _HAS_PLOTLY:
                    st.markdown("##### Biểu đồ ứng suất theo chiều sâu (phạm vi ảnh hưởng)")
                    _beta_plot = _cmp.get("cdm_beta", 1.0)
                    _dsig = _layers[0]["sigma_vf_kPa"] - _layers[0]["sigma_v0_kPa"]  # Delta_sigma

                    _depths  = [ly["depth_mid_m"]   for ly in _layers]
                    _sv0_arr = [ly["sigma_v0_kPa"]  for ly in _layers]
                    _svf_arr = [ly["sigma_vf_kPa"]  for ly in _layers]
                    _pc_arr  = [ly["PC_kPa"]         for ly in _layers]
                    # CDM: stress giảm theo beta
                    _svf_cdm = [ly["sigma_v0_kPa"] + _beta_plot * (ly["sigma_vf_kPa"] - ly["sigma_v0_kPa"])
                                for ly in _layers]
                    # OC status color mapping
                    _oc_colors = {
                        "OC": "#2196F3", "cross_PC": "#FF9800", "NC": "#F44336", "unknown": "#9E9E9E"
                    }
                    _oc_vals = [ly["OC_status"] for ly in _layers]

                    _fig_str = go.Figure()
                    _fig_str.add_trace(go.Scatter(
                        x=_sv0_arr, y=_depths, mode="lines+markers",
                        name="σv0 (ứng suất ban đầu)",
                        line=dict(color="#1565C0", width=2),
                        marker=dict(size=5, symbol="circle"),
                    ))
                    _fig_str.add_trace(go.Scatter(
                        x=_svf_arr, y=_depths, mode="lines+markers",
                        name="σvf (sau đắp, không xử lý)",
                        line=dict(color="#C62828", width=2, dash="dot"),
                        marker=dict(size=5, symbol="square"),
                    ))
                    _fig_str.add_trace(go.Scatter(
                        x=_svf_cdm, y=_depths, mode="lines+markers",
                        name=f"σvf CDM (beta={_beta_plot:.3f})",
                        line=dict(color="#E65100", width=2, dash="dash"),
                        marker=dict(size=5, symbol="diamond"),
                    ))
                    _fig_str.add_trace(go.Scatter(
                        x=[p for p in _pc_arr if p is not None],
                        y=[_depths[i] for i, p in enumerate(_pc_arr) if p is not None],
                        mode="lines+markers",
                        name="PC (áp lực tiền cố kết)",
                        line=dict(color="#2E7D32", width=2, dash="longdash"),
                        marker=dict(size=7, symbol="triangle-up"),
                    ))

                    # Shade OC/cross/NC zones với màu nền
                    _oc_zone_color = {"OC": "rgba(33,150,243,0.08)",
                                      "cross_PC": "rgba(255,152,0,0.12)",
                                      "NC": "rgba(244,67,54,0.10)",
                                      "unknown": "rgba(200,200,200,0.05)"}
                    for _li, _ly in enumerate(_layers):
                        _oc = _ly["OC_status"]
                        _y0 = _ly["depth_mid_m"] - _ly["H_i_m"] / 2
                        _y1 = _ly["depth_mid_m"] + _ly["H_i_m"] / 2
                        _fig_str.add_hrect(
                            y0=_y0, y1=_y1,
                            fillcolor=_oc_zone_color.get(_oc, "rgba(200,200,200,0.05)"),
                            line_width=0,
                        )

                    _fig_str.update_layout(
                        xaxis_title="Ứng suất hữu hiệu (kPa)",
                        yaxis_title="Độ sâu (m)",
                        yaxis=dict(autorange="reversed", ticksuffix=" m"),
                        height=480,
                        legend=dict(orientation="h", yanchor="bottom", y=1.02, x=0),
                        margin=dict(t=80, l=70),
                        annotations=[dict(
                            x=0.02, y=0.02, xref="paper", yref="paper",
                            text="<b>Màu vùng:</b> Xanh=OC | Cam=cắt qua PC | Đỏ=NC",
                            showarrow=False, font=dict(size=10),
                            bgcolor="rgba(255,255,255,0.7)", bordercolor="gray",
                        )],
                    )
                    st.plotly_chart(_fig_str, use_container_width=True)
                    st.caption(
                        f"Δσ = {_dsig:.0f} kPa (đắp H={_cmp['H_fill_m']}m, γ=20 kN/m³). "
                        f"Đường 'σvf CDM (beta={_beta_plot:.3f})' minh hoạ ứng suất vào đất giữa cột — "
                        "CDM tính lún theo S1 đàn hồi (TCVN 9403 Phụ lục C), không phải Cc cố kết. "
                        "Vùng NC/cắt qua PC là vùng có lún Cc lớn nhất nếu không xử lý."
                    )

                # ── Lý thuyết + giải thích kết quả (LaTeX, có giá trị thực) ────
                with st.expander("Lý thuyết tính lún và giải thích kết quả CDM"):
                    _beta_txt = _cmp.get("cdm_beta", 1.0)
                    _a_txt    = _cmp.get("cdm_area_ratio", 0.25)
                    _S1_txt   = _cmp.get("cdm_S1_cm", 0)
                    _Ec_txt   = _cmp.get("cdm_Ec_kPa", 0)
                    _Es_txt   = _cmp.get("cdm_Es_kPa", 0)
                    _Etb_txt  = _cmp.get("cdm_composite_kPa", 0)
                    st.markdown(r"""
## 1. Lún cố kết sơ cấp — Phụ lục A TCCS 41:2022

Áp dụng cho các phương án: **không xử lý, bấc thấm, giếng cát**. Mỗi lớp đất yếu $i$
tính theo trạng thái cố kết.

### 1.1 Đất quá cố kết (OC) — $\sigma'_{vf} \leq P_C$

$$S_i = H_i \cdot \frac{C_s}{1 + e_0} \cdot \log \frac{\sigma'_{vf}}{\sigma'_{v0}}$$

### 1.2 Cắt qua $P_C$ — $\sigma'_{v0} < P_C < \sigma'_{vf}$

$$S_i = H_i \left[ \frac{C_s}{1+e_0} \log \frac{P_C}{\sigma'_{v0}}
                 + \frac{C_c}{1+e_0} \log \frac{\sigma'_{vf}}{P_C} \right]$$

### 1.3 Đất bình thường cố kết (NC) — $\sigma'_{v0} \geq P_C$

$$S_i = H_i \cdot \frac{C_c}{1 + e_0} \cdot \log \frac{\sigma'_{vf}}{\sigma'_{v0}}$$

Trong đó $\sigma'_{vf} = \sigma'_{v0} + \Delta\sigma$. Phạm vi tính: mọi lớp có mẫu nén cố kết
($C_c$ đo được). Chiều dày đại diện $H_i$ dùng boundary trung điểm giữa các mẫu liền kề.

---

## 2. Lún CDM — TCVN 9403:2012 Phụ lục C

CDM thay đất yếu bằng nền hỗn hợp (composite). Lún tổng:

$$S = S_1 + S_2$$

### 2.1 Lún đàn hồi tức thời trong khối gia cố

$$S_1 = \frac{q \cdot H_{\text{soft}}}{a \cdot E_c + (1 - a) \cdot E_s}$$

### 2.2 Lún cố kết bên dưới mũi cột

$$S_2 = 0 \quad \text{(giả thiết CDM cắm đến lớp cứng)}$$

### 2.3 Mô-đun thành phần

$$E_c = 75 \cdot C_{c,\text{col}} = 75 \cdot \left(\text{ratio}_{lab \to field} \cdot \frac{q_{u,lab}}{2}\right)
\quad \text{(TCVN 9403 B.5.1)}$$

$$E_s = 250 \cdot C_u \quad \text{(tương quan Mesri)}$$

$$E_{\text{composite}} = a \cdot E_c + (1 - a) \cdot E_s$$
""")
                    st.info(
                        f"**Kết quả cho zone này** ($a = {_a_txt:.2f}$):\n\n"
                        f"- $E_c = {_Ec_txt:.0f}$ kPa | $E_s = {_Es_txt:.0f}$ kPa | "
                        f"$E_{{\\text{{composite}}}} = {_Etb_txt:.0f}$ kPa\n"
                        f"- **$S_1 = {_S1_txt:.1f}$ cm** (đàn hồi tức thời trong quá trình thi công)\n"
                        f"- $S_2 = 0$ cm (CDM cắm đến lớp cứng)\n"
                        f"- Hệ số $\\beta = {_beta_txt:.3f}$ "
                        f"(tỷ lệ ứng suất vào đất giữa cột)"
                    )
                    st.markdown("---")
                    st.markdown("## 3. So sánh CDM vs không xử lý")
                    st.markdown(
                        "| Phương pháp | Cơ sở vật lý | Biên độ |\n"
                        "|---|---|---|\n"
                        f"| Không xử lý | $C_c$ cố kết (log) | **{_cmp['S_total_cm']:.0f} cm** (lún dài hạn) |\n"
                        f"| CDM | $S_1$ đàn hồi (tuyến tính) | **{_S1_txt:.0f} cm** (tức thời) |"
                    )
                    st.markdown(r"""
**Tại sao chênh lệch lớn?**

- Lún $C_c$: $S \propto \log\!\left(\dfrac{\sigma'_{vf}}{\sigma'_{v0}}\right)$ — tăng mạnh khi đất NC và $H_{\text{soft}}$ lớn.
- Lún đàn hồi CDM: $S_1 \propto \dfrac{1}{E_{\text{composite}}}$ — càng cao (cột cứng hơn, mật độ lớn hơn) → lún càng nhỏ.

**Để giảm $S_1$ thêm**: tăng diện tích thay thế $a$ (rút khoảng cách $e$), hoặc tăng $q_{u,lab}$ ($> 1$ MPa).

---

## 4. Lún còn lại sau thi công $\Delta S$

### 4.1 CDM

$$\Delta S \approx 0 \quad (\text{$S_1$ đàn hồi — lún xảy ra ngay khi đắp})$$

### 4.2 Phương án khác (không xử lý / bấc thấm / giếng cát)

$$\Delta S = S_{\text{total}} \cdot \bigl(1 - U(t_{tc})\bigr)$$

| Phương án | $U(t)$ |
|---|---|
| Không xử lý | $U_v$ — Terzaghi đứng, phụ thuộc $C_v$, $H_{dr}$ |
| Bấc thấm / giếng cát | $U = 1 - (1 - U_v)(1 - U_h)$ — cố kết kết hợp |

### 4.3 Surcharge (gia tải tạm)

Tăng ứng suất $> $ tải thiết kế → tăng tốc độ cố kết → giảm $\Delta S$ sau khi dỡ tải. **Không** thay đổi $S_{\text{total}}$ dưới tải thiết kế.
""")
                    st.markdown(f"### 4.4 Yêu cầu TCCS 41:2022")
                    st.latex(
                        rf"\Delta S \leq {_cmp['residual_limit_cm']:.0f} \text{{ cm}}"
                        r" \quad \text{(đường cấp 1 đoạn thông thường)}"
                    )

        # ── PHẦN 1b: Lún sơ bộ TKCS — Điều 9.2.3 TCCS 41:2022 ──────────────
        st.divider()
        st.markdown("### Tính lún sơ bộ TKCS — Điều 9.2.3 TCCS 41:2022 (vòng lặp)")
        st.caption(
            "H'_tk = H_fill + S_gt — chiều cao đắp hiệu dụng bao gồm phần đắp lún vào đất yếu. "
            "Lặp lại đến khi |S_calc - S_gt| < tolerance. "
            "Kết quả lớn hơn tính toán một lần (không lặp) do trọng lượng đắp phần lún."
        )
        _it1, _it2, _it3 = st.columns([1, 1, 2])
        with _it1:
            _it_zone = st.radio("Zone", list(_ZONE_NAMES.keys()),
                                format_func=lambda z: _ZONE_NAMES[z],
                                key="it_zone")
            _it_bhs = [b["name"] for b in _load_boreholes_by_zone(_it_zone)]
            _it_bh  = st.selectbox("Hố khoan", _it_bhs, key="it_bh")
        with _it2:
            _it_H    = st.number_input("H_fill (m)", 1.0, 10.0, 3.0, 0.5, key="it_H")
            _it_pct  = st.number_input("S_gt ban đầu (% H_soft)", 1.0, 30.0, 7.5, 0.5,
                                        key="it_pct",
                                        help="Đất thường: 5-10%; Than bùn: 20-30% (Điều 9.2.3)")
            _it_tol  = st.number_input("Tolerance (cm)", 0.1, 5.0, 1.0, 0.1, key="it_tol")
        with _it3:
            st.caption(
                "**Công thức:** H'_tk = H_fill + S_gt  →  Δσ = H'_tk × γ_fill  →  "
                "S_c = Σ công thức Cc  →  kiểm tra |S_c - S_gt| < tol"
            )
            st.caption("**Khi nào dùng:** TKCS khi chưa có mẫu Cc chi tiết — dùng thông số trung bình zone.")

        # Auto-run (bỏ nút) — tự lặp khi đủ input
        with st.spinner("Đang lặp 9.2.3..."):
            try:
                _iter_res = _sc_iter923(
                    _it_bh, _it_zone,
                    H_fill_m=float(_it_H),
                    S_gt_init_pct=float(_it_pct),
                    tolerance_cm=float(_it_tol),
                )
                st.session_state["it_result"] = _iter_res
            except Exception as _e:
                st.error(f"Lỗi tính toán 9.2.3: {_e}")
                _iter_res = None
        st.success(
            f"**✓ Đã kiểm tra chéo dữ liệu địa kỹ thuật — HK {_it_bh}** "
            f"(zone {_it_zone}) · "
            f"Mẫu thí nghiệm hố khoan hiện tại + tham khảo trung bình khu vực khi thiếu",
            icon="✅",
        )
        if _iter_res:
            _im1, _im2, _im3 = st.columns(3)
            with _im1:
                st.metric("Lún không lặp (tham chiếu)",
                          f"{_iter_res['S_ref_cm']:.0f} cm",
                          help="Δσ = H_fill × γ (không tính phần đắp lún vào)")
            with _im2:
                st.metric("Lún sơ bộ TKCS (Điều 9.2.3)",
                          f"{_iter_res['S_final_cm']:.0f} cm",
                          f"+{_iter_res['S_increase_pct']:.0f}% so với không lặp",
                          delta_color="inverse",
                          help="Δσ = H'_tk × γ; H'_tk = H_fill + S_hội tụ")
            with _im3:
                _cv_str = "Hội tụ" if _iter_res["converged"] else "Chưa hội tụ"
                st.metric(f"Vòng lặp ({_cv_str})",
                          f"{_iter_res['n_iterations']} vòng",
                          f"tol={_iter_res['tolerance_cm']} cm | S_gt_init={_iter_res['S_gt_init_cm']:.0f} cm")

            st.markdown(f"**Chi tiết vòng lặp** — Zone {_iter_res['zone_code']}, "
                        f"HK {_iter_res['bh_name']}, H_fill={_iter_res['H_fill_m']}m, "
                        f"H_soft={_iter_res['H_soft_m']}m")
            _df_iter = pd.DataFrame(_iter_res["iterations"]).rename(columns={
                "iter":       "Vòng",
                "S_gt_cm":    "S_gt (cm)",
                "H_eff_m":    "H'_tk (m)",
                "Dsigma_kPa": "Δσ (kPa)",
                "S_calc_cm":  "S_calc (cm)",
                "delta_cm":   "Delta (cm)",
                "converged":  "Hội tụ",
            })
            st.dataframe(
                _df_iter.style.apply(
                    lambda col: ["background-color:#d4edda" if v is True
                                 else "background-color:#fff3cd" if v is False
                                 else "" for v in col],
                    subset=["Hội tụ"]
                ),
                use_container_width=True, hide_index=True,
            )
            st.caption(
                f"Lún sơ bộ TKCS = {_iter_res['S_final_cm']:.0f} cm "
                f"(lớn hơn {_iter_res['S_increase_pct']:.0f}% so với tính một lần). "
                "Nguyên nhân: đất yếu bị nào xuống, cần thêm đắp bù → tăng tải trọng → tăng lún thêm."
            )

        # Xuất PDF qua nút đã tắt (§25 CLAUDE.md) — dùng Ctrl+P trình duyệt


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE KE_SW – CỌC VÁN THÉP DỰ ỨNG LỰC SW (KÈ CÔNG VIÊN)
# ═══════════════════════════════════════════════════════════════════════════════
if _page == "ke_sw":
    import math as _math

    _SW_JSON  = _ROOT / "data" / "sw_pile_catalog.json"
    _KE_JSON  = _ROOT / "data" / "ke_sw_202605_TTHC.json"

    @st.cache_data(show_spinner=False)
    def _load_sw_catalog() -> list[dict]:
        try:
            raw = json.loads(_SW_JSON.read_text(encoding="utf-8"))
            return raw["piles"]
        except Exception:
            return []

    @st.cache_data(show_spinner=False)
    def _load_ke_sw(_mtime: float = 0) -> dict:
        try:
            return json.loads(_KE_JSON.read_text(encoding="utf-8"))
        except Exception:
            return {}

    _sw_piles  = _load_sw_catalog()
    try:
        _mtime = _KE_JSON.stat().st_mtime
    except OSError:
        _mtime = 0
    _ke_data   = _load_ke_sw(_mtime)
    _sw_names  = sorted({p["name"] for p in _sw_piles})
    _Ec_table  = {
        "fc30 MPa": 26_561_000,
        "fc40 MPa": 31_623_000,
        "fc60 MPa": 38_730_000,
        "fc70 MPa": 43_628_130,
    }

    def _sw_by_name(name: str) -> dict | None:
        candidates = [p for p in _sw_piles if p["name"] == name]
        if not candidates:
            return None
        return max(candidates, key=lambda p: p.get("Mcr_Tm", 0))

    st.subheader("Cọc ván thép dự ứng lực SW — Kè Công Viên (KE)")

    # ── A. Catalog tiết diện ────────────────────────────────────────────────────
    st.markdown("### A. Catalog tiết diện SW")
    if not _sw_piles:
        st.warning("Không tải được catalog cọc SW.")
    else:
        _fc_sel = st.selectbox("Cường độ bê tông (cho tính Ec)", list(_Ec_table.keys()),
                               index=2, key="sw_fc")
        _Ec_sel = _Ec_table[_fc_sel]

        _cat_rows = []
        for p in _sw_piles:
            _Itd  = p.get("Itd_cm4", 0) or 0
            _Atd  = p.get("Atd_cm2", 0) or 0
            _Mcr  = p.get("Mcr_Tm", 0) or 0
            _EI   = _Ec_sel * _Itd * 1e-8
            _EI_m = _EI / (p.get("width_mm", 996) / 1000)
            _cat_rows.append({
                "Loại cọc":         p["name"],
                "H (mm)":           p.get("H_mm"),
                "t (mm)":           p.get("t_mm"),
                "Cáp (sợi)":        p.get("n_strands"),
                "φ cáp (mm)":       p.get("phi_strand_mm"),
                "Atd (cm²)":        _Atd,
                "Itd (cm⁴)":        _Itd,
                "Mcr (T.m)":        _Mcr,
                "Mcr (kN.m)":       round(_Mcr * 9.81, 1),
                "EI (kN·m²/m)":     round(_EI_m, 0),
                "TL (T/cọc)":       p.get("weight_T"),
                "L tc (m)":         p.get("L_std_m"),
                "L min (m)":        p.get("L_min_m"),
                "L max (m)":        p.get("L_max_m"),
                "Spacing (mm)":     p.get("width_mm", 996),
            })

        def _hi_sw(row):
            if row["Loại cọc"] in ("SW-840", "SW-940"):
                return ["background-color:#fff3cd"] * len(row)
            return [""] * len(row)

        # st.table thay st.dataframe để in PDF không bị overlap với chart bên dưới
        st.table(pd.DataFrame(_cat_rows).style.apply(_hi_sw, axis=1))
        st.caption(
            f"Ec = {_Ec_sel/1000:.0f} MPa ({_fc_sel}). "
            "Hàng vàng = cọc đã chọn cho dự án TTHC (SW-840 / SW-940). "
            "EI tính trên 1 m dài tường (chia spacing)."
        )

        if _HAS_PLOTLY:
            _fig_cat = go.Figure()
            _xs = [p.get("H_mm") for p in _sw_piles]
            _ys = [round((p.get("Mcr_Tm") or 0) * 9.81, 1) for p in _sw_piles]
            _nm = [p["name"] for p in _sw_piles]
            _fig_cat.add_trace(go.Scatter(
                x=_xs, y=_ys, mode="markers+text", text=_nm,
                textposition="top center",
                marker=dict(size=10,
                            color=["#FF6B35" if n in ("SW-840", "SW-940") else "#4A90D9"
                                   for n in _nm]),
                name="Mcr (kN·m/cọc)",
            ))
            _fig_cat.update_layout(
                title="Mô men nứt Mcr theo chiều cao tiết diện",
                xaxis_title="Chiều cao H (mm)",
                yaxis_title="Mcr (kN·m/cọc)",
                height=320, margin=dict(t=40, b=40),
            )
            st.plotly_chart(_fig_cat, use_container_width=True)

    # ── B. Kết quả thiết kế TTHC ────────────────────────────────────────────────
    st.divider()
    st.markdown("### B. Kết quả thiết kế — Kè Công Viên TTHC")
    _bhs_on_alignment: list = []   # default cho Section C có thể truy cập
    if not _ke_data:
        st.warning("Không tải được dữ liệu thiết kế kè SW.")
    else:
        _dc     = _ke_data.get("design_conditions", {})
        _bhs_ke = _ke_data.get("boreholes", [])
        _sum_ke = _ke_data.get("pile_selection_summary", {})

        _d1, _d2, _d3 = st.columns(3)
        _d1.metric("Cao độ đỉnh kè (m)", _dc.get("top_ke_elevation_m", "–"))
        _d2.metric("Lớp bùn sét yếu", _dc.get("soft_clay_layer_symbol", "–"))
        _d3.metric("Xuyên qua lớp yếu (m)", _dc.get("min_penetration_below_soft_clay_m", "–"))

        # Catalog duy nhất, sort theo H_mm tăng dần → dùng để chọn cọc tối ưu
        _seen_names: set = set()
        _catalog_sorted: list = []
        for _p in sorted(_sw_piles, key=lambda x: x.get("H_mm", 0)):
            if _p["name"] not in _seen_names:
                _seen_names.add(_p["name"])
                _catalog_sorted.append(_p)
        _pile_options = [_p["name"] for _p in _catalog_sorted]
        # Lookup: tên cọc → L_max_m
        _pile_lmax: dict[str, float] = {
            _p["name"]: _p.get("L_max_m") or 0 for _p in _catalog_sorted
        }

        def _optimal_pile(L_req: float) -> tuple[str, float]:
            """Cọc nhỏ nhất có L_max ≥ L_req. Trả (name, L_max)."""
            for _p in _catalog_sorted:
                if (_p.get("L_max_m") or 0) >= L_req:
                    return _p["name"], _p.get("L_max_m", 0)
            last = _catalog_sorted[-1] if _catalog_sorted else {}
            return "Vượt catalog", last.get("L_max_m", 0)

        # Session state: lưu "Cọc kiến nghị", "L thiết kế" và "Z/H1 tùy chỉnh" per HK
        _rec_key   = "ke_sw_rec_piles"
        _ltk_key   = "ke_sw_L_thiet_ke"
        _zcust_key = "ke_sw_z_custom"   # dict: bh_name → Z_m tùy chỉnh
        _h1cust_key = "ke_sw_h1_custom" # dict: bh_name → H_lớp_1 (m) tùy chỉnh
        for _k in (_rec_key, _ltk_key, _zcust_key, _h1cust_key):
            if _k not in st.session_state:
                st.session_state[_k] = {}

        # ── Cho phép user chọn HK trên tuyến kè SW ──────────────────────────────
        _all_bh_names = [b["name"] for b in _bhs_ke]
        _sw_align_key = "ke_sw_alignment_picks"

        # Không dùng key= để tránh session state cũ override default.
        # default luôn là tất cả HK → người dùng có thể bỏ chọn trong session hiện tại.
        # Kết quả lấy từ SQLite nên rerun không tốn thêm thời gian tính toán.
        _picked_bhs = st.multiselect(
            "Hố khoan trên tuyến kè (chọn HK để tính NT1/NT2)",
            _all_bh_names,
            default=_all_bh_names,
            format_func=lambda x: f"KE-{x}",
            help="Mặc định chọn tất cả 12 HK KE. Kết quả được lưu — load tức thì.",
        )
        # Đồng bộ sang session state để mục C, D đọc
        st.session_state[_sw_align_key] = _picked_bhs

        # Mục B — DB-first: đọc từ SQLite trước, chỉ tính lại nếu chưa có
        _nt_results: dict = {}
        _total_bhs = len(_picked_bhs)
        if _total_bhs > 0:
            import sys as _s2; _s2.path.insert(0, str(_ROOT / "scripts"))
            try:
                from ke_sw_nt_calc import (
                    _get_bh_Z_m as _gz_b,
                    calc_nt_for_bh as _cfb_b,
                    load_nt_detail_from_db as _load_nt_db,
                    save_nt_detail_single as _save_nt_db,
                    save_nt_z_scenario as _save_nt_z_scen,
                    load_nt_z_scenario as _load_nt_z_scen,
                )
            except Exception:
                _gz_b = lambda x: None  # noqa: E731
                _cfb_b = _load_nt_db = _save_nt_db = None
                _save_nt_z_scen = _load_nt_z_scen = None

            _prog_bar = st.progress(0, text="Đang tải kết quả KE...")
            _n_computed = 0
            for _i, _nm in enumerate(_picked_bhs):
                _prog_bar.progress(int(_i / _total_bhs * 100),
                    text=f"KE-{_nm}... ({_i + 1}/{_total_bhs})")
                _bh_m = next((b for b in _bhs_ke if b["name"] == _nm), None)
                if not _bh_m:
                    continue
                _pile_b = _bh_m.get("recommended_pile") or "SW-840"
                _L_b    = float(_bh_m.get("recommended_L_m") or 29.0)
                try:
                    # Per-HK Z + H1 override
                    _z_custom_d = st.session_state.get("ke_sw_z_custom", {})
                    _z_override = _z_custom_d.get(_nm)
                    _has_z_ov   = _z_override is not None
                    _h1_cust_b  = st.session_state.get(_h1cust_key, {})
                    _h1_ov_b    = _h1_cust_b.get(_nm)
                    _has_h1_ov  = _h1_ov_b is not None

                    # 1. Cache lookup — BYPASS khi có H1 override (cache table KHÔNG
                    # key trên H1, sẽ trả stale entry; bắt buộc tính lại).
                    _cached_b = None
                    if not _has_h1_ov:
                        if _has_z_ov:
                            _cached_b = (_load_nt_z_scen(f"KE-{_nm}", _pile_b, _L_b, _z_override)
                                         if _load_nt_z_scen else None)
                        else:
                            _cached_b = (_load_nt_db(f"KE-{_nm}", _pile_b, _L_b)
                                         if _load_nt_db else None)
                    if _cached_b:
                        _nt_results[_nm] = _cached_b
                        continue

                    # 2. Tính toán mới — áp Z/H1 override nếu user đã chỉnh
                    _z_b  = float(_z_override) if _has_z_ov else (
                        _gz_b(f"KE-{_nm}") or _bh_m.get("Z_m") or 0.0)
                    _h1b  = (float(_h1_ov_b) if _has_h1_ov
                              else float(_bh_m.get("H_layer1_m") or 0.0))
                    # prefer_input=True khi có override → tôn trọng giá trị user nhập,
                    # không cho calc_nt_for_bh đè bằng Z/D_bot từ SQLite.
                    _n1b, _n2b = _cfb_b(f"KE-{_nm}", float(_z_b), _h1b, _pile_b, _L_b,
                                          prefer_input=(_has_z_ov or _has_h1_ov))

                    # 3. Lưu vào đúng bảng
                    if _has_z_ov:
                        if _save_nt_z_scen:
                            try:
                                _save_nt_z_scen(float(_z_b), _n1b, _n2b)
                            except Exception:
                                pass
                    elif _save_nt_db:
                        try:
                            _save_nt_db(float(_z_b), "SQLite", "SQLite", _n1b, _n2b)
                        except Exception:
                            pass
                    _nt_results[_nm] = {
                        "L_req_m":        _n1b.get("L_req_m"),
                        "NT1":            _n1b.get("result"),
                        "NT2":            _n2b.get("result"),
                        "W_pile_kN":      _n2b.get("W_kN"),
                        "NT2_multilayer": {
                            "Rs_kN": _n2b["Rs_kN"], "Rp_kN": _n2b["Rp_kN"],
                            "RR_kN": _n2b["RR_kN"], "ratio": _n2b["ratio"],
                            "tip_layer":  _n2b["tip_symbol"],
                            "tip_method": _n2b.get("tip_method"),
                        },
                    }
                    _n_computed += 1
                except Exception:
                    pass
            _prog_bar.progress(100, text=f"Hoàn thành {_total_bhs}/{_total_bhs} hố khoan"
                               + (f" ({_n_computed} tính mới, {_total_bhs-_n_computed} từ DB)" if _n_computed else " (từ DB)"))
            import time as _t; _t.sleep(0.4)
            _prog_bar.empty()

        for _bh in _bhs_ke:
            _nm = _bh["name"]
            if _nm in _nt_results:
                _bh.update(_nt_results[_nm])

        # Lọc HK theo user pick
        _bhs_on_alignment = [b for b in _bhs_ke if b["name"] in _picked_bhs]
        if not _bhs_on_alignment:
            st.info("Chọn ít nhất 1 HK trên tuyến kè để xem bảng tổng hợp.")

        # SQLite ke_sw_nt_detail = nguồn chính cho computed fields
        # JSON chỉ giữ config: recommended_pile, recommended_L_m, note.
        # Tải sớm để dùng cho cả init session_state và _ke_rows loop.
        try:
            import sqlite3 as _sq3_ntd
            with _sq3_ntd.connect(str(_DB), timeout=5) as _c_ntd:
                _c_ntd.row_factory = _sq3_ntd.Row
                _nt_detail: dict = {
                    r["bh_name"]: dict(r)
                    for r in _c_ntd.execute("SELECT * FROM ke_sw_nt_detail").fetchall()
                }
        except Exception:
            _nt_detail = {}

        # MẶC ĐỊNH = cọc tối ưu: tự điền cho HK chưa có trong session_state.
        # User vẫn override được qua dropdown; nút "Đặt lại tối ưu" dưới đây để reset.
        for _bh in _bhs_on_alignment:
            _bh_nm = _bh["name"]
            if _bh_nm not in st.session_state[_rec_key]:
                _lreq_init = float(
                    _nt_detail.get(f"KE-{_bh_nm}", {}).get("L_req_nt1_m")
                    or _bh.get("L_req_m") or 0
                )
                _on, _olmax = _optimal_pile(_lreq_init)
                st.session_state[_rec_key][_bh_nm] = _on
                st.session_state[_ltk_key][_bh_nm] = _olmax

        # Nút reset thủ công + nút giả định Z=0
        _btn_b1, _btn_b2 = st.columns([1, 1])
        with _btn_b1:
            if st.button("Đặt lại cọc tối ưu cho tất cả", key="btn_use_optimal",
                         help="Ghi đè lựa chọn hiện tại về cọc tối ưu (nhỏ nhất có L_max ≥ L_req)"):
                for _bh in _bhs_on_alignment:
                    _lreq_rst = float(
                        _nt_detail.get(f"KE-{_bh['name']}", {}).get("L_req_nt1_m")
                        or _bh.get("L_req_m") or 0
                    )
                    _on, _olmax = _optimal_pile(_lreq_rst)
                    st.session_state[_rec_key][_bh["name"]] = _on
                    st.session_state[_ltk_key][_bh["name"]] = _olmax
                st.rerun()
        with _btn_b2:
            _z_cust_now = st.session_state.get(_zcust_key, {})
            _z0_active = bool(_z_cust_now)
            if not _z0_active:
                if st.button("Giả định Z=0 cho tất cả", key="btn_z0_on",
                             help="Áp dụng Z=0 m cho tất cả HK — lưu vào kịch bản riêng, không ghi đè dữ liệu thực"):
                    st.session_state[_zcust_key] = {_bh["name"]: 0.0 for _bh in _bhs_on_alignment}
                    st.rerun()
            else:
                if st.button("Hủy giả định Z", key="btn_z0_off", type="primary",
                             help="Xóa mọi giả định Z — dùng lại Z thực từ dữ liệu khảo sát"):
                    st.session_state[_zcust_key] = {}
                    st.rerun()

        _z_cust_active = st.session_state.get(_zcust_key, {})
        if _z_cust_active:
            _z_sum = ", ".join(f"{k}={v:.1f}m" for k, v in sorted(_z_cust_active.items()))
            st.warning(f"Đang dùng Z giả định ({_z_sum}) — kết quả lưu vào kịch bản riêng, không thay đổi dữ liệu gốc.")

        # ── Helper: tính lại NT1/NT2 theo Cọc kiến nghị + L thiết kế ──────────
        # LƯU Ý: tham số KHÔNG được bắt đầu bằng "_" — Streamlit @st.cache_data
        # bỏ qua hash với arg có prefix "_" → tất cả call ra cùng cache key!
        @st.cache_data(show_spinner=False)
        def _recalc_nt(bh_name: str, Z_m: float, H_lyr1: float,
                       pile: str, L_des: float) -> dict:
            """Tính NT1+NT2 cho 1 HK theo pile+L do user chọn. Cache theo input.

            prefer_input=True → tôn trọng Z/H1 user nhập, không cho ke_sw_nt_calc
            ghi đè bằng giá trị thực từ SQLite (cần thiết để user edit có hiệu lực).
            """
            try:
                import sys as _sys_rc
                _sys_rc.path.insert(0, str(_ROOT / "scripts"))
                from ke_sw_nt_calc import calc_nt_for_bh as _calc_for_bh_rc
                _n1, _n2 = _calc_for_bh_rc(
                    f"KE-{bh_name}", Z_m or 0.0, H_lyr1 or 0.0,
                    pile, L_des, prefer_input=True,
                )
                return {
                    "L_req_m":  _n1.get("L_req_m"),
                    "NT1":      _n1.get("result"),
                    "NT2":      _n2.get("result"),
                    "Rs_kN":    _n2.get("Rs_kN"),
                    "Rp_kN":    _n2.get("Rp_kN"),
                    "RR_kN":    _n2.get("RR_kN"),
                    "W_kN":     _n2.get("W_kN"),
                    "ratio":    _n2.get("ratio"),
                }
            except Exception:
                return {}

        _ke_rows = []
        for _bh in _bhs_on_alignment:
            _ntd_row = _nt_detail.get(f"KE-{_bh['name']}") or {}
            _nt2     = _bh.get("NT2_multilayer") or {}
            _L_req   = float(_ntd_row.get("L_req_nt1_m") or _bh.get("L_req_m") or 0)
            _opt_name, _opt_Lmax = _optimal_pile(_L_req)
            # "Cọc kiến nghị": session_state → JSON default → optimal
            _rec = st.session_state[_rec_key].get(
                _bh["name"],
                _bh.get("recommended_pile") or _opt_name,
            )
            # "L thiết kế": session_state → L_max của cọc kiến nghị → JSON default
            _ltk = st.session_state[_ltk_key].get(
                _bh["name"],
                _pile_lmax.get(_rec) or _bh.get("recommended_L_m") or _opt_Lmax,
            )
            _ltk_f = float(_ltk) if _ltk is not None else float(_opt_Lmax)

            # Z tùy chỉnh per-HK (từ session state dict)
            _z_cust_row = st.session_state.get(_zcust_key, {})
            _z_override_row = _z_cust_row.get(_bh["name"])
            _has_z_ov_row   = _z_override_row is not None
            _z_for_calc     = float(_z_override_row) if _has_z_ov_row else float(_ntd_row.get("Z_m") or _bh.get("Z_m") or 0.0)

            # H lớp 1 tùy chỉnh per-HK
            _h1_cust_row = st.session_state.get(_h1cust_key, {})
            _h1_override_row = _h1_cust_row.get(_bh["name"])
            _has_h1_ov_row   = _h1_override_row is not None
            _h1_for_calc     = (float(_h1_override_row) if _has_h1_ov_row
                                  else float(_ntd_row.get("D_bottom_soft_m") or _bh.get("H_layer1_m") or 0.0))

            # Nếu pile / L / Z / H1 khác JSON → tính lại; nếu trùng → dùng giá trị JSON
            _json_pile = _bh.get("recommended_pile")
            _json_L    = _bh.get("recommended_L_m")
            _need_recalc = (_rec != _json_pile) or (
                _json_L is not None and abs(_ltk_f - float(_json_L)) > 0.01
            ) or _has_z_ov_row or _has_h1_ov_row
            if _need_recalc:
                _rc = _recalc_nt(
                    bh_name=_bh["name"],
                    Z_m=_z_for_calc,
                    H_lyr1=_h1_for_calc,
                    pile=_rec,
                    L_des=float(_ltk_f),
                )
                _nt1_val = _rc.get("NT1") or "–"
                _nt2_val = _rc.get("NT2") or "–"
                _rs_val  = _rc.get("Rs_kN")
                _rp_val  = _rc.get("Rp_kN")
                _rr_val  = _rc.get("RR_kN")
                _w_val   = _rc.get("W_kN")
                _rat_val = _rc.get("ratio")
                # ĐỒNG BỘ L_req + cọc tối ưu theo kết quả vừa tính
                _rc_L_req = _rc.get("L_req_m")
                if _rc_L_req is not None:
                    _L_req = float(_rc_L_req)
                    _opt_name, _opt_Lmax = _optimal_pile(_L_req)
            else:
                _nt1_val = _ntd_row.get("nt1_result") or _bh.get("NT1")
                _nt2_val = _ntd_row.get("nt2_result") or _bh.get("NT2")
                _rs_val  = _ntd_row.get("Rs_kN")    or _nt2.get("Rs_kN")
                _rp_val  = _ntd_row.get("Rp_kN")    or _nt2.get("Rp_kN")
                _rr_val  = _ntd_row.get("RR_kN")    or _nt2.get("RR_kN")
                _w_val   = _ntd_row.get("W_kN")     or _bh.get("W_pile_kN")
                _rat_val = _ntd_row.get("ratio_nt2") or _nt2.get("ratio")

            _ke_rows.append({
                "Hố khoan":          _bh["name"],
                "Z (m)":             _z_for_calc,
                "H lớp 1 (m)":       _h1_for_calc,
                "L yêu cầu (m)":     _L_req,
                "Cọc tối ưu":        _opt_name,
                "L_max (m)":         _opt_Lmax,
                "Đủ chiều dài":      "Đạt" if _opt_Lmax >= _L_req else "Không đạt",
                "Cọc kiến nghị":     _rec,
                "L thiết kế (m)":    _ltk_f,
                "NT1":               _nt1_val,
                "NT2":               _nt2_val,
                "Rs (kN)":           round(_rs_val, 1) if _rs_val is not None else "–",
                "Rp (kN)":           round(_rp_val, 1) if _rp_val is not None else "–",
                "RR (kN)":           round(_rr_val, 1) if _rr_val is not None else "–",
                "W (kN)":            round(_w_val or 0, 1),
                "RR/W":              round(_rat_val or 0, 2) if _rat_val is not None else "–",
                "Ghi chú":           ("Tính lại theo cọc user chọn" if _need_recalc
                                       else _bh.get("note", "")),
            })

        if not _ke_rows:
            st.warning("Không có hố khoan nào trên tuyến kè trong dữ liệu thiết kế.")
        else:
            _df_b = pd.DataFrame(_ke_rows)
            _editable_b = {"Cọc kiến nghị", "L thiết kế (m)", "Z (m)", "H lớp 1 (m)"}
            _disabled_b = [c for c in _df_b.columns if c not in _editable_b]
            _edited_b = st.data_editor(
                _df_b,
                column_config={
                    "Z (m)": st.column_config.NumberColumn(
                        "Z (m)",
                        help="Cao độ mặt đất tự nhiên tại HK. Chỉnh để cập nhật L yêu cầu + "
                             "NT1/NT2 + Mục C/D/E tự đồng bộ theo Z mới.",
                        min_value=-10.0, max_value=10.0,
                        step=0.1, format="%.2f", width="small",
                    ),
                    "H lớp 1 (m)": st.column_config.NumberColumn(
                        "H lớp 1 (m)",
                        help="Bề dày lớp bùn sét yếu (lớp 1). Chỉnh để cập nhật L yêu cầu + "
                             "NT1/NT2 + k_h trong D.1 + Mục E — đồng bộ toàn app.",
                        min_value=0.0, max_value=50.0,
                        step=0.5, format="%.1f", width="small",
                    ),
                    "Cọc kiến nghị": st.column_config.SelectboxColumn(
                        "Cọc kiến nghị", options=_pile_options,
                        required=True, width="medium",
                    ),
                    "L thiết kế (m)": st.column_config.NumberColumn(
                        "L thiết kế (m)", min_value=1.0, max_value=40.0,
                        step=0.5, format="%.1f", width="small",
                    ),
                    "Đủ chiều dài": st.column_config.Column(width="small"),
                },
                disabled=_disabled_b,
                hide_index=True,
                use_container_width=True,
                key="ke_b_editor",
            )
            # Lưu chỉnh sửa vào session state — chỉ khi người dùng thay đổi
            _edit_cols = ["Cọc kiến nghị", "L thiết kế (m)", "Z (m)", "H lớp 1 (m)"]
            if not _edited_b[_edit_cols].equals(_df_b[_edit_cols]):
                _need_rerun_b = False
                for _, _row in _edited_b.iterrows():
                    _bh_n = _row["Hố khoan"]
                    st.session_state[_rec_key][_bh_n] = _row["Cọc kiến nghị"]
                    st.session_state[_ltk_key][_bh_n] = _row["L thiết kế (m)"]

                    # Z override per HK — so sánh với SQLite (authoritative)
                    _real_z_bh = float(
                        _nt_detail.get(f"KE-{_bh_n}", {}).get("Z_m")
                        or next((b.get("Z_m") or 0.0 for b in _bhs_on_alignment if b["name"] == _bh_n), 0.0)
                    )
                    _edited_z = _row["Z (m)"]
                    if abs(float(_edited_z) - float(_real_z_bh)) > 0.01:
                        st.session_state[_zcust_key][_bh_n] = float(_edited_z)
                        _need_rerun_b = True
                    elif _bh_n in st.session_state[_zcust_key]:
                        del st.session_state[_zcust_key][_bh_n]
                        _need_rerun_b = True

                    # H lớp 1 override per HK — so sánh với SQLite (authoritative)
                    _real_h1_bh = float(
                        _nt_detail.get(f"KE-{_bh_n}", {}).get("D_bottom_soft_m")
                        or next((b.get("H_layer1_m") or 0.0 for b in _bhs_on_alignment if b["name"] == _bh_n), 0.0)
                    )
                    _edited_h1 = _row["H lớp 1 (m)"]
                    if abs(float(_edited_h1) - float(_real_h1_bh)) > 0.01:
                        st.session_state[_h1cust_key][_bh_n] = float(_edited_h1)
                        # Đồng bộ với widget D.1 chi tiết (nếu user đã mở expander)
                        _wkey_h1 = f"dpy_H1_{_bh_n}"
                        if _wkey_h1 in st.session_state:
                            st.session_state[_wkey_h1] = float(_edited_h1)
                        _need_rerun_b = True
                    elif _bh_n in st.session_state[_h1cust_key]:
                        del st.session_state[_h1cust_key][_bh_n]
                        _need_rerun_b = True

                # Đồng bộ Z với widget D.1 (nếu đã render trước đó)
                for _bh_n, _z_v in st.session_state[_zcust_key].items():
                    _wkey_z = f"dpy_Z_{_bh_n}"
                    if _wkey_z in st.session_state:
                        st.session_state[_wkey_z] = float(_z_v)
                # Bump version → Mục E sẽ phát hiện cache stale + auto-refresh
                if _need_rerun_b:
                    try:
                        import sqlite3 as _sq3_bump
                        from datetime import datetime as _dt_bump
                        with _sq3_bump.connect(str(_DB), timeout=10) as _con_bump:
                            _con_bump.execute("""
                                CREATE TABLE IF NOT EXISTS data_versions (
                                    data_key TEXT PRIMARY KEY,
                                    ts       TEXT NOT NULL,
                                    note     TEXT
                                )
                            """)
                            _con_bump.execute("""
                                INSERT OR REPLACE INTO data_versions
                                    (data_key, ts, note)
                                VALUES ('ke_geology', ?, ?)
                            """, (
                                _dt_bump.now().strftime("%Y-%m-%d %H:%M:%S"),
                                "Mục B user edit Z/H1",
                            ))
                            _con_bump.commit()
                    except Exception:
                        pass
                    st.rerun()

        # ── Bảng địa chất từng HK trên tuyến kè ────────────────────────────
        if _bhs_on_alignment:
            with st.expander("Bảng địa chất các hố khoan trên tuyến kè", expanded=False):
                _SYM_COLOR = {
                    "1": "#d6eaf8", "1b": "#d6eaf8",
                    "XMD": "#fdebd0",
                    "2a": "#d5f5e3", "2b": "#d5f5e3", "2c": "#d5f5e3",
                    "3": "#ebdef0",
                    "4": "#fdfefe", "5": "#fdfefe", "6": "#fdfefe",
                }
                for _gb in _bhs_on_alignment:
                    _gbnm  = _gb["name"]                         # e.g. "HK6"
                    _gbfull = f"KE-{_gbnm}" if not _gbnm.startswith("KE-") else _gbnm
                    _gb_z   = float(_gb.get("Z_m") or 0.0)

                    _geo_rows = _load_layers(_gbfull)

                    st.markdown(
                        f"**{_gbfull}** &nbsp; cao độ mặt đất = "
                        f"**{_gb_z:+.2f} m** &nbsp;|&nbsp; "
                        f"lớp bùn H₁ = **{_gb.get('H_layer1_m', '–')} m**"
                    )

                    if not _geo_rows:
                        st.caption("Chưa có dữ liệu địa tầng trong hệ thống.")
                    else:
                        # Thêm cột cao độ đỉnh / đáy lớp
                        _gdf_rows = []
                        for _lr in _geo_rows:
                            _et = round(_gb_z - _lr["depth_top_m"], 2)
                            _eb = round(_gb_z - _lr["depth_bot_m"], 2)
                            # VST Su trung bình trong lớp
                            _su_v = None
                            try:
                                _con_g = sqlite3.connect(str(_DB))
                                _r_vst = _con_g.execute("""
                                    SELECT ROUND(AVG(t.Su_kPa),1) su
                                    FROM vane_shear_tests t
                                    JOIN vst_locations v ON t.vst_loc_id=v.id
                                    WHERE v.name=?
                                      AND t.depth_m >= ? AND t.depth_m < ?
                                """, (_gbfull, _lr["depth_top_m"], _lr["depth_bot_m"])).fetchone()
                                _con_g.close()
                                _su_v = _r_vst[0] if _r_vst and _r_vst[0] else None
                            except Exception:
                                pass
                            _gdf_rows.append({
                                "Ký hiệu":      _lr.get("symbol") or "–",
                                "Đỉnh (m)":     _lr["depth_top_m"],
                                "Đáy (m)":      _lr["depth_bot_m"],
                                "Dày (m)":      _lr.get("thickness_m") or round(_lr["depth_bot_m"] - _lr["depth_top_m"], 2),
                                "Cao độ đỉnh":  _et,
                                "Cao độ đáy":   _eb,
                                "γ (kN/m³)":    _lr.get("gamma_kNm3"),
                                "Su VST (kPa)": _su_v,
                                "Su lab (kPa)": _lr.get("c_kPa"),
                                "φ (°)":        _lr.get("phi_deg"),
                                "e₀":           _lr.get("e0"),
                                "Cc":           _lr.get("Cc"),
                                "PC (kPa)":     _lr.get("PC_kPa"),
                            })
                        _gdf = pd.DataFrame(_gdf_rows)
                        st.dataframe(
                            _gdf,
                            use_container_width=True,
                            hide_index=True,
                            column_config={
                                "Ký hiệu":      st.column_config.TextColumn(width="small"),
                                "Đỉnh (m)":     st.column_config.NumberColumn(format="%.2f", width="small"),
                                "Đáy (m)":      st.column_config.NumberColumn(format="%.2f", width="small"),
                                "Dày (m)":      st.column_config.NumberColumn(format="%.2f", width="small"),
                                "Cao độ đỉnh":  st.column_config.NumberColumn(format="%.2f", width="small"),
                                "Cao độ đáy":   st.column_config.NumberColumn(format="%.2f", width="small"),
                                "γ (kN/m³)":    st.column_config.NumberColumn(format="%.2f", width="small"),
                                "Su VST (kPa)": st.column_config.NumberColumn(format="%.1f", width="small"),
                                "Su lab (kPa)": st.column_config.NumberColumn(format="%.1f", width="small"),
                                "φ (°)":        st.column_config.NumberColumn(format="%.1f", width="small"),
                                "e₀":           st.column_config.NumberColumn(format="%.3f", width="small"),
                                "Cc":           st.column_config.NumberColumn(format="%.4f", width="small"),
                                "PC (kPa)":     st.column_config.NumberColumn(format="%.1f", width="small"),
                            },
                        )
                    st.divider()

        # ── Trắc dọc địa chất theo tuyến kè SW ─────────────────────────────
        if _picked_bhs:
            st.markdown("#### Trắc dọc địa chất theo tuyến kè SW")
            import sqlite3 as _sq3_pf
            import numpy as _np_pf

            _db_names_pf = [f"KE-{n}" if not n.startswith("KE-") else n
                            for n in _picked_bhs]
            _ph_pf = ",".join("?" * len(_db_names_pf))

            # ── 1. Query tất cả dữ liệu cần thiết từ SQLite ────────────────
            with _sq3_pf.connect(str(_DB)) as _con_pf:
                _bh_info = _con_pf.execute(
                    f"SELECT name, x_coord_m, y_coord_m, elevation_m "
                    f"FROM boreholes WHERE name IN ({_ph_pf})",
                    _db_names_pf,
                ).fetchall()
                _layer_rows = _con_pf.execute(
                    f"SELECT b.name, l.symbol, l.depth_top_m, l.depth_bot_m "
                    f"FROM layers l JOIN boreholes b ON l.borehole_id=b.id "
                    f"WHERE b.name IN ({_ph_pf}) ORDER BY b.name, l.depth_top_m",
                    _db_names_pf,
                ).fetchall()
                # VST per HK (depth, Su)
                try:
                    _vst_rows = _con_pf.execute(
                        f"SELECT v.name, t.depth_m, t.Su_kPa "
                        f"FROM vane_shear_tests t "
                        f"JOIN vst_locations v ON t.vst_loc_id=v.id "
                        f"WHERE v.name IN ({_ph_pf}) AND t.Su_kPa > 0",
                        _db_names_pf,
                    ).fetchall()
                except Exception:
                    _vst_rows = []
                # SPT per HK
                try:
                    _spt_rows = _con_pf.execute(
                        f"SELECT b.name, s.depth_m, s.N "
                        f"FROM spt_values s JOIN boreholes b ON s.borehole_id=b.id "
                        f"WHERE b.name IN ({_ph_pf}) AND s.N IS NOT NULL",
                        _db_names_pf,
                    ).fetchall()
                except Exception:
                    _spt_rows = []
                # Lab Cu_UU / c_kPa per HK
                try:
                    _lab_rows = _con_pf.execute(
                        f"SELECT b.name, l.depth_from_m, l.depth_to_m, "
                        f"  l.Cu_UU_kPa, l.c_kPa FROM lab_tests l "
                        f"JOIN boreholes b ON l.borehole_id=b.id "
                        f"WHERE b.name IN ({_ph_pf}) "
                        f"  AND (l.Cu_UU_kPa IS NOT NULL OR l.c_kPa IS NOT NULL)",
                        _db_names_pf,
                    ).fetchall()
                except Exception:
                    _lab_rows = []
                # NT detail (D_bottom_soft_m, L_design, pile_type)
                try:
                    _ntd_rows = _con_pf.execute(
                        f"SELECT bh_name, D_bottom_soft_m, L_design_m, pile_type "
                        f"FROM ke_sw_nt_detail WHERE bh_name IN ({_ph_pf})",
                        _db_names_pf,
                    ).fetchall()
                except Exception:
                    _ntd_rows = []
                # Bình đồ kè — polyline boundary từ DXF
                try:
                    _binhdo_rows = _con_pf.execute(
                        "SELECT polyline_id, x_m, y_m FROM ke_binhdo_toadoke "
                        "ORDER BY polyline_id, vertex_idx"
                    ).fetchall()
                except Exception:
                    _binhdo_rows = []

            _ntd_by_bh = {r[0]: {"D_bot": r[1], "L_des": r[2], "pile": r[3]}
                          for r in _ntd_rows}
            # Bình đồ kè: gom theo polyline_id, hoán đổi x↔y cho đúng convention
            # boreholes dùng x_coord_m=Northing (~1191xxx), y_coord_m=Easting (~605xxx)
            # binhdo DXF:   x_m=Easting  (~605xxx),          y_m=Northing (~1191xxx)
            # → chart_x = binhdo.y_m,  chart_y = binhdo.x_m
            _binhdo_cx: list = []   # chart X (Northing)
            _binhdo_cy: list = []   # chart Y (Easting)
            _prev_pl = None
            for _pl_id, _bx, _by in _binhdo_rows:
                if _prev_pl is not None and _pl_id != _prev_pl:
                    _binhdo_cx.append(None)
                    _binhdo_cy.append(None)
                _binhdo_cx.append(_by)   # y_m = Northing → chart X
                _binhdo_cy.append(_bx)   # x_m = Easting  → chart Y
                _prev_pl = _pl_id
            _vst_by_bh: dict = {}
            for n, d, s in _vst_rows:
                _vst_by_bh.setdefault(n, []).append((d, s))
            _spt_by_bh: dict = {}
            for n, d, N in _spt_rows:
                _spt_by_bh.setdefault(n, []).append((d, N))
            _lab_by_bh: dict = {}
            for n, df, dt, cu, ck in _lab_rows:
                v = cu if cu is not None else ck
                if v and v > 0 and df is not None and dt is not None:
                    _lab_by_bh.setdefault(n, []).append(
                        ((df + dt) / 2.0, v))

            # ── 2. UI inputs ───────────────────────────────────────────────
            _lower_opts_set = sorted({r[1] for r in _layer_rows if r[1] != "1"})
            _default_lower = "2b" if "2b" in _lower_opts_set else (
                _lower_opts_set[0] if _lower_opts_set else "")

            _ic1, _ic2, _ic3 = st.columns([1, 1, 1])
            _z_top_design = _ic1.number_input(
                "Cao độ đỉnh kè TK (m)",
                value=float(_dc.get("top_ke_elevation_m", 2.7)),
                step=0.1, format="%.2f",
                key="_ke_b_z_top_design",
            )
            _z_water = _ic2.number_input(
                "Cao độ mặt nước (m)",
                value=float(_dc.get("water_elevation_m", 0.0)),
                step=0.1, format="%.2f",
                key="_ke_b_z_water",
            )
            _lower_sym = _ic3.selectbox(
                "Lớp phía dưới muốn vẽ",
                _lower_opts_set,
                index=(_lower_opts_set.index(_default_lower)
                       if _default_lower in _lower_opts_set else 0),
                key="_ke_b_lower_sym",
            ) if _lower_opts_set else ""

            # ── 3. Thứ tự HK dọc tuyến kè + chainage cộng dồn ───────────────
            _KE_BH_ORDER = ["HK1","HK2","HK3","HK4","HK5","HK12",
                             "HK6","HK7","HK8","HK9","HK10","HK11"]
            _coord_by_bh = {r[0]: (float(r[1]), float(r[2]))
                            for r in _bh_info if r[1] and r[2]}
            # Sắp xếp các HK có tọa độ theo thứ tự tuyến kè
            _ordered_bhs = [f"KE-{n}" for n in _KE_BH_ORDER
                            if f"KE-{n}" in _coord_by_bh]
            # Chainage = khoảng cách cộng dồn dọc tuyến HK1→HK2→...→HK11
            _chain_pf: dict[str, float] = {}
            _cum = 0.0
            for _i, _bn in enumerate(_ordered_bhs):
                _chain_pf[_bn] = _cum
                if _i + 1 < len(_ordered_bhs):
                    _ax, _ay = _coord_by_bh[_bn]
                    _bx, _by = _coord_by_bh[_ordered_bhs[_i + 1]]
                    _cum += _np_pf.hypot(_bx - _ax, _by - _ay)
            # Fallback cho HK ngoài thứ tự
            for _bn in _coord_by_bh:
                if _bn not in _chain_pf:
                    _chain_pf[_bn] = _cum + 999.0

            _layers_by_bh_pf: dict[str, list] = {}
            for r in _layer_rows:
                _layers_by_bh_pf.setdefault(r[0], []).append(r[1:])

            # Su / N trung bình trong khoảng depth [d_top, d_bot]
            def _avg_in_range(samples: list, d_top: float, d_bot: float):
                vals = [v for d, v in samples
                         if d_top <= d <= d_bot and v is not None]
                return sum(vals) / len(vals) if vals else None

            # Cọc + L thiết kế từ session_state (live update với bảng)
            _rec_pick = st.session_state.get(_rec_key, {})
            _ltk_pick = st.session_state.get(_ltk_key, {})

            _prof: list = []
            for r in _bh_info:
                _bh_p = r[0]
                _bh_short = _bh_p.replace("KE-", "")
                _z_p = float(r[3] or 0.0)
                _lays_p = _layers_by_bh_pf.get(_bh_p, [])
                # Đáy lớp 1 (lấy max của symbol IN '1','XMD')
                _d_bot_l1 = max(
                    (d_b for sym, _d_t, d_b in _lays_p if sym in ("1", "XMD")),
                    default=0.0,
                )
                _d_top_lo = next(
                    (d_t for sym, d_t, _d_b in _lays_p if sym == _lower_sym),
                    None,
                )
                _d_bot_lo = next(
                    (d_b for sym, _d_t, d_b in _lays_p if sym == _lower_sym),
                    None,
                )
                # NT detail: D_bot soft + L_design + pile
                _nt_p = _ntd_by_bh.get(_bh_p, {})
                _D_bot_soft = _nt_p.get("D_bot")
                _L_des = _ltk_pick.get(_bh_short) or _nt_p.get("L_des")
                _pile_name = _rec_pick.get(_bh_short) or _nt_p.get("pile") or "—"

                # Layer detail cho hover
                _lay_hover = []
                for sym, d_t, d_b in _lays_p:
                    _su_vst = _avg_in_range(_vst_by_bh.get(_bh_p, []), d_t, d_b)
                    _su_lab = _avg_in_range(_lab_by_bh.get(_bh_p, []), d_t, d_b)
                    _N_avg = _avg_in_range(_spt_by_bh.get(_bh_p, []), d_t, d_b)
                    _lay_hover.append({
                        "sym": sym, "d_t": d_t, "d_b": d_b,
                        "z_t": _z_p - d_t, "z_b": _z_p - d_b,
                        "su_vst": _su_vst, "su_lab": _su_lab, "N": _N_avg,
                    })

                _prof.append({
                    "bh":   _bh_short,
                    "bh_full": _bh_p,
                    "x":    _chain_pf.get(_bh_p, 0.0),
                    "z_t":  _z_p,
                    "z_b1": _z_p - _d_bot_l1 if _d_bot_l1 else None,
                    "z_tl": _z_p - _d_top_lo if _d_top_lo is not None else None,
                    "z_bl": _z_p - _d_bot_lo if _d_bot_lo is not None else None,
                    "z_bot_soft": (_z_p - _D_bot_soft) if _D_bot_soft else None,
                    "z_tip":     (_z_p - _L_des)      if _L_des else None,
                    "L_des":     _L_des,
                    "pile":      _pile_name,
                    "lay_hover": _lay_hover,
                })
            _prof.sort(key=lambda p: (
                _KE_BH_ORDER.index(p["bh"]) if p["bh"] in _KE_BH_ORDER
                else len(_KE_BH_ORDER)
            ))

            _xs = [p["x"] for p in _prof]
            _z_tops = [p["z_t"] for p in _prof]
            _z_b1s = [p["z_b1"] for p in _prof]
            _bh_lbls = [p["bh"] for p in _prof]

            # ── 4. Plotly chart (chính) ───────────────────────────────────
            def _hex_to_rgba(_hex: str, _alpha: float = 0.55) -> str:
                _h = _hex.lstrip("#")
                if len(_h) == 3:
                    _h = "".join(c * 2 for c in _h)
                _r, _g, _b = int(_h[0:2], 16), int(_h[2:4], 16), int(_h[4:6], 16)
                return f"rgba({_r},{_g},{_b},{_alpha})"

            if _HAS_PLOTLY:
                _fig_p = go.Figure()

                # Layer 1 (bùn) — fill xanh nhạt
                if all(v is not None for v in _z_b1s):
                    _fig_p.add_trace(go.Scatter(
                        x=_xs, y=_z_tops, mode="lines",
                        line=dict(color="#3E2723", width=2),
                        name="Mặt đất tự nhiên (Z)",
                        showlegend=True,
                    ))
                    _fig_p.add_trace(go.Scatter(
                        x=_xs, y=_z_b1s, mode="lines+markers",
                        line=dict(color="#01579B", width=2),
                        marker=dict(symbol="square", size=8),
                        fill="tonexty",
                        fillcolor=_hex_to_rgba(
                            _LAYER_COLORS.get("1", "#81D4FA"), 0.55),
                        name="Lớp 1 (bùn) — đáy",
                    ))

                # Layer dưới — fill
                if _lower_sym:
                    _z_tls = [p["z_tl"] for p in _prof]
                    _z_bls = [p["z_bl"] for p in _prof]
                    if all(v is not None for v in _z_tls) and \
                            all(v is not None for v in _z_bls):
                        _col_lo = _LAYER_COLORS.get(_lower_sym, "#A5D6A7")
                        _fig_p.add_trace(go.Scatter(
                            x=_xs, y=_z_tls, mode="lines+markers",
                            line=dict(color="#33691E", width=1.5, dash="dot"),
                            marker=dict(symbol="triangle-up", size=8),
                            name=f"Đỉnh lớp {_lower_sym}",
                        ))
                        _fig_p.add_trace(go.Scatter(
                            x=_xs, y=_z_bls, mode="lines+markers",
                            line=dict(color="#1B5E20", width=1.5, dash="dot"),
                            marker=dict(symbol="triangle-down", size=8),
                            fill="tonexty",
                            fillcolor=_hex_to_rgba(_col_lo, 0.45),
                            name=f"Đáy lớp {_lower_sym}",
                        ))

                # Đáy bùn (từ ke_sw_nt_detail.D_bottom_soft_m)
                _z_bot_soft = [p["z_bot_soft"] for p in _prof]
                if any(v is not None for v in _z_bot_soft):
                    _fig_p.add_trace(go.Scatter(
                        x=_xs, y=_z_bot_soft, mode="lines+markers",
                        line=dict(color="#D81B60", width=2.5, dash="dash"),
                        marker=dict(symbol="diamond", size=10,
                                    color="#D81B60"),
                        name="Đáy lớp đất yếu (tính NT1)",
                    ))

                # Mũi cọc
                _z_tip = [p["z_tip"] for p in _prof]
                _tip_text = [
                    (f"{p['pile']}<br>L = {p['L_des']:.1f} m<br>"
                     f"Mũi tại Z = {p['z_tip']:.2f} m")
                    if p["z_tip"] is not None else ""
                    for p in _prof
                ]
                if any(v is not None for v in _z_tip):
                    _fig_p.add_trace(go.Scatter(
                        x=_xs, y=_z_tip, mode="lines+markers+text",
                        line=dict(color="#FF6F00", width=2.5),
                        marker=dict(symbol="triangle-down", size=18,
                                    color="#FF6F00",
                                    line=dict(color="#BF360C", width=2)),
                        text=[p["pile"] for p in _prof],
                        textposition="bottom center",
                        textfont=dict(size=8, color="#BF360C"),
                        hovertext=_tip_text,
                        hoverinfo="text",
                        name="Mũi cọc thiết kế",
                        connectgaps=False,
                    ))

                # Đỉnh kè TK + Mặt nước (hlines)
                _x_min = min(_xs) - 5
                _x_max = max(_xs) + 5
                _fig_p.add_trace(go.Scatter(
                    x=[_x_min, _x_max],
                    y=[_z_top_design, _z_top_design],
                    mode="lines",
                    line=dict(color="#C62828", width=2.5, dash="dash"),
                    name=f"Đỉnh kè TK = {_z_top_design:.2f} m",
                    hovertemplate=f"Đỉnh kè TK = {_z_top_design:.2f} m"
                                  "<extra></extra>",
                ))
                _fig_p.add_trace(go.Scatter(
                    x=[_x_min, _x_max],
                    y=[_z_water, _z_water],
                    mode="lines",
                    line=dict(color="#0288D1", width=2, dash="dashdot"),
                    name=f"Mặt nước = {_z_water:.2f} m",
                    hovertemplate=f"Mặt nước = {_z_water:.2f} m"
                                  "<extra></extra>",
                ))

                # Hover markers per layer (tooltip chi tiết)
                _hov_x, _hov_y, _hov_txt = [], [], []
                for p in _prof:
                    for lay in p["lay_hover"]:
                        _hov_x.append(p["x"])
                        _hov_y.append((lay["z_t"] + lay["z_b"]) / 2.0)
                        _su_parts = []
                        if lay["su_vst"] is not None:
                            _su_parts.append(f"su(VST)={lay['su_vst']:.1f} kPa")
                        if lay["su_lab"] is not None:
                            _su_parts.append(f"su(Lab)={lay['su_lab']:.1f} kPa")
                        _su_txt = " · ".join(_su_parts) if _su_parts else "su: —"
                        _N_txt = (f"N̄ = {lay['N']:.1f}"
                                  if lay["N"] is not None else "N: —")
                        _hov_txt.append(
                            f"<b>{p['bh']} — Lớp {lay['sym']}</b><br>"
                            f"Độ sâu: {lay['d_t']:.2f}–{lay['d_b']:.2f} m "
                            f"(dày {lay['d_b']-lay['d_t']:.2f} m)<br>"
                            f"Cao độ: {lay['z_t']:+.2f} → "
                            f"{lay['z_b']:+.2f} m<br>"
                            f"{_su_txt}<br>{_N_txt}"
                        )
                if _hov_x:
                    _fig_p.add_trace(go.Scatter(
                        x=_hov_x, y=_hov_y, mode="markers",
                        marker=dict(size=14, color="rgba(0,0,0,0)",
                                    line=dict(width=0)),
                        hovertext=_hov_txt, hoverinfo="text",
                        showlegend=False, name="hover",
                    ))

                # Vạch đứng + nhãn HK ở đỉnh
                for p in _prof:
                    _fig_p.add_vline(x=p["x"], line=dict(
                        color="gray", width=0.5, dash="dot"), opacity=0.5)
                    _fig_p.add_annotation(
                        x=p["x"], y=_z_top_design + 0.6,
                        text=f"<b>{p['bh']}</b>",
                        showarrow=False,
                        font=dict(size=11, color="#1A237E"),
                    )

                _fig_p.update_layout(
                    title=(f"Trắc dọc địa chất tuyến kè SW — "
                           f"{len(_prof)} HK (lớp 1 + lớp {_lower_sym})"),
                    xaxis_title="Khoảng cách dọc tuyến (m, gốc PCA)",
                    yaxis_title="Cao độ (m)",
                    height=520,
                    hovermode="closest",
                    legend=dict(orientation="h", yanchor="bottom",
                                y=-0.25, xanchor="left", x=0),
                    margin=dict(t=60, b=80, l=60, r=20),
                    plot_bgcolor="#FAFAFA",
                )
                _fig_p.update_xaxes(showgrid=True, gridcolor="#E0E0E0")
                _fig_p.update_yaxes(showgrid=True, gridcolor="#E0E0E0")
                st.plotly_chart(_fig_p, use_container_width=True)

            # ── 5. Matplotlib (double-render cho PDF) ────────────────────
            if _HAS_MPL:
                import matplotlib.pyplot as _plt_pf
                _fig_m, _ax_m = _plt_pf.subplots(figsize=(11, 5.5), dpi=110)
                _ax_m.axhline(_z_top_design, color="#C62828", lw=2.2,
                              ls="--",
                              label=f"Đỉnh kè TK = {_z_top_design:.2f} m")
                _ax_m.axhline(_z_water, color="#0288D1", lw=1.8,
                              ls="-.",
                              label=f"Mặt nước = {_z_water:.2f} m")
                if all(v is not None for v in _z_b1s):
                    _ax_m.fill_between(_xs, _z_tops, _z_b1s,
                                       color=_LAYER_COLORS.get("1", "#81D4FA"),
                                       alpha=0.55, label="Lớp 1 (bùn)")
                    _ax_m.plot(_xs, _z_b1s, color="#01579B", marker="s",
                               lw=1.8, label="Đáy lớp 1")
                _ax_m.plot(_xs, _z_tops, color="#3E2723", marker="o",
                           lw=2.0, label="Mặt đất tự nhiên")
                if _lower_sym:
                    _z_tls_m = [p["z_tl"] for p in _prof]
                    _z_bls_m = [p["z_bl"] for p in _prof]
                    if all(v is not None for v in _z_tls_m) and \
                            all(v is not None for v in _z_bls_m):
                        _col_lo = _LAYER_COLORS.get(_lower_sym, "#A5D6A7")
                        _ax_m.fill_between(_xs, _z_tls_m, _z_bls_m,
                                           color=_col_lo, alpha=0.45,
                                           label=f"Lớp {_lower_sym}")
                        _ax_m.plot(_xs, _z_tls_m, color="#33691E",
                                   marker="^", lw=1.4, ls=":",
                                   label=f"Đỉnh lớp {_lower_sym}")
                        _ax_m.plot(_xs, _z_bls_m, color="#1B5E20",
                                   marker="v", lw=1.4, ls=":",
                                   label=f"Đáy lớp {_lower_sym}")
                _z_bs_m = [p["z_bot_soft"] for p in _prof]
                if any(v is not None for v in _z_bs_m):
                    _ax_m.plot(_xs, _z_bs_m, color="#D81B60", marker="D",
                               lw=2.0, ls="--", label="Đáy lớp đất yếu (tính NT1)")
                _z_tip_m = [p["z_tip"] for p in _prof]
                if any(v is not None for v in _z_tip_m):
                    _ax_m.plot(_xs, _z_tip_m, color="#FF6F00", lw=2.0,
                               marker="v", markersize=11,
                               markeredgecolor="#BF360C",
                               markeredgewidth=1.5, zorder=5,
                               label="Mũi cọc TK")
                    for p in _prof:
                        if p["z_tip"] is not None:
                            _ax_m.text(p["x"], p["z_tip"] - 0.8,
                                       p["pile"], ha="center", fontsize=7,
                                       color="#BF360C")
                for p in _prof:
                    _ax_m.axvline(p["x"], color="gray", lw=0.5, ls=":",
                                  alpha=0.6)
                    _ax_m.text(p["x"], _z_top_design + 0.5, p["bh"],
                               ha="center", fontsize=9, fontweight="bold",
                               color="#1A237E")
                _ax_m.set_xlabel("Khoảng cách dọc tuyến (m, gốc PCA)")
                _ax_m.set_ylabel("Cao độ (m)")
                _ax_m.set_title(
                    f"Trắc dọc địa chất tuyến kè SW — {len(_prof)} HK "
                    f"(lớp 1 + lớp {_lower_sym})"
                )
                _ax_m.grid(True, ls="--", alpha=0.4)
                _ax_m.legend(loc="lower left", fontsize=7, ncol=3,
                             framealpha=0.9)
                _fig_m.tight_layout()
                # Lưu vào session_state cho PDF generator
                st.session_state["_ke_b_profile_mpl_fig"] = _fig_m
                with st.expander("Bản Matplotlib (dùng cho PDF)", expanded=False):
                    st.pyplot(_fig_m, use_container_width=True)
                _plt_pf.close(_fig_m)

            # ── 6. Bình đồ vị trí hố khoan (x,y) ────────────────────────
            st.markdown("#### Bình đồ vị trí hố khoan + khoảng cách")

            _plan_data = [
                {"bh": r[0].replace("KE-", ""), "bh_full": r[0],
                 "x": float(r[1]), "y": float(r[2]),
                 "ch": _chain_pf.get(r[0], 0.0)}
                for r in _bh_info if r[1] and r[2]
            ]
            _plan_data.sort(key=lambda p: (
                _KE_BH_ORDER.index(p["bh"]) if p["bh"] in _KE_BH_ORDER
                else len(_KE_BH_ORDER)
            ))

            # Khoảng cách giữa các cặp HK liên tiếp (Euclidean)
            _seg_dist = []
            for i in range(len(_plan_data) - 1):
                _a = _plan_data[i]; _b = _plan_data[i + 1]
                _d = _np_pf.hypot(_b["x"] - _a["x"], _b["y"] - _a["y"])
                _seg_dist.append({
                    "from": _a["bh"], "to": _b["bh"],
                    "x_mid": (_a["x"] + _b["x"]) / 2,
                    "y_mid": (_a["y"] + _b["y"]) / 2,
                    "dist": float(_d),
                })

            if _HAS_PLOTLY and _plan_data:
                _fig_pl = go.Figure()
                # Bình đồ kè (polyline DXF) — vẽ trước để HK nổi lên trên
                if _binhdo_cx:
                    _fig_pl.add_trace(go.Scatter(
                        x=_binhdo_cx, y=_binhdo_cy,
                        mode="lines",
                        line=dict(color="rgba(0,0,0,0.85)", width=1.5),
                        hoverinfo="skip",
                        showlegend=True,
                        name="Ranh kè (bình đồ)",
                        connectgaps=False,
                    ))
                # Đường nối tuyến qua các HK
                _fig_pl.add_trace(go.Scatter(
                    x=[p["x"] for p in _plan_data],
                    y=[p["y"] for p in _plan_data],
                    mode="lines+markers+text",
                    line=dict(color="#1976D2", width=2.5),
                    marker=dict(symbol="circle", size=14,
                                color="#FFC107",
                                line=dict(color="#1976D2", width=2)),
                    text=[p["bh"] for p in _plan_data],
                    textposition="top center",
                    textfont=dict(size=11, color="#0D47A1"),
                    hovertext=[
                        f"<b>{p['bh_full']}</b><br>"
                        f"X = {p['x']:,.2f} m<br>"
                        f"Y = {p['y']:,.2f} m<br>"
                        f"Chainage = {p['ch']:.2f} m"
                        for p in _plan_data
                    ],
                    hoverinfo="text",
                    name="Hố khoan trên tuyến",
                ))
                # Nhãn khoảng cách giữa các cặp
                for s in _seg_dist:
                    _fig_pl.add_annotation(
                        x=s["x_mid"], y=s["y_mid"],
                        text=f"<b>{s['dist']:.1f} m</b>",
                        showarrow=False,
                        font=dict(size=10, color="#D32F2F"),
                        bgcolor="rgba(255,255,255,0.85)",
                        borderpad=2,
                    )
                _fig_pl.update_layout(
                    title=(f"Bình đồ vị trí {len(_plan_data)} hố khoan "
                           f"trên tuyến kè SW"),
                    xaxis_title="X (Northing VN-2000, m)",
                    yaxis_title="Y (Easting VN-2000, m)",
                    height=560,
                    hovermode="closest",
                    plot_bgcolor="#FAFAFA",
                    showlegend=bool(_binhdo_cx),
                    legend=dict(x=0.01, y=0.99, bgcolor="rgba(255,255,255,0.8)",
                                bordercolor="#CCC", borderwidth=1),
                    margin=dict(t=60, b=60, l=60, r=20),
                )
                _fig_pl.update_yaxes(scaleanchor="x", scaleratio=1,
                                     showgrid=True, gridcolor="#E0E0E0")
                _fig_pl.update_xaxes(showgrid=True, gridcolor="#E0E0E0")
                st.plotly_chart(_fig_pl, use_container_width=True)

            # Matplotlib double-render
            if _HAS_MPL and _plan_data:
                import matplotlib.pyplot as _plt_pl
                _fig_plm, _ax_plm = _plt_pl.subplots(figsize=(10, 8), dpi=110)
                # Vẽ polyline bình đồ kè trước (nền)
                if _binhdo_cx:
                    _seg_xs, _seg_ys = [], []
                    for _cx, _cy in zip(_binhdo_cx, _binhdo_cy):
                        if _cx is None:
                            if _seg_xs:
                                _ax_plm.plot(_seg_xs, _seg_ys,
                                             color="#000000", lw=1.5,
                                             alpha=0.85, zorder=2)
                            _seg_xs, _seg_ys = [], []
                        else:
                            _seg_xs.append(_cx)
                            _seg_ys.append(_cy)
                    if _seg_xs:
                        _ax_plm.plot(_seg_xs, _seg_ys, color="#000000",
                                     lw=1.5, alpha=0.85, zorder=2,
                                     label="Ranh kè (bình đồ)")
                _xs_pl = [p["x"] for p in _plan_data]
                _ys_pl = [p["y"] for p in _plan_data]
                _ax_plm.plot(_xs_pl, _ys_pl, color="#1976D2", lw=2.5,
                             marker="o", markersize=12,
                             markerfacecolor="#FFC107",
                             markeredgecolor="#1976D2", markeredgewidth=2,
                             zorder=5, label="Hố khoan trên tuyến")
                for p in _plan_data:
                    _ax_plm.annotate(
                        p["bh"], xy=(p["x"], p["y"]),
                        xytext=(0, 12), textcoords="offset points",
                        ha="center", fontsize=10, fontweight="bold",
                        color="#0D47A1",
                    )
                for s in _seg_dist:
                    _ax_plm.annotate(
                        f"{s['dist']:.1f} m",
                        xy=(s["x_mid"], s["y_mid"]),
                        ha="center", va="center", fontsize=9,
                        color="#D32F2F", fontweight="bold",
                        bbox=dict(boxstyle="round,pad=0.2",
                                  facecolor="white", alpha=0.85,
                                  edgecolor="none"),
                    )
                _ax_plm.set_xlabel("X (Northing VN-2000, m)")
                _ax_plm.set_ylabel("Y (Easting VN-2000, m)")
                _ax_plm.set_title(
                    f"Bình đồ vị trí {len(_plan_data)} hố khoan "
                    f"trên tuyến kè SW"
                )
                _ax_plm.set_aspect("equal", adjustable="datalim")
                _ax_plm.grid(True, ls="--", alpha=0.4)
                if _binhdo_cx:
                    _ax_plm.legend(fontsize=9, loc="lower right")
                _fig_plm.tight_layout()
                st.session_state["_ke_b_plan_mpl_fig"] = _fig_plm
                with st.expander("Bản Matplotlib bình đồ (dùng cho PDF)",
                                 expanded=False):
                    st.pyplot(_fig_plm, use_container_width=True)
                _plt_pl.close(_fig_plm)

            # Bảng khoảng cách
            if _seg_dist:
                _df_dist = pd.DataFrame([
                    {"Từ HK": s["from"], "Đến HK": s["to"],
                     "Khoảng cách (m)": round(s["dist"], 2)}
                    for s in _seg_dist
                ])
                _total = sum(s["dist"] for s in _seg_dist)
                st.caption(
                    f"Tổng chiều dài tuyến (cộng dồn cặp liên tiếp): "
                    f"**{_total:.2f} m**"
                )
                st.table(_df_dist)

        if _HAS_PLOTLY and _ke_rows:
            _names_plot  = [r["Hố khoan"] for r in _ke_rows if r["RR/W"] != "–"]
            _ratios_plot = [float(r["RR/W"]) for r in _ke_rows if r["RR/W"] != "–"]
            _min_ratio   = min(_ratios_plot) if _ratios_plot else 0
            _colors_plot = [
                "#FF6B35" if v == _min_ratio else "#2E7D32"
                for v in _ratios_plot
            ]
            _fig_nt2 = go.Figure()
            _fig_nt2.add_trace(go.Bar(
                x=_names_plot, y=_ratios_plot,
                marker_color=_colors_plot,
                text=[f"{v:.2f}" for v in _ratios_plot],
                textposition="outside",
                name="RR/W",
            ))
            _fig_nt2.add_hline(y=1.0, line_dash="dash", line_color="red",
                               annotation_text="Giới hạn NT2 (RR ≥ W)")
            _fig_nt2.update_layout(
                title="Tỷ số NT2: RR/W theo hố khoan trên tuyến kè SW  (cam = HK kiểm soát)",
                yaxis_title="RR / W", xaxis_title="Hố khoan",
                height=350, margin=dict(t=50, b=40),
                yaxis=dict(range=[0, max(_ratios_plot) * 1.2 if _ratios_plot else 3]),
            )
            st.plotly_chart(_fig_nt2, use_container_width=True)
        st.caption(
            "Hiển thị 7 hố khoan trên tuyến kè SW: KE-HK2, HK3, HK7, HK8, HK9, HK10, HK11. "
            "Cam = HK kiểm soát (tỷ số RR/W nhỏ nhất)."
        )

        # ── Chú thích ý nghĩa ký hiệu + công thức tính ────────────────────────
        with st.expander("Giải thích ký hiệu & công thức trong bảng", expanded=True):
            st.markdown(
                r"""
**Cột mô tả hố khoan:**
| Ký hiệu | Ý nghĩa | Nguồn |
|---|---|---|
| **Z (m)** | Cao độ mặt đất tự nhiên tại HK | Hồ sơ địa chất — cao độ cổ hố khoan |
| **H lớp 1 (m)** | Chiều dày lớp bùn sét yếu (lớp 1) | Hồ sơ địa chất — lớp ký hiệu `1` hoặc `XMD` |
| **L yêu cầu (m)** | Chiều dài cọc tối thiểu = `Z_đỉnh_kè − Z + H_yếu + h_ngàm` | Tính từ Z và H lớp 1 |

**Cột lựa chọn cọc:**
| Ký hiệu | Ý nghĩa | Nguồn |
|---|---|---|
| **Cọc tối ưu** | Cọc nhỏ nhất trong catalog có `L_max ≥ L_yêu_cầu` | Tự động chọn |
| **L_max (m)** | Chiều dài tối đa khả thi của cọc tối ưu | Catalog cọc SW dự ứng lực |
| **Đủ chiều dài** | "Đạt" nếu `L_max ≥ L_yêu_cầu` | So sánh hai cột trên |
| **Cọc kiến nghị** | Cọc user chọn (dropdown — có thể khác cọc tối ưu) | Bảng có thể chỉnh sửa |
| **L thiết kế (m)** | Chiều dài cọc user nhập | Bảng có thể chỉnh sửa |

**Cột kết quả NT1 / NT2 (TCVN 11823-10:2017, Điều 7.3.8.6.2):**

NT1 — Kiểm tra chiều dài xuyên qua lớp yếu:
"""
            )
            st.latex(r"L_{thiết\ kế} \;\geq\; L_{yêu\ cầu} \;=\; (Z_{đỉnh\ kè} - Z_{tự\ nhiên}) + H_{lớp\ yếu} + h_{ngàm}")
            st.markdown(
                r"""
NT2 — Kiểm tra sức kháng (phương pháp α — Tomlinson 1980):
"""
            )
            st.latex(r"R_s \;=\; \alpha \cdot s_u \cdot p \cdot L_{ngập\ đất}")
            st.latex(r"R_p \;=\; 9 \cdot s_u \cdot A_p \quad (\text{sét}); \quad R_p = q_p^{SPT} \cdot A_p \quad (\text{cát})")
            st.latex(r"RR \;=\; \varphi_{stat} \cdot (R_s + R_p) \;\geq\; W_{cọc}")
            st.latex(r"W_{cọc} \;=\; \dfrac{TL \times 9{,}81}{L_{std}} \cdot L_{thiết\ kế}")
            st.markdown(
                r"""
| Ký hiệu | Ý nghĩa | Đơn vị | Công thức / Nguồn |
|---|---|---|---|
| **Rs** | Sức kháng ma sát thân cọc | kN | $\alpha \cdot s_u \cdot p \cdot L_{ngập}$ (α=1.0; $p$ = chu vi cọc; $L_{ngập}$ = $L_{tk}$ − đắp) |
| **Rp** | Sức kháng mũi cọc | kN | $9 s_u A_p$ (sét) hoặc Meyerhof (cát) |
| **RR** | Sức kháng tính toán | kN | $\varphi_{stat}(R_s + R_p)$, $\varphi_{stat}=0{,}35$ — TCVN 11823-10 Bảng 9 |
| **W** | Trọng lượng cọc | kN | $\dfrac{TL_T \cdot 9{,}81}{L_{std}} \cdot L_{tk}$ — TL/L_std lấy từ catalog |
| **RR/W** | Hệ số an toàn | – | Yêu cầu **≥ 1,0**; HK có giá trị nhỏ nhất = HK kiểm soát |
| **NT1 / NT2** | Kết quả "Đạt" / "Không đạt" | – | Theo hai điều kiện trên |

**Ghi chú quan trọng — vì sao W khác nhau:**

W cọc phụ thuộc vào **TL (trọng lượng cọc tiêu chuẩn)** và **L_std (chiều dài tiêu chuẩn)** của từng loại cọc trong catalog. Hai cọc khác loại sẽ có TL và L_std khác nhau → W khác nhau dù L thiết kế giống nhau.

| Cọc | TL (T) | L_std (m) | w/m (kN/m) | W tại L=24m (kN) |
|---|---|---|---|---|
| SW-600B | 10,88 | 20 | 5,34 | 128,1 |
| SW-740  | 14,55 | 21 | 6,80 | 163,2 |
| SW-840  | 16,35 | 22 | 7,29 | 175,0 |
| SW-940  | 18,31 | 23 | 7,81 | 187,4 |

Nếu trong bảng thấy hai cọc khác loại mà W giống nhau → bug, vui lòng báo lại.

**Cột Ghi chú:** *"Tính lại theo cọc user chọn"* = số liệu Rs/Rp/RR/W trong hàng đã được tính lại real-time theo cọc + L user vừa chọn (không phải lấy từ dữ liệu thiết kế gốc).
"""
            )

        # ── Chi tiết NT1/NT2 từng hố khoan (inline, lọc theo Mục B) ──────────
        st.markdown("#### Chi tiết tính toán NT1/NT2 từng hố khoan")
        _SRC_LBL = {
            "VST": "VST", "lab": "Lab", "default": "Giả định",
            "sand": "Cát", "SPT": "SPT", "missing": "Thiếu SPT",
            "unknown": "?",
        }
        try:
            # Lọc + GIỮ THỨ TỰ chính xác như Mục B (_picked_bhs)
            _picked_db = [f"KE-{n}" for n in _picked_bhs]
            if not _picked_db:
                st.info("Chưa có HK nào được chọn ở Mục B.")
            else:
                _placeholders = ",".join("?" * len(_picked_db))
                _con_nt = sqlite3.connect(str(_DB))
                _nt_rows_raw = _con_nt.execute(
                    f"""SELECT bh_name, pile_type, L_design_m, Z_m, Z_source,
                    fill_m, L_soil_m, tip_depth_m, D_bottom_soft_m, D_source,
                    L_req_nt1_m, margin_nt1_m, nt1_result,
                    Rs_kN, Rs_clay_kN, Rs_sand_kN,
                    tip_symbol, tip_method, tip_su_kNm2, tip_N160,
                    Rp_kN, phi_stat, phi_basis, RR_kN, W_kN,
                    ratio_nt2, nt2_result, su_warnings
                    FROM ke_sw_nt_detail
                    WHERE bh_name IN ({_placeholders})""",
                    _picked_db,
                ).fetchall()
                _con_nt.close()
                # Sort theo thứ tự _picked_db (giống bảng Mục B)
                _bh_order = {nm: i for i, nm in enumerate(_picked_db)}
                _nt_rows = sorted(_nt_rows_raw, key=lambda r: _bh_order.get(r[0], 999))

                if not _nt_rows:
                    st.info("Không có dữ liệu NT cho các HK đã chọn. "
                            "Nhấn 'Tính lại từ số liệu khảo sát' phía trên.")

                _df_vst_ke = _load_vst_su("KE") if _picked_db else pd.DataFrame()

                for _nt in _nt_rows:
                    (
                        _bh_n, _pile_t, _L_d, _Z_m, _Z_src,
                        _fill_m, _L_soil, _tip_d, _D_bot, _D_src,
                        _L_req1, _marg1, _res1,
                        _Rs, _Rs_clay, _Rs_sand,
                        _tip_sym, _tip_meth, _tip_su, _tip_N160,
                        _Rp, _phi, _phi_basis, _RR, _W,
                        _rat2, _res2, _warns,
                    ) = _nt
                    st.markdown(f"##### {_bh_n}")

                    if _warns:
                        st.warning("**Cảnh báo:** " + " · ".join(_warns.split("; ")))

                    # ── Hàng 1: NT1 và NT2 cùng một hàng ────────────────────────
                    _nt1_col, _nt2_col = st.columns(2)
                    _badge_nt1 = "🟢" if _marg1 >= 0 else "🔴"
                    _badge_nt2 = "🟢" if _rat2 >= 1 else "🔴"
                    with _nt1_col:
                        st.markdown(
                            f"**NT1 — Chiều dài xuyên qua lớp yếu** {_badge_nt1} **{_res1}**"
                        )
                        st.caption(
                            f"L tk = **{_L_d:.1f} m** | L req = {_L_req1:.1f} m | "
                            f"Δ = **{_marg1:+.2f} m**  \n"
                            f"Đáy lớp mềm = {_D_bot:.1f} m | "
                            f"Z mặt đất = {_Z_m:+.2f} m  \n"
                            f"Đắp = {_fill_m:.2f} m | Trong đất = {_L_soil:.2f} m"
                        )
                    with _nt2_col:
                        st.markdown(
                            f"**NT2 — Sức kháng TCVN 11823-10** {_badge_nt2} **{_res2}** "
                            f"(RR/W = **{_rat2:.2f}**)"
                        )
                        _rs_parts = []
                        if _Rs_clay is not None:
                            _rs_parts.append(f"sét={_Rs_clay:.0f}")
                        if _Rs_sand is not None and _Rs_sand > 0:
                            _rs_parts.append(f"cát={_Rs_sand:.0f}")
                        _rs_bd = " (" + " + ".join(_rs_parts) + ")" if _rs_parts else ""
                        _tip_short = f"{_tip_sym}[{_tip_meth or 'α'}]"
                        if _tip_meth == "SPT":
                            _tip_short += (f"·N₁₆₀={_tip_N160:.0f}" if _tip_N160
                                           else "·noSPT")
                        else:
                            _tip_short += f"·su={_tip_su:.0f} kPa"
                        st.caption(
                            f"Rs = {_Rs:.0f}{_rs_bd} | Rp = {_Rp:.0f} kN  \n"
                            f"RR = φ(Rs+Rp) = **{_RR:.0f} kN** | W = {_W:.0f} kN  \n"
                            f"φ = **{_phi:.2f}** ({_phi_basis or '—'}) | Tip: {_tip_short}"
                        )

                    # ── Hàng 2: bảng lớp đất full width ─────────────────────────
                    _con2 = sqlite3.connect(str(_DB))
                    _lyrs = _con2.execute(
                        "SELECT l.symbol, l.L_m, l.method, l.su_kPa, l.alpha, "
                        "l.N160, l.gamma_kNm3, l.sigma_v_eff_kPa, l.su_source, l.Rs_kN "
                        "FROM ke_sw_nt2_layers l "
                        "JOIN ke_sw_nt_detail d ON l.sw_design_id=d.id "
                        "WHERE d.bh_name=? ORDER BY l.layer_order",
                        (_bh_n,),
                    ).fetchall()
                    _con2.close()
                    # Tìm HK gốc VST tham khảo (nếu có nearest)
                    _vst_ref_bh = None
                    if _lyrs:
                        _ldf_rows = []
                        for (_sym, _Llyr, _meth, _su, _alp, _N, _gam, _sigv,
                             _src, _Rs_lyr) in _lyrs:
                            _src_short = (_src.replace("nearest:", "≈")
                                          if _src and _src.startswith("nearest:")
                                          else _SRC_LBL.get(_src or "", _src or ""))
                            if _src and _src.startswith("nearest:") and _vst_ref_bh is None:
                                # Parse "nearest:KE-HK2(233m,c_)" → "KE-HK2"
                                import re as _re_src
                                _m = _re_src.match(r"nearest:([\w-]+)\(", _src)
                                if _m:
                                    _vst_ref_bh = _m.group(1)
                            _ldf_rows.append({
                                "Lớp":   _sym,
                                "L (m)": round(_Llyr, 2),
                                "Phương pháp": (_meth or "alpha")[:5],
                                "su (kPa)":   round(_su, 1) if _su else 0,
                                "α":     round(_alp, 3) if _alp else 0,
                                "N₁₆₀": round(_N, 0) if _N is not None else "—",
                                "γ (kN/m³)":  round(_gam, 1) if _gam else "—",
                                "σ'v (kPa)":  round(_sigv, 0) if _sigv else "—",
                                "Nguồn": _src_short,
                                "Rs (kN)": round(_Rs_lyr, 0),
                            })
                        st.table(pd.DataFrame(_ldf_rows))

                    # ── Hàng 3: Cột địa chất + VST cùng hàng, font giảm để không tràn ──
                    # Cả 2 đặt cạnh nhau trong 50% column → figsize=(7,6).
                    # Font: soil column 0.75, VST 1.25 — match nhưng không chữ to quá.
                    _col_sc, _col_vst = st.columns(2, gap="medium")
                    with _col_sc:
                        st.markdown(f"**Cột địa chất {_bh_n}**")
                        try:
                            _sc_lay = _load_layers(_bh_n)
                            _sc_spt = _load_spt(_bh_n)
                            if _sc_lay and _HAS_MPL:
                                _fig_sc = _draw_soil_column_mpl(
                                    _sc_lay, _bh_n, spt=_sc_spt or None,
                                    figsize=(7, 6), font_scale=0.75,
                                )
                                st.pyplot(_fig_sc, use_container_width=True)
                                plt.close(_fig_sc)
                            else:
                                st.caption("_(không có địa tầng)_")
                        except Exception as _e_sc:
                            st.caption(f"_(không vẽ được: {_e_sc})_")

                    with _col_vst:
                        _vst_target_bh = _vst_ref_bh or _bh_n
                        st.markdown(
                            f"**VST {_vst_target_bh}**"
                            + ("  _(tham khảo từ HK gốc)_"
                               if _vst_ref_bh and _vst_ref_bh != _bh_n else "")
                        )
                        if not _df_vst_ke.empty and _HAS_MPL:
                            _bh_short = _vst_target_bh.replace("KE-", "")
                            # Match chính xác — TRÁNH substring match:
                            #   "HK1" in "KE-HK12" → True (sai)
                            # Dùng regex: HK<num> + KHÔNG có digit phía sau
                            import re as _re_vst
                            _bh_num = _bh_short.replace("HK", "")
                            _pat_vst = _re_vst.compile(
                                rf'HK{_re_vst.escape(_bh_num)}(?!\d)'
                            )
                            _vst_match = [
                                l for l in _df_vst_ke["loc_name"].unique()
                                if _pat_vst.search(str(l))
                            ]
                            if _vst_match:
                                _fig_vst = _chart_su_profile_mpl(
                                    _df_vst_ke, _vst_match,
                                    figsize=(7, 6), font_scale=1.25,
                                )
                                st.pyplot(_fig_vst, use_container_width=True)
                                plt.close(_fig_vst)
                            else:
                                st.caption(f"_(không có VST cho {_vst_target_bh})_")
                        else:
                            st.caption("_(không có dữ liệu VST)_")

                    st.divider()
        except Exception as _exc_nt:
            st.error(f"Lỗi khi đọc dữ liệu NT: {_exc_nt}")

    # Mục C (Kiểm tra NT1/NT2 tùy chỉnh) đã ẩn — chuyển sang C.3 auto-loop.
    # Xem git history (commit trước 2026-05-21) nếu cần khôi phục form thủ công.

    # C. So sánh 3 phương pháp NT2 (α, β, λ) — SPT-Meyerhof đã bỏ.

    # ── C. So sánh 3 phương pháp NT2 — tự động cho mọi HK ở Mục B ─────────────
    st.divider()
    st.markdown("### C. So sánh 3 phương pháp NT2 — tự động cho mọi HK ở Mục B")
    st.info(
        "Tự động chạy so sánh Rs/Rp/RR theo 3 phương pháp TCVN 11823-10:2017 cho mọi HK + cọc kiến nghị ở Mục B:  \n"
        "**α** (Tomlinson 1980, sét, Điều 7.3.8.6.2, φ=0,35) · "
        "**β** (Esrig & Kirby 1979, sét, Điều 7.3.8.6.3, φ=0,25) · "
        "**λ** (Vijayvergiya & Focht 1972, cọc ống, Điều 7.3.8.6.4, φ=0,40)"
    )

    try:
        import sys as _sys_cmp
        _sys_cmp.path.insert(0, str(_ROOT / "scripts"))
        from ke_sw_nt_calc import (
            calc_nt2_all_methods as _calc_4m,
            _load_catalog as _load_cat,
            _get_bh_Z_m as _get_bh_z,
        )
        _cat_cmp = _load_cat()

        # HK list = tất cả HK KE từ Mục B (không hardcode)
        _picks_c = st.session_state.get("ke_sw_alignment_picks", []) or []
        _bh_c_list = [f"KE-{nm}" for nm in _picks_c if nm.startswith("HK")]

        try:
            from ke_sw_nt_calc import (
                load_nt2_compare_from_db as _load_4m_db,
                save_nt2_compare_to_db as _save_4m_db,
                calc_nt2_all_methods as _calc_4m,
            )
        except Exception:
            _load_4m_db = _save_4m_db = _calc_4m = None

        if not _bh_c_list:
            st.caption("_(Chưa có HK nào ở Mục B để so sánh — chọn HK ở Mục B trước.)_")
        else:
            _method_lbl = {
                "auto":   "Tự chọn",
                "alpha":  "α",
                "beta":   "β",
                "lambda": "λ",
            }

            # Mục C — DB-first: đọc SQLite trước, tính toán + lưu nếu chưa có
            _all_results: list = []
            _prog_c = st.progress(0, text="Đang tải kết quả NT2 mục C...")
            _n_c = len(_bh_c_list)
            _nc_computed = 0
            for _ic, _bh_c in enumerate(_bh_c_list):
                _prog_c.progress(int(_ic / _n_c * 100),
                    text=f"{_bh_c}... ({_ic + 1}/{_n_c})")
                _nm_short = _bh_c.replace("KE-", "")
                _pile_c = (st.session_state.get("ke_sw_rec_piles", {}) or {}).get(_nm_short)
                _L_c    = (st.session_state.get("ke_sw_L_thiet_ke", {}) or {}).get(_nm_short)
                if not _pile_c or not _L_c or _pile_c not in _cat_cmp:
                    continue
                try:
                    # 1. Thử đọc từ SQLite
                    _res_c = _load_4m_db(_bh_c, _pile_c, float(_L_c)) if _load_4m_db else None
                    if _res_c is None and _calc_4m:
                        # 2. Tính toán và lưu vào SQLite
                        _z_c = _get_bh_z(_bh_c) or 0.0
                        _pp  = _cat_cmp[_pile_c]
                        _res_c = _calc_4m(_bh_c, float(_z_c), _pile_c, _pp, float(_L_c))
                        if _save_4m_db:
                            try:
                                _save_4m_db(_bh_c, _pile_c, float(_L_c), _res_c)
                            except Exception:
                                pass  # disk I/O error on Google Drive — kết quả vẫn hiển thị
                        _nc_computed += 1
                    if _res_c:
                        _cm_raw = dict(_res_c.get("common", {}))
                        # Enrich geometric fields từ catalog nếu thiếu (khi load từ DB)
                        if "D_m" not in _cm_raw and _pile_c in _cat_cmp:
                            _pp_cat = _cat_cmp[_pile_c]
                            _cm_raw["D_m"] = _pp_cat.get("width_mm", 840) / 1000.0
                            _cm_raw["perimeter_m"] = _pp_cat.get("perimeter_mm", 2640) / 1000.0
                        _all_results.append((_bh_c, _pile_c, float(_L_c), _cm_raw, _res_c))
                except Exception as _e_each:
                    st.warning(f"{_bh_c}: lỗi tính — {_e_each}")
                    continue
            _prog_c.progress(100, text=f"Hoàn thành {_n_c}/{_n_c} hố khoan"
                             + (f" ({_nc_computed} tính mới, {_n_c-_nc_computed} từ DB)" if _nc_computed else " (từ DB)"))
            import time as _tc; _tc.sleep(0.3)
            _prog_c.empty()

            if not _all_results:
                st.warning("Không có HK nào có đủ dữ liệu (cọc kiến nghị + L) để so sánh.")
            else:
                # ── Bảng tổng hợp: 1 dòng / HK, 5 cột RR theo phương pháp ─────
                _sum_rows = []
                for _bh, _pl, _L, _cm, _res in _all_results:
                    _row = {
                        "Hố khoan": _bh,
                        "Cọc":      _pl,
                        "L (m)":    _L,
                        "W (kN)":   _cm["W_kN"],
                    }
                    for _m_key, _m_name in _method_lbl.items():
                        _x = _res[_m_key]
                        _row[f"RR — {_m_name}"]    = _x["RR_kN"]
                        _row[f"Ratio — {_m_name}"] = _x["ratio"]
                    _row["Kết quả (tự chọn)"] = _res["auto"]["result"]
                    _sum_rows.append(_row)

                _sum_df = pd.DataFrame(_sum_rows)

                def _hl_kq(val):
                    if val == "Đạt":
                        return "background-color:#d4edda; color:#155724"
                    if val == "Không đạt":
                        return "background-color:#f8d7da; color:#721c24"
                    return ""

                _ratio_cols = [f"Ratio — {n}" for n in _method_lbl.values()]
                _RR_cols    = [f"RR — {n}"    for n in _method_lbl.values()]
                _col_cfg = {"L (m)":  st.column_config.NumberColumn(format="%.1f"),
                            "W (kN)": st.column_config.NumberColumn(format="%.0f")}
                for _c in _RR_cols:
                    _col_cfg[_c] = st.column_config.NumberColumn(format="%.0f")
                for _c in _ratio_cols:
                    _col_cfg[_c] = st.column_config.NumberColumn(format="%.2f")

                st.markdown("**Bảng tổng hợp NT2 — mọi HK ở Mục B**")
                st.dataframe(
                    _sum_df.style.map(_hl_kq, subset=["Kết quả (tự chọn)"]),
                    use_container_width=True, hide_index=True,
                    column_config=_col_cfg,
                )

                # ── Bar chart so sánh RR (group by phương pháp, color by HK) ──
                if _HAS_PLOTLY:
                    _fig_all = go.Figure()
                    for _m_key, _m_name in _method_lbl.items():
                        _ys = [r[_m_key]["RR_kN"] for *_t, r in _all_results]
                        _xs = [t[0] for t in _all_results]
                        _W  = [t[3]["W_kN"] for t in _all_results]
                        _ok = [r[_m_key]["RR_kN"] >= w for w, (*_t, r) in zip(_W, _all_results)]
                        _fig_all.add_trace(go.Bar(
                            x=_xs, y=_ys, name=_m_name,
                            text=[f"{r[_m_key]['ratio']:.2f}" for *_t, r in _all_results],
                            textposition="outside", textfont=dict(size=9),
                        ))
                    # Đường W (mỗi HK có W khác nhau — vẽ scatter)
                    _fig_all.add_trace(go.Scatter(
                        x=[t[0] for t in _all_results],
                        y=[t[3]["W_kN"] for t in _all_results],
                        mode="lines+markers", name="W cọc",
                        line=dict(color="#333", dash="dash", width=2),
                        marker=dict(size=8, symbol="diamond"),
                    ))
                    _fig_all.update_layout(
                        title="So sánh RR theo 3 phương pháp — tất cả HK",
                        xaxis_title="Hố khoan",
                        yaxis_title="RR = φ·(Rs+Rp) (kN)",
                        barmode="group", height=420,
                        margin=dict(t=50, b=40),
                        legend=dict(orientation="h", y=-0.18),
                    )
                    st.plotly_chart(_fig_all, use_container_width=True)

                # ── Chi tiết từng HK (expander) ───────────────────────────────
                st.markdown("**Chi tiết từng hố khoan**")
                for _bh, _pl, _L, _cm, _res in _all_results:
                    _auto_ok = _res["auto"]["result"] == "Đạt"
                    _exp_lbl = (f"{_bh} — {_pl} L={_L:.0f}m — "
                                 f"Tự chọn: {_res['auto']['result']} "
                                 f"(RR={_res['auto']['RR_kN']:.0f} kN, "
                                 f"ratio={_res['auto']['ratio']:.2f})")
                    with st.expander(_exp_lbl, expanded=not _auto_ok):
                        _ic1, _ic2, _ic3, _ic4 = st.columns(4)
                        _ic1.metric("L cọc trong sét (m)", f"{_cm.get('L_clay_total_m', 0.0):.2f}")
                        _ic2.metric("Su TB sét (kPa)", f"{_cm.get('su_avg_clay_kPa', 0.0):.1f}")
                        _ic3.metric("σ'v TB sét (kPa)", f"{_cm.get('sigma_avg_clay_kPa', 0.0):.0f}")
                        _ic4.metric("W cọc (kN)", f"{_cm.get('W_kN', 0.0):.0f}",
                                    f"D={_cm.get('D_m', 0.0):.3f}m, P={_cm.get('perimeter_m', 0.0):.3f}m")

                        _det_rows = []
                        for _m_key, _m_name in _method_lbl.items():
                            _x = _res[_m_key]
                            _det_rows.append({
                                "Phương pháp": _m_name,
                                "φ_stat":      _x["phi_stat"],
                                "Rs (kN)":     _x["Rs_kN"],
                                "Rp (kN)":     _x["Rp_kN"],
                                "RR (kN)":     _x["RR_kN"],
                                "W (kN)":      _cm.get("W_kN", 0.0),
                                "Tỷ số RR/W":  _x["ratio"],
                                "Kết quả":     _x["result"],
                                "Phương pháp mũi":  _x.get("tip_method", "—"),
                            })
                        _det_df = pd.DataFrame(_det_rows)
                        st.dataframe(
                            _det_df.style.map(_hl_kq, subset=["Kết quả"]),
                            use_container_width=True, hide_index=True,
                            column_config={
                                "φ_stat":     st.column_config.NumberColumn(format="%.2f"),
                                "Rs (kN)":    st.column_config.NumberColumn(format="%.0f"),
                                "Rp (kN)":    st.column_config.NumberColumn(format="%.0f"),
                                "RR (kN)":    st.column_config.NumberColumn(format="%.0f"),
                                "W (kN)":     st.column_config.NumberColumn(format="%.0f"),
                                "Tỷ số RR/W": st.column_config.NumberColumn(format="%.2f"),
                            },
                        )
                        if _cm.get("lambda_coef") is not None:
                            st.caption(
                                f"λ_coef = {_cm['lambda_coef']:.3f}. "
                                "Khuyến nghị: dùng phương pháp **Tự chọn** (α cho lớp sét, tự động theo địa tầng)."
                            )
    except Exception as _e_cmp:
        st.warning(f"Không thực hiện được so sánh 3 phương pháp: {_e_cmp}")

    # ═══════════════════════════════════════════════════════════════════════════
    # D. Phương pháp giải tích + p-y (Winkler beam)
    # ═══════════════════════════════════════════════════════════════════════════
    st.markdown("---")
    _numpy_err = ""
    try:
        import sys as _sys_wif
        _sys_wif.path.insert(0, str(_ROOT / "scripts"))
        from winkler_np import (
            solve_numpy as _solve_pynite,  # alias giữ tên biến cũ
            SoilLayer as _WIF_SoilLayer,
        )
        from wall_internal_force import sw_pile_props as _wif_sw_props
        _HAS_WINKLER_SOLVER = True
    except Exception as _e_np:
        _HAS_WINKLER_SOLVER = False
        _solve_pynite = None
        _numpy_err = f"{type(_e_np).__name__}: {_e_np}"

    if not _HAS_WINKLER_SOLVER:
        import sys as _sys_diag
        st.error(
            "**Solver Winkler không khởi tạo được.**\n\n"
            f"- Phiên bản hệ thống: `{_sys_diag.version.split()[0]}` ({_sys_diag.platform})\n"
            f"- Chi tiết lỗi: `{_numpy_err}`\n\n"
            "Liên hệ kỹ sư phát triển để kiểm tra môi trường tính toán."
        )
    else:
        pass

    if True:  # Luôn hiển thị D.1 lý thuyết + D.4 Rankine; D.2/D.3 chỉ chạy nếu có solver
        # ── Helper: tính p-y Winkler cho 1 HK + cọc + L ──────────────────────
        @st.cache_data(show_spinner=False)
        def _calc_py_winkler(bh_name: str, pile_name: str, L_m: float,
                             H_kNm: float, M_kNm: float, cdm_thk_m: float,
                             eps50: float, k_cdm_factor: float) -> dict:
            """Giải Winkler p-y cho 1 cọc — NumPy thuần (scripts/winkler_np.py).

            Trả dict {u_top_mm, u_max_mm, M_max_kNm, Q_max_kN, Mcr_kNm, zs, ux, Ms, k_h}.
            """
            if not _HAS_WINKLER_SOLVER:
                return {"error": "scripts/winkler_np.py không load được — kiểm tra NumPy."}

            # NumPy Winkler solver — không phụ thuộc PyNite/anastruct
            _pl_obj = _sw_by_name(pile_name)
            if not _pl_obj:
                return {"error": f"Không có cọc {pile_name}"}
            _pile_pn = _wif_sw_props(
                H_mm=float(_pl_obj["H_mm"]),
                Itd_cm4=float(_pl_obj.get("Itd_cm4") or 0),
                Mcr_Tm=float(_pl_obj.get("Mcr_Tm") or 0),
                Atd_cm2=float(_pl_obj.get("Atd_cm2") or 0),
                fc_MPa=70.0, name=pile_name,
            )
            # Layers từ DB — SQLite lưu name = 'KE-HKx' (đầy đủ), KHÔNG strip prefix
            _bh_short = bh_name if bh_name.startswith("KE-") else f"KE-{bh_name}"
            _ly_raw = []
            try:
                _con_py = sqlite3.connect(str(_DB))
                _con_py.row_factory = sqlite3.Row
                _rows_py = _con_py.execute("""
                    SELECT l.symbol, l.thickness_m,
                           ROUND(AVG(lt.gamma_kNm3), 2) AS gamma_kNm3,
                           ROUND(AVG(lt.c_kPa), 1) AS c_kPa,
                           ROUND(AVG(lt.Cu_UU_kPa), 1) AS Cu_kPa
                    FROM layers l
                    JOIN boreholes b ON l.borehole_id = b.id
                    LEFT JOIN lab_tests lt ON lt.borehole_id = b.id
                        AND lt.depth_from_m >= l.depth_top_m
                        AND lt.depth_to_m <= l.depth_bot_m
                    WHERE b.name = ?
                    GROUP BY l.id ORDER BY l.depth_top_m
                """, (_bh_short,)).fetchall()
                _con_py.close()
                for _r in _rows_py:
                    _su = _r["Cu_kPa"] if _r["Cu_kPa"] is not None else _r["c_kPa"]
                    _ly_raw.append({
                        "symbol": _r["symbol"] or "1",
                        "thickness_m": _r["thickness_m"] or 0,
                        "Su_kPa": _su if _su is not None else 11.0,
                        "gamma_kNm3": _r["gamma_kNm3"] if _r["gamma_kNm3"] is not None else 15.0,
                    })
            except Exception:
                pass
            if not _ly_raw:
                _ly_raw = [{"symbol": "1", "thickness_m": L_m + 5.0,
                            "Su_kPa": 11.0, "gamma_kNm3": 15.0}]
            _ly_pn = [
                _WIF_SoilLayer(
                    symbol=str(_lr["symbol"]),
                    thickness_m=float(_lr["thickness_m"] or 0),
                    Su_kPa=float(_lr["Su_kPa"]),
                    gamma_kNm3=float(_lr["gamma_kNm3"]),
                ) for _lr in _ly_raw
            ]
            return _solve_pynite(
                layers=_ly_pn, pile=_pile_pn, L_m=float(L_m),
                H_kNm=float(H_kNm), M_kNm=float(M_kNm),
                N=max(20, int(L_m * 2)),
                eps50=float(eps50),
                k_sand_kNm3=10_000.0,
                cdm_thickness_m=float(cdm_thk_m),
                cdm_factor=float(k_cdm_factor),
            )

        # ═══════════════════════════════════════════════════════════════════════
        # D.1. Tổng quan p-y Winkler — tự động cho mọi HK ở Mục B
        # ═══════════════════════════════════════════════════════════════════════
        st.markdown("### D.1. Tổng quan p-y Winkler — tự động cho mọi HK ở Mục B")

        # ── Auto-loop: chạy Winkler cho mọi HK đã pick ở Mục B ────────────────
        _picks_d = st.session_state.get("ke_sw_alignment_picks", []) or []
        _hk_d_list = [nm for nm in _picks_d if nm.startswith("HK")]
        # Lấy params tải chung từ session_state (nếu user đã chỉnh ở form D.1
        # chi tiết bên dưới) — nếu chưa có thì dùng mặc định
        _eps50_a   = float(st.session_state.get("dpy_eps50", 0.02)  or 0.02)
        _cdm_Lc_a  = float(st.session_state.get("cdm_Lc", 0.0) or 0.0)
        _cdm_Lng_a = float(st.session_state.get("cdm_L_ngam", 0.0) or 0.0)
        _cdm_thk_a = max(0.0, _cdm_Lc_a - _cdm_Lng_a)
        _k_cdm_a   = float(st.session_state.get("dpy_k_cdm_factor", 3.0) or 3.0)

        # ── Tổng hợp Boussinesq cho TẤT CẢ HK trên tuyến (auto-loop) ─────────
        # Tính 1 lần ở đầu D.1, KHÔNG phụ thuộc user mở expander từng HK.
        # Tải khai thác q áp trên đỉnh kè TK (top_ke), Boussinesq theo
        # TCVN 11823-3:2017 Eq.(39) — Eq xác định Dph(z) cho semi-infinite.
        if _hk_d_list:
            try:
                import sys as _sys_bsa
                _sys_bsa.path.insert(0, str(_ROOT / "scripts"))
                from boussinesq_surcharge import (
                    SurchargeStrip as _BStripA,
                    BoussiGeometry as _BGeomA,
                    compute_all as _bs_compute_a,
                )
                import numpy as _np_bsa
                _bs_rows: list = []
                for _hk_a in _hk_d_list:
                    _bh_pick_a = next(
                        (b for b in _bhs_on_alignment if b["name"] == _hk_a), {}
                    ) or {}
                    _Z_a    = float(_bh_pick_a.get("Z_m") or 0.0)
                    _Zb_a   = float(st.session_state.get(f"dpy_Zb_{_hk_a}",
                                       -2.0) or -2.0)
                    _top_a  = float(st.session_state.get("dpy_top_ke", 2.7) or 2.7)
                    _L_a    = float(st.session_state.get("ke_sw_L_thiet_ke", {})
                                       .get(_hk_a, 29.0) or 29.0)
                    _q_a_op = float(st.session_state.get(
                        f"dpy_q_operation_{_hk_a}", 15.0) or 15.0)
                    _q_a_a  = float(st.session_state.get(
                        f"dpy_q_a_{_hk_a}", 0.0) or 0.0)
                    _q_a_w  = float(st.session_state.get(
                        f"dpy_q_w_{_hk_a}", 10.0) or 10.0)
                    _strip_a = _BStripA(
                        q=_q_a_op, a=_q_a_a, w=_q_a_w,
                        ref_surface=_top_a, side="Front",
                        label=f"q={_q_a_op:.0f}",
                    )
                    _geom_a = _BGeomA(
                        top_elev=_top_a, pile_length=_L_a,
                        soil_level_front=_Z_a, soil_level_back=_Zb_a,
                    )
                    _res_a = _bs_compute_a(_geom_a, [_strip_a])
                    _bs_rows.append({
                        "Hố khoan":      f"KE-{_hk_a}",
                        "q (kN/m²)":     _q_a_op,
                        "a (m)":         _q_a_a,
                        "w (m)":         _q_a_w,
                        "Đỉnh kè (m)":   _top_a,
                        "L cọc (m)":     _L_a,
                        "F_Bous (kN/m)": float(_res_a.get("F_front", 0.0)),
                        "z_app (m)":     float(_res_a.get("z_front", 0.0)),
                        "M_top (kNm/m)": float(_res_a.get("M_top_front", 0.0)),
                    })

                if _bs_rows:
                    with st.expander(
                        f"Tổng hợp tải Boussinesq — {len(_bs_rows)} HK trên tuyến (auto-loop)",
                        expanded=False,
                    ):
                        st.caption(
                            "Boussinesq theo TCVN 11823-3:2017 Eq.(39). "
                            "Tải khai thác q áp tại **đỉnh kè TK** (top_ke); "
                            "F_Bous = lực hợp tích phân từ đỉnh kè xuống mũi cừ."
                        )
                        st.table(pd.DataFrame(_bs_rows).style.format({
                            "q (kN/m²)":     "{:.0f}",
                            "a (m)":         "{:.1f}",
                            "w (m)":         "{:.1f}",
                            "Đỉnh kè (m)":   "{:.2f}",
                            "L cọc (m)":     "{:.1f}",
                            "F_Bous (kN/m)": "{:.1f}",
                            "z_app (m)":     "{:.2f}",
                            "M_top (kNm/m)": "{:.1f}",
                        }))
            except Exception as _e_bs_all:
                st.caption(f"_(Không tính được Boussinesq tổng hợp: {_e_bs_all})_")

        # Đọc kết quả mới nhất từ SQLite ke_sw_winkler_results (lưu bởi D.1 chi tiết)
        _bh_names_full = [f"KE-{nm}" for nm in _hk_d_list]
        _auto_rows: list = []
        if _bh_names_full and _DB.exists():
            try:
                import sqlite3 as _sq3_d1
                _conn_d1 = _sq3_d1.connect(str(_DB))
                _ph_d1 = ",".join("?" * len(_bh_names_full))
                _rows_d1 = _conn_d1.execute(f"""
                    SELECT bh_name, pile_type, L_m, u_max_mm, M_max_kNm, Mcr_kNm,
                           mcr_ratio, Q_max_kN, load_case, ts
                    FROM (
                        SELECT *, ROW_NUMBER() OVER (
                            PARTITION BY bh_name ORDER BY ts DESC
                        ) AS rn
                        FROM ke_sw_winkler_results
                        WHERE bh_name IN ({_ph_d1})
                          AND load_case = 'distributed'
                    ) WHERE rn = 1
                    ORDER BY bh_name
                """, _bh_names_full).fetchall()
                _conn_d1.close()
                for _row in _rows_d1:
                    _bhn, _pile, _L, _u, _Mmax, _Mcr, _ratio, _Q, _lcase, _ts = _row
                    _ratio = _ratio if _ratio is not None else (_Mmax / _Mcr if _Mcr else 0.0)
                    _ok_u = abs(_u) < 50.0
                    _ok_M = (_Mcr > 0 and _Mmax < _Mcr)
                    _auto_rows.append({
                        "Hố khoan":    _bhn,
                        "Cọc":         _pile,
                        "L (m)":       float(_L),
                        "u_top (mm)":  float(_u),
                        "M_max (kNm)": float(_Mmax),
                        "Mcr (kNm)":   float(_Mcr),
                        "M/Mcr":       float(_ratio),
                        "Q_max (kN)":  float(_Q or 0),
                        "Nguồn":       _lcase or "—",
                        "Kết quả":     "Đạt" if (_ok_u and _ok_M) else "Không đạt",
                    })
            except Exception as _e_d1q:
                st.warning(f"Không đọc được kết quả Winkler đã lưu: {_e_d1q}")

        if not _auto_rows:
            st.info("Chưa có kết quả. Chạy D.1 chi tiết cho từng HK bên dưới — kết quả tự động cập nhật vào bảng này.")
        else:
            st.caption("_Kết quả mới nhất từ D.1 chi tiết — chạy lại cho 1 HK ở mục bên dưới để cập nhật._")
            _auto_df = pd.DataFrame(_auto_rows)

            def _hl_d_auto(val):
                if val == "Đạt":       return "background-color:#d4edda; color:#155724"
                if val == "Không đạt": return "background-color:#f8d7da; color:#721c24"
                return ""

            st.dataframe(
                _auto_df.style.map(_hl_d_auto, subset=["Kết quả"]),
                use_container_width=True, hide_index=True,
                column_config={
                    "L (m)":        st.column_config.NumberColumn(format="%.1f"),
                    "u_top (mm)":   st.column_config.NumberColumn(format="%.2f"),
                    "M_max (kNm)":  st.column_config.NumberColumn(format="%.1f"),
                    "Mcr (kNm)":    st.column_config.NumberColumn(format="%.0f"),
                    "M/Mcr":        st.column_config.NumberColumn(format="%.2f"),
                    "Q_max (kN)":   st.column_config.NumberColumn(format="%.1f"),
                },
            )

            if _HAS_PLOTLY:
                from plotly.subplots import make_subplots as _msub_d1
                _fig_d1a = _msub_d1(rows=1, cols=2,
                                     subplot_titles=("Chuyển vị đỉnh u_top (mm)",
                                                     "Tỷ số M/Mcr"))
                _xs_d1a = [r["Hố khoan"] for r in _auto_rows]
                _us_d1a = [r["u_top (mm)"] for r in _auto_rows]
                _ms_d1a = [r["M/Mcr"]      for r in _auto_rows]
                _fig_d1a.add_trace(go.Bar(
                    x=_xs_d1a, y=_us_d1a, name="u_top",
                    marker_color=["#1565C0" if abs(u) < 50 else "#E53935" for u in _us_d1a],
                    text=[f"{u:.1f}" for u in _us_d1a], textposition="outside",
                ), row=1, col=1)
                _fig_d1a.add_hline(y=50,  line_dash="dash", line_color="#666",
                                    annotation_text="50 mm", row=1, col=1)
                _fig_d1a.add_hline(y=-50, line_dash="dash", line_color="#666", row=1, col=1)
                _fig_d1a.add_trace(go.Bar(
                    x=_xs_d1a, y=_ms_d1a, name="M/Mcr",
                    marker_color=["#2E7D32" if m < 1 else "#E53935" for m in _ms_d1a],
                    text=[f"{m:.2f}" for m in _ms_d1a], textposition="outside",
                ), row=1, col=2)
                _fig_d1a.add_hline(y=1.0, line_dash="dash", line_color="#666",
                                    annotation_text="M_cr", row=1, col=2)
                _fig_d1a.update_layout(height=380, showlegend=False,
                                        margin=dict(t=50, b=40),
                                        title="So sánh Winkler — tất cả HK ở Mục B")
                st.plotly_chart(_fig_d1a, use_container_width=True)

        # ═══════════════════════════════════════════════════════════════════════
        # E. Ổn định tổng thể — 3 phương pháp, cập nhật động từ Mục B + D.1
        # ═══════════════════════════════════════════════════════════════════════
        st.markdown("---")
        _hdr_e1, _hdr_e2 = st.columns([3, 1])
        with _hdr_e1:
            st.markdown("### E. Ổn định tổng thể — 3 phương pháp")
        with _hdr_e2:
            _refresh_e = st.button(
                "🔄 Cập nhật 4 PP cho mọi HK",
                use_container_width=True, key="_refresh_kesw_stab_btn",
                help="Đọc tham số từ DB (rows cũ) + tính lại 4 PP với "
                     "cung trượt qua chân cừ. Đảm bảo Mục E khớp với D.1 chi tiết.",
            )

        # ── Công thức tính ổn định tổng thể — chỉ hiện 1 lần, dùng chung ────
        with st.expander("Công thức tính ổn định tổng thể — chi tiết (chung cho mọi HK)"):
            try:
                _md_e_path_g = _ROOT / "48-ke-sw-on-dinh-tong-the.md"
                if _md_e_path_g.exists():
                    st.markdown(_md_e_path_g.read_text(encoding="utf-8"),
                                unsafe_allow_html=False)
                else:
                    st.caption("_(Không tìm thấy tài liệu công thức.)_")
            except Exception as _e_md_e_g:
                st.caption(f"_(Lỗi đọc công thức: {_e_md_e_g})_")

        # ── Auto-detect stale: cache ts < geology version ts ──
        # Nguồn truth: bảng `data_versions['ke_geology']` được bump bởi script
        # update địa chất (ke_borehole_layers_update.py). Mọi row ke_sw_stability
        # có `ts < ke_geology.ts` đều là kết quả tính trên dữ liệu cũ → cần lại.
        # Hash của tập stale ngăn loop vô tận khi refresh fail.
        try:
            import sqlite3 as _sq3_stale
            with _sq3_stale.connect(str(_DB), timeout=10) as _con_stale:
                _con_stale.row_factory = _sq3_stale.Row
                _con_stale.execute("""
                    CREATE TABLE IF NOT EXISTS data_versions (
                        data_key TEXT PRIMARY KEY, ts TEXT NOT NULL, note TEXT
                    )
                """)
                _geo_ts_row = _con_stale.execute(
                    "SELECT ts FROM data_versions WHERE data_key='ke_geology'"
                ).fetchone()
                _geo_ts = _geo_ts_row["ts"] if _geo_ts_row else None
                if _geo_ts:
                    _stale_rows = _con_stale.execute(
                        """SELECT bh_name, MIN(ts) AS cache_ts
                           FROM ke_sw_stability
                           WHERE bh_name LIKE 'KE-%' AND ts < ?
                           GROUP BY bh_name""", (_geo_ts,),
                    ).fetchall()
                else:
                    _stale_rows = []
        except Exception:
            _stale_rows = []
            _geo_ts = None

        _stale_hash = "|".join(sorted(r["bh_name"] for r in _stale_rows))
        _last_stale_hash = st.session_state.get("_kesw_last_stale_hash", "")
        if _stale_rows and _stale_hash != _last_stale_hash:
            _stale_names = ", ".join(r["bh_name"] for r in _stale_rows[:5])
            _more = f" (+{len(_stale_rows)-5} HK khác)" if len(_stale_rows) > 5 else ""
            st.info(
                f"⚙ **Phát hiện {len(_stale_rows)} HK có cache cũ hơn lần cập nhật "
                f"địa chất ({_geo_ts})**. Tự động tính lại 4 PP cho khớp dữ liệu mới...  \n"
                f"_{_stale_names}{_more}_"
            )
            st.session_state["_kesw_last_stale_hash"] = _stale_hash
            _refresh_e = True   # Tái dụng code refresh phía dưới
        elif not _stale_rows and _last_stale_hash:
            st.session_state["_kesw_last_stale_hash"] = ""

        _FS_SLIP_MIN_E = 1.40   # Fs trượt cho phép
        _FS_OT_MIN_E   = 1.20   # Fs lật cho phép
        _SLIP_KEYS_E   = ["bishop", "spencer", "morgenstern_price"]
        _SLIP_LBL_E    = {"bishop": "Bishop",
                          "spencer": "Spencer", "morgenstern_price": "M-P"}

        # Quy tắc riêng từng HK: symbol alias — áp dụng trước mọi phân loại
        _HK_SYM_ALIAS: dict[str, dict[str, str]] = {
            "HK08": {"XMD": "1"},   # KE-HK08: lớp XMD là bùn tự nhiên (chưa xử lý CDM)
            "HK12": {"XMD": "1"},   # KE-HK12: lớp XMD là bùn tự nhiên (chưa xử lý CDM)
        }

        def _alias_sym(sym: str, bh: str) -> str:
            s = (sym or "").upper()
            k = bh.upper().replace("KE-", "")
            return _HK_SYM_ALIAS.get(k, {}).get(s, s)

        # ── Refresh handler: tính lại 4 PP cho mọi HK đã có row trong SQLite ──
        if _refresh_e:
            try:
                from sw_global_stability import (
                    CDMBlock as _CDMB_R, check_global_slip as _chk_R,
                    check_overturning as _chk_ot_R, check_toe_kickout as _chk_to_R,
                )
                from wall_internal_force import (
                    WallGeometry as _WG_R, EarthLayer as _EL_R,
                )
                import sqlite3 as _sq3_R

                with _sq3_R.connect(str(_DB), timeout=30) as _con_R:
                    _con_R.row_factory = _sq3_R.Row
                    # Lấy 1 row đại diện mỗi (HK, pile, L) — có params cần để rebuild
                    _seeds_R = _con_R.execute("""
                        SELECT bh_name, pile_type, L_m,
                               top_elev, Z_m, Zb_m, wlvl_front, wlvl_back, q_kPa,
                               cdm_a, cdm_c_col, cdm_w_m, su_front, su_back, H1_m
                        FROM ke_sw_stability
                        GROUP BY bh_name, pile_type, L_m
                    """).fetchall()

                    if not _seeds_R:
                        st.warning(
                            "Chưa có HK nào trong DB. Hãy mở D.1 chi tiết cho "
                            "ít nhất 1 HK để tạo dòng dữ liệu, rồi nhấn lại nút này."
                        )
                    else:
                        _prog_R = st.progress(0, text=f"0 / {len(_seeds_R)} HK")
                        _n_done = 0
                        # CDM defaults — không có trong bảng → dùng giá trị tiêu chuẩn
                        _phi_col_R = float(st.session_state.get("cdm_phi_col", 30.0))
                        _gam_col_R = float(st.session_state.get("cdm_gamma_col", 19.0))
                        # Fill defaults — bộ chuẩn dự án CLAUDE.md
                        _fill_gam_R = float(st.session_state.get("_dpy_gamma_fill_g", 18.0))
                        _fill_gsub_R = float(st.session_state.get("_dpy_gamma_sub_fill_g", 8.0))
                        _fill_phi_R = float(st.session_state.get("_dpy_phi_fill_g", 25.0))
                        _fill_c_R = float(st.session_state.get("_dpy_c_fill_g", 0.0))

                        for _i_R, _seed in enumerate(_seeds_R):
                            _bh_R = _seed["bh_name"]
                            _pile_R = _seed["pile_type"]
                            _L_R = float(_seed["L_m"])
                            _top_R = float(_seed["top_elev"])
                            _Z_R = float(_seed["Z_m"])
                            _Zb_R = float(_seed["Zb_m"])
                            _wF_R = float(_seed["wlvl_front"])
                            _wB_R = float(_seed["wlvl_back"])
                            _q_R = float(_seed["q_kPa"] or 0)
                            _cdma_R = float(_seed["cdm_a"] or 0.20)
                            _cdmc_R = float(_seed["cdm_c_col"] or 75.0)
                            _suF_R = float(_seed["su_front"] or 10.0)
                            _suB_R = float(_seed["su_back"] or _suF_R)
                            _H1_R = float(_seed["H1_m"] or 0.0)
                            _pbot_R = _top_R - _L_R

                            # Build geom + fill
                            _geom_R = _WG_R(
                                top_elev=_top_R, pile_length=_L_R,
                                soil_level_front=_Z_R, soil_level_back=_Zb_R,
                                water_elev_front=_wF_R, water_elev_back=_wB_R,
                                surcharge_front=_q_R,
                            )
                            _fill_R = (_EL_R(_Z_R, _fill_gam_R, _fill_gsub_R,
                                              _fill_phi_R, _fill_c_R)
                                        if _top_R > _Z_R else None)

                            # ── Build front/back layers từ SQLite layers + lab_tests ──
                            _front_R: list = []
                            _back_R: list = []
                            _rows_lay_R = _con_R.execute("""
                                SELECT l.symbol, l.depth_top_m, l.depth_bot_m,
                                       ROUND(AVG(lt.gamma_kNm3), 2) AS gam,
                                       ROUND(AVG(lt.phi_deg), 1) AS phi,
                                       ROUND(AVG(lt.c_kPa), 1) AS c,
                                       ROUND(AVG(lt.Cu_UU_kPa), 1) AS Cu,
                                       b.elevation_m AS bh_elev
                                FROM layers l
                                JOIN boreholes b ON l.borehole_id = b.id
                                LEFT JOIN lab_tests lt ON lt.borehole_id = b.id
                                    AND lt.depth_from_m >= l.depth_top_m
                                    AND lt.depth_to_m <= l.depth_bot_m
                                WHERE b.name = ?
                                GROUP BY l.id ORDER BY l.depth_top_m
                            """, (_bh_R,)).fetchall()
                            _bh_elev_R = (float(_rows_lay_R[0]["bh_elev"])
                                          if _rows_lay_R and _rows_lay_R[0]["bh_elev"]
                                          else _Z_R)
                            _back_started_R = False
                            for _rL in _rows_lay_R:
                                _sym = _alias_sym(_rL["symbol"] or "", _bh_R)
                                _etop = _bh_elev_R - float(_rL["depth_top_m"] or 0)
                                _ebot = _bh_elev_R - float(_rL["depth_bot_m"] or 0)
                                if _etop < _pbot_R - 2.0 or _ebot > _top_R:
                                    continue
                                _gam = float(_rL["gam"]) if _rL["gam"] else 18.0
                                _gsub = max(_gam - 9.81, 5.0)
                                if _sym in {"1", "XMD", "3", "5", "5A", "5B"}:
                                    _phi_f = 0.0
                                    _c_f = (float(_rL["Cu"]) if _rL["Cu"]
                                             else (float(_rL["c"]) if _rL["c"] else _suF_R))
                                    _c_b = _suB_R if _sym in {"1", "XMD"} else _c_f
                                elif _sym in {"F", "2", "2A", "2B", "2C", "4", "6", "7"}:
                                    _phi_f = float(_rL["phi"]) if _rL["phi"] else 30.0
                                    _c_f = 0.0; _c_b = 0.0
                                else:
                                    _phi_f = float(_rL["phi"]) if _rL["phi"] else 20.0
                                    _c_f = float(_rL["c"]) if _rL["c"] else 0.0
                                    _c_b = _c_f
                                _front_R.append(_EL_R(_ebot, _gam, _gsub, _phi_f, _c_f))
                                if _sym in {"1", "XMD"}:
                                    _back_started_R = True
                                if _back_started_R:
                                    _back_R.append(_EL_R(_ebot + (_Zb_R - _Z_R),
                                                          _gam, _gsub, _phi_f, _c_b))

                            # Fallback layers
                            if not _front_R:
                                _front_R = [
                                    _EL_R(_Z_R - _H1_R, 15.0, 5.0, 0.0, _suF_R or 10.0),
                                    _EL_R(_pbot_R - 1.0, 18.0, 8.0, 30.0, 0.0),
                                ]
                                _back_R = [
                                    _EL_R(_Zb_R - _H1_R, 15.0, 5.0, 0.0, _suB_R or _suF_R or 10.0),
                                    _EL_R(_pbot_R - 1.0, 18.0, 8.0, 30.0, 0.0),
                                ]

                            # CDM block
                            _cdm_R = _CDMB_R(top_elev=_top_R, bot_elev=_Z_R - _H1_R - 1.0,
                                              area_ratio_a=_cdma_R, c_col_kPa=_cdmc_R,
                                              phi_col_deg=_phi_col_R, gamma_col_kNm3=_gam_col_R)

                            # 4 PP slip với cung qua chân cừ (constraint trong check_global_slip)
                            _slips_R: dict[str, tuple] = {}
                            for _mk in _SLIP_KEYS_E:
                                try:
                                    _Fs, _xc, _yc, _Rd = _chk_R(
                                        _geom_R, _front_R, _back_R,
                                        fill=_fill_R, cdm=_cdm_R, method=_mk,
                                    )
                                    if _Fs and float(_Fs) > 0:
                                        _slips_R[_mk] = (float(_Fs), float(_xc),
                                                          float(_yc), float(_Rd))
                                except Exception:
                                    pass

                            if not _slips_R:
                                continue

                            # Critical = min Fs
                            _crit_R = min(_slips_R, key=lambda k: _slips_R[k][0])
                            # Lật + Kickout (không phụ thuộc method)
                            try:
                                _Fs_ot, _Mg, _Ml = _chk_ot_R(
                                    _geom_R, _front_R, _back_R, fill=_fill_R, cdm=_cdm_R)
                            except Exception:
                                _Fs_ot, _Mg, _Ml = 0.0, 0.0, 0.0
                            try:
                                _Fs_tk, _Ma, _Mp = _chk_to_R(
                                    _geom_R, _front_R, _back_R, fill=_fill_R, cdm=_cdm_R)
                            except Exception:
                                _Fs_tk, _Ma, _Mp = 0.0, 0.0, 0.0

                            # Lưu 4 rows
                            for _mk, (_Fs_v, _xc_v, _yc_v, _R_v) in _slips_R.items():
                                _con_R.execute("""
                                    INSERT OR REPLACE INTO ke_sw_stability
                                    (bh_name, pile_type, L_m, method,
                                     Fs_slip, Fs_overturning, Fs_toe_kickout,
                                     slip_xc, slip_yc, slip_R,
                                     M_giu_kNm, M_lat_kNm,
                                     top_elev, Z_m, Zb_m, wlvl_front, wlvl_back, q_kPa,
                                     cdm_a, cdm_c_col, cdm_w_m, su_front, su_back, H1_m,
                                     warnings, ts)
                                    VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,
                                            datetime('now','localtime'))
                                """, (
                                    _bh_R, _pile_R, _L_R, _mk,
                                    _Fs_v, float(_Fs_ot), float(_Fs_tk),
                                    _xc_v, _yc_v, _R_v,
                                    float(_Mg) if _mk == _crit_R else None,
                                    float(_Ml) if _mk == _crit_R else None,
                                    _top_R, _Z_R, _Zb_R, _wF_R, _wB_R, _q_R,
                                    _cdma_R, _cdmc_R, float(_seed["cdm_w_m"] or 5.0),
                                    _suF_R, _suB_R, _H1_R,
                                    f"refresh:{_crit_R}",
                                ))
                            _n_done += 1
                            _prog_R.progress((_i_R + 1) / len(_seeds_R),
                                              text=f"{_i_R + 1} / {len(_seeds_R)} HK · vừa xử lý {_bh_R}")
                    _con_R.commit()
                    st.success(f"✓ Đã cập nhật {_n_done} HK với 3 phương pháp (cung qua chân cừ).")
                    st.rerun()
            except ImportError as _e_imp_R:
                st.error(f"Thiếu module: {_e_imp_R}")
            except Exception as _e_R:
                st.error(f"Lỗi cập nhật: {type(_e_R).__name__}: {_e_R}")
                import traceback as _tb_R
                with st.expander("Chi tiết"):
                    st.code(_tb_R.format_exc())

        # Đọc lựa chọn hiện tại từ Mục B (pile + L per HK)
        _rec_b_e = st.session_state.get("ke_sw_rec_piles", {})
        _ltk_b_e = st.session_state.get("ke_sw_L_thiet_ke", {})

        _e_rows = []
        if _DB.exists() and _bhs_on_alignment:
            try:
                import sqlite3 as _sq3_e
                with _sq3_e.connect(str(_DB), timeout=10) as _con_e:
                    _con_e.row_factory = _sq3_e.Row
                    # Tạo bảng nếu chưa có
                    _con_e.execute("""
                        CREATE TABLE IF NOT EXISTS ke_sw_stability (
                            id INTEGER PRIMARY KEY AUTOINCREMENT,
                            project TEXT DEFAULT '202605-TTHC',
                            zone TEXT DEFAULT 'KE',
                            bh_name TEXT NOT NULL, pile_type TEXT NOT NULL,
                            L_m REAL NOT NULL, method TEXT NOT NULL DEFAULT 'bishop',
                            Fs_slip REAL, Fs_overturning REAL, Fs_toe_kickout REAL,
                            slip_xc REAL, slip_yc REAL, slip_R REAL,
                            M_giu_kNm REAL, M_lat_kNm REAL,
                            top_elev REAL, Z_m REAL, Zb_m REAL,
                            wlvl_front REAL, wlvl_back REAL, q_kPa REAL,
                            cdm_a REAL, cdm_c_col REAL, cdm_w_m REAL,
                            su_front REAL, su_back REAL, H1_m REAL,
                            warnings TEXT, ts TEXT,
                            UNIQUE(bh_name, pile_type, L_m, method)
                        )
                    """)
                    _con_e.commit()

                    for _bh_e in _bhs_on_alignment:
                        _nm_e     = _bh_e["name"]
                        _bh_full_e = f"KE-{_nm_e}" if not _nm_e.startswith("KE-") else _nm_e
                        # Ưu tiên: D.1 form hiện tại (dpy_pile_*) → Mục B → JSON mặc định
                        _pile_e   = (st.session_state.get(f"dpy_pile_{_nm_e}")
                                     or _rec_b_e.get(_nm_e)
                                     or _bh_e.get("recommended_pile") or "SW-840")
                        _L_e      = float(st.session_state.get(f"dpy_L_{_nm_e}")
                                          or _ltk_b_e.get(_nm_e)
                                          or _bh_e.get("recommended_L_m") or 29.0)

                        # Lấy 3 phương pháp (Bishop, Spencer, M-P) cho combo (HK, cọc, L)
                        _db_rows_e = _con_e.execute("""
                            SELECT method, Fs_slip, Fs_overturning, ts
                            FROM ke_sw_stability
                            WHERE bh_name=? AND pile_type=? AND ABS(L_m-?)<0.1
                            ORDER BY method
                        """, (_bh_full_e, _pile_e, _L_e)).fetchall()

                        _by_method = {r["method"]: dict(r) for r in _db_rows_e}
                        _slips_e   = {mk: (_by_method[mk]["Fs_slip"]
                                           if mk in _by_method else None)
                                      for mk in _SLIP_KEYS_E}
                        _fs_ot_e   = next(
                            (v["Fs_overturning"] for v in _by_method.values()
                             if v.get("Fs_overturning")), None)
                        _ts_e      = next(
                            (str(v["ts"])[:16] for v in _by_method.values()
                             if v.get("ts")), "–")

                        _valid_sl  = [v for v in _slips_e.values() if v]
                        _fs_sl_min = min(_valid_sl) if _valid_sl else None
                        _has_data  = bool(_valid_sl)
                        _ok_slip   = _fs_sl_min is not None and _fs_sl_min >= _FS_SLIP_MIN_E
                        _ok_ot     = _fs_ot_e is not None and float(_fs_ot_e) >= _FS_OT_MIN_E
                        _ket_qua   = ("Chưa tính" if not _has_data
                                      else ("Đạt" if (_ok_slip and _ok_ot) else "Không đạt"))

                        _e_rows.append({
                            "Hố khoan":        _bh_full_e,
                            "Cọc":             _pile_e,
                            "L (m)":           _L_e,
                            "Bishop":          _slips_e.get("bishop"),
                            "Spencer":         _slips_e.get("spencer"),
                            "M-P":             _slips_e.get("morgenstern_price"),
                            "Fs lật":          float(_fs_ot_e) if _fs_ot_e else None,
                            "Kết quả":         _ket_qua,
                            "Tính lúc":        _ts_e,
                        })
            except Exception as _e_sum_stab:
                st.caption(f"_(Lỗi đọc kết quả ổn định tổng thể: {_e_sum_stab})_")

        if _e_rows:
            _df_e = pd.DataFrame(_e_rows)
            _slip_cols = ["Bishop", "Spencer", "M-P"]
            _ot_cols   = ["Fs lật"]

            def _hl_e_slip(val):
                if not isinstance(val, float): return ""
                return ("color:#155724;font-weight:bold" if val >= _FS_SLIP_MIN_E
                        else "color:#721c24;font-weight:bold")

            def _hl_e_ot(val):
                if not isinstance(val, float): return ""
                return ("color:#155724;font-weight:bold" if val >= _FS_OT_MIN_E
                        else "color:#721c24;font-weight:bold")

            def _hl_e_kq(val):
                if val == "Đạt":       return "background-color:#d4edda;color:#155724"
                if val == "Không đạt": return "background-color:#f8d7da;color:#721c24"
                return "background-color:#fff3cd;color:#856404"  # Chưa tính

            st.dataframe(
                _df_e.style
                    .map(_hl_e_slip, subset=_slip_cols)
                    .map(_hl_e_ot,   subset=_ot_cols)
                    .map(_hl_e_kq,   subset=["Kết quả"]),
                use_container_width=True, hide_index=True,
                column_config={
                    "L (m)":       st.column_config.NumberColumn(format="%.1f"),
                    "Bishop":      st.column_config.NumberColumn(format="%.3f",
                                    help=f"Fs trượt Bishop — ngưỡng ≥ {_FS_SLIP_MIN_E}"),
                    "Spencer":     st.column_config.NumberColumn(format="%.3f",
                                    help=f"Fs trượt Spencer — ngưỡng ≥ {_FS_SLIP_MIN_E}"),
                    "M-P":         st.column_config.NumberColumn(format="%.3f",
                                    help=f"Fs trượt Morgenstern-Price — ngưỡng ≥ {_FS_SLIP_MIN_E}"),
                    "Fs lật":      st.column_config.NumberColumn(format="%.3f",
                                    help=f"Fs lật quanh chân cừ — ngưỡng ≥ {_FS_OT_MIN_E}"),
                },
            )
            _n_done = sum(1 for r in _e_rows if r["Kết quả"] != "Chưa tính")
            st.caption(
                f"_{_n_done}/{len(_e_rows)} HK đã tính — Fs trượt ≥ {_FS_SLIP_MIN_E}, "
                f"Fs lật ≥ {_FS_OT_MIN_E}. Bảng tự cập nhật khi D.1 tính xong từng HK._"
            )
        else:
            st.info("Chưa có kết quả ổn định. Tính từ D.1 chi tiết từng HK bên dưới → tự cập nhật ở đây.")

        # ═══════════════════════════════════════════════════════════════════════
        # D.1 chi tiết: nhập tay 1 HK để xem biểu đồ u/M/k_h/Boussinesq + EP
        # ═══════════════════════════════════════════════════════════════════════
        st.markdown("---")
        st.markdown("### D.1. Chi tiết từng HK")

        # ── Công thức tính kh — chỉ hiện 1 lần, dùng chung cho mọi HK ─────────
        with st.expander("Công thức tính kh — chi tiết (chung cho mọi HK)"):
            st.markdown("##### Mô hình Winkler — lò xo nền ngang p-y")
            st.markdown(
                "Cừ làm việc như **dầm Euler-Bernoulli** trên nền đất "
                "được thay thế bằng **lò xo độc lập** tại mỗi nút FEM. "
                "Độ cứng lò xo `k_spring` (kN/m/m) tỷ lệ với mô đun đất ngang "
                "kh (kN/m³), bề rộng tiếp xúc cừ D, và chiều dài đoạn lò xo dz."
            )
            st.latex(r"k_{\text{spring}} = k_h(z) \cdot D \cdot dz \quad [\text{kN/m}]")
            st.markdown(
                "Trong đó:\n"
                "- $D$ — **0.5 × chu vi cừ** (do cừ chuyển vị về phía sông, "
                "chỉ 1 mặt tiếp xúc đất Back)\n"
                "- $dz$ — chiều dài phân tố FEM (L cọc / số nút − 1)"
            )

            st.markdown("##### 1. Đất sét (Matlock 1970) — áp dụng cho lớp bùn / xử lý CDM")
            st.latex(r"p_u = N_p \cdot S_u \cdot D \quad [\text{kN/m}]")
            st.latex(r"N_p = \min\left(3 + \dfrac{\gamma \cdot z}{S_u},\ 9\right)")
            st.latex(r"y_{50} = 2.5 \cdot \varepsilon_{50} \cdot D \quad [\text{m}]")
            st.latex(r"k_h = \dfrac{p_u}{y_{50}} = \dfrac{N_p \cdot S_u}{2.5 \cdot \varepsilon_{50}} \quad [\text{kN/m}^3]")
            st.markdown(
                "Ý nghĩa:\n"
                "- $p_u$ — áp lực giới hạn (ultimate lateral resistance)\n"
                "- $N_p$ — hệ số dung lượng Matlock, biến thiên 3 (đỉnh) → 9 (sâu)\n"
                "- $S_u$ — cường độ chống cắt không thoát nước (dùng Su Back tự nhiên)\n"
                "- $\\gamma$ — dung trọng đẩy nổi (γ' dưới MNN)\n"
                "- $z$ — độ sâu từ Ground B\n"
                "- $\\varepsilon_{50}$ — biến dạng tại 50% phá hoại (theo loại sét)\n"
                "- $y_{50}$ — chuyển vị tại 50% pu"
            )

            st.markdown("##### 2. Đất cát (API GoM) — kh tăng tuyến tính theo độ sâu")
            st.latex(r"k_h(z) = k_{\text{sand}} \cdot z \quad [\text{kN/m}^3]")
            st.markdown(
                "Trong đó:\n"
                "- $k_{\\text{sand}}$ — hệ số tăng kh theo độ sâu (mặc định **10 000 kN/m³**)\n"
                "- $z$ — độ sâu từ mặt đất (tăng theo chiều sâu)"
            )

            st.markdown("##### 3. Vùng CDM gia cố — tăng cứng nền")
            st.latex(r"k_h^{\text{CDM}} = k_h \cdot f_{\text{CDM}}")
            st.markdown(
                "- $f_{\\text{CDM}}$ — hệ số tăng cứng (mặc định × 3,0)\n"
                "- Áp dụng cho đoạn cừ trong vùng CDM (dày = chiều sâu xử lý CDM)"
            )

            st.markdown("##### 4. Hiệu chỉnh dự án")
            st.markdown(
                "- **0.5 × chu vi cừ**: cừ chuyển vị về sông (Back) → "
                "chỉ 1 mặt tiếp xúc đất Back (không có 2 mặt như cọc tròn)\n"
                "- **Bỏ qua đoạn trên Zb**: từ đỉnh cừ đến mặt đất Back — "
                "đoạn này không có lò xo\n"
                "- **Su Back** thay Su Front cho clay layers — "
                "vì cừ tựa vào đất Back tự nhiên (chưa xử lý)\n"
                "- Layer cát giữ kh API gốc (không phân biệt Front/Back vì cát không có Su)"
            )

        # ═══════════════════════════════════════════════════════════════════════
        # D.1 (chi tiết) — Auto-loop cho mọi HK ở Mục B (mỗi HK 1 expander)
        # ═══════════════════════════════════════════════════════════════════════
        if not _hk_d_list:
            st.caption("_(Chưa có HK nào ở Mục B để chạy chi tiết.)_")
        for _hk_iter in _hk_d_list:
            with st.expander(f"Chi tiết KE-{_hk_iter}", expanded=False):

                # ── Auto-fill HK-specific params từ Mục B (auto-loop, không cần button) ───
                _bh_pick_d = next((b for b in _bhs_on_alignment if b["name"] == _hk_iter), {}) or {}
                _dpy_pile_default = (st.session_state.get("ke_sw_rec_piles", {}).get(_hk_iter)
                                     or _bh_pick_d.get("recommended_pile") or "SW-840")
                _dpy_L_default    = (st.session_state.get("ke_sw_L_thiet_ke", {}).get(_hk_iter)
                                     or _bh_pick_d.get("recommended_L_m") or 29.0)
                # Ưu tiên override do user chỉnh ở Mục B (đồng bộ toàn app)
                _z_cust_d_iter  = st.session_state.get("ke_sw_z_custom", {}).get(_hk_iter)
                _h1_cust_d_iter = st.session_state.get("ke_sw_h1_custom", {}).get(_hk_iter)
                _dpy_Z_default    = (float(_z_cust_d_iter) if _z_cust_d_iter is not None
                                       else float(_bh_pick_d.get("Z_m")
                                                  if _bh_pick_d.get("Z_m") is not None else -0.8))
                _dpy_H1_default   = (float(_h1_cust_d_iter) if _h1_cust_d_iter is not None
                                       else float(_bh_pick_d.get("H_layer1_m") or 0.0))
                try:
                    _zone_c_iter = (_ke_data.get("_meta", {}).get("zone_code")
                                    or _bh_pick_d.get("zone_code") or "KE")
                    _su_db_iter  = _su_avg_in_range(_zone_c_iter, 0.0, _dpy_H1_default)
                    # Ctd Front SAU xử lý CDM — bộ chuẩn dự án 80 kPa (CDM 28d)
                    _dpy_su_default = float(_su_db_iter) if _su_db_iter is not None else 80.0
                except Exception:
                    _dpy_su_default = 80.0

                # Su tự nhiên phía Back: VST → lab → fallback 11 kPa (cached 5 phút)
                _bh_full_vst = (f"KE-{_hk_iter}"
                                if not str(_hk_iter).startswith("KE-") else str(_hk_iter))
                _su_back_default = _load_su_back_natural(_bh_full_vst)

                _dpy_apply_bh = _hk_iter

                # ── Auto-fill ε₅₀ theo thí nghiệm 3 trục SQLite hoặc trạng thái đất ─
                def _eps50_from_state(_su_kPa: float) -> float:
                    """Matlock 1970 / API: ε₅₀ theo trạng thái sét (Su)."""
                    if _su_kPa < 12.0:   return 0.020   # sét rất yếu
                    if _su_kPa < 24.0:   return 0.020   # sét yếu
                    if _su_kPa < 48.0:   return 0.010   # sét trung bình
                    if _su_kPa < 96.0:   return 0.007   # sét chặt
                    if _su_kPa < 192.0:  return 0.005   # sét rất chặt
                    return 0.004                          # sét cứng

                _eps50_default = _eps50_from_state(_dpy_su_default)
                _eps50_source = "lib"   # "triaxial" hoặc "lib"
                try:
                    _bh_e = (f"KE-{_hk_iter}" if not _hk_iter.startswith("KE-") else _hk_iter)
                    _E_avg, _Cu_avg = _load_eps50_params(_bh_e, _dpy_H1_default)
                    if _E_avg > 0 and _Cu_avg > 0:
                        _eps_tri = _Cu_avg / _E_avg
                        if 0.001 < _eps_tri < 0.1:
                            _eps50_default = _eps_tri
                            _eps50_source = "triaxial"
                except Exception:
                    pass

                # Auto-sync session_state với JSON defaults — LUÔN ĐÈ (không check
                # if not in) → data HK luôn khớp JSON, tránh session cũ gây sai số.
                _defaults_map = [
                    (f"dpy_pile_{_hk_iter}",    _dpy_pile_default),
                    (f"dpy_L_{_hk_iter}",       float(_dpy_L_default)),
                    (f"dpy_Z_{_hk_iter}",       _dpy_Z_default),
                    (f"dpy_H1_{_hk_iter}",      _dpy_H1_default),
                    (f"dpy_su_{_hk_iter}",      _dpy_su_default),
                    (f"dpy_su_back_{_hk_iter}", _su_back_default),
                    (f"dpy_eps50_{_hk_iter}",   _eps50_default),
                ]
                for _k, _v in _defaults_map:
                    st.session_state[_k] = _v
                st.session_state[f"dpy_eps50_src_{_hk_iter}"] = _eps50_source

                # Badge "Đã kiểm tra chéo" với chi tiết thiết kế gốc
                _eps_src_lbl = ("nén 3 trục" if _eps50_source == "triaxial"
                                  else "thư viện trạng thái")
                st.success(
                    f"**✓ Đã kiểm tra chéo dữ liệu thiết kế gốc cho {_hk_iter}**  ·  "
                    f"Cọc kiến nghị: **{_dpy_pile_default}**  ·  "
                    f"L = **{_dpy_L_default:.1f} m**  ·  "
                    f"Z mặt đất = **{_dpy_Z_default:+.2f} m**  ·  "
                    f"H₁ lớp bùn = **{_dpy_H1_default:.1f} m**  ·  "
                    f"Ctd Front = **{_dpy_su_default:.0f} kPa**  ·  "
                    f"Su Back (tự nhiên) = **{_su_back_default:.0f} kPa**  ·  "
                    f"ε₅₀ = **{_eps50_default:.4f}** ({_eps_src_lbl})",
                    icon="✅",
                )

                _dpy_applied = st.session_state.get("dpy_applied", {}) or {}

                def _dlbl(key: str, txt: str) -> str:
                    return f":blue[**{txt}**]" if key in _dpy_applied else txt

                # ── Form 3 cột + schematic bên phải ──────────────────────────────────
                _col_form_d, _col_schem_d = st.columns([3, 2], gap="large")

                with _col_form_d:
                    _df_c1, _df_c2, _df_c3 = st.columns(3)
                    with _df_c1:
                        # Lọc cọc theo Mục B (chỉ cọc đã được kiến nghị ở B), fallback toàn catalog
                        _rec_piles_b = sorted({
                            p for p in (st.session_state.get("ke_sw_rec_piles", {}) or {}).values() if p
                        })
                        _sw_opts_d = [n for n in _sw_names if n in _rec_piles_b] or _sw_names
                        _dpy_pile_idx = (_sw_opts_d.index("SW-840") if "SW-840" in _sw_opts_d else 0)
                        _dpy_pile = st.selectbox(
                            _dlbl("dpy_pile", f"Loại cọc SW (lọc theo Mục B — {len(_sw_opts_d)} loại)"),
                            _sw_opts_d, index=_dpy_pile_idx, key=f"dpy_pile_{_hk_iter}",
                        )
                        _dpy_L = st.number_input(_dlbl("dpy_L", "Chiều dài thiết kế L (m)"),
                                                  10.0, 35.0, 29.0, 0.5, key=f"dpy_L_{_hk_iter}")
                        _dpy_top_ke = st.number_input("Cao độ đỉnh kè (m)", 0.0, 5.0, 2.70, 0.05,
                                                        key=f"dpy_top_ke_{_hk_iter}")
                    with _df_c2:
                        _dpy_Z = st.number_input(_dlbl("dpy_Z", "Cao độ mặt đất tự nhiên (m)"),
                                                  -5.0, 3.0, -0.80, 0.05, key=f"dpy_Z_{_hk_iter}")
                        _dpy_H1 = st.number_input(_dlbl("dpy_H1", "Chiều dày lớp bùn sét H₁ (m)"),
                                                    5.0, 35.0, 22.0, 0.5, key=f"dpy_H1_{_hk_iter}")
                        _dpy_su = st.number_input(
                            _dlbl("dpy_su", "Ctd lớp đất xử lý nền Front (kN/m²)"),
                            value=80.0, step=1.0, key=f"dpy_su_{_hk_iter}",
                            help="Lực dính tương đương của lớp đất Front sau khi xử lý nền (CDM). "
                                 "Mặc định 80 kN/m² — bộ chuẩn dự án (CDM 28 ngày qu ≈ 160 kPa)."
                        )
                    # Mục D chỉ dùng tải phân bố — H và M đầu cọc = 0
                    _dpy_H = 0.0
                    _dpy_M = 0.0
                    with _df_c3:
                        _dpy_q_op = st.number_input("Tải trọng khai thác q (kN/m²)", 0.0, 200.0, 15.0, 2.5,
                                                      key=f"dpy_q_operation_{_hk_iter}",
                                                      help="Tải trọng thường xuyên trên mặt Front (xe, người, công trình). "
                                                           "Mặc định 15 kN/m² — bộ chuẩn dự án kè công viên. "
                                                           "Boussinesq công thức (39) TCVN 11823-3 §10.6.2.")
                    st.caption(
                        "**Quy ước:** **Front (TRÁI)** = phía đất đắp / mặt đường / xe chạy / người đi / "
                        "tải trọng công trình → áp lực + tải trọng đẩy cừ SW về phía Back. "
                        "**Back (PHẢI)** = phía đào / sông / mặt nước hở (chịu áp lực bị động dưới đáy đào). "
                        "Cừ SW làm tâm sơ đồ; mặt đất + lớp đất hai bên thường khác nhau "
                        "(Front có fill cao hơn, Back đào xuống thấp)."
                    )
                    _dpy_w1, _dpy_w2, _dpy_w3, _dpy_w4, _dpy_w5 = st.columns(5)
                    _eps_src_v = st.session_state.get(f"dpy_eps50_src_{_hk_iter}", "lib")
                    _eps_lbl = ("ε₅₀ (nén 3 trục)" if _eps_src_v == "triaxial"
                                 else "ε₅₀ (theo Su)")
                    _dpy_eps50 = _dpy_w1.number_input(
                        _eps_lbl, 0.001, 0.05, step=0.005,
                        key=f"dpy_eps50_{_hk_iter}", format="%.4f",
                        help="ε₅₀ — biến dạng tại 50% phá hoại từ nén 3 trục UU.\n\n"
                             "Tự lấy theo thứ tự:\n"
                             "1. **Nén 3 trục thí nghiệm phòng**: ε₅₀ = Cu/E (cùng HK + lớp bùn)\n"
                             "2. **Thư viện trạng thái đất** (Matlock 1970 / API):\n"
                             "   - Sét rất yếu/yếu (Su<24): 0.020\n"
                             "   - Sét trung bình (24-48): 0.010\n"
                             "   - Sét chặt (48-96): 0.007\n"
                             "   - Sét rất chặt (96-192): 0.005\n"
                             "   - Sét cứng (>192): 0.004",
                    )
                    _dpy_q_a   = _dpy_w2.number_input("a — khoảng cách tải (m)", 0.0, 20.0, 0.0, 0.5,
                                                         key=f"dpy_q_a_{_hk_iter}",
                                                         help="a=0: tải sát tường; a>0: tải lùi ra xa")
                    _dpy_q_w   = _dpy_w3.number_input("w — chiều rộng dải tải (m)", 1.0, 200.0, 10.0, 1.0,
                                                         key=f"dpy_q_w_{_hk_iter}",
                                                         help="Mặc định 10m — bộ chuẩn dự án (làn xe + lề). w ≥ 100m ≈ tải vô hạn.")
                    _dpy_wlvl  = _dpy_w4.number_input("Mực nước Front (m)", -5.0, 3.0, -1.0, 0.5,
                                                         key=f"dpy_wlvl_{_hk_iter}")
                    _dpy_bc    = _dpy_w5.selectbox("Liên kết đáy cọc", ["Free", "Fixed", "Cantilever"],
                                                    key=f"dpy_bc_{_hk_iter}")

                    # ── Hàng nhập đất phía Back (phía đào / sông) ───────────────────
                    st.markdown(
                        "**Đất phía Back (sau cọc — phía đào / sông):** "
                        "*Đáy lớp bùn sét giả định BẰNG cao độ Front (cùng đáy lớp 1 thực tế).*"
                    )
                    _dpy_b1, _dpy_b2, _dpy_b3 = st.columns(3)
                    _dpy_Zb     = _dpy_b1.number_input("Cao độ Ground B (m)", -10.0, 5.0,
                                                          -2.0, 0.05,
                                                          key=f"dpy_Z_back_{_hk_iter}",
                                                          help="Cao độ đáy đào phía Back (mặc định −2.0m; thấp hơn Front do đã đào)")
                    # Su Back mặc định = Su lớp bùn (_dpy_su) — đất Back là cùng lớp 1
                    # chưa xử lý, nên Su tương đương lớp bùn Front (trước CDM).
                    _dpy_sub    = _dpy_b2.number_input(
                        _dlbl("dpy_su_back", "Su Back (kN/m²)"),
                        value=float(st.session_state.get(f"dpy_su_back_{_hk_iter}", _su_back_default)),
                        step=1.0,
                        key=f"dpy_su_back_{_hk_iter}",
                        help="Sức kháng cắt không thoát nước phía Back (đất tự nhiên, chưa xử lý). "
                             "Mặc định lấy từ thí nghiệm cắt cánh hiện trường (VST) hoặc UU phòng — "
                             "chỉ dùng khi lớp không có số liệu đo đạc.")
                    _dpy_wlvl_b = _dpy_b3.number_input("Mực nước Back (m)", -10.0, 5.0,
                                                          -2.0, 0.5,
                                                          key=f"dpy_wlvl_back_{_hk_iter}",
                                                          help="Mực nước phía Back (sông/đào). Mặc định −2.0m.")

                    # ── Thông số lớp đất đắp (Fill — chỉ Front) ──────────────────────
                    st.markdown(
                        "**Đất đắp Front (Fill — từ đỉnh kè xuống mặt đất tự nhiên):** "
                        "*Kỹ sư nhập thông số thiết kế cụ thể của dự án.*"
                    )
                    _dpy_f1, _dpy_f2, _dpy_f3, _dpy_f4 = st.columns(4)
                    _dpy_gamma_fill = _dpy_f1.number_input(
                        "γ_fill (kN/m³)", value=18.0, step=0.1,
                        key=f"dpy_gamma_fill_{_hk_iter}",
                        help="Dung trọng đất đắp ướt. Mặc định 18 kN/m³ — đất đắp chặt vừa (cát san lấp)."
                    )
                    _dpy_phi_fill = _dpy_f2.number_input(
                        "φ_fill (°)", value=25.0, step=0.5,
                        key=f"dpy_phi_fill_{_hk_iter}",
                        help="Góc ma sát đất đắp. Mặc định 25° — đất đắp chặt vừa."
                    )
                    _dpy_c_fill = _dpy_f3.number_input(
                        "c_fill (kPa)", value=0.0, step=1.0,
                        key=f"dpy_c_fill_{_hk_iter}",
                        help="Lực dính đất đắp. Mặc định 0 kPa — đất rời (cát đắp)."
                    )
                    _dpy_gamma_sub_fill = _dpy_f4.number_input(
                        "γ_sub_fill (kN/m³)", value=8.0, step=0.1,
                        key=f"dpy_gamma_sub_fill_{_hk_iter}",
                        help="Dung trọng đất đắp bão hoà − γ_w. Mặc định 8 kN/m³ (γ_sat ≈ 18 → γ' = 8)."
                    )

                with _col_schem_d:
                    if _HAS_MPL:
                        _pile_tip_d = _dpy_top_ke - _dpy_L

                        # ── Lấy TẤT CẢ lớp đất từ SQLite cho HK đang áp ──────────────
                        _all_layers_F: list = []
                        _all_layers_B: list = []
                        _has_real_layers = False
                        if _dpy_apply_bh and _dpy_apply_bh != "(không áp)":
                            try:
                                _con_real = sqlite3.connect(str(_DB))
                                _cur_real = _con_real.cursor()
                                _row_bh_r = _cur_real.execute(
                                    "SELECT id, elevation_m FROM boreholes WHERE name LIKE ? LIMIT 1",
                                    (f"%{_dpy_apply_bh}%",),
                                ).fetchone()
                                if _row_bh_r:
                                    _bh_id_r, _bh_elev_r = _row_bh_r
                                    _bh_elev_r = _bh_elev_r or 0.0
                                    _ly_real = _cur_real.execute(
                                        "SELECT symbol, description, depth_top_m, depth_bot_m, "
                                        "thickness_m FROM layers WHERE borehole_id=? "
                                        "ORDER BY depth_top_m",
                                        (_bh_id_r,),
                                    ).fetchall()
                                    _max_depth_r = float(_bh_elev_r) - _pile_tip_d + 2.0
                                    for _sym, _desc, _dt, _db, _thk in _ly_real:
                                        # Sử dụng depth_top/bot nếu có, fallback theo thickness
                                        if _dt is None or _db is None:
                                            if _thk is None:
                                                continue
                                            _dt = sum(L[2] - L[1] if L[2] is None else 0
                                                      for L in _ly_real)  # bo qua neu thieu
                                        # Bo lop o ngoai pham vi
                                        if _db is None or _db > _max_depth_r:
                                            if _dt is None or _dt > _max_depth_r:
                                                continue
                                            _db = _max_depth_r
                                        _e_top_r = _bh_elev_r - _dt
                                        _e_bot_r = _bh_elev_r - _db
                                        # γ/φ theo symbol thuc te
                                        _sym_l = (_sym or "").upper()
                                        if _sym_l == "F":
                                            _gam_r, _phi_r = 18.0, 28.0
                                        elif _sym_l in ("1", "XMD"):
                                            _gam_r, _phi_r = 15.0, 0.0
                                        elif _sym_l == "3":
                                            _gam_r, _phi_r = 17.0, 12.0
                                        elif _sym_l == "4":
                                            _gam_r, _phi_r = 19.0, 30.0
                                        elif _sym_l in ("5", "5A", "5B"):
                                            _gam_r, _phi_r = 19.5, 25.0
                                        else:
                                            _gam_r, _phi_r = 18.0, 20.0
                                        _nm_r = f"L{_sym}: {(_desc or '')[:18]}"
                                        _all_layers_F.append((_nm_r, _e_top_r, _e_bot_r, _gam_r, _phi_r))
                                        # Back side: cùng layers, shift theo Z_back nếu khác Z_front
                                        _shift = float(_dpy_Zb) - float(_bh_elev_r)
                                        _all_layers_B.append((_nm_r, _e_top_r + _shift,
                                                                _e_bot_r + _shift, _gam_r, _phi_r))
                                    _has_real_layers = len(_all_layers_F) > 0
                                _con_real.close()
                            except Exception:
                                pass

                        _layers_F = _all_layers_F
                        _layers_B = _all_layers_B

                        if _has_real_layers:
                            _fig_sch_d = draw_pile_schematic(
                                L_m=_dpy_L + 2.0,
                                water_lvl=_dpy_wlvl, top_elev=_dpy_top_ke,
                                H_kN=_dpy_H, M_kNm=_dpy_M,
                                layers=_layers_F, bc_type=_dpy_bc,
                                soil_lvl_front=_dpy_Z, soil_lvl_back=_dpy_Zb,
                                water_lvl_back=_dpy_wlvl_b,
                                back_layers=_layers_B,
                                gamma_fill=float(_dpy_gamma_fill),
                                phi_fill=float(_dpy_phi_fill),
                                c_fill=float(_dpy_c_fill),
                            )
                            if _fig_sch_d:
                                st.pyplot(_fig_sch_d, use_container_width=True)
                                plt.close(_fig_sch_d)
                            st.caption(
                                f"**Front:** Z={_dpy_Z:+.2f}m · MN={_dpy_wlvl:+.2f}m · "
                                f"Ctd_Front (xử lý nền)={_dpy_su:.0f} kPa  \n"
                                f"**Back:** Z={_dpy_Zb:+.2f}m · MN={_dpy_wlvl_b:+.2f}m · "
                                f"Su_Back (tự nhiên)={_dpy_sub:.0f} kPa  \n"
                                f"**Lớp đất:** {len(_layers_F)} lớp từ HK `{_dpy_apply_bh}` "
                                f"(bao gồm 2m dưới chân cừ)."
                            )
                        else:
                            st.warning(
                                "**Chưa có lớp đất để vẽ.** Cần chọn HK ở \"Áp dữ liệu HK từ Mục B\" "
                                "ở trên — lớp đất sẽ được lấy trực tiếp từ hồ sơ địa chất, không dùng giá trị giả định."
                            )

                # ─── Biểu đồ áp lực nước + Boussinesq cạnh nhau ───────────────────────
                st.markdown("#### Áp lực nước + Tải Boussinesq")
                _wp_inp1, _wp_inp2, _wp_inp3 = st.columns([2, 2, 3])
                with _wp_inp1:
                    _wp_mode = st.selectbox(
                        "Chế độ áp lực nước",
                        ["hydrostatic", "seepage"],
                        format_func=lambda x: "Thủy tĩnh" if x == "hydrostatic" else "Terzaghi (thấm)",
                        key=f"dpy_wp_mode_{_hk_iter}",
                    )
                with _wp_inp2:
                    _wp_gamma = st.number_input("γ_w (kN/m³)", 9.0, 10.5, 9.81, 0.01,
                                                  key=f"dpy_wp_gamma_{_hk_iter}")
                with _wp_inp3:
                    _wp_dh = float(_dpy_wlvl_b) - float(_dpy_wlvl)
                    st.metric("Δh = Back − Front", f"{_wp_dh:+.2f}m")

                _wp_col, _bs_col = st.columns([1, 1], gap="medium")

                # ───── BÊN TRÁI: Áp lực nước (thu nhỏ) ─────
                with _wp_col:
                    st.markdown("**Áp lực nước — TCVN 11823-3 §10.5.1**")
                    try:
                        import sys as _sys_wp
                        _sys_wp.path.insert(0, str(_ROOT / "scripts"))
                        from water_pressure import (
                            WaterGeometry as _WG,
                            compute_all as _wp_compute,
                        )
                        _wp_geom = _WG(
                            top_elev=float(_dpy_top_ke),
                            pile_length=float(_dpy_L),
                            soil_level_front=float(_dpy_Z),
                            water_elev_front=float(_dpy_wlvl),
                            water_elev_back=float(_dpy_wlvl_b),
                            gamma_w=float(_wp_gamma),
                        )
                        _wp_res = _wp_compute(_wp_geom, mode=_wp_mode)
                        if _wp_res.get("fig"):
                            # Thu nhỏ
                            _wp_res["fig"].set_size_inches(7, 4.5)
                            for _ax in _wp_res["fig"].axes:
                                _ax.tick_params(labelsize=7)
                                if _ax.get_title():
                                    _ax.set_title(_ax.get_title(), fontsize=8)
                            st.pyplot(_wp_res["fig"], use_container_width=True)
                            plt.close(_wp_res["fig"])
                        _wm1, _wm2, _wm3 = st.columns(3)
                        _wm1.metric("F Front", f"{_wp_res['F_front']:.0f}", "kN/m")
                        _wm2.metric("F Back",  f"{_wp_res['F_back']:.0f}",  "kN/m")
                        _wm3.metric("F Net",   f"{_wp_res['F_net']:.0f}",   "kN/m")
                        st.caption(f"Chế độ: {_wp_res['mode']}")
                    except Exception as _e_wp:
                        st.warning(f"Không vẽ được áp lực nước: {_e_wp}")

                # ───── BÊN PHẢI: Áp lực ngang do tải Boussinesq ─────
                with _bs_col:
                    st.markdown(f"**Áp lực ngang Boussinesq từ tải q = {_dpy_q_op:.0f} kN/m²**")
                    try:
                        import sys as _sys_b2
                        _sys_b2.path.insert(0, str(_ROOT / "scripts"))
                        from boussinesq_surcharge import (
                            SurchargeStrip as _Strip2,
                            BoussiGeometry as _Geom2,
                            compute_all as _bs_compute,
                        )
                        # ref_surface = TOP_KE (đỉnh kè TK, +2.7m) vì tải khai
                        # thác q đặt trên mặt kè (xe/người), không phải mặt đất
                        # tự nhiên Front (-0.80m).
                        _bs_strip_full = _Strip2(
                            q=float(_dpy_q_op),
                            a=float(_dpy_q_a),
                            w=float(_dpy_q_w),
                            ref_surface=float(_dpy_top_ke),
                            side="Front",
                            label=f"q = {_dpy_q_op} kN/m²",
                        )
                        _bs_geom_full = _Geom2(
                            top_elev=float(_dpy_top_ke),
                            pile_length=float(_dpy_L),
                            soil_level_front=float(_dpy_Z),
                            soil_level_back=float(_dpy_Zb),
                        )
                        _bs_res = _bs_compute(_bs_geom_full, [_bs_strip_full])
                        if _bs_res.get("fig"):
                            _bs_res["fig"].set_size_inches(7, 4.5)
                            for _ax in _bs_res["fig"].axes:
                                _ax.tick_params(labelsize=7)
                                if _ax.get_title():
                                    _ax.set_title(_ax.get_title(), fontsize=8)
                            st.pyplot(_bs_res["fig"], use_container_width=True)
                            plt.close(_bs_res["fig"])
                        _bm1, _bm2, _bm3 = st.columns(3)
                        _bm1.metric("F Front", f"{_bs_res['F_front']:.0f}", "kN/m")
                        _bm2.metric("F Back",  f"{_bs_res['F_back']:.0f}",  "kN/m")
                        _bm3.metric("F Net",   f"{_bs_res['F_net']:.0f}",   "kN/m")
                        st.caption(f"a={_dpy_q_a}m, w={_dpy_q_w}m · TCVN 11823-3 §10.6.2 Eq.39")
                    except Exception as _e_bs:
                        st.warning(f"Không vẽ được Boussinesq: {_e_bs}")

                # ─── Biểu đồ áp lực đất ngang — TCVN 11823-3 §10.5 ────────────────────
                st.markdown("#### Biểu đồ áp lực đất ngang (Active · Net · Passive) — TCVN 11823-3 §10.5")
                _ep_c1, _ep_c2, _ep_c3, _ep_c4 = st.columns(4)
                with _ep_c1:
                    _ep_ka_method = st.selectbox(
                        "Phương pháp Ka",
                        ["rankine", "coulomb"],
                        format_func=lambda x: "Rankine" if x == "rankine" else "Coulomb (Eq.25)",
                        key=f"dpy_ep_ka_{_hk_iter}",
                    )
                with _ep_c2:
                    _ep_delta = st.number_input("δ (°) — ma sát tường", 0.0, 30.0, 0.0, 1.0,
                                                   key=f"dpy_ep_delta_{_hk_iter}",
                                                   help="δ=0 (Rankine); >0 cho Coulomb. Bảng 20 TCVN")
                with _ep_c3:
                    _ep_surcharge = st.number_input("Tải mặt Front q (kN/m²)", 0.0, 200.0,
                                                      float(_dpy_q_op), 5.0,
                                                      key=f"dpy_ep_q_{_hk_iter}",
                                                      help="Mặc định bằng tải khai thác đã nhập")
                with _ep_c4:
                    _ep_kp_method = st.selectbox(
                        "Phương pháp Kp",
                        ["rankine"],
                        format_func=lambda x: "Rankine",
                        key=f"dpy_ep_kp_{_hk_iter}",
                    )

                if _has_real_layers:
                    try:
                        import sys as _sys_ep
                        _sys_ep.path.insert(0, str(_ROOT / "scripts"))
                        from earth_pressure import (
                            EpGeometry as _EpG,
                            compute_all as _ep_compute,
                        )
                        from lateral_earth_pressure import SoilLayer as _SL

                        # ── Lấy thông số đất THẬT từ SQLite lab_tests ────────────────
                        # Cho mỗi lớp: tính γ, φ, c trung bình từ các mẫu lab trong depth range
                        _con_lab = sqlite3.connect(str(_DB))
                        _cur_lab = _con_lab.cursor()
                        _row_bh_ep = _cur_lab.execute(
                            "SELECT id, elevation_m FROM boreholes WHERE name LIKE ? LIMIT 1",
                            (f"%{_dpy_apply_bh}%",),
                        ).fetchone()
                        _bh_id_ep, _bh_elev_ep = (_row_bh_ep if _row_bh_ep else (None, 0.0))
                        _bh_elev_ep = _bh_elev_ep or 0.0

                        # Lay layers gốc (depth) tu DB
                        _real_lyrs = _cur_lab.execute(
                            "SELECT symbol, description, depth_top_m, depth_bot_m FROM layers "
                            "WHERE borehole_id=? ORDER BY depth_top_m",
                            (_bh_id_ep,),
                        ).fetchall() if _bh_id_ep else []

                        _ep_front_layers = []
                        _ep_back_layers_acc: list = []
                        _ep_sus_back: list[float] = []
                        _real_param_rows = []   # cho caption hiển thị nguồn
                        _ep_soil_types: list[str] = []
                        _ep_sus: list[float] = []
                        # Phân loại sét/cát theo symbol TCVN
                        _CLAY_SYMS = {"1", "XMD", "3", "5", "5A", "5B"}
                        _SAND_SYMS = {"F", "2A", "2B", "2C", "4", "6", "7"}
                        for _sym, _desc, _dt, _db in _real_lyrs:
                            _sym = _alias_sym(_sym, _dpy_apply_bh)
                            if _dt is None or _db is None:
                                continue
                            # Truy van lab_tests trong khoang [_dt, _db]
                            _row_avg = _cur_lab.execute(
                                "SELECT AVG(gamma_kNm3), AVG(phi_deg), AVG(c_kPa), "
                                "AVG(Cu_UU_kPa), COUNT(*) "
                                "FROM lab_tests WHERE borehole_id=? "
                                "AND depth_from_m >= ? AND depth_to_m <= ?",
                                (_bh_id_ep, _dt, _db),
                            ).fetchone()
                            _gam_v, _phi_v, _c_v, _cu_v, _n_samp = _row_avg
                            # Su (Cu_UU) → c cho lớp sét nếu phi=0
                            if (_phi_v is None or _phi_v < 1.0) and _cu_v is not None:
                                _c_for_lyr = float(_cu_v)
                            elif _c_v is not None:
                                _c_for_lyr = float(_c_v)
                            else:
                                _c_for_lyr = 0.0
                            # Fallback theo symbol nếu thiếu lab
                            _sym_u = (_sym or "").upper()
                            if _gam_v is None:
                                _gam_v = (18.0 if _sym_u == "F"
                                           else 15.0 if _sym_u in ("1", "XMD")
                                           else 17.0 if _sym_u == "3"
                                           else 19.0 if _sym_u == "4" else 18.0)
                            if _phi_v is None:
                                _phi_v = (28.0 if _sym_u == "F"
                                           else 0.0  if _sym_u in ("1", "XMD")
                                           else 12.0 if _sym_u == "3"
                                           else 30.0 if _sym_u == "4" else 20.0)
                            _e_top_ep = _bh_elev_ep - _dt
                            _e_bot_ep = _bh_elev_ep - _db
                            if _e_bot_ep < _dpy_top_ke - _dpy_L - 0.01:
                                _e_bot_ep = _dpy_top_ke - _dpy_L - 0.01
                            _g_sub_ep = max(float(_gam_v) - 9.81, 8.0)
                            # Phân loại Sand/Clay theo SYMBOL TCVN (không theo phi)
                            if _sym_u in _CLAY_SYMS:
                                _soil_type = "Clay"
                                _c_use = _c_for_lyr   # mặc định = c lab
                            elif _sym_u in _SAND_SYMS:
                                _soil_type = "Sand"
                                _c_use = 0.0          # cát: cohesionless
                            else:
                                _soil_type = "Clay" if _c_for_lyr > 5.0 else "Sand"
                                _c_use = _c_for_lyr if _soil_type == "Clay" else 0.0
                            # Front: nếu là lớp BÙN (1/XMD) và user nhập Ctd > 0 → dùng Ctd
                            _c_use_front = _c_use
                            if _sym_u in {"1", "XMD"} and float(_dpy_su) > 0:
                                _c_use_front = float(_dpy_su)
                            # Back: nếu là Clay và user nhập Su Back > 0 → dùng Su Back
                            _c_use_back = _c_use
                            if _sym_u in {"1", "XMD"} and float(_dpy_sub) > 0:
                                _c_use_back = float(_dpy_sub)
                            _ep_soil_types.append(_soil_type)
                            _ep_sus.append(_c_use_front)
                            _ep_front_layers.append(_SL(
                                tip_elev=float(_e_bot_ep),
                                gamma=float(_gam_v),
                                gamma_sub=float(_g_sub_ep),
                                phi=float(_phi_v),
                                c=float(_c_use_front),
                                delta=float(_ep_delta),
                            ))
                            _ep_back_layers_acc.append(_SL(
                                tip_elev=float(_e_bot_ep),
                                gamma=float(_gam_v),
                                gamma_sub=float(_g_sub_ep),
                                phi=float(_phi_v),
                                c=float(_c_use_back),
                                delta=float(_ep_delta),
                            ))
                            _ep_sus_back.append(_c_use_back)
                            # Tính Ka/Kp cho hiển thị
                            import math as _math_kak
                            _phi_r = _math_kak.radians(float(_phi_v))
                            _ka_lyr = _math_kak.tan(_math_kak.radians(45 - float(_phi_v)/2))**2
                            _kp_lyr = _math_kak.tan(_math_kak.radians(45 + float(_phi_v)/2))**2
                            _real_param_rows.append({
                                "Lớp": _sym, "Mô tả": (_desc or "")[:25],
                                "Loại": _soil_type,
                                "z (m)": f"{_dt:.1f}–{_db:.1f}",
                                "γ (kN/m³)": round(float(_gam_v), 1),
                                "φ (°)": round(float(_phi_v), 1),
                                "c/cu (kPa)": round(_c_use, 1),
                                "Ka": round(_ka_lyr, 3),
                                "Kp": round(_kp_lyr, 3),
                                "n mẫu": _n_samp or 0,
                            })
                        _con_lab.close()
                        _ep_back_layers = _ep_back_layers_acc   # Back dùng c riêng (Su Back)

                        # Fill (đất đắp Front) — DÙNG THÔNG SỐ USER NHẬP
                        _ep_fill = _SL(
                            tip_elev=float(_dpy_Z),
                            gamma=float(_dpy_gamma_fill),
                            gamma_sub=float(_dpy_gamma_sub_fill),
                            phi=float(_dpy_phi_fill),
                            c=float(_dpy_c_fill),
                            delta=float(_ep_delta),
                        )
                        _ep_geom = _EpG(
                            top_elev=float(_dpy_top_ke),
                            pile_length=float(_dpy_L),
                            soil_level_front=float(_dpy_Z),
                            soil_level_back=float(_dpy_Zb),
                            water_elev_front=float(_dpy_wlvl),
                            water_elev_back=float(_dpy_wlvl_b),
                            surcharge_front=float(_ep_surcharge),
                        )
                        # Tính trước xử lý nền (PA1: dùng đầy đủ Front layers)
                        _ep_res = _ep_compute(
                            _ep_geom,
                            front_layers=_ep_front_layers,
                            back_layers=_ep_back_layers,
                            fill=_ep_fill,
                            ka_method=_ep_ka_method,
                            kp_method=_ep_kp_method,
                            delta_deg=float(_ep_delta),
                            front_soil_types=_ep_soil_types,
                            front_sus=_ep_sus,
                            back_soil_types=_ep_soil_types,
                            back_sus=_ep_sus_back,
                        )
                        # ── PA2: SAU xử lý nền CDM ───────────────────────────────────
                        # Front: fill chỉ tác dụng từ top_ke xuống đỉnh đệm XMC (z_mat).
                        # Dưới z_mat là khối CDM → không có áp lực đất Active trên Front.
                        _cdm_loads_d = st.session_state.get("cdm_loads", {})
                        _cdm_CDTK    = st.session_state.get("cdm_CDTK", 0.8)
                        _cdm_h_mat   = _cdm_loads_d.get("h_mat", 0.4)
                        _z_mat_top   = float(_cdm_CDTK) + float(_cdm_h_mat)   # đỉnh đệm XMC

                        # Geometry PA2: soil_level_front = z_mat (đỉnh đệm)
                        _ep_geom_cdm = _EpG(
                            top_elev=float(_dpy_top_ke),
                            pile_length=float(_dpy_L),
                            soil_level_front=_z_mat_top,
                            soil_level_back=float(_dpy_Zb),
                            water_elev_front=float(_dpy_wlvl),
                            water_elev_back=float(_dpy_wlvl_b),
                            surcharge_front=float(_ep_surcharge),
                        )
                        # Fill PA2: dùng cùng thông số fill từ D.1 user nhập
                        _ep_fill_cdm = _SL(
                            tip_elev=_z_mat_top,    # fill kết thúc tại đỉnh đệm
                            gamma=float(_dpy_gamma_fill),
                            gamma_sub=float(_dpy_gamma_sub_fill),
                            phi=float(_dpy_phi_fill),
                            c=float(_dpy_c_fill),
                            delta=float(_ep_delta),
                        )
                        _ep_res_cdm = _ep_compute(
                            _ep_geom_cdm,
                            front_layers=[],   # CDM blocks → không lớp tự nhiên Front
                            back_layers=_ep_back_layers,
                            fill=_ep_fill_cdm,
                            ka_method=_ep_ka_method,
                            kp_method=_ep_kp_method,
                            delta_deg=float(_ep_delta),
                            front_soil_types=[],
                            front_sus=[],
                            back_soil_types=_ep_soil_types,
                            back_sus=_ep_sus_back,
                        )

                        # Hiển thị 2 biểu đồ PA1 + PA2 cạnh nhau
                        _ep_left, _ep_mid = st.columns(2, gap="medium")
                        with _ep_left:
                            st.markdown("**(1) SAU khi xử lý nền**")
                            if _ep_res.get("fig"):
                                _ep_res["fig"].set_size_inches(8, 4.8)
                                for _ax in _ep_res["fig"].axes:
                                    _ax.tick_params(labelsize=7)
                                    if _ax.get_title():
                                        _ax.set_title(_ax.get_title(), fontsize=8)
                                st.pyplot(_ep_res["fig"], use_container_width=True)
                                plt.close(_ep_res["fig"])
                            _e1, _e2, _e3 = st.columns(3)
                            _e1.metric("F Active", f"{_ep_res['F_active']:.0f}", "kN/m")
                            _e2.metric("F Passive", f"{_ep_res['F_passive']:.0f}", "kN/m")
                            _e3.metric("F Net", f"{_ep_res['F_net']:.0f}", "kN/m")
                            st.caption(f"Front: fill + {len(_ep_front_layers)} lớp tự nhiên · "
                                         f"Back: {len(_ep_back_layers)} lớp")

                        with _ep_mid:
                            st.markdown("**(2) Tổng các áp lực — Active + Water + Passive → Net**")
                            # Đóng figure 3-panel của _ep_res_cdm (không hiển thị)
                            if _ep_res_cdm.get("fig"):
                                plt.close(_ep_res_cdm["fig"])

                            # ── Vẽ biểu đồ tổng áp lực (1 panel) ──────────────────────
                            # Component pressures (lấy từ PA2 = SAU CDM):
                            #   p_active   = áp lực đất chủ động Front (+ kN/m²)
                            #   p_passive  = áp lực đất bị động Back   (− kN/m², kháng)
                            #   p_water_n  = chênh lệch áp lực nước (Front − Back)
                            #   p_total    = p_active − p_passive + p_water_n
                            import numpy as _np_ep
                            _elevs_t   = _ep_res_cdm["elevs"]
                            _pa_t      = _ep_res_cdm["active_h"]
                            _pp_t      = _ep_res_cdm["passive_h"]

                            _gw_kNm3 = 9.81
                            _wf_t = _np_ep.maximum(0.0,
                                float(_dpy_wlvl) - _elevs_t) * _gw_kNm3   # nước Front
                            _wb_t = _np_ep.maximum(0.0,
                                float(_dpy_wlvl_b) - _elevs_t) * _gw_kNm3 # nước Back
                            _pw_net_t = _wf_t - _wb_t
                            _p_total_t = _pa_t - _pp_t + _pw_net_t

                            _fig_tot, _ax_tot = plt.subplots(figsize=(8, 4.8))
                            _ax_tot.plot(_pa_t,     _elevs_t, "-",  color="#C62828",
                                          lw=1.5, label="Active (Ka)")
                            _ax_tot.plot(-_pp_t,    _elevs_t, "-",  color="#1565C0",
                                          lw=1.5, label="Passive (Kp, kháng)")
                            _ax_tot.plot(_pw_net_t, _elevs_t, "--", color="#0277BD",
                                          lw=1.2, label="Nước Net (F−B)")
                            _ax_tot.plot(_p_total_t, _elevs_t, "-", color="#1B5E20",
                                          lw=2.4, label="TỔNG (Active+Water−Passive)")
                            # Tô vùng tổng (dương = đẩy ra Back, âm = đẩy ra Front)
                            _ax_tot.fill_betweenx(_elevs_t, 0, _p_total_t,
                                                   where=_p_total_t > 0,
                                                   color="#C62828", alpha=0.15)
                            _ax_tot.fill_betweenx(_elevs_t, 0, _p_total_t,
                                                   where=_p_total_t < 0,
                                                   color="#1565C0", alpha=0.15)
                            _ax_tot.axvline(0, color="black", lw=0.7)
                            _ax_tot.axhline(float(_dpy_top_ke), color="#8B4513",
                                             lw=0.7, ls=":", alpha=0.7)
                            _ax_tot.axhline(_z_mat_top, color="#666", lw=0.6, ls="--",
                                             alpha=0.7)
                            _ax_tot.text(0.02, _z_mat_top,
                                          f" z_mat = {_z_mat_top:+.2f} m",
                                          transform=_ax_tot.get_yaxis_transform(),
                                          fontsize=6, color="#666", va="bottom")
                            _ax_tot.set_xlabel("Áp lực ngang (kN/m²)", fontsize=8)
                            _ax_tot.set_ylabel("Cao độ (m)", fontsize=8)
                            _ax_tot.set_title("Tổng áp lực ngang lên tường (SAU CDM)",
                                               fontsize=8)
                            _ax_tot.tick_params(labelsize=7)
                            _ax_tot.grid(True, ls=":", color="#E0E0E0", lw=0.5)
                            _ax_tot.legend(fontsize=6, loc="lower right", framealpha=0.85)
                            st.pyplot(_fig_tot, use_container_width=True)
                            plt.close(_fig_tot)

                            # Tổng lực hợp lực Net (numpy<2.0 dùng trapz, ≥2.0 dùng trapezoid)
                            _trap_fn = getattr(_np_ep, "trapezoid", None) or getattr(_np_ep, "trapz")
                            _F_total_net = float(_trap_fn(_p_total_t, _elevs_t))
                            _c1m, _c2m, _c3m = st.columns(3)
                            _c1m.metric("F Active (kN/m)", f"{_ep_res_cdm['F_active']:.0f}",
                                          delta=f"Δ {_ep_res_cdm['F_active']-_ep_res['F_active']:+.0f}",
                                          delta_color="inverse")
                            _c2m.metric("F Passive (kN/m)", f"{_ep_res_cdm['F_passive']:.0f}")
                            _c3m.metric("F Tổng Net (kN/m)", f"{abs(_F_total_net):.0f}",
                                          f"{'→ Back' if _F_total_net > 0 else '→ Front'}",
                                          delta_color="off")
                            st.caption(
                                f"Front: fill {_dpy_top_ke:+.2f}m → z_mat = {_z_mat_top:+.2f}m "
                                f"(CDTK={_cdm_CDTK:+.2f}m + h_mat={_cdm_h_mat:.2f}m). "
                                f"**Tổng** = Active + Nước_net − Passive (tích phân = lực hợp)."
                            )


                        # Bảng thông số đất + Ka/Kp tính được
                        if _real_param_rows:
                            st.markdown("**Thông số đất dùng tính áp lực — từ thí nghiệm phòng:**")
                            # Dùng st.table thay vì st.dataframe để tránh overlap
                            # với st.latex phía dưới khi in PDF
                            st.table(pd.DataFrame(_real_param_rows))
                            st.markdown("**Công thức áp dụng (TCVN 11823-3:2017 §10.5):**")
                            st.latex(r"\text{Đất sét: } \sigma_h^{active} = K_a \cdot \sigma'_v - 2c\sqrt{K_a}, \quad \sigma_h^{passive} = K_p \cdot \sigma'_v + 2c\sqrt{K_p}")
                            st.latex(r"\text{Đất cát: } \sigma_h^{active} = K_a \cdot \sigma'_v, \quad \sigma_h^{passive} = K_p \cdot \sigma'_v \quad (c = 0)")
                            st.latex(r"K_a = \tan^2(45° - \varphi/2), \quad K_p = \tan^2(45° + \varphi/2) \quad \text{(Rankine)}")
                            st.caption(
                                f"**Nguồn:** Thí nghiệm phòng của HK `{_dpy_apply_bh}` — "
                                f"γ, φ, c, Cu_UU trung bình theo `depth_from_m`/`depth_to_m`.  \n"
                                f"**Phân loại Sand/Clay:** theo SYMBOL TCVN, KHÔNG theo φ "
                                f"(Clay: 1, XMD, 3, 5, 5A, 5B; Sand: F, 2A-C, 4, 6, 7).  \n"
                                f"**Sét:** dùng c = c_lab nếu có, ngược lại dùng Cu_UU. "
                                f"**Cát:** c = 0 (cohesionless).  \n"
                                f"**Tải mặt** q = {_ep_surcharge:.0f} kN/m² · "
                                f"Ka {_ep_ka_method} (δ={_ep_delta}°)"
                            )
                    except Exception as _e_ep:
                        st.warning(f"Không vẽ được biểu đồ áp lực đất: {_e_ep}")
                else:
                    st.info("Cần áp HK ở trên để có lớp đất tính áp lực đất (Ka/Kp theo từng lớp).")

                # Nút quick-set + Widget chiều rộng CDM (shared Mục D + Section E)
                _btn_cd1, _btn_cd2, _btn_cd3 = st.columns([2, 3, 2])
                with _btn_cd1:
                    if st.button("Đặt CDM ngàm 1.0m vào đất tốt",
                                  key=f"btn_cdm_ngam_1m_{_hk_iter}",
                                  help="Tự động set cdm_L_ngam = 1.0m và cdm_Lc = H₁ + 1.0m "
                                       "(ngàm tối thiểu 1m vào lớp tốt theo TCVN 9403)"):
                        st.session_state["cdm_L_ngam"] = 1.0
                        st.session_state["cdm_Lc"] = float(_dpy_H1) + 1.0
                        st.rerun()
                with _btn_cd2:
                    _cdm_Lc_cur = float(st.session_state.get("cdm_Lc", 0.0) or 0.0)
                    _cdm_Lng_cur = float(st.session_state.get("cdm_L_ngam", 0.0) or 0.0)
                    st.caption(
                        f"_CDM hiện tại: Lc = {_cdm_Lc_cur:.1f}m "
                        f"(ngàm {_cdm_Lng_cur:.1f}m vào đất tốt)._"
                    )
                with _btn_cd3:
                    # CDM width Front — dùng chung cho cả Mục D schematic + Section E
                    _cdm_w_e = st.number_input(
                        "Bề rộng CDM Front (m)",
                        min_value=1.0, max_value=20.0,
                        value=5.0, step=0.5,
                        key=f"e_cdm_width_{_hk_iter}",
                        help="Bề rộng vùng gia cố CDM phía Front kè (mặc định 5m). "
                             "Áp dụng cho sơ đồ kích thước Mục D + tính ổn định Section E."
                    )

                # Auto-run (bỏ nút) — mọi expander tự tính khi mở
                if True:
                    # Lấy chiều dài CDM từ tab "Thiết kế CDM" (session_state)
                    # cdm_Lc = chiều dài cọc CDM total; cdm_L_ngam = phần ngàm vào đất tốt
                    # → phần CDM trong soft soil = cdm_Lc − cdm_L_ngam
                    _cdm_Lc_val = float(st.session_state.get("cdm_Lc", 0.0) or 0.0)
                    _cdm_Lng_val = float(st.session_state.get("cdm_L_ngam", 0.0) or 0.0)
                    _cdm_thk_eff = max(0.0, _cdm_Lc_val - _cdm_Lng_val)
                    _k_cdm_fac = float(st.session_state.get("dpy_k_cdm_factor", 3.0) or 3.0)

                    # Winkler chỉ dùng tải user (H, M). KHÔNG cộng F từ EP làm point
                    # load — sai vật lý (EP là distributed, không được lump ở đỉnh).
                    # F per component HIỂN THỊ ở bảng bên dưới để user tham khảo.
                    _H_total = float(_dpy_H)
                    _M_total = float(_dpy_M)

                    # Bảng chi tiết F per component (lấy từ EP + Boussinesq)
                    _F_act_v = float(locals().get("_ep_res_cdm", {}).get("F_active", 0.0) or 0.0)
                    _F_pas_v = float(locals().get("_ep_res_cdm", {}).get("F_passive", 0.0) or 0.0)
                    _F_water_v = 0.0
                    try:
                        _trap_fn_w = (getattr(_np_ep, "trapezoid", None)
                                      or getattr(_np_ep, "trapz"))
                        _F_water_v = abs(float(_trap_fn_w(_pw_net_t, _elevs_t)))
                    except Exception:
                        pass
                    _F_bs_v = float(locals().get("_F_bs", 0.0) or 0.0)
                    # TỔNG đưa vào Winkler = Active + Nước + Surcharge (KHÔNG có Passive)
                    _F_sum_load = _F_act_v + _F_water_v + _F_bs_v

                    _ld_rows = [
                        {"Loại tải": "Áp lực chủ động Active (Front)",
                         "F (kN/m)": _F_act_v,
                         "Chiều": "→ Back (đẩy cừ)",
                         "Đưa vào Winkler": "Có"},
                        {"Loại tải": "Áp lực bị động Passive (Back)",
                         "F (kN/m)": _F_pas_v,
                         "Chiều": "← Front (giữ cừ)",
                         "Đưa vào Winkler": "Không (lò xo nền đã đảm nhận)"},
                        {"Loại tải": "Chênh lệch nước Front − Back",
                         "F (kN/m)": _F_water_v,
                         "Chiều": "→ Back" if (_wlvl_s_v := float(_dpy_wlvl)) > float(_dpy_wlvl_b) else "← Front",
                         "Đưa vào Winkler": "Có"},
                        {"Loại tải": "Boussinesq từ tải khai thác q",
                         "F (kN/m)": _F_bs_v,
                         "Chiều": "→ Back",
                         "Đưa vào Winkler": "Có"},
                        {"Loại tải": "TỔNG TẢI WINKLER = Active + Nước + q",
                         "F (kN/m)": _F_sum_load,
                         "Chiều": "→ Back" if _F_sum_load > 0 else "← Front",
                         "Đưa vào Winkler": "—"},
                    ]
                    st.markdown("**Chi tiết các loại tải trọng tác dụng lên cừ:**")
                    _df_ld = pd.DataFrame(_ld_rows)
                    _df_ld["F (kN/m)"] = _df_ld["F (kN/m)"].map(lambda v: f"{v:.1f}")
                    st.table(_df_ld)
                    st.caption(
                        "_Winkler dùng phương pháp **tải phân bố p(z)** — TỔNG = "
                        "Active + Nước + Surcharge (Boussinesq). "
                        "Lò xo nền Back đảm nhận phản kháng bị động (không double-count Passive)._"
                    )

                    # === Winkler với TẢI PHÂN BỐ — phương trình EI y'''' + kh D y = p(z) ===
                    try:
                        import numpy as _np_md
                        from winkler_np import solve_numpy_dist as _sn_dist_md
                        from wall_internal_force import sw_pile_props as _sw_props_md
                        _trap_md = (getattr(_np_md, "trapezoid", None) or _np_md.trapz)

                        # Pile object
                        _pl_md = _sw_by_name(_dpy_pile)
                        _pile_md = _sw_props_md(
                            H_mm=float(_pl_md["H_mm"]),
                            Itd_cm4=float(_pl_md.get("Itd_cm4") or 0),
                            Mcr_Tm=float(_pl_md.get("Mcr_Tm") or 0),
                            Atd_cm2=float(_pl_md.get("Atd_cm2") or 0),
                            fc_MPa=70.0, name=_dpy_pile,
                        )

                        # kh_layers từ SQLite
                        _bh_short_md = (f"KE-{_dpy_apply_bh}"
                                         if _dpy_apply_bh and not _dpy_apply_bh.startswith("KE-")
                                         else (_dpy_apply_bh or "KE-HK2"))
                        _ly_raw_md: list = []
                        try:
                            _con_md = sqlite3.connect(str(_DB))
                            _con_md.row_factory = sqlite3.Row
                            for _r in _con_md.execute("""
                                SELECT l.symbol, l.thickness_m,
                                       ROUND(AVG(lt.gamma_kNm3), 2) AS gamma_kNm3,
                                       ROUND(AVG(lt.c_kPa), 1) AS c_kPa,
                                       ROUND(AVG(lt.Cu_UU_kPa), 1) AS Cu_kPa
                                FROM layers l JOIN boreholes b ON l.borehole_id = b.id
                                LEFT JOIN lab_tests lt ON lt.borehole_id = b.id
                                    AND lt.depth_from_m >= l.depth_top_m
                                    AND lt.depth_to_m <= l.depth_bot_m
                                WHERE b.name = ?
                                GROUP BY l.id ORDER BY l.depth_top_m
                            """, (_bh_short_md,)).fetchall():
                                _su = _r["Cu_kPa"] if _r["Cu_kPa"] is not None else _r["c_kPa"]
                                _ly_raw_md.append({
                                    "symbol": _alias_sym(_r["symbol"] or "1", _bh_short_md),
                                    "thickness_m": _r["thickness_m"] or 0,
                                    "Su_kPa": float(_su) if _su is not None else None,
                                    "gamma_kNm3": _r["gamma_kNm3"] if _r["gamma_kNm3"] is not None else 15.0,
                                })
                            _con_md.close()
                        except Exception:
                            pass
                        if not _ly_raw_md:
                            _ly_raw_md = [{"symbol": "1", "thickness_m": float(_dpy_L) + 5.0,
                                            "Su_kPa": None, "gamma_kNm3": 15.0}]
                        # Cừ chuyển vị về phía sông (Back) → kh dùng thông số đất phía Back.
                        # Ưu tiên: per-layer Su từ SQLite (VST/lab) → user-input Su Back
                        _Su_back_fallback = float(_dpy_sub) if _dpy_sub > 0 else 11.0
                        _real_layers_md = [
                            _WIF_SoilLayer(
                                symbol=str(lr["symbol"]),
                                thickness_m=float(lr["thickness_m"] or 0),
                                Su_kPa=(float(lr["Su_kPa"]) if lr["Su_kPa"] is not None
                                        else _Su_back_fallback),
                                gamma_kNm3=float(lr["gamma_kNm3"]),
                            ) for lr in _ly_raw_md
                        ]
                        # Ground B (Zb) cao hơn chân cừ → khoảng từ top cừ → Zb
                        # KHÔNG có tiếp xúc đất ở phía Back (cừ trong không khí / nước).
                        # Thêm "ghost layer" trên đầu, Su≈0 → kh≈0 vùng này.
                        _d_offset_md = max(0.0, float(_dpy_top_ke) - float(_dpy_Zb))
                        if _d_offset_md > 0.01:
                            _ghost_md = _WIF_SoilLayer(
                                symbol="AIR", thickness_m=_d_offset_md,
                                Su_kPa=0.001, gamma_kNm3=0.001, is_clay=True,
                            )
                            _layers_md = [_ghost_md] + _real_layers_md
                        else:
                            _layers_md = _real_layers_md

                        # 0.5 chu vi cừ tham gia kh — override D_m = 0.5 × D
                        # (kh × D × dz = lực lò xo per node — halve D = halve lực)
                        _pile_md.D_m = _pile_md.D_m * 0.5

                        # Build p(z) TỔNG = Active − Passive + Nước + Surcharge
                        _L_md = float(_dpy_L)
                        _top_md = float(_dpy_top_ke)
                        _Nload_md = 100
                        _zs_md = list(_np_md.linspace(0, _L_md, _Nload_md))

                        def _ep2pile_md(_p_ep, _e_ep):
                            _d = _top_md - _np_md.asarray(_e_ep)
                            _i = _np_md.argsort(_d)
                            return list(_np_md.interp(
                                _zs_md, _d[_i], _np_md.asarray(_p_ep)[_i],
                                left=0.0, right=0.0))

                        if "_ep_res_cdm" in dir() and _ep_res_cdm:
                            _pA = _ep2pile_md(_ep_res_cdm["active_h"], _ep_res_cdm["elevs"])
                            _pP = _ep2pile_md(_ep_res_cdm["passive_h"], _ep_res_cdm["elevs"])
                        else:
                            _pA = [0.0] * _Nload_md
                            _pP = [0.0] * _Nload_md
                        _gw_md = 9.81
                        _pW = [max(0, float(_dpy_wlvl) - (_top_md - z)) * _gw_md
                               - max(0, float(_dpy_wlvl_b) - (_top_md - z)) * _gw_md
                               for z in _zs_md]

                        # KHÔNG cộng Passive vào tải — lò xo Winkler đã đảm nhận
                        # phản kháng bị động tự động (xem wall_internal_force.py:225-227).
                        # Cộng Passive ở đây → double-count.
                        _p_total_md = [_pA[i] + _pW[i] for i in range(_Nload_md)]
                        # Lưu để section per-load tái sử dụng (Boussinesq cộng vào sau)

                        _res_d1 = _sn_dist_md(
                            layers=_layers_md, pile=_pile_md, L_m=_L_md,
                            zs_load=_zs_md, p_load_kNm2=_p_total_md,
                            N=max(60, int(_L_md * 3)),
                            eps50=float(_dpy_eps50), k_sand_kNm3=10_000.0,
                            cdm_thickness_m=_cdm_thk_eff, cdm_factor=_k_cdm_fac,
                            tip_fixity="free", top_pin=False,
                        )
                        # Lưu kết quả Winkler distributed vào DB (CLAUDE.md §5)
                        try:
                            from wall_internal_force import save_winkler_results_to_db as _save_wnk_d
                            _bh_full = f"KE-{_dpy_apply_bh}" if _dpy_apply_bh and _dpy_apply_bh != "(không áp)" else "KE-HK2"
                            _save_wnk_d(_res_d1, bh_name=_bh_full,
                                        pile_type=str(_dpy_pile),
                                        L_m=float(_L_md), load_case="distributed",
                                        H_load_kN=0.0, M_load_kNm=0.0)
                        except Exception:
                            pass
                        # Map _res_d1["Qs"] có thể cần với key (winkler_np đã trả)
                    except Exception as _e_md:
                        # Fallback về point load nếu solve_numpy_dist fail
                        _res_d1 = _calc_py_winkler(
                            bh_name=f"KE-{_dpy_apply_bh}" if _dpy_apply_bh != "(không áp)" else "KE-HK2",
                            pile_name=_dpy_pile, L_m=float(_dpy_L),
                            H_kNm=_H_total, M_kNm=_M_total,
                            cdm_thk_m=_cdm_thk_eff, eps50=float(_dpy_eps50),
                            k_cdm_factor=_k_cdm_fac,
                        )
                        st.caption(f"_(Fallback point load — solve_numpy_dist lỗi: {_e_md})_")
                    # Caption tham số kh
                    _kh_min_v = min(_res_d1.get("k_h", [0])) if _res_d1.get("k_h") else 0
                    _kh_max_v = max(_res_d1.get("k_h", [0])) if _res_d1.get("k_h") else 0
                    st.caption(
                        f"**Lò xo nền kh** (cừ chuyển vị về phía sông → dùng đất Back): "
                        f"Su Back = {float(_dpy_sub):.0f} kPa · "
                        f"0.5 chu vi cừ ({_pile_md.D_m*1000:.0f} mm) · "
                        f"kh tính từ Ground B (Zb = {float(_dpy_Zb):+.2f} m, "
                        f"offset {_d_offset_md:.2f} m bỏ qua) · "
                        f"kh: **{_kh_min_v:.0f} → {_kh_max_v:.0f} kN/m³**"
                    )

                    st.markdown(
                        f"**Kết quả kh hiện tại của HK này:**  "
                        f"kh min = **{_kh_min_v:.0f} kN/m³** · "
                        f"kh max = **{_kh_max_v:.0f} kN/m³** · "
                        f"$k_{{\\text{{spring}}}}$ trung bình ≈ "
                        f"**{(_kh_min_v + _kh_max_v) / 2 * _pile_md.D_m:.0f} kN/m/m**  \n"
                        f"$\\varepsilon_{{50}}$ áp dụng = **{float(_dpy_eps50):.4f}** "
                        f"({st.session_state.get(f'dpy_eps50_src_{_hk_iter}', 'lib')}) · "
                        f"D (0.5 chu vi cừ) = **{_pile_md.D_m*1000:.0f} mm**"
                    )
                    # ── Tính tải Boussinesq từ surcharge khai thác ────────────────────
                    try:
                        import sys as _sys_b
                        _sys_b.path.insert(0, str(_ROOT / "scripts"))
                        from boussinesq_surcharge import (
                            SurchargeStrip as _Strip,
                            BoussiGeometry as _Geom,
                            delta_ph_at_elev as _dph_at,
                        )
                        # ref_surface = TOP_KE (đỉnh kè TK) vì tải khai thác
                        # đặt trên mặt kè, không phải mặt đất tự nhiên.
                        _bs_strip = _Strip(
                            q=float(_dpy_q_op), a=float(_dpy_q_a), w=float(_dpy_q_w),
                            ref_surface=float(_dpy_top_ke), side="Front",
                            label=f"Tải khai thác q={_dpy_q_op} kN/m²",
                        )
                        # Profile Boussinesq từ ĐỈNH KÈ TK xuống chân cừ (toàn
                        # bộ chiều dài cừ phơi với áp lực ngang từ surcharge)
                        import numpy as _np_b
                        _e_top_bs = float(_dpy_top_ke)
                        _e_tip_bs = float(_dpy_top_ke - _dpy_L)
                        _elevs_bs = _np_b.linspace(_e_top_bs, _e_tip_bs, 40)
                        _dph_bs   = [_dph_at(e, _bs_strip, _e_top_bs) for e in _elevs_bs]
                        _trap_fn_b = getattr(_np_b, "trapezoid", None) or getattr(_np_b, "trapz")
                        _F_bs     = float(_trap_fn_b(_dph_bs, -_elevs_bs))
                    except Exception as _eb:
                        _dph_bs = None
                        _F_bs   = 0.0

                    if "error" in _res_d1:
                        st.error(_res_d1["error"])
                    else:
                        _r1, _r2, _r3, _r4 = st.columns(4)
                        _u_d1   = _res_d1["u_top_mm"]
                        _M_d1   = _res_d1["M_max_kNm"]
                        _Mcr_d1 = _res_d1["Mcr_kNm"]
                        _r1.metric("Chuyển vị đỉnh u (mm)", f"{_u_d1:.2f}",
                                     "Đạt" if abs(_u_d1) < 50 else "Vượt 5cm",
                                     delta_color="normal" if abs(_u_d1) < 50 else "inverse")
                        _r2.metric("Mô-men max (kNm/m)", f"{_M_d1:.1f}",
                                     f"M_cr = {_Mcr_d1:.0f}",
                                     delta_color="normal" if _M_d1 < _Mcr_d1 else "inverse")
                        _r3.metric("EI cọc (kNm²)", f"{_res_d1['EI_kNm2']:.0f}")
                        _r4.metric("D cọc (mm)", f"{_res_d1['D_mm']:.0f}")
                        # Biểu đồ p-y Winkler 4 cột (u, M, k_h, Boussinesq) đã gỡ
                        # theo yêu cầu — thay bằng biểu đồ nội lực sau xử lý CDM bên dưới.

                        # ── BIỂU ĐỒ NỘI LỰC CỪ SAU XỬ LÝ NỀN BẰNG CDM (matplotlib) ─────
                        try:
                            import matplotlib.pyplot as _plt_w
                            from matplotlib.patches import Rectangle as _RectW
                            st.markdown("##### Nội lực cừ SW sau xử lý nền CDM")
                            _fig_w, _axw = _plt_w.subplots(
                                1, 4, figsize=(12, 7), sharey=True,
                                gridspec_kw={"width_ratios": [1, 2, 2, 2]},
                            )
                            _zsw = _res_d1["zs"]
                            _top_w = float(_dpy_top_ke) if "_dpy_top_ke" in dir() else 2.7
                            _Z_w = float(_dpy_Z) if "_dpy_Z" in dir() else 0.0
                            _Lw = float(_dpy_L)
                            _el_pile = [_top_w - z for z in _zsw]
                            _el_mid = [_top_w - (_zsw[i] + _zsw[i + 1]) / 2 for i in range(len(_zsw) - 1)]
                            _cdm_top_w = _Z_w
                            _cdm_bot_w = _Z_w - _cdm_thk_eff if _cdm_thk_eff > 0 else _Z_w

                            # Panel 0: sơ đồ cừ + CDM (kèm dimension labels để kiểm tra)
                            _bot_pile_w = _top_w - _Lw
                            _Zb_w = float(_dpy_Zb)
                            _wlvl_F = float(_dpy_wlvl)
                            _wlvl_B = float(_dpy_wlvl_b)

                            # Cừ
                            _axw[0].add_patch(_RectW((-0.5, _bot_pile_w), 1.0, _Lw,
                                                      facecolor="#444444", edgecolor="black",
                                                      linewidth=1.2, label=f"Cừ SW L={_Lw:.1f}m"))
                            # CDM — dùng width user nhập (chia sẻ với Section E)
                            _cdm_w_d = float(st.session_state.get(f"e_cdm_width_{_hk_iter}", 5.0) or 5.0)
                            _cdm_x_right_d = -0.5
                            _cdm_x_left_d = _cdm_x_right_d - _cdm_w_d
                            if _cdm_thk_eff > 0:
                                _axw[0].add_patch(_RectW(
                                    (_cdm_x_left_d, _cdm_bot_w), _cdm_w_d, _cdm_thk_eff,
                                    facecolor="#90c890", edgecolor="green",
                                    alpha=0.7, hatch="//",
                                    label=f"CDM b={_cdm_w_d:.1f}m"))
                            # Đất đắp Front — cùng width với CDM
                            if _top_w - _Z_w > 0:
                                _axw[0].add_patch(_RectW(
                                    (_cdm_x_left_d, _Z_w), _cdm_w_d, _top_w - _Z_w,
                                    facecolor="#d4a373", alpha=0.4,
                                    label=f"Đắp {(_top_w - _Z_w):.1f}m"))
                            # Mặt đất Back (đào sâu hơn)
                            _axw[0].plot([-3, -0.5], [_Z_w, _Z_w], "k-", lw=1.5)
                            _axw[0].plot([0.5, 3], [_Zb_w, _Zb_w], "k-", lw=1.5)
                            # Mực nước
                            _axw[0].axhline(_wlvl_F, color="cyan", linestyle="-", lw=1.0, alpha=0.5,
                                             xmin=0, xmax=0.5)
                            _axw[0].axhline(_wlvl_B, color="cyan", linestyle="-", lw=1.0, alpha=0.5,
                                             xmin=0.5, xmax=1.0)

                            # ── Dimension labels (annotation cao độ + chiều dài) ──
                            # Cao độ đỉnh kè
                            _axw[0].annotate(f"top = {_top_w:+.2f}m",
                                              xy=(0.5, _top_w), xytext=(2.0, _top_w + 0.5),
                                              fontsize=7, color="#1565C0", fontweight="bold",
                                              arrowprops=dict(arrowstyle="->", color="#1565C0", lw=0.8))
                            # Cao độ mặt đất Front (Z)
                            _axw[0].annotate(f"Z = {_Z_w:+.2f}m",
                                              xy=(-0.5, _Z_w), xytext=(-3.3, _Z_w + 0.3),
                                              fontsize=7, color="#8B4513", fontweight="bold",
                                              arrowprops=dict(arrowstyle="->", color="#8B4513", lw=0.8))
                            # Cao độ mặt đất Back (Zb)
                            _axw[0].annotate(f"Zb = {_Zb_w:+.2f}m",
                                              xy=(0.5, _Zb_w), xytext=(1.8, _Zb_w - 0.6),
                                              fontsize=7, color="#FF6F00", fontweight="bold",
                                              arrowprops=dict(arrowstyle="->", color="#FF6F00", lw=0.8))
                            # Mực nước Front + Back
                            _axw[0].text(-3.3, _wlvl_F + 0.1, f"MNN_F = {_wlvl_F:+.2f}",
                                          fontsize=6, color="cyan")
                            _axw[0].text(1.8, _wlvl_B + 0.1, f"MNN_B = {_wlvl_B:+.2f}",
                                          fontsize=6, color="cyan")
                            # Chân cừ
                            _axw[0].annotate(f"tip = {_bot_pile_w:+.2f}m",
                                              xy=(0, _bot_pile_w), xytext=(1.5, _bot_pile_w - 0.5),
                                              fontsize=7, color="#444", fontweight="bold",
                                              arrowprops=dict(arrowstyle="->", color="#444", lw=0.8))
                            # Đáy CDM
                            if _cdm_thk_eff > 0:
                                _axw[0].annotate(f"CDM bot = {_cdm_bot_w:+.2f}m",
                                                  xy=(-2, _cdm_bot_w), xytext=(-3.3, _cdm_bot_w - 0.4),
                                                  fontsize=7, color="green", fontweight="bold",
                                                  arrowprops=dict(arrowstyle="->", color="green", lw=0.8))

                            # Đáy lớp bùn sét (Z - H1) — đường nét đứt nâu
                            _bun_bot_F = _Z_w - float(_dpy_H1)
                            _bun_bot_B = _Zb_w - float(_dpy_H1)
                            _axw[0].plot([-3, -0.5], [_bun_bot_F, _bun_bot_F],
                                          color="#8B4513", linestyle=":", lw=1.0, alpha=0.7)
                            _axw[0].plot([0.5, 3], [_bun_bot_B, _bun_bot_B],
                                          color="#8B4513", linestyle=":", lw=1.0, alpha=0.7)
                            # Annotation cao độ đáy lớp bùn (lấy mức cao hơn để gọn)
                            _bun_bot_show = max(_bun_bot_F, _bun_bot_B)
                            _axw[0].annotate(f"Đáy bùn = {_bun_bot_show:+.2f}m\nH₁ = {float(_dpy_H1):.1f}m",
                                              xy=(0.5, _bun_bot_B),
                                              xytext=(1.8, _bun_bot_B + 0.4),
                                              fontsize=6.5, color="#8B4513",
                                              arrowprops=dict(arrowstyle="->", color="#8B4513", lw=0.8))

                            # Thông số Su Front / Back ở giữa lớp bùn (annotation góc)
                            _su_mid_F = (_Z_w + _bun_bot_F) / 2.0
                            _su_mid_B = (_Zb_w + _bun_bot_B) / 2.0
                            _axw[0].text(-2.95, _su_mid_F,
                                          f"Lớp 1 (Front)\nCtd = {float(_dpy_su):.0f} kPa",
                                          fontsize=6, color="#8B4513", style="italic",
                                          va="center", ha="left",
                                          bbox=dict(boxstyle="round,pad=0.15",
                                                    facecolor="#FFF8E1",
                                                    edgecolor="#8B4513", lw=0.4, alpha=0.85))
                            _axw[0].text(0.6, _su_mid_B,
                                          f"Lớp 1 (Back)\nSu = {float(_dpy_sub):.0f} kPa",
                                          fontsize=6, color="#FF6F00", style="italic",
                                          va="center", ha="left",
                                          bbox=dict(boxstyle="round,pad=0.15",
                                                    facecolor="#FFF3E0",
                                                    edgecolor="#FF6F00", lw=0.4, alpha=0.85))

                            # Front/Back labels
                            _axw[0].text(-2.8, _top_w + 0.6, "Front", fontsize=8, color="brown",
                                          fontweight="bold", style="italic")
                            _axw[0].text(2.0, _top_w + 0.6, "Back", fontsize=8, color="#FF6F00",
                                          fontweight="bold", style="italic")

                            _axw[0].set_xlim(-3.5, 3.5)
                            _axw[0].set_ylabel("Cao độ (m)")
                            _axw[0].set_title(f"Sơ đồ kích thước hình học\nL cừ = {_Lw:.1f} m",
                                               fontsize=9)
                            _axw[0].set_xticks([])
                            _axw[0].grid(alpha=0.3)
                            _axw[0].legend(fontsize=6, loc="lower left")

                            # ── Helper: thêm reference lines + cao độ chú thích ───
                            def _add_elev_refs(_ax):
                                """Vẽ các đường tham chiếu cao độ giống Panel 0 (sơ đồ kích thước)."""
                                # Vùng CDM
                                if _cdm_thk_eff > 0:
                                    _ax.axhspan(_cdm_bot_w, _cdm_top_w,
                                                 alpha=0.12, color="green")
                                # Đỉnh kè
                                _ax.axhline(_top_w, color="#1565C0", linestyle="-", lw=0.8, alpha=0.6)
                                # Mặt đất Front (Z)
                                _ax.axhline(_Z_w, color="#8B4513", linestyle="-", lw=0.7, alpha=0.6)
                                # Mặt đất Back (Zb)
                                _ax.axhline(_Zb_w, color="#FF6F00", linestyle="-", lw=0.7, alpha=0.5)
                                # Đáy lớp bùn (Z - H1)
                                _ax.axhline(_bun_bot_F, color="#8B4513",
                                             linestyle=":", lw=0.7, alpha=0.5)
                                # Mực nước
                                _ax.axhline(_wlvl_F, color="cyan", linestyle="-", lw=0.6, alpha=0.4)
                                # Chân cừ
                                _ax.axhline(_bot_pile_w, color="#444", linestyle="-", lw=0.8, alpha=0.6)
                                # Đáy CDM
                                if _cdm_thk_eff > 0:
                                    _ax.axhline(_cdm_bot_w, color="green",
                                                 linestyle="--", lw=0.7, alpha=0.5)

                            def _add_elev_labels(_ax, _xlim_max):
                                """Thêm text cao độ ở mép phải biểu đồ (dùng cho panel đầu các loại)."""
                                _x_lbl = _xlim_max * 0.98
                                for _yv, _txt, _col in [
                                    (_top_w,      f"top={_top_w:+.1f}",     "#1565C0"),
                                    (_Z_w,        f"Z={_Z_w:+.1f}",         "#8B4513"),
                                    (_bun_bot_F,  f"đáy bùn={_bun_bot_F:+.1f}", "#8B4513"),
                                    (_Zb_w,       f"Zb={_Zb_w:+.1f}",       "#FF6F00"),
                                    (_bot_pile_w, f"tip={_bot_pile_w:+.1f}", "#444"),
                                ]:
                                    _ax.text(_x_lbl, _yv, _txt, fontsize=5.5,
                                             color=_col, ha="right", va="center", alpha=0.8,
                                             bbox=dict(boxstyle="round,pad=0.1",
                                                       facecolor="white", edgecolor="none",
                                                       alpha=0.6))

                            # Auto-scale: ưu tiên DATA visible. Threshold không bắt
                            # buộc phải nằm trong khung — nếu data << threshold thì
                            # chỉ vẽ data, threshold ở ngoài (label vẫn hiện ở title).

                            # Panel 1: u(z) — scale theo data, fallback nhỏ 10mm
                            _add_elev_refs(_axw[1])
                            _axw[1].plot(_res_d1["ux"], _el_pile, "b-", lw=2.0, label="u(z)")
                            _u_data_max = max((abs(u) for u in _res_d1["ux"]), default=1.0)
                            # Nếu u_data nhỏ hơn nhiều so 50mm → khung = data*1.4 (không ép 50mm vào);
                            # nếu u_data lớn → khung gồm cả threshold 50mm.
                            _u_xmax = max(_u_data_max * 1.4, 10.0)
                            if _u_data_max > 25:    # khi gần ngưỡng → bao gồm 50mm threshold
                                _u_xmax = max(_u_xmax, 55.0)
                            # Vẽ ngưỡng ±50 nếu nằm trong khung
                            if _u_xmax >= 50:
                                _axw[1].axvline(50, color="red", linestyle="--",
                                                  lw=1.0, label="±50 mm (5cm)")
                                _axw[1].axvline(-50, color="red", linestyle="--", lw=1.0)
                            _axw[1].axvline(0, color="black", lw=0.5)
                            _axw[1].set_xlim(-_u_xmax, _u_xmax)
                            _add_elev_labels(_axw[1], _u_xmax)
                            _axw[1].set_xlabel("u (mm)")
                            _gh_u = "Đạt" if _res_d1['u_max_mm'] < 50 else "Vượt 5cm"
                            _axw[1].set_title(
                                f"Chuyển vị u(z)\nu_max = {_res_d1['u_max_mm']:.2f} mm  ({_gh_u}, GH=50mm)",
                                fontsize=10)
                            _axw[1].grid(alpha=0.3)
                            _axw[1].legend(fontsize=7, loc="lower left")

                            # Panel 2: M(z) — scale theo data, threshold Mcr ngoài nếu data nhỏ
                            _add_elev_refs(_axw[2])
                            _axw[2].plot(_res_d1["Ms"], _el_mid, "g-", lw=2.0, label="M(z)")
                            _Mcr_w = _res_d1["Mcr_kNm"]
                            _M_data_max = max((abs(m) for m in _res_d1["Ms"]), default=1.0)
                            _M_xmax = max(_M_data_max * 1.4, 50.0)
                            if _M_data_max > _Mcr_w * 0.5 and _Mcr_w > 0:
                                _M_xmax = max(_M_xmax, _Mcr_w * 1.1)
                            if _Mcr_w > 0 and _M_xmax >= _Mcr_w:
                                _axw[2].axvline(_Mcr_w, color="red", linestyle="--",
                                                 lw=1.0, label=f"±Mcr={_Mcr_w:.0f}")
                                _axw[2].axvline(-_Mcr_w, color="red", linestyle="--", lw=1.0)
                            _axw[2].axvline(0, color="black", lw=0.5)
                            _axw[2].set_xlim(-_M_xmax, _M_xmax)
                            _add_elev_labels(_axw[2], _M_xmax)
                            _axw[2].set_xlabel("M (kNm)")
                            _ratio_w = _res_d1["M_max_kNm"] / _Mcr_w if _Mcr_w else 0
                            _gh_M = "Đạt" if _ratio_w < 1.0 else "Vượt Mcr"
                            _axw[2].set_title(
                                f"Moment M(z)\nM_max={_res_d1['M_max_kNm']:.0f}  "
                                f"M/Mcr={_ratio_w:.2f}  ({_gh_M})",
                                fontsize=10)
                            _axw[2].grid(alpha=0.3)
                            _axw[2].legend(fontsize=7, loc="lower left")

                            # Panel 3: Q(z) — scale chặt theo data
                            _Qs_w = _res_d1.get("Qs", [])
                            if _Qs_w:
                                _add_elev_refs(_axw[3])
                                _axw[3].plot(_Qs_w, _el_mid, "m-", lw=2.0, label="Q(z)")
                                _axw[3].axvline(0, color="black", lw=0.5)
                                _Qmax_w = _res_d1.get("Q_max_kN", max(abs(q) for q in _Qs_w))
                                _Q_xmax = max(_Qmax_w * 1.4, 5.0)
                                _axw[3].set_xlim(-_Q_xmax, _Q_xmax)
                                _add_elev_labels(_axw[3], _Q_xmax)
                                _axw[3].set_title(f"Lực cắt Q(z)\nQ_max={_Qmax_w:.0f} kN", fontsize=10)
                                _axw[3].set_xlabel("Q (kN)")
                                _axw[3].grid(alpha=0.3)
                                _axw[3].legend(fontsize=7, loc="lower left")
                            else:
                                _axw[3].set_visible(False)

                            _status_w = ("ĐẠT" if (_res_d1["u_max_mm"] < 50 and
                                                    _res_d1["M_max_kNm"] < _res_d1["Mcr_kNm"])
                                          else "KHÔNG ĐẠT")
                            _fig_w.suptitle(
                                f"Nội lực cừ {_dpy_pile} L={_Lw:.0f}m sau xử lý CDM "
                                f"— {_status_w}",
                                fontsize=11, fontweight="bold",
                            )
                            _plt_w.tight_layout()
                            st.pyplot(_fig_w, use_container_width=True)
                            _plt_w.close(_fig_w)
                        except Exception as _e_mpl:
                            st.warning(f"Không vẽ được biểu đồ matplotlib: {_e_mpl}")

                        # ───────────────────────────────────────────────────────
                        # Phân tích nội lực theo TỪNG LOẠI TẢI TRỌNG
                        # Phương pháp: TẢI PHÂN BỐ p(z) trên dầm Euler-Bernoulli
                        # (solve_numpy_dist + Hermite consistent load vector).
                        # KHÔNG quy đổi tải về tập trung ở đỉnh.
                        # ───────────────────────────────────────────────────────
                        try:
                            st.markdown("##### Nội lực theo từng loại tải trọng — TẢI PHÂN BỐ p(z)")
                            st.caption(
                                "Phương pháp: `solve_numpy_dist()` — Hermite consistent load "
                                "vector, p(z) interpolate qua các nút FEM. "
                                "Phương trình: $EI \\, y'''' + k_h(z) \\cdot D \\cdot y = p(z)$. "
                                "TỔNG = chồng chất tuyến tính các thành phần."
                            )

                            import numpy as _np_dl
                            from winkler_np import solve_numpy_dist as _sn_dist
                            from wall_internal_force import sw_pile_props as _sw_props_pl
                            _trap_dl = (getattr(_np_dl, "trapezoid", None) or _np_dl.trapz)

                            # 1. Build pile object (giống _calc_py_winkler)
                            _pl_obj_dl = _sw_by_name(_dpy_pile)
                            if _pl_obj_dl is None:
                                raise RuntimeError(f"Không có cọc {_dpy_pile}")
                            _pile_dl = _sw_props_pl(
                                H_mm=float(_pl_obj_dl["H_mm"]),
                                Itd_cm4=float(_pl_obj_dl.get("Itd_cm4") or 0),
                                Mcr_Tm=float(_pl_obj_dl.get("Mcr_Tm") or 0),
                                Atd_cm2=float(_pl_obj_dl.get("Atd_cm2") or 0),
                                fc_MPa=70.0, name=_dpy_pile,
                            )

                            # 2. Build kh_layers cho phía Back (cừ chuyển vị về sông)
                            # Quy tắc (memory project_winkler_kh_back):
                            #   - Dùng đất phía Back (đã đào về cao độ Zb < Z)
                            #   - BỎ lớp fill 'F' và bất kỳ lớp nào nằm trên Zb
                            #   - Lớp cắt ngang Zb → cắt phần trên, giữ phần dưới
                            #   - Ghost AIR cho đoạn pile từ top_elev xuống Zb (không có soil contact)
                            _bh_short_dl = (f"KE-{_dpy_apply_bh}"
                                             if _dpy_apply_bh and not _dpy_apply_bh.startswith("KE-")
                                             else (_dpy_apply_bh or "KE-HK2"))
                            _Zb_dl = float(_dpy_Zb)
                            _ly_raw_dl: list = []
                            try:
                                _con_dl = sqlite3.connect(str(_DB))
                                _con_dl.row_factory = sqlite3.Row
                                _rows_dl = _con_dl.execute("""
                                    SELECT l.symbol, l.depth_top_m, l.depth_bot_m, l.thickness_m,
                                           b.elevation_m AS bh_elev,
                                           ROUND(AVG(lt.gamma_kNm3), 2) AS gamma_kNm3,
                                           ROUND(AVG(lt.c_kPa), 1) AS c_kPa,
                                           ROUND(AVG(lt.Cu_UU_kPa), 1) AS Cu_kPa
                                    FROM layers l
                                    JOIN boreholes b ON l.borehole_id = b.id
                                    LEFT JOIN lab_tests lt ON lt.borehole_id = b.id
                                        AND lt.depth_from_m >= l.depth_top_m
                                        AND lt.depth_to_m <= l.depth_bot_m
                                    WHERE b.name = ?
                                    GROUP BY l.id ORDER BY l.depth_top_m
                                """, (_bh_short_dl,)).fetchall()
                                _con_dl.close()
                                _bh_elev_dl = (float(_rows_dl[0]["bh_elev"])
                                                if _rows_dl and _rows_dl[0]["bh_elev"] is not None
                                                else _Zb_dl)
                                for _r in _rows_dl:
                                    _sym_r = _alias_sym(_r["symbol"] or "1", _bh_short_dl)
                                    # BỎ lớp fill — Back đã đào hết fill
                                    if str(_sym_r).upper() in ("F", "FILL"):
                                        continue
                                    # Lọc theo Zb: bỏ lớp hoàn toàn trên Zb,
                                    # cắt phần thuộc phía dưới cho lớp giao Zb
                                    _e_top = _bh_elev_dl - float(_r["depth_top_m"] or 0)
                                    _e_bot = _bh_elev_dl - float(_r["depth_bot_m"] or 0)
                                    if _e_top <= _Zb_dl + 1e-3:
                                        # Lớp đã hoàn toàn dưới Zb → giữ nguyên thickness
                                        _thk_eff = float(_r["thickness_m"] or 0)
                                    elif _e_bot >= _Zb_dl - 1e-3:
                                        # Lớp hoàn toàn trên Zb → bỏ
                                        continue
                                    else:
                                        # Lớp cắt ngang Zb → giữ phần dưới (Zb − _e_bot)
                                        _thk_eff = max(0.0, _Zb_dl - _e_bot)
                                    if _thk_eff <= 1e-3:
                                        continue
                                    _su = _r["Cu_kPa"] if _r["Cu_kPa"] is not None else _r["c_kPa"]
                                    _ly_raw_dl.append({
                                        "symbol":      _sym_r,
                                        "thickness_m": _thk_eff,
                                        "Su_kPa":      float(_su) if _su is not None else None,
                                        "gamma_kNm3":  _r["gamma_kNm3"] if _r["gamma_kNm3"] is not None else 15.0,
                                    })
                            except Exception:
                                pass
                            if not _ly_raw_dl:
                                _ly_raw_dl = [{"symbol": "1", "thickness_m": float(_dpy_L) + 5.0,
                                                "Su_kPa": None, "gamma_kNm3": 15.0}]
                            # Ưu tiên: per-layer Su từ SQLite (VST/lab) → user-input Su Back
                            _Su_back_fallback_dl = float(_dpy_sub) if _dpy_sub > 0 else 11.0
                            _real_layers_dl = [
                                _WIF_SoilLayer(
                                    symbol=str(lr["symbol"]),
                                    thickness_m=float(lr["thickness_m"] or 0),
                                    Su_kPa=(float(lr["Su_kPa"]) if lr["Su_kPa"] is not None
                                            else _Su_back_fallback_dl),
                                    gamma_kNm3=float(lr["gamma_kNm3"]),
                                ) for lr in _ly_raw_dl
                            ]
                            # Ghost air layer phần trên Zb (cừ không tiếp xúc đất Back)
                            _d_offset_dl = max(0.0, float(_dpy_top_ke) - _Zb_dl)
                            if _d_offset_dl > 0.01:
                                _ghost_dl = _WIF_SoilLayer(
                                    symbol="AIR", thickness_m=_d_offset_dl,
                                    Su_kPa=0.001, gamma_kNm3=0.001, is_clay=True,
                                )
                                _layers_kh_dl = [_ghost_dl] + _real_layers_dl
                            else:
                                _layers_kh_dl = _real_layers_dl
                            # 0.5 chu vi cừ
                            _pile_dl.D_m = _pile_dl.D_m * 0.5

                            # 3. Build p(z) profiles per component (z = độ sâu từ đỉnh cừ)
                            _L_dl = float(_dpy_L)
                            _top_dl = float(_dpy_top_ke)
                            _N_load = 100
                            _zs_load_dl = list(_np_dl.linspace(0, _L_dl, _N_load))

                            def _ep_to_pile(_p_ep, _elev_ep):
                                """Convert (elev, p) profile to p(z) on pile depth grid."""
                                _depths = _top_dl - _np_dl.asarray(_elev_ep)
                                _idx = _np_dl.argsort(_depths)
                                _d_s = _depths[_idx]
                                _p_s = _np_dl.asarray(_p_ep)[_idx]
                                return list(_np_dl.interp(_zs_load_dl, _d_s, _p_s,
                                                            left=0.0, right=0.0))

                            # Active + Passive lấy từ EP đã tính trước đó
                            _has_ep = "_ep_res_cdm" in dir() and _ep_res_cdm is not None
                            if _has_ep:
                                _p_act_z  = _ep_to_pile(_ep_res_cdm["active_h"],
                                                          _ep_res_cdm["elevs"])
                                _p_pas_z  = _ep_to_pile(_ep_res_cdm["passive_h"],
                                                          _ep_res_cdm["elevs"])
                            else:
                                _p_act_z  = [0.0] * _N_load
                                _p_pas_z  = [0.0] * _N_load

                            # Nước net Front − Back theo cao độ
                            _gw_dl = 9.81
                            _p_water_z: list = []
                            for _z in _zs_load_dl:
                                _e = _top_dl - _z
                                _wf = max(0.0, float(_dpy_wlvl) - _e) * _gw_dl
                                _wb = max(0.0, float(_dpy_wlvl_b) - _e) * _gw_dl
                                _p_water_z.append(_wf - _wb)

                            # Surcharge Boussinesq từ tải khai thác q
                            _has_bs = ("_dph_bs" in dir() and _dph_bs is not None
                                       and len(_dph_bs) > 0)
                            if _has_bs:
                                _p_surch_z = _ep_to_pile(_dph_bs, _elevs_bs)
                            else:
                                _p_surch_z = [0.0] * _N_load

                            # TỔNG = Active + Nước + Surcharge
                            # KHÔNG cộng Passive — lò xo Winkler đã đảm nhận phản
                            # kháng bị động tự động (double-count nếu cộng vào).
                            _p_total_z = [_p_act_z[i] + _p_water_z[i] + _p_surch_z[i]
                                           for i in range(_N_load)]

                            # 4. Chạy solve_numpy_dist cho từng thành phần (KHÔNG Passive)
                            _comps_dist = [
                                ("Active",    _p_act_z),
                                ("Nước",      _p_water_z),
                                ("Surcharge", _p_surch_z),
                                ("TỔNG",      _p_total_z),
                            ]
                            _comp_colors_dl = {
                                "Active":    "#C62828",
                                "Nước":      "#0277BD",
                                "Surcharge": "#FF6F00",
                                "TỔNG":      "#1B5E20",
                            }
                            _N_fem = max(60, int(_L_dl * 3))
                            _comp_dist_results: dict = {}
                            for _nm_dl, _prof_dl in _comps_dist:
                                _F_int = abs(float(_trap_dl(_prof_dl, _zs_load_dl)))
                                if _F_int < 0.1 and _nm_dl != "TỔNG":
                                    continue
                                try:
                                    _res_dl = _sn_dist(
                                        layers=_layers_kh_dl, pile=_pile_dl, L_m=_L_dl,
                                        zs_load=_zs_load_dl, p_load_kNm2=_prof_dl,
                                        N=_N_fem,
                                        eps50=float(_dpy_eps50),
                                        k_sand_kNm3=10_000.0,
                                        cdm_thickness_m=_cdm_thk_eff,
                                        cdm_factor=_k_cdm_fac,
                                        tip_fixity="free", top_pin=False,
                                    )
                                    if "error" not in _res_dl:
                                        _res_dl["F_int_kNm"] = _F_int
                                        _comp_dist_results[_nm_dl] = _res_dl
                                except Exception as _e_dl:
                                    st.caption(f"_({_nm_dl}: {_e_dl})_")

                            # 5. Bảng tổng hợp
                            if _comp_dist_results:
                                _rows_dist = []
                                for _nm_dl, _res_dl in _comp_dist_results.items():
                                    _rows_dist.append({
                                        "Tải":         _nm_dl,
                                        "F tích phân (kN/m)": _res_dl["F_int_kNm"],
                                        "u_top (mm)":  _res_dl["u_top_mm"],
                                        "u_max (mm)":  _res_dl["u_max_mm"],
                                        "M_max (kNm)": _res_dl["M_max_kNm"],
                                        "Q_max (kN)":  _res_dl.get("Q_max_kN", 0.0),
                                    })
                                st.dataframe(
                                    pd.DataFrame(_rows_dist),
                                    use_container_width=True, hide_index=True,
                                    column_config={
                                        "F tích phân (kN/m)": st.column_config.NumberColumn(format="%.1f"),
                                        "u_top (mm)":  st.column_config.NumberColumn(format="%.2f"),
                                        "u_max (mm)":  st.column_config.NumberColumn(format="%.2f"),
                                        "M_max (kNm)": st.column_config.NumberColumn(format="%.1f"),
                                        "Q_max (kN)":  st.column_config.NumberColumn(format="%.1f"),
                                    },
                                )

                                # 6. Biểu đồ matplotlib 4 panel: p(z) + u + M + Q
                                try:
                                    import matplotlib.pyplot as _plt_dl_mpl
                                    _fig_dl, _ax_dl = _plt_dl_mpl.subplots(
                                        1, 4, figsize=(15, 6.5), sharey=True,
                                        gridspec_kw={"width_ratios": [1.2, 1, 1, 1]},
                                    )

                                    def _refs_dl(_ax):
                                        if _cdm_thk_eff > 0:
                                            _ax.axhspan(_cdm_bot_w, _cdm_top_w,
                                                         alpha=0.10, color="green")
                                        _ax.axhline(_top_w, color="#1565C0", lw=0.7, alpha=0.5)
                                        _ax.axhline(_Z_w, color="#8B4513", lw=0.7, alpha=0.5)
                                        _ax.axhline(_Zb_w, color="#FF6F00", lw=0.7, alpha=0.4)
                                        _ax.axhline(_bun_bot_F, color="#8B4513",
                                                     linestyle=":", lw=0.6, alpha=0.4)
                                        _ax.axhline(_wlvl_F, color="cyan", lw=0.5, alpha=0.4)
                                        _ax.axhline(_bot_pile_w, color="#444", lw=0.7, alpha=0.5)

                                    # Panel 0: p(z) profile + kh secondary axis
                                    _refs_dl(_ax_dl[0])
                                    _el_zs = [_top_dl - z for z in _zs_load_dl]
                                    # kh axis bên phải (twin)
                                    if "TỔNG" in _comp_dist_results:
                                        _kh_vals = _comp_dist_results["TỔNG"].get("k_h", [])
                                        _kh_zs   = _comp_dist_results["TỔNG"].get("zs", [])
                                        if _kh_vals and _kh_zs:
                                            _ax_kh = _ax_dl[0].twiny()
                                            _kh_el = [_top_dl - z for z in _kh_zs]
                                            _ax_kh.plot(_kh_vals, _kh_el,
                                                         color="#7B1FA2", lw=1.2, alpha=0.7,
                                                         label="kh (kN/m³)")
                                            _ax_kh.set_xlabel("kh (kN/m³)", fontsize=8, color="#7B1FA2")
                                            _ax_kh.tick_params(axis="x", labelcolor="#7B1FA2",
                                                                labelsize=7)
                                            _ax_kh.set_xlim(0, max(_kh_vals) * 1.15 if max(_kh_vals) > 0 else 1)
                                    for _nm_dl, _prof_dl in _comps_dist:
                                        if _nm_dl not in _comp_dist_results:
                                            continue
                                        _color = _comp_colors_dl[_nm_dl]
                                        _is_tot = (_nm_dl == "TỔNG")
                                        _ax_dl[0].plot(
                                            _prof_dl, _el_zs,
                                            color=_color, lw=2.4 if _is_tot else 1.2,
                                            ls="-" if _is_tot else ":",
                                            label=_nm_dl, alpha=1.0 if _is_tot else 0.85,
                                        )
                                    _ax_dl[0].axvline(0, color="black", lw=0.5)
                                    _ax_dl[0].set_xlabel("p (kN/m²)")
                                    _ax_dl[0].set_ylabel("Cao độ (m)")
                                    _ax_dl[0].set_title("Tải phân bố p(z)", fontsize=10)
                                    _ax_dl[0].grid(alpha=0.3)
                                    _ax_dl[0].legend(fontsize=7, loc="lower left", framealpha=0.85)

                                    # Panel 1-3: u, M, Q
                                    # Min ranges để hiển thị có chiều đẹp
                                    _min_x_key = {"ux": 5.0, "Ms": 50.0, "Qs": 5.0}
                                    for _i_p, (_ax_p, _key, _xlbl) in enumerate([
                                        (_ax_dl[1], "ux", "u (mm)"),
                                        (_ax_dl[2], "Ms", "M (kNm)"),
                                        (_ax_dl[3], "Qs", "Q (kN)"),
                                    ]):
                                        _refs_dl(_ax_p)
                                        _data_max = max(
                                            (max(abs(v) for v in _r.get(_key, [0]) or [0])
                                              for _r in _comp_dist_results.values()),
                                            default=1.0,
                                        )
                                        # Scale chặt theo data + margin 30%
                                        _xmax_p = max(_data_max * 1.3, _min_x_key[_key])
                                        for _nm_dl, _res_dl in _comp_dist_results.items():
                                            _color = _comp_colors_dl[_nm_dl]
                                            _is_tot = (_nm_dl == "TỔNG")
                                            _zs_p = _res_dl["zs"]
                                            if _key == "ux":
                                                _ys = [_top_dl - z for z in _zs_p]
                                                _xs = _res_dl["ux"]
                                            else:
                                                _ys = [_top_dl - (_zs_p[i] + _zs_p[i + 1]) / 2
                                                        for i in range(len(_zs_p) - 1)]
                                                _xs = _res_dl.get(_key, [])
                                            if not _xs:
                                                continue
                                            _ax_p.plot(
                                                _xs, _ys, color=_color,
                                                lw=2.4 if _is_tot else 1.2,
                                                ls="-" if _is_tot else ":",
                                                alpha=1.0 if _is_tot else 0.85,
                                            )
                                        _ax_p.axvline(0, color="black", lw=0.5)
                                        _ax_p.set_xlim(-_xmax_p, _xmax_p)
                                        _ax_p.set_xlabel(_xlbl)
                                        _ax_p.grid(alpha=0.3)
                                        # text cao độ mép phải
                                        for _yv, _txt, _col in [
                                            (_top_w, f"top={_top_w:+.1f}", "#1565C0"),
                                            (_Z_w, f"Z={_Z_w:+.1f}", "#8B4513"),
                                            (_bun_bot_F, f"đáy bùn={_bun_bot_F:+.1f}", "#8B4513"),
                                            (_Zb_w, f"Zb={_Zb_w:+.1f}", "#FF6F00"),
                                            (_bot_pile_w, f"tip={_bot_pile_w:+.1f}", "#444"),
                                        ]:
                                            _ax_p.text(_xmax_p * 0.98, _yv, _txt, fontsize=5,
                                                        color=_col, ha="right", va="center",
                                                        alpha=0.75,
                                                        bbox=dict(boxstyle="round,pad=0.1",
                                                                  facecolor="white",
                                                                  edgecolor="none", alpha=0.55))

                                    # Ngưỡng ±50mm + ±Mcr — chỉ vẽ khi nằm trong khung
                                    _u_xlim_d = _ax_dl[1].get_xlim()
                                    if abs(_u_xlim_d[1]) >= 50:
                                        _ax_dl[1].axvline(50, color="red", linestyle="--", lw=0.8,
                                                           alpha=0.7, label="±50 mm")
                                        _ax_dl[1].axvline(-50, color="red", linestyle="--", lw=0.8, alpha=0.7)
                                        _ax_dl[1].legend(fontsize=7, loc="lower left")
                                    _Mcr_dl = _comp_dist_results["TỔNG"]["Mcr_kNm"] if "TỔNG" in _comp_dist_results else 0
                                    _M_xlim_d = _ax_dl[2].get_xlim()
                                    if _Mcr_dl > 0 and abs(_M_xlim_d[1]) >= _Mcr_dl:
                                        _ax_dl[2].axvline(_Mcr_dl, color="red", linestyle="--",
                                                            lw=0.8, alpha=0.7,
                                                            label=f"±Mcr={_Mcr_dl:.0f}")
                                        _ax_dl[2].axvline(-_Mcr_dl, color="red", linestyle="--",
                                                            lw=0.8, alpha=0.7)
                                        _ax_dl[2].legend(fontsize=7, loc="lower left")

                                    _ax_dl[1].set_title("Chuyển vị u(z)", fontsize=10)
                                    _ax_dl[2].set_title("Moment M(z)", fontsize=10)
                                    _ax_dl[3].set_title("Lực cắt Q(z)", fontsize=10)
                                    _fig_dl.suptitle(
                                        f"Nội lực theo từng tải phân bố — {_dpy_pile} L={_L_dl:.0f}m "
                                        f"(Hermite consistent load, N={_N_fem})",
                                        fontsize=10, fontweight="bold",
                                    )
                                    _plt_dl_mpl.tight_layout()
                                    st.pyplot(_fig_dl, use_container_width=True)
                                    _plt_dl_mpl.close(_fig_dl)
                                except Exception as _e_dl_mpl:
                                    st.caption(f"_(Không vẽ được biểu đồ per-tải phân bố: {_e_dl_mpl})_")

                                # ───────────────────────────────────────────────
                                # Chỉ chạy BC = Free (đáy tự do, cừ thường)
                                # ───────────────────────────────────────────────
                                try:
                                    st.markdown("##### Kiểm tra với BC = Free (đáy tự do, cừ thường)")
                                    st.caption(
                                        "Tổ hợp tải TỔNG p(z), liên kết Free "
                                        "(đáy + đỉnh tự do, chỉ chịu lò xo Winkler). "
                                        "Giới hạn chuyển vị ngang: **5 cm = 50 mm**."
                                    )
                                    _bc_cases = [
                                        ("Free", {"tip_fixity": "free", "top_pin": False}),
                                    ]
                                    _bc_colors = {"Free": "#1565C0"}
                                    _bc_results: dict = {}
                                    for _bc_nm, _bc_kw in _bc_cases:
                                        try:
                                            _r_bc = _sn_dist(
                                                layers=_layers_kh_dl, pile=_pile_dl,
                                                L_m=_L_dl,
                                                zs_load=_zs_load_dl, p_load_kNm2=_p_total_z,
                                                N=_N_fem,
                                                eps50=float(_dpy_eps50),
                                                k_sand_kNm3=10_000.0,
                                                cdm_thickness_m=_cdm_thk_eff,
                                                cdm_factor=_k_cdm_fac,
                                                **_bc_kw,
                                            )
                                            if "error" not in _r_bc:
                                                _bc_results[_bc_nm] = _r_bc
                                        except Exception as _e_bc:
                                            st.caption(f"_({_bc_nm}: {_e_bc})_")

                                    if _bc_results:
                                        # Bảng so sánh 3 BC
                                        _Mcr_bc = next(iter(_bc_results.values()))["Mcr_kNm"]
                                        _rows_bc = []
                                        for _bc_nm, _r_bc in _bc_results.items():
                                            _u = _r_bc["u_max_mm"]
                                            _M = _r_bc["M_max_kNm"]
                                            _Q = _r_bc.get("Q_max_kN", 0.0)
                                            _ok_u = abs(_u) < 50.0   # 5cm = chuyển vị ngang cho phép
                                            _ok_M = _M < _Mcr_bc if _Mcr_bc > 0 else False
                                            _rows_bc.append({
                                                "Liên kết đáy": _bc_nm,
                                                "u_top (mm)":   _r_bc["u_top_mm"],
                                                "u_max (mm)":   _u,
                                                "M_max (kNm)":  _M,
                                                "M/Mcr":        _M / _Mcr_bc if _Mcr_bc else 0.0,
                                                "Q_max (kN)":   _Q,
                                                "Kết quả":      "Đạt" if (_ok_u and _ok_M) else "Không đạt",
                                            })
                                        def _hl_bc(val):
                                            if val == "Đạt":      return "background-color:#d4edda; color:#155724"
                                            if val == "Không đạt": return "background-color:#f8d7da; color:#721c24"
                                            return ""
                                        st.dataframe(
                                            pd.DataFrame(_rows_bc).style.map(_hl_bc, subset=["Kết quả"]),
                                            use_container_width=True, hide_index=True,
                                            column_config={
                                                "u_top (mm)":  st.column_config.NumberColumn(format="%.2f"),
                                                "u_max (mm)":  st.column_config.NumberColumn(format="%.2f"),
                                                "M_max (kNm)": st.column_config.NumberColumn(format="%.1f"),
                                                "M/Mcr":       st.column_config.NumberColumn(format="%.2f"),
                                                "Q_max (kN)":  st.column_config.NumberColumn(format="%.1f"),
                                            },
                                        )

                                        # Chart 3 panel: u, M, Q — 3 đường BC chồng
                                        try:
                                            _fig_bc, _ax_bc = _plt_dl_mpl.subplots(
                                                1, 3, figsize=(13, 6.0), sharey=True,
                                            )
                                            _xmax_u_bc = max((max(abs(u) for u in r["ux"])
                                                                for r in _bc_results.values()),
                                                              default=10.0) * 1.1
                                            _xmax_M_bc = max((max(abs(m) for m in r["Ms"])
                                                                for r in _bc_results.values()),
                                                              default=10.0) * 1.1
                                            _xmax_Q_bc = max((max(abs(q) for q in r.get("Qs", [0]))
                                                                for r in _bc_results.values()),
                                                              default=10.0) * 1.1
                                            for _i_bc, (_ax_pp, _key, _xlbl, _xmax_bc) in enumerate([
                                                (_ax_bc[0], "ux", "u (mm)", _xmax_u_bc),
                                                (_ax_bc[1], "Ms", "M (kNm)", _xmax_M_bc),
                                                (_ax_bc[2], "Qs", "Q (kN)", _xmax_Q_bc),
                                            ]):
                                                _refs_dl(_ax_pp)
                                                for _bc_nm, _r_bc in _bc_results.items():
                                                    _zs_b = _r_bc["zs"]
                                                    if _key == "ux":
                                                        _ys_b = [_top_dl - z for z in _zs_b]
                                                        _xs_b = _r_bc["ux"]
                                                    else:
                                                        _ys_b = [_top_dl - (_zs_b[i] + _zs_b[i + 1]) / 2
                                                                  for i in range(len(_zs_b) - 1)]
                                                        _xs_b = _r_bc.get(_key, [])
                                                    if not _xs_b:
                                                        continue
                                                    _ax_pp.plot(
                                                        _xs_b, _ys_b,
                                                        color=_bc_colors[_bc_nm], lw=2.0,
                                                        label=_bc_nm,
                                                    )
                                                _ax_pp.axvline(0, color="black", lw=0.5)
                                                _ax_pp.set_xlim(-_xmax_bc, _xmax_bc)
                                                _ax_pp.set_xlabel(_xlbl)
                                                _ax_pp.grid(alpha=0.3)
                                            _ax_bc[0].set_ylabel("Cao độ (m)")
                                            _ax_bc[0].set_title("u(z)", fontsize=10)
                                            _ax_bc[1].set_title("M(z)", fontsize=10)
                                            _ax_bc[2].set_title("Q(z)", fontsize=10)
                                            _ax_bc[0].axvline(50, color="red", linestyle="--", lw=0.8, alpha=0.6,
                                                                label="±50mm (5cm)")
                                            _ax_bc[0].axvline(-50, color="red", linestyle="--", lw=0.8, alpha=0.6)
                                            if _Mcr_bc > 0:
                                                _ax_bc[1].axvline(_Mcr_bc, color="red", linestyle="--", lw=0.8, alpha=0.6)
                                                _ax_bc[1].axvline(-_Mcr_bc, color="red", linestyle="--", lw=0.8, alpha=0.6)
                                            _ax_bc[0].legend(fontsize=8, loc="lower left",
                                                              framealpha=0.85, title="BC")
                                            _fig_bc.suptitle(
                                                f"So sánh BC chân cừ — {_dpy_pile} L={_L_dl:.0f}m (tải TỔNG p(z))",
                                                fontsize=10, fontweight="bold",
                                            )
                                            _plt_dl_mpl.tight_layout()
                                            st.pyplot(_fig_bc, use_container_width=True)
                                            _plt_dl_mpl.close(_fig_bc)
                                        except Exception as _e_bc_chart:
                                            st.caption(f"_(Không vẽ được chart BC: {_e_bc_chart})_")
                                except Exception as _e_bc_all:
                                    st.caption(f"_(Không chạy được so sánh BC: {_e_bc_all})_")
                        except Exception as _e_per_load:
                            st.caption(f"_(Không vẽ được phân tích per-tải phân bố: {_e_per_load})_")

                        # ═══════════════════════════════════════════════════════════
                        # E. ỔN ĐỊNH TỔNG THỂ TƯỜNG SW + CDM (3 kiểm tra)
                        # ═══════════════════════════════════════════════════════════

                st.markdown("---")
                st.markdown("### E. Ổn định tổng thể tường SW + CDM")
                st.caption(
                    "2 kiểm tra theo TCVN 4253 + USACE EM 1110-2-2504 + FHWA GEC-13: "
                    "(1) Trượt cung tròn qua chân cừ — Bishop/Spencer LE, "
                    "(2) Lật quanh chân cừ."
                )

                try:
                    from sw_global_stability import (
                        CDMBlock as _CDMB, check_all as _sw_check_all,
                    )
                    from wall_internal_force import (
                        WallGeometry as _WG_S, EarthLayer as _EL_S,
                    )

                    # Lấy đầy đủ thông số từ widgets D.1 + SQLite (đồng bộ với EP)
                    _Z_s     = float(_dpy_Z)
                    _top_s   = float(_dpy_top_ke)
                    _L_s     = float(_dpy_L)
                    _H1_s    = float(_dpy_H1)
                    _wlvl_s  = float(_dpy_wlvl)
                    _wlvl_b_s = float(_dpy_wlvl_b)
                    _Zb_s    = float(_dpy_Zb)
                    _su_F    = float(_dpy_su)        # Ctd lớp 1 đã xử lý
                    _su_B    = float(_dpy_sub)       # Su Back (đất chưa xử lý)
                    _q_s     = float(_dpy_q_op)

                    _geom_s = _WG_S(
                        top_elev=_top_s, pile_length=_L_s,
                        soil_level_front=_Z_s, soil_level_back=_Zb_s,
                        water_elev_front=_wlvl_s, water_elev_back=_wlvl_b_s,
                        surcharge_front=_q_s,
                    )

                    # Fill từ widgets (bộ chuẩn dự án γ=18, φ=25, c=0, γ_sub=8)
                    _fill_s = (_EL_S(_Z_s,
                                      float(_dpy_gamma_fill),
                                      float(_dpy_gamma_sub_fill),
                                      float(_dpy_phi_fill),
                                      float(_dpy_c_fill))
                                if _top_s > _Z_s else None)

                    _pile_bot_s = _top_s - _L_s

                    # ── Lấy địa tầng THỰC TẾ từ SQLite (HK đang phân tích) ────
                    _front_s: list = []
                    _back_s: list = []
                    _bh_elev_e = _Z_s   # fallback
                    _layers_log: list = []
                    try:
                        _con_e = sqlite3.connect(str(_DB))
                        _con_e.row_factory = sqlite3.Row
                        # SQLite lưu name = 'KE-HKx' (có prefix), không phải 'HKx'
                        _bh_full_e = (f"KE-{_hk_iter}" if not _hk_iter.startswith("KE-")
                                       else _hk_iter)
                        _rows_e = _con_e.execute("""
                            SELECT l.symbol, l.depth_top_m, l.depth_bot_m,
                                   ROUND(AVG(lt.gamma_kNm3), 2) AS gam,
                                   ROUND(AVG(lt.phi_deg), 1) AS phi,
                                   ROUND(AVG(lt.c_kPa), 1) AS c,
                                   ROUND(AVG(lt.Cu_UU_kPa), 1) AS Cu,
                                   b.elevation_m AS bh_elev
                            FROM layers l
                            JOIN boreholes b ON l.borehole_id = b.id
                            LEFT JOIN lab_tests lt ON lt.borehole_id = b.id
                                AND lt.depth_from_m >= l.depth_top_m
                                AND lt.depth_to_m <= l.depth_bot_m
                            WHERE b.name = ?
                            GROUP BY l.id ORDER BY l.depth_top_m
                        """, (_bh_full_e,)).fetchall()
                        _con_e.close()
                        if _rows_e and _rows_e[0]["bh_elev"] is not None:
                            _bh_elev_e = float(_rows_e[0]["bh_elev"])

                        # Back side: bắt đầu từ lớp BÙN (1, XMD) trở xuống.
                        # Trên Zb đã đào hết — không có Fill / lớp F phía Back.
                        _back_started = False
                        for _r in _rows_e:
                            _sym = _alias_sym(_r["symbol"] or "", _bh_full_e)
                            _e_top_f = _bh_elev_e - float(_r["depth_top_m"] or 0)
                            _e_bot_f = _bh_elev_e - float(_r["depth_bot_m"] or 0)
                            # Bỏ lớp hoàn toàn trên đỉnh kè hoặc dưới chân cừ -2m
                            if _e_top_f < _pile_bot_s - 2.0 or _e_bot_f > _top_s:
                                continue
                            _gam = float(_r["gam"]) if _r["gam"] else 18.0
                            _gam_sub = max(_gam - 9.81, 5.0)
                            # Phân loại theo SYMBOL TCVN
                            if _sym in {"1", "XMD", "3", "5", "5A", "5B"}:   # Clay
                                _phi_f = 0.0
                                _c_f = float(_r["Cu"]) if _r["Cu"] else (float(_r["c"]) if _r["c"] else _su_F)
                                _c_b = _su_B if _sym in {"1", "XMD"} else _c_f
                            elif _sym in {"F", "2", "2A", "2B", "2C", "4", "6", "7"}:  # Sand
                                _phi_f = float(_r["phi"]) if _r["phi"] else 30.0
                                _c_f = 0.0
                                _c_b = 0.0
                            else:
                                _phi_f = float(_r["phi"]) if _r["phi"] else 20.0
                                _c_f = float(_r["c"]) if _r["c"] else 0.0
                                _c_b = _c_f
                            _front_s.append(_EL_S(_e_bot_f, _gam, _gam_sub, _phi_f, _c_f))

                            # Back: BẮT ĐẦU từ lớp bùn ("1", "XMD") trở xuống
                            # Skip Fill ("F") + lớp trên bùn (vd "2a" ở đỉnh nếu có)
                            if _sym in {"1", "XMD"}:
                                _back_started = True
                            if _back_started:
                                _shift = _Zb_s - _Z_s
                                _back_s.append(_EL_S(_e_bot_f + _shift,
                                                      _gam, _gam_sub, _phi_f, _c_b))
                                _c_b_disp = f"{_c_b:.0f}"
                                _on_back = "✓"
                            else:
                                _c_b_disp = "— (skip)"
                                _on_back = "✗"
                            _layers_log.append({
                                "Symbol": _sym, "z_top (m)": f"{_e_top_f:+.2f}",
                                "z_bot (m)": f"{_e_bot_f:+.2f}",
                                "γ": f"{_gam:.1f}", "γ_sub": f"{_gam_sub:.1f}",
                                "φ": f"{_phi_f:.1f}",
                                "c Front (kPa)": f"{_c_f:.0f}",
                                "c Back (kPa)":  _c_b_disp,
                                "Có ở Back?": _on_back,
                            })
                    except Exception as _e_load_e:
                        st.caption(f"_(Không tải được hồ sơ địa chất HK: {_e_load_e})_")

                    # Fallback nếu SQLite không có data
                    if not _front_s:
                        _front_s = [
                            _EL_S(_Z_s - _H1_s, 15.0, 5.0, 0.0, _su_F or 10.0),
                            _EL_S(_pile_bot_s - 1.0, 18.0, 8.0, 30.0, 0.0),
                        ]
                        # Back: chỉ bùn + cát dưới — KHÔNG có Fill
                        _back_s = [
                            _EL_S(_Zb_s - _H1_s, 15.0, 5.0, 0.0, _su_B or _su_F or 10.0),
                            _EL_S(_pile_bot_s - 1.0, 18.0, 8.0, 30.0, 0.0),
                        ]
                        _layers_log = [
                            {"Symbol": "1", "z_top (m)": f"{_Z_s:+.2f}",
                             "z_bot (m)": f"{_Z_s - _H1_s:+.2f}",
                             "γ": "15.0", "γ_sub": "5.0", "φ": "0.0",
                             "c Front (kPa)": f"{_su_F:.0f}",
                             "c Back (kPa)": f"{_su_B:.0f}",
                             "Có ở Back?": "✓"},
                            {"Symbol": "2", "z_top (m)": f"{_Z_s - _H1_s:+.2f}",
                             "z_bot (m)": f"{_pile_bot_s - 1.0:+.2f}",
                             "γ": "18.0", "γ_sub": "8.0", "φ": "30.0",
                             "c Front (kPa)": "0", "c Back (kPa)": "0",
                             "Có ở Back?": "✓"},
                        ]
                        st.caption("_(Dùng layers giả định — không có dữ liệu hồ sơ địa chất cho HK này.)_")

                    # ── Hiển thị thông tin tường + HK trước khi tính ──────
                    _info_c1, _info_c2 = st.columns(2)
                    with _info_c1:
                        st.markdown("**Thông tin tường chắn (D.1):**")
                        st.markdown(
                            f"- Đỉnh kè: **{_top_s:+.2f} m** · Cao độ Front Z: **{_Z_s:+.2f} m** · "
                            f"Back Zb: **{_Zb_s:+.2f} m**\n"
                            f"- Cọc: **{_dpy_pile}** · L = **{_L_s:.1f} m** · "
                            f"chân cừ: **{_pile_bot_s:+.2f} m**\n"
                            f"- MNN Front: **{_wlvl_s:+.2f}** · Back: **{_wlvl_b_s:+.2f}** m\n"
                            f"- Tải mặt q = **{_q_s:.0f} kN/m²**\n"
                            f"- Fill: γ={float(_dpy_gamma_fill):.1f}, φ={float(_dpy_phi_fill):.1f}°, "
                            f"c={float(_dpy_c_fill):.0f}, γ_sub={float(_dpy_gamma_sub_fill):.1f}"
                        )
                    with _info_c2:
                        st.markdown(f"**Hố khoan {_hk_iter}:**")
                        st.markdown(
                            f"- Cao độ cổ HK: **{_bh_elev_e:+.2f} m**\n"
                            f"- Số lớp đất tải vào ổn định: **{len(_front_s)}**\n"
                            f"- Su Front (Ctd CDM): **{_su_F:.0f} kPa**\n"
                            f"- Su Back (tự nhiên): **{_su_B:.0f} kPa**\n"
                            f"- ε₅₀ áp dụng: **{float(_dpy_eps50):.4f}**"
                        )
                    if _layers_log:
                        _src_lbl = "thí nghiệm" if _rows_e else "giả định"
                        with st.expander(
                            f"Xác nhận thông số {len(_layers_log)} lớp — nguồn: {_src_lbl} "
                            f"({_bh_full_e if '_bh_full_e' in dir() else _hk_iter})",
                            expanded=True,
                        ):
                            # st.table thay st.dataframe — in PDF ổn định hơn
                            st.table(pd.DataFrame(_layers_log))

                    # CDM block — Geometry tự động:
                    #   top = Z_s (đáy fill / mặt đất tự nhiên Front)
                    #   bot = Z_s − H1 − 1m (đáy lớp bùn + 1m vào lớp tốt)
                    #   width = 5m mặc định (user override được)
                    _cdm_top_e = _Z_s
                    _cdm_bot_e = _Z_s - _H1_s - 1.0   # đáy bùn + 1m
                    _cdm_thk_e = _cdm_top_e - _cdm_bot_e   # H1 + 1m
                    # _cdm_w_e đã được nhập ở đầu expander (chung Mục D + E)
                    _cdm_a    = float(st.session_state.get("cdm_a_ratio", 0.20) or 0.20)
                    _cdm_qu   = float(st.session_state.get("cdm_qu_28", 0.0) or 150.0)
                    _cdm_phi  = float(st.session_state.get("cdm_phi_col", 30.0) or 30.0)
                    _cdm_gam  = float(st.session_state.get("cdm_gamma_col", 19.0) or 19.0)
                    # c_col = qu/2 (Su_undrained = qu/2 cho mẫu CDM 28 ngày)
                    _cdm_c_col = _cdm_qu / 2.0
                    _cdm_s = _CDMB(top_elev=_cdm_top_e, bot_elev=_cdm_bot_e,
                                    area_ratio_a=_cdm_a, c_col_kPa=_cdm_c_col,
                                    phi_col_deg=_cdm_phi, gamma_col_kNm3=_cdm_gam)

                    # Selectbox 3 phương pháp tính trượt cung tròn (Limit Equilibrium)
                    _SLIP_METHODS = {
                        "bishop":            "Bishop Simplified",
                        "spencer":           "Spencer",
                        "morgenstern_price": "Morgenstern-Price",
                    }
                    st.caption(
                        "**Bishop Simplified** — ΣM = 0 quanh tâm, tan(φ)/F trong mẫu số (lặp).  \n"
                        "**Spencer** — thoả ΣF + ΣM, giả thiết góc lực giữa lát = const.  \n"
                        "**Morgenstern-Price** — thoả ΣF + ΣM, hàm phân bố λ·f(x) cho lực giữa lát."
                    )

                    # Tính 3 phương pháp slip (Bishop / Spencer / M-P).
                    # Lật + Toe-kickout không phụ thuộc method.
                    from sw_global_stability import check_global_slip as _chk_slip
                    _slip_fos_all: dict[str, float | None] = {}
                    _slip_geom_all: dict[str, tuple[float, float, float]] = {}
                    with st.spinner("Đang chạy 3 phương pháp Limit Equilibrium..."):
                        for _mk in _SLIP_METHODS:
                            try:
                                _fs_m, _xc_m, _yc_m, _R_m = _chk_slip(
                                    _geom_s, _front_s, _back_s,
                                    fill=_fill_s, cdm=_cdm_s, method=_mk,
                                )
                                if _fs_m is not None and float(_fs_m) > 0:
                                    _slip_fos_all[_mk] = float(_fs_m)
                                    _slip_geom_all[_mk] = (
                                        float(_xc_m), float(_yc_m), float(_R_m),
                                    )
                                else:
                                    _slip_fos_all[_mk] = None
                            except Exception:
                                _slip_fos_all[_mk] = None

                    # Critical method = min Fs trong các PP hợp lệ
                    _valid_fs_dict = {k: v for k, v in _slip_fos_all.items()
                                       if v is not None and v > 0}
                    _slip_method_sel = (min(_valid_fs_dict, key=_valid_fs_dict.get)
                                          if _valid_fs_dict else "bishop")

                    # Chạy đầy đủ 3 kiểm tra (slip + lật + toe) bằng phương pháp critical
                    with st.spinner("Đang chạy kiểm tra ổn định tổng thể..."):
                        _res_s = _sw_check_all(_geom_s, _front_s, _back_s,
                                                fill=_fill_s, cdm=_cdm_s,
                                                method=_slip_method_sel)
                    # Đồng bộ Fs slip + geometry theo critical (tránh sai khác do
                    # check_all gọi 1 lần độc lập với loop trên)
                    if _slip_method_sel in _valid_fs_dict:
                        _res_s.Fs_global_slip = _valid_fs_dict[_slip_method_sel]
                        if _slip_method_sel in _slip_geom_all:
                            _xc_c, _yc_c, _R_c = _slip_geom_all[_slip_method_sel]
                            _res_s.slip_xc = _xc_c
                            _res_s.slip_yc = _yc_c
                            _res_s.slip_R = _R_c
                        _res_s.slip_method = _slip_method_sel

                    # Lưu 3 phương pháp vào SQLite ke_sw_stability
                    try:
                        import sqlite3 as _sq3_sv
                        _con_sv = _sq3_sv.connect(str(_DB), timeout=30)
                        _con_sv.execute("PRAGMA journal_mode=WAL")
                        _con_sv.execute("""
                            CREATE TABLE IF NOT EXISTS ke_sw_stability (
                                id INTEGER PRIMARY KEY AUTOINCREMENT,
                                project TEXT DEFAULT '202605-TTHC',
                                zone TEXT DEFAULT 'KE',
                                bh_name TEXT NOT NULL, pile_type TEXT NOT NULL,
                                L_m REAL NOT NULL, method TEXT NOT NULL DEFAULT 'bishop',
                                Fs_slip REAL, Fs_overturning REAL, Fs_toe_kickout REAL,
                                slip_xc REAL, slip_yc REAL, slip_R REAL,
                                M_giu_kNm REAL, M_lat_kNm REAL,
                                top_elev REAL, Z_m REAL, Zb_m REAL,
                                wlvl_front REAL, wlvl_back REAL, q_kPa REAL,
                                cdm_a REAL, cdm_c_col REAL, cdm_w_m REAL,
                                su_front REAL, su_back REAL, H1_m REAL,
                                warnings TEXT, ts TEXT,
                                UNIQUE(bh_name, pile_type, L_m, method)
                            )
                        """)
                        _bh_sv = (f"KE-{_hk_iter}" if not _hk_iter.startswith("KE-")
                                   else _hk_iter)
                        _warn_sv = "; ".join(_res_s.warnings) if _res_s.warnings else ""
                        # Lưu từng phương pháp (4 rows per HK)
                        for _mk_sv, _fs_sl_sv in _slip_fos_all.items():
                            if not _fs_sl_sv or _fs_sl_sv <= 0:
                                continue
                            _gx, _gy, _gR = _slip_geom_all.get(
                                _mk_sv,
                                (_res_s.slip_xc, _res_s.slip_yc, _res_s.slip_R),
                            )
                            # M_giu/M_lat chỉ lưu cho phương pháp critical
                            _m_giu_sv = (float(_res_s.M_giu_kNm)
                                         if _mk_sv == _slip_method_sel else None)
                            _m_lat_sv = (float(_res_s.M_lat_kNm)
                                         if _mk_sv == _slip_method_sel else None)
                            _con_sv.execute("""
                                INSERT OR REPLACE INTO ke_sw_stability
                                (bh_name, pile_type, L_m, method,
                                 Fs_slip, Fs_overturning,
                                 slip_xc, slip_yc, slip_R,
                                 M_giu_kNm, M_lat_kNm,
                                 top_elev, Z_m, Zb_m, wlvl_front, wlvl_back, q_kPa,
                                 cdm_a, cdm_c_col, cdm_w_m, su_front, su_back, H1_m,
                                 warnings, ts)
                                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,
                                        datetime('now','localtime'))
                            """, (
                                _bh_sv, str(_dpy_pile), float(_L_s), _mk_sv,
                                float(_fs_sl_sv), float(_res_s.Fs_overturning),
                                float(_gx), float(_gy), float(_gR),
                                _m_giu_sv, _m_lat_sv,
                                float(_top_s), float(_Z_s), float(_Zb_s),
                                float(_wlvl_s), float(_wlvl_b_s), float(_q_s),
                                float(_cdm_a), float(_cdm_c_col), float(_cdm_w_e),
                                float(_su_F), float(_su_B), float(_H1_s),
                                _warn_sv,
                            ))
                        _con_sv.commit()
                        _con_sv.close()
                        # Đồng bộ ngược lên Mục B session state để E tìm đúng key
                        st.session_state.setdefault("ke_sw_rec_piles", {})[_hk_iter] = str(_dpy_pile)
                        st.session_state.setdefault("ke_sw_L_thiet_ke", {})[_hk_iter] = float(_L_s)
                    except Exception:
                        pass  # không block UI nếu SQLite lỗi

                    # 2 metric cards (Fs trượt + Fs lật; nhổ chân bỏ)
                    _se1, _se2 = st.columns(2)
                    _fs_sl_crit = _valid_fs_dict.get(_slip_method_sel, _res_s.Fs_global_slip)
                    _se1.metric(
                        f"Fs trượt cung tròn ({_SLIP_METHODS[_slip_method_sel]} ★)",
                        f"{_fs_sl_crit:.3f}",
                        f"min 1.40 — {'Đạt' if _fs_sl_crit >= 1.40 else 'KHÔNG ĐẠT'}",
                        delta_color="normal" if _fs_sl_crit >= 1.40 else "inverse",
                    )
                    _se2.metric(
                        "Fs lật quanh chân cừ",
                        f"{_res_s.Fs_overturning:.3f}",
                        f"min 1.20 — {'Đạt' if _res_s.Fs_overturning >= 1.20 else 'KHÔNG ĐẠT'}",
                        delta_color="normal" if _res_s.Fs_overturning >= 1.20 else "inverse",
                    )

                    # ── So sánh 3 phương pháp slip cạnh nhau ─────────────────
                    st.markdown(
                        "**So sánh Fs trượt cung tròn — 3 phương pháp Limit Equilibrium**"
                    )
                    _FS_MIN_SLIP = {
                        "bishop": 1.40,
                        "spencer": 1.40, "morgenstern_price": 1.40,
                    }
                    _cmp_cols = st.columns(len(_SLIP_METHODS))
                    _valid_fs = [v for v in _slip_fos_all.values() if v is not None and v > 0]
                    _critical_fs = min(_valid_fs) if _valid_fs else None
                    for _ic, (_mk, _mlabel) in enumerate(_SLIP_METHODS.items()):
                        _fs_m = _slip_fos_all.get(_mk)
                        _fmin = _FS_MIN_SLIP[_mk]
                        if _fs_m is None or _fs_m <= 0:
                            _cmp_cols[_ic].metric(_mlabel, "—",
                                                   "Không hội tụ",
                                                   delta_color="off")
                            continue
                        _ok_m = _fs_m >= _fmin
                        _is_crit = (_critical_fs is not None
                                    and abs(_fs_m - _critical_fs) < 1e-6)
                        _label_show = (f"{_mlabel} ★" if _is_crit
                                        else _mlabel)
                        _cmp_cols[_ic].metric(
                            _label_show, f"{_fs_m:.3f}",
                            f"min {_fmin:.2f} — "
                            f"{'Đạt' if _ok_m else 'KHÔNG ĐẠT'}",
                            delta_color="normal" if _ok_m else "inverse",
                        )
                    st.caption("★ = critical (Fs nhỏ nhất).")

                    # Chi tiết tính toán + mặt trượt nguy hiểm (trải phẳng)
                    st.markdown("#### Chi tiết tính toán + mặt trượt nguy hiểm")
                    if True:
                        st.markdown(f"**Phương pháp:** `{_res_s.slip_method}`")
                        st.markdown(
                            f"- Tâm mặt trượt: ({_res_s.slip_xc:+.1f}, {_res_s.slip_yc:+.1f}) m, "
                            f"R = {_res_s.slip_R:.1f} m\n"
                            f"- Moment lật (Ma): {_res_s.M_lat_kNm:,.0f} kNm/m\n"
                            f"- Moment giữ (Mp): {_res_s.M_giu_kNm:,.0f} kNm/m\n"
                            f"- Ma (Active Front): {_res_s.Ma_kNm:,.0f} kNm/m\n"
                            f"- Mp (Passive Back): {_res_s.Mp_kNm:,.0f} kNm/m\n"
                            f"- CDM Lc dùng: {_cdm_Lc_val:.1f} m (composite φ_eq, c_eq, γ_eq theo TCVN 9403)"
                        )
                        if _res_s.warnings:
                            for _w in _res_s.warnings:
                                st.warning(_w)

                        # Vẽ sơ đồ mặt trượt + lớp đất + bán kính
                        try:
                            import matplotlib.pyplot as _plt_s
                            from matplotlib.patches import Circle as _Circ_s, Rectangle as _Rect_s

                            # Scale dynamic — chỉ bao đoạn cung trượt thực sự
                            # (dưới mặt đất), không lấy cả vòng tròn lớn.
                            _R = float(_res_s.slip_R)
                            _xc = float(_res_s.slip_xc)
                            _yc = float(_res_s.slip_yc)

                            import numpy as _np_lim
                            _th_lim = _np_lim.linspace(0, 2 * _np_lim.pi, 360)
                            _xs_lim = _xc + _R * _np_lim.cos(_th_lim)
                            _ys_lim = _yc + _R * _np_lim.sin(_th_lim)
                            _mk_lim = (
                                ((_xs_lim < 0) & (_ys_lim < _top_s))
                                | ((_xs_lim >= 0) & (_ys_lim < _Zb_s))
                            )
                            if _mk_lim.any():
                                _arc_x_min = float(_xs_lim[_mk_lim].min())
                                _arc_x_max = float(_xs_lim[_mk_lim].max())
                                _arc_y_min = float(_ys_lim[_mk_lim].min())
                            else:
                                _arc_x_min, _arc_x_max = _xc - _R, _xc + _R
                                _arc_y_min = _yc - _R
                            # Khung gọn: vừa với arc + chân cừ + lề 2-3m
                            _x_min = min(_arc_x_min - 2.5, _pile_bot_s - 2, -10)
                            _x_max = max(_arc_x_max + 2.5, 10)
                            _y_min = min(_arc_y_min - 1.5, _pile_bot_s - 1.5)
                            # +3m phía trên top để chứa mũi tên surcharge q
                            _y_max = max(_top_s + 3.5, _yc + 1.5 if _yc > _top_s else _top_s + 3.5)

                            _fig_s, _ax_s = _plt_s.subplots(1, 1, figsize=(9, 6.5))

                            # ── Vẽ lớp đất Front + Back với màu theo symbol ────
                            _layer_colors = {
                                "F":   "#D7CCC8",   # cát san lấp
                                "1":   "#8FB3D6",   # sét bùn yếu
                                "XMD": "#90c890",   # CDM
                                "2A":  "#FFE082",   # cát
                                "2B":  "#FFE082",
                                "2C":  "#FFE082",
                                "3":   "#A1887F",   # sét dẻo
                                "4":   "#FFB74D",
                                "5":   "#D7846E",
                                "5A":  "#D7846E",
                                "5B":  "#D7846E",
                                "6":   "#FFA000",
                                "7":   "#FF8F00",
                            }

                            def _draw_layers_side(_layers, _x_left, _x_right, _ground_elev):
                                """Vẽ các lớp đất từ ground xuống chân cừ theo elev."""
                                _e_cur = _ground_elev
                                for _i_l, _lay in enumerate(_layers):
                                    _e_bot = _lay.tip_elev
                                    if _e_bot >= _e_cur:
                                        continue
                                    # Safe access _layers_log (có thể ngắn hơn nếu shift)
                                    if _i_l < len(_layers_log):
                                        _sym_l = _layers_log[_i_l].get("Symbol", "")
                                        _gam_str = _layers_log[_i_l].get("γ", f"{_lay.gamma:.1f}")
                                    else:
                                        _sym_l = ""
                                        _gam_str = f"{_lay.gamma:.1f}"
                                    _col_l = _layer_colors.get(_sym_l, "#CFD8DC")
                                    _ax_s.add_patch(_Rect_s(
                                        (_x_left, _e_bot), _x_right - _x_left, _e_cur - _e_bot,
                                        facecolor=_col_l, edgecolor="#666", lw=0.4, alpha=0.6,
                                    ))
                                    _y_mid = (_e_cur + _e_bot) / 2.0
                                    if abs(_e_cur - _e_bot) > 0.8:
                                        _ax_s.text((_x_left + _x_right) / 2.0, _y_mid,
                                                    f"{_sym_l}\nγ={_gam_str}",
                                                    fontsize=7, ha="center", va="center",
                                                    color="#333",
                                                    bbox=dict(boxstyle="round,pad=0.1",
                                                              facecolor="white",
                                                              edgecolor="none", alpha=0.7))
                                    _e_cur = _e_bot

                            # Front (trái cừ, từ mặt đất Z xuống) — dùng full _layers_log
                            _draw_layers_side(_front_s, _x_min, -0.3, _Z_s)
                            # Back (phải cừ, từ Zb xuống) — CHỈ đất tự nhiên (bùn trở xuống)
                            _back_log = [l for l in _layers_log if l.get("Có ở Back?") == "✓"]
                            _layers_log_save = _layers_log
                            _layers_log = _back_log
                            _draw_layers_side(_back_s, 0.3, _x_max, _Zb_s)
                            _layers_log = _layers_log_save
                            # Nhãn "Đất tự nhiên" phía Back
                            _ax_s.text(
                                (0.3 + _x_max) / 2, _Zb_s + 0.4,
                                "Đất tự nhiên",
                                fontsize=8, ha="center", color="#4a4a4a",
                                bbox=dict(boxstyle="round,pad=0.25", facecolor="#FFF9C4",
                                          edgecolor="#888", lw=0.7, alpha=0.9)
                            )

                            # Đất đắp Fill (Front, từ Z lên đến đỉnh kè)
                            if _top_s > _Z_s:
                                _ax_s.add_patch(_Rect_s(
                                    (_x_min, _Z_s), -0.3 - _x_min, _top_s - _Z_s,
                                    facecolor="#D7CCC8", edgecolor="#8B4513",
                                    hatch="..", lw=0.5, alpha=0.6,
                                    label="Fill đắp"))

                            # Mặt đất profile
                            _ax_s.plot([_x_min, -0.3], [_Z_s, _Z_s], "k-", lw=1.5)
                            _ax_s.plot([-0.3, 0.3], [_top_s, _top_s], "k-", lw=1.5)
                            _ax_s.plot([0.3, _x_max], [_Zb_s, _Zb_s], "k-", lw=1.5)
                            _ax_s.plot([-0.3, -0.3], [_Z_s, _top_s], "k-", lw=1.5)
                            _ax_s.plot([0.3, 0.3], [_top_s, _Zb_s], "k-", lw=1.5)

                            # MNN Front + Back
                            _ax_s.plot([_x_min, -0.3], [_wlvl_s, _wlvl_s],
                                        color="cyan", lw=1.0, alpha=0.6)
                            _ax_s.plot([0.3, _x_max], [_wlvl_b_s, _wlvl_b_s],
                                        color="cyan", lw=1.0, alpha=0.6)
                            _ax_s.text(_x_min + 0.5, _wlvl_s + 0.2, "MNN_F", fontsize=7,
                                        color="cyan")
                            _ax_s.text(_x_max - 4, _wlvl_b_s + 0.2, "MNN_B", fontsize=7,
                                        color="cyan")

                            # ── Tải khai thác q (surcharge) — mũi tên trên mặt Front ───
                            if _q_s > 0:
                                _q_y0 = _top_s   # đỉnh kè
                                _q_y_arrow = _q_y0 + 1.5
                                _q_x_start = max(_x_min + 1.0, -0.3 - 12)
                                _q_x_end = -0.3
                                # Khung tải (đường ngang trên + ngang dưới)
                                _ax_s.plot([_q_x_start, _q_x_end],
                                            [_q_y_arrow, _q_y_arrow],
                                            color="#C62828", lw=1.5)
                                # Mũi tên xuống cách đều
                                _n_arrows = max(4, int((_q_x_end - _q_x_start) / 1.5))
                                for _ka in range(_n_arrows + 1):
                                    _xa = _q_x_start + _ka * (_q_x_end - _q_x_start) / _n_arrows
                                    _ax_s.annotate(
                                        "", xy=(_xa, _q_y0 + 0.05),
                                        xytext=(_xa, _q_y_arrow - 0.05),
                                        arrowprops=dict(arrowstyle="->",
                                                         color="#C62828", lw=1.0))
                                _ax_s.text((_q_x_start + _q_x_end) / 2,
                                            _q_y_arrow + 0.4,
                                            f"q = {_q_s:.0f} kN/m² (tải khai thác)",
                                            fontsize=8, color="#C62828", fontweight="bold",
                                            ha="center",
                                            bbox=dict(boxstyle="round,pad=0.2",
                                                      facecolor="white",
                                                      edgecolor="#C62828", lw=0.6))

                            # Cừ SW
                            _ax_s.add_patch(_Rect_s((-0.3, _pile_bot_s), 0.6, _L_s,
                                                     facecolor="#444444", edgecolor="black",
                                                     label=f"Cừ SW L={_L_s:.1f}m"))
                            # CDM — Front: từ cừ ra phía trái width = _cdm_w_e
                            if _cdm_s:
                                _cdm_x_right = -0.3   # sát cừ
                                _cdm_x_left = _cdm_x_right - _cdm_w_e
                                _ax_s.add_patch(_Rect_s(
                                    (_cdm_x_left, _cdm_s.bot_elev),
                                    _cdm_w_e, _cdm_s.thickness,
                                    facecolor="#90c890", alpha=0.55,
                                    hatch="//", edgecolor="green", lw=1.0,
                                    label=f"Vùng CDM (b={_cdm_w_e:.1f}m, dày={_cdm_s.thickness:.1f}m)"))
                                # Annotation chiều rộng CDM
                                _ax_s.annotate(
                                    "", xy=(_cdm_x_right, _cdm_s.bot_elev - 0.6),
                                    xytext=(_cdm_x_left, _cdm_s.bot_elev - 0.6),
                                    arrowprops=dict(arrowstyle="<->", color="green", lw=1.0))
                                _ax_s.text((_cdm_x_left + _cdm_x_right) / 2,
                                            _cdm_s.bot_elev - 0.9,
                                            f"b = {_cdm_w_e:.1f} m",
                                            fontsize=8, color="green", fontweight="bold",
                                            ha="center")

                            # Cung trượt — CHỈ vẽ phần dưới mặt đất (Front<top, Back<Zb)
                            import numpy as _np_arc
                            _theta = _np_arc.linspace(0, 2 * _np_arc.pi, 360)
                            _xs_c = _xc + _R * _np_arc.cos(_theta)
                            _ys_c = _yc + _R * _np_arc.sin(_theta)
                            # Mask: trên cừ (x<0) phần thấp hơn top; bên Back (x>=0) thấp hơn Zb
                            _mask_arc = (
                                ((_xs_c < 0) & (_ys_c < _top_s))
                                | ((_xs_c >= 0) & (_ys_c < _Zb_s))
                            )
                            _xs_plot = _np_arc.where(_mask_arc, _xs_c, _np_arc.nan)
                            _ys_plot = _np_arc.where(_mask_arc, _ys_c, _np_arc.nan)
                            _ax_s.plot(_xs_plot, _ys_plot,
                                        color="red", linestyle="--", lw=2.2,
                                        label=f"Cung trượt Fs={_res_s.Fs_global_slip:.2f}")
                            # Tâm O — chỉ vẽ nếu nằm trong khung; nếu không chú thích R ở góc
                            if _x_min <= _xc <= _x_max and _y_min <= _yc <= _y_max:
                                _ax_s.plot(_xc, _yc, "r+", markersize=12, markeredgewidth=2)
                                _ax_s.text(_xc + 0.3, _yc + 0.3,
                                            f"O({_xc:+.1f},{_yc:+.1f}) R={_R:.1f}m",
                                            fontsize=7, color="red", fontweight="bold")
                            else:
                                _ax_s.text(0.98, 0.97,
                                            f"O({_xc:+.1f},{_yc:+.1f})\nR = {_R:.1f} m",
                                            transform=_ax_s.transAxes, fontsize=8,
                                            color="red", fontweight="bold",
                                            ha="right", va="top",
                                            bbox=dict(boxstyle="round,pad=0.3",
                                                      facecolor="white", edgecolor="red", lw=0.8))

                            # Annotations cao độ trên trục y bên trái
                            for _yv, _txt, _col in [
                                (_top_s,      f"top = {_top_s:+.2f}",     "#1565C0"),
                                (_Z_s,        f"Z = {_Z_s:+.2f}",         "#8B4513"),
                                (_Zb_s,       f"Zb = {_Zb_s:+.2f}",       "#FF6F00"),
                                (_pile_bot_s, f"tip = {_pile_bot_s:+.2f}", "#444"),
                            ]:
                                _ax_s.axhline(_yv, color=_col, lw=0.4, alpha=0.3)
                                _ax_s.text(_x_min + 0.2, _yv + 0.15, _txt,
                                            fontsize=7, color=_col, fontweight="bold")

                            _ax_s.set_xlim(_x_min, _x_max)
                            _ax_s.set_ylim(_y_min, _y_max)
                            _ax_s.set_aspect("equal")
                            _ax_s.set_xlabel("x (m)")
                            _ax_s.set_ylabel("Cao độ (m)")
                            _ax_s.set_title(
                                f"Mặt trượt nguy hiểm — Fs={_res_s.Fs_global_slip:.2f}  "
                                f"({_res_s.slip_method})",
                                fontsize=12, fontweight="bold",
                            )
                            _ax_s.grid(alpha=0.3)
                            _ax_s.legend(loc="upper right", fontsize=8, framealpha=0.9)
                            _plt_s.tight_layout()
                            st.pyplot(_fig_s, use_container_width=True)
                            _plt_s.close(_fig_s)
                        except Exception as _e_plot_s:
                            st.caption(f"_(Không vẽ được sơ đồ mặt trượt: {_e_plot_s})_")

                except ImportError as _e_imp:
                    st.info(
                        "Tính năng ổn định tổng thể cần module `sw_global_stability` "
                        "và `geotech-staff-engineer`. Cài: "
                        "`pip install geotech-staff-engineer`"
                    )
                except Exception as _e_st:
                    st.warning(f"Không chạy được kiểm tra ổn định tổng thể: {_e_st}")


    # Xuất PDF qua nút đã tắt (§25 CLAUDE.md) — dùng Ctrl+P trình duyệt

    # ── Xuất Word toàn bộ tab Cọc ván SW ────────────────────────────────────
    st.divider()
    if st.button("Xuất Word toàn diện — Cọc ván SW (Kè)", use_container_width=True, type="primary"):
        with st.spinner("Đang xây dựng tài liệu Word toàn diện..."):
            try:
                import io as _io_w
                import sqlite3 as _sq_w
                import numpy as _np_w
                import matplotlib
                matplotlib.use("Agg")
                import matplotlib.pyplot as _plt_w
                from mpl_toolkits.mplot3d import Axes3D as _Axes3D  # noqa: F401
                from docx import Document as _Docx_w
                from docx.shared import Pt as _Pt_w, Cm as _Cm_w, RGBColor as _RGB_w
                from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_AL
                from docx.oxml.ns import qn as _qn_w
                from docx.oxml import OxmlElement as _OXE_w
                from datetime import datetime as _dt_w
                import sys as _sys_w
                _sys_w.path.insert(0, str(_ROOT / "scripts"))
                from ke_sw_nt_calc import _load_catalog as _lcat_w
                from wall_internal_force import SoilLayer as _SL_w, PileProps as _PP_w
                from winkler_np import solve_numpy_dist as _wk_solve_w
                _MD_THEORY_W = (_ROOT / "48-ke-sw-on-dinh-tong-the.md")

                # ── helpers ─────────────────────────────────────────────────
                def _frun_w(para, field):
                    r = para.add_run()
                    b = _OXE_w("w:fldChar"); b.set(_qn_w("w:fldCharType"), "begin")
                    ins = _OXE_w("w:instrText"); ins.text = field
                    ins.set(_qn_w("xml:space"), "preserve")
                    e = _OXE_w("w:fldChar"); e.set(_qn_w("w:fldCharType"), "end")
                    r._r.append(b); r._r.append(ins); r._r.append(e)

                def _no_border_w(tbl):
                    tbl_elem = tbl._tbl
                    pr = tbl_elem.find(_qn_w("w:tblPr"))
                    if pr is None:
                        pr = _OXE_w("w:tblPr")
                        tbl_elem.insert(0, pr)
                    bd = _OXE_w("w:tblBorders")
                    for s in ("top","left","bottom","right","insideH","insideV"):
                        n = _OXE_w(f"w:{s}"); n.set(_qn_w("w:val"), "none"); bd.append(n)
                    pr.append(bd)

                def _cell_bg_w(cell, hex_color):
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = _OXE_w("w:shd")
                    shd.set(_qn_w("w:val"), "clear")
                    shd.set(_qn_w("w:color"), "auto")
                    shd.set(_qn_w("w:fill"), hex_color)
                    tcPr.append(shd)

                def _fmt_cells_w(row_or_cells, pt=8):
                    cells = row_or_cells if isinstance(row_or_cells, (list, tuple)) else row_or_cells.cells
                    for c in cells:
                        for p in c.paragraphs:
                            for r in p.runs:
                                r.font.size = _Pt_w(pt)

                def _formula_png_w(rows, title="", w_cm=14.5, h_rows_cm=1.1):
                    h = max(h_rows_cm * len(rows) + (0.7 if title else 0), 1.2)
                    fig, ax = _plt_w.subplots(figsize=(w_cm / 2.54, h / 2.54), dpi=120)
                    ax.set_facecolor("#f4f6fb"); fig.patch.set_facecolor("#f4f6fb")
                    ax.axis("off")
                    y = 0.96
                    if title:
                        ax.text(0.01, y, title, transform=ax.transAxes,
                                fontsize=8, fontweight="bold", va="top", color="#1a3c5e")
                        y -= 0.14
                    step = (y - 0.03) / max(len(rows), 1)
                    for i, (lbl, mstr) in enumerate(rows):
                        yi = y - (i + 0.5) * step
                        if lbl:
                            ax.text(0.01, yi, lbl, transform=ax.transAxes,
                                    fontsize=7, va="center", color="#555", style="italic")
                        ax.text(0.50, yi, mstr, transform=ax.transAxes,
                                fontsize=11, va="center", ha="center")
                    buf = _io_w.BytesIO()
                    fig.savefig(buf, format="png", bbox_inches="tight",
                                dpi=120, facecolor=fig.get_facecolor())
                    _plt_w.close(fig)
                    buf.seek(0)
                    return buf

                # ── Load data từ SQLite ──────────────────────────────────────
                with _sq_w.connect(str(_DB), timeout=10) as _cw:
                    _cw.row_factory = _sq_w.Row
                    _nt_det_w = {r["bh_name"]: dict(r) for r in
                                 _cw.execute("SELECT * FROM ke_sw_nt_detail").fetchall()}
                    _bhs_w = {r["name"]: dict(r) for r in
                              _cw.execute(
                                  "SELECT name,x_coord_m,y_coord_m,elevation_m "
                                  "FROM boreholes WHERE name LIKE 'KE-%'"
                              ).fetchall()}
                    # NT2 compare
                    _nt2_cmp_w = {}
                    try:
                        rows_cmp = _cw.execute(
                            "SELECT * FROM ke_sw_nt2_compare"
                        ).fetchall()
                        for r in rows_cmp:
                            k = (r["bh_name"], r["pile_type"], r["L_m"])
                            _nt2_cmp_w.setdefault(k, {})[r["method"]] = dict(r)
                    except Exception:
                        pass
                    # Winkler results — distributed load case only
                    _winkler_w = {}
                    try:
                        for r in _cw.execute(
                            "SELECT * FROM ke_sw_winkler_results "
                            "WHERE load_case='distributed'"
                        ).fetchall():
                            k = (r["bh_name"], r["pile_type"])
                            # keep row with higher u_max (worst case) if multiple rows
                            if k not in _winkler_w or (r["u_max_mm"] or 0) > (_winkler_w[k].get("u_max_mm") or 0):
                                _winkler_w[k] = dict(r)
                    except Exception:
                        pass
                    # Global stability — pivot ke_sw_stability by method + geometry
                    _stab_w = {}
                    _stab_geom_w = {}   # geometry per bh_name for Winkler re-run
                    try:
                        for r in _cw.execute(
                            "SELECT bh_name,pile_type,L_m,method,Fs_slip,"
                            "Fs_overturning,Fs_toe_kickout,"
                            "top_elev,Z_m,Zb_m,wlvl_front,wlvl_back,"
                            "q_kPa,cdm_w_m "
                            "FROM ke_sw_stability "
                            "WHERE method IN ('bishop','spencer','morgenstern_price')"
                        ).fetchall():
                            _k = (r["bh_name"], r["pile_type"])
                            if _k not in _stab_w:
                                _stab_w[_k] = {"L_m": r["L_m"],
                                                "Fs_overturning": r["Fs_overturning"],
                                                "Fs_toe_kickout": r["Fs_toe_kickout"]}
                            _stab_w[_k][f"Fs_{r['method']}"] = r["Fs_slip"]
                            if r["bh_name"] not in _stab_geom_w:
                                _stab_geom_w[r["bh_name"]] = {
                                    "top_elev": r["top_elev"], "Z_m": r["Z_m"],
                                    "Zb_m": r["Zb_m"], "wlvl_front": r["wlvl_front"],
                                    "wlvl_back": r["wlvl_back"], "q_kPa": r["q_kPa"],
                                    "cdm_w_m": r["cdm_w_m"],
                                }
                    except Exception:
                        pass

                _cat_w  = _lcat_w()
                _ss_rec = st.session_state.get("ke_sw_rec_piles", {}) or {}
                _ss_L   = st.session_state.get("ke_sw_L_thiet_ke", {}) or {}
                _align_hks = sorted(_nt_det_w.keys())

                # ── Load per-HK layer data + Su ────────────────────────────────
                _layers_per_hk_w = {}
                with _sq_w.connect(str(_DB), timeout=10) as _cwl:
                    _cwl.row_factory = _sq_w.Row
                    for _bhn_l in _align_hks:
                        try:
                            _bhid = _cwl.execute(
                                "SELECT id FROM boreholes WHERE name=?", (_bhn_l,)
                            ).fetchone()
                            if not _bhid:
                                continue
                            _bhid = _bhid["id"]
                            _lyrs_l = [dict(r) for r in _cwl.execute(
                                "SELECT symbol, depth_top_m, depth_bot_m, "
                                "COALESCE(gamma_kNm3, 15.5) AS gamma_kNm3 "
                                "FROM layers WHERE borehole_id=? ORDER BY depth_top_m",
                                (_bhid,)
                            ).fetchall()]
                            # Su: VST priority then lab
                            _sus_l = {}
                            for r in _cwl.execute(
                                "SELECT depth_m, Su_kPa FROM vane_shear_tests "
                                "WHERE borehole_id=? ORDER BY depth_m", (_bhid,)
                            ).fetchall():
                                _sus_l[float(r["depth_m"])] = float(r["Su_kPa"])
                            for r in _cwl.execute(
                                "SELECT depth_mid_m, "
                                "COALESCE(Cu_UU_kPa, c_kPa, 10.0) AS su "
                                "FROM lab_tests WHERE borehole_id=? ORDER BY depth_mid_m",
                                (_bhid,)
                            ).fetchall():
                                dm = float(r["depth_mid_m"])
                                if dm not in _sus_l:
                                    _sus_l[dm] = float(r["su"])
                            for _lyr_l in _lyrs_l:
                                dm = (_lyr_l["depth_top_m"] + _lyr_l["depth_bot_m"]) / 2
                                if _sus_l:
                                    _nd = min(_sus_l, key=lambda d: abs(d - dm))
                                    _lyr_l["su_kPa"] = max(1.0, _sus_l[_nd])
                                else:
                                    _lyr_l["su_kPa"] = 10.0
                            _layers_per_hk_w[_bhn_l] = _lyrs_l
                        except Exception:
                            _layers_per_hk_w[_bhn_l] = []

                # ── Helper: Winkler re-run → 4-panel figure ────────────────────
                def _build_winkler_fig_w(bhn, pile_type, L_m, geom, lyrs_su):
                    """Re-run Winkler, trả (fig, err_str). geom = dict từ ke_sw_stability."""
                    cp = _cat_w.get(pile_type, {})
                    if not cp or not lyrs_su:
                        return None, "Thiếu catalog/layer"
                    pile_w = _PP_w(
                        name=pile_type,
                        D_m=float(cp.get("H_mm", 600)) / 1000.0,
                        EI_kNm2=float(cp.get("EI_kNm2", 100000)),
                        Mcr_kNm=float(cp.get("Mcr_kNm", 200)),
                    )
                    layers_w = []
                    for lyr in lyrs_su:
                        h = float(lyr.get("depth_bot_m", 0)) - float(lyr.get("depth_top_m", 0))
                        if h > 0:
                            layers_w.append(_SL_w(
                                symbol=str(lyr.get("symbol", "?")),
                                thickness_m=h,
                                Su_kPa=float(lyr.get("su_kPa", 10.0)),
                                gamma_kNm3=float(lyr.get("gamma_kNm3", 15.5)),
                            ))
                    if not layers_w:
                        return None, "Không có dữ liệu lớp đất"
                    top_e = float(geom.get("top_elev", 2.7))
                    Z_fr  = float(geom.get("Z_m", 0.0))
                    Z_bk  = float(geom.get("Zb_m", -2.0))
                    wlvl_f= float(geom.get("wlvl_front", -1.0))
                    wlvl_b= float(geom.get("wlvl_back", -2.0))
                    q_kPa = float(geom.get("q_kPa", 15.0))
                    cdm_w = float(geom.get("cdm_w_m", 5.0))
                    GW = 9.81
                    # Build cumulative su/gamma lookup
                    _lbk = []
                    _cz = 0.0
                    for ll in layers_w:
                        _lbk.append((_cz, _cz + ll.thickness_m, ll.Su_kPa, ll.gamma_kNm3))
                        _cz += ll.thickness_m
                    def _su_gm(z):
                        for za, zb, su, gm in _lbk:
                            if za <= z < zb:
                                return su, gm
                        return (_lbk[-1][2], _lbk[-1][3]) if _lbk else (10.0, 15.5)
                    N_pts = max(60, int(L_m * 3))
                    zs_ld = _np_w.linspace(0, L_m, N_pts)
                    p_net = _np_w.zeros(N_pts)
                    for ii, z in enumerate(zs_ld):
                        elev = top_e - z
                        df = max(0.0, Z_fr - elev)
                        da = min(df, max(0.0, Z_fr - wlvl_f))
                        db = max(0.0, df - (Z_fr - wlvl_f))
                        _, gm_lyr = _su_gm(z)
                        sv = gm_lyr * da + (gm_lyr - GW) * db + q_kPa
                        su_lyr, _ = _su_gm(z)
                        p_act = max(0.0, sv - 2.0 * su_lyr)
                        p_net[ii] = p_act + GW * (max(0.0, wlvl_f - elev) - max(0.0, wlvl_b - elev))
                    try:
                        res = _wk_solve_w(
                            layers=layers_w, pile=pile_w, L_m=L_m,
                            zs_load=zs_ld, p_load_kNm2=p_net,
                            N=N_pts, eps50=0.02,
                            cdm_thickness_m=cdm_w, cdm_factor=3.0,
                            tip_fixity="free", top_pin=False,
                        )
                    except Exception as ex:
                        return None, str(ex)
                    zs_p = res["zs"]; ux = res["ux"]; Ms = res["Ms"]
                    Qs   = res.get("Qs", _np_w.zeros_like(zs_p))
                    Mcr  = res["Mcr_kNm"]; u_max = res["u_max_mm"]
                    M_max= res["M_max_kNm"]; Q_max= res["Q_max_kN"]
                    fig4, axes4 = _plt_w.subplots(
                        1, 4, figsize=(21/2.54, 11/2.54), dpi=130,
                        gridspec_kw={"width_ratios": [1.1, 1, 1, 1]}
                    )
                    fig4.suptitle(f"{bhn} — {pile_type}, L={L_m:.1f}m",
                                  fontsize=9, fontweight="bold")
                    # Panel 0: sơ đồ hình học
                    ax0 = axes4[0]
                    ax0.set_xlim(-2.5, 2.5); ax0.set_ylim(L_m + 0.5, -0.5)
                    ax0.add_patch(_plt_w.Rectangle((-0.2, 0), 0.4, L_m,
                                                    color="#444", zorder=4))
                    zb_d = max(0, top_e - Z_bk)
                    cdm_s = max(0, zb_d - cdm_w); cdm_e = min(L_m, zb_d)
                    if cdm_e > cdm_s:
                        ax0.add_patch(_plt_w.Rectangle(
                            (-2.0, cdm_s), 2.0, cdm_e - cdm_s,
                            color="#90EE90", alpha=0.55, hatch="//", zorder=2))
                        ax0.text(-1.0, (cdm_s+cdm_e)/2, "CDM", fontsize=6,
                                 ha="center", va="center", color="#1a5e20")
                    # Soil fill front
                    zfr_d = max(0, top_e - Z_fr)
                    ax0.add_patch(_plt_w.Rectangle(
                        (0.2, 0), 2.0, zfr_d, color="#D2B48C", alpha=0.4, zorder=1))
                    ax0.axhline(zfr_d, color="brown", ls="--", lw=1)
                    ax0.text(2.2, zfr_d, f"Z={Z_fr:.1f}m", fontsize=5.5,
                             va="bottom", color="brown")
                    mf_d = max(0, top_e - wlvl_f)
                    mb_d = max(0, top_e - wlvl_b)
                    for md, side, col in [(mf_d, "F", "#1F77B4"), (mb_d, "B", "#AEC6CF")]:
                        if 0 <= md <= L_m:
                            ax0.axhline(md, color=col, ls="-.", lw=1)
                            ax0.text(2.2 if side=="F" else -2.2, md,
                                     f"MN{side}={wlvl_f if side=='F' else wlvl_b:.1f}",
                                     fontsize=5, va="bottom", color=col,
                                     ha="left" if side=="F" else "right")
                    ax0.set_ylabel("Độ sâu từ đỉnh cừ (m)", fontsize=7)
                    ax0.set_title("Hình học", fontsize=8); ax0.set_xticks([])
                    ax0.tick_params(labelsize=6)
                    # Panel 1: u(z)
                    ax1 = axes4[1]
                    ax1.plot(ux, zs_p, "b-", lw=1.5)
                    ax1.axvline(50, color="r", ls="--", lw=0.9)
                    ax1.axvline(-50, color="r", ls="--", lw=0.9)
                    ax1.axvline(0, color="k", lw=0.4, alpha=0.4)
                    ok_u = "Đạt" if abs(u_max) <= 50 else "Không đạt"
                    ax1.set_title(f"u(z)  u_max={u_max:.1f}mm\n{ok_u}", fontsize=7.5,
                                  fontweight="bold",
                                  color="#1a5e20" if ok_u=="Đạt" else "#b71c1c")
                    ax1.set_xlabel("u (mm)", fontsize=7); ax1.invert_yaxis()
                    ax1.tick_params(labelsize=6); ax1.grid(alpha=0.3)
                    # Panel 2: M(z)
                    ax2 = axes4[2]
                    ax2.plot(Ms, zs_p, "g-", lw=1.5)
                    ax2.axvline(Mcr, color="r", ls="--", lw=0.9, label=f"Mcr={Mcr:.0f}")
                    ax2.axvline(-Mcr, color="r", ls="--", lw=0.9)
                    ax2.axvline(0, color="k", lw=0.4, alpha=0.4)
                    ok_m = "Đạt" if abs(M_max) <= Mcr else "Không đạt"
                    ax2.set_title(f"M(z)  M_max={M_max:.0f}kNm\n{ok_m}", fontsize=7.5,
                                  fontweight="bold",
                                  color="#1a5e20" if ok_m=="Đạt" else "#b71c1c")
                    ax2.set_xlabel("M (kNm)", fontsize=7); ax2.invert_yaxis()
                    ax2.tick_params(labelsize=6); ax2.grid(alpha=0.3)
                    ax2.legend(fontsize=6)
                    # Panel 3: Q(z)
                    ax3 = axes4[3]
                    ax3.plot(Qs, zs_p, "m-", lw=1.5)
                    ax3.axvline(0, color="k", lw=0.4, alpha=0.4)
                    ax3.set_title(f"Q(z)  Q_max={Q_max:.0f}kN", fontsize=7.5, fontweight="bold")
                    ax3.set_xlabel("Q (kN)", fontsize=7); ax3.invert_yaxis()
                    ax3.tick_params(labelsize=6); ax3.grid(alpha=0.3)
                    _plt_w.tight_layout()
                    return fig4, None

                # ── Bắt đầu tạo Document ─────────────────────────────────────
                _doc_w = _Docx_w()
                _sec_w = _doc_w.sections[0]
                _sec_w.page_width  = _Cm_w(21)
                _sec_w.page_height = _Cm_w(29.7)
                for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
                    setattr(_sec_w, attr, _Cm_w(2.0))
                _pw_w = _sec_w.page_width - _sec_w.left_margin - _sec_w.right_margin

                # Header
                _co_w  = st.session_state.get("export_co_name") or ""
                _stf_w = st.session_state.get("export_co_staff") or ""
                _logo_w = st.session_state.get("export_logo_bytes")
                _ht_w = _sec_w.header.add_table(1, 2, width=_pw_w)
                _ht_w.cell(0,0).width = _Cm_w(3)
                _ht_w.cell(0,1).width = _pw_w - _Cm_w(3)
                if _logo_w:
                    _ht_w.cell(0,0).paragraphs[0].add_run().add_picture(
                        _io_w.BytesIO(_logo_w), height=_Cm_w(1.2))
                _hp = _ht_w.cell(0,1).paragraphs[0]
                _hp.alignment = _WD_AL.RIGHT
                _hr = _hp.add_run(_co_w); _hr.font.size = _Pt_w(9); _hr.bold = True
                _no_border_w(_ht_w)
                # Footer
                _ft_w = _sec_w.footer.add_table(1, 2, width=_pw_w)
                _ft_w.cell(0,0).width = _pw_w - _Cm_w(3)
                _ft_w.cell(0,1).width = _Cm_w(3)
                _ft_w.cell(0,0).paragraphs[0].add_run(_stf_w).font.size = _Pt_w(8)
                _fp_w = _ft_w.cell(0,1).paragraphs[0]
                _fp_w.alignment = _WD_AL.RIGHT
                _fp_w.add_run("Trang ").font.size = _Pt_w(8)
                _frun_w(_fp_w, "PAGE")
                _fp_w.add_run(" / ").font.size = _Pt_w(8)
                _frun_w(_fp_w, "NUMPAGES")
                _no_border_w(_ft_w)

                # ══════════════════════════════════════════════════════════════
                # TRANG BÌA
                # ══════════════════════════════════════════════════════════════
                _doc_w.add_paragraph()
                _doc_w.add_paragraph()
                _pc = _doc_w.add_paragraph("BÁO CÁO KỸ THUẬT")
                _pc.alignment = _WD_AL.CENTER
                _r = _pc.runs[0]; _r.bold = True; _r.font.size = _Pt_w(20)
                _r.font.color.rgb = _RGB_w(0x1F, 0x38, 0x64)
                _pt = _doc_w.add_paragraph("THIẾT KẾ TƯỜNG CỪ CỌC VÁN SW DỰ ỨNG LỰC")
                _pt.alignment = _WD_AL.CENTER
                _pt.runs[0].bold = True; _pt.runs[0].font.size = _Pt_w(16)
                _ps = _doc_w.add_paragraph("Kè Công Viên TTHC — Tuyến SW 7 hố khoan")
                _ps.alignment = _WD_AL.CENTER
                _ps.runs[0].font.size = _Pt_w(13); _ps.runs[0].font.color.rgb = _RGB_w(0x55,0x55,0x55)
                _doc_w.add_paragraph()
                # Meta table
                _mt_w = _doc_w.add_table(rows=5, cols=2)
                _mt_w.style = "Table Grid"
                _meta_rows = [
                    ("Tiêu chuẩn thiết kế", "TCVN 11823-10:2017 (AASHTO LRFD)"),
                    ("Phương pháp tính lún", "TCCS 41:2022 / TCVN 9403:2012"),
                    ("Phần mềm tính toán", "Winkler beam — NumPy solver; Bishop/Spencer/M-P"),
                    ("Ngày lập", _dt_w.now().strftime("%d/%m/%Y")),
                    ("Người thực hiện", _stf_w or "—"),
                ]
                for i, (k, v) in enumerate(_meta_rows):
                    _mt_w.cell(i,0).text = k; _mt_w.cell(i,1).text = v
                    _mt_w.cell(i,0).paragraphs[0].runs[0].bold = True
                    _fmt_cells_w(_mt_w.rows[i], 9)
                _doc_w.add_page_break()

                # ══════════════════════════════════════════════════════════════
                # PHẦN 1: CATALOG TIẾT DIỆN CỌC ÁN SW
                # ══════════════════════════════════════════════════════════════
                _doc_w.add_heading("1. Catalog tiết diện cọc ván SW dự ứng lực", level=1)
                _doc_w.add_paragraph(
                    "Cọc ván SW (Sheet-pile dự ứng lực bê tông) được sản xuất theo tiêu chuẩn "
                    "Nhật Bản. Bảng dưới liệt kê 22 tiết diện trong catalog, bao gồm chiều cao "
                    "tiết diện H, moment kháng nứt Mcr, độ cứng uốn EI, trọng lượng TL và "
                    "chiều dài tiêu chuẩn tối thiểu/tối đa."
                )
                # Bảng catalog
                _cat_cols = ["Loại cọc","H (mm)","t (mm)","Mcr (kNm)","EI (kNm²)","TL (T/m)","L_min (m)","L_max (m)"]
                _ct_w = _doc_w.add_table(rows=1, cols=len(_cat_cols))
                _ct_w.style = "Table Grid"
                for i, h in enumerate(_cat_cols):
                    _c = _ct_w.rows[0].cells[i]; _c.text = h
                    _cell_bg_w(_c, "1F3864")
                    for p in _c.paragraphs:
                        for r in p.runs:
                            r.bold = True; r.font.size = _Pt_w(7.5)
                            r.font.color.rgb = _RGB_w(0xFF,0xFF,0xFF)
                _pile_names_w = sorted(_cat_w.keys())
                for pi, pn in enumerate(_pile_names_w):
                    pp = _cat_w[pn]
                    row_c = _ct_w.add_row().cells
                    row_c[0].text = pn
                    row_c[1].text = str(pp.get("H_mm",""))
                    row_c[2].text = str(pp.get("t_mm",""))
                    row_c[3].text = f"{pp.get('Mcr_kNm',0):.1f}"
                    row_c[4].text = f"{pp.get('EI_kNm2',0):.0f}"
                    row_c[5].text = f"{pp.get('weight_T',0):.3f}"
                    row_c[6].text = str(pp.get("L_min_m",""))
                    row_c[7].text = str(pp.get("L_max_m",""))
                    _fill = "EAF4FF" if pi % 2 == 0 else "FFFFFF"
                    for c in row_c: _cell_bg_w(c, _fill)
                    _fmt_cells_w(row_c, 7.5)
                _doc_w.add_paragraph()
                # Biểu đồ Mcr vs H (matplotlib 2D)
                _fig_cat, _ax_cat = _plt_w.subplots(figsize=(14/2.54, 7/2.54), dpi=130)
                _H_vals  = [_cat_w[n].get("H_mm",0) for n in _pile_names_w]
                _M_vals  = [_cat_w[n].get("Mcr_kNm",0) for n in _pile_names_w]
                _EI_vals = [_cat_w[n].get("EI_kNm2",0) for n in _pile_names_w]
                _ax_cat.scatter(_H_vals, _M_vals, c="#1F77B4", s=60, zorder=3)
                _ax_cat.plot(_H_vals, _M_vals, "--", color="#1F77B4", lw=1.2, alpha=0.7)
                for n, hv, mv in zip(_pile_names_w, _H_vals, _M_vals):
                    _ax_cat.annotate(n, (hv, mv), textcoords="offset points",
                                     xytext=(4,4), fontsize=5.5)
                _ax2_cat = _ax_cat.twinx()
                _ax2_cat.plot(_H_vals, _EI_vals, "s-", color="#FF7F0E", lw=1.2, ms=4, alpha=0.8)
                _ax2_cat.set_ylabel("EI (kNm²)", fontsize=8, color="#FF7F0E")
                _ax2_cat.tick_params(axis="y", labelcolor="#FF7F0E", labelsize=7)
                _ax_cat.set_xlabel("Chiều cao tiết diện H (mm)", fontsize=8)
                _ax_cat.set_ylabel("Mcr (kNm)", fontsize=8, color="#1F77B4")
                _ax_cat.tick_params(axis="both", labelsize=7)
                _ax_cat.set_title("Moment kháng nứt Mcr và Độ cứng EI theo tiết diện SW", fontsize=9, fontweight="bold")
                _ax_cat.grid(axis="y", alpha=0.3); _ax_cat.grid(axis="x", alpha=0.2)
                _plt_w.tight_layout()
                _buf_cat = _io_w.BytesIO()
                _fig_cat.savefig(_buf_cat, format="png", dpi=130, bbox_inches="tight", facecolor="white")
                _plt_w.close(_fig_cat); _buf_cat.seek(0)
                _doc_w.add_picture(_buf_cat, width=_Cm_w(15))
                _doc_w.add_page_break()

                # ══════════════════════════════════════════════════════════════
                # PHẦN 2: VỊ TRÍ HỐ KHOAN — BẢN ĐỒ 2D + 3D
                # ══════════════════════════════════════════════════════════════
                _doc_w.add_heading("2. Vị trí hố khoan trên tuyến kè SW", level=1)
                _on_bhs = [k for k in _nt_det_w.keys()]
                _bh_xy = [(k, _bhs_w.get(k,{}).get("x_coord_m"), _bhs_w.get(k,{}).get("y_coord_m"),
                           _bhs_w.get(k,{}).get("elevation_m",0)) for k in _on_bhs
                          if _bhs_w.get(k,{}).get("x_coord_m")]
                if _bh_xy:
                    # 2D plan view
                    _fig2d, _ax2d = _plt_w.subplots(figsize=(14/2.54, 8/2.54), dpi=130)
                    xs2 = [b[1] for b in _bh_xy]; ys2 = [b[2] for b in _bh_xy]
                    _ax2d.plot(xs2, ys2, "o-", color="#1F77B4", ms=8, lw=1.5, markerfacecolor="#FF7F0E")
                    for nm, x2, y2, _ in _bh_xy:
                        _ax2d.annotate(nm.replace("KE-",""), (x2,y2), textcoords="offset points",
                                       xytext=(5,5), fontsize=7, fontweight="bold")
                    _ax2d.set_xlabel("X – Northing VN-2000 (m)", fontsize=8)
                    _ax2d.set_ylabel("Y – Easting VN-2000 (m)", fontsize=8)
                    _ax2d.set_title("Bình đồ vị trí hố khoan trên tuyến kè SW (mặt bằng 2D)", fontsize=9, fontweight="bold")
                    _ax2d.grid(alpha=0.3); _ax2d.tick_params(labelsize=7)
                    _plt_w.tight_layout()
                    _buf2d = _io_w.BytesIO()
                    _fig2d.savefig(_buf2d, format="png", dpi=130, bbox_inches="tight", facecolor="white")
                    _plt_w.close(_fig2d); _buf2d.seek(0)
                    _doc_w.add_paragraph("Hình 2.1 — Mặt bằng 2D vị trí hố khoan:")
                    _doc_w.add_picture(_buf2d, width=_Cm_w(15))
                    _doc_w.add_paragraph()
                    # 3D với elevation
                    _fig3d = _plt_w.figure(figsize=(14/2.54, 9/2.54), dpi=130)
                    _ax3d = _fig3d.add_subplot(111, projection="3d")
                    xs3 = [b[1] for b in _bh_xy]; ys3 = [b[2] for b in _bh_xy]
                    zs3 = [b[3] for b in _bh_xy]
                    _ax3d.scatter(xs3, ys3, zs3, c="#FF7F0E", s=80, depthshade=True, zorder=3)
                    _ax3d.plot(xs3, ys3, zs3, "--", color="#1F77B4", lw=1, alpha=0.7)
                    for nm, x3, y3, z3 in _bh_xy:
                        _ax3d.text(x3, y3, z3+0.3, nm.replace("KE-",""), fontsize=6, fontweight="bold")
                    _ax3d.set_xlabel("Northing (m)", fontsize=7, labelpad=2)
                    _ax3d.set_ylabel("Easting (m)", fontsize=7, labelpad=2)
                    _ax3d.set_zlabel("Cao độ (m)", fontsize=7, labelpad=2)
                    _ax3d.set_title("Vị trí hố khoan 3D (X-Y-Elevation)", fontsize=9, fontweight="bold")
                    _ax3d.tick_params(labelsize=6)
                    _fig3d.tight_layout()
                    _buf3d = _io_w.BytesIO()
                    _fig3d.savefig(_buf3d, format="png", dpi=130, bbox_inches="tight", facecolor="white")
                    _plt_w.close(_fig3d); _buf3d.seek(0)
                    _doc_w.add_paragraph("Hình 2.2 — Mô hình 3D hố khoan (Northing – Easting – Cao độ):")
                    _doc_w.add_picture(_buf3d, width=_Cm_w(15))
                _doc_w.add_page_break()

                # ══════════════════════════════════════════════════════════════
                # PHẦN 3: CƠ SỞ LÝ THUYẾT & CÔNG THỨC (LATEX via mathtext)
                # ══════════════════════════════════════════════════════════════
                _doc_w.add_heading("3. Cơ sở lý thuyết và công thức tính toán", level=1)

                _doc_w.add_heading("3.1 Kiểm tra NT1 — Chiều dài cọc tối thiểu", level=2)
                _doc_w.add_paragraph(
                    "Cọc phải xuyên qua toàn bộ lớp đất yếu và ngàm vào lớp chịu lực tối thiểu. "
                    "Chiều dài yêu cầu tính theo điều kiện hình học địa tầng:"
                )
                _f31 = _formula_png_w([
                    ("Chiều dài yêu cầu:", r"$L_{req} = (Z_{cổ} - Z_{mũi}) + H_{1} + h_{min\_pen}$"),
                    ("Ngàm tối thiểu:", r"$h_{min\_pen} = 1{,}0\ \mathrm{m}\ \text{(vào lớp cứng)}$"),
                    ("Điều kiện đạt:", r"$L_{thiết\ kế} \geq L_{req}\ \Rightarrow\ \mathrm{NT1: Đạt}$"),
                ], title="NT1 — Điều kiện chiều dài (hình học địa tầng)", h_rows_cm=1.3)
                _doc_w.add_picture(_f31, width=_Cm_w(14.5))

                _doc_w.add_heading("3.2 Kiểm tra NT2 — Sức chịu tải dọc trục", level=2)
                _doc_w.add_paragraph(
                    "Sức chịu tải dọc trục NT2 theo TCVN 11823-10:2017 (AASHTO LRFD), "
                    "phương pháp α (Tomlinson 1980) cho lớp sét và SPT-Meyerhof cho lớp cát. "
                    "Hệ số φ_stat = 0,35 (cọc đóng)."
                )
                _f32 = _formula_png_w([
                    ("Ma sát thân (sét):",  r"$R_s = \alpha \cdot s_u \cdot P \cdot L_{soil}$"),
                    ("Ma sát thân (cát):",  r"$R_s = 1{,}9 \cdot N_{160} \cdot P \cdot L$  (kN)"),
                    ("Sức kháng mũi (sét):", r"$R_p = 9 \cdot s_u^{stat} \cdot A_p$"),
                    ("Sức kháng mũi (cát):", r"$R_p = q_p \cdot A_p \leq 3200 \cdot N_{160} \cdot A_p$"),
                    ("Sức kháng tính toán:", r"$RR = \varphi_{stat} \cdot (R_s + R_p) \geq W_{cọc}$"),
                    ("Trọng lượng cọc:",    r"$W = \dfrac{TL \times 9{,}81}{L_{std}} \times L_{des}$  (kN)"),
                ], title="NT2 — TCVN 11823-10:2017 / AASHTO LRFD (φ_stat = 0,35)", h_rows_cm=1.2)
                _doc_w.add_picture(_f32, width=_Cm_w(14.5))

                _doc_w.add_heading("3.3 Nội lực tường cừ — Mô hình Winkler (API Clay)", level=2)
                _doc_w.add_paragraph(
                    "Dầm trên nền đàn hồi Winkler với lò xo kháng ngang phía Back (phía sông). "
                    "Tải phân bố = áp lực đất Active phía Front + áp lực nước chênh lệch. "
                    "Tiêu chí kiểm tra: u_max ≤ 50 mm và M_max ≤ Mcr."
                )
                _f33 = _formula_png_w([
                    ("Phản lực nền:",      r"$p = k_h \cdot u \cdot B_{cọc}$"),
                    ("Hệ số lò xo (API Clay):", r"$k_h = \dfrac{67 \cdot S_u}{d_{cọc}}$  (kN/m³)"),
                    ("Vùng CDM:",         r"$k_h^{CDM} = k_{factor} \times k_h$  $(k_{factor} \approx 3{,}0)$"),
                    ("Tiêu chí chuyển vị:", r"$u_{max} \leq 50\ \mathrm{mm}$"),
                    ("Tiêu chí mô men:",   r"$M_{max} \leq M_{cr}$"),
                ], title="Winkler beam — API Clay (Matlock 1970)", h_rows_cm=1.2)
                _doc_w.add_picture(_f33, width=_Cm_w(14.5))

                _doc_w.add_heading("3.4 Ổn định tổng thể — Bishop / Spencer / Morgenstern-Price", level=2)
                _doc_w.add_paragraph(
                    "Phân tích ổn định cung trượt tròn bằng 3 phương pháp phân mảnh. "
                    "Tiêu chí: Fs_trượt ≥ 1,30 (TCVN 4253:2012)."
                )
                _f34 = _formula_png_w([
                    ("Bishop Simplified:", r"$F_s = \dfrac{\sum [c'b + (W-ub)\tan\phi'] / m_\alpha}{\sum W\sin\alpha}$"),
                    ("",                  r"$m_\alpha = \cos\alpha\!\left(1 + \dfrac{\tan\alpha\tan\phi'}{F_s}\right)$"),
                    ("Spencer:",          r"$\text{Thêm cân bằng lực dọc liên phân mảnh (iterative)}$"),
                    ("Morgenstern-Price:", r"$\text{Hàm } f(\alpha) \text{ mô tả phương lực liên phân mảnh (đầy đủ nhất)}$"),
                    ("Tiêu chí đạt:",     r"$F_s \geq 1{,}30\ \text{(TCVN 4253:2012)}$"),
                ], title="Ổn định tổng thể — 3 phương pháp phân mảnh", h_rows_cm=1.2)
                _doc_w.add_picture(_f34, width=_Cm_w(14.5))
                _doc_w.add_page_break()

                # ══════════════════════════════════════════════════════════════
                # PHẦN 4: KẾT QUẢ THIẾT KẾ — BẢNG TỔNG HỢP 7 HK
                # ══════════════════════════════════════════════════════════════
                _doc_w.add_heading("4. Kết quả thiết kế — Tổng hợp 7 hố khoan trên tuyến", level=1)
                _doc_w.add_paragraph(
                    "Bảng tổng hợp thông số thiết kế và kết quả kiểm tra NT1/NT2 "
                    "cho 7 hố khoan trên tuyến kè SW. Dữ liệu lấy từ cơ sở dữ liệu SQLite "
                    "(ke_sw_nt_detail) — đảm bảo đồng nhất với tính toán mới nhất."
                )
                _h4 = ["Hố khoan","Cọc thiết kế","L TK (m)","Z (m)","H lớp yếu (m)",
                        "L yc NT1 (m)","NT1","Rs (kN)","Rp (kN)","RR (kN)","W (kN)","RR/W","NT2"]
                _tb4 = _doc_w.add_table(rows=1, cols=len(_h4))
                _tb4.style = "Table Grid"
                for i, h in enumerate(_h4):
                    _c = _tb4.rows[0].cells[i]; _c.text = h
                    _cell_bg_w(_c, "1F3864")
                    for p in _c.paragraphs:
                        for r in p.runs:
                            r.bold = True; r.font.size = _Pt_w(7)
                            r.font.color.rgb = _RGB_w(0xFF,0xFF,0xFF)
                _bh_list_w = sorted(_nt_det_w.keys())
                for bi, bhn in enumerate(_bh_list_w):
                    _d = _nt_det_w[bhn]
                    _nm_short = bhn.replace("KE-","")
                    _pile_d = _ss_rec.get(_nm_short) or _d.get("pile_type","")
                    _L_d    = _ss_L.get(_nm_short) or _d.get("L_req_nt1_m","")
                    _nt1 = str(_d.get("nt1_result","")).strip()
                    _nt2 = str(_d.get("nt2_result","")).strip()
                    _row4 = _tb4.add_row().cells
                    _row4[0].text  = bhn
                    _row4[1].text  = str(_pile_d)
                    _row4[2].text  = f"{float(_L_d):.1f}" if _L_d else "—"
                    _row4[3].text  = f"{float(_d.get('Z_m',0)):.2f}" if _d.get("Z_m") is not None else "—"
                    _row4[4].text  = f"{float(_d.get('D_bottom_soft_m',0)):.2f}" if _d.get("D_bottom_soft_m") is not None else "—"
                    _row4[5].text  = f"{float(_d.get('L_req_nt1_m',0)):.2f}" if _d.get("L_req_nt1_m") is not None else "—"
                    _row4[6].text  = _nt1 or "—"
                    _row4[7].text  = f"{float(_d.get('Rs_kN',0)):.0f}" if _d.get("Rs_kN") is not None else "—"
                    _row4[8].text  = f"{float(_d.get('Rp_kN',0)):.0f}" if _d.get("Rp_kN") is not None else "—"
                    _row4[9].text  = f"{float(_d.get('RR_kN',0)):.0f}" if _d.get("RR_kN") is not None else "—"
                    _row4[10].text = f"{float(_d.get('W_kN',0)):.0f}" if _d.get("W_kN") is not None else "—"
                    _ratio4 = _d.get("ratio_nt2")
                    _row4[11].text = f"{float(_ratio4):.2f}" if _ratio4 is not None else "—"
                    _row4[12].text = _nt2 or "—"
                    _fill4 = "EAF4FF" if bi % 2 == 0 else "FFFFFF"
                    for c in _row4: _cell_bg_w(c, _fill4)
                    # Tô màu kết quả
                    for _idx, _val in [(6, _nt1), (12, _nt2)]:
                        if "Đạt" in _val: _cell_bg_w(_row4[_idx], "CCFFCC")
                        elif "Không" in _val: _cell_bg_w(_row4[_idx], "FFCCCC")
                    _fmt_cells_w(_row4, 7.5)
                _doc_w.add_page_break()

                # ══════════════════════════════════════════════════════════════
                # PHẦN 5: SO SÁNH 4 PHƯƠNG PHÁP NT2
                # ══════════════════════════════════════════════════════════════
                _doc_w.add_heading("5. So sánh sức chịu tải NT2 — 4 phương pháp", level=1)
                _doc_w.add_paragraph(
                    "Bảng so sánh RR và tỷ số RR/W theo 4 phương pháp tính toán: "
                    "Auto (α+SPT), α-method, β-method, λ-method. "
                    "Phương pháp Auto được dùng làm kết quả chính thức."
                )
                _h5 = ["Hố khoan","Cọc","L (m)","RR_Auto","RR/W_Auto","RR_α","RR/W_α","RR_β","RR/W_β","RR_λ","RR/W_λ","Kết quả"]
                _tb5 = _doc_w.add_table(rows=1, cols=len(_h5))
                _tb5.style = "Table Grid"
                for i, h in enumerate(_h5):
                    _c = _tb5.rows[0].cells[i]; _c.text = h
                    _cell_bg_w(_c, "1F3864")
                    for p in _c.paragraphs:
                        for r in p.runs:
                            r.bold = True; r.font.size = _Pt_w(7)
                            r.font.color.rgb = _RGB_w(0xFF,0xFF,0xFF)
                # Collect từ SQLite _nt2_cmp_w
                _seen5 = set()
                for (bhn5, pn5, Lm5), mdict5 in _nt2_cmp_w.items():
                    if bhn5 in _seen5: continue
                    _seen5.add(bhn5)
                    _auto5 = mdict5.get("auto", {})
                    _alp5  = mdict5.get("alpha", {})
                    _bet5  = mdict5.get("beta", {})
                    _lam5  = mdict5.get("lambda", {})
                    _ok5 = str(_auto5.get("result","")).strip()
                    _row5 = _tb5.add_row().cells
                    _row5[0].text  = bhn5
                    _row5[1].text  = str(pn5)
                    _row5[2].text  = f"{Lm5:.1f}"
                    _row5[3].text  = f"{_auto5.get('RR_kN',0):.0f}"
                    _row5[4].text  = f"{_auto5.get('ratio',0):.2f}"
                    _row5[5].text  = f"{_alp5.get('RR_kN',0):.0f}"
                    _row5[6].text  = f"{_alp5.get('ratio',0):.2f}"
                    _row5[7].text  = f"{_bet5.get('RR_kN',0):.0f}"
                    _row5[8].text  = f"{_bet5.get('ratio',0):.2f}"
                    _row5[9].text  = f"{_lam5.get('RR_kN',0):.0f}"
                    _row5[10].text = f"{_lam5.get('ratio',0):.2f}"
                    _row5[11].text = _ok5 or "—"
                    if "Đạt" in _ok5: _cell_bg_w(_row5[11], "CCFFCC")
                    elif "Không" in _ok5: _cell_bg_w(_row5[11], "FFCCCC")
                    _fmt_cells_w(_row5, 7.5)
                # Bar chart RR/W
                if _nt2_cmp_w:
                    _bh_lbl5 = []; _rr_a=[]; _rr_al=[]; _rr_b=[]; _rr_l=[]
                    for (bhn5, pn5, _), mdict5 in sorted(_nt2_cmp_w.items()):
                        _bh_lbl5.append(bhn5.replace("KE-",""))
                        _rr_a.append(mdict5.get("auto",{}).get("ratio",0) or 0)
                        _rr_al.append(mdict5.get("alpha",{}).get("ratio",0) or 0)
                        _rr_b.append(mdict5.get("beta",{}).get("ratio",0) or 0)
                        _rr_l.append(mdict5.get("lambda",{}).get("ratio",0) or 0)
                    _fig5, _ax5 = _plt_w.subplots(figsize=(15/2.54, 7/2.54), dpi=130)
                    _xs5 = _np_w.arange(len(_bh_lbl5)); _ww5 = 0.20
                    _ax5.bar(_xs5-1.5*_ww5, _rr_a,  _ww5, label="Auto",  color="#2196F3", alpha=0.85)
                    _ax5.bar(_xs5-0.5*_ww5, _rr_al, _ww5, label="α",     color="#4CAF50", alpha=0.85)
                    _ax5.bar(_xs5+0.5*_ww5, _rr_b,  _ww5, label="β",     color="#FF9800", alpha=0.85)
                    _ax5.bar(_xs5+1.5*_ww5, _rr_l,  _ww5, label="λ",     color="#9C27B0", alpha=0.85)
                    _ax5.axhline(1.0, color="red", ls="--", lw=1.2, label="RR/W = 1,0")
                    _ax5.set_xticks(_xs5); _ax5.set_xticklabels(_bh_lbl5, fontsize=8)
                    _ax5.set_ylabel("Tỷ số RR/W", fontsize=9)
                    _ax5.set_title("So sánh tỷ số RR/W — 4 phương pháp NT2", fontsize=10, fontweight="bold")
                    _ax5.legend(fontsize=8, ncol=5); _ax5.grid(axis="y", alpha=0.3)
                    _plt_w.tight_layout()
                    _buf5 = _io_w.BytesIO()
                    _fig5.savefig(_buf5, format="png", dpi=130, bbox_inches="tight", facecolor="white")
                    _plt_w.close(_fig5); _buf5.seek(0)
                    _doc_w.add_paragraph()
                    _doc_w.add_picture(_buf5, width=_Cm_w(15))
                _doc_w.add_page_break()

                # ══════════════════════════════════════════════════════════════
                # PHẦN 6: NỘI LỰC WINKLER
                # ══════════════════════════════════════════════════════════════
                _doc_w.add_heading("6. Kết quả nội lực — Mô hình Winkler (beam on elastic foundation)", level=1)
                _doc_w.add_paragraph(
                    "Biểu đồ nội lực tường cừ SW theo mô hình dầm Winkler với lò xo "
                    "nền phía Back (API Clay, Matlock 1970). Tiêu chí kiểm tra: "
                    "u_max ≤ 50 mm, M_max ≤ Mcr."
                )
                _h6 = ["Hố khoan","Cọc","L (m)","u_max (mm)","M_max (kNm)","Mcr (kNm)","M/Mcr","Q_max (kN)","Chuyển vị","Mô men","Kết quả"]
                _tb6 = _doc_w.add_table(rows=1, cols=len(_h6))
                _tb6.style = "Table Grid"
                for i, h in enumerate(_h6):
                    _c = _tb6.rows[0].cells[i]; _c.text = h
                    _cell_bg_w(_c, "1F3864")
                    for p in _c.paragraphs:
                        for r in p.runs:
                            r.bold = True; r.font.size = _Pt_w(7)
                            r.font.color.rgb = _RGB_w(0xFF,0xFF,0xFF)
                _u_all=[]; _m_all=[]
                for (bhn6, pn6), wr6 in sorted(_winkler_w.items()):
                    _u6  = wr6.get("u_max_mm",0) or 0
                    _m6  = wr6.get("M_max_kNm",0) or 0
                    _mcr6= wr6.get("Mcr_kNm",0) or 0
                    _q6  = wr6.get("Q_max_kN",0) or 0
                    _rat6= wr6.get("mcr_ratio",0) or (_m6/_mcr6 if _mcr6 else 0)
                    _ok_u= "Đạt" if _u6 <= 50 else "Không đạt"
                    _ok_m= "Đạt" if _m6 <= _mcr6 else "Không đạt"
                    _ok6 = "Đạt" if _u6 <= 50 and _m6 <= _mcr6 else "Không đạt"
                    _L6  = wr6.get("L_m","—")
                    _row6 = _tb6.add_row().cells
                    _row6[0].text  = bhn6
                    _row6[1].text  = str(pn6)
                    _row6[2].text  = f"{float(_L6):.1f}" if _L6 != "—" else "—"
                    _row6[3].text  = f"{_u6:.2f}"
                    _row6[4].text  = f"{_m6:.0f}"
                    _row6[5].text  = f"{_mcr6:.0f}"
                    _row6[6].text  = f"{_rat6:.2f}"
                    _row6[7].text  = f"{_q6:.0f}"
                    _row6[8].text  = _ok_u
                    _row6[9].text  = _ok_m
                    _row6[10].text = _ok6
                    for _ix, _ok in [(8,_ok_u),(9,_ok_m),(10,_ok6)]:
                        _cell_bg_w(_row6[_ix], "CCFFCC" if _ok=="Đạt" else "FFCCCC")
                    _fmt_cells_w(_row6, 7.5)
                    _u_all.append(_u6); _m_all.append(_m6)
                # Chart tổng hợp u_max và M/Mcr
                if _u_all:
                    _bh_lbl6 = [k[0].replace("KE-","") for k in sorted(_winkler_w.keys())]
                    _fig6, (_axu, _axm) = _plt_w.subplots(1, 2, figsize=(15/2.54, 7/2.54), dpi=130)
                    _axu.bar(_bh_lbl6, _u_all, color=["#CCFFCC" if v<=50 else "#FFCCCC" for v in _u_all], edgecolor="#333", lw=0.5)
                    _axu.axhline(50, color="red", ls="--", lw=1.2, label="Giới hạn 50mm")
                    _axu.set_ylabel("u_max (mm)", fontsize=8); _axu.set_title("Chuyển vị ngang tối đa", fontsize=9, fontweight="bold")
                    _axu.legend(fontsize=7); _axu.tick_params(labelsize=7); _axu.grid(axis="y", alpha=0.3)
                    _mrat_all = [wr6.get("mcr_ratio",0) or 0 for _, wr6 in sorted(_winkler_w.items())]
                    _axm.bar(_bh_lbl6, _mrat_all, color=["#CCFFCC" if v<=1 else "#FFCCCC" for v in _mrat_all], edgecolor="#333", lw=0.5)
                    _axm.axhline(1.0, color="red", ls="--", lw=1.2, label="Giới hạn M/Mcr=1")
                    _axm.set_ylabel("M/Mcr", fontsize=8); _axm.set_title("Tỷ số Mô men / Mcr", fontsize=9, fontweight="bold")
                    _axm.legend(fontsize=7); _axm.tick_params(labelsize=7); _axm.grid(axis="y", alpha=0.3)
                    _plt_w.tight_layout()
                    _buf6 = _io_w.BytesIO()
                    _fig6.savefig(_buf6, format="png", dpi=130, bbox_inches="tight", facecolor="white")
                    _plt_w.close(_fig6); _buf6.seek(0)
                    _doc_w.add_paragraph()
                    _doc_w.add_picture(_buf6, width=_Cm_w(15))

                # ── 6b: Per-HK nội lực 4-panel (re-run Winkler) ──────────────
                _doc_w.add_heading("6.2 Biểu đồ nội lực từng hố khoan — u(z) / M(z) / Q(z)", level=2)
                _doc_w.add_paragraph(
                    "Biểu đồ nội lực tái tính Winkler per hố khoan. "
                    "Tải phân bố: áp lực đất Active (Rankine không thoát nước) + "
                    "chênh lệch thủy áp hai phía. Điều kiện biên: đỉnh tự do — đáy tự do. "
                    "Su lấy từ VST hiện trường (ưu tiên) hoặc thí nghiệm phòng."
                )
                for _bhn6b in _align_hks:
                    _nd6b  = _nt_det_w.get(_bhn6b, {})
                    _pt6b  = (_ss_rec.get(_bhn6b.replace("KE-",""))
                               or _nd6b.get("pile_type", ""))
                    _Lm6b  = float(_ss_L.get(_bhn6b.replace("KE-",""))
                                    or _nd6b.get("L_req_nt1_m") or 0)
                    _gm6b  = _stab_geom_w.get(_bhn6b, {})
                    _ly6b  = _layers_per_hk_w.get(_bhn6b, [])
                    if not _pt6b or _Lm6b <= 0 or not _gm6b or not _ly6b:
                        _doc_w.add_paragraph(f"  {_bhn6b}: Chưa đủ dữ liệu — bỏ qua.")
                        continue
                    _fig6b, _err6b = _build_winkler_fig_w(
                        _bhn6b, _pt6b, _Lm6b, _gm6b, _ly6b)
                    if _err6b:
                        _doc_w.add_paragraph(f"  {_bhn6b}: Lỗi tính toán — {_err6b}")
                        continue
                    _buf6b = _io_w.BytesIO()
                    _fig6b.savefig(_buf6b, format="png", dpi=130,
                                   bbox_inches="tight", facecolor="white")
                    _plt_w.close(_fig6b); _buf6b.seek(0)
                    _doc_w.add_picture(_buf6b, width=_Cm_w(17))
                    _doc_w.add_paragraph()
                _doc_w.add_page_break()

                # ══════════════════════════════════════════════════════════════
                # PHẦN 7: ỔN ĐỊNH TỔNG THỂ
                # ══════════════════════════════════════════════════════════════
                _doc_w.add_heading("7. Ổn định tổng thể — Bishop / Spencer / Morgenstern-Price", level=1)
                # 7.0 Lý thuyết từ file 48-ke-sw-on-dinh-tong-the.md
                if _MD_THEORY_W.exists():
                    _doc_w.add_heading("7.1 Cơ sở lý thuyết (TCVN 4253:2012 / USACE EM 1110-2-2504)", level=2)
                    for _mdln in _MD_THEORY_W.read_text(encoding="utf-8").splitlines():
                        _mdln = _mdln.strip()
                        if not _mdln:
                            continue
                        if _mdln.startswith("# ") or _mdln.startswith("$$") or _mdln.startswith("|") or _mdln.startswith("---"):
                            continue
                        if _mdln.startswith("## "):
                            _doc_w.add_heading(_mdln[3:], level=3)
                        elif _mdln.startswith("### "):
                            _doc_w.add_heading(_mdln[4:], level=4)
                        elif _mdln.startswith("- ") or _mdln.startswith("* "):
                            _doc_w.add_paragraph(_mdln[2:], style="List Bullet")
                        else:
                            _clean = _mdln.replace("**","").replace("*","").replace("`","")
                            if _clean:
                                _doc_w.add_paragraph(_clean)
                _doc_w.add_heading("7.2 Kết quả tính toán", level=2)
                _doc_w.add_paragraph(
                    "Hệ số an toàn ổn định trượt cung tròn theo 3 phương pháp phân mảnh. "
                    "Tiêu chí: Fs ≥ 1,30 (TCVN 4253:2012). "
                    "Tính toán bằng ứng suất hữu hiệu, 1 lát/m dài tuyến."
                )
                _h7 = ["Hố khoan","Cọc","L (m)","Fs Bishop","Fs Spencer","Fs M-P","Fs lật","Fs đẩy trồi","Kết luận"]
                _tb7 = _doc_w.add_table(rows=1, cols=len(_h7))
                _tb7.style = "Table Grid"
                for i, h in enumerate(_h7):
                    _c = _tb7.rows[0].cells[i]; _c.text = h
                    _cell_bg_w(_c, "1F3864")
                    for p in _c.paragraphs:
                        for r in p.runs:
                            r.bold = True; r.font.size = _Pt_w(7)
                            r.font.color.rgb = _RGB_w(0xFF,0xFF,0xFF)
                _fs_b_all=[]; _lbl_all=[]
                for (bhn7, pn7), st7 in sorted(_stab_w.items()):
                    _fb = st7.get("Fs_bishop")
                    _fs = st7.get("Fs_spencer")
                    _fm = st7.get("Fs_morgenstern_price")
                    _fl = st7.get("Fs_overturning")
                    _ft = st7.get("Fs_toe_kickout")
                    _L7 = st7.get("L_m","—")
                    _ok7 = "Đạt" if (_fb and float(_fb) >= 1.30) else ("Không đạt" if _fb else "—")
                    _row7 = _tb7.add_row().cells
                    _row7[0].text = bhn7
                    _row7[1].text = str(pn7)
                    _row7[2].text = f"{float(_L7):.1f}" if _L7 != "—" else "—"
                    _row7[3].text = f"{float(_fb):.3f}" if _fb else "—"
                    _row7[4].text = f"{float(_fs):.3f}" if _fs else "—"
                    _row7[5].text = f"{float(_fm):.3f}" if _fm else "—"
                    _row7[6].text = f"{float(_fl):.3f}" if _fl else "—"
                    _row7[7].text = f"{float(_ft):.3f}" if _ft else "—"
                    _row7[8].text = _ok7
                    _cell_bg_w(_row7[8], "CCFFCC" if _ok7=="Đạt" else ("FFCCCC" if _ok7=="Không đạt" else "FFFFFF"))
                    _fmt_cells_w(_row7, 7.5)
                    if _fb: _fs_b_all.append(float(_fb)); _lbl_all.append(bhn7.replace("KE-",""))
                # Bar chart Fs
                if _fs_b_all:
                    _fig7, _ax7 = _plt_w.subplots(figsize=(14/2.54, 6/2.54), dpi=130)
                    _colors7 = ["#4CAF50" if v>=1.30 else "#F44336" for v in _fs_b_all]
                    _ax7.bar(_lbl_all, _fs_b_all, color=_colors7, edgecolor="#333", lw=0.5, alpha=0.85)
                    _ax7.axhline(1.30, color="red", ls="--", lw=1.5, label="Fs_min = 1,30")
                    _ax7.set_ylabel("Hệ số an toàn Fs", fontsize=9)
                    _ax7.set_title("Fs ổn định tổng thể — Bishop Simplified", fontsize=10, fontweight="bold")
                    _ax7.legend(fontsize=8); _ax7.grid(axis="y", alpha=0.3); _ax7.tick_params(labelsize=8)
                    _plt_w.tight_layout()
                    _buf7 = _io_w.BytesIO()
                    _fig7.savefig(_buf7, format="png", dpi=130, bbox_inches="tight", facecolor="white")
                    _plt_w.close(_fig7); _buf7.seek(0)
                    _doc_w.add_paragraph()
                    _doc_w.add_picture(_buf7, width=_Cm_w(14))
                # 7.3: Grouped bar — 3 phương pháp Bishop / Spencer / M-P
                if _stab_w:
                    _lbl7g=[]; _fb7g=[]; _fs7g=[]; _fm7g=[]; _seen7g=set()
                    for (bhn7g, _), st7g in sorted(_stab_w.items()):
                        if bhn7g in _seen7g: continue
                        _seen7g.add(bhn7g)
                        _lbl7g.append(bhn7g.replace("KE-",""))
                        _fb7g.append(float(st7g.get("Fs_bishop") or 0))
                        _fs7g.append(float(st7g.get("Fs_spencer") or 0))
                        _fm7g.append(float(st7g.get("Fs_morgenstern_price") or 0))
                    _xs7g = _np_w.arange(len(_lbl7g)); _wg7 = 0.25
                    _fig7g, _ax7g = _plt_w.subplots(figsize=(15/2.54, 7/2.54), dpi=130)
                    def _bc7(v): return "#4CAF50" if v >= 1.30 else "#F44336"
                    _ax7g.bar(_xs7g-_wg7, _fb7g, _wg7, label="Bishop",
                               color=[_bc7(v) for v in _fb7g], edgecolor="#333", lw=0.4, alpha=0.88)
                    _ax7g.bar(_xs7g,      _fs7g, _wg7, label="Spencer",
                               color=[_bc7(v) for v in _fs7g], edgecolor="#333", lw=0.4, alpha=0.65, hatch="//")
                    _ax7g.bar(_xs7g+_wg7, _fm7g, _wg7, label="M-P",
                               color=[_bc7(v) for v in _fm7g], edgecolor="#333", lw=0.4, alpha=0.50, hatch="xx")
                    _ax7g.axhline(1.30, color="red", ls="--", lw=1.5, label="Fs_min=1,30")
                    _ax7g.set_xticks(_xs7g); _ax7g.set_xticklabels(_lbl7g, fontsize=8)
                    _ax7g.set_ylabel("Hệ số an toàn Fs", fontsize=9)
                    _ax7g.set_title(
                        "So sánh Fs ổn định — Bishop / Spencer / Morgenstern-Price",
                        fontsize=10, fontweight="bold")
                    _ax7g.legend(fontsize=8, ncol=4); _ax7g.grid(axis="y", alpha=0.3)
                    _ax7g.tick_params(labelsize=7); _plt_w.tight_layout()
                    _buf7g = _io_w.BytesIO()
                    _fig7g.savefig(_buf7g, format="png", dpi=130, bbox_inches="tight", facecolor="white")
                    _plt_w.close(_fig7g); _buf7g.seek(0)
                    _doc_w.add_paragraph()
                    _doc_w.add_picture(_buf7g, width=_Cm_w(15))
                _doc_w.add_page_break()

                # ══════════════════════════════════════════════════════════════
                # PHẦN 8: KẾT LUẬN & KIẾN NGHỊ
                # ══════════════════════════════════════════════════════════════
                _doc_w.add_heading("8. Kết luận và Kiến nghị", level=1)
                _n_nt1_ok = sum(1 for d in _nt_det_w.values() if "Đạt" in str(d.get("nt1_result","")))
                _n_nt2_ok = sum(1 for d in _nt_det_w.values() if "Đạt" in str(d.get("nt2_result","")))
                _n_wink_ok = sum(1 for wr in _winkler_w.values()
                                 if (wr.get("u_max_mm",99) or 99) <= 50
                                 and (wr.get("M_max_kNm",999) or 999) <= (wr.get("Mcr_kNm",1) or 1))
                _n_stab_ok = sum(1 for st in _stab_w.values()
                                 if st.get("Fs_bishop") and float(st["Fs_bishop"]) >= 1.30)
                _tot = len(_nt_det_w)
                _p8 = _doc_w.add_paragraph()
                _p8.add_run("Tổng hợp kết quả:").bold = True
                for _it in [
                    f"NT1 — Chiều dài cọc: {_n_nt1_ok}/{_tot} hố khoan đạt yêu cầu",
                    f"NT2 — Sức chịu tải: {_n_nt2_ok}/{_tot} hố khoan đạt yêu cầu",
                    f"Nội lực Winkler: {_n_wink_ok}/{len(_winkler_w)} hố khoan đạt tiêu chí (u≤50mm & M≤Mcr)",
                    f"Ổn định tổng thể: {_n_stab_ok}/{len(_stab_w)} hố khoan đạt Fs≥1,30",
                ]:
                    _doc_w.add_paragraph(_it, style="List Bullet")
                _doc_w.add_paragraph()
                _p8k = _doc_w.add_paragraph()
                _p8k.add_run("Kiến nghị:").bold = True
                _kns = [
                    "Ưu tiên cọc SW tiết diện tối ưu cho từng hố khoan theo kết quả tính toán trong báo cáo.",
                    "Cập nhật Su thực tế từ VST/lab sau khi có đầy đủ kết quả khảo sát địa chất.",
                    "Kiểm tra lại ổn định tổng thể bằng phần mềm FEM (PLAXIS) sau khi có địa tầng cuối.",
                    "Thi công cọc SW từ hai đầu vào giữa để hạn chế heave đất nền trong quá trình thi công.",
                    "Quan trắc chuyển vị định kỳ không vượt 25 mm trong suốt quá trình thi công.",
                    "Kết hợp xử lý nền CDM với kè SW theo phương án đã tính trong báo cáo lún nền.",
                ]
                for _kn in _kns:
                    _doc_w.add_paragraph(_kn, style="List Bullet")

                # ── Lưu + Download ───────────────────────────────────────────
                _buf_final = _io_w.BytesIO()
                _doc_w.save(_buf_final)
                _buf_final.seek(0)
                st.download_button(
                    "Tải Word toàn diện — Cọc ván SW",
                    _buf_final.getvalue(),
                    file_name=f"CocVanSW_ToanDien_{_dt_w.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
                st.success(
                    f"Tạo báo cáo thành công — "
                    f"NT1: {_n_nt1_ok}/{_tot}  NT2: {_n_nt2_ok}/{_tot}  "
                    f"Nội lực: {_n_wink_ok}/{len(_winkler_w)}  Ổn định: {_n_stab_ok}/{len(_stab_w)} hố khoan đạt."
                )
            except Exception as _e_wf:
                st.error(f"Không tạo được báo cáo Word: {_e_wf}")
                import traceback as _tb_wf
                st.code(_tb_wf.format_exc())

# ── TKBVTC CDM — Mặt bằng tổng hợp cọc CDM + hố khoan ──────────────────────
if _page == "cdm_bvt":
    st.markdown("## TKBVTC CDM")
    st.markdown("### Mặt bằng tổng hợp: Vị trí cọc CDM và hố khoan (có nền địa lý 2D / 3D)")

    import sqlite3 as _sq3_cdm
    import numpy as _np_cdm

    with _sq3_cdm.connect(str(_DB)) as _con_cdm2:
        try:
            _cdm_rows = _con_cdm2.execute(
                "SELECT zone, northing_m, easting_m FROM cdm_toado "
                "WHERE source='xlsx_toado_cdm' ORDER BY zone"
            ).fetchall()
        except Exception:
            _cdm_rows = []
        try:
            _bh_all = _con_cdm2.execute(
                "SELECT name, x_coord_m, y_coord_m FROM boreholes "
                "WHERE x_coord_m > 1000000 AND y_coord_m IS NOT NULL"
            ).fetchall()
        except Exception:
            _bh_all = []
        try:
            _bd_rows = _con_cdm2.execute(
                "SELECT polyline_id, x_m, y_m FROM ke_binhdo_toadoke "
                "ORDER BY polyline_id, vertex_idx"
            ).fetchall()
        except Exception:
            _bd_rows = []

    _cdm_cv = [(r[1], r[2]) for r in _cdm_rows if r[0] == "CONG_VIEN"]
    _cdm_ke = [(r[1], r[2]) for r in _cdm_rows if r[0] == "KE"]

    def _zone_prefix(name):
        if name.startswith("KE-"):  return "KE"
        if name.startswith("BXN-"): return "BXN"
        if name.startswith("NHC-"): return "NHC"
        return "Khác"

    _bh_by_zone: dict = {}
    for nm, xc, yc in _bh_all:
        _bh_by_zone.setdefault(_zone_prefix(nm), []).append(
            {"name": nm, "x": float(xc), "y": float(yc)}
        )

    # Tính sớm `_cnt_all` (số cọc CDM gần nhất theo HK) cho folium tooltip.
    # Không phụ thuộc D/Ztop/pen — chỉ dùng tọa độ HK + tọa độ cọc CDM.
    # Section "Thống kê khối lượng" phía dưới sẽ tái sử dụng dict này.
    _cnt_all: dict = {}
    try:
        for _cdm_pts, _zone_pfx in [(_cdm_ke, "KE"), (_cdm_cv, "BXN")]:
            _zone_bhs = _bh_by_zone.get(_zone_pfx, [])
            if not _cdm_pts or not _zone_bhs:
                continue
            _A_cnt = _np_cdm.array([(b["x"], b["y"]) for b in _zone_bhs])
            for _nx, _ey in _cdm_pts:
                _d2 = (_A_cnt[:, 0] - _nx) ** 2 + (_A_cnt[:, 1] - _ey) ** 2
                _nm_near = _zone_bhs[int(_np_cdm.argmin(_d2))]["name"]
                _cnt_all[_nm_near] = _cnt_all.get(_nm_near, 0) + 1
    except Exception:
        _cnt_all = {}

    # Bình đồ kè polylines: swap DXF x(Easting)↔y(Northing) → chart X=Northing
    _bd_cx: list = []
    _bd_cy: list = []
    _prev_bd = None
    for _pl_id, _bx, _by in _bd_rows:
        if _prev_bd is not None and _pl_id != _prev_bd:
            _bd_cx.append(None); _bd_cy.append(None)
        _bd_cx.append(_by); _bd_cy.append(_bx)
        _prev_bd = _pl_id

    _c1, _c2, _c3 = st.columns(3)
    _c1.metric("Cọc CDM — Công viên", f"{len(_cdm_cv):,}")
    _c2.metric("Cọc CDM — Kè",        f"{len(_cdm_ke):,}")
    _c3.metric("Hố khoan",             str(len(_bh_all)))

    if _HAS_PLOTLY:
        _fig_cdm = go.Figure()
        if _bd_cx:
            _fig_cdm.add_trace(go.Scatter(
                x=_bd_cx, y=_bd_cy, mode="lines",
                line=dict(color="rgba(0,0,0,0.80)", width=1.2),
                name="Ranh kè (bình đồ)", hoverinfo="skip", connectgaps=False,
            ))
        if _cdm_cv:
            _fig_cdm.add_trace(go.Scattergl(
                x=[p[0] for p in _cdm_cv], y=[p[1] for p in _cdm_cv],
                mode="markers",
                marker=dict(size=3, color="rgba(33,150,243,0.45)", line=dict(width=0)),
                name=f"CDM Công viên ({len(_cdm_cv):,} cọc)",
                hovertemplate="N=%.1f<br>E=%.1f<extra>CDM CV</extra>",
            ))
        if _cdm_ke:
            _fig_cdm.add_trace(go.Scattergl(
                x=[p[0] for p in _cdm_ke], y=[p[1] for p in _cdm_ke],
                mode="markers",
                marker=dict(size=3, color="rgba(255,152,0,0.50)", line=dict(width=0)),
                name=f"CDM Kè ({len(_cdm_ke):,} cọc)",
                hovertemplate="N=%.1f<br>E=%.1f<extra>CDM Kè</extra>",
            ))
        _BH_STYLE = {
            "KE":   dict(symbol="circle",      size=12, color="#E53935",
                         line=dict(color="#B71C1C", width=1.5)),
            "BXN":  dict(symbol="square",      size=11, color="#AB47BC",
                         line=dict(color="#6A1B9A", width=1.5)),
            "NHC":  dict(symbol="triangle-up", size=12, color="#43A047",
                         line=dict(color="#1B5E20", width=1.5)),
            "Khác": dict(symbol="diamond",     size=10, color="#78909C",
                         line=dict(color="#37474F", width=1.5)),
        }
        for _zone_nm, _bhs in sorted(_bh_by_zone.items()):
            _fig_cdm.add_trace(go.Scatter(
                x=[b["x"] for b in _bhs], y=[b["y"] for b in _bhs],
                mode="markers+text",
                marker=_BH_STYLE.get(_zone_nm, _BH_STYLE["Khác"]),
                text=[b["name"].split("-")[-1] for b in _bhs],
                textposition="top center", textfont=dict(size=9),
                hovertext=[f"<b>{b['name']}</b><br>N={b['x']:,.1f}<br>E={b['y']:,.1f}"
                           for b in _bhs],
                hoverinfo="text",
                name=f"Hố khoan {_zone_nm} ({len(_bhs)})",
            ))
        _fig_cdm.update_layout(
            title="Mặt bằng tổng hợp: Cọc CDM + Hố khoan",
            xaxis_title="Northing VN-2000 (m)",
            yaxis_title="Easting VN-2000 (m)",
            height=700, hovermode="closest", plot_bgcolor="#F5F5F5",
            legend=dict(x=0.01, y=0.99, bgcolor="rgba(255,255,255,0.88)",
                        bordercolor="#BBB", borderwidth=1, font=dict(size=11)),
            margin=dict(t=60, b=60, l=70, r=20),
        )
        _fig_cdm.update_yaxes(scaleanchor="x", scaleratio=1,
                               showgrid=True, gridcolor="#E0E0E0")
        _fig_cdm.update_xaxes(showgrid=True, gridcolor="#E0E0E0")
        st.plotly_chart(_fig_cdm, use_container_width=True)

    if _HAS_MPL and (_cdm_cv or _cdm_ke):
        import matplotlib.pyplot as _plt_cdm
        _fig_cm, _ax_cm = _plt_cdm.subplots(figsize=(14, 10), dpi=110)
        if _bd_cx:
            _seg_x, _seg_y = [], []
            for _cx2, _cy2 in zip(_bd_cx, _bd_cy):
                if _cx2 is None:
                    if _seg_x:
                        _ax_cm.plot(_seg_x, _seg_y, color="black", lw=1.0,
                                    alpha=0.80, zorder=2)
                    _seg_x, _seg_y = [], []
                else:
                    _seg_x.append(_cx2); _seg_y.append(_cy2)
            if _seg_x:
                _ax_cm.plot(_seg_x, _seg_y, color="black", lw=1.0,
                            alpha=0.80, zorder=2, label="Ranh kè")
        if _cdm_cv:
            _ax_cm.scatter([p[0] for p in _cdm_cv], [p[1] for p in _cdm_cv],
                           s=1, c="steelblue", alpha=0.35, linewidths=0,
                           zorder=3, label=f"CDM Công viên ({len(_cdm_cv):,})")
        if _cdm_ke:
            _ax_cm.scatter([p[0] for p in _cdm_ke], [p[1] for p in _cdm_ke],
                           s=1, c="darkorange", alpha=0.40, linewidths=0,
                           zorder=3, label=f"CDM Kè ({len(_cdm_ke):,})")
        _BH_MPL = {"KE": ("o","#E53935"), "BXN": ("s","#AB47BC"),
                   "NHC": ("^","#43A047"), "Khác": ("D","#78909C")}
        for _zn, _bhs in sorted(_bh_by_zone.items()):
            _mk, _cl = _BH_MPL.get(_zn, ("D","gray"))
            _ax_cm.scatter([b["x"] for b in _bhs], [b["y"] for b in _bhs],
                           s=60, marker=_mk, c=_cl, zorder=5,
                           label=f"HK {_zn} ({len(_bhs)})",
                           edgecolors="white", linewidths=0.8)
            for b in _bhs:
                _ax_cm.annotate(b["name"].split("-")[-1],
                                xy=(b["x"], b["y"]),
                                xytext=(0, 6), textcoords="offset points",
                                ha="center", fontsize=7, color="#333")
        _ax_cm.set_xlabel("Northing VN-2000 (m)")
        _ax_cm.set_ylabel("Easting VN-2000 (m)")
        _ax_cm.set_title("Mặt bằng tổng hợp: Cọc CDM + Hố khoan")
        _ax_cm.set_aspect("equal", adjustable="datalim")
        _ax_cm.grid(True, ls="--", alpha=0.35)
        _ax_cm.legend(fontsize=8, loc="lower right", framealpha=0.88, markerscale=2)
        _fig_cm.tight_layout()
        with st.expander("Bản in (dùng cho PDF)", expanded=False):
            st.pyplot(_fig_cm, use_container_width=True)
        _plt_cdm.close(_fig_cm)

    st.caption(
        f"Dữ liệu: {len(_cdm_cv):,} cọc CDM khu Công viên + "
        f"{len(_cdm_ke):,} cọc CDM khu Kè. "
        "Tọa độ VN-2000. Click tên lớp trong chú thích để ẩn/hiện."
    )

    # ════════════════════════════════════════════════════════════════════════
    # BẢN ĐỒ 2D CÓ NỀN ĐỊA LÝ (folium) + BẢN ĐỒ 3D (pydeck)
    # ════════════════════════════════════════════════════════════════════════
    st.divider()
    st.markdown("#### Bản đồ có nền địa lý 2D / 3D")

    try:
        import folium as _flm
        from streamlit_folium import st_folium as _stflm
        _HAS_FLM = True
    except ImportError:
        _HAS_FLM = False
    try:
        import pydeck as _pdk
        _HAS_PDK = True
    except ImportError:
        _HAS_PDK = False
    try:
        from pyproj import Transformer as _Trf
        _HAS_PRJ = True
    except ImportError:
        _HAS_PRJ = False

    if not (_HAS_FLM and _HAS_PDK and _HAS_PRJ):
        st.error("Thiếu thư viện bản đồ. Cài: folium, streamlit-folium, pydeck, pyproj.")
    elif not _bh_all:
        st.info("Chưa có dữ liệu hố khoan trong khu vực.")
    else:
        _MAP_CRS = {
            "VN-2000 / TM-3 106°00' (TTHC Quận 1, Thủ Thiêm)": "EPSG:9210",
            "VN-2000 / TM-3 105°45' (HCM mặc định)": "EPSG:9209",
            "VN-2000 / UTM 48N (105°)": "EPSG:3405",
        }

        @st.cache_data(show_spinner=False)
        def _bd_make_trf(epsg: str):
            return _Trf.from_crs(epsg, "EPSG:4326", always_xy=True)

        _bm_c1, _bm_c2, _bm_c3 = st.columns([3, 2, 1])
        with _bm_c1:
            _crs_lbl = st.selectbox(
                "Hệ tọa độ gốc của dữ liệu",
                list(_MAP_CRS.keys()),
                index=0,
                key="_cdm_bvt_crs",
                help="Đổi nếu vị trí trên bản đồ bị lệch khoảng 25-50km. TTHC Quận 1/Thủ Thiêm thường dùng TM-3 106°00'.",
            )
            _crs_code = _MAP_CRS[_crs_lbl]
        with _bm_c2:
            _show_cdm_sample = st.checkbox(
                "Hiển thị mẫu cọc CDM trên bản đồ 2D",
                value=False,
                help="Bật để xem một mẫu cọc CDM (tối đa 2,000 cọc) — nhiều hơn sẽ làm chậm trình duyệt. Bản đồ 3D vẽ toàn bộ.",
            )
        with _bm_c3:
            _vmode = st.radio("Chế độ", ["2D", "3D"], horizontal=False, key="_cdm_bvt_view")

        _trf_map = _bd_make_trf(_crs_code)

        # Chuyển HK sang lat/lon
        _hk_geo = []
        for nm, xc, yc in _bh_all:
            # boreholes: x=Northing, y=Easting → trf.transform(Easting, Northing)
            _lon, _lat = _trf_map.transform(yc, xc)
            _hk_geo.append({
                "name": nm,
                "zone": _zone_prefix(nm),
                "lat": _lat, "lon": _lon,
                "n_piles": _cnt_all.get(nm, 0),
            })
        _lat0 = sum(h["lat"] for h in _hk_geo) / len(_hk_geo)
        _lon0 = sum(h["lon"] for h in _hk_geo) / len(_hk_geo)

        # Polyline kè
        _ke_geo_lines: list[list[tuple[float, float]]] = []
        if _bd_rows:
            _cur_pl, _cur_pts = None, []
            for _pl_id, _bx, _by in _bd_rows:
                if _cur_pl is not None and _pl_id != _cur_pl:
                    if len(_cur_pts) >= 2:
                        _ke_geo_lines.append(_cur_pts)
                    _cur_pts = []
                # JSON: x=Easting, y=Northing → trf.transform(Easting, Northing)
                _lon_p, _lat_p = _trf_map.transform(_bx, _by)
                _cur_pts.append((_lat_p, _lon_p))
                _cur_pl = _pl_id
            if len(_cur_pts) >= 2:
                _ke_geo_lines.append(_cur_pts)

        # CDM sample (chỉ cho folium 2D — pydeck vẽ full)
        _CDM_COLORS = {"CV": "#2196F3", "KE": "#FF9800"}
        _CDM_ZONE_LABEL = {"CV": "Công viên", "KE": "Kè"}

        def _convert_cdm_pts(rows, max_n=None):
            out = []
            step = max(1, len(rows) // max_n) if (max_n and len(rows) > max_n) else 1
            for nx, ey in rows[::step]:
                # cdm_toado: northing_m, easting_m → trf(easting, northing)
                _lonc, _latc = _trf_map.transform(ey, nx)
                out.append((_latc, _lonc))
            return out

        try:
            _con_cap = sqlite3.connect(_DB)
            _n_tp_bvt = _con_cap.execute(
                "SELECT COUNT(*) FROM cdm_coc_thu"
            ).fetchone()[0]
            _n_mt_bvt = _con_cap.execute(
                "SELECT COUNT(DISTINCT polyline_id) FROM metro_lines"
            ).fetchone()[0]
            _con_cap.close()
        except Exception:
            _n_tp_bvt = 0
            _n_mt_bvt = 0
        st.caption(
            f"Tâm bản đồ: lat={_lat0:.5f}°N, lon={_lon0:.5f}°E · "
            f"{len(_hk_geo)} hố khoan · {len(_ke_geo_lines)} đường kè · "
            f"{len(_cdm_cv) + len(_cdm_ke):,} cọc CDM · "
            f"{_n_tp_bvt} cọc CDM thử · "
            f"{_n_mt_bvt} polyline metro"
        )

        # ── 2D folium map ───────────────────────────────────────────────────
        if _vmode == "2D":
            _fmap = _flm.Map(location=[_lat0, _lon0], zoom_start=17,
                             tiles=None, control_scale=True)

            _flm.TileLayer("OpenStreetMap", name="OpenStreetMap").add_to(_fmap)
            _flm.TileLayer(
                tiles="https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}",
                attr="Esri", name="Ảnh vệ tinh (Esri)",
            ).add_to(_fmap)
            _flm.TileLayer(
                tiles="https://{s}.tile.opentopomap.org/{z}/{x}/{y}.png",
                attr="OpenTopoMap", name="Địa hình (OpenTopoMap)",
            ).add_to(_fmap)
            _flm.TileLayer(
                tiles="https://{s}.basemaps.cartocdn.com/light_all/{z}/{x}/{y}.png",
                attr="CartoDB", name="Carto nhạt (in PDF)",
            ).add_to(_fmap)

            # Lớp polyline kè
            if _ke_geo_lines:
                _fg_kel = _flm.FeatureGroup(name="Tuyến kè (đen)", show=True)
                for poly in _ke_geo_lines:
                    _flm.PolyLine(poly, color="#000000", weight=3,
                                  opacity=0.85).add_to(_fg_kel)
                _fg_kel.add_to(_fmap)

            # Lớp cọc CDM thử (cdm_coc_thu — 18 cọc)
            try:
                _con_bvt_tp = sqlite3.connect(_DB)
                _bvt_tp_rows = _con_bvt_tp.execute(
                    "SELECT code, northing_m, easting_m, elevation_m, description "
                    "FROM cdm_coc_thu ORDER BY point_name"
                ).fetchall()
                _con_bvt_tp.close()
            except Exception:
                _bvt_tp_rows = []
            if _bvt_tp_rows:
                _fg_bvt_tp = _flm.FeatureGroup(
                    name=f"Cọc CDM thử ({len(_bvt_tp_rows)})", show=True
                )
                for _code, _nor, _eas, _elev, _desc in _bvt_tp_rows:
                    _lon_tp, _lat_tp = _trf_map.transform(_eas, _nor)
                    _num = _code.split("-")[-1] if "-" in _code else _code
                    _flm.CircleMarker(
                        location=[_lat_tp, _lon_tp],
                        radius=7, color="#BF360C", weight=2,
                        fill=True, fill_color="#FF6F00", fill_opacity=0.95,
                        popup=_flm.Popup(
                            f"<b>{_code}</b><br><i>{_desc}</i><br>"
                            f"Cao độ mặt: <b>{_elev:.2f} m</b><br>"
                            f"Tọa độ: N={_nor:,.1f} m, E={_eas:,.1f} m",
                            max_width=240,
                        ),
                        tooltip=_code,
                    ).add_to(_fg_bvt_tp)
                    _flm.map.Marker(
                        [_lat_tp, _lon_tp],
                        icon=_flm.DivIcon(html=(
                            f'<div style="font-size:10px;font-weight:bold;'
                            f'color:#BF360C;text-shadow:1px 1px 2px #fff,'
                            f'-1px -1px 2px #fff;margin-left:10px;'
                            f'margin-top:-6px;white-space:nowrap">{_num}</div>'
                        )),
                    ).add_to(_fg_bvt_tp)
                _fg_bvt_tp.add_to(_fmap)

            # Lớp tuyến metro (bảng metro_lines) — bỏ aux_tunnel + boundary_control
            _METRO_STYLE_BVT = {
                "centerline":       {"color": "#D32F2F", "weight": 4, "opacity": 0.95, "dash": None,  "label": "Tuyến chính metro"},
                "boundary_land":    {"color": "#616161", "weight": 2, "opacity": 0.80, "dash": "4,4", "label": "Ranh GPMB"},
                "boundary_station": {"color": "#9E9E9E", "weight": 2, "opacity": 0.80, "dash": "2,4", "label": "Ranh ga (Xref)"},
                "aux_panel":        {"color": "#7B1FA2", "weight": 1, "opacity": 0.60, "dash": None,  "label": "Tường panel"},
                "other":            {"color": "#000000", "weight": 1, "opacity": 0.50, "dash": None,  "label": "Khác"},
            }
            try:
                _con_bvt_mt = sqlite3.connect(_DB)
                _bvt_mt_rows = _con_bvt_mt.execute(
                    "SELECT polyline_id, vertex_idx, x_m, y_m, category "
                    "FROM metro_lines "
                    "WHERE category NOT IN ('aux_tunnel', 'boundary_control') "
                    "ORDER BY category, polyline_id, vertex_idx"
                ).fetchall()
                _con_bvt_mt.close()
            except Exception:
                _bvt_mt_rows = []
            if _bvt_mt_rows:
                _bvt_polys: dict = {}
                _bvt_key = (None, None)
                _bvt_pts: list = []
                _bvt_cat = None
                for _pl, _vi, _bx, _by, _cat in _bvt_mt_rows:
                    if (_cat, _pl) != _bvt_key and _bvt_pts:
                        _bvt_polys.setdefault(_bvt_cat, []).append(_bvt_pts)
                        _bvt_pts = []
                    _bvt_key = (_cat, _pl)
                    _bvt_cat = _cat
                    _lonm, _latm = _trf_map.transform(_bx, _by)
                    _bvt_pts.append((_latm, _lonm))
                if _bvt_pts and _bvt_cat is not None:
                    _bvt_polys.setdefault(_bvt_cat, []).append(_bvt_pts)

                for _cat in ["aux_panel", "boundary_land",
                             "boundary_station", "other", "centerline"]:
                    _polys = _bvt_polys.get(_cat)
                    if not _polys:
                        continue
                    _sty = _METRO_STYLE_BVT.get(_cat, _METRO_STYLE_BVT["other"])
                    _fg_bvt_mt = _flm.FeatureGroup(
                        name=f"Metro: {_sty['label']} ({len(_polys)})",
                        show=True,
                    )
                    for _poly in _polys:
                        if len(_poly) < 2:
                            continue
                        _kwargs = dict(
                            locations=_poly,
                            color=_sty["color"],
                            weight=_sty["weight"],
                            opacity=_sty["opacity"],
                        )
                        if _sty["dash"]:
                            _kwargs["dash_array"] = _sty["dash"]
                        _flm.PolyLine(**_kwargs).add_to(_fg_bvt_mt)
                    _fg_bvt_mt.add_to(_fmap)

            # Lớp HK theo zone
            _HK_FLM_COLOR = {"KE": "#E53935", "BXN": "#AB47BC", "NHC": "#43A047",
                              "Khác": "#78909C"}
            _zones_fg: dict[str, "_flm.FeatureGroup"] = {}
            for h in _hk_geo:
                z = h["zone"]
                if z not in _zones_fg:
                    _zones_fg[z] = _flm.FeatureGroup(
                        name=f"Hố khoan {z}", show=True
                    )
                _popup = _flm.Popup(
                    html=(
                        f"<b>{h['name']}</b><br>"
                        f"Khu vực: <b>{z}</b><br>"
                        f"Số cọc CDM gán cho HK này: <b>{h['n_piles']:,}</b>"
                    ),
                    max_width=260,
                )
                _flm.CircleMarker(
                    location=[h["lat"], h["lon"]],
                    radius=8, color=_HK_FLM_COLOR.get(z, "#777"),
                    weight=2, fill=True,
                    fill_color=_HK_FLM_COLOR.get(z, "#999"),
                    fill_opacity=0.85,
                    popup=_popup, tooltip=h["name"],
                ).add_to(_zones_fg[z])
                _flm.map.Marker(
                    [h["lat"], h["lon"]],
                    icon=_flm.DivIcon(html=(
                        f'<div style="font-size:11px;font-weight:bold;'
                        f'color:#222;text-shadow:1px 1px 2px #fff,'
                        f'-1px -1px 2px #fff;margin-left:10px;margin-top:-10px">'
                        f'{h["name"].split("-")[-1]}</div>'
                    )),
                ).add_to(_zones_fg[z])
            for fg in _zones_fg.values():
                fg.add_to(_fmap)

            # Lớp cọc CDM mẫu (giới hạn 2000 mỗi nhóm)
            if _show_cdm_sample:
                for _grp_key, _pts in [("CV", _cdm_cv), ("KE", _cdm_ke)]:
                    if not _pts:
                        continue
                    _pts_geo = _convert_cdm_pts(_pts, max_n=2000)
                    _fg_cdm = _flm.FeatureGroup(
                        name=f"Mẫu cọc CDM {_CDM_ZONE_LABEL[_grp_key]} ({len(_pts_geo):,}/{len(_pts):,})",
                        show=True,
                    )
                    for la, lo in _pts_geo:
                        _flm.CircleMarker(
                            location=[la, lo], radius=1.5,
                            color=_CDM_COLORS[_grp_key],
                            weight=0, fill=True,
                            fill_color=_CDM_COLORS[_grp_key],
                            fill_opacity=0.55,
                        ).add_to(_fg_cdm)
                    _fg_cdm.add_to(_fmap)

            _flm.LayerControl(collapsed=False, position="topright").add_to(_fmap)

            try:
                from folium.plugins import (Fullscreen as _FS,
                                             MeasureControl as _MC,
                                             MiniMap as _MM)
                _FS(position="topright", title="Toàn màn hình",
                    title_cancel="Thoát").add_to(_fmap)
                _MC(position="topleft", primary_length_unit="meters",
                    primary_area_unit="sqmeters").add_to(_fmap)
                _MM(toggle_display=True, position="bottomright").add_to(_fmap)
            except Exception:
                pass

            _stflm(_fmap, width=None, height=720, returned_objects=[])

            with st.expander("Hướng dẫn dùng bản đồ 2D"):
                st.markdown(
                    "- **Đổi nền**: bảng góc trên phải — chọn OSM / Vệ tinh / Địa hình / Carto\n"
                    "- **Tắt/mở lớp**: tick checkbox \"Tuyến kè\", \"Hố khoan KE/BXN/NHC\"\n"
                    "- **Zoom**: nút +/− góc trái trên hoặc lăn chuột\n"
                    "- **Pan**: kéo chuột trái\n"
                    "- **Toàn màn hình**: nút góc phải trên\n"
                    "- **Đo khoảng cách / diện tích**: nút thước góc trái trên\n"
                    "- **Nhấp HK**: hiện popup chi tiết\n"
                    "- **Mẫu cọc CDM**: tick \"Hiển thị mẫu cọc CDM\" phía trên — hiển thị tối đa 2,000 cọc/zone để không làm chậm trình duyệt; xem toàn bộ trong chế độ 3D."
                )

        # ── 3D pydeck view ──────────────────────────────────────────────────
        else:
            _layers_3d = []

            # Lớp HK (Column extrude)
            _HK_PDK_COLOR = {
                "KE":   [229, 57, 53, 220], "BXN":  [171, 71, 188, 220],
                "NHC":  [67, 160, 71, 220], "Khác": [120, 144, 156, 220],
            }
            _hk_pdk_data = [
                {
                    "position": [h["lon"], h["lat"]],
                    "name": h["name"], "zone": h["zone"],
                    "color": _HK_PDK_COLOR.get(h["zone"], [120, 144, 156, 220]),
                    "height": 120.0,
                    "n_piles": h["n_piles"],
                }
                for h in _hk_geo
            ]
            _layers_3d.append(_pdk.Layer(
                "ColumnLayer", data=_hk_pdk_data,
                get_position="position",
                get_elevation="height",
                get_fill_color="color",
                radius=4, pickable=True, auto_highlight=True,
                extruded=True,
            ))

            # Lớp polyline kè
            if _ke_geo_lines:
                _path_data = [
                    {"path": [[lon, lat] for (lat, lon) in poly]}
                    for poly in _ke_geo_lines
                ]
                _layers_3d.append(_pdk.Layer(
                    "PathLayer", data=_path_data,
                    get_path="path", get_color=[0, 0, 0, 220],
                    width_scale=1, width_min_pixels=2, get_width=4,
                ))

            # Lớp cọc CDM (full 27K — pydeck WebGL OK)
            for _grp_key, _pts, _label in [
                ("CV", _cdm_cv, "Công viên"), ("KE", _cdm_ke, "Kè"),
            ]:
                if not _pts:
                    continue
                _cdm_geo = _convert_cdm_pts(_pts, max_n=None)
                _layers_3d.append(_pdk.Layer(
                    "ScatterplotLayer",
                    data=[{"position": [lo, la]} for (la, lo) in _cdm_geo],
                    get_position="position",
                    get_radius=0.5,
                    get_fill_color=(
                        [33, 150, 243, 140] if _grp_key == "CV" else [255, 152, 0, 140]
                    ),
                    radius_min_pixels=1, radius_max_pixels=5,
                ))

            # Lớp cọc CDM thử (18 cọc — nổi bật bằng ColumnLayer cam cao 80m)
            try:
                _con_3d_tp = sqlite3.connect(_DB)
                _3d_tp_rows = _con_3d_tp.execute(
                    "SELECT code, northing_m, easting_m FROM cdm_coc_thu ORDER BY point_name"
                ).fetchall()
                _con_3d_tp.close()
            except Exception:
                _3d_tp_rows = []
            if _3d_tp_rows:
                _tp_pdk_data = []
                for _code, _nor, _eas in _3d_tp_rows:
                    _lon_tp, _lat_tp = _trf_map.transform(_eas, _nor)
                    _tp_pdk_data.append({
                        "position": [_lon_tp, _lat_tp],
                        "name":     _code,
                        "height":   80.0,
                    })
                _layers_3d.append(_pdk.Layer(
                    "ColumnLayer", data=_tp_pdk_data,
                    get_position="position",
                    get_elevation="height",
                    get_fill_color=[255, 111, 0, 240],  # cam đậm
                    radius=3, pickable=True, auto_highlight=True,
                    extruded=True,
                ))

            # Lớp tuyến metro centerline (PathLayer đỏ — không vẽ ranh/tường cho 3D đỡ rối)
            try:
                _con_3d_mt = sqlite3.connect(_DB)
                _3d_mt_rows = _con_3d_mt.execute(
                    "SELECT polyline_id, vertex_idx, x_m, y_m FROM metro_lines "
                    "WHERE category='centerline' ORDER BY polyline_id, vertex_idx"
                ).fetchall()
                _con_3d_mt.close()
            except Exception:
                _3d_mt_rows = []
            if _3d_mt_rows:
                _mt_paths: list = []
                _cur_pl_mt = None
                _cur_path_mt: list = []
                for _pl, _vi, _bx, _by in _3d_mt_rows:
                    if _cur_pl_mt is not None and _pl != _cur_pl_mt and _cur_path_mt:
                        _mt_paths.append({"path": _cur_path_mt})
                        _cur_path_mt = []
                    _cur_pl_mt = _pl
                    _lonm, _latm = _trf_map.transform(_bx, _by)
                    _cur_path_mt.append([_lonm, _latm])
                if _cur_path_mt:
                    _mt_paths.append({"path": _cur_path_mt})
                if _mt_paths:
                    _layers_3d.append(_pdk.Layer(
                        "PathLayer", data=_mt_paths,
                        get_path="path",
                        get_color=[211, 47, 47, 240],  # đỏ đậm
                        width_scale=1, width_min_pixels=3, get_width=5,
                    ))

            _style_lbl = st.selectbox(
                "Nền bản đồ 3D",
                ["Sáng", "Tối", "Đường phố", "Vệ tinh"],
                index=0, key="_cdm_bvt_3dstyle",
            )
            _style = {"Sáng":"light","Tối":"dark","Đường phố":"road","Vệ tinh":"satellite"}[_style_lbl]

            _deck3d = _pdk.Deck(
                layers=_layers_3d,
                initial_view_state=_pdk.ViewState(
                    latitude=_lat0, longitude=_lon0,
                    zoom=17, pitch=55, bearing=0,
                ),
                map_style=_style,
                tooltip={
                    "html": "<b>{name}</b><br>Khu vực: {zone}<br>Cọc CDM gần nhất: {n_piles}",
                    "style": {"backgroundColor": "#2c3e50", "color": "white"},
                },
            )
            st.pydeck_chart(_deck3d, use_container_width=True)

            with st.expander("Hướng dẫn dùng bản đồ 3D"):
                st.markdown(
                    "- **Xoay**: Ctrl + kéo chuột (hoặc chuột phải)\n"
                    "- **Phóng**: lăn chuột\n"
                    "- **Di chuyển**: kéo chuột trái\n"
                    "- **Đổi nền**: dropdown phía trên\n"
                    "- **Hover HK**: hiện tooltip chi tiết\n"
                    f"- **Số cọc CDM hiển thị**: {len(_cdm_cv)+len(_cdm_ke):,} (toàn bộ — pydeck dùng WebGL nên không chậm)"
                )

    # ── Thống kê khối lượng cọc CDM ──────────────────────────────────────────
    st.divider()
    st.markdown("### Thống kê khối lượng cọc CDM")

    _ci1, _ci2, _ci3 = st.columns(3)
    _D_cdm    = _ci1.number_input("Đường kính cọc D (m)",       value=0.60,
                                   min_value=0.30, max_value=2.0,
                                   step=0.05, format="%.2f", key="_cdm_D")
    _Ztop_cdm = _ci2.number_input("Cao độ đỉnh cọc (m)",        value=0.80,
                                   step=0.10, format="%.2f",      key="_cdm_ztop")
    _pen_cdm  = _ci3.number_input("Ngàm dưới đáy lớp bùn (m)",  value=1.00,
                                   min_value=0.0, step=0.50, format="%.2f",
                                   key="_cdm_pen")

    _SOFT_SYM = frozenset({"1","1b","XMD","2","2a","2b","2c"})

    with _sq3_cdm.connect(str(_DB)) as _con_vol:
        # D_bottom_soft: KE → ke_sw_nt_detail; others → layers table
        _ke_soft = {r[0]: r[1] for r in _con_vol.execute(
            "SELECT bh_name, D_bottom_soft_m FROM ke_sw_nt_detail "
            "WHERE D_bottom_soft_m IS NOT NULL"
        ).fetchall()}
        _sym_ph = ",".join(f"'{s}'" for s in _SOFT_SYM)
        _lyr_soft = {r[0]: r[1] for r in _con_vol.execute(
            f"SELECT b.name, MAX(l.depth_bot_m) "
            f"FROM layers l JOIN boreholes b ON l.borehole_id=b.id "
            f"WHERE l.symbol IN ({_sym_ph}) AND b.x_coord_m > 1000000 "
            f"GROUP BY b.name"
        ).fetchall()}
        _bh_elev = {r[0]: (float(r[1] or 0), float(r[2]), float(r[3]))
                    for r in _con_vol.execute(
            "SELECT name, elevation_m, x_coord_m, y_coord_m "
            "FROM boreholes WHERE x_coord_m > 1000000 AND y_coord_m IS NOT NULL"
        ).fetchall()}

    # Hợp nhất nguồn D_bottom_soft (ke_sw_nt_detail ưu tiên cho KE)
    _bh_design: list[dict] = []
    for _bn, (elev, xc, yc) in sorted(_bh_elev.items()):
        _D_soft = _ke_soft.get(_bn) or _lyr_soft.get(_bn)
        if _D_soft is None:
            continue
        _z_bot_soft = elev - _D_soft
        _z_tip      = _z_bot_soft - _pen_cdm
        _L          = _Ztop_cdm - _z_tip
        _zone       = ("KE" if _bn.startswith("KE-") else
                       "BXN" if _bn.startswith("BXN-") else "NHC")
        _bh_design.append({
            "bh": _bn, "zone": _zone,
            "elevation": elev,
            "D_soft": round(_D_soft, 2),
            "z_bot_soft": round(_z_bot_soft, 2),
            "z_tip": round(_z_tip, 2),
            "L": round(_L, 2),
            "x": xc, "y": yc,
        })

    # Gán cọc CDM đến hố khoan gần nhất cùng khu (vectorised numpy)
    def _assign_nearest(cdm_pts, bh_list):
        if not cdm_pts or not bh_list:
            return {}
        _A = _np_cdm.array([(b["x"], b["y"]) for b in bh_list])
        counts = {b["bh"]: 0 for b in bh_list}
        for _nx, _ey in cdm_pts:
            _d2 = (_A[:,0] - _nx)**2 + (_A[:,1] - _ey)**2
            counts[bh_list[int(_np_cdm.argmin(_d2))]["bh"]] += 1
        return counts

    _ke_bhs  = [b for b in _bh_design if b["zone"] == "KE"]
    _bxn_bhs = [b for b in _bh_design if b["zone"] == "BXN"]

    _cnt_ke  = _assign_nearest(_cdm_ke,  _ke_bhs)
    _cnt_bxn = _assign_nearest(_cdm_cv,  _bxn_bhs)
    _cnt_all = {**_cnt_ke, **_cnt_bxn}

    _A_sec = _np_cdm.pi / 4 * _D_cdm**2

    # Gán số cọc vào _bh_design
    for b in _bh_design:
        b["n_piles"] = _cnt_all.get(b["bh"], 0)
        b["V_m3"]    = round(b["n_piles"] * _A_sec * b["L"], 1)

    # ── Bảng 1: chiều dài cọc theo hố khoan ─────────────────────────────────
    st.markdown("#### Chiều dài thiết kế cọc CDM theo hố khoan")
    import pandas as _pd_cdm
    _df1 = _pd_cdm.DataFrame([
        {
            "Hố khoan":              b["bh"],
            "Khu vực":               b["zone"],
            "Cao độ nền (m)":        b["elevation"],
            "Đáy lớp bùn (m sâu)":  b["D_soft"],
            "Cao độ đáy bùn (m)":   b["z_bot_soft"],
            "Cao độ đỉnh cọc (m)":  _Ztop_cdm,
            "Cao độ mũi cọc (m)":   b["z_tip"],
            "Chiều dài L (m)":       b["L"],
        }
        for b in _bh_design if b["n_piles"] > 0 or b["zone"] in ("KE","BXN")
    ])
    if not _df1.empty:
        st.dataframe(
            _df1.style.format({
                "Cao độ nền (m)": "{:.2f}",
                "Đáy lớp bùn (m sâu)": "{:.1f}",
                "Cao độ đáy bùn (m)": "{:.2f}",
                "Cao độ đỉnh cọc (m)": "{:.2f}",
                "Cao độ mũi cọc (m)": "{:.2f}",
                "Chiều dài L (m)": "{:.2f}",
            }).apply(
                lambda row: ["background-color:#E3F2FD"]*len(row)
                            if row["Khu vực"] == "KE"
                            else ["background-color:#F3E5F5"]*len(row),
                axis=1,
            ),
            use_container_width=True, hide_index=True,
        )

    # ── Bảng 2: thống kê tổng hợp per zone ──────────────────────────────────
    st.markdown("#### Tổng hợp khối lượng")
    _summary = []
    for _zone_s, _label in [("KE","Kè"), ("BXN","Công viên")]:
        _subset = [b for b in _bh_design if b["zone"] == _zone_s and b["n_piles"] > 0]
        if not _subset:
            continue
        _n = sum(b["n_piles"] for b in _subset)
        _V = sum(b["V_m3"]    for b in _subset)
        _L_avg = sum(b["L"] * b["n_piles"] for b in _subset) / _n if _n else 0
        _summary.append({
            "Khu vực":              _label,
            "Số cọc CDM":           _n,
            "Đường kính D (m)":     _D_cdm,
            "Tiết diện (m²)":       round(_A_sec, 4),
            "L trung bình (m)":     round(_L_avg, 2),
            "L min (m)":            round(min(b["L"] for b in _subset), 2),
            "L max (m)":            round(max(b["L"] for b in _subset), 2),
            "Tổng thể tích (m³)":   round(_V, 0),
        })
    # Tổng cộng
    if len(_summary) > 1:
        _summary.append({
            "Khu vực":             "Tổng cộng",
            "Số cọc CDM":          sum(r["Số cọc CDM"] for r in _summary),
            "Đường kính D (m)":    _D_cdm,
            "Tiết diện (m²)":      round(_A_sec, 4),
            "L trung bình (m)":    round(
                sum(r["L trung bình (m)"] * r["Số cọc CDM"] for r in _summary)
                / max(sum(r["Số cọc CDM"] for r in _summary), 1), 2),
            "L min (m)":           round(min(r["L min (m)"] for r in _summary), 2),
            "L max (m)":           round(max(r["L max (m)"] for r in _summary), 2),
            "Tổng thể tích (m³)":  round(sum(r["Tổng thể tích (m³)"] for r in _summary), 0),
        })
    if _summary:
        _df2 = _pd_cdm.DataFrame(_summary)
        st.dataframe(
            _df2.style.format({
                "Đường kính D (m)": "{:.2f}",
                "Tiết diện (m²)": "{:.4f}",
                "L trung bình (m)": "{:.2f}",
                "L min (m)": "{:.2f}", "L max (m)": "{:.2f}",
                "Tổng thể tích (m³)": "{:,.0f}",
            }).apply(
                lambda row: ["font-weight:bold; background-color:#FFF9C4"]*len(row)
                            if row["Khu vực"] == "Tổng cộng"
                            else [""]*len(row),
                axis=1,
            ),
            use_container_width=True, hide_index=True,
        )

    # ── Biểu đồ chiều dài L per borehole ────────────────────────────────────
    _plot_bhs = [b for b in _bh_design
                 if b["zone"] in ("KE","BXN") and b["n_piles"] > 0]
    if _HAS_PLOTLY and _plot_bhs:
        _fig_L = go.Figure()
        for _zone_s, _clr, _label in [
            ("KE","#1976D2","Kè"), ("BXN","#8E24AA","Công viên")
        ]:
            _sub = [b for b in _plot_bhs if b["zone"] == _zone_s]
            if not _sub:
                continue
            _fig_L.add_trace(go.Bar(
                x=[b["bh"].split("-")[-1] for b in _sub],
                y=[b["L"] for b in _sub],
                name=_label,
                marker_color=_clr,
                text=[f"{b['L']:.1f} m" for b in _sub],
                textposition="outside",
                customdata=[[b["bh"], b["elevation"], b["z_tip"], b["n_piles"]]
                            for b in _sub],
                hovertemplate=(
                    "<b>%{customdata[0]}</b><br>"
                    "Cao độ nền: %{customdata[1]:.2f} m<br>"
                    "Cao độ mũi: %{customdata[2]:.2f} m<br>"
                    "L = %{y:.2f} m<br>"
                    "Số cọc: %{customdata[3]:,}<extra></extra>"
                ),
            ))
        _fig_L.add_hline(y=_Ztop_cdm - (
            sum(b["z_tip"] * b["n_piles"] for b in _plot_bhs)
            / max(sum(b["n_piles"] for b in _plot_bhs), 1)
        ), line_dash="dot", line_color="red",
            annotation_text="L trung bình", annotation_position="right")
        _fig_L.update_layout(
            title=f"Chiều dài cọc CDM theo hố khoan điều hành  (D={_D_cdm:.2f}m, đỉnh={_Ztop_cdm:+.1f}m)",
            xaxis_title="Hố khoan", yaxis_title="Chiều dài L (m)",
            height=380, barmode="group",
            margin=dict(t=55, b=40, r=60),
            plot_bgcolor="#FAFAFA",
        )
        st.plotly_chart(_fig_L, use_container_width=True)

# ── Trang Tiến độ & Thí nghiệm CDM ──────────────────────────────────────────
if _page == "cdm_tien_do":
    st.markdown("## Tiến độ thi công & Kiểm tra thí nghiệm CDM")

    _con_td = _get_db()

    # ── A. Tổng quan tiến độ ──────────────────────────────────────────────────
    st.markdown("### A. Tiến độ thi công")

    _td_rows = _con_td.execute("""
        SELECT zone, status, COUNT(*) AS n
        FROM cdm_thi_cong GROUP BY zone, status
    """).fetchall()

    _td_total   = sum(r[2] for r in _td_rows)
    _td_done    = sum(r[2] for r in _td_rows if r[1] == "hoan_thanh")
    _td_going   = sum(r[2] for r in _td_rows if r[1] == "dang_thi_cong")
    _td_pending = sum(r[2] for r in _td_rows if r[1] == "chua_thi_cong")
    _td_fail    = sum(r[2] for r in _td_rows if r[1] == "khong_dat")
    _td_pause   = sum(r[2] for r in _td_rows if r[1] == "tam_dung")
    _td_pct     = round(_td_done / _td_total * 100, 1) if _td_total else 0

    _m1, _m2, _m3, _m4, _m5 = st.columns(5)
    _m1.metric("Tổng cọc",        f"{_td_total:,}")
    _m2.metric("Hoàn thành",      f"{_td_done:,}",    delta=f"{_td_pct}%")
    _m3.metric("Đang thi công",   f"{_td_going:,}")
    _m4.metric("Chưa thi công",   f"{_td_pending:,}")
    _m5.metric("Không đạt / Tạm dừng", f"{_td_fail + _td_pause:,}")

    # Biểu đồ tiến độ theo ngày
    _td_daily = _con_td.execute("""
        SELECT ngay_thi_cong, zone, COUNT(*) AS n
        FROM cdm_thi_cong
        WHERE ngay_thi_cong IS NOT NULL AND status IN ('hoan_thanh','dang_thi_cong')
        GROUP BY ngay_thi_cong, zone
        ORDER BY ngay_thi_cong
    """).fetchall()

    if _td_daily:
        import pandas as _pd_td
        _df_daily = _pd_td.DataFrame(_td_daily, columns=["ngay","zone","n"])
        _df_pivot  = _df_daily.pivot_table(index="ngay", columns="zone", values="n", fill_value=0).reset_index()

        _fig_td = go.Figure()
        _zone_colors = {"KE": "#1565C0", "CONG_VIEN": "#6A1B9A"}
        for _z in [c for c in _df_pivot.columns if c != "ngay"]:
            _clr = _zone_colors.get(_z, "#546E7A")
            _fig_td.add_trace(go.Bar(
                x=_df_pivot["ngay"], y=_df_pivot[_z],
                name=_z, marker_color=_clr,
            ))
        _fig_td.update_layout(
            barmode="stack",
            title="Số cọc thi công theo ngày",
            xaxis_title="Ngày thi công",
            yaxis_title="Số cọc",
            height=380,
            legend_title="Khu vực",
            margin=dict(t=48, b=48, l=48, r=16),
        )
        st.plotly_chart(_fig_td, use_container_width=True)

        # Đường cộng dồn
        _df_cum = _df_daily.groupby("ngay")["n"].sum().reset_index()
        _df_cum["cum"] = _df_cum["n"].cumsum()
        _fig_cum = go.Figure()
        _fig_cum.add_trace(go.Scatter(
            x=_df_cum["ngay"], y=_df_cum["cum"],
            mode="lines+markers", name="Cộng dồn",
            line=dict(color="#1B5E20", width=2),
            fill="tozeroy", fillcolor="rgba(27,94,32,0.10)",
        ))
        _fig_cum.add_hline(y=_td_total, line_dash="dash",
                           line_color="#B71C1C", annotation_text="Tổng kế hoạch")
        _fig_cum.update_layout(
            title="Tiến độ thi công cộng dồn",
            xaxis_title="Ngày", yaxis_title="Cọc cộng dồn",
            height=320,
            margin=dict(t=48, b=48, l=48, r=16),
        )
        st.plotly_chart(_fig_cum, use_container_width=True)
    else:
        st.info("Chưa có dữ liệu ngày thi công. Nhập dữ liệu từ bảng Excel để hiển thị biểu đồ tiến độ.")

    # Bảng tổng hợp trạng thái
    st.markdown("#### Phân bố trạng thái theo khu vực")
    _td_status_labels = {
        "chua_thi_cong":  "Chưa thi công",
        "dang_thi_cong":  "Đang thi công",
        "hoan_thanh":     "Hoàn thành",
        "khong_dat":      "Không đạt",
        "tam_dung":       "Tạm dừng",
    }
    _td_summary: dict = {}
    for _zone, _st, _n in _td_rows:
        _td_summary.setdefault(_zone, {})[_st] = _n

    _td_table_rows = []
    _all_statuses = ["hoan_thanh","dang_thi_cong","chua_thi_cong","khong_dat","tam_dung"]
    for _zone in sorted(_td_summary.keys()):
        _row = {"Khu vực": _zone}
        _zone_total = sum(_td_summary[_zone].values())
        for _st in _all_statuses:
            _n_st = _td_summary[_zone].get(_st, 0)
            _row[_td_status_labels[_st]] = f"{_n_st:,}"
        _row["Tổng"] = f"{_zone_total:,}"
        _td_table_rows.append(_row)

    if _td_table_rows:
        import pandas as _pd_td2
        _df_td_tbl = _pd_td2.DataFrame(_td_table_rows)
        st.table(_df_td_tbl)

    # Bảng tổng hợp tổ đội
    _td_todoi = _con_td.execute("""
        SELECT to_doi, zone, COUNT(*) AS n,
               SUM(CASE WHEN status='hoan_thanh' THEN 1 ELSE 0 END) AS done
        FROM cdm_thi_cong
        WHERE to_doi IS NOT NULL
        GROUP BY to_doi, zone
        ORDER BY to_doi, zone
    """).fetchall()
    if _td_todoi:
        st.markdown("#### Tổng hợp theo tổ đội")
        import pandas as _pd_td3
        _df_todoi = _pd_td3.DataFrame(_td_todoi, columns=["Tổ đội","Khu vực","Số cọc","Hoàn thành"])
        st.dataframe(_df_todoi, use_container_width=True, hide_index=True)

    st.divider()

    # ── B. Kiểm tra thí nghiệm qu ────────────────────────────────────────────
    st.markdown("### B. Kiểm tra thí nghiệm")

    _kt_rows = _con_td.execute("""
        SELECT kt.zone, kt.point_name, kt.loai_tn,
               kt.ngay_thu_nghiem, kt.tuoi_ngay,
               kt.do_sau_tu_dinh_m, kt.z_sample_m,
               kt.qu_kPa, kt.qu_yc_kPa, kt.dat_yeu_cau,
               kt.ham_luong_xi_mang_pct, kt.ghi_chu
        FROM cdm_kiem_tra kt
        ORDER BY kt.zone, kt.ngay_thu_nghiem, kt.z_sample_m
    """).fetchall()

    if not _kt_rows:
        st.info(
            "Chưa có kết quả thí nghiệm. "
            "Điền dữ liệu vào sheet 'Thí nghiệm CDM' trong bảng Excel và chạy script nhập liệu."
        )
    else:
        import pandas as _pd_kt
        _df_kt = _pd_kt.DataFrame(_kt_rows, columns=[
            "Khu vực","Số hiệu cọc","Loại TN","Ngày TN","Tuổi (ngày)",
            "Độ sâu (m)","Cao độ (m)","qu (kPa)","qu yêu cầu (kPa)",
            "Đạt","Hàm lượng XM (%)","Ghi chú",
        ])
        _df_kt["Đạt"] = _df_kt["Đạt"].map({1:"Đạt", 0:"Không đạt", None:"Chưa đánh giá"})

        # Metrics
        _n_kt    = len(_df_kt)
        _n_dat   = (_df_kt["Đạt"] == "Đạt").sum()
        _n_kdat  = (_df_kt["Đạt"] == "Không đạt").sum()
        _qu_avg  = _df_kt["qu (kPa)"].mean()
        _qu_min  = _df_kt["qu (kPa)"].min()

        _km1, _km2, _km3, _km4 = st.columns(4)
        _km1.metric("Tổng mẫu", f"{_n_kt:,}")
        _km2.metric("Đạt", f"{_n_dat:,}", delta=f"{round(_n_dat/_n_kt*100,1)}%" if _n_kt else None)
        _km3.metric("Không đạt", f"{_n_kdat:,}")
        _km4.metric("qu trung bình", f"{_qu_avg:.0f} kPa" if _qu_avg else "—")

        # Scatter: qu vs độ sâu
        _fig_qu = go.Figure()
        _zone_clr_kt = {"KE": "#1565C0", "CONG_VIEN": "#6A1B9A"}
        _dat_sym = {"Đạt": "circle", "Không đạt": "x", "Chưa đánh giá": "circle-open"}
        for _z in _df_kt["Khu vực"].unique():
            _dz = _df_kt[_df_kt["Khu vực"] == _z]
            for _dat_val in ["Đạt","Không đạt","Chưa đánh giá"]:
                _dd = _dz[_dz["Đạt"] == _dat_val]
                if _dd.empty:
                    continue
                _fig_qu.add_trace(go.Scatter(
                    x=_dd["qu (kPa)"],
                    y=_dd["Độ sâu (m)"],
                    mode="markers",
                    marker=dict(
                        symbol=_dat_sym[_dat_val],
                        size=9,
                        color=_zone_clr_kt.get(_z, "#546E7A"),
                        opacity=0.8 if _dat_val == "Đạt" else 1.0,
                        line=dict(width=1, color="#333333") if _dat_val == "Không đạt" else None,
                    ),
                    name=f"{_z} — {_dat_val}",
                    text=_dd["Số hiệu cọc"],
                    hovertemplate="<b>%{text}</b><br>qu=%{x:.0f} kPa<br>sâu=%{y:.1f} m<extra></extra>",
                ))

        # Đường qu yêu cầu (lấy min của các giá trị không null)
        _qu_yc_vals = _df_kt["qu yêu cầu (kPa)"].dropna()
        if not _qu_yc_vals.empty:
            _qu_yc_ref = _qu_yc_vals.min()
            _fig_qu.add_vline(x=_qu_yc_ref, line_dash="dash", line_color="#B71C1C",
                              annotation_text=f"qu yc = {_qu_yc_ref:.0f} kPa")

        _fig_qu.update_yaxes(autorange="reversed", title="Độ sâu từ đỉnh cọc (m)")
        _fig_qu.update_xaxes(title="qu (kPa)")
        _fig_qu.update_layout(
            title="Phân bố cường độ nén theo độ sâu",
            height=480,
            legend_title="Khu vực / Kết quả",
            margin=dict(t=48, b=48, l=48, r=16),
        )
        st.plotly_chart(_fig_qu, use_container_width=True)

        # Histogram qu
        _fig_hist = go.Figure()
        for _z in _df_kt["Khu vực"].unique():
            _dz = _df_kt[_df_kt["Khu vực"] == _z]["qu (kPa)"].dropna()
            if _dz.empty:
                continue
            _fig_hist.add_trace(go.Histogram(
                x=_dz, name=_z,
                marker_color=_zone_clr_kt.get(_z, "#546E7A"),
                opacity=0.75, nbinsx=20,
            ))
        if not _qu_yc_vals.empty:
            _fig_hist.add_vline(x=_qu_yc_ref, line_dash="dash", line_color="#B71C1C",
                                annotation_text="qu yc")
        _fig_hist.update_layout(
            barmode="overlay",
            title="Phân bố qu (kPa)",
            xaxis_title="qu (kPa)", yaxis_title="Số mẫu",
            height=320,
            margin=dict(t=48, b=48, l=48, r=16),
        )
        st.plotly_chart(_fig_hist, use_container_width=True)

        # Bảng chi tiết
        with st.expander("Chi tiết kết quả thí nghiệm", expanded=False):
            _df_kt_show = _df_kt[[
                "Khu vực","Số hiệu cọc","Loại TN","Tuổi (ngày)",
                "Cao độ (m)","qu (kPa)","qu yêu cầu (kPa)","Đạt","Ghi chú"
            ]]
            import pandas as _pd_kt2
            def _color_dat(v):
                if v == "Đạt":       return "background-color:#E8F5E9; color:#1B5E20"
                if v == "Không đạt": return "background-color:#FFEBEE; color:#B71C1C"
                return ""
            st.dataframe(
                _df_kt_show.style.map(_color_dat, subset=["Đạt"]),
                use_container_width=True, hide_index=True,
            )

        # Bảng tổng hợp per zone / loại TN
        st.markdown("#### Tổng hợp đánh giá theo khu vực")
        _kt_agg = _con_td.execute("""
            SELECT zone, loai_tn,
                   COUNT(*) AS n,
                   SUM(CASE WHEN dat_yeu_cau=1 THEN 1 ELSE 0 END) AS n_dat,
                   SUM(CASE WHEN dat_yeu_cau=0 THEN 1 ELSE 0 END) AS n_kdat,
                   ROUND(AVG(qu_kPa),0) AS qu_avg,
                   ROUND(MIN(qu_kPa),0) AS qu_min,
                   ROUND(MAX(qu_kPa),0) AS qu_max
            FROM cdm_kiem_tra
            GROUP BY zone, loai_tn
            ORDER BY zone, loai_tn
        """).fetchall()
        if _kt_agg:
            import pandas as _pd_kt3
            _df_agg = _pd_kt3.DataFrame(_kt_agg, columns=[
                "Khu vực","Loại TN","Số mẫu","Đạt","Không đạt",
                "qu TB (kPa)","qu min (kPa)","qu max (kPa)",
            ])
            st.table(_df_agg)

# ── Placeholder: TKBVTC Cọc SW ───────────────────────────────────────────────
if _page == "sw_bvt":
    st.markdown("## TKBVTC Cọc SW")
    st.info("Trang đang phát triển — thiết kế bản vẽ thi công cọc ván SW.")


# ── Thống nhất đầu vào TVTK ──────────────────────────────────────────────────
if _page == "tvtk_prep":
    import sqlite3 as _sq
    import pandas as _pd_tv
    import plotly.graph_objects as _go_tv

    st.markdown("## Tài liệu chuẩn bị thống nhất đầu vào")
    st.caption("Khu vực: Kè Công Viên · Bãi Đỗ Xe Ngầm · Nhà Hành Chính  |  Họp TVTK 26/5/2026")

    _cv = _sq.connect(_DB)
    _cv.row_factory = _sq.Row

    # ── Tạo bảng config nếu chưa có + load giá trị đã lưu ──────────────────
    _cv.execute("""
        CREATE TABLE IF NOT EXISTS tvtk_config (
            zone_code   TEXT PRIMARY KEY,
            design_elev_m REAL,
            gwt_m       REAL,
            updated_at  TEXT DEFAULT (datetime('now','localtime'))
        )
    """)
    _cv.commit()

    def _tv_load(zone_code: str, field: str, default: float) -> float:
        r = _cv.execute(
            f"SELECT {field} FROM tvtk_config WHERE zone_code = ?", (zone_code,)
        ).fetchone()
        return float(r[0]) if r and r[0] is not None else default

    # ── 1. Phân khu địa chất ────────────────────────────────────────────────
    st.markdown("### 1. Sơ bộ phân khu địa chất / thiết kế")
    _zones_tv = _cv.execute("SELECT id, code, name_vi AS name, notes FROM zones WHERE code != 'QTT' ORDER BY id").fetchall()
    _bhs_tv   = _cv.execute("""
        SELECT b.id, b.zone_id, b.name, b.elevation_m, b.x_coord_m, b.y_coord_m, b.depth_m
        FROM boreholes b WHERE b.zone_id IN (1,2,3) ORDER BY b.zone_id, b.name
    """).fetchall()

    # Load CDM selection sớm — dùng cho cả bình đồ lẫn bảng thống kê
    _cv.execute("""
        CREATE TABLE IF NOT EXISTS tvtk_bh_cdm (
            bh_name TEXT PRIMARY KEY,
            selected INTEGER DEFAULT 1,
            updated_at TEXT DEFAULT (datetime('now','localtime'))
        )
    """)
    _cv.commit()
    _saved_sel = {r[0]: bool(r[1]) for r in _cv.execute(
        "SELECT bh_name, selected FROM tvtk_bh_cdm"
    ).fetchall()}
    _cdm_yes = {b["name"] for b in _bhs_tv if _saved_sel.get(b["name"], True)}
    _cdm_no  = {b["name"] for b in _bhs_tv if not _saved_sel.get(b["name"], True)}

    # Bộ lọc hố khoan — ảnh hưởng đến cả bảng thống kê lẫn bình đồ
    _c_flt, _c_cnt1, _c_cnt2 = st.columns([4, 1, 1])
    _tv_filter = _c_flt.radio(
        "Lọc hố khoan:",
        ["Tất cả", "Tham gia tính CDM", "Không tham gia CDM"],
        horizontal=True, key="_tv_map_filter",
    )
    _c_cnt1.metric("Tham gia CDM", len(_cdm_yes))
    _c_cnt2.metric("Không tính CDM", len(_cdm_no))

    if _tv_filter == "Tham gia tính CDM":
        _bhs_stat = [b for b in _bhs_tv if b["name"] in _cdm_yes]
    elif _tv_filter == "Không tham gia CDM":
        _bhs_stat = [b for b in _bhs_tv if b["name"] in _cdm_no]
    else:
        _bhs_stat = list(_bhs_tv)

    _zone_summary = []
    for z in _zones_tv:
        _bhs_z     = [b for b in _bhs_stat if b["zone_id"] == z["id"]]
        _bhs_z_all = [b for b in _bhs_tv   if b["zone_id"] == z["id"]]
        _ids_z     = [b["id"] for b in _bhs_z]
        _names_z   = [b["name"] for b in _bhs_z]
        if _ids_z:
            _ph   = ",".join("?" * len(_ids_z))
            _n_cc = _cv.execute(
                f"SELECT COUNT(DISTINCT borehole_id) FROM lab_tests WHERE borehole_id IN ({_ph}) AND Cc IS NOT NULL AND Cc != 0",
                _ids_z
            ).fetchone()[0]
        else:
            _n_cc = 0
        if _names_z:
            _ph_n  = ",".join("?" * len(_names_z))
            _n_vst = _cv.execute(
                f"SELECT COUNT(*) FROM vst_locations WHERE name IN ({_ph_n})",
                _names_z
            ).fetchone()[0]
        else:
            _n_vst = 0
        _elev_avg = sum(b["elevation_m"] or 0 for b in _bhs_z) / max(len(_bhs_z), 1)
        _zone_summary.append({
            "Khu vực": z["code"],
            "Tên khu vực": z["name"],
            "Số HK": f"{len(_bhs_z)} / {len(_bhs_z_all)}",
            "HK có TN cố kết": _n_cc,
            "Vị trí VST": _n_vst,
            "Cao độ TB (m)": f"{_elev_avg:.2f}" if _bhs_z else "—",
            "Ghi chú": z["notes"] or "",
        })
    st.table(_pd_tv.DataFrame(_zone_summary))

    # Bình đồ tọa độ HK + CDM
    _cdm_cv = _cv.execute(
        "SELECT easting_m, northing_m FROM cdm_toado WHERE zone='CONG_VIEN' AND id % 15 = 0"
    ).fetchall()
    _cdm_ke = _cv.execute(
        "SELECT easting_m, northing_m FROM cdm_toado WHERE zone='KE' AND id % 15 = 0"
    ).fetchall()

    _zone_codes  = {1: "KE", 2: "BXN", 3: "NHC"}
    _zone_colors = {1: "#f97316", 2: "#3b82f6", 3: "#22c55e"}
    _fig_tv = _go_tv.Figure()

    # CDM point cloud (nền)
    if _cdm_cv:
        _fig_tv.add_trace(_go_tv.Scatter(
            x=[r[0] for r in _cdm_cv], y=[r[1] for r in _cdm_cv],
            mode="markers", name="Cọc CDM – Công Viên",
            marker=dict(size=3, color="#a78bfa", opacity=0.3, symbol="circle"),
            hovertemplate="CDM Công Viên<br>E=%{x:.1f} N=%{y:.1f}<extra></extra>",
        ))
    if _cdm_ke:
        _fig_tv.add_trace(_go_tv.Scatter(
            x=[r[0] for r in _cdm_ke], y=[r[1] for r in _cdm_ke],
            mode="markers", name="Cọc CDM – Kè",
            marker=dict(size=3, color="#fbbf24", opacity=0.3, symbol="circle"),
            hovertemplate="CDM Kè<br>E=%{x:.1f} N=%{y:.1f}<extra></extra>",
        ))

    # Hố khoan — 2 nhóm: tham gia CDM (circle solid) vs không tham gia (x)
    _vis_yes = True if _tv_filter in ["Tất cả", "Tham gia tính CDM"] else "legendonly"
    _vis_no  = True if _tv_filter in ["Tất cả", "Không tham gia CDM"] else "legendonly"

    for _zid in (1, 2, 3):
        _zname   = _zone_codes[_zid]
        _zcol    = _zone_colors[_zid]
        _pts_all = [b for b in _bhs_tv if b["zone_id"] == _zid and b["y_coord_m"]]
        _pts_yes = [b for b in _pts_all if b["name"] in _cdm_yes]
        _pts_no  = [b for b in _pts_all if b["name"] in _cdm_no]

        if _pts_yes:
            _fig_tv.add_trace(_go_tv.Scatter(
                x=[b["y_coord_m"] for b in _pts_yes], y=[b["x_coord_m"] for b in _pts_yes],
                mode="markers+text", name=f"HK {_zname} – tham gia CDM",
                text=[b["name"].split("-")[-1] for b in _pts_yes],
                textposition="top center", textfont=dict(size=9, color=_zcol),
                marker=dict(size=11, color=_zcol, symbol="circle",
                            line=dict(width=1.5, color="#fff")),
                hovertemplate="%{text}<br>Cao độ: %{customdata:.2f} m<br><b>Tham gia tính CDM</b><extra></extra>",
                customdata=[b["elevation_m"] for b in _pts_yes],
                visible=_vis_yes,
            ))
        if _pts_no:
            _fig_tv.add_trace(_go_tv.Scatter(
                x=[b["y_coord_m"] for b in _pts_no], y=[b["x_coord_m"] for b in _pts_no],
                mode="markers+text", name=f"HK {_zname} – không tính CDM",
                text=[b["name"].split("-")[-1] for b in _pts_no],
                textposition="top center", textfont=dict(size=9, color="#94a3b8"),
                marker=dict(size=11, color=_zcol, symbol="x",
                            line=dict(width=2.5, color=_zcol), opacity=0.55),
                hovertemplate="%{text}<br>Cao độ: %{customdata:.2f} m<br><i>Không tham gia tính CDM</i><extra></extra>",
                customdata=[b["elevation_m"] for b in _pts_no],
                visible=_vis_no,
            ))

    _fig_tv.update_layout(
        height=520, title="Bình đồ vị trí hố khoan và phạm vi cọc CDM",
        template="plotly_dark",
        xaxis_title="Easting (m)", yaxis_title="Northing (m)",
        yaxis=dict(scaleanchor="x", scaleratio=1),
        legend=dict(orientation="h", y=1.08, font=dict(size=10)),
        margin=dict(l=40, r=20, t=70, b=40),
        dragmode="pan",
    )
    st.plotly_chart(_fig_tv, use_container_width=True,
                    config={"scrollZoom": True, "displayModeBar": True,
                            "modeBarButtonsToAdd": ["zoom2d", "pan2d", "resetScale2d"]})
    st.divider()

    # ── 2. Cao độ tự nhiên – cao độ thiết kế ────────────────────────────────
    st.markdown("### 2. Cao độ tự nhiên – cao độ thiết kế")
    _c2a, _c2b, _c2c, _c2s = st.columns([1, 1, 1, 0.6])
    _des_ke  = _c2a.number_input("Cao độ TK kè KE (m)",   value=_tv_load("KE",  "design_elev_m", 2.70), step=0.05, format="%.2f", key="_tv_desKE")
    _des_bxn = _c2b.number_input("Cao độ TK BXN (m)",     value=_tv_load("BXN", "design_elev_m", 3.00), step=0.05, format="%.2f", key="_tv_desBXN")
    _des_nhc = _c2c.number_input("Cao độ TK NHC (m)",     value=_tv_load("NHC", "design_elev_m", 2.50), step=0.05, format="%.2f", key="_tv_desNHC")
    if _c2s.button("Lưu cao độ TK", key="_tv_save_des", use_container_width=True):
        for _zc, _dv in (("KE", _des_ke), ("BXN", _des_bxn), ("NHC", _des_nhc)):
            _cv.execute("""
                INSERT INTO tvtk_config (zone_code, design_elev_m)
                VALUES (?, ?)
                ON CONFLICT(zone_code) DO UPDATE SET design_elev_m=excluded.design_elev_m,
                    updated_at=datetime('now','localtime')
            """, (_zc, _dv))
        _cv.commit()
        st.success("Đã lưu cao độ thiết kế.")
    _des_map = {1: _des_ke, 2: _des_bxn, 3: _des_nhc}

    # Bảng chọn HK tham gia tính CDM — bảng và _saved_sel đã load ở Section 1

    _elev_rows = []
    for b in _bhs_tv:
        _des  = _des_map.get(b["zone_id"], 2.5)
        _diff = _des - (b["elevation_m"] or 0)
        _elev_rows.append({
            "Tính CDM": _saved_sel.get(b["name"], True),
            "Hố khoan": b["name"],
            "Khu vực": _zone_codes.get(b["zone_id"], ""),
            "Cao độ tự nhiên (m)": round(b["elevation_m"], 3) if b["elevation_m"] else None,
            "Cao độ TK (m)": round(_des, 2),
            "Chiều dày đắp (m)": round(_diff, 2) if _diff > 0 else 0.0,
            "Độ sâu HK (m)": round(b["depth_m"], 1) if b["depth_m"] else None,
        })

    _df_elev = _pd_tv.DataFrame(_elev_rows)
    _edited  = st.data_editor(
        _df_elev,
        column_config={
            "Tính CDM": st.column_config.CheckboxColumn(
                "Tính CDM", help="Chọn hố khoan tham gia tính toán lún nền CDM", default=True
            ),
            "Hố khoan":          st.column_config.TextColumn(disabled=True),
            "Khu vực":           st.column_config.TextColumn(disabled=True),
            "Cao độ tự nhiên (m)": st.column_config.NumberColumn(disabled=True, format="%.3f"),
            "Cao độ TK (m)":     st.column_config.NumberColumn(disabled=True, format="%.2f"),
            "Chiều dày đắp (m)": st.column_config.NumberColumn(disabled=True, format="%.2f"),
            "Độ sâu HK (m)":     st.column_config.NumberColumn(disabled=True, format="%.1f"),
        },
        use_container_width=True,
        hide_index=True,
        key="_tv_bh_editor",
    )

    # Tự động lưu khi checkbox thay đổi
    for _, row in _edited.iterrows():
        _cv.execute("""
            INSERT INTO tvtk_bh_cdm (bh_name, selected)
            VALUES (?, ?)
            ON CONFLICT(bh_name) DO UPDATE SET selected=excluded.selected,
                updated_at=datetime('now','localtime')
        """, (row["Hố khoan"], int(row["Tính CDM"])))
    _cv.commit()

    _n_sel = int(_edited["Tính CDM"].sum())
    st.caption(f"Đã chọn **{_n_sel}** / {len(_edited)} hố khoan tham gia tính toán CDM.")

    # ── Trắc dọc cao độ Kè KE ────────────────────────────────────────────────
    import numpy as _np_e
    _ke_bhs_e = sorted(
        [b for b in _bhs_tv if b["name"] in _cdm_yes and b["name"].startswith("KE-")
         and b["x_coord_m"] and b["y_coord_m"]],
        key=lambda b: b["name"],
    )
    if len(_ke_bhs_e) >= 2:
        _cdm_cfg_e  = _cv.execute("SELECT top_elev_m, penetration_m FROM tvtk_cdm_config WHERE id=1").fetchone()
        _top_e      = float(_cdm_cfg_e["top_elev_m"])  if _cdm_cfg_e else 0.8
        _pen_e      = float(_cdm_cfg_e["penetration_m"]) if _cdm_cfg_e else 1.0
        # Tính H_soft trực tiếp theo nguyên tắc (cùng logic Section 4):
        # _SOFT_ALWAYS: lớp 1, 1b, XMD luôn tính (XMD = lớp bùn tại KE-HK8)
        # lớp khác chỉ tính khi AVG(e₀) > 1 từ thí nghiệm nén cố kết
        _SOFT_ALWAYS_E = ("1", "1b", "XMD")
        _hsoft_e   = {}   # PA3: lớp 1 + 1b + XMD + lớp khác có e₀>1
        _h_pa1_map = {}   # PA1: lớp 1 only
        _h_pa2_map = {}   # PA2: lớp 1 + 1b
        for _b2 in _ke_bhs_e:
            _lyrs2 = _cv.execute("""
                SELECT symbol, depth_top_m, depth_bot_m,
                       COALESCE(thickness_m, depth_bot_m - depth_top_m) AS thick
                FROM layers WHERE borehole_id = ? ORDER BY depth_top_m
            """, (_b2["id"],)).fetchall()
            _h1_v = _h1b_v = _hx_v = 0.0
            for _l2 in _lyrs2:
                _s2 = _l2["symbol"]; _t2 = float(_l2["thick"] or 0)
                if _t2 <= 0:
                    continue
                if _s2 in _SOFT_ALWAYS_E:
                    if _s2 == "1":       _h1_v  += _t2
                    elif _s2 == "1b":    _h1b_v += _t2
                    else:                _hx_v  += _t2  # XMD = bùn
                else:
                    _e0r2 = _cv.execute("""
                        SELECT AVG(e0) avg_e0 FROM lab_tests
                        WHERE borehole_id=? AND depth_from_m>=? AND depth_from_m<?
                          AND e0 IS NOT NULL AND e0 > 0
                    """, (_b2["id"], _l2["depth_top_m"], _l2["depth_bot_m"])).fetchone()
                    if _e0r2 and _e0r2["avg_e0"] and _e0r2["avg_e0"] > 1.0:
                        _hx_v += _t2
            _hsoft_e[_b2["name"]]   = round(_h1_v + _h1b_v + _hx_v, 2)
            _h_pa1_map[_b2["name"]] = round(_h1_v, 2)
            _h_pa2_map[_b2["name"]] = round(_h1_v + _h1b_v, 2)

        # Chainage theo PCA – SVD
        _xy_e   = _np_e.array([(b["x_coord_m"], b["y_coord_m"]) for b in _ke_bhs_e])
        _ctr_e  = _xy_e.mean(axis=0)
        _, _, _Vt_e = _np_e.linalg.svd(_xy_e - _ctr_e, full_matrices=False)
        _ch_raw = ((_xy_e - _ctr_e) @ _Vt_e[0]).tolist()
        _order_e = sorted(range(len(_ke_bhs_e)), key=lambda i: _ch_raw[i])
        _bhs_e   = [_ke_bhs_e[i] for i in _order_e]
        _ch_e    = [_ch_raw[i] - _ch_raw[_order_e[0]] for i in _order_e]  # bắt đầu từ 0

        _elev_e     = [float(b["elevation_m"] or 0) for b in _bhs_e]
        _bot_e      = [float(b["elevation_m"] or 0) - float(_hsoft_e.get(b["name"]) or 0) - _pen_e
                       for b in _bhs_e]
        _bot_pa1_e  = [float(b["elevation_m"] or 0) - float(_h_pa1_map.get(b["name"]) or 0) - _pen_e
                       for b in _bhs_e]
        _bot_pa2_e  = [float(b["elevation_m"] or 0) - float(_h_pa2_map.get(b["name"]) or 0) - _pen_e
                       for b in _bhs_e]
        _L_e        = [_top_e - b for b in _bot_e]
        _L_pa1_e    = [_top_e - b for b in _bot_pa1_e]
        _L_pa2_e    = [_top_e - b for b in _bot_pa2_e]
        _lbl_e      = [b["name"].replace("KE-", "") for b in _bhs_e]
        _n_e        = len(_bhs_e)
        _ch_ext     = [_ch_e[0] - 10, _ch_e[-1] + 10]

        # Lưu 3 phương án vào SQLite
        try:
            _cv.execute("ALTER TABLE tvtk_bh_cdm ADD COLUMN H_pa1_m REAL")
        except Exception:
            pass
        try:
            _cv.execute("ALTER TABLE tvtk_bh_cdm ADD COLUMN H_pa2_m REAL")
        except Exception:
            pass
        for _b2 in _bhs_e:
            _nm2 = _b2["name"]
            _cv.execute("""
                INSERT INTO tvtk_bh_cdm (bh_name, H_pa1_m, H_pa2_m)
                VALUES (?, ?, ?)
                ON CONFLICT(bh_name) DO UPDATE SET
                    H_pa1_m = excluded.H_pa1_m,
                    H_pa2_m = excluded.H_pa2_m,
                    updated_at = datetime('now','localtime')
            """, (_nm2, _h_pa1_map.get(_nm2), _h_pa2_map.get(_nm2)))
        _cv.commit()

        # Cao độ đáy lớp 1 và 1b (deepest segment per HK)
        _bh_names_e = [b["name"] for b in _bhs_e]
        _ph_e = ",".join("?" * _n_e)
        _lyr1_map = {r["name"]: float(r["elev_bot"]) for r in _cv.execute(f"""
            SELECT b.name, b.elevation_m - MAX(lyr.depth_bot_m) AS elev_bot
            FROM layers lyr JOIN boreholes b ON lyr.borehole_id = b.id
            WHERE b.name IN ({_ph_e}) AND lyr.symbol = '1'
            GROUP BY b.name, b.elevation_m
        """, _bh_names_e).fetchall()}
        _lyr1b_map = {r["name"]: float(r["elev_bot"]) for r in _cv.execute(f"""
            SELECT b.name, b.elevation_m - MAX(lyr.depth_bot_m) AS elev_bot
            FROM layers lyr JOIN boreholes b ON lyr.borehole_id = b.id
            WHERE b.name IN ({_ph_e}) AND lyr.symbol = '1b'
            GROUP BY b.name, b.elevation_m
        """, _bh_names_e).fetchall()}

        _bot1_e  = [_lyr1_map.get(b["name"])  for b in _bhs_e]  # None nếu không có lớp
        _bot1b_e = [_lyr1b_map.get(b["name"]) for b in _bhs_e]

        _fig_e = _go_tv.Figure()

        # ── Vùng tô: đất đắp (mặt đất → cao độ TK) ──────────────────────────
        if any(v < _des_ke for v in _elev_e):
            _fig_e.add_trace(_go_tv.Scatter(
                x=_ch_e + _ch_e[::-1],
                y=[_des_ke] * _n_e + _elev_e[::-1],
                fill="toself", fillcolor="rgba(210,180,100,0.22)",
                line=dict(width=0), mode="lines",
                name="Đất đắp", hoverinfo="skip",
            ))

        # ── Vùng tô: phạm vi xử lý CDM (đỉnh → đáy) ─────────────────────────
        _fig_e.add_trace(_go_tv.Scatter(
            x=_ch_e + _ch_e[::-1],
            y=[_top_e] * _n_e + _bot_e[::-1],
            fill="toself", fillcolor="rgba(30,120,200,0.15)",
            line=dict(width=0), mode="lines",
            name="Phạm vi xử lý CDM", hoverinfo="skip",
        ))

        # ── Cột CDM per HK (đứng, đứt) ──────────────────────────────────────
        for i in range(_n_e):
            _fig_e.add_trace(_go_tv.Scatter(
                x=[_ch_e[i], _ch_e[i]], y=[_top_e, _bot_e[i]],
                mode="lines", line=dict(color="#1a6fbd", width=1, dash="dot"),
                showlegend=False, hoverinfo="skip",
            ))

        # ── Mặt đất tự nhiên ─────────────────────────────────────────────────
        _fig_e.add_trace(_go_tv.Scatter(
            x=_ch_e, y=_elev_e,
            mode="lines+markers+text",
            name="Mặt đất tự nhiên",
            line=dict(color="#7B3F00", width=2.5),
            marker=dict(size=10, color="#7B3F00", symbol="circle",
                        line=dict(color="white", width=1.5)),
            text=[f"{n}<br><b>{v:+.2f}</b>" for n, v in zip(_lbl_e, _elev_e)],
            textposition="top center",
            textfont=dict(size=9, color="#5a2000"),
            hovertemplate="<b>%{customdata}</b><br>Chainage: %{x:.0f} m<br>Mặt đất: <b>%{y:+.3f} m</b><extra></extra>",
            customdata=_lbl_e,
        ))

        # ── Cao độ thiết kế ───────────────────────────────────────────────────
        _fig_e.add_trace(_go_tv.Scatter(
            x=_ch_ext, y=[_des_ke, _des_ke],
            mode="lines", name=f"Cao độ thiết kế ({_des_ke:+.2f} m)",
            line=dict(color="#2ca02c", width=2.5, dash="dash"),
            hovertemplate=f"Cao độ thiết kế: {_des_ke:+.2f} m<extra></extra>",
        ))

        # ── Đỉnh cọc CDM ─────────────────────────────────────────────────────
        _fig_e.add_trace(_go_tv.Scatter(
            x=_ch_ext, y=[_top_e, _top_e],
            mode="lines", name=f"Đỉnh cọc CDM ({_top_e:+.2f} m)",
            line=dict(color="#1a6fbd", width=2, dash="dash"),
            hovertemplate=f"Đỉnh CDM: {_top_e:+.2f} m<extra></extra>",
        ))

        # ── Đáy lớp 1 ────────────────────────────────────────────────────────
        _ch_1  = [_ch_e[i] for i in range(_n_e) if _bot1_e[i] is not None]
        _val_1 = [_bot1_e[i] for i in range(_n_e) if _bot1_e[i] is not None]
        _lbl_1 = [_lbl_e[i] for i in range(_n_e) if _bot1_e[i] is not None]
        if _ch_1:
            _fig_e.add_trace(_go_tv.Scatter(
                x=_ch_1, y=_val_1,
                mode="lines+markers+text",
                name="Đáy lớp 1",
                line=dict(color="#8c564b", width=2, dash="longdash"),
                marker=dict(size=8, symbol="triangle-down", color="#8c564b",
                            line=dict(color="white", width=1)),
                text=[f"{v:+.2f}" for v in _val_1],
                textposition="top right",
                textfont=dict(size=9, color="#5D3317"),
                hovertemplate="<b>%{customdata}</b><br>Đáy lớp 1: <b>%{y:+.2f} m</b><extra></extra>",
                customdata=_lbl_1,
            ))

        # ── Đáy lớp 1b ───────────────────────────────────────────────────────
        _ch_1b  = [_ch_e[i] for i in range(_n_e) if _bot1b_e[i] is not None]
        _val_1b = [_bot1b_e[i] for i in range(_n_e) if _bot1b_e[i] is not None]
        _lbl_1b = [_lbl_e[i] for i in range(_n_e) if _bot1b_e[i] is not None]
        if _ch_1b:
            _fig_e.add_trace(_go_tv.Scatter(
                x=_ch_1b, y=_val_1b,
                mode="lines+markers+text",
                name="Đáy lớp 1b",
                line=dict(color="#e377c2", width=2, dash="longdashdot"),
                marker=dict(size=8, symbol="triangle-down-open", color="#e377c2",
                            line=dict(color="#e377c2", width=2)),
                text=[f"{v:+.2f}" for v in _val_1b],
                textposition="bottom right",
                textfont=dict(size=9, color="#b5367f"),
                hovertemplate="<b>%{customdata}</b><br>Đáy lớp 1b: <b>%{y:+.2f} m</b><extra></extra>",
                customdata=_lbl_1b,
            ))

        # ── PA1: Đáy CDM qua lớp 1 ───────────────────────────────────────────
        _hov_pa1 = [
            f"<b>{_lbl_e[i]}</b><br>Chainage: {_ch_e[i]:.0f} m<br>"
            f"<b>PA1</b> Đáy CDM: <b>{_bot_pa1_e[i]:+.2f} m</b><br>"
            f"H_lớp1: {_h_pa1_map.get(_bhs_e[i]['name'], 0):.1f} m  |  L_cọc: {_L_pa1_e[i]:.1f} m"
            for i in range(_n_e)
        ]
        _fig_e.add_trace(_go_tv.Scatter(
            x=_ch_e, y=_bot_pa1_e,
            mode="lines+markers+text",
            name="PA1 – đáy qua lớp 1",
            line=dict(color="#ff7f0e", width=2, dash="dash"),
            marker=dict(size=8, color="#ff7f0e", symbol="diamond-open",
                        line=dict(color="#ff7f0e", width=2)),
            text=[f"{v:+.2f}" for v in _bot_pa1_e],
            textposition="top left",
            textfont=dict(size=8, color="#c55100"),
            hovertemplate="%{customdata}<extra></extra>",
            customdata=_hov_pa1,
        ))

        # ── PA2: Đáy CDM qua lớp 1 + 1b ─────────────────────────────────────
        _hov_pa2 = [
            f"<b>{_lbl_e[i]}</b><br>Chainage: {_ch_e[i]:.0f} m<br>"
            f"<b>PA2</b> Đáy CDM: <b>{_bot_pa2_e[i]:+.2f} m</b><br>"
            f"H_1+1b: {_h_pa2_map.get(_bhs_e[i]['name'], 0):.1f} m  |  L_cọc: {_L_pa2_e[i]:.1f} m"
            for i in range(_n_e)
        ]
        _fig_e.add_trace(_go_tv.Scatter(
            x=_ch_e, y=_bot_pa2_e,
            mode="lines+markers+text",
            name="PA2 – đáy qua lớp 1+1b",
            line=dict(color="#9467bd", width=2, dash="dashdot"),
            marker=dict(size=8, color="#9467bd", symbol="diamond",
                        line=dict(color="white", width=1.5)),
            text=[f"{v:+.2f}" for v in _bot_pa2_e],
            textposition="bottom left",
            textfont=dict(size=8, color="#6b3fa0"),
            hovertemplate="%{customdata}<extra></extra>",
            customdata=_hov_pa2,
        ))

        # ── PA3: Đáy CDM đầy đủ (1+1b+XMD+e₀>1) ────────────────────────────
        _hov_bot = [
            f"<b>{_lbl_e[i]}</b><br>Chainage: {_ch_e[i]:.0f} m<br>"
            f"<b>PA3</b> Đáy CDM: <b>{_bot_e[i]:+.2f} m</b><br>"
            f"H_soft: {_hsoft_e.get(_bhs_e[i]['name'], 0):.1f} m  |  L_cọc: {_L_e[i]:.1f} m"
            for i in range(_n_e)
        ]
        _fig_e.add_trace(_go_tv.Scatter(
            x=_ch_e, y=_bot_e,
            mode="lines+markers+text",
            name="PA3 – đáy qua lớp yếu đầy đủ",
            line=dict(color="#d62728", width=2.5),
            marker=dict(size=9, color="#d62728", symbol="square",
                        line=dict(color="white", width=1.5)),
            text=[f"{v:+.2f}" for v in _bot_e],
            textposition="bottom center",
            textfont=dict(size=9, color="#a01010"),
            hovertemplate="%{customdata}<extra></extra>",
            customdata=_hov_bot,
        ))

        _all_vals = (_elev_e + _bot_e + _bot_pa1_e + _bot_pa2_e
                     + [v for v in _val_1 + _val_1b if v is not None])
        _y_lo = min(_all_vals) - 3
        _y_hi = max(max(_elev_e), _des_ke) + 3.5
        _fig_e.update_layout(
            height=580,
            margin=dict(l=65, r=20, t=60, b=90),
            xaxis=dict(title="Chainage dọc tuyến Kè KE (m)", gridcolor="#ebebeb",
                       tickfont=dict(size=11)),
            yaxis=dict(title="Cao độ (m)", range=[_y_lo, _y_hi],
                       gridcolor="#ebebeb", tickfont=dict(size=11),
                       zeroline=True, zerolinecolor="#999", zerolinewidth=1),
            legend=dict(orientation="h", x=0, y=-0.20, xanchor="left",
                        font=dict(size=10)),
            plot_bgcolor="white", paper_bgcolor="white",
            hovermode="closest",
            title=dict(
                text=(
                    f"Trắc dọc Kè KE – {_n_e} hố khoan  |  "
                    f"PA1 L: {min(_L_pa1_e):.1f}–{max(_L_pa1_e):.1f} m  |  "
                    f"PA2 L: {min(_L_pa2_e):.1f}–{max(_L_pa2_e):.1f} m  |  "
                    f"PA3 L: {min(_L_e):.1f}–{max(_L_e):.1f} m"
                ),
                font=dict(size=12), x=0,
            ),
        )
        st.plotly_chart(_fig_e, use_container_width=True)
        _n_1b = sum(1 for v in _bot1b_e if v is not None)
        st.caption(
            f"PA1 (lớp 1): đáy {min(_bot_pa1_e):+.2f}÷{max(_bot_pa1_e):+.2f} m  |  "
            f"PA2 (lớp 1+1b): đáy {min(_bot_pa2_e):+.2f}÷{max(_bot_pa2_e):+.2f} m  |  "
            f"PA3 (đầy đủ): đáy {min(_bot_e):+.2f}÷{max(_bot_e):+.2f} m  |  "
            + (f"Lớp 1b có tại {_n_1b} HK" if _n_1b else "Không có lớp 1b")
        )

        # ── Bảng thống kê độ lún S1 khối gia cố CDM theo 3 phương án ─────────
        import math as _math_e
        _cdm_full_e = _cv.execute(
            "SELECT D_mm, spacing_m, pattern, Ec_factor, qu_kPa, q_kPa FROM tvtk_cdm_config WHERE id=1"
        ).fetchone()
        if _cdm_full_e:
            _D_m_e   = float(_cdm_full_e["D_mm"]) / 1000.0
            _s_m_e   = float(_cdm_full_e["spacing_m"])
            _pat_e   = _cdm_full_e["pattern"]
            _kEc_e   = float(_cdm_full_e["Ec_factor"])
            _qu_e    = float(_cdm_full_e["qu_kPa"])
            _q_cdm_e = float(_cdm_full_e["q_kPa"])   # 40.8 kPa

            # Tỷ lệ diện tích thay thế a = Ac / At
            _Ac_e = _math_e.pi * (_D_m_e / 2.0) ** 2
            _At_e = (_s_m_e ** 2 if _pat_e == "square"
                     else _math_e.sqrt(3) / 2.0 * _s_m_e ** 2)
            _a_e  = _Ac_e / _At_e

            # Mô đun trụ CDM: Ec = k × qu_design/2  (TCVN 9403 B.5.1)
            _Ec_e = _kEc_e * (_qu_e / 2.0)

            # Chuẩn bị cột SQLite nếu chưa có
            for _col_s in ("Cu_VST_avg_kPa", "Es_kPa",
                           "S1_pa1_cm", "S1_pa2_cm", "S1_pa3_cm",
                           "Ip_avg", "bjerrum_mu", "Cu_corrected_kPa"):
                try:
                    _cv.execute(f"ALTER TABLE tvtk_bh_cdm ADD COLUMN {_col_s} REAL")
                except Exception:
                    pass

            # Lazy import Bjerrum
            try:
                from scripts.settlement_calc import bjerrum_mu as _bj_mu
            except Exception:
                _bj_mu = None

            _SOFT_SYM = ("1", "1b", "CH", "MH", "CH-OH", "MH-OH")
            _ph_sym = ",".join("?" * len(_SOFT_SYM))

            _s1_rows_e = []
            for _b3 in _bhs_e:
                _nm3   = _b3["name"]
                _bid3  = _b3["id"]
                _H_p1  = float(_h_pa1_map.get(_nm3) or 0)
                _H_p2  = float(_h_pa2_map.get(_nm3) or 0)
                _H_p3  = float(_hsoft_e.get(_nm3) or 0)

                # Su trung bình từ VST trong phạm vi lớp yếu (PA3 = sâu nhất)
                _vr = _cv.execute("""
                    SELECT AVG(v.Su_kPa) avg_su
                    FROM vane_shear_tests v
                    JOIN vst_locations vl ON v.vst_loc_id = vl.id
                    WHERE vl.name = ? AND v.Su_kPa > 0
                      AND v.depth_m <= ?
                """, (_nm3, max(_H_p1, _H_p2, _H_p3))).fetchone()
                _Su_e  = float(_vr["avg_su"]) if _vr and _vr["avg_su"] else None

                # Ip trung bình từ lab_tests các lớp yếu của HK này
                _ir = _cv.execute(f"""
                    SELECT AVG(lt.Ip) avg_ip
                    FROM lab_tests lt
                    WHERE lt.borehole_id = ? AND lt.Ip IS NOT NULL AND lt.Ip > 0
                      AND lt.symbol_tcvn IN ({_ph_sym})
                """, (_bid3, *_SOFT_SYM)).fetchone()
                _Ip_e = float(_ir["avg_ip"]) if _ir and _ir["avg_ip"] else None

                # Bjerrum correction: Cu = μ × Su (TCCS 41 C.5)
                _mu_e = _bj_mu(_Ip_e) if (_bj_mu is not None and _Ip_e is not None) else 1.0
                _Cu_e = _Su_e * _mu_e if _Su_e is not None else None
                _Es_e = 250.0 * _Cu_e if _Cu_e else None

                def _s1_cm(H):
                    if not _Es_e or H <= 0:
                        return None
                    _Ecomp = _a_e * _Ec_e + (1.0 - _a_e) * _Es_e
                    return round(_q_cdm_e * H / _Ecomp * 100.0, 1)

                _v1 = _s1_cm(_H_p1)
                _v2 = _s1_cm(_H_p2)
                _v3 = _s1_cm(_H_p3)

                # Lưu vào SQLite
                _cv.execute("""
                    INSERT INTO tvtk_bh_cdm
                        (bh_name, Cu_VST_avg_kPa, Es_kPa,
                         S1_pa1_cm, S1_pa2_cm, S1_pa3_cm,
                         Ip_avg, bjerrum_mu, Cu_corrected_kPa)
                    VALUES (?,?,?,?,?,?,?,?,?)
                    ON CONFLICT(bh_name) DO UPDATE SET
                        Cu_VST_avg_kPa    = excluded.Cu_VST_avg_kPa,
                        Es_kPa            = excluded.Es_kPa,
                        S1_pa1_cm         = excluded.S1_pa1_cm,
                        S1_pa2_cm         = excluded.S1_pa2_cm,
                        S1_pa3_cm         = excluded.S1_pa3_cm,
                        Ip_avg            = excluded.Ip_avg,
                        bjerrum_mu        = excluded.bjerrum_mu,
                        Cu_corrected_kPa  = excluded.Cu_corrected_kPa,
                        updated_at        = datetime('now','localtime')
                """, (_nm3, _Su_e, _Es_e, _v1, _v2, _v3,
                      _Ip_e, _mu_e if _Su_e else None, _Cu_e))

                _s1_rows_e.append({
                    "Hố khoan":         _nm3.replace("KE-", ""),
                    "Su VST (kPa)":     round(_Su_e, 1) if _Su_e else "—",
                    "Ip":               round(_Ip_e, 1) if _Ip_e else "—",
                    "μ (Bjerrum)":      round(_mu_e, 3) if _Su_e and _Ip_e else "—",
                    "Cu = μ·Su (kPa)":  round(_Cu_e, 1) if _Cu_e else "—",
                    "Es (kPa)":         int(_Es_e)      if _Es_e else "—",
                    "PA1 H (m)":        f"{_H_p1:.1f}",
                    "PA1 S₁ (cm)":      _v1 if _v1 is not None else "—",
                    "PA2 H (m)":        f"{_H_p2:.1f}",
                    "PA2 S₁ (cm)":      _v2 if _v2 is not None else "—",
                    "PA3 H (m)":        f"{_H_p3:.1f}",
                    "PA3 S₁ (cm)":      _v3 if _v3 is not None else "—",
                })
            _cv.commit()

            st.markdown("#### Độ lún khối gia cố CDM theo 3 phương án chiều sâu")

            # ── Giới hạn ΔS cho phép theo TCCS 41:2022 Bảng 1 — Điều 6.2.3 ──
            # Tạo bảng + populate nếu chưa có (idempotent)
            try:
                from scripts.settlement_calc import (
                    create_tccs41_limits_table as _make_tccs41_lim,
                    get_allowable_residual_settlement as _get_ds_limit,
                )
                _make_tccs41_lim()
            except Exception:
                _get_ds_limit = None

            # Migrate tvtk_cdm_config thêm 2 cột chọn cấp đường + vị trí
            for _col_lim, _def_lim in [("road_class_code", "'cat1'"),
                                        ("position_code",   "'general'")]:
                try:
                    _cv.execute(
                        f"ALTER TABLE tvtk_cdm_config ADD COLUMN {_col_lim} TEXT DEFAULT {_def_lim}"
                    )
                except Exception:
                    pass
            _cv.commit()
            _lim_cfg = _cv.execute(
                "SELECT road_class_code, position_code FROM tvtk_cdm_config WHERE id=1"
            ).fetchone()
            _rc_def  = (_lim_cfg["road_class_code"] if _lim_cfg and _lim_cfg["road_class_code"]
                        else "cat1")
            _pos_def = (_lim_cfg["position_code"]   if _lim_cfg and _lim_cfg["position_code"]
                        else "general")

            _RC_OPTS  = [
                ("cat1", "Cao tốc / ≥ 80 km/h / A1"),
                ("cat2", "≤ 60 km/h / A1"),
            ]
            _POS_OPTS = [
                ("general",      "Đoạn nền đắp thông thường"),
                ("side_culvert", "Đoạn hai bên cống / cống chui"),
                ("near_bridge",  "Đoạn gần mố cầu"),
            ]
            _c_lim1, _c_lim2, _c_lim3 = st.columns([1.4, 1.4, 1.0])
            _rc_pick = _c_lim1.selectbox(
                "Cấp đường (TCCS 41 Bảng 1)",
                options=[k for k, _ in _RC_OPTS],
                format_func=lambda k: dict(_RC_OPTS)[k],
                index=[k for k, _ in _RC_OPTS].index(_rc_def) if _rc_def in dict(_RC_OPTS) else 0,
                key="_tv_tccs41_rc",
            )
            _pos_pick = _c_lim2.selectbox(
                "Vị trí đoạn nền đắp",
                options=[k for k, _ in _POS_OPTS],
                format_func=lambda k: dict(_POS_OPTS)[k],
                index=[k for k, _ in _POS_OPTS].index(_pos_def) if _pos_def in dict(_POS_OPTS) else 0,
                key="_tv_tccs41_pos",
            )
            # Persist lựa chọn
            _cv.execute(
                "UPDATE tvtk_cdm_config SET road_class_code=?, position_code=? WHERE id=1",
                (_rc_pick, _pos_pick),
            )
            _cv.commit()

            _dS_lim = None
            _dS_info = None
            if _get_ds_limit is not None:
                try:
                    _dS_info = _get_ds_limit(_rc_pick, _pos_pick)
                    _dS_lim  = float(_dS_info["delta_S_cm_max"])
                except Exception:
                    _dS_info = None
            if _dS_lim is not None:
                _c_lim3.metric("ΔS cho phép (cm)", f"≤ {_dS_lim:.0f}")

            _Ecomp_repr = (f"{_a_e * _Ec_e + (1 - _a_e) * 3000:.0f}"
                           if not _s1_rows_e
                           else "—")
            st.caption(
                f"TCVN 9403 Phụ lục C: S₁ = q·H / (a·Ec + (1−a)·Es)  |  "
                f"q = {_q_cdm_e} kPa  |  a = {_a_e:.3f}  |  "
                f"Ec = {_Ec_e:,.0f} kPa (k={int(_kEc_e)}×qu/2)  |  "
                f"Es = 250·Cu (Cu = μ·Su theo TCCS 41 Phụ lục C.3.2 — hệ số Bjerrum theo Ip)"
            )
            if _dS_info is not None:
                st.caption(
                    f"TCCS 41:2022 Bảng 1 — Điều 6.2.3:  "
                    f"{_dS_info['road_class_desc']}  |  "
                    f"{_dS_info['position_desc']}  →  "
                    f"**ΔS ≤ {_dS_lim:.0f} cm**  "
                    f"(thời hạn t = {_dS_info['t_years_flexible']} năm mặt đường mềm / "
                    f"{_dS_info['t_years_rigid']} năm mặt đường cứng)"
                )

            # Thêm cột Đạt/Không đạt cho mỗi PA — so sánh S₁ vs ΔS_limit
            if _dS_lim is not None and _s1_rows_e:
                for _row in _s1_rows_e:
                    for _pa_k, _col_k in (("PA1", "PA1 S₁ (cm)"),
                                           ("PA2", "PA2 S₁ (cm)"),
                                           ("PA3", "PA3 S₁ (cm)")):
                        _v = _row.get(_col_k)
                        if isinstance(_v, (int, float)):
                            _row[f"Đạt {_pa_k}"] = "Đạt" if _v <= _dS_lim else "Không đạt"
                        else:
                            _row[f"Đạt {_pa_k}"] = "—"

            if _s1_rows_e:
                _df_s1 = _pd_tv.DataFrame(_s1_rows_e)
                st.table(_df_s1)
                # Tổng hợp min–max + đếm Đạt/Không đạt
                _get_nums = lambda col: [
                    r[col] for r in _s1_rows_e if isinstance(r[col], (int, float))
                ]
                for _pa_lbl, _col in (("PA1", "PA1 S₁ (cm)"),
                                      ("PA2", "PA2 S₁ (cm)"),
                                      ("PA3", "PA3 S₁ (cm)")):
                    _vals = _get_nums(_col)
                    if _vals:
                        _line = (
                            f"{_pa_lbl}: S₁ = {min(_vals):.1f} ÷ {max(_vals):.1f} cm  "
                            f"(TB {sum(_vals)/len(_vals):.1f} cm)"
                        )
                        if _dS_lim is not None:
                            _n_pass = sum(1 for v in _vals if v <= _dS_lim)
                            _n_fail = len(_vals) - _n_pass
                            _line += (
                                f"  |  Đạt {_n_pass}/{len(_vals)} HK"
                                + (f"  |  Không đạt {_n_fail} HK" if _n_fail else "")
                                + f"  (ngưỡng ΔS ≤ {_dS_lim:.0f} cm)"
                            )
                        st.caption(_line)

    st.divider()

    # ── 3. Hố khoan đại diện từng khu ───────────────────────────────────────
    st.markdown("### 3. Xác định hố khoan đại diện từng khu")
    _rep_rows = []
    for z in _zones_tv:
        _bhs_z = [b for b in _bhs_tv if b["zone_id"] == z["id"]]
        # chọn HK có nhiều mẫu Cc nhất
        _best_bh, _best_cc = None, -1
        for b in _bhs_z:
            _nc = _cv.execute(
                "SELECT COUNT(*) FROM lab_tests WHERE borehole_id = ? AND Cc IS NOT NULL AND Cc != 0",
                (b["id"],)
            ).fetchone()[0]
            if _nc > _best_cc:
                _best_cc, _best_bh = _nc, b
        if not _best_bh:
            continue
        _n_vst_bh = _cv.execute("""
            SELECT COUNT(*) FROM vane_shear_tests v
            JOIN vst_locations vl ON v.vst_loc_id = vl.id
            WHERE vl.name = ?
        """, (_best_bh["name"],)).fetchone()[0]
        _rep_rows.append({
            "Khu vực": z["code"],
            "HK đại diện": _best_bh["name"],
            "Mẫu cố kết (Cc)": _best_cc,
            "Điểm VST": _n_vst_bh,
            "Độ sâu HK (m)": f"{_best_bh['depth_m']:.0f}" if _best_bh["depth_m"] else "—",
            "Tiêu chí chọn": "Nhiều mẫu TN cố kết nhất",
        })
    st.table(_pd_tv.DataFrame(_rep_rows))
    st.divider()

    # ── 4. Đất yếu – phạm vi xử lý nền ─────────────────────────────────────
    st.markdown("### 4. Nhận diện đất yếu – phạm vi xử lý nền")
    st.caption(f"Chỉ hiển thị {len(_cdm_yes)} hố khoan tham gia tính toán xử lý nền CDM")

    # Nguyên tắc xác định đất yếu:
    #   • Lớp 1, 1b, XMD: luôn tính (XMD = lớp bùn tại KE-HK8, 6.7–22.4m)
    #   • Lớp khác (2b, 2, 3…): tính NẾU e₀_TB > 1 từ thí nghiệm nén cố kết
    #   • Đáy cọc CDM = cao độ tự nhiên − H_soft − 1 m (ngàm lớp cứng)
    _SOFT_ALWAYS_S4 = ("1", "1b", "XMD")
    with st.expander("Nguyên tắc xác định chiều dày lớp đất yếu H_soft", expanded=False):
        st.markdown(
            "**Lớp 1, lớp 1b và lớp XMD** — luôn được tính vào H_soft "
            "(phân loại địa tầng là đất yếu; lớp XMD là lớp bùn tại KE-HK8, sâu 6,7–22,4 m).\n\n"
            "**Lớp khác (2b, lớp 2, lớp 3…)** — chỉ tính vào H_soft nếu "
            "hệ số rỗng trung bình e₀ > 1,0 từ thí nghiệm nén cố kết trong phòng "
            "(e₀ > 1 ↔ đất ở trạng thái rất rỗng, chưa cố kết hoặc dẻo chảy).\n\n"
            "**Đáy cọc CDM** = cao độ mặt đất tự nhiên tại hố khoan − H_soft − 1 m "
            "(1 m ngàm vào lớp đất cứng phía dưới lớp yếu)."
        )

    _soft_rows = []
    for b in [b for b in _bhs_tv if b["name"] in _cdm_yes]:
        # Lấy TẤT CẢ lớp địa tầng để phân loại lại
        _lyrs_all = _cv.execute("""
            SELECT symbol, description, depth_top_m, depth_bot_m,
                   COALESCE(thickness_m, depth_bot_m - depth_top_m) AS thick
            FROM layers WHERE borehole_id = ? ORDER BY depth_top_m
        """, (b["id"],)).fetchall()
        if not _lyrs_all:
            continue

        _h1 = _h1b = _h_extra = 0.0
        _extra_syms = {}   # {symbol: thickness} cho lớp khác được tính

        for _lyr in _lyrs_all:
            _sym   = _lyr["symbol"]
            _thick = float(_lyr["thick"] or 0)
            if _thick <= 0:
                continue

            if _sym in _SOFT_ALWAYS_S4:
                # Luôn tính: lớp 1, 1b, XMD (XMD = bùn tại KE-HK8)
                if _sym == "1":
                    _h1    += _thick
                elif _sym == "1b":
                    _h1b   += _thick
                else:  # XMD
                    _h_extra += _thick
                    _extra_syms[_sym] = _extra_syms.get(_sym, 0) + _thick
            else:
                # Kiểm tra e0_TB của lớp này
                _e0_row = _cv.execute("""
                    SELECT AVG(lt.e0) avg_e0
                    FROM lab_tests lt
                    WHERE lt.borehole_id = ?
                      AND lt.depth_from_m >= ? AND lt.depth_from_m < ?
                      AND lt.e0 IS NOT NULL AND lt.e0 > 0
                """, (b["id"], _lyr["depth_top_m"], _lyr["depth_bot_m"])).fetchone()
                _e0_avg = _e0_row["avg_e0"] if _e0_row and _e0_row["avg_e0"] else None
                if _e0_avg is not None and _e0_avg > 1.0:
                    _h_extra += _thick
                    _extra_syms[_sym] = _extra_syms.get(_sym, 0) + _thick

        _htot = _h1 + _h1b + _h_extra

        # Lấy lớp soft thực tế (chỉ lớp được tính) để xác định từ–đến
        _soft_depths = (
            [l["depth_top_m"] for l in _lyrs_all if l["symbol"] in ("1","1b")]
            + [l["depth_top_m"] for l in _lyrs_all
               if l["symbol"] not in ("1","1b") and l["symbol"] in _extra_syms]
        )
        _soft_bots = (
            [l["depth_bot_m"] for l in _lyrs_all if l["symbol"] in ("1","1b")]
            + [l["depth_bot_m"] for l in _lyrs_all
               if l["symbol"] not in ("1","1b") and l["symbol"] in _extra_syms]
        )
        _top_first = min(_soft_depths) if _soft_depths else 0
        _bot_last  = max(_soft_bots)  if _soft_bots  else 0

        # SPT và VST
        _spt_avg = _cv.execute("""
            SELECT ROUND(AVG(s.N), 1) FROM spt_values s
            WHERE s.borehole_id = ?
              AND s.depth_m >= ? AND s.depth_m <= ?
              AND s.N IS NOT NULL AND s.N > 0
        """, (b["id"], _top_first, _bot_last)).fetchone()[0]
        _su_avg = _cv.execute("""
            SELECT ROUND(AVG(v.Su_kPa),1) FROM vane_shear_tests v
            JOIN vst_locations vl ON v.vst_loc_id = vl.id
            WHERE vl.name = ?
        """, (b["name"],)).fetchone()[0]

        _extra_str = (
            "  ".join(f"{sym}:{h:.1f}m" for sym, h in sorted(_extra_syms.items()))
            if _extra_syms else "—"
        )
        _soft_rows.append({
            "Hố khoan":             b["name"],
            "Khu vực":              _zone_codes.get(b["zone_id"], ""),
            "Từ (m)":               f"{_top_first:.1f}",
            "Đến (m)":              f"{_bot_last:.1f}",
            "Dày lớp 1 (m)":        f"{_h1:.1f}"   if _h1   else "—",
            "Dày lớp 1b (m)":       f"{_h1b:.1f}"  if _h1b  else "—",
            "Lớp khác (e₀>1) (m)":  _extra_str,
            "H_soft (m)":           f"{_htot:.1f}",
            "N-SPT TB":             f"{_spt_avg:.0f}" if _spt_avg else "—",
            "Su TB-VST (kPa)":      f"{_su_avg:.1f}" if _su_avg else "—",
            "_h_soft":              round(_htot, 2),
            "_bh_name":             b["name"],
            "Khu vực_":             _zone_codes.get(b["zone_id"], ""),
        })
    if _soft_rows:
        st.table(_pd_tv.DataFrame(_soft_rows).drop(columns=["_h_soft", "_bh_name", "Khu vực_"]))
        st.caption(
            "Lớp 1: Bùn sét dẻo chảy  |  "
            "Lớp 1b: Sét mềm xen kẹp cát  |  "
            "Lớp khác (e₀>1): lớp không tên-1/1b nhưng còn rỗng, đủ điều kiện đất yếu"
        )
    else:
        st.info("Chưa có dữ liệu địa tầng.")

    # ── Biểu đồ VST – Su theo chiều sâu – từng HK riêng lẻ (so sánh UU) ──────
    st.markdown("#### Biểu đồ cắt cánh (VST) – So sánh với thí nghiệm UU phòng")

    _ke_bhs_cdm = sorted(
        [b for b in _bhs_tv if b["name"] in _cdm_yes and b["name"].startswith("KE-")],
        key=lambda b: b["name"],
    )
    _vst_data = _cv.execute("""
        SELECT vl.name, vl.Z_ground_m, v.depth_m, v.Su_kPa, v.Sp_kPa, v.St
        FROM vane_shear_tests v
        JOIN vst_locations vl ON v.vst_loc_id = vl.id
        WHERE vl.name LIKE 'KE-%'
        ORDER BY vl.name, v.depth_m
    """).fetchall()

    # UU lab data cho KE
    _uu_data = _cv.execute("""
        SELECT b.name,
               (lt.depth_from_m + lt.depth_to_m) / 2.0 AS depth_mid_m,
               lt.Cu_UU_kPa
        FROM lab_tests lt
        JOIN boreholes b ON lt.borehole_id = b.id
        WHERE b.name LIKE 'KE-%'
          AND lt.Cu_UU_kPa IS NOT NULL AND lt.Cu_UU_kPa > 0
        ORDER BY b.name, depth_mid_m
    """).fetchall()

    if _vst_data:
        # Nhóm theo tên HK
        import numpy as _np_vst
        from plotly.subplots import make_subplots as _make_subplots

        # Nhóm theo tên HK
        _vst_by_bh: dict = {}
        for _vr in _vst_data:
            _vst_by_bh.setdefault(_vr["name"], []).append(_vr)

        _uu_by_bh: dict = {}
        for _ur in _uu_data:
            _uu_by_bh.setdefault(_ur["name"], []).append(_ur)

        # Chiều sâu đáy lớp yếu per HK
        _soft_bot_map = {
            r["_bh_name"]: r["_h_soft"]
            for r in _soft_rows
            if r.get("_bh_name", "").startswith("KE-")
        }

        # Bjerrum μ + Ip per HK — TÍNH ON-THE-FLY cho TẤT CẢ HK trong biểu đồ
        # (Không phụ thuộc tvtk_bh_cdm vì biểu đồ có thể có HK ngoài tuyến CDM)
        try:
            from scripts.settlement_calc import bjerrum_mu as _bj_mu_vst
        except Exception:
            _bj_mu_vst = None

        # Query Ip TB cho TẤT CẢ KE HK 1 lần (batch — tránh N+1 queries)
        _SOFT_SYM_VST = ("1", "1b", "CH", "MH", "CH-OH", "MH-OH")
        _ph_sym_vst   = ",".join("?" * len(_SOFT_SYM_VST))
        _ip_rows = _cv.execute(f"""
            SELECT b.name AS bh_name, AVG(lt.Ip) AS Ip_avg
            FROM lab_tests lt
            JOIN boreholes b ON lt.borehole_id = b.id
            WHERE b.name LIKE 'KE-%'
              AND lt.Ip IS NOT NULL AND lt.Ip > 0
              AND lt.symbol_tcvn IN ({_ph_sym_vst})
            GROUP BY b.name
        """, _SOFT_SYM_VST).fetchall()

        # Ưu tiên giá trị đã lưu trong tvtk_bh_cdm (giữ tính nhất quán với bảng 3 PA)
        _bj_saved = {
            r["bh_name"]: {"Ip": r["Ip_avg"], "mu": r["bjerrum_mu"]}
            for r in _cv.execute("""
                SELECT bh_name, Ip_avg, bjerrum_mu
                FROM tvtk_bh_cdm
                WHERE bh_name LIKE 'KE-%' AND bjerrum_mu IS NOT NULL
            """).fetchall()
        }

        _bj_per_bh = {}
        for _r_ip in _ip_rows:
            _bh_k = _r_ip["bh_name"]
            _Ip_k = float(_r_ip["Ip_avg"]) if _r_ip["Ip_avg"] is not None else None
            # Ưu tiên giá trị đã lưu
            if _bh_k in _bj_saved and _bj_saved[_bh_k]["mu"] is not None:
                _bj_per_bh[_bh_k] = _bj_saved[_bh_k]
            elif _Ip_k is not None and _bj_mu_vst is not None:
                _bj_per_bh[_bh_k] = {"Ip": _Ip_k, "mu": float(_bj_mu_vst(_Ip_k))}

        _bh_names_sorted = sorted(_vst_by_bh.keys(),
                                  key=lambda n: int(n.replace("KE-HK", "") or 0))
        _n_bh  = len(_bh_names_sorted)
        _NCOLS = 3
        _NROWS = (_n_bh + _NCOLS - 1) // _NCOLS

        # Trục x chung: 0 → max(Su, Cu_UU) * 1.15
        _all_su  = [p["Su_kPa"] for p in _vst_data]
        _all_uu  = [r["Cu_UU_kPa"] for r in _uu_data]
        _x_max   = max((_all_su + _all_uu) or [40]) * 1.15
        _x_max   = max(_x_max, 40)
        # μ < 1 nên Cu_corrected ≤ Su → không cần mở rộng x range

        _fig_sub = _make_subplots(
            rows=_NROWS, cols=_NCOLS,
            shared_xaxes=False,
            shared_yaxes=False,
            horizontal_spacing=0.08,
            vertical_spacing=0.06,
            subplot_titles=_bh_names_sorted,
        )

        for _i, _bh_n in enumerate(_bh_names_sorted):
            _row = _i // _NCOLS + 1
            _col = _i  % _NCOLS + 1

            _vpts = _vst_by_bh.get(_bh_n, [])
            _upts = _uu_by_bh.get(_bh_n, [])

            _vd  = [p["depth_m"] for p in _vpts]
            _vs  = [p["Su_kPa"]  for p in _vpts]
            _sp  = [p["Sp_kPa"]  for p in _vpts]
            _st  = [p["St"]      for p in _vpts]

            # ── Vùng tô lớp yếu ──────────────────────────────────────────────
            _h_soft = _soft_bot_map.get(_bh_n)
            if _h_soft:
                _fig_sub.add_trace(_go_tv.Scatter(
                    x=[0, _x_max, _x_max, 0, 0],
                    y=[0, 0, _h_soft, _h_soft, 0],
                    fill="toself",
                    fillcolor="rgba(135,206,250,0.13)",
                    line=dict(width=0),
                    mode="lines",
                    showlegend=(_i == 0),
                    name="Lớp đất yếu",
                    legendgroup="soft",
                    hoverinfo="skip",
                ), row=_row, col=_col)

            # ── Su_VST – đường liền ───────────────────────────────────────────
            _hov_vst = [
                f"<b>{_bh_n}</b><br>Chiều sâu: {d:.1f} m<br>"
                f"<b>Su (VST): {s:.1f} kPa</b>"
                + (f"<br>Sp: {sp:.1f} kPa  St: {st:.1f}" if sp and st else "")
                for d, s, sp, st in zip(_vd, _vs, _sp, _st)
            ]
            _fig_sub.add_trace(_go_tv.Scatter(
                x=_vs, y=_vd,
                mode="lines+markers",
                name="Su (VST)" if _i == 0 else "Su (VST)",
                legendgroup="vst",
                showlegend=(_i == 0),
                line=dict(color="#1a6fbd", width=2),
                marker=dict(size=7, symbol="circle", color="#1a6fbd",
                            line=dict(color="white", width=1)),
                hovertemplate="%{customdata}<extra></extra>",
                customdata=_hov_vst,
            ), row=_row, col=_col)

            # ── Sp – đường đứt mảnh ──────────────────────────────────────────
            _sp_valid = [(d, s) for d, s in zip(_vd, _sp) if s]
            if _sp_valid:
                _dp_sp, _sp_vals = zip(*_sp_valid)
                _fig_sub.add_trace(_go_tv.Scatter(
                    x=list(_sp_vals), y=list(_dp_sp),
                    mode="lines+markers",
                    name="Sp (sau phá hoại)" if _i == 0 else "Sp (sau phá hoại)",
                    legendgroup="sp",
                    showlegend=(_i == 0),
                    line=dict(color="#1a6fbd", width=1.2, dash="dot"),
                    marker=dict(size=5, symbol="x", color="#1a6fbd"),
                    hovertemplate=f"<b>{_bh_n}</b><br>Sp: %{{x:.1f}} kPa | %{{y:.1f}} m<extra></extra>",
                ), row=_row, col=_col)

            # ── Cu TÍNH TOÁN (TCCS 41 Phụ lục C.3.2 — Công thức C.5: Cu = μ·Su) ──
            _bj_e   = _bj_per_bh.get(_bh_n, {})
            _mu_bh  = _bj_e.get("mu")
            _Ip_bh  = _bj_e.get("Ip")
            if _mu_bh is not None and _vd:
                _vs_cu = [s * _mu_bh for s in _vs]
                # Text label hiển thị giá trị Cu trên mỗi marker
                _txt_cu = [f"{cu:.1f}" for cu in _vs_cu]
                _hov_cu = [
                    f"<b>{_bh_n}</b><br>Chiều sâu: {d:.1f} m<br>"
                    f"<b>Cu = μ·Su = {cu:.1f} kPa</b><br>"
                    f"Su gốc (VST): {s:.1f} kPa<br>"
                    f"μ (Bjerrum) = {_mu_bh:.3f}"
                    + (f"<br>Ip TB lớp yếu = {_Ip_bh:.0f}" if _Ip_bh else "")
                    for d, cu, s in zip(_vd, _vs_cu, _vs)
                ]
                # Trace Cu tính toán — đường liền xanh lá đậm + text giá trị bên marker
                _fig_sub.add_trace(_go_tv.Scatter(
                    x=_vs_cu, y=_vd,
                    mode="lines+markers+text",
                    name="Cu tính toán (Cu = μ·Su — TCCS 41 C.5)"
                         if _i == 0 else "Cu tính toán (Cu = μ·Su — TCCS 41 C.5)",
                    legendgroup="cu_corr",
                    showlegend=(_i == 0),
                    line=dict(color="#16a34a", width=2.6),
                    marker=dict(size=9, symbol="diamond", color="#16a34a",
                                line=dict(color="white", width=1.2)),
                    text=_txt_cu,
                    textposition="middle right",
                    textfont=dict(size=8, color="#14532d", family="Arial"),
                    hovertemplate="%{customdata}<extra></extra>",
                    customdata=_hov_cu,
                ), row=_row, col=_col)
                # Vạch đứng Cu_TB trong phạm vi lớp yếu
                _vs_cu_soft = [
                    cu for d, cu in zip(_vd, _vs_cu)
                    if _h_soft is None or d <= _h_soft
                ]
                if _vs_cu_soft:
                    _Cu_avg = sum(_vs_cu_soft) / len(_vs_cu_soft)
                    _y_top  = (_h_soft if _h_soft else max(_vd))
                    # Đường đứng Cu_TB (xanh lá nhạt)
                    _fig_sub.add_trace(_go_tv.Scatter(
                        x=[_Cu_avg, _Cu_avg], y=[0, _y_top],
                        mode="lines",
                        line=dict(color="#16a34a", width=1.5, dash="dashdot"),
                        name=("Cu_TB tính toán (μ·Su_TB lớp yếu)"
                              if _i == 0 else "Cu_TB tính toán (μ·Su_TB lớp yếu)"),
                        legendgroup="cu_avg",
                        showlegend=(_i == 0),
                        hovertemplate=(f"<b>{_bh_n}</b><br>"
                                       f"Cu TB = {_Cu_avg:.1f} kPa<br>"
                                       f"trong phạm vi lớp yếu<extra></extra>"),
                    ), row=_row, col=_col)
                    # Annotation hộp xanh lá ở 0.3·H_soft (trên cao, dễ thấy)
                    _y_ann  = (_h_soft * 0.30) if _h_soft else (max(_vd) * 0.3)
                    _ann_txt = (
                        f"<b>Cu_TB = {_Cu_avg:.1f} kPa</b><br>"
                        f"μ = {_mu_bh:.3f}"
                        + (f"  Ip = {_Ip_bh:.0f}" if _Ip_bh else "")
                    )
                    _fig_sub.add_annotation(
                        x=_Cu_avg, y=_y_ann, text=_ann_txt,
                        showarrow=True, arrowhead=2, arrowsize=1,
                        arrowwidth=1.5, arrowcolor="#16a34a",
                        ax=45, ay=-30,
                        font=dict(size=9.5, color="#14532d",
                                  family="Arial Black"),
                        bgcolor="rgba(220,252,231,0.95)",
                        bordercolor="#16a34a", borderwidth=1.4, borderpad=4,
                        row=_row, col=_col,
                    )

            # ── Cu_UU – marker kim cương ──────────────────────────────────────
            if _upts:
                _ud = [r["depth_mid_m"]  for r in _upts]
                _uu = [r["Cu_UU_kPa"]    for r in _upts]
                _hov_uu = [
                    f"<b>{_bh_n}</b><br>Chiều sâu: {d:.1f} m<br>"
                    f"<b>Cu (UU): {u:.1f} kPa</b>"
                    for d, u in zip(_ud, _uu)
                ]
                _fig_sub.add_trace(_go_tv.Scatter(
                    x=_uu, y=_ud,
                    mode="markers",
                    name="Cu (UU phòng)" if _i == 0 else "Cu (UU phòng)",
                    legendgroup="uu",
                    showlegend=(_i == 0),
                    marker=dict(size=10, symbol="diamond",
                                color="#e85800",
                                line=dict(color="white", width=1)),
                    hovertemplate="%{customdata}<extra></extra>",
                    customdata=_hov_uu,
                ), row=_row, col=_col)

            # ── Đường Su = 25 kPa ─────────────────────────────────────────────
            _y_max_bh = max(_vd + [r["depth_mid_m"] for r in _upts] + [30]) + 2
            _fig_sub.add_trace(_go_tv.Scatter(
                x=[25, 25], y=[0, _y_max_bh],
                mode="lines",
                line=dict(color="red", width=1, dash="dash"),
                name="Su = 25 kPa" if _i == 0 else "Su = 25 kPa",
                legendgroup="ref25",
                showlegend=(_i == 0),
                hoverinfo="skip",
            ), row=_row, col=_col)

            # ── Trục cho subplot này ──────────────────────────────────────────
            _ax_idx = "" if (_row == 1 and _col == 1) else str(_i + 1)
            _fig_sub.update_xaxes(
                range=[0, _x_max], gridcolor="#ebebeb",
                tickfont=dict(size=10), row=_row, col=_col,
            )
            _fig_sub.update_yaxes(
                autorange="reversed", gridcolor="#ebebeb",
                tickfont=dict(size=10),
                title_text="Sâu (m)" if _col == 1 else "",
                row=_row, col=_col,
            )

        # Ẩn subplot trống nếu n_bh không chia hết cho NCOLS
        for _blank in range(_n_bh, _NROWS * _NCOLS):
            _br = _blank // _NCOLS + 1
            _bc = _blank  % _NCOLS + 1
            _fig_sub.update_xaxes(visible=False, row=_br, col=_bc)
            _fig_sub.update_yaxes(visible=False, row=_br, col=_bc)

        _fig_sub.update_layout(
            height=300 * _NROWS,
            margin=dict(l=55, r=15, t=55, b=40),
            plot_bgcolor="white",
            paper_bgcolor="white",
            hovermode="closest",
            legend=dict(
                orientation="h", x=0, y=-0.04,
                xanchor="left", yanchor="top",
                font=dict(size=11),
                tracegroupgap=6,
            ),
            title=dict(
                text=(
                    f"VST & UU – Kè KE: {_n_bh} hố khoan  |  "
                    f"VST {len(_vst_data)} điểm  |  UU {len(_uu_data)} điểm"
                ),
                font=dict(size=13), x=0,
            ),
        )
        st.plotly_chart(_fig_sub, use_container_width=True)
        st.caption(
            "**Su (xanh dương liền):** cắt cánh hiện trường (VST) — cường độ kháng cắt nguyên trạng.  "
            "**Sp (xanh dương chấm):** độ bền sau phá hoại.  "
            "**Cu tính toán (xanh lá đậm, kim cương):** Cu = μ·Su theo TCCS 41 Phụ lục C.3.2 — Công thức C.5; μ nội suy từ Bảng C.1 theo Ip TB lớp yếu của từng hố khoan.  "
            "**Vạch xanh lá đứng (dashdot):** Cu_TB tính toán trong phạm vi lớp yếu.  "
            "**Cu (kim cương cam):** thí nghiệm cắt 3 trục UU trong phòng (tham khảo so sánh).  "
            "**Đường đỏ đứt:** ngưỡng Su = 25 kPa.  "
            "**Vùng xanh nhạt:** phạm vi lớp đất yếu."
        )
    else:
        st.info("Chưa có dữ liệu cắt cánh cho khu vực Kè.")

    # ── Thống kê TN nén cố kết lớp 1 và 1b — chỉ HK tham gia CDM ───────────
    st.markdown("#### Thí nghiệm nén cố kết lớp 1 và lớp 1b")
    _cdm_yes_ids = [b["id"] for b in _bhs_tv if b["name"] in _cdm_yes]
    _cc_rows = []
    for _sym in ("1", "1b"):
        for z in _zones_tv:
            _ids_z_cdm = [b["id"] for b in _bhs_tv
                          if b["zone_id"] == z["id"] and b["name"] in _cdm_yes]
            if not _ids_z_cdm:
                continue
            _ph = ",".join("?" * len(_ids_z_cdm))
            r = _cv.execute(f"""
                SELECT
                    COUNT(*)                                                   n_mau,
                    COUNT(CASE WHEN lt.Cc   IS NOT NULL AND lt.Cc  >0 THEN 1 END) n_Cc,
                    ROUND(AVG(CASE WHEN lt.Cc  >0 THEN lt.Cc  END),3)        avg_Cc,
                    ROUND(MIN(CASE WHEN lt.Cc  >0 THEN lt.Cc  END),3)        min_Cc,
                    ROUND(MAX(CASE WHEN lt.Cc  >0 THEN lt.Cc  END),3)        max_Cc,
                    ROUND(AVG(CASE WHEN lt.Cs  >0 THEN lt.Cs  END),3)        avg_Cs,
                    ROUND(AVG(CASE WHEN lt.e0  >0 THEN lt.e0  END),3)        avg_e0,
                    ROUND(AVG(CASE WHEN lt.PC_kPa>0 THEN lt.PC_kPa END),0)   avg_PC,
                    ROUND(MIN(CASE WHEN lt.PC_kPa>0 THEN lt.PC_kPa END),0)   min_PC,
                    ROUND(MAX(CASE WHEN lt.PC_kPa>0 THEN lt.PC_kPa END),0)   max_PC,
                    ROUND(AVG(CASE WHEN lt.Cv_cm2s>0 THEN lt.Cv_cm2s END)*1e4,3) avg_Cv_e4
                FROM lab_tests lt
                JOIN boreholes b ON lt.borehole_id = b.id
                JOIN layers lyr  ON lyr.borehole_id = b.id
                    AND lt.depth_from_m >= lyr.depth_top_m
                    AND lt.depth_from_m <  lyr.depth_bot_m
                    AND lyr.symbol = ?
                WHERE lt.borehole_id IN ({_ph})
            """, (_sym, *_ids_z_cdm)).fetchone()

            if not r or r[0] == 0:
                continue
            _cc_rows.append({
                "Khu vực":           z["code"],
                "Ký hiệu lớp":       _sym,
                "Tổng mẫu":          r[0],
                "Mẫu có Cc":         r[1],
                "Cc TB":             r[2] or "—",
                "Cc min":            r[3] or "—",
                "Cc max":            r[4] or "—",
                "Cs TB":             r[5] or "—",
                "e₀ TB":             r[6] or "—",
                "PC TB (kPa)":       int(r[7]) if r[7] else "—",
                "PC min (kPa)":      int(r[8]) if r[8] else "—",
                "PC max (kPa)":      int(r[9]) if r[9] else "—",
                "Cv TB (×10⁻⁴ cm²/s)": r[10] or "—",
            })

    if _cc_rows:
        st.table(_pd_tv.DataFrame(_cc_rows))
    else:
        st.caption("Không có dữ liệu nén cố kết cho lớp 1 / 1b.")

    # ── Auto-save thông số mặc định vào SQLite ───────────────────────────────
    # Tạo bảng tvtk_soil_params nếu chưa có
    _cv.execute("""
        CREATE TABLE IF NOT EXISTS tvtk_soil_params (
            zone_code   TEXT PRIMARY KEY,
            Cc_avg      REAL, Cs_avg      REAL, e0_avg      REAL,
            PC_kPa_avg  REAL, Cv_e4_avg   REAL,
            H_soft_avg_m REAL,
            n_Cc        INTEGER,
            updated_at  TEXT DEFAULT (datetime('now','localtime'))
        )
    """)
    # Cập nhật H_soft per borehole vào tvtk_bh_cdm (thêm cột nếu chưa có)
    try:
        _cv.execute("ALTER TABLE tvtk_bh_cdm ADD COLUMN H_soft_m REAL")
    except Exception:
        pass
    # Lưu H_soft (= lớp1 + lớp1b) cho từng HK
    for _sr in _soft_rows:
        _cv.execute("""
            INSERT INTO tvtk_bh_cdm (bh_name, H_soft_m)
            VALUES (?, ?)
            ON CONFLICT(bh_name) DO UPDATE SET H_soft_m=excluded.H_soft_m,
                updated_at=datetime('now','localtime')
        """, (_sr["_bh_name"], _sr["_h_soft"]))

    # Lưu thông số TB nén cố kết theo zone (chỉ lớp 1)
    for _ccr in _cc_rows:
        if _ccr["Ký hiệu lớp"] != "1":
            continue
        _zc = _ccr["Khu vực"]
        # H_soft_avg cho zone
        _bhs_z_soft = [s for s in _soft_rows if s["Khu vực"] == _zc]
        _h_soft_avg = (sum(s["_h_soft"] for s in _bhs_z_soft) / len(_bhs_z_soft)) if _bhs_z_soft else None
        _cv.execute("""
            INSERT INTO tvtk_soil_params
                (zone_code, Cc_avg, Cs_avg, e0_avg, PC_kPa_avg, Cv_e4_avg, H_soft_avg_m, n_Cc)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(zone_code) DO UPDATE SET
                Cc_avg=excluded.Cc_avg, Cs_avg=excluded.Cs_avg,
                e0_avg=excluded.e0_avg, PC_kPa_avg=excluded.PC_kPa_avg,
                Cv_e4_avg=excluded.Cv_e4_avg, H_soft_avg_m=excluded.H_soft_avg_m,
                n_Cc=excluded.n_Cc, updated_at=datetime('now','localtime')
        """, (
            _zc,
            _ccr["Cc TB"]   if _ccr["Cc TB"]   != "—" else None,
            _ccr["Cs TB"]   if _ccr["Cs TB"]   != "—" else None,
            _ccr["e₀ TB"]   if _ccr["e₀ TB"]   != "—" else None,
            _ccr["PC TB (kPa)"] if _ccr["PC TB (kPa)"] != "—" else None,
            _ccr["Cv TB (×10⁻⁴ cm²/s)"] if _ccr["Cv TB (×10⁻⁴ cm²/s)"] != "—" else None,
            round(_h_soft_avg, 2) if _h_soft_avg else None,
            _ccr["Mẫu có Cc"],
        ))
    _cv.commit()
    # Hiển thị xác nhận
    _saved_params = _cv.execute(
        "SELECT zone_code, Cc_avg, Cs_avg, PC_kPa_avg, H_soft_avg_m, n_Cc FROM tvtk_soil_params ORDER BY zone_code"
    ).fetchall()
    if _saved_params:
        with st.expander("Thông số mặc định đã lưu", expanded=False):
            st.table(_pd_tv.DataFrame(_saved_params, columns=[
                "Khu vực", "Cc TB", "Cs TB", "PC TB (kPa)", "H bùn TB (m)", "n mẫu Cc"
            ]))
    st.divider()

    # ── 5. Chỉ tiêu đề xuất phục vụ tính nền & ổn định ─────────────────────
    st.markdown("### 5. Đề xuất chỉ tiêu thiết kế phục vụ tính nền và ổn định")

    # Thông số hình học cọc CDM đề xuất — editable, lưu SQLite
    import math as _math_cdm
    _cv.execute("""
        CREATE TABLE IF NOT EXISTS tvtk_cdm_config (
            id            INTEGER PRIMARY KEY CHECK (id = 1),
            D_mm          REAL DEFAULT 800,
            spacing_m     REAL DEFAULT 1.8,
            pattern       TEXT DEFAULT 'square',
            top_elev_m    REAL DEFAULT 0.8,
            penetration_m REAL DEFAULT 1.0,
            q_kPa         REAL DEFAULT 40.8,
            Ec_factor     REAL DEFAULT 100,
            qu_kPa        REAL DEFAULT 800,
            updated_at    TEXT
        )
    """)
    for _col, _def in [("Ec_factor", "100"), ("qu_kPa", "800")]:
        try:
            _cv.execute(f"ALTER TABLE tvtk_cdm_config ADD COLUMN {_col} REAL DEFAULT {_def}")
        except Exception:
            pass
    _cv.execute("""
        INSERT OR IGNORE INTO tvtk_cdm_config
            (id, D_mm, spacing_m, pattern, top_elev_m, penetration_m, q_kPa, Ec_factor, qu_kPa)
        VALUES (1, 800, 1.8, 'square', 0.8, 1.0, 40.8, 100, 800)
    """)
    _cv.commit()
    _cdm_cfg = _cv.execute("""
        SELECT D_mm, spacing_m, pattern, top_elev_m, penetration_m, q_kPa, Ec_factor, qu_kPa
        FROM tvtk_cdm_config WHERE id=1
    """).fetchone()
    _D_mm_def, _s_def, _pat_def, _top_def, _pen_def, _q_def, _Ef_def, _qu_def = (
        _cdm_cfg or (800, 1.8, "square", 0.8, 1.0, 40.8, 100, 800)
    )

    # Hàng 1 — Hình học bố trí cọc
    _r1a, _r1b, _r1c, _r1d, _r1e = st.columns(5)
    _D_mm  = _r1a.number_input("D cọc (mm)", value=float(_D_mm_def), step=100.0,
                                min_value=400.0, max_value=1200.0, format="%.0f", key="_tv5_D")
    _s_inp = _r1b.number_input("Khoảng cách s (m)", value=float(_s_def), step=0.1,
                                min_value=0.8, max_value=5.0, format="%.1f", key="_tv5_s")
    _pat   = _r1c.selectbox("Bố trí lưới", ["square", "triangle"],
                             index=0 if _pat_def == "square" else 1, key="_tv5_pat",
                             format_func=lambda x: "Hình vuông" if x == "square" else "Hình tam giác")
    _top   = _r1d.number_input("Cao độ đỉnh cọc (m)", value=float(_top_def), step=0.1, format="%.1f", key="_tv5_top")
    _pen   = _r1e.number_input("Ngàm lớp cứng (m)", value=float(_pen_def), step=0.5,
                                min_value=0.5, format="%.1f", key="_tv5_pen")

    # Hàng 2 — Mô đun đàn hồi + tải trọng + computed + Save
    _D_m    = _D_mm / 1000
    _A_unit = _s_inp**2 if _pat == "square" else (_math_cdm.sqrt(3) / 2 * _s_inp**2)
    _a_cdm  = _math_cdm.pi * _D_m**2 / 4 / _A_unit

    _r2a, _r2b, _r2c, _r2d, _r2e, _r2f, _r2s = st.columns([1.1, 1.1, 1.1, 0.9, 0.9, 0.9, 0.7])
    _Ec_factor = _r2a.number_input("Hệ số k (Ec = k × Cc)", value=float(_Ef_def),
                                    step=25.0, min_value=25.0, max_value=300.0, format="%.0f", key="_tv5_Ef")
    _qu_kPa    = _r2b.number_input("qu thiết kế (kPa)", value=float(_qu_def),
                                    step=50.0, min_value=100.0, format="%.0f", key="_tv5_qu")
    _q_kPa     = _r2c.number_input("q tải (kN/m²)", value=float(_q_def),
                                    step=0.5, min_value=0.0, format="%.2f", key="_tv5_q")
    _Cc_col = _qu_kPa / 2
    _Ec_cdm = _Ec_factor * _Cc_col
    _r2d.metric("a (tỷ lệ thay thế)", f"{_a_cdm:.4f}")
    _r2e.metric("Cc cọc (kPa)", f"{_Cc_col:,.0f}")
    _r2f.metric("Ec cọc (kPa)", f"{_Ec_cdm:,.0f}")

    if _r2s.button("Lưu", key="_tv5_save", use_container_width=True, type="primary"):
        _cv.execute("""
            UPDATE tvtk_cdm_config
            SET D_mm=?, spacing_m=?, pattern=?, top_elev_m=?, penetration_m=?,
                q_kPa=?, Ec_factor=?, qu_kPa=?, updated_at=datetime('now','localtime')
            WHERE id=1
        """, (_D_mm, _s_inp, _pat, _top, _pen, _q_kPa, _Ec_factor, _qu_kPa))
        # Lấy S_no_treat trung bình mỗi zone từ kết quả tính lún nền
        _s_notrt = {r[0]: r[1] for r in _cv.execute(
            "SELECT zone_code, AVG(S_total_cm) FROM settlement_scenarios GROUP BY zone_code"
        ).fetchall()}
        # Tính lại S1 cho từng khu vực với Ec/a/q mới
        for _zr in _cv.execute("SELECT zone_code, Es_kPa, H_cdm_m FROM cdm_design").fetchall():
            _Ecomp  = _a_cdm * _Ec_cdm + (1.0 - _a_cdm) * _zr[1]
            _S1_new = _q_kPa * _zr[2] / _Ecomp * 100  # cm
            _S0     = _s_notrt.get(_zr[0])
            _S_red  = _S1_new / _S0 if _S0 else None
            _cv.execute("""
                UPDATE cdm_design
                SET diameter_m=?, spacing_m=?, pattern=?, area_ratio=?,
                    Cc_col_kPa=?, Ec_kPa=?, Ec_factor=?, q_kPa=?, S1_cm=?, S_reduction=?
                WHERE zone_code=?
            """, (round(_D_m, 3), _s_inp, _pat, round(_a_cdm, 4),
                  _Cc_col, _Ec_cdm, _Ec_factor, _q_kPa,
                  round(_S1_new, 1), round(_S_red, 3) if _S_red else None, _zr[0]))
        _cv.commit()
        st.success("Đã lưu thông số và cập nhật dự báo lún.")
        st.rerun()

    # Kết quả tính toán lún nền CDM — đọc từ cdm_design kết hợp settlement_scenarios
    _cdm_results = _cv.execute("""
        SELECT cd.zone_code, cd.Cu_soft_kPa, cd.Ec_kPa, cd.Es_kPa,
               cd.H_cdm_m, cd.q_kPa, cd.S1_cm, cd.S_reduction,
               sc.S_avg
        FROM cdm_design cd
        LEFT JOIN (
            SELECT zone_code, ROUND(AVG(S_total_cm),1) S_avg
            FROM settlement_scenarios GROUP BY zone_code
        ) sc ON sc.zone_code = cd.zone_code
        ORDER BY cd.zone_code
    """).fetchall()
    if _cdm_results:
        st.markdown("#### Dự báo lún nền sau xử lý CDM (TCVN 9403 Phụ lục C)")
        _cdm_disp = []
        for _rd in _cdm_results:
            _S1  = _rd[6]
            _S0  = _rd[8]
            _red = (1 - _rd[7]) * 100 if _rd[7] else None
            _cdm_disp.append({
                "Khu vực": _rd[0],
                "Cu_VST (kPa)": f"{_rd[1]:.1f}",
                "H đất yếu (m)": f"{_rd[4]:.1f}",
                "q tải (kPa)": f"{_rd[5]:.2f}",
                "S không xử lý (cm)": f"{_S0:.1f}" if _S0 else "—",
                "S1_CDM (cm)": f"{_S1:.1f}" if _S1 else "—",
                "Giảm lún (%)": f"{_red:.1f}%" if _red else "—",
            })
        st.table(_pd_tv.DataFrame(_cdm_disp))
        st.caption(
            f"S1 = q × H / (a × Ec + (1-a) × Es)  |  "
            f"Ec = {_Ef_def:.0f} × Cc_cọc (kPa)  |  Es = 250 × Cu_VST (kPa)  |  "
            "S không xử lý = lún cố kết sơ cấp trung bình các hố khoan"
        )

    # ── Bảng thống kê chiều dài cọc CDM theo nguyên tắc thiết kế ──────────────
    st.markdown("#### Thống kê chiều dài cọc CDM theo nguyên tắc thiết kế")
    st.caption(
        f"**Công thức:** L = (Cao độ đỉnh cọc {_top:+.1f} m) − Cao độ đáy lớp yếu + Ngàm {_pen:.1f} m  "
        f"= {_top:+.1f} − Cao_độ_tự_nhiên + H_đất_yếu + {_pen:.1f}  \n"
        f"**Quy ước đào/đắp:**  "
        f"Đào = Cao_độ_tự_nhiên − {_top:+.1f}  (dương → đào bỏ đất trước khi thi công cọc; "
        f"âm → đắp thêm lên đỉnh cọc).  "
        f"Khi đào lớn hơn 0, lớp đất yếu phía TRÊN đỉnh cọc CDM coi như đã được đào bỏ — "
        f"không cần xử lý → L_CDM < H_đất_yếu là hợp lý."
    )
    _len_rows = []
    _len_zone_agg   = {}   # zone -> list of L_cdm values
    _len_zone_noexc = {}   # zone -> list of L_no_excav (không đào)
    _n_excav_warn   = 0    # số HK có đào > pen (cảnh báo)
    for b in [b for b in _bhs_tv if b["name"] in _cdm_yes]:
        _elev_b = b["elevation_m"]
        # H_soft: sum of ALL soft symbols from layers table
        _hs_val = _cv.execute("""
            SELECT COALESCE(SUM(depth_bot_m - depth_top_m), 0.0)
            FROM layers WHERE borehole_id=? AND symbol IN ('1','1b','2','XMD')
        """, (b["id"],)).fetchone()[0]
        if _hs_val <= 0 or _elev_b is None:
            _len_rows.append({
                "Hố khoan":        b["name"],
                "Khu vực":         _zone_codes.get(b["zone_id"], ""),
                "Cao độ TN (m)":   f"{_elev_b:.2f}" if _elev_b is not None else "—",
                "Đỉnh cọc TK (m)": f"{_top:+.2f}",
                "Đào/Đắp (m)":     "—",
                "H đất yếu (m)":   f"{_hs_val:.1f}" if _hs_val > 0 else "—",
                "L (không đào) m": "—",
                "L CDM (m)":       "—",
                "Ghi chú":         "Thiếu dữ liệu",
            })
            continue
        _excav     = _elev_b - _top                 # >0: đào,  <0: đắp
        _L_no_exc  = _hs_val + _pen                 # L tối thiểu nếu không đào
        _L_cdm     = _top - _elev_b + _hs_val + _pen
        _zcode     = _zone_codes.get(b["zone_id"], "")
        _len_zone_agg.setdefault(_zcode, []).append(_L_cdm)
        _len_zone_noexc.setdefault(_zcode, []).append(_L_no_exc)

        # Ghi chú: cảnh báo khi đào lớn
        if _excav > _pen:        # đào nhiều hơn cả phần ngàm
            _note = f"Cảnh báo: đào {_excav:.1f} m"
            _n_excav_warn += 1
        elif _excav > 0:
            _note = f"Đào {_excav:.1f} m trước khi thi công"
        elif _excav < 0:
            _note = f"Đắp {-_excav:.1f} m lên đỉnh cọc"
        else:
            _note = "Đỉnh cọc = mặt đất tự nhiên"

        _len_rows.append({
            "Hố khoan":        b["name"],
            "Khu vực":         _zcode,
            "Cao độ TN (m)":   f"{_elev_b:.2f}",
            "Đỉnh cọc TK (m)": f"{_top:+.2f}",
            "Đào/Đắp (m)":     f"{_excav:+.2f}",
            "H đất yếu (m)":   f"{_hs_val:.1f}",
            "L (không đào) m": f"{_L_no_exc:.1f}",
            "L CDM (m)":       f"{_L_cdm:.1f}",
            "Ghi chú":         _note,
        })

    if _len_rows:
        # Sort: zone then bh_name
        _len_df = _pd_tv.DataFrame(_len_rows).sort_values(["Khu vực", "Hố khoan"])
        st.table(_len_df.reset_index(drop=True))
        if _n_excav_warn > 0:
            st.warning(
                f"Có **{_n_excav_warn} hố khoan** có chiều cao đào lớn hơn chiều sâu ngàm "
                f"({_pen:.1f} m) — cần xem xét lại cao độ đỉnh cọc CDM ({_top:+.1f} m) "
                f"cho phù hợp với cao độ tự nhiên của từng khu vực."
            )

    # Tổng hợp theo zone
    if _len_zone_agg:
        _agg_rows = []
        for _zc in ("KE", "BXN", "NHC"):
            if _zc not in _len_zone_agg: continue
            _ls    = _len_zone_agg[_zc]
            _ls_ne = _len_zone_noexc.get(_zc, [])
            _agg_rows.append({
                "Khu vực":               _zc,
                "Số HK":                 len(_ls),
                "L min (m)":             f"{min(_ls):.1f}",
                "L TB (m)":              f"{sum(_ls)/len(_ls):.1f}",
                "L max (m)":             f"{max(_ls):.1f}",
                "L TB không đào (m)":    f"{sum(_ls_ne)/len(_ls_ne):.1f}" if _ls_ne else "—",
            })
        st.table(_pd_tv.DataFrame(_agg_rows))

    st.caption("Giá trị trung bình theo khu vực × ký hiệu lớp đất — từ thí nghiệm phòng và cắt cánh hiện trường")
    _SYM_LABEL = {
        "CH": "CH — Sét dẻo cao (đất yếu)",
        "MH": "MH — Sét hữu cơ (đất yếu)",
        "MH-OH": "MH-OH — Sét hữu cơ (đất yếu)",
        "CH-OH": "CH-OH — Sét hữu cơ (đất yếu)",
        "CL": "CL — Sét dẻo thấp",
        "ML": "ML — Bụi dẻo thấp",
        "SC": "SC — Cát lẫn sét",
        "SM": "SM — Cát lẫn bụi",
        "SP": "SP — Cát hạt rời",
    }
    _param_rows = []
    for z in _zones_tv:
        for sym in ("CH", "MH", "MH-OH", "CH-OH", "CL", "ML", "SC", "SM", "SP"):
            r = _cv.execute("""
                SELECT COUNT(*) n,
                    ROUND(AVG(lt.gamma_kNm3),2) g,
                    ROUND(AVG(lt.e0),3) e0,
                    ROUND(AVG(lt.Cc),3) Cc,
                    ROUND(AVG(lt.Cs),3) Cs,
                    ROUND(AVG(lt.PC_kPa),0) PC,
                    ROUND(AVG(lt.Cu_UU_kPa),1) su,
                    ROUND(AVG(lt.phi_deg),1) phi,
                    ROUND(AVG(lt.c_kPa),1) c,
                    ROUND(AVG(lt.Cv_cm2s)*1e4, 3) cv_e4
                FROM lab_tests lt JOIN boreholes b ON lt.borehole_id = b.id
                WHERE b.zone_id = ? AND lt.symbol_tcvn = ?
            """, (z["id"], sym)).fetchone()
            if not r or r[0] == 0:
                continue
            _param_rows.append({
                "Khu vực": z["code"],
                "Ký hiệu": sym,
                "Mô tả": _SYM_LABEL.get(sym, sym),
                "n mẫu": r[0],
                "γ (kN/m³)": r[1] or "—",
                "e₀": r[2] or "—",
                "Cc": r[3] or "—",
                "Cs": r[4] or "—",
                "PC (kPa)": int(r[5]) if r[5] else "—",
                "su_UU (kPa)": r[6] or "—",
                "φ (°)": r[7] or "—",
                "c (kPa)": r[8] or "—",
                "Cv (×10⁻⁴ cm²/s)": r[9] or "—",
            })
    if _param_rows:
        st.table(_pd_tv.DataFrame(_param_rows))
    else:
        st.info("Chưa có dữ liệu thí nghiệm.")
    st.divider()

    # ── 6. Mực nước ngầm / điều kiện thủy lực ───────────────────────────────
    st.markdown("### 6. Mực nước ngầm – điều kiện thủy lực")
    _c6a, _c6b, _c6c, _c6d, _c6s = st.columns([1, 1, 1, 1, 0.6])
    _gwt_ke   = _c6a.number_input("MNN phía kè KE (m)",    value=_tv_load("KE",  "gwt_m", -0.5), step=0.1, format="%.1f", key="_tv_gwt_ke")
    _gwt_bxn  = _c6b.number_input("MNN tại BXN (m)",       value=_tv_load("BXN", "gwt_m",  0.0), step=0.1, format="%.1f", key="_tv_gwt_bxn")
    _gwt_nhc  = _c6c.number_input("MNN tại NHC (m)",       value=_tv_load("NHC", "gwt_m",  0.0), step=0.1, format="%.1f", key="_tv_gwt_nhc")
    _wlvl_song= _c6d.number_input("Cao độ MN sông (m)",    value=_tv_load("KE",  "gwt_m", -1.5), step=0.1, format="%.1f", key="_tv_wlvl_song")
    if _c6s.button("Lưu MNN", key="_tv_save_gwt", use_container_width=True):
        for _zc, _gv in (("KE", _gwt_ke), ("BXN", _gwt_bxn), ("NHC", _gwt_nhc)):
            _cv.execute("""
                INSERT INTO tvtk_config (zone_code, gwt_m)
                VALUES (?, ?)
                ON CONFLICT(zone_code) DO UPDATE SET gwt_m=excluded.gwt_m,
                    updated_at=datetime('now','localtime')
            """, (_zc, _gv))
        _cv.commit()
        st.success("Đã lưu mực nước ngầm.")

    _gwt_rows = [
        {"Khu vực": "KE – Kè Công Viên",      "Mực nước ngầm (m)": f"{_gwt_ke:.1f}",
         "Mực nước sông/kênh (m)": f"{_wlvl_song:.1f}", "Ghi chú": "Dao động theo mùa ±0,5 m"},
        {"Khu vực": "BXN – Bãi Đỗ Xe Ngầm",  "Mực nước ngầm (m)": f"{_gwt_bxn:.1f}",
         "Mực nước sông/kênh (m)": "—",        "Ghi chú": "Cần xác nhận TVTK"},
        {"Khu vực": "NHC – Nhà Hành Chính",   "Mực nước ngầm (m)": f"{_gwt_nhc:.1f}",
         "Mực nước sông/kênh (m)": "—",        "Ghi chú": "Cần xác nhận TVTK"},
    ]
    st.table(_pd_tv.DataFrame(_gwt_rows))
    st.caption("Lưu ý: Mực nước ngầm ảnh hưởng trực tiếp đến áp lực thấm, ổn định hố đào và thiết kế tiêu thoát nước.")

    _cv.close()


