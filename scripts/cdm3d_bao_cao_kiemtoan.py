"""Sinh bao cao Word 'Kiem toan coc CDM — Mo hinh 3D Mohr-Coulomb (CalculiX)'.

Tong hop toan bo chuoi khao sat trong phien lam viec: xac dinh ap luc vua gay
Ux=80cm, giam Es dat yeu (P co dinh), Ecmd=Es, kiem toan noi luc M/Q/N va
kiem toan coc THAT (Ec=40000kPa) tai 40->80cm.

Chay: python scripts/cdm3d_bao_cao_kiemtoan.py
Output: BaoCao_KiemToan_CocCDM_MohrCoulomb_<ngay>.docx (project root)
"""
from __future__ import annotations

import datetime as _dt
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(r"g:\My Drive\AI-SUC TAI COC THEO DAT NEN")
IMG = ROOT / "images"
OUT = ROOT / f"BaoCao_KiemToan_CocCDM_MohrCoulomb_{_dt.date.today():%Y%m%d}.docx"


def _fmt(run, size=11, bold=False, italic=False, color=None, name="Times New Roman"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = name
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    _fmt(p.add_run(text), size=16, bold=True, color=(0x1B, 0x3A, 0x5C))
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    _fmt(p.add_run(text), size=13, bold=True, color=(0x2E, 0x5C, 0x8A))
    return p


def body(doc, text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _fmt(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        _fmt(p.add_run(bold_lead), bold=True)
        _fmt(p.add_run(text))
    else:
        _fmt(p.add_run(text))
    return p


def warn_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    _fmt(p.add_run("CẢNH BÁO / GIẢ ĐỊNH: "), bold=True, color=(0xB7, 0x1C, 0x1C))
    _fmt(p.add_run(text), italic=True, color=(0xB7, 0x1C, 0x1C))
    return p


def figure(doc, img_name, caption, width_cm=15.5):
    path = IMG / img_name
    if not path.exists():
        body(doc, f"[Thiếu ảnh: {img_name}]", color=(0xB0, 0x00, 0x00))
        return
    doc.add_picture(str(path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _fmt(p.add_run(caption), italic=True, size=10, color=(0x44, 0x44, 0x44))
    p.paragraph_format.space_after = Pt(14)


def table_simple(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers):
        _fmt(t.rows[0].cells[i].paragraphs[0].add_run(htext), bold=True, size=10.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            _fmt(cells[i].paragraphs[0].add_run(str(val)), size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.0)

# ============================== BÌA ==============================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
_fmt(p.add_run("BÁO CÁO KIỂM TOÁN CỌC CDM"), size=22, bold=True, color=(0x1B, 0x3A, 0x5C))
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
_fmt(p2.add_run("Mô hình 3D Cọc–Đất Mohr-Coulomb (Gmsh + CalculiX)"), size=15, italic=True, color=(0x55, 0x55, 0x55))
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
_fmt(p3.add_run(f"Khu vực: KE (Kè Công viên TTHC)  —  Giai đoạn tải: GD5  —  Ngày lập: {_dt.date.today():%d/%m/%Y}"), size=11, color=(0x66, 0x66, 0x66))
doc.add_paragraph()

# ============================== 1. MỤC ĐÍCH ==============================
h1(doc, "1. Mục đích và phạm vi")
body(doc, "Báo cáo này tổng hợp chuỗi khảo sát bằng mô hình phần tử hữu hạn 3D "
          "(gmsh dựng lưới + CalculiX giải, vật liệu lớp đất yếu theo mô hình Mohr-Coulomb) "
          "nhằm tìm cơ chế vật lý hợp lý giải thích chênh lệch giữa chuyển vị ngang cọc CDM "
          "quan trắc thực địa (~80cm) và kết quả mô hình đàn hồi tuyến tính đơn giản ban đầu "
          "(~5-9mm) tại giai đoạn tải GD5, đồng thời kiểm toán khả năng chịu lực (nén, kéo do uốn) "
          "của bản thân cọc CDM dưới các kịch bản chuyển vị giả định.")
body(doc, "Toàn bộ tính toán dùng công cụ mã nguồn mở (gmsh, CalculiX, Python) theo yêu cầu "
          "dự án — không sử dụng phần mềm thương mại.")
bullet(doc, "Khu vực KE, nhóm 9 cọc CDM lưới vuông 3×3, D=800mm, khoảng cách 1,8m (minh hoạ tương tác, không phải toàn bộ lưới cọc thực tế).")
bullet(doc, "3 lớp: đất đắp / đất yếu (chứa cọc CDM) / lớp cứng dưới mũi cọc.")
bullet(doc, "Giai đoạn tải khảo sát: GD5 — phương án tải chỉ đắp nửa miền +X (load_footprint='half_pos').")

# ============================== 2. PHƯƠNG PHÁP ==============================
h1(doc, "2. Mô hình và phương pháp")

h2(doc, "2.1. Vật liệu — Mô hình Mohr-Coulomb cho đất yếu")
body(doc, "Lớp đất yếu được gán vật liệu *MOHR COULOMB (φ=ψ=0° → thu về tiêu chuẩn Tresca, "
          "khớp quy ước dự án cho đất sét không thoát nước c=Su, φ=0). Cường độ kháng cắt Su "
          "trong các kịch bản kiểm toán được lấy RẤT CAO (Su=5000kPa) để không chi phối "
          "(non-binding) — mục đích là cô lập ảnh hưởng của độ cứng Es/Ec và tải, không phải "
          "khảo sát phá hoại dẻo của đất trong báo cáo này (phá hoại dẻo với Su thật đã khảo sát "
          "riêng, xem hồ sơ phiên làm việc — không lặp lại ở đây).")

h2(doc, "2.2. Cơ chế tải \"áp lực vữa\" trên 1 hàng cọc")
body(doc, "Giả định áp lực vữa (grout pressure) P tác dụng như một lực ngang tương đương "
          "(P×D×H_cọc, phân bố đều lên các nút) trên MỘT HÀNG CỌC (3 trụ cùng toạ độ x, khác y) "
          "trong suốt chiều dài cọc, cộng dồn với tải GD5 thật. Hàng mục tiêu (x=0, ngay ranh giới "
          "tải GD5) được chọn vì thể hiện chuyển vị tự nhiên lớn nhất khi chỉ có tải GD5 thật.")
warn_box(doc, "Đây là mô hình hoá ĐƠN GIẢN HOÁ cho lực ngang giả định — không phải cơ chế bơm vữa "
              "thực tế (áp lực vữa thật tác dụng dạng bơm ép/nứt thuỷ lực cục bộ, phức tạp hơn nhiều). "
              "Dùng để khảo sát độ nhạy, KHÔNG dùng trực tiếp cho thiết kế thi công.")

h2(doc, "2.3. Trích nội lực N, M, Q từ trường ứng suất khối 3D")
body(doc, "Vì cọc được mô hình dạng khối liên tục (phần tử tứ diện C3D4), không phải phần tử dầm, "
          "nội lực dầm tương đương được ước lượng bằng hồi quy tuyến tính ứng suất dọc trục qua từng "
          "mặt cắt ngang cọc trung tâm:")
body(doc, "σ_zz(x, y) ≈ a + b·x + c·y  ⟹  N(z) = a·A,  M(z) = b·I,  Q(z) = -dM/dz", italic=True)
body(doc, "trong đó A, I là diện tích và mô men quán tính tiết diện tròn cọc (D=0,8m).")

h2(doc, "2.4. Tiêu chí kiểm toán")
table_simple(doc,
    ["Chỉ tiêu", "Công thức / giá trị", "Nguồn"],
    [
        ["Nén giới hạn qu", "800 kPa", "SQLite tvtk_cdm_config (KE), Ec_factor=100, Ec=k·Cc"],
        ["Kéo cho phép σ_ba", "σ_ba = 0,25·qu/F_sem = 166,7 kPa  (F_sem=1,2)", "R14 KIỂM TOÁN LỚP BTXM — công thức cho VẬT LIỆU XI MĂNG ĐẤT (đệm cát-XM), đã verify khớp 100% hồ sơ Excel dự án (66-cushion-params-tthc.md)"],
    ])
warn_box(doc, "Công thức σ_ba=0,25·qu/F_sem được ưu tiên hơn công thức ACI/TCVN 11823 fr=0,63√qu[MPa] "
              "(cho BÊ TÔNG, fr≈563kPa) vì cùng bản chất vật liệu xi măng đất với CDM — công thức bê tông "
              "ngoại suy xuống qu thấp (0,8MPa) cho tỷ lệ fr/qu≈70% (cao bất thường so với thông lệ bê tông "
              "~10-15%), không đáng tin cậy. Cả hai công thức đều CHƯA có thí nghiệm kéo/uốn trực tiếp trên "
              "mẫu CDM dự án để đối chứng — kết quả kiểm toán trong báo cáo cần kỹ sư xác nhận thêm bằng "
              "thí nghiệm thực tế trước khi dùng cho thiết kế chính thức.")

# ============================== 3. KẾT QUẢ ==============================
h1(doc, "3. Kết quả theo từng kịch bản")

h2(doc, "3.1. Xác định áp lực vữa gây Ux=80cm (cọc thật, Ec=40000kPa)")
body(doc, "Với cột thật Ec=40000kPa, vật liệu đàn hồi tuyến tính hoá hoàn toàn (quan hệ P–Ux là "
          "đường thẳng chính xác tuyệt đối, R²=1,000), áp lực vữa cần thiết trên hàng x=0 để đạt "
          "Ux=80cm tại đỉnh cọc là P≈5921kPa (~5,9 MPa).")
warn_box(doc, "P≈5,9 MPa rất cao so với áp lực vữa thi công thực tế thông thường (~0,2–2 MPa) — cơ chế "
              "\"áp lực vữa đơn thuần trên 1 hàng cọc\" một mình khó giải thích trọn vẹn 80cm ở mức tải "
              "vữa khả thi thực tế.")
figure(doc, "cdm3d_KE_grout_row_P_vs_Ux.png", "Hình 3.1a — Quan hệ áp lực vữa P và chuyển vị ngang Ux tại đỉnh hàng cọc x=0.")
figure(doc, "cdm3d_KE_grout_row_matcat_XZ_Ux_80cm.png", "Hình 3.1b — Mặt cắt 2D X-Z tại tim cọc, trường Ux, tại P hiệu chỉnh cho Ux=80cm.")

h2(doc, "3.2. Giảm Es đất yếu, giữ áp lực vữa cố định P=2MPa")
body(doc, "Giữ nguyên độ cứng cọc thật (Ecmd=40000kPa cố định), giảm dần Es lớp đất yếu để tìm Es "
          "cho Ux=40cm. Kết quả cho thấy hiện tượng BÃO HOÀ: Ux đạt trần ~37,3cm quanh Es≈550kPa rồi "
          "giảm nhẹ, và tại Es=300kPa nghiệm phân kỳ (vật liệu Mohr-Coulomb bắt đầu chạm ngưỡng chảy "
          "cục bộ gần điểm đặt lực dù Su=5000kPa rất cao). Ux=40cm KHÔNG đạt được qua cơ chế này — "
          "cọc cứng \"gánh\" tải thay đất, giới hạn hiệu ứng làm mềm đất đơn thuần.")
figure(doc, "cdm3d_KE_grout2mpa_Es_vs_Ux.png", "Hình 3.2 — Quan hệ Es (đất yếu) và Ux khi giữ Ec cọc, P vữa=2MPa cố định — thể hiện hiện tượng bão hoà và phân kỳ.")

h2(doc, "3.3. Kịch bản Ecmd=Es (cọc mềm hoá theo đất) — loại trừ hiệu ứng \"cọc cứng gánh tải\"")
body(doc, "Để kiểm tra xem trần bão hoà ở mục 3.2 có phải do cọc cứng cố định gây ra hay không, "
          "kịch bản này cho Ecmd=Es tại MỖI mức khảo sát (cọc và đất luôn cùng độ cứng), P vữa=1,5MPa "
          "cố định, quan sát tại 3 độ sâu (đỉnh cọc, cách đỉnh 5m, 10m).")
table_simple(doc,
    ["Es=Ecmd (kPa)", "Đỉnh cọc (cm)", "-5m (cm)", "-10m (cm)", "Đáng tin?"],
    [
        ["3450 (thật)", "23,4", "28,4", "35,3", "Có"],
        ["1500", "28,8", "64,6", "81,9", "Ranh giới"],
        ["800", "29,8", "120,1", "154,2", "Không"],
        ["400", "26,3", "237,3", "308,9", "Không"],
        ["200", "16,6", "467,0", "616,2", "Không"],
        ["100", "0,6", "914,3", "1224,7", "Không"],
    ])
warn_box(doc, "Khi Es=Ecmd giảm dưới ~1500kPa, chuyển vị tại độ sâu -5m/-10m tăng phi thực tế (đến hơn "
              "12m ở Es=100kPa) trong khi đỉnh cọc lại giảm — đã kiểm chứng không phải kỳ dị số cục bộ "
              "(chuyển vị lớn tương tự xuất hiện cả ở nút đất xa tải và hàng cọc khác cùng độ sâu) mà là "
              "hiệu ứng lan rộng khi cọc và đất đồng thời quá mềm, kết hợp cách quy tải nút đều dọc thân "
              "cọc — KHÔNG đáng tin để dùng kỹ thuật ở vùng Es<1500kPa.")
figure(doc, "cdm3d_KE_Ecmd_eq_Es_profile_full_pile.png", "Hình 3.3 — Profile Ux dọc suốt chiều dài cọc tại 6 mức Es=Ecmd (thang thường + log).")

h2(doc, "3.4. Kiểm toán nội lực M/Q/N — so sánh Ecmd=Es và cọc thật")
body(doc, "Với kịch bản Ecmd=Es (Es=3450 và 1500kPa, vùng còn đáng tin), nội lực trong cọc RẤT NHỎ "
          "(M~4-9 kNm) — cọc không \"hút\" thêm nội lực vì cùng độ cứng với đất xung quanh. Kiểm toán "
          "cho kết quả ĐẠT cả nén lẫn kéo, với tỷ lệ sử dụng kéo 91-99% (rất sát ngưỡng ở Es=3450).")
figure(doc, "cdm3d_KE_MQN_kiemtoan_coc_CDM.png", "Hình 3.4 — Nội lực N(z), M(z), Q(z) và kiểm toán ứng suất mép tiết diện, kịch bản Ecmd=Es.")
body(doc, "Đây LÀ kịch bản giả định (cọc mềm bằng đất) dùng để cô lập hiệu ứng — KHÔNG phản ánh cấu "
          "hình cọc CDM thật (Ec thiết kế = 40000kPa, cứng hơn đất yếu ~11,6 lần).")

h2(doc, "3.5. Kiểm toán cọc THẬT (Ec=40000kPa) tại Ux=40 → 80cm — kết quả quyết định")
body(doc, "Dùng đúng độ cứng cọc thiết kế thật (Ec=40000kPa), điều chỉnh áp lực vữa P để đạt lần lượt "
          "Ux đỉnh = 40/50/60/70/80cm, trích nội lực và kiểm toán tại mỗi mức:")
table_simple(doc,
    ["Ux đỉnh (cm)", "M_max (kNm)", "σ kéo max (kPa)", "Tỷ lệ kéo (%)", "Kết luận"],
    [
        ["40", "66,5", "1293,4", "776,1", "KHÔNG ĐẠT (nứt)"],
        ["50", "83,6", "1641,5", "984,9", "KHÔNG ĐẠT (nứt)"],
        ["60", "100,7", "1989,7", "1193,8", "KHÔNG ĐẠT (nứt)"],
        ["70", "117,8", "2337,8", "1402,7", "KHÔNG ĐẠT (nứt)"],
        ["80", "134,9", "2685,9", "1611,5", "KHÔNG ĐẠT (nứt)"],
    ])
body(doc, "Nội suy tuyến tính từ 5 điểm cho thấy ngưỡng phá hoại thực tế còn thấp hơn nhiều:", bold=False)
bullet(doc, "vượt 100% (nứt kéo) tại Ux đỉnh ≈ 7,6 cm.", bold_lead="KÉO ")
bullet(doc, "vượt 100% (vỡ nén) tại Ux đỉnh ≈ 23,3 cm.", bold_lead="NÉN ")
body(doc, "KÉO chi phối (xảy ra trước) — với cấu hình cọc CDM thật, cọc được dự đoán nứt do uốn ở mức "
          "chuyển vị nhỏ hơn nhiều so với 40cm, chứ chưa nói đến 80cm quan trắc thực địa.", bold=True)
figure(doc, "cdm3d_KE_kiemtoan_40_80cm.png", "Hình 3.5 — Mô men M(z), kiểm toán kéo (thang log) và bảng tổng hợp tại 5 mức chuyển vị, cọc thật Ec=40000kPa.")

# ============================== 4. KẾT LUẬN ==============================
h1(doc, "4. Kết luận và khuyến nghị")
body(doc, "Kết quả quan trọng nhất của chuỗi khảo sát: với độ cứng cọc CDM THẬT (Ec=40000kPa), mô hình "
          "3D Mohr-Coulomb dự đoán cọc sẽ NỨT DO KÉO (uốn) tại mức chuyển vị ngang chỉ khoảng 7,6cm — "
          "rất nhỏ so với 80cm quan trắc thực địa. Điều này có ý nghĩa kỹ thuật quan trọng và mở ra một số "
          "khả năng cần làm rõ thêm:")
bullet(doc, "vật liệu CDM thực tế có thể có ứng xử dẻo hơn nhiều so với giả thiết đàn hồi-giòn tuyến tính "
             "trong mô hình này (Mohr-Coulomb ở đây chỉ áp dụng cho ĐẤT, cọc vẫn thuần đàn hồi) — nếu bản "
             "thân vật liệu cọc có khả năng biến dạng dẻo/dư ứng suất, ngưỡng nứt thực tế có thể cao hơn.",
        bold_lead="(a) ")
bullet(doc, "cọc có thể ĐÃ nứt cục bộ ở chuyển vị nhỏ nhưng đất xung quanh vẫn \"ôm giữ\" không sập/không "
             "quan sát được hư hỏng rõ rệt từ bên ngoài (nứt bên trong, chưa lộ ra ngoài mặt).",
        bold_lead="(b) ")
bullet(doc, "cơ chế chuyển vị 80cm quan trắc thực địa có thể KHÔNG đến từ uốn cọc đơn thuần dưới tải tĩnh, "
             "mà từ trượt/xoay khối đất-cọc lớn (block sliding/rotation), cố kết lệch tâm theo giai đoạn thi "
             "công, hoặc tổ hợp nhiều cơ chế — không phải nội lực uốn cục bộ trong thân 1 cọc.",
        bold_lead="(c) ")
body(doc, "Khuyến nghị bước tiếp theo:")
bullet(doc, "bổ sung thí nghiệm kéo/uốn trực tiếp trên mẫu CDM dự án để xác định σ_ba thực tế, thay cho "
             "công thức ngoại suy (R14 đệm cát-XM) dùng trong báo cáo này.")
bullet(doc, "khảo sát ứng xử dẻo/dư ứng suất của vật liệu cọc CDM (không chỉ đất) nếu có số liệu phù hợp.")
bullet(doc, "đối chiếu cơ chế trượt khối/xoay khối với dữ liệu quan trắc thực địa (inclinometer theo độ sâu) "
             "để phân biệt biến dạng uốn cục bộ và chuyển vị khối.")

# ============================== 5. PHỤ LỤC ==============================
h1(doc, "5. Phụ lục — Danh mục giả định cần kỹ sư xác nhận")
bullet(doc, "γ đất đắp=19,0 kN/m³, E đất đắp=8000kPa: GIẢ ĐỊNH, không có trong hồ sơ khảo sát dự án.")
bullet(doc, "Hệ số Poisson (đất yếu 0,35, đất đắp 0,30, cọc CDM 0,25): GIẢ ĐỊNH giá trị địa kỹ thuật thông thường.")
bullet(doc, "γ cọc CDM=18,0 kN/m³: GIẢ ĐỊNH, chưa có thí nghiệm dung trọng vữa xi măng trộn thực tế.")
bullet(doc, "Lớp cứng dưới mũi cọc: E=8× Es đất yếu: GIẢ ĐỊNH, chưa có khoan sâu qua mũi cọc.")
bullet(doc, "σ_ba=0,25·qu/F_sem: công thức của lớp đệm cát-xi măng (mục đích khác), ngoại suy áp dụng cho cọc CDM.")
bullet(doc, "Cơ chế \"áp lực vữa trên 1 hàng cọc\" là mô hình hoá đơn giản hoá cho mục đích khảo sát độ nhạy, "
             "không phải cơ chế bơm vữa/tải trọng thi công thực tế đã được xác nhận hiện trường.")
bullet(doc, "Mọi kết quả trong báo cáo dựa trên mô hình đàn hồi tuyến tính (cọc) + Mohr-Coulomb (đất, "
             "Su=5000kPa không chi phối) — CHƯA xét vật liệu cọc CDM có thể nứt/dẻo, CHƯA xét tương tác "
             "trượt cọc-đất (bonded), CHƯA có bước khởi tạo ứng suất geostatic K0.")

doc.save(str(OUT))
print(f"Da luu: {OUT}")
