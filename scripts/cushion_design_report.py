"""
cushion_design_report.py — Báo cáo Word docx kiểm toán đệm cát-XM (ALiCC PWRI).

Cấu trúc:
  Trang bìa
  1. Tổng quan + cơ sở lý thuyết ALiCC PWRI
  2. Thông số đầu vào dự án TTHC (verify từ hồ sơ E.2)
  3. Công thức kiểm toán
  4. Tính toán hiện trạng (Hse=0.40, q_uckse=600) — verdict
  5. Ma trận tối ưu hoá Hse × q_uckse + biểu đồ
  6. Phân tích Pareto chi phí - hệ số an toàn
  7. So sánh 3 phương án đề xuất
  8. Kết luận + khuyến nghị

Header: logo + tên công ty (phải)
Footer: nhân sự thực hiện (trái) + Trang X / Y (phải)
"""
from __future__ import annotations
import io
import math
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches

import matplotlib.pyplot as plt

from cdm_cushion_params import check_alicc, find_Hse_min, Q_UCKSE_KPA, FS_PUNCHING, THETA_DEG
from cushion_design_charts import (
    chart_heatmap_ratio, chart_ratio_vs_quckse_fixed_Hse,
    chart_ratio_vs_Hse_fixed_quckse, chart_compare_options,
    chart_alicc_schematic, chart_pareto_cost_ratio,
)
from core.report_style import style_tbl_for_docx, BODY_FONT_PT, HEADER_FONT_PT
from core.cost import cushion_cost_per_m2, cushion_cost_delta_pct

# ════════ HEADER/FOOTER HELPERS ════════
def _field_run(para, field: str) -> None:
    run = para.add_run()
    fc_b = OxmlElement("w:fldChar"); fc_b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.text = field
    instr.set(qn("xml:space"), "preserve")
    fc_e = OxmlElement("w:fldChar"); fc_e.set(qn("w:fldCharType"), "end")
    run._r.append(fc_b); run._r.append(instr); run._r.append(fc_e)


def _set_no_tbl_border(tbl):
    """Xoá border cho bảng layout (header/footer)."""
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none"); e.set(qn("w:sz"), "0")
        tblBorders.append(e)
    tbl._tbl.tblPr.append(tblBorders)


def _setup_header_footer(doc: Document, co_name: str = "", co_staff: str = "",
                          logo_bytes: Optional[bytes] = None):
    sec = doc.sections[0]
    pw = sec.page_width - sec.left_margin - sec.right_margin

    # Header — logo trái + tên công ty phải
    ht = sec.header.add_table(1, 2, width=pw)
    _set_no_tbl_border(ht)
    ht.cell(0, 0).width = Cm(4); ht.cell(0, 1).width = pw - Cm(4)
    if logo_bytes:
        try:
            r = ht.cell(0, 0).paragraphs[0].add_run()
            r.add_picture(io.BytesIO(logo_bytes), height=Cm(1.2))
        except Exception:
            pass
    p = ht.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(co_name); r.font.size = Pt(10); r.bold = True

    # Footer — staff trái + Trang X / Y phải
    ft = sec.footer.add_table(1, 2, width=pw)
    _set_no_tbl_border(ft)
    ft.cell(0, 0).width = pw - Cm(4); ft.cell(0, 1).width = Cm(4)
    p1 = ft.cell(0, 0).paragraphs[0]
    r = p1.add_run(co_staff); r.font.size = Pt(9)
    p2 = ft.cell(0, 1).paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p2.add_run("Trang "); r.font.size = Pt(9)
    _field_run(p2, "PAGE")
    r = p2.add_run(" / "); r.font.size = Pt(9)
    _field_run(p2, "NUMPAGES")


# ════════ STYLE HELPERS ════════
def _apply_default_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(BODY_FONT_PT)


def _h1(doc, text: str, page_break: bool = False):
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(16); r.bold = True
    r.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)


def _h2(doc, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0x37, 0x47, 0x57)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)


def _para(doc, text: str, bold: bool = False, italic: bool = False, size_pt: int = None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if size_pt:
        r.font.size = Pt(size_pt)
    r.bold = bold; r.italic = italic
    return p


def _add_image_from_fig(doc, fig, width_cm: float = 16.0, caption: str = ""):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=160, bbox_inches="tight")
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(buf, width=Cm(width_cm))
    plt.close(fig)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        r.italic = True; r.font.size = Pt(10)
        cp.paragraph_format.space_after = Pt(8)


def _verdict_cell(cell, ok: bool):
    """Tô màu cell theo verdict."""
    color = "C8E6C9" if ok else "FFCDD2"
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shd)
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32) if ok else RGBColor(0xC6, 0x28, 0x28)


# ════════ COST MODEL — đọc từ data/cost_model.json qua core.cost ════════
def _cost(Hse, q_uckse):
    return cushion_cost_per_m2(Hse, float(q_uckse))["cost_total_vnd"]


SIGMA_H_TOTAL = 1.10  # He + Hse (constraint hình học, vẫn dùng cho compute He)


# ════════ BUILD REPORT ════════
def build_cushion_report(
    co_name: str = "",
    co_staff: str = "",
    logo_bytes: Optional[bytes] = None,
    project_name: str = "Trung tâm Hành chính TP.HCM",
    location: str = "Thành phố Hồ Chí Minh",
    Hse_current: float = 0.40,
    qu_current: float = 600.0,
    D: float = 0.80,
    s: float = 1.80,
    options: Optional[List[Dict]] = None,
) -> bytes:
    """Tạo báo cáo Word đầy đủ. Trả về bytes file docx.

    options: list 3 phương án đề xuất, mỗi dict có:
        {"label", "Hse", "q_uckse", "note"}
    """
    if options is None:
        options = [
            {"label": "Phương án A — Tăng cường độ",
             "Hse": 0.40, "q_uckse": 800,
             "note": "Giữ Hse=0.40 m, tăng q_uckse 600 → 800 kPa (tăng hàm lượng XM trong mix)"},
            {"label": "Phương án B — Tăng chiều dày",
             "Hse": 0.55, "q_uckse": 600,
             "note": "Giữ q_uckse=600 kPa, tăng Hse 0.40 → 0.55 m (giảm He cát tương ứng)"},
            {"label": "Phương án C — Đệm mỏng cường độ cao",
             "Hse": 0.30, "q_uckse": 1000,
             "note": "Giảm Hse xuống 0.30 m, tăng q_uckse lên 1000 kPa (đệm mỏng + XM cao)"},
        ]

    doc = Document()
    _apply_default_style(doc)
    _setup_header_footer(doc, co_name=co_name, co_staff=co_staff, logo_bytes=logo_bytes)

    # ─── TRANG BÌA ───
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("BÁO CÁO KIỂM TOÁN"); r.font.size = Pt(22); r.bold = True
    r.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CHỌC THỦNG LỚP ĐỆM CÁT - XI MĂNG\n(Phương pháp ALiCC – PWRI Japan)")
    r.font.size = Pt(18); r.bold = True
    r.font.color.rgb = RGBColor(0x37, 0x47, 0x57)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run(f"Dự án: {project_name}"); r.font.size = Pt(14); r.bold = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Địa điểm: {location}"); r.font.size = Pt(12)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run("Ngày lập:")
    r.font.size = Pt(11); r.italic = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(_today_vn()); r.font.size = Pt(11)

    # ─── 1. TỔNG QUAN ───
    _h1(doc, "1. TỔNG QUAN VÀ CƠ SỞ LÝ THUYẾT", page_break=True)

    _h2(doc, "1.1. Vai trò của lớp đệm cát - xi măng trong hệ móng CDM")
    _para(doc,
          "Lớp đệm cát - xi măng (cement-treated sand cushion) được bố trí trực tiếp trên đầu các trụ "
          "đất xi măng (CDM) với vai trò là lớp truyền tải - phân bố ứng suất từ nền đắp xuống các trụ. "
          "Cơ chế truyền tải chính qua hiệu ứng vòm đất (soil arching): "
          "tải trọng nền đắp được tập trung vào đầu cọc nhờ vòm đất hình thành giữa các trụ CDM, "
          "phần tải còn lại tác dụng trực tiếp lên lớp đệm gây ra ứng suất cắt theo chu vi đầu cọc. "
          "Lớp đệm cần đủ cường độ kháng cắt để không bị chọc thủng dưới ứng suất này.")

    _h2(doc, "1.2. Phương pháp ALiCC (PWRI Japan)")
    _para(doc,
          "ALiCC (Advanced Low Improvement Cement Column) là phương pháp kiểm tra chọc thủng lớp đệm "
          "xi măng do Public Works Research Institute (PWRI, Nhật Bản) đề xuất. "
          "Phương pháp được trình bày trong tài liệu 'Technical Manual of ALiCC method for soft soil improvement'. "
          "Tiêu chí kiểm tra:")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("τ_se ≤ τ_ase")
    r.font.size = Pt(14); r.bold = True; r.italic = True
    r.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)

    # Sơ đồ
    fig_sch = chart_alicc_schematic(D=D, s=s, Hse=Hse_current, He=1.10 - Hse_current)
    _add_image_from_fig(doc, fig_sch, width_cm=15.5,
                         caption="Hình 1.1. Sơ đồ cơ chế truyền tải qua đệm cát - XM (ALiCC PWRI)")

    # ─── 2. THÔNG SỐ ĐẦU VÀO ───
    _h1(doc, "2. THÔNG SỐ ĐẦU VÀO", page_break=True)
    _para(doc,
          "Các thông số được lấy trực tiếp từ hồ sơ thiết kế (bảng E.2 'Thông số đầu vào cọc xi măng đất "
          "(CDM)'). Đối với q_uckse và Fs cụ thể cho lớp đệm, theo bảng E.2:")

    rows_inp = [
        ("Cường độ kháng nén thiết kế của đệm cát-XM", "q_uckse", f"{qu_current:.0f}", "kPa", "Hồ sơ E.2"),
        ("Hệ số an toàn cường độ cắt", "F_s", f"{FS_PUNCHING}", "—", "Hồ sơ E.2"),
        ("Góc đàn hồi dẻo (plastic arch angle)", "θ", f"{THETA_DEG:.0f}", "độ", "Hồ sơ E.2"),
        ("Đường kính trụ CDM", "D", f"{D:.2f}", "m", "tvtk_cdm_config"),
        ("Khoảng cách tâm trụ (lưới vuông)", "s", f"{s:.2f}", "m", "tvtk_cdm_config"),
        ("Bề dày đệm hiện tại", "H_se", f"{Hse_current:.2f}", "m", "tvtk_fill_composition"),
        ("Chiều cao cát He (trên đệm)", "H_e", f"{1.10 - Hse_current:.2f}", "m", "= 1.10 - H_se"),
        ("Bề dày áo đường", "h_aod", "0.80", "m", "tvtk_fill_composition"),
        ("Dung trọng áo đường", "γ_aod", "24.0", "kN/m³", "tvtk_fill_composition"),
        ("Dung trọng cát He", "γ_he", "18.0", "kN/m³", "tvtk_fill_composition"),
        ("Dung trọng đệm cát-XM", "γ_hse", "22.5", "kN/m³", "tvtk_fill_composition"),
    ]
    t = doc.add_table(rows=1, cols=5)
    hdrs = t.rows[0].cells
    for i, h in enumerate(["Tham số", "Ký hiệu", "Giá trị", "Đơn vị", "Nguồn"]):
        hdrs[i].text = h
    for row in rows_inp:
        cells = t.add_row().cells
        for i, v in enumerate(row): cells[i].text = str(v)
    style_tbl_for_docx(t)
    _para(doc, " ")

    # ─── 3. CÔNG THỨC KIỂM TOÁN ───
    _h1(doc, "3. CÔNG THỨC KIỂM TOÁN", page_break=True)

    _h2(doc, "3.1. Ứng suất cắt cho phép")
    _para(doc,
          "Cường độ cắt không thoát nước của đệm xi măng theo Mohr-Coulomb (φ=0) lấy bằng một nửa "
          "cường độ kháng nén nở hông, chia thêm cho hệ số an toàn F_s = 3:")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("τ_ase = q_uckse / (2 · F_s)"); r.bold = True; r.font.size = Pt(13)

    _h2(doc, "3.2. Hình học vòm đất (arching)")
    _para(doc, "Chiều cao vùng vòm đất hình thành giữa 2 trụ CDM kề nhau:")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("H_0 = (s − D) · tan(θ/2)"); r.bold = True; r.font.size = Pt(13)
    _para(doc, "So sánh H_0 với chiều cao đất đắp (h_aod + H_e):")
    _para(doc, "• CT(1) — H_0 ≤ He_total: vòm đất hình thành hoàn toàn trong nền đắp", italic=True)
    _para(doc, "• CT(2) — H_0 > He_total: nền đắp không đủ cao để hình thành vòm đầy đủ", italic=True)

    _h2(doc, "3.3. Thể tích đất truyền tải lên đệm")
    _para(doc, "Trường hợp CT(1):")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("V_soil = [((s−D)/2)·s² − π(s³−D³)/24 + (4−π)(√2−1)·s³/24] · tan(θ)")
    r.font.size = Pt(11); r.bold = True
    _para(doc, "Trường hợp CT(2):")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("V_soil = H_e·s² − (1/3)·[π·(H_e/tanθ + D/2)²·(H_e + D/2·tanθ) − π·(D/2·tanθ)]")
    r.font.size = Pt(11); r.bold = True

    _para(doc, "Thể tích phần đệm xi măng nằm dưới vùng truyền tải (V_CGCXM):")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("V_CGCXM = H_se·s² − (1/3)·[π·(H_se/tanθ + D/2)²·(H_se + D/2·tanθ) − π·(D/2·tanθ)]")
    r.font.size = Pt(11); r.bold = True

    _h2(doc, "3.4. Áp lực thẳng đứng trên đệm — P_soil")
    _para(doc, "Áp lực trung bình tác dụng lên phần đệm KHÔNG được trụ CDM đỡ:")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("P_soil = [(V_soil − V_CGCXM)·γ_fill + V_CGCXM·γ_hse] / A_unit")
    r.font.size = Pt(12); r.bold = True
    _para(doc, "Trong đó:", italic=True)
    _para(doc, "• A_unit = s² − π·D²/4 — diện tích phần không được trụ đỡ")
    _para(doc, "• γ_fill = (h_aod·γ_aod + H_e·γ_he) / (h_aod + H_e) — dung trọng TB trọng số đất đắp")

    _h2(doc, "3.5. Ứng suất cắt thực tế và tiêu chí kiểm tra")
    _para(doc, "Ứng suất cắt phát sinh trên mặt phá hoại chu vi cọc:")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("τ_se = (P_soil − q_a) · A_unit / (π · D · H_se)")
    r.font.size = Pt(13); r.bold = True
    _para(doc, "Tiêu chí ĐẠT: τ_se ≤ τ_ase  ⇔  ratio = τ_se/τ_ase ≤ 1")

    # ─── 4. TÍNH TOÁN HIỆN TRẠNG ───
    _h1(doc, "4. KẾT QUẢ TÍNH TOÁN HIỆN TRẠNG", page_break=True)
    _para(doc, f"Áp dụng cho cấu tạo hiện tại của dự án: "
                f"H_se = {Hse_current:.2f} m, q_uckse = {qu_current:.0f} kPa.")

    r0 = check_alicc(Hse_current, q_uckse=qu_current)
    result_rows = [
        ("Bề dày đệm H_se", f"{r0['Hse_m']:.3f}", "m"),
        ("Chiều cao đất đắp tương đương He_total", f"{r0['He_total_m']:.3f}", "m"),
        ("Dung trọng đắp TB (trọng số)", f"{r0['gamma_fill_TB']:.2f}", "kN/m³"),
        ("Chiều cao vòm H_0", f"{r0['H0_m']:.3f}", "m"),
        ("Trường hợp công thức", r0["case"], "—"),
        ("Thể tích đất V_soil", f"{r0['V_soil_m3']:.4f}", "m³"),
        ("Thể tích đệm V_CGCXM", f"{r0['V_CGCXM_m3']:.4f}", "m³"),
        ("Diện tích A_unit", f"{r0['A_unit_m2']:.4f}", "m²"),
        ("Áp lực P_soil", f"{r0['P_soil_kPa']:.2f}", "kPa"),
        ("Ứng suất cắt thực tế τ_se", f"{r0['tau_se_kPa']:.2f}", "kPa"),
        ("Ứng suất cắt cho phép τ_ase", f"{r0['tau_ase_kPa']:.2f}", "kPa"),
        ("Hệ số ratio = τ_se/τ_ase", f"{r0['ratio']:.3f}", "—"),
    ]
    t = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["Đại lượng", "Giá trị", "Đơn vị"]):
        t.rows[0].cells[i].text = h
    for row in result_rows:
        c = t.add_row().cells
        for i, v in enumerate(row): c[i].text = str(v)
    style_tbl_for_docx(t)

    # Verdict
    _para(doc, " ")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if r0["ok"]:
        r = p.add_run(f"KẾT LUẬN: ĐẠT  (ratio = {r0['ratio']:.3f} ≤ 1.000)")
        r.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    else:
        excess = (r0["ratio"] - 1.0) * 100
        r = p.add_run(f"KẾT LUẬN: KHÔNG ĐẠT  (ratio = {r0['ratio']:.3f} > 1.000, vượt {excess:.1f}%)")
        r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    r.bold = True; r.font.size = Pt(14)

    # ─── 5. MA TRẬN TỐI ƯU ───
    _h1(doc, "5. MA TRẬN TỐI ƯU HOÁ Hse × q_uckse", page_break=True)
    _para(doc,
          "Quét toàn bộ tổ hợp H_se ∈ [0.30, 0.70] m và q_uckse ∈ [500, 1200] kPa "
          "để xác định không gian thiết kế. Mỗi ô là giá trị ratio = τ_se/τ_ase. "
          "Ô có ratio ≤ 1 (xanh) → đạt; ô đỏ → không đạt.")

    fig_hm = chart_heatmap_ratio(Hse_current=Hse_current, qu_current=int(qu_current))
    _add_image_from_fig(doc, fig_hm, width_cm=16.5,
                         caption="Hình 5.1. Heatmap ratio = τ_se/τ_ase theo H_se và q_uckse")

    _h2(doc, "5.1. Sensitivity — Ratio theo q_uckse (Hse cố định)")
    fig_q = chart_ratio_vs_quckse_fixed_Hse(Hse_fixed=Hse_current)
    _add_image_from_fig(doc, fig_q, width_cm=16.0,
                         caption=f"Hình 5.2. ratio vs q_uckse tại H_se = {Hse_current:.2f} m")

    _h2(doc, "5.2. Sensitivity — Ratio theo Hse (q_uckse cố định)")
    fig_h = chart_ratio_vs_Hse_fixed_quckse(qu_fixed=int(qu_current))
    _add_image_from_fig(doc, fig_h, width_cm=16.0,
                         caption=f"Hình 5.3. ratio vs H_se tại q_uckse = {qu_current:.0f} kPa")

    # ─── 6. PARETO ───
    _h1(doc, "6. PHÂN TÍCH PARETO — Chi phí vs Hệ số an toàn", page_break=True)
    _para(doc,
          "Mỗi điểm trong biểu đồ Pareto là một tổ hợp (H_se, q_uckse) ứng với chi phí và ratio. "
          "Đường Pareto biểu diễn các phương án không-bị-thống-trị: bất kỳ điểm nào nằm trên đường này "
          "đều là 'tối ưu' theo nghĩa không thể giảm cả chi phí và ratio cùng lúc.")
    fig_p = chart_pareto_cost_ratio()
    _add_image_from_fig(doc, fig_p, width_cm=16.0,
                         caption="Hình 6.1. Đường Pareto: chi phí - hệ số an toàn (72 tổ hợp)")

    _para(doc,
          "Chú thích chi phí: tính tương đối cho 1 m² mặt nền, gồm (i) cát He thường ~200 nghìn VNĐ/m³, "
          "(ii) đệm cát-XM cơ bản (q_uckse = 600 kPa) ~700 nghìn VNĐ/m³, "
          "(iii) tăng 100 kPa cường độ ≈ +52,5 nghìn VNĐ/m³. "
          "Đây là mô hình minh hoạ — cần đối chiếu báo giá vật liệu thực tế.",
          italic=True, size_pt=10)

    # ─── 7. SO SÁNH 3 PHƯƠNG ÁN ───
    _h1(doc, "7. SO SÁNH 3 PHƯƠNG ÁN ĐỀ XUẤT", page_break=True)

    c_base = _cost(Hse_current, qu_current)
    enriched = []
    for o in options:
        ra = check_alicc(o["Hse"], q_uckse=float(o["q_uckse"]))
        c = _cost(o["Hse"], o["q_uckse"])
        enriched.append({**o, "ratio": ra["ratio"], "ok": ra["ok"],
                         "cost_delta_pct": (c - c_base) / c_base * 100,
                         "cost_vnd": c})

    # Bảng so sánh
    t = doc.add_table(rows=1, cols=7)
    for i, h in enumerate(["Phương án", "H_se (m)", "q_uckse (kPa)", "He (m)",
                           "ratio", "Đánh giá", "Δ chi phí (%)"]):
        t.rows[0].cells[i].text = h
    for o in enriched:
        c = t.add_row().cells
        c[0].text = o["label"]
        c[1].text = f"{o['Hse']:.2f}"
        c[2].text = f"{o['q_uckse']:.0f}"
        c[3].text = f"{SIGMA_H_TOTAL - o['Hse']:.2f}"
        c[4].text = f"{o['ratio']:.3f}"
        c[5].text = "Đạt" if o["ok"] else "Không đạt"
        c[6].text = f"{o['cost_delta_pct']:+.1f}%"
    style_tbl_for_docx(t)
    # Tô màu cột verdict
    for i, o in enumerate(enriched, start=1):
        _verdict_cell(t.rows[i].cells[5], o["ok"])

    # Biểu đồ so sánh
    chart_opts = [{"label": o["label"].split("—")[1].strip() if "—" in o["label"] else o["label"],
                    "Hse": o["Hse"], "q_uckse": o["q_uckse"],
                    "ratio": o["ratio"], "cost_delta_pct": o["cost_delta_pct"]}
                   for o in enriched]
    fig_cmp = chart_compare_options(chart_opts)
    _add_image_from_fig(doc, fig_cmp, width_cm=16.0,
                         caption="Hình 7.1. So sánh ratio và chi phí giữa 3 phương án")

    _h2(doc, "7.1. Mô tả từng phương án")
    for i, o in enumerate(enriched):
        _para(doc, f"{o['label']}", bold=True)
        _para(doc, f"   • {o['note']}")
        _para(doc, f"   • Kết quả: ratio = {o['ratio']:.3f} → "
                    f"{'ĐẠT' if o['ok'] else 'KHÔNG ĐẠT'}; "
                    f"Δ chi phí = {o['cost_delta_pct']:+.1f}%")

    # ─── 8. KẾT LUẬN ───
    _h1(doc, "8. KẾT LUẬN VÀ KHUYẾN NGHỊ", page_break=True)

    _h2(doc, "8.1. Đánh giá hiện trạng")
    if r0["ok"]:
        _para(doc, f"Cấu tạo hiện tại với H_se = {Hse_current:.2f} m và q_uckse = "
                    f"{qu_current:.0f} kPa ĐẠT yêu cầu chống chọc thủng theo ALiCC PWRI "
                    f"(ratio = {r0['ratio']:.3f}).")
    else:
        excess = (r0["ratio"] - 1.0) * 100
        _para(doc, f"Cấu tạo hiện tại KHÔNG ĐẠT yêu cầu — ratio = {r0['ratio']:.3f} > 1.0, "
                    f"vượt giới hạn {excess:.1f}%. Cần điều chỉnh thiết kế.")

    _h2(doc, "8.2. Khuyến nghị phương án thiết kế")
    # Recommend cheapest passing
    ok_opts = [o for o in enriched if o["ok"]]
    if ok_opts:
        cheapest = min(ok_opts, key=lambda o: o["cost_delta_pct"])
        _para(doc, f"Khuyến nghị Phương án {cheapest['label']}:", bold=True)
        _para(doc, f"   • {cheapest['note']}")
        _para(doc, f"   • Đảm bảo dự trữ an toàn ({(1 - cheapest['ratio'])*100:.1f}% so với ngưỡng)")
        _para(doc, f"   • Chi phí tăng tối thiểu ({cheapest['cost_delta_pct']:+.1f}% so với hiện trạng)")

    _h2(doc, "8.3. Lưu ý kỹ thuật")
    _para(doc, "• Mọi giá trị tính đều dùng dung trọng đắp γ_fill là trung bình trọng số "
                "(áo đường + cát He), KHÔNG dùng γ_he đơn lẻ. Đây là cách tiếp cận đúng theo "
                "lý thuyết ALiCC do V_soil bao gồm thể tích đất đắp toàn bộ vùng vòm.")
    _para(doc, "• Khi áp dụng tăng q_uckse, cần thí nghiệm trộn thử (mix design) để xác định "
                "hàm lượng xi măng tương ứng theo từng nguồn cát thực tế.")
    _para(doc, "• Khi áp dụng tăng H_se, cần điều chỉnh tương ứng cao độ He cát thường để "
                "giữ tổng chiều dày Σh = 1.90 m (tính từ cao độ đỉnh CDM đến cao độ thiết kế).")
    _para(doc, "• Báo cáo này chỉ kiểm chọc thủng (ALiCC). Các tính toán khác (lún cố kết, ổn định, "
                "sức chịu tải cọc) đã trình bày trong các báo cáo chuyên đề riêng.")

    # Save
    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf.getvalue()


def _today_vn() -> str:
    import datetime
    d = datetime.date.today()
    return f"Ngày {d.day:02d} tháng {d.month:02d} năm {d.year}"


# ════════ DEMO ════════
if __name__ == "__main__":
    import sys
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    data = build_cushion_report(
        co_name="Công ty TNHH Kỹ thuật Địa kỹ thuật TTHC",
        co_staff="Người lập: KS. ___________",
    )
    out = Path("scratch/cushion_report_demo.docx")
    out.parent.mkdir(exist_ok=True)
    out.write_bytes(data)
    print(f"OK: {out} ({len(data)/1024:.1f} KB)")
