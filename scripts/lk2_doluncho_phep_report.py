# -*- coding: utf-8 -*-
"""Xuất Word: Lý thuyết độ lún cố kết còn lại cho phép (TCCS 41:2022) +
kiến nghị chọn ΔS = 15 cm cho khu vực Bờ kè KE.

Tuân quy tắc dự án: tiếng Việt có dấu, không emoji, header bảng in đậm,
body 12pt, header/footer + đánh số trang (PAGE/NUMPAGES).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

_ROOT = Path(__file__).resolve().parent.parent
_OUT = _ROOT / "LyThuyet_DoLunChoPhep_BoKe_15cm.docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREEN = RGBColor(0x2E, 0x7D, 0x32)


def _field(para, field: str) -> None:
    run = para.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = field
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(instr); run._r.append(e)


def _no_border(tbl) -> None:
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}"); e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def _shade(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:fill"), hex_color)
    tcPr.append(sh)


def _set_body_font(doc) -> None:
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.4


def _h(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = NAVY
    r.font.size = Pt(15 if level == 1 else 13)
    p.paragraph_format.space_before = Pt(10)
    return p


def _para(doc, text, bold=False, italic=False, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def _table(doc, header, rows, widths=None, highlight_row=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, htext in enumerate(header):
        c = t.rows[0].cells[j]
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = c.paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(11)
        _shade(c, "1F3A5F")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        even = (i % 2 == 1)
        for j, val in enumerate(row):
            cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = cells[j].paragraphs[0].add_run(str(val))
            run.font.size = Pt(11)
            if highlight_row is not None and i == highlight_row:
                run.bold = True; _shade(cells[j], "DCFCE7")
            elif even:
                _shade(cells[j], "F3F4F6")
    if widths:
        for j, w in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = Cm(w)
    return t


def build() -> Path:
    doc = Document()
    _set_body_font(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2)
    pw = sec.page_width - sec.left_margin - sec.right_margin

    # Header / Footer
    ht = sec.header.add_table(1, 2, width=pw); _no_border(ht)
    ht.cell(0, 0).paragraphs[0].add_run("DỰ ÁN TRUNG TÂM HÀNH CHÍNH").bold = True
    rp = ht.cell(0, 1).paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.add_run("Bờ kè — Cọc xi măng đất CDM")
    ft = sec.footer.add_table(1, 2, width=pw); _no_border(ft)
    ft.cell(0, 0).paragraphs[0].add_run("Bản tính địa kỹ thuật")
    fp = ft.cell(0, 1).paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.add_run("Trang "); _field(fp, "PAGE"); fp.add_run(" / "); _field(fp, "NUMPAGES")

    # Title
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("ĐỘ LÚN CỐ KẾT CÒN LẠI CHO PHÉP\nVÀ KIẾN NGHỊ TIÊU CHÍ THIẾT KẾ CHO KHU VỰC BỜ KÈ")
    tr.bold = True; tr.font.size = Pt(17); tr.font.color.rgb = NAVY
    _para(doc, "Cơ sở: TCCS 41:2022 — Áo đường mềm và nền đường đắp trên đất yếu",
          italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # 1. Khái niệm
    _h(doc, "1. Khái niệm độ lún cố kết còn lại")
    _para(doc,
          "Độ lún cố kết còn lại (ΔS) là phần độ lún cố kết tiếp tục xảy ra của nền đất yếu "
          "SAU thời điểm đưa công trình vào khai thác. Theo TCCS 41:2022 (Điều 6.2.3), nền "
          "đường đắp trên đất yếu phải được xử lý sao cho độ lún cố kết còn lại tại trục tim "
          "đường trong thời hạn khai thác không vượt quá giới hạn cho phép, nhằm bảo đảm độ "
          "bằng phẳng và an toàn khai thác.")
    _para(doc, "Công thức xác định (công thức 36 — TCCS 41:2022):", bold=True)
    fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("ΔS = Sᴄ × (1 − Uₜ)")
    fr.bold = True; fr.font.size = Pt(13)
    _para(doc, "Trong đó:")
    for s in [
        "Sᴄ — tổng độ lún cố kết cuối cùng của nền đất yếu (cm);",
        "Uₜ — độ cố kết đạt được tại thời điểm t (tính theo lý thuyết Terzaghi);",
        "t — thời hạn tính toán: 15 năm với mặt đường cấp cao A1 (mặt đường mềm); "
        "30 năm với mặt đường cứng.",
    ]:
        p = doc.add_paragraph(s, style="List Bullet"); p.runs[0].font.size = Pt(12)

    # 2. Bảng giới hạn
    _h(doc, "2. Giới hạn độ lún cố kết còn lại cho phép — Bảng 1 (Điều 6.2.3)")
    _para(doc, "Giá trị giới hạn ΔS (cm) tại trục tim đường, theo cấp đường và vị trí đoạn nền đắp:")
    _table(
        doc,
        ["Loại cấp đường", "Gần mố cầu", "Hai bên cống / đường dân sinh", "Đoạn thông thường"],
        [
            ["Cao tốc và đường có v ≥ 80 km/h, mặt cấp cao A1", "10", "20", "30"],
            ["Đường có v ≤ 60 km/h, mặt cấp cao A1", "20", "30", "40"],
        ],
        widths=[7.5, 2.6, 4.0, 2.9],
    )
    _para(doc, "Ghi chú: vị trí càng nhạy cảm với chênh lún (gần mố cầu, hai bên cống) thì "
               "giới hạn càng nghiêm ngặt.", italic=True)

    # 3. Đặc thù bờ kè
    _h(doc, "3. Đặc thù khu vực Bờ kè và cơ sở chọn tiêu chí nghiêm ngặt")
    for s in [
        "Tuyến kè nằm sát sông, chịu đồng thời tải trọng đứng (đắp + hoạt tải) và tải "
        "trọng ngang (áp lực đất, áp lực nước, dao động mực nước triều) — chênh lún gây mất "
        "ổn định mái và nứt kết cấu mặt.",
        "Nền đất yếu rất dày (lớp bùn sét 1 dày trung bình ~24–27 m, e₀ ≈ 2,0, "
        "Su ≈ 14 kPa, sau hiệu chỉnh Bjerrum Cᵤ ≈ 12,5 kPa) — tiềm năng lún lớn.",
        "Kè là công trình lâu dài, yêu cầu an toàn và tuổi thọ cao; hư hỏng do lún khó "
        "sửa chữa và ảnh hưởng trực tiếp đến an toàn dân sinh ven sông.",
        "Mặt kè kết hợp đường dạo / hạ tầng kỹ thuật, đòi hỏi độ bằng phẳng cao tương "
        "đương các đoạn nhạy cảm.",
    ]:
        p = doc.add_paragraph(s, style="List Bullet"); p.runs[0].font.size = Pt(12)

    # 4. Kiến nghị
    _h(doc, "4. Kiến nghị tiêu chí thiết kế: ΔS = 15 cm")
    kp = doc.add_paragraph()
    kr = kp.add_run("KIẾN NGHỊ: Chọn độ lún cố kết còn lại cho phép ΔS = 15 cm cho toàn "
                    "khu vực Bờ kè.")
    kr.bold = True; kr.font.size = Pt(13); kr.font.color.rgb = GREEN
    _para(doc, "Luận cứ:", bold=True)
    for s in [
        "ΔS = 15 cm nằm giữa mức 10 cm (đoạn gần mố cầu) và 20 cm (hai bên cống) của nhóm "
        "đường cao tốc / cấp cao A1 — phản ánh mức độ quan trọng cao của tuyến kè nhưng vẫn "
        "khả thi về khối lượng xử lý.",
        "Nghiêm ngặt hơn mức 20 cm (đoạn hai bên cống, cấp cao) và mức 20 cm (gần mố cầu "
        "đường cấp ≤ 60 km/h), bảo đảm biên an toàn cho công trình ven sông.",
        "Phù hợp định hướng kiểm soát chênh lún dọc tuyến kè, hạn chế nứt mặt và mất ổn "
        "định mái trong suốt thời hạn khai thác 15 năm.",
    ]:
        p = doc.add_paragraph(s, style="List Bullet"); p.runs[0].font.size = Pt(12)
    _para(doc, "Bảng định vị mức 15 cm so với các giá trị trong Bảng 1:", bold=True)
    _table(
        doc,
        ["Mức ΔS (cm)", "Ý nghĩa tham chiếu", "Áp dụng"],
        [
            ["10", "Cao tốc/A1 — gần mố cầu", "Nghiêm ngặt nhất"],
            ["15", "KIẾN NGHỊ cho Bờ kè", "Chọn thiết kế"],
            ["20", "Cao tốc/A1 — hai bên cống", "Tham chiếu trên"],
            ["30", "Cao tốc/A1 — đoạn thông thường", "Ít nghiêm ngặt"],
            ["40", "Đường v ≤ 60 km/h — đoạn thông thường", "Ít nghiêm ngặt nhất"],
        ],
        widths=[3.0, 8.0, 4.0],
        highlight_row=1,
    )

    # 5. Hệ quả thiết kế
    _h(doc, "5. Hệ quả đối với thiết kế cọc xi măng đất (CDM)")
    _para(doc,
          "Với ΔS = 15 cm, chiều dài và cường độ cọc CDM phải bảo đảm tổng độ lún còn lại "
          "(lún khối gia cố S₁ cộng lún cố kết phần dưới mũi cọc S₂) không vượt 15 cm sau "
          "15 năm. Do lớp bùn rất dày và yếu, để đạt ΔS = 15 cm thường cần cọc xuyên hết lớp "
          "bùn vào lớp đất tốt, kết hợp tăng cường độ trụ qᵤ hoặc giảm khoảng cách bố trí "
          "trụ. Các phương án chiều dài và bố trí trụ được kiểm toán riêng theo từng vùng.")
    _para(doc,
          "Khuyến nghị: với những vị trí mà chỉ tiêu 15 cm khó đạt do bề dày bùn quá lớn, "
          "cần phân tích bù bằng giải pháp tăng qᵤ, giảm khoảng cách trụ s, hoặc gia tải "
          "trước, đồng thời kiểm soát chênh lún giữa các vùng kề nhau theo điều kiện độ bằng "
          "phẳng (Phụ lục E — TCCS 41:2022).", italic=True)

    doc.add_paragraph()
    _para(doc, "Tài liệu viện dẫn: TCCS 41:2022, Điều 6.2.3, Bảng 1, công thức (36); "
               "Phụ lục C (cố kết Terzaghi).", italic=True)

    doc.save(str(_OUT))
    return _OUT


if __name__ == "__main__":
    out = build()
    print(f"Da xuat: {out}  ({out.stat().st_size/1024:.0f} KB)")
