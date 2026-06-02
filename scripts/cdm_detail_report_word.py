"""Word builder — Bản tính chi tiết CDM per hố khoan, gom theo 6 vùng Bờ kè KE.

Theo mẫu điển hình (data/cdm_report_template_dienhinh.json): mỗi HK 1 mục gồm
5 phần (thông số · S1 · S2 · lún theo thời gian · sức chịu tải · kiểm đệm).

Dùng chung dict từ cdm_detail_report.build_6zone_detail (Rule 6 — UI ↔ Word parity).

API: build_6zone_detail_docx(zones, dS, co_name="", logo_bytes=None) -> bytes
     build_zone_detail_to_file(out_path, dS)  # tiện CLI

Quy tắc: tiếng Việt có dấu, header bảng bold, body 12pt, KHÔNG emoji.
"""
from __future__ import annotations

import io
from typing import Optional

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

_GREEN = RGBColor(0x2E, 0x7D, 0x32)
_RED = RGBColor(0xC6, 0x28, 0x28)
_NAVY = RGBColor(0x10, 0x2A, 0x43)
_HDR_BG = "D9E1F2"
_ZEBRA_BG = "F3F4F6"


# ─────────────────────────────────────────────────────────────────────────────
# Low-level helpers
# ─────────────────────────────────────────────────────────────────────────────
def _shade(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell(cell, text, bold=False, size=11, color=None, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def _add_table(doc, headers, rows, widths=None, zebra=True, size=10):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for j, h in enumerate(headers):
        _set_cell(tbl.rows[0].cells[j], h, bold=True, size=size,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(tbl.rows[0].cells[j], _HDR_BG)
    for i, r in enumerate(rows):
        cells = tbl.add_row().cells
        for j, v in enumerate(r):
            _set_cell(cells[j], v, size=size)
            if zebra and i % 2 == 1:
                _shade(cells[j], _ZEBRA_BG)
    if widths:
        for row in tbl.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Cm(w)
    return tbl


def _field_run(para, field: str) -> None:
    r = para.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = field
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    r._r.append(b); r._r.append(instr); r._r.append(e)


def _no_border(tbl) -> None:
    t = tbl._tbl
    tblPr = t.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tblPr.append(borders)


def _setup_header_footer(doc, co_name: str, logo_bytes, dS: float) -> None:
    sec = doc.sections[0]
    pw = sec.page_width - sec.left_margin - sec.right_margin
    # Header
    ht = sec.header.add_table(1, 2, width=pw)
    ht.cell(0, 0).width = Cm(3); ht.cell(0, 1).width = pw - Cm(3)
    if logo_bytes:
        try:
            ht.cell(0, 0).paragraphs[0].add_run().add_picture(
                io.BytesIO(logo_bytes), height=Cm(1.1))
        except Exception:
            pass
    pr = ht.cell(0, 1).paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = pr.add_run(co_name or "")
    run.bold = True; run.font.size = Pt(9)
    _no_border(ht)
    # Footer
    ft = sec.footer.add_table(1, 2, width=pw)
    ft.cell(0, 0).width = pw - Cm(3); ft.cell(0, 1).width = Cm(3)
    fl = ft.cell(0, 0).paragraphs[0]
    fl.add_run(f"Bản tính chi tiết CDM 6 vùng — ΔS = {dS:.0f} cm").font.size = Pt(8)
    fr = ft.cell(0, 1).paragraphs[0]
    fr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr.add_run("Trang ").font.size = Pt(8)
    _field_run(fr, "PAGE"); fr.add_run(" / ").font.size = Pt(8)
    _field_run(fr, "NUMPAGES")
    _no_border(ft)


def _time_chart_png(time_curve, dS_allow):
    """PNG biểu đồ lún-thời gian (matplotlib Agg). None nếu lỗi."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        xs = [r["t_years"] for r in time_curve]
        ys = [r["S_total_t_cm"] for r in time_curve]
        fig, ax = plt.subplots(figsize=(6.0, 2.6), dpi=130)
        ax.plot(xs, ys, "-o", color="#1565C0", lw=1.8, ms=4)
        for x, y in zip(xs, ys):
            ax.annotate(f"{y:.1f}", (x, y), textcoords="offset points",
                        xytext=(0, 5), ha="center", fontsize=7)
        ax.axhline(dS_allow, ls=":", color="#C62828",
                   label=f"[S] = {dS_allow:.0f} cm")
        ax.set_xlabel("Thời gian (năm)"); ax.set_ylabel("Độ lún (cm)")
        ax.set_title("Độ lún tổng theo thời gian", fontsize=9)
        ax.legend(fontsize=7); ax.grid(alpha=0.3)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png"); plt.close(fig)
        buf.seek(0)
        return buf
    except Exception:
        return None


# ─────────────────────────────────────────────────────────────────────────────
# Render 1 hố khoan
# ─────────────────────────────────────────────────────────────────────────────
def _render_hk(doc, d: dict) -> None:
    g, m, s1, s2 = d["geometry"], d["material"], d["S1"], d["S2"]
    t, b = d["time"], d["bearing"]
    doc.add_heading(f"Hố khoan {d['bh_name']}", level=2)

    # Verdict line
    p = doc.add_paragraph()
    for label, ok in (("Lún", t.get("ok")), ("Sức chịu tải", b.get("ok")),
                      ("Đệm", (d["cushion"] or {}).get("ok"))):
        p.add_run(f"{label}: ").bold = True
        rr = p.add_run(("Đạt" if ok else "Không đạt") if ok is not None else "—")
        rr.bold = True
        if ok is not None:
            rr.font.color.rgb = _GREEN if ok else _RED
        p.add_run("    ")

    # Phần 1.1 + 2.1
    doc.add_heading("1. Thông số tính toán", level=3)
    _add_table(doc,
               ["Thông số hình học", "Giá trị", "Thông số vật liệu", "Giá trị"],
               [
                   ("Bề rộng nhóm W (m)", g["W_m"], "Cường độ trụ qu (kPa)", m["qu_kPa"]),
                   ("Đường kính D (m)", g["D_m"], "Mô đun trụ Ec (kPa)", f"{m['Ec_kPa']:.0f}"),
                   ("Khoảng cách S (m)", g["S_m"], "Mô đun đất Es (kPa)", f"{m['Es_kPa']:.0f}"),
                   ("Hệ số cải tạo a (%)", g["a_pct"], "Eeq (kPa)", f"{m['Eeq_kPa']:.0f}"),
                   ("CDTK (m)", g["CDTK_m"], "Bjerrum μ", m["mu"]),
                   ("CDld (m)", g["CDld_m"], "Ip lớp yếu (%)", m["Ip"]),
                   ("CDTN (m)", g["CDTN_m"], "cu = μ·Su (kPa)", m["cu_used_kPa"]),
                   ("CDNN (m)", g["CDNN_m"], "Tải lún q (kPa)", d["loads"]["q_static_kPa"]),
                   ("CD1 đỉnh trụ (m)", g["CD1_m"], "Chiều dài trụ L (m)", g["L_m"]),
                   ("CD2 đáy trụ (m)", g["CD2_m"], "H đất yếu (m)", g["H_soft_m"]),
               ], widths=[4.5, 2.2, 4.5, 2.2])

    if d.get("soil_layers"):
        doc.add_paragraph().add_run("1.2 Thông số địa chất theo lớp").bold = True
        rows = [(l.get("symbol"), l.get("description", "")[:28],
                 l.get("depth_top_m"), l.get("depth_bot_m"),
                 l.get("gamma"), l.get("e0"), l.get("Cc"), l.get("Cs"), l.get("PC"))
                for l in d["soil_layers"]]
        _add_table(doc, ["Lớp", "Mô tả", "Từ (m)", "Đến (m)", "γ", "e0",
                          "Cc", "Cs", "PC"], rows, size=9)

    # Phần 2.2 S1
    doc.add_heading("2. Tính lún dưới chân trụ", level=3)
    doc.add_paragraph(
        f"2.2 Độ lún khối gia cố S1 = {s1['S1_cm']:.2f} cm "
        f"(Công thức C.2 phụ lục C TCVN 9403:2012; Σ per-lớp = "
        f"{s1['S1_layered_cm']:.2f} cm).")
    if s1.get("layers"):
        rows = [(l["depth_top_m"], l["depth_bot_m"], l["Cu_kPa"], l["Cu_corr_kPa"],
                 f"{l['Esoil_kPa']:.0f}", f"{l['Eeq_kPa']:.0f}", f"{l['Si_cm']:.3f}")
                for l in s1["layers"]]
        _add_table(doc, ["Từ (m)", "Đến (m)", "Su", "cu=μSu", "Esoil",
                          "Eeq", "Si (cm)"], rows, size=9)

    # Phần 2.3 S2
    doc.add_paragraph(
        f"2.3 Độ lún cố kết dưới mũi S2 = {s2['S2_cm']:.2f} cm "
        f"(lún dư thiết kế {s2['S2_residual_cm']:.2f} cm; {s2['n_layers']} phân tố, "
        f"dừng tại z = {s2['stop_depth_m']} m — Δσ/σ'v < 10%).")
    if s2.get("layers"):
        rows = []
        for l in s2["layers"]:
            rows.append((
                round(l.get("depth_mid_m", 0), 1), l.get("symbol"),
                "Cát" if l.get("is_sand") else "Sét",
                round(l.get("sigma_v0_kPa", 0), 1), round(l.get("dsigma_kPa", 0), 1),
                l.get("Cc"), l.get("method"), round(l.get("Si_cm", 0), 3)))
        _add_table(doc, ["z (m)", "Lớp", "Loại", "σ'v0", "Δσ", "Cc",
                          "PP", "Si (cm)"], rows, size=9)

    # Phần III — lún theo thời gian
    doc.add_heading("III. Độ lún cố kết theo thời gian", level=3)
    if t.get("curve"):
        rows = [(r["t_years"], f"{r['Uv_pct']:.1f}", f"{r['S2_t_cm']:.2f}",
                 f"{r['S_total_t_cm']:.2f}") for r in t["curve"]]
        _add_table(doc, ["t (năm)", "Uv (%)", "S2(t) (cm)", "S tổng (cm)"],
                   rows, size=10)
        png = _time_chart_png(t["curve"], t["dS_allow_cm"])
        if png is not None:
            doc.add_picture(png, width=Cm(13))
    p = doc.add_paragraph()
    p.add_run("Kết luận lún: ").bold = True
    r = p.add_run(
        f"S thiết kế = {t['S_design_cm']:.2f} cm "
        f"{'≤' if t['ok'] else '>'} [S] = {t['dS_allow_cm']:.0f} cm → "
        + ("Đạt" if t["ok"] else "Không đạt") + " (điều 6.3.2 TCCS 41:2022)")
    r.bold = True; r.font.color.rgb = _GREEN if t["ok"] else _RED

    # Phần 4 — sức chịu tải
    doc.add_heading("IV. Sức chịu tải cọc xi măng đất", level=3)
    _add_table(doc, ["Đại lượng", "Giá trị (kN)"], [
        ("Tải tác dụng 1 cọc N", round(b.get("N_applied_kN", 0), 1)),
        ("Theo vật liệu Nvl", round(b.get("Q_ult_material_kN", 0), 1)),
        ("Theo nền Nđn", round(b.get("Q_ult_soil_kN", 0), 1)),
        ("Tối thiểu [N]", round(b.get("Q_ult_min_kN", 0), 1)),
        ("Cho phép Qa (FS=2.5)", round(b.get("Q_allow_kN") or 0, 1)),
    ], widths=[8.0, 4.0])
    p = doc.add_paragraph(); p.add_run("Kết luận sức chịu tải: ").bold = True
    okb = b.get("ok")
    rr = p.add_run(("Đạt" if okb else "Không đạt") if okb is not None else "—")
    rr.bold = True
    if okb is not None:
        rr.font.color.rgb = _GREEN if okb else _RED

    # Phần 5 — đệm
    if d.get("cushion"):
        c = d["cushion"]
        doc.add_heading("V. Kiểm chọc thủng đệm cát - xi măng (ALiCC)", level=3)
        p = doc.add_paragraph(
            f"τ_se = {c['tau_se_kPa']:.1f} kPa vs τ_ase = {c['tau_ase_kPa']:.1f} kPa "
            f"(tỉ số = {c['ratio']:.2f}) → ")
        rr = p.add_run("Đạt" if c["ok"] else "Không đạt")
        rr.bold = True; rr.font.color.rgb = _GREEN if c["ok"] else _RED

    if d.get("warnings"):
        doc.add_paragraph("Lưu ý: " + "; ".join(d["warnings"][:3])).italic = True


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────
def build_6zone_detail_docx(zones: list, dS: float,
                            co_name: str = "", logo_bytes: Optional[bytes] = None
                            ) -> bytes:
    doc = Document()
    # Body 12pt (Rule 5)
    try:
        doc.styles["Normal"].font.size = Pt(12)
    except Exception:
        pass
    _setup_header_footer(doc, co_name, logo_bytes, dS)

    # Trang bìa
    h = doc.add_heading("BẢN TÍNH CHI TIẾT CỌC ĐẤT GIA CỐ XI MĂNG (CDM)", level=0)
    for r in h.runs:
        r.font.color.rgb = _NAVY
    sub = doc.add_paragraph("Khu vực Bờ kè — Phân vùng gia cố theo 6 vùng")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(
        f"Độ lún cố kết còn lại cho phép ΔS = {dS:.0f} cm · "
        "Tiêu chuẩn: TCVN 9403:2012, TCCS 41:2022/TCĐBVN")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Bảng tổng hợp 6 vùng
    doc.add_heading("Tổng hợp 6 vùng", level=1)
    seen = set()
    sum_rows = []
    for z in zones:
        bhs = ", ".join(mm["bh_name"] for mm in z["members"])
        sum_rows.append((z["zone_name"], bhs, f"{z['total_length_m']:.1f}"))
    _add_table(doc, ["Vùng", "Hố khoan", "Tổng tuyến (m)"], sum_rows,
               widths=[3.0, 9.0, 3.0])

    # Chi tiết theo vùng
    for z in zones:
        doc.add_heading(
            f"{z['zone_name']} — " + ", ".join(mm["bh_name"] for mm in z["members"]),
            level=1)
        for d in z["details"]:
            bn = d.get("bh_name")
            if not d.get("ok"):
                doc.add_paragraph(f"{bn}: {d.get('error', 'thiếu dữ liệu')}")
                continue
            if bn in seen:
                doc.add_paragraph(
                    f"{bn} — đã trình bày chi tiết ở vùng trước (xem mục tương ứng).")
                continue
            seen.add(bn)
            _render_hk(doc, d)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_zone_detail_to_file(out_path, dS: float = 30.0, db_path=None) -> str:
    from cdm_detail_report import build_6zone_detail
    zones = build_6zone_detail(float(dS), db_path=db_path)
    data = build_6zone_detail_docx(zones, dS)
    with open(out_path, "wb") as f:
        f.write(data)
    return str(out_path)


if __name__ == "__main__":
    import sys
    out = sys.argv[1] if len(sys.argv) > 1 else "plaxis_out/BanTinhCDM_6Vung_dS30.docx"
    ds = float(sys.argv[2]) if len(sys.argv) > 2 else 30.0
    from pathlib import Path
    Path(out).parent.mkdir(parents=True, exist_ok=True)
    p = build_zone_detail_to_file(out, ds)
    print("Saved:", p)
