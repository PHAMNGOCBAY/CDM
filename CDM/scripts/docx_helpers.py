# -*- coding: utf-8 -*-
"""
docx_helpers.py – Thư viện tạo tài liệu Word cho báo cáo địa kỹ thuật.
Không chứa dữ liệu dự án – dùng lại cho mọi dự án tương tự.

Lịch sử:
  v1  2026-05-13  Tách từ gen_thuyet_minh.py monolithic; tất cả helper Word
  v2  2026-05-13  BytesIO thay temp file; gộp caption/_set_cell; thêm vn_float;
                  _counters dict thay global _tbl/_fig
"""
import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker

from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Bộ đếm bảng / hình ───────────────────────────────────────────────────────
_counters = {'tbl': 0, 'fig': 0}

def reset_counters():
    _counters['tbl'] = 0
    _counters['fig'] = 0

# ── Thiết lập trang ──────────────────────────────────────────────────────────
def setup_page(doc, w_cm=21, h_cm=29.7, left=2.5, right=2.0, top=2.5, bottom=2.0):
    sec = doc.sections[0]
    sec.page_width    = Cm(w_cm);   sec.page_height    = Cm(h_cm)
    sec.left_margin   = Cm(left);   sec.right_margin   = Cm(right)
    sec.top_margin    = Cm(top);    sec.bottom_margin  = Cm(bottom)

# ── Font ──────────────────────────────────────────────────────────────────────
def _fmt(run, bold=False, size=12, italic=False, color=None):
    run.bold       = bold
    run.italic     = italic
    run.font.size  = Pt(size)
    run.font.name  = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    if color:
        run.font.color.rgb = RGBColor(*color)

# ── Số thực định dạng tiếng Việt (dấu phẩy thập phân) ───────────────────────
def vn_float(s):
    return float(str(s).replace(',', '.'))

# ── Đoạn văn bản ─────────────────────────────────────────────────────────────
def heading(doc, text, size=13, bold=True, center=False, space_before=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    _fmt(p.add_run(text), bold=bold, size=size)
    return p

def body(doc, text, indent=0, bold=False, italic=False, justify=True, sb=2, sa=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    _fmt(p.add_run(text), bold=bold, size=12, italic=italic)
    return p

def formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    _fmt(p.add_run(text), italic=True, size=12)

def body_sym(doc, sym, desc, indent=0.8):
    """Dòng định nghĩa ký hiệu: ký hiệu đậm nghiêng + mô tả thường."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    _fmt(p.add_run(f'{sym}: '), bold=True, italic=True, size=11)
    _fmt(p.add_run(desc),                              size=11)

# ── Mục lục tự động ──────────────────────────────────────────────────────────
def add_toc(doc):
    heading(doc, 'MỤC LỤC', size=13, bold=True, center=True, space_before=0)
    p = doc.add_paragraph()
    r1 = OxmlElement('w:r')
    fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin')
    r1.append(fc1); p._p.append(r1)

    r2 = OxmlElement('w:r')
    ins = OxmlElement('w:instrText'); ins.set(qn('xml:space'), 'preserve')
    ins.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2.append(ins); p._p.append(r2)

    r3 = OxmlElement('w:r')
    fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'end')
    r3.append(fc2); p._p.append(r3)
    doc.add_page_break()

# ── Caption bảng / hình ───────────────────────────────────────────────────────
def _caption(doc, prefix, key, title, space_before, space_after):
    _counters[key] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    _fmt(p.add_run(f'{prefix} {_counters[key]}. {title}'), bold=True, italic=True, size=11)

def tbl_caption(doc, title):
    _caption(doc, 'Bảng', 'tbl', title, space_before=6, space_after=2)

def fig_caption(doc, title):
    _caption(doc, 'Hình', 'fig', title, space_before=4, space_after=8)

# ── Ô bảng ────────────────────────────────────────────────────────────────────
def _shd(cell, hex_color):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), hex_color); s.set(qn('w:val'), 'clear')
    pr.append(s)

def _set_cell(cell, text, bold=False, size=10, center=True, bg=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    _fmt(p.add_run(text), bold=bold, size=size)
    if bg:
        _shd(cell, bg)

def hdr_cell(cell, text, bg='D9E1F2', bold=True, size=10, center=True):
    _set_cell(cell, text, bold=bold, size=size, center=center, bg=bg)

def dat_cell(cell, text, bold=False, size=10, center=True):
    _set_cell(cell, text, bold=bold, size=size, center=center)

def make_table(doc, headers, rows, highlight_rows=None, bg_hdr='D9E1F2'):
    n_col = len(headers)
    t = doc.add_table(rows=len(rows) + 1, cols=n_col)
    t.style     = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        hdr_cell(t.rows[0].cells[j], h, bg=bg_hdr)
    for i, row in enumerate(rows):
        hl = highlight_rows and i in highlight_rows
        bg = 'E2EFDA' if hl else None
        for j, val in enumerate(row):
            c = t.rows[i + 1].cells[j]
            if bg:
                hdr_cell(c, str(val), bg=bg, bold=(j == len(row) - 1), size=10)
            else:
                dat_cell(c, str(val), center=(j != 1 if n_col > 3 else True))
    return t

# ── Biểu đồ cột ──────────────────────────────────────────────────────────────
DEFAULT_COLORS = ['#4472C4', '#ED7D31', '#A9D18E', '#FF5050', '#7030A0']

def add_chart(doc, title, x_labels, datasets, ylabel,
              colors=None, width_cm=14, annotate=True):
    fig, ax = plt.subplots(figsize=(10, 5))
    n = len(x_labels); m = len(datasets)
    bw = min(0.65 / m, 0.35)
    for i, ds in enumerate(datasets):
        off  = (i - m / 2 + 0.5) * bw
        col  = colors[i] if colors and i < len(colors) else DEFAULT_COLORS[i % 5]
        bars = ax.bar([x + off for x in range(n)], ds['values'], width=bw,
                      label=ds['label'], color=col, edgecolor='black', linewidth=0.5)
        if annotate:
            for b in bars:
                ax.text(b.get_x() + b.get_width() / 2, b.get_height() * 1.01,
                        f"{b.get_height():.2f}", ha='center', va='bottom', fontsize=8)
    ax.set_title(title, fontsize=11, fontweight='bold', pad=10)
    ax.set_xticks(list(range(n))); ax.set_xticklabels(x_labels, fontsize=10)
    ax.set_ylabel(ylabel, fontsize=10)
    ax.legend(fontsize=9, loc='upper left')
    ax.yaxis.set_minor_locator(ticker.AutoMinorLocator())
    ax.grid(axis='y', linestyle='--', alpha=0.4)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    doc.add_picture(buf, width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_caption(doc, title)


# ═══════════════════════════════════════════════════════════════════════════════
# NÂNG CAO — Công thức LaTeX, biểu đồ matplotlib, bảng styled
# ═══════════════════════════════════════════════════════════════════════════════
def add_latex_formula(doc, latex_str: str, center=True, width_cm: float | None = None):
    """Render công thức LaTeX qua matplotlib mathtext → embed PNG vào Word.
    Sử dụng khi cần công thức chuyên nghiệp (vd phân số, tích phân, σ').

    Ví dụ:
        add_latex_formula(doc, r"S_1 = \\frac{q \\cdot H_{soft}}{a E_c + (1-a) E_s}")
    """
    fig = plt.figure(figsize=(6, 0.9))
    fig.patch.set_alpha(0)
    plt.axis('off')
    # mathtext không cần TeX runtime — built-in matplotlib
    plt.text(0.5, 0.5, f"${latex_str}$", fontsize=18, ha='center', va='center',
             usetex=False)  # mathtext built-in
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=200, bbox_inches='tight',
                transparent=True, pad_inches=0.05)
    plt.close(fig)
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(buf, width=Cm(width_cm) if width_cm else Cm(9))


def add_mpl_figure(doc, mpl_fig, title: str | None = None, width_cm: float = 14):
    """Embed matplotlib Figure → Word (dpi cao, có caption tự đánh số)."""
    buf = io.BytesIO()
    mpl_fig.savefig(buf, format='png', dpi=160, bbox_inches='tight',
                    facecolor='white')
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Cm(width_cm))
    if title:
        fig_caption(doc, title)


def styled_table(doc, headers, rows, highlight_rows=None,
                 col_widths_cm: list | None = None,
                 header_bg='1F4E79', header_color=(255, 255, 255),
                 zebra=True):
    """Bảng dữ liệu chuyên nghiệp:
    - Header xanh đậm, chữ trắng
    - Zebra (xen kẽ trắng/xám nhạt) khi zebra=True
    - Hàng highlight nền vàng nhạt"""
    n_col = len(headers)
    t = doc.add_table(rows=len(rows) + 1, cols=n_col)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths_cm:
        for j, w in enumerate(col_widths_cm[:n_col]):
            for r in range(len(rows) + 1):
                t.rows[r].cells[j].width = Cm(w)
    # Header
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        _fmt(run, bold=True, size=10, color=header_color)
        _shd(c, header_bg)
    # Data rows
    for i, row in enumerate(rows):
        is_hl = highlight_rows and i in highlight_rows
        for j, val in enumerate(row):
            c = t.rows[i + 1].cells[j]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            _fmt(run, bold=is_hl, size=10)
            if is_hl:
                _shd(c, 'FFF2CC')   # vàng nhạt cho hàng kiến nghị
            elif zebra and i % 2 == 1:
                _shd(c, 'F4F6F9')   # xám nhạt zebra
    return t


def list_bullet(doc, text: str, indent: float = 1.0, size: int = 11):
    """Đoạn văn dấu gạch ngang đúng chuẩn báo cáo kỹ thuật VN."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _fmt(p.add_run(f"— {text}"), size=size)


def info_box(doc, title: str, lines: list[str], bg='E3F2FD'):
    """Khung thông tin nổi bật (vd kết quả, kết luận)."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    _shd(c, bg)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    _fmt(c.paragraphs[0].add_run(title), bold=True, size=11)
    for ln in lines:
        p = c.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        _fmt(p.add_run(ln), size=10)
    return t


def conclusion_box(doc, ok: bool, text: str):
    """Hộp kết luận xanh (đạt) hoặc đỏ (không đạt)."""
    bg = 'E2EFDA' if ok else 'F8CBAD'
    info_box(doc, ("KẾT LUẬN: ĐẠT" if ok else "KẾT LUẬN: KHÔNG ĐẠT"),
             [text], bg=bg)
