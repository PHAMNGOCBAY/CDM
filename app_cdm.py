# -*- coding: utf-8 -*-
"""
app_cdm.py  –  Ứng dụng thiết kế cọc đất xi măng (CDM)
Streamlit, cấu trúc tương tự app_coc_tai_ngang.py

Chạy:  streamlit run app_cdm.py --server.port 8503
"""
import io
import json
import math
import sys
from pathlib import Path

import pandas as pd
import streamlit as st

try:
    import plotly.graph_objects as go
    _HAS_PLOTLY = True
except ImportError:
    _HAS_PLOTLY = False

import sqlite3

# ── Đường dẫn ────────────────────────────────────────────────────────────────
_ROOT    = Path(__file__).parent.parent
_DB      = _ROOT / "data" / "TTHC.sqlite"
_CDM_SC  = _ROOT / "CDM" / "scripts"
_THEORY  = _ROOT / "35-ly-thuyet-cdm.md"


@st.cache_data(show_spinner=False)
def _load_theory() -> str:
    try:
        return _THEORY.read_text(encoding="utf-8")
    except FileNotFoundError:
        return ""

st.set_page_config(
    page_title="CDM Design Tool",
    page_icon="🏗",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── CSS nhỏ để khớp style app hiện tại ───────────────────────────────────────
st.markdown("""
<style>
[data-testid="stSidebar"] { min-width: 220px; max-width: 280px; }
div[data-testid="stMetricValue"] { font-size: 1.1rem; }
#MainMenu { visibility: hidden; }
[data-testid="stToolbar"] { display: none; }
[data-testid="stDecoration"] { display: none; }
footer { visibility: hidden; }
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════════
# CONSTANTS
# ═══════════════════════════════════════════════════════════════════════════════
_CLAY_SYMBOLS = {        # symbol lớp bùn sét yếu theo zone
    "KE":  ["1", "1b"],
    "BXN": ["2"],
    "NHC": ["2", "1"],   # NHC chưa có layers – fallback
}
_ZONE_NAMES = {"KE": "Kè Công Viên (KE)", "BXN": "Bãi Đỗ Xe Ngầm (BXN)", "NHC": "Nhà Hành Chính (NHC)"}

_LAYER_COLORS: dict[str, str] = {
    "F":    "#9E9E9E",
    "1":    "#81D4FA",
    "1b":   "#4FC3F7",
    "2":    "#1565C0",
    "2a":   "#66BB6A",
    "2b":   "#FDD835",
    "2c":   "#C5E1A5",
    "3":    "#8D6E63",
    "3a":   "#A5D6A7",
    "3b":   "#69B578",
    "3c":   "#43A047",
    "4":    "#5C85D6",
    "5":    "#FFA726",
    "5a":   "#FB8C00",
    "5b":   "#E65100",
    "6":    "#6D4C41",
    "7":    "#795548",
    "8":    "#BCAAA4",
    "TK6a": "#FFD54F",
    "TK6b": "#FFB300",
    "XMD":  "#E91E63",
}
_LAYER_DEFAULT_COLOR = "#EEEEEE"
_ZONE_MARKER: dict[str, str] = {"KE": "circle", "BXN": "square", "NHC": "diamond"}

_PAGES = ["Địa chất", "Thông số", "So sánh PA", "Kết quả", "Xuất"]


# ═══════════════════════════════════════════════════════════════════════════════
# DB HELPERS
# ═══════════════════════════════════════════════════════════════════════════════
def _db() -> sqlite3.Connection:
    con = sqlite3.connect(_DB)
    con.row_factory = sqlite3.Row
    return con


@st.cache_data(ttl=300)
def _load_boreholes_by_zone(zone_code: str) -> list[dict]:
    with _db() as con:
        rows = con.execute("""
            SELECT b.name, b.elevation_m, b.depth_m
            FROM boreholes b JOIN zones z ON b.zone_id=z.id
            WHERE z.code=? ORDER BY b.name
        """, (zone_code,)).fetchall()
    return [dict(r) for r in rows]


@st.cache_data(ttl=300)
def _load_layers(bh_name: str) -> list[dict]:
    with _db() as con:
        rows = con.execute("""
            SELECT l.symbol, l.description, l.depth_top_m, l.depth_bot_m, l.thickness_m,
                   ROUND(AVG(lt.gamma_kNm3), 2) AS gamma_kNm3,
                   ROUND(AVG(lt.c_kPa),      1) AS c_kPa,
                   ROUND(AVG(lt.phi_deg),     1) AS phi_deg
            FROM layers l
            JOIN boreholes b ON l.borehole_id = b.id
            LEFT JOIN lab_tests lt
                   ON lt.borehole_id = b.id
                  AND lt.depth_from_m < l.depth_bot_m
                  AND lt.depth_to_m   > l.depth_top_m
            WHERE b.name = ?
            GROUP BY l.id
            ORDER BY l.depth_top_m
        """, (bh_name,)).fetchall()
    return [dict(r) for r in rows]


@st.cache_data(ttl=300)
def _load_vst_su(zone_code: str) -> pd.DataFrame:
    """VST Su toàn zone, dùng để vẽ profile và tính trung bình."""
    with _db() as con:
        rows = con.execute("""
            SELECT vt.depth_m, vt.Su_kPa, vl.name loc_name
            FROM vane_shear_tests vt
            JOIN vst_locations vl ON vt.vst_loc_id=vl.id
            JOIN zones z ON vl.zone_id=z.id
            WHERE z.code=? AND vt.Su_kPa IS NOT NULL
            ORDER BY vt.depth_m
        """, (zone_code,)).fetchall()
    return pd.DataFrame([dict(r) for r in rows])


@st.cache_data(ttl=300)
def _load_lab(bh_name: str) -> pd.DataFrame:
    with _db() as con:
        rows = con.execute("""
            SELECT depth_from_m, depth_to_m, gamma_kNm3, e0,
                   Cu_UU_kPa, c_kPa, phi_deg, Cc, Cs, E_kPa
            FROM lab_tests WHERE borehole_id=(
                SELECT id FROM boreholes WHERE name=?
            ) ORDER BY depth_from_m
        """, (bh_name,)).fetchall()
    return pd.DataFrame([dict(r) for r in rows])


@st.cache_data(ttl=300)
def _load_cdm_tests() -> pd.DataFrame:
    with _db() as con:
        rows = con.execute("""
            SELECT ct.*, b.name bh_name
            FROM cdm_tests ct JOIN boreholes b ON ct.borehole_id=b.id
            ORDER BY ct.cement_type, ct.WC_ratio, ct.dosage_kgm3
        """).fetchall()
    return pd.DataFrame([dict(r) for r in rows])


def _clay_params(bh_name: str, zone_code: str) -> dict:
    """Trả về h_clay, elevation, depth_clay_bot từ layers."""
    layers = _load_layers(bh_name)
    syms = _CLAY_SYMBOLS.get(zone_code, ["1", "2"])
    clay = [l for l in layers if l["symbol"] in syms]
    if not clay:
        return {}
    h_clay = sum(l["thickness_m"] or 0 for l in clay)
    depth_bot = max(l["depth_bot_m"] for l in clay)
    with _db() as con:
        row = con.execute("SELECT elevation_m FROM boreholes WHERE name=?", (bh_name,)).fetchone()
    elev = row["elevation_m"] if row and row["elevation_m"] else 0.0
    return {"h_clay": round(h_clay, 2), "depth_clay_bot": round(depth_bot, 2),
            "elev_clay_bot": round(elev - depth_bot, 2), "elevation_m": elev}


def _su_avg_in_range(zone_code: str, depth_top: float, depth_bot: float) -> float | None:
    df = _load_vst_su(zone_code)
    if df.empty:
        return None
    sub = df[(df.depth_m >= depth_top) & (df.depth_m <= depth_bot)]
    return round(float(sub.Su_kPa.mean()), 2) if not sub.empty else None


def _gamma_avg(bh_name: str) -> float:
    df = _load_lab(bh_name)
    if df.empty or df.gamma_kNm3.dropna().empty:
        return 15.0
    return round(float(df.gamma_kNm3.dropna().mean()), 3)


# ═══════════════════════════════════════════════════════════════════════════════
# 3D BOREHOLE MAP
# ═══════════════════════════════════════════════════════════════════════════════
@st.cache_data(ttl=300, show_spinner=False)
def _load_borehole_3d_data() -> tuple[list[dict], list[dict]]:
    if not _DB.exists():
        return [], []
    conn = sqlite3.connect(_DB)
    conn.row_factory = sqlite3.Row
    bhs = [dict(r) for r in conn.execute(
        "SELECT b.id, b.name, z.code zone, b.elevation_m, b.depth_m, "
        "b.x_coord_m, b.y_coord_m "
        "FROM boreholes b JOIN zones z ON z.id=b.zone_id "
        "WHERE b.x_coord_m IS NOT NULL "
        "ORDER BY z.id, b.name"
    ).fetchall()]
    ids = [b["id"] for b in bhs]
    if not ids:
        conn.close()
        return [], []
    ph = ",".join("?" * len(ids))
    lays = [dict(r) for r in conn.execute(
        f"SELECT borehole_id, symbol, description, depth_top_m, depth_bot_m "
        f"FROM layers WHERE borehole_id IN ({ph}) ORDER BY borehole_id, depth_top_m",
        ids,
    ).fetchall()]
    conn.close()
    return bhs, lays


@st.cache_data(ttl=300, show_spinner=False)
def _load_clay_bottom_3d() -> list[dict]:
    """Đáy lớp bùn sét yếu cho mỗi hố khoan có tọa độ."""
    if not _DB.exists():
        return []
    conn = sqlite3.connect(_DB)
    conn.row_factory = sqlite3.Row
    rows = [dict(r) for r in conn.execute("""
        SELECT b.name, z.code zone, b.elevation_m, b.x_coord_m, b.y_coord_m,
               ROUND(MAX(l.depth_bot_m), 3) clay_bot_depth,
               ROUND(b.elevation_m - MAX(l.depth_bot_m), 3) clay_bot_elev
        FROM boreholes b
        JOIN zones z ON z.id=b.zone_id
        JOIN layers l ON l.borehole_id=b.id
        WHERE b.x_coord_m IS NOT NULL
          AND ((z.code='KE'  AND l.symbol IN ('1','1b'))
            OR (z.code IN ('BXN','NHC') AND l.symbol='2'))
        GROUP BY b.id
        ORDER BY z.id, b.name
    """).fetchall()]
    conn.close()
    return rows


def _draw_boreholes_3d(
    selected_zones: list[str],
    show_clay_bottom: bool = False,
    show_cdm_top: bool = False,
    cdm_top_z: float = 2.7,
    focus_bh: str | None = None,
) -> "go.Figure":
    from collections import defaultdict

    bhs, lays = _load_borehole_3d_data()
    bhs  = [b for b in bhs if b["zone"] in selected_zones]
    if focus_bh:
        bhs_focus = [b for b in bhs if b["name"] == focus_bh]
        if bhs_focus:
            bhs = bhs_focus
    bh_ids = {b["id"] for b in bhs}
    lays = [l for l in lays if l["borehole_id"] in bh_ids]

    fig = go.Figure()
    if not bhs:
        fig.update_layout(title="Không có dữ liệu tọa độ hố khoan")
        return fig

    lay_by_bh: dict[int, list[dict]] = defaultdict(list)
    for l in lays:
        lay_by_bh[l["borehole_id"]].append(l)

    # Một trace / symbol → gộp legend
    all_syms: list[str] = []
    for l in lays:
        if l["symbol"] not in all_syms:
            all_syms.append(l["symbol"])

    added_legends: set[str] = set()
    for sym in all_syms:
        color = _LAYER_COLORS.get(sym, _LAYER_DEFAULT_COLOR)
        xs, ys, zs, texts = [], [], [], []
        for bh in bhs:
            for l in lay_by_bh[bh["id"]]:
                if l["symbol"] != sym:
                    continue
                z0 = bh["elevation_m"] - l["depth_top_m"]
                z1 = bh["elevation_m"] - l["depth_bot_m"]
                desc = l.get("description") or sym
                hover = (
                    f"<b>{bh['name']}</b><br>"
                    f"Lớp {sym}: {desc}<br>"
                    f"Sâu {l['depth_top_m']:.1f}–{l['depth_bot_m']:.1f} m<br>"
                    f"Cao độ {z0:.2f} → {z1:.2f} m"
                )
                xs += [bh["x_coord_m"], bh["x_coord_m"], None]
                ys += [bh["y_coord_m"], bh["y_coord_m"], None]
                zs += [z0, z1, None]
                texts += [hover, hover, None]

        show_leg = sym not in added_legends
        added_legends.add(sym)
        fig.add_trace(go.Scatter3d(
            x=xs, y=ys, z=zs, mode="lines",
            line=dict(color=color, width=10),
            name=f"Lớp {sym}", legendgroup=sym, showlegend=show_leg,
            hovertemplate="%{text}<extra></extra>", text=texts,
        ))

    # Labels tên hố khoan
    _zone_colors = {"KE": "#E53935", "BXN": "#1565C0", "NHC": "#2E7D32"}
    fig.add_trace(go.Scatter3d(
        x=[b["x_coord_m"] for b in bhs],
        y=[b["y_coord_m"] for b in bhs],
        z=[b["elevation_m"] + 3 for b in bhs],
        mode="text+markers",
        text=[b["name"].replace("BXN-CV-", "").replace("KE-", "") for b in bhs],
        textposition="top center",
        textfont=dict(size=9, color="#212121"),
        marker=dict(
            size=5,
            color=[_zone_colors.get(b["zone"], "#333") for b in bhs],
            symbol=[_ZONE_MARKER.get(b["zone"], "circle") for b in bhs],
        ),
        name="Vị trí HK", showlegend=True,
        hovertemplate="<b>%{text}</b><extra></extra>",
    ))

    # Đáy lớp bùn
    if show_clay_bottom:
        _clay = [c for c in _load_clay_bottom_3d() if c["zone"] in selected_zones]
        if len(_clay) >= 3:
            cx = [c["x_coord_m"] for c in _clay]
            cy = [c["y_coord_m"] for c in _clay]
            cz = [c["clay_bot_elev"] for c in _clay]
            ct = [
                f"<b>{c['name']}</b><br>"
                f"Đáy bùn sâu: {c['clay_bot_depth']:.1f} m<br>"
                f"Cao độ đáy: {c['clay_bot_elev']:.2f} m"
                for c in _clay
            ]
            fig.add_trace(go.Mesh3d(
                x=cx, y=cy, z=cz, delaunayaxis="z",
                intensity=cz,
                colorscale=[[0, "#0D47A1"], [0.5, "#42A5F5"], [1, "#BBDEFB"]],
                showscale=False, opacity=0.55,
                name="Đáy lớp bùn", legendgroup="clay_bot", showlegend=True,
                hovertemplate="%{text}<extra></extra>", text=ct,
            ))
            fig.add_trace(go.Scatter3d(
                x=cx, y=cy, z=cz, mode="markers",
                marker=dict(size=6, color="#0D47A1", symbol="diamond"),
                name="Điểm đáy bùn", legendgroup="clay_bot", showlegend=False,
                hovertemplate="%{text}<extra></extra>", text=ct,
            ))
            vx, vy, vz, vt = [], [], [], []
            for c in _clay:
                vx += [c["x_coord_m"], c["x_coord_m"], None]
                vy += [c["y_coord_m"], c["y_coord_m"], None]
                vz += [c["elevation_m"], c["clay_bot_elev"], None]
                lbl = f"<b>{c['name']}</b><br>Bùn dày {c['clay_bot_depth']:.1f} m"
                vt += [lbl, lbl, None]
            fig.add_trace(go.Scatter3d(
                x=vx, y=vy, z=vz, mode="lines",
                line=dict(color="#1565C0", width=3, dash="dot"),
                name="Chiều dày bùn", legendgroup="clay_bot", showlegend=False,
                hovertemplate="%{text}<extra></extra>", text=vt,
            ))

    # Mặt phẳng đỉnh trụ CDM
    if show_cdm_top:
        all_x = [b["x_coord_m"] for b in bhs]
        all_y = [b["y_coord_m"] for b in bhs]
        pad = 20.0
        mx, Mx = min(all_x) - pad, max(all_x) + pad
        my, My = min(all_y) - pad, max(all_y) + pad
        fig.add_trace(go.Mesh3d(
            x=[mx, Mx, Mx, mx], y=[my, my, My, My],
            z=[cdm_top_z]*4, i=[0, 0], j=[1, 2], k=[2, 3],
            color="#E91E63", opacity=0.25,
            name=f"Đỉnh trụ CDM +{cdm_top_z:.2f} m",
            showlegend=True, flatshading=True,
            hovertemplate=f"Đỉnh trụ CDM: <b>{cdm_top_z:.2f} m</b><extra></extra>",
        ))
        fig.add_trace(go.Scatter3d(
            x=[mx, Mx, Mx, mx, mx], y=[my, my, My, My, my],
            z=[cdm_top_z]*5, mode="lines",
            line=dict(color="#C2185B", width=3),
            name="Viền đỉnh CDM", showlegend=False,
        ))

    # Layout
    all_x = [b["x_coord_m"] for b in bhs]
    all_y = [b["y_coord_m"] for b in bhs]
    span_x = max(all_x) - min(all_x) or 1
    span_y = max(all_y) - min(all_y) or 1
    max_z = max(b["elevation_m"] for b in bhs) + 5
    min_z = min(b["elevation_m"] - b["depth_m"] for b in bhs) - 2

    fig.update_layout(
        height=660,
        margin=dict(l=0, r=0, t=30, b=0),
        legend=dict(
            x=1.01, y=0.95,
            bgcolor="rgba(255,255,255,0.85)",
            bordercolor="#ccc", borderwidth=1,
            font=dict(size=11),
        ),
        scene=dict(
            xaxis=dict(title="X — Easting (m)", tickformat=".0f"),
            yaxis=dict(title="Y — Northing (m)", tickformat=".0f"),
            zaxis=dict(title="Cao độ (m)", range=[min_z, max_z + 5]),
            aspectmode="manual",
            aspectratio=dict(
                x=span_x / max(span_x, span_y),
                y=span_y / max(span_x, span_y),
                z=0.35,
            ),
            camera=dict(eye=dict(x=1.4, y=-1.4, z=0.8)),
        ),
        paper_bgcolor="#FAFAFA",
    )
    return fig


_VN2000_CM  = 105.75          # kinh tuyến trục TP.HCM: 105°45'E
_VN2000_FE  = 500_000.0       # False Easting
_M_PER_DEG  = 110_574.0       # m/độ vĩ tuyến (xích đạo ≈ 110,574 m)


def _vn2000_to_latlon(northing: float, easting: float) -> tuple[float, float]:
    """VN-2000 TP.HCM (CM=105°45', FE=500000) → WGS84 (lat, lon).
    Sai số <1 m trong phạm vi TP.HCM; đủ dùng cho overlay bản đồ.
    """
    lat = northing / _M_PER_DEG
    lon = _VN2000_CM + (easting - _VN2000_FE) / (_M_PER_DEG * math.cos(math.radians(lat)))
    return lat, lon


def _project_to_latlon(
    bhs: list[dict],
    ref_bh_name: str = "",
    ref_lat: float = 0.0,
    ref_lon: float = 0.0,
    use_gps_ref: bool = False,
) -> list[dict] | None:
    """Chuyển tọa độ VN-2000 → lat/lon.

    Mặc định dùng công thức trực tiếp CM=105°45'.
    Nếu use_gps_ref=True và có điểm tham chiếu GPS → hiệu chỉnh thêm offset.
    """
    result = []
    for b in bhs:
        if b.get("x_coord_m") is None:
            continue
        lat, lon = _vn2000_to_latlon(b["x_coord_m"], b["y_coord_m"])
        result.append({**b, "lat": lat, "lon": lon})

    if use_gps_ref and ref_bh_name and result:
        ref_calc = next((b for b in result if b["name"] == ref_bh_name), None)
        if ref_calc:
            dlat = ref_lat - ref_calc["lat"]
            dlon = ref_lon - ref_calc["lon"]
            result = [{**b, "lat": b["lat"] + dlat, "lon": b["lon"] + dlon}
                      for b in result]
    return result or None


_MAP_STYLES = {
    "Đường phố (OSM)": "open-street-map",
    "Bản đồ sáng": "carto-positron",
    "Vệ tinh (Esri)": "__esri__",
}


def _draw_map_2d(bhs_ll: list[dict], style_key: str) -> "go.Figure":
    """Bản đồ 2D vị trí hố khoan, hỗ trợ OSM / Esri satellite."""
    _zone_colors = {"KE": "#E53935", "BXN": "#1565C0", "NHC": "#2E7D32"}
    fig = go.Figure()

    for zone, color in _zone_colors.items():
        grp = [b for b in bhs_ll if b["zone"] == zone]
        if not grp:
            continue
        fig.add_trace(go.Scattermap(
            lat=[b["lat"] for b in grp],
            lon=[b["lon"] for b in grp],
            mode="markers+text",
            text=[b["name"].replace("BXN-CV-", "").replace("KE-", "") for b in grp],
            textposition="top right",
            textfont=dict(size=10),
            marker=dict(size=13, color=color),
            name=zone,
            hovertemplate=(
                "<b>%{text}</b><br>"
                "lat=%{lat:.6f}<br>lon=%{lon:.6f}"
                "<extra></extra>"
            ),
        ))

    c_lat = sum(b["lat"] for b in bhs_ll) / len(bhs_ll)
    c_lon = sum(b["lon"] for b in bhs_ll) / len(bhs_ll)

    if style_key == "__esri__":
        map_cfg = dict(
            style="white-bg",
            layers=[{
                "below": "traces",
                "sourcetype": "raster",
                "source": [
                    "https://server.arcgisonline.com/ArcGIS/rest/services/"
                    "World_Imagery/MapServer/tile/{z}/{y}/{x}"
                ],
                "sourceattribution": "Esri World Imagery",
            }],
            center=dict(lat=c_lat, lon=c_lon),
            zoom=17,
        )
    else:
        map_cfg = dict(
            style=style_key,
            center=dict(lat=c_lat, lon=c_lon),
            zoom=17,
        )

    fig.update_layout(
        map=map_cfg,
        height=600,
        margin=dict(l=0, r=0, t=10, b=0),
        legend=dict(
            orientation="h", x=0.01, y=0.01,
            bgcolor="rgba(255,255,255,0.85)",
            bordercolor="#ccc", borderwidth=1,
        ),
        paper_bgcolor="#FAFAFA",
    )
    return fig


# ═══════════════════════════════════════════════════════════════════════════════
# CDM CALCULATIONS
# ═══════════════════════════════════════════════════════════════════════════════
def calc_a(D: float, e: float, arrangement: str = "triangle") -> float:
    if arrangement == "triangle":
        return math.pi * D**2 / (2 * math.sqrt(3) * e**2)
    return math.pi * D**2 / (4 * e**2)  # square


def calc_Ec(qu_field: float) -> float:
    return 100 * (qu_field / 2)   # Ec = 100·Cc, Cc = qu/2


def calc_Es(Cu: float) -> float:
    return 250 * Cu


def calc_Etb(a: float, Ec: float, Es: float) -> float:
    return a * Ec + (1 - a) * Es


def calc_S1(q: float, Lc: float, Etb: float) -> float:
    """Độ lún S1 (cm)."""
    return q * Lc / Etb * 100


def calc_sigma_col(q: float, Ec: float, Etb: float) -> float:
    """Ứng suất tập trung lên đầu cọc (kN/m²)."""
    return (Ec / Etb) * q


def calc_Pcol(sigma_col: float, D: float) -> float:
    """Lực nén lên 1 trụ CDM (kN)."""
    return sigma_col * math.pi * D**2 / 4


def calc_bearing(D: float, Lc: float, Cc: float, Cu: float, FS: float = 2.0) -> dict:
    """Sức chịu tải trụ CDM theo phương pháp AIT (TCVN 9403:2012 Phụ lục B)."""
    Ac = math.pi * D**2 / 4
    Qult_col  = 9 * Cc * Ac                         # sức chịu ở mũi (kN)
    Qult_skin = math.pi * D * Lc * Cu               # ma sát thân (kN)
    Qult      = Qult_col + Qult_skin
    Qa        = Qult / FS
    return {"Qult": round(Qult, 1), "Qa": round(Qa, 1),
            "Qult_col": round(Qult_col, 1), "Qult_skin": round(Qult_skin, 1)}


def calc_punching_check(
    D: float, e: float, Hse: float, He: float,
    quckse: float, Fs: float, theta_deg: float,
    gamma_fill: float, gamma_mat: float, qa: float = 0.0,
) -> dict:
    """Kiểm tra chọc thủng lớp đệm xi măng – phương pháp ALiCC (PWRI Japan)."""
    tan_t2 = math.tan(math.radians(theta_deg / 2))   # tan(θ/2) cho Ho
    tan_t  = math.tan(math.radians(theta_deg))        # tan(θ) cho thể tích
    tase   = quckse / (2.0 * Fs)
    Ho     = (e - D) * tan_t2

    if Ho <= He:
        Vsoil = (((e - D) * 0.5 * e**2
                  - math.pi * (e**3 - D**3) / 24.0
                  + (4.0 - math.pi) * (math.sqrt(2) - 1.0) * e**3 / 24.0)
                 * tan_t)
    else:
        r0    = He / tan_t + D / 2.0
        Vsoil = (He * e**2
                 - (1.0/3.0) * (math.pi * r0**2 * (He + D/2.0 * tan_t)
                                - math.pi * D / 2.0 * tan_t))

    r_mat  = Hse / tan_t + D / 2.0
    VCGCXM = (Hse * e**2
              - (1.0/3.0) * (math.pi * r_mat**2 * (Hse + D/2.0 * tan_t)
                             - math.pi * D / 2.0 * tan_t))

    A_unit = e**2 - math.pi * D**2 / 4.0
    PSoil  = ((Vsoil - VCGCXM) * gamma_fill + VCGCXM * gamma_mat) / A_unit
    tse    = (PSoil - qa) * A_unit / (math.pi * D * Hse)

    return {
        "tase_kPa":  round(tase, 2),
        "Ho_m":      round(Ho, 3),
        "He_m":      round(He, 3),
        "cong_thuc": "CT(1) Ho≤He" if Ho <= He else "CT(2) Ho>He",
        "Vsoil_m3":  round(Vsoil, 4),
        "VCGCXM_m3": round(VCGCXM, 4),
        "PSoil_kPa": round(PSoil, 3),
        "tse_kPa":   round(tse, 3),
        "check_CT":  "Đạt" if tse <= tase else "Không đạt",
    }


def build_scenarios(
    D: float, Lc: float, qu_field: float, Cu: float,
    q: float, spacings: list[float], arrangement: str
) -> list[dict]:
    Ec = calc_Ec(qu_field)
    Es = calc_Es(Cu)
    Cc = qu_field / 2
    rows = []
    for e in spacings:
        a    = calc_a(D, e, arrangement)
        Etb  = calc_Etb(a, Ec, Es)
        S1   = calc_S1(q, Lc, Etb)
        sc   = calc_sigma_col(q, Ec, Etb)
        Pcol = calc_Pcol(sc, D)
        bc   = calc_bearing(D, Lc, Cc, Cu)
        rows.append({
            "e (m)": e, "a": round(a, 4), "a (%)": round(a*100, 1),
            "Ec (kN/m²)": int(Ec), "Es (kN/m²)": int(Es),
            "Etb (kN/m²)": round(Etb, 0),
            "S₁ (cm)": round(S1, 2),
            "σ_col (kN/m²)": round(sc, 1),
            "Pcol (kN)": round(Pcol, 1),
            "Qa (kN)": bc["Qa"],
            "Đạt SCT": "Đạt" if Pcol < bc["Qa"] else "Không đạt",
        })
    return rows


def q_total(loads: dict) -> float:
    return (loads.get("q_traffic", 20)
            + loads.get("h_road", 0.8) * loads.get("g_road", 24)
            + loads.get("h_fill", 1.5) * loads.get("g_fill", 18)
            + loads.get("h_mat",  0.4) * loads.get("g_mat",  22.5))


# ═══════════════════════════════════════════════════════════════════════════════
# SESSION STATE INIT
# ═══════════════════════════════════════════════════════════════════════════════
_DEFAULTS = {
    "cdm_zone": "KE",
    "cdm_bh": None,
    "cdm_h_clay": 24.8,
    "cdm_Su": 11.2,
    "cdm_gamma": 15.0,
    "cdm_elevation": 0.0,
    "cdm_D": 0.8,
    "cdm_Lc": 26.2,
    "cdm_CDTK": 0.8,
    "cdm_qu": 800.0,
    "cdm_FS_lab": 2.0,
    "cdm_arrangement": "triangle",
    "cdm_cement_type": "Hoàng Thạch PCB40",
    "cdm_dosage": 240,
    "cdm_WC": 1.0,
    "cdm_spacings": [1.4, 1.6, 1.8],
    "cdm_rec_idx": 1,
    "cdm_geo_view": "3D địa chất",
    "cdm_map_ref_bh": "",
    "cdm_map_ref_lat": 10.7770,
    "cdm_map_ref_lon": 106.6800,
    "cdm_map_style": "Đường phố (OSM)",
    "cdm_vst_sel": [],
    "cdm_quckse": 600.0,
    "cdm_Fs_mat": 3.0,
    "cdm_theta":  80.0,
    "cdm_qa_mat": 0.0,
    "cdm_loads": {
        "q_traffic": 20.0,
        "h_road": 0.8,   "g_road": 24.0,
        "h_fill": 1.5,   "g_fill": 18.0,
        "h_mat":  0.4,   "g_mat":  22.5,
    },
}
for _k, _v in _DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v


def _get(key):
    return st.session_state.get(key, _DEFAULTS.get(key))


# ═══════════════════════════════════════════════════════════════════════════════
# CHART HELPERS
# ═══════════════════════════════════════════════════════════════════════════════
def _chart_scenarios(df: pd.DataFrame, rec_idx: int) -> go.Figure:
    """3 biểu đồ con: a(%), Etb, S1."""
    x = [f"e={r['e (m)']}m" for _, r in df.iterrows()]
    colors = ["#ED7D31" if i == rec_idx else "#4472C4" for i in range(len(df))]

    fig = go.Figure()
    # S1
    fig.add_trace(go.Bar(name="S₁ (cm)", x=x, y=df["S₁ (cm)"],
                         marker_color=colors, text=df["S₁ (cm)"].apply(lambda v: f"{v:.2f}"),
                         textposition="outside"))
    fig.update_layout(title="Độ lún S₁ theo khoảng cách trụ",
                      yaxis_title="S₁ (cm)", height=340,
                      legend=dict(orientation="h"), margin=dict(t=40, b=20))
    return fig


def _chart_etb(df: pd.DataFrame, rec_idx: int) -> go.Figure:
    x = [f"e={r['e (m)']}m" for _, r in df.iterrows()]
    colors = ["#ED7D31" if i == rec_idx else "#5B9BD5" for i in range(len(df))]
    fig = go.Figure(go.Bar(name="Etb", x=x, y=df["Etb (kN/m²)"],
                            marker_color=colors,
                            text=df["Etb (kN/m²)"].apply(lambda v: f"{int(v):,}"),
                            textposition="outside"))
    fig.update_layout(title="Mô đun tương đương Etb", yaxis_title="kN/m²",
                      height=300, margin=dict(t=40, b=20))
    return fig


def _linreg(xs: list, ys: list) -> tuple[float, float]:
    """Hồi quy tuyến tính đơn giản, trả về (a, b) với y = a*x + b."""
    n = len(xs)
    if n < 2:
        return 0.0, (ys[0] if ys else 0.0)
    sx  = sum(xs);  sy  = sum(ys)
    sxy = sum(xi*yi for xi, yi in zip(xs, ys))
    sxx = sum(xi**2 for xi in xs)
    denom = n * sxx - sx * sx
    if abs(denom) < 1e-12:
        return 0.0, sy / n
    a = (n * sxy - sx * sy) / denom
    b = (sy - a * sx) / n
    return a, b


def _chart_su_profile(
    df_vst: pd.DataFrame,
    selected_locs: list[str] | None = None,
) -> go.Figure:
    if df_vst.empty:
        return go.Figure()

    all_locs = sorted(df_vst["loc_name"].unique())
    show_locs = selected_locs if selected_locs else all_locs

    _colors = [
        "#E53935","#1565C0","#2E7D32","#F57F17","#6A1B9A",
        "#00838F","#AD1457","#558B2F","#4527A0","#00695C",
    ]

    fig = go.Figure()
    reg_lines = []   # (loc, a, b, color, y_min, y_max)

    for i, loc in enumerate(all_locs):
        if loc not in show_locs:
            continue
        grp   = df_vst[df_vst["loc_name"] == loc].sort_values("depth_m")
        color = _colors[i % len(_colors)]
        depths = grp["depth_m"].tolist()
        sus    = grp["Su_kPa"].tolist()
        y_plot = [-d for d in depths]

        # Đường + markers + nhãn giá trị
        fig.add_trace(go.Scatter(
            x=sus, y=y_plot,
            mode="lines+markers+text",
            name=loc,
            line=dict(color=color, width=1.8),
            marker=dict(size=7, color=color),
            text=[f"{v:.1f}" for v in sus],
            textposition="middle right",
            textfont=dict(size=9, color=color),
            hovertemplate=(
                f"<b>{loc}</b><br>"
                "Sâu: %{customdata:.1f} m<br>"
                "Su: %{x:.1f} kPa<extra></extra>"
            ),
            customdata=depths,
        ))

        # Hồi quy tuyến tính: Su = a*depth + b
        if len(depths) >= 2:
            a, b = _linreg(depths, sus)
            reg_lines.append((loc, a, b, color, min(depths), max(depths)))

    # Vẽ đường hồi quy + annotation phương trình
    for loc, a, b, color, d_min, d_max in reg_lines:
        d_ext = (d_max - d_min) * 0.1
        dd = [d_min - d_ext, d_max + d_ext]
        su_reg = [a * d + b for d in dd]
        sign  = "+" if b >= 0 else "−"
        eq    = f"Su={a:+.2f}z{sign}{abs(b):.2f}".replace("+-", "−")
        fig.add_trace(go.Scatter(
            x=su_reg, y=[-d for d in dd],
            mode="lines",
            name=f"Hồi quy {loc}",
            line=dict(color=color, width=1.2, dash="dash"),
            hovertemplate=f"<b>Hồi quy {loc}</b><br>{eq}<extra></extra>",
            showlegend=False,
        ))
        # Nhãn phương trình ở cuối đường
        fig.add_annotation(
            x=su_reg[-1], y=-dd[-1],
            text=eq,
            showarrow=False,
            font=dict(size=9, color=color),
            xanchor="left",
            bgcolor="rgba(255,255,255,0.75)",
            bordercolor=color,
            borderwidth=1,
        )

    fig.update_layout(
        title="Biểu đồ Su – Cắt cánh (VST)" + (
            f" — {', '.join(show_locs)}" if selected_locs else " — Toàn bộ"
        ),
        xaxis_title="Su (kPa)",
        yaxis_title="Cao độ (m)",
        height=500,
        legend=dict(font_size=10, orientation="v", x=1.01, y=1),
        margin=dict(t=50, b=20, r=160),
    )
    return fig


def _chart_combined(df: pd.DataFrame, rec_idx: int) -> go.Figure:
    x = [f"e={r['e (m)']}m" for _, r in df.iterrows()]
    colors = ["#ED7D31" if i == rec_idx else "#4472C4" for i in range(len(df))]
    fig = go.Figure()
    fig.add_trace(go.Bar(name="a (%)", x=x, y=df["a (%)"],
                         marker_color=colors, yaxis="y1",
                         offsetgroup=0,
                         text=df["a (%)"].apply(lambda v: f"{v:.1f}%"),
                         textposition="outside"))
    fig.add_trace(go.Scatter(name="S₁ (cm)", x=x, y=df["S₁ (cm)"],
                              mode="lines+markers+text",
                              text=df["S₁ (cm)"].apply(lambda v: f"{v:.2f}"),
                              textposition="top center",
                              yaxis="y2", line=dict(color="#FF0000", width=2),
                              marker=dict(size=8, color="#FF0000")))
    fig.update_layout(
        title="Tổng hợp: tỷ lệ thay thế a(%) và độ lún S₁(cm)",
        yaxis=dict(title="a (%)", side="left"),
        yaxis2=dict(title="S₁ (cm)", side="right", overlaying="y"),
        barmode="group", height=340,
        legend=dict(orientation="h"), margin=dict(t=40, b=20),
    )
    return fig


def _draw_cdm_section(
    D: float, Lc: float, CDTK: float,
    h_clay: float, h_road: float, h_fill: float, h_mat: float,
    q_traffic: float,
):
    """Mặt cắt ngang CDM: các lớp đất, trụ CDM, mũi tên tải, kích thước D/Lc."""
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches

    z_mat      = CDTK + h_mat
    z_fill     = z_mat + h_fill
    z_road     = z_fill + h_road
    z_clay_bot = CDTK - h_clay
    z_pile_bot = CDTK - Lc
    z_base     = min(z_clay_bot, z_pile_bot) - 1.5

    W  = max(4.0 * D, 2.5)
    cx = W / 2
    r  = D / 2

    fig, ax = plt.subplots(figsize=(5, 8))
    fig.patch.set_facecolor("#FAFAFA")
    ax.set_facecolor("white")
    ax.set_xlim(0, W * 1.2)
    ax.set_ylim(z_base, z_road + 2.5)
    ax.set_ylabel("Cao độ (m)", fontsize=9)
    ax.set_xticks([])
    ax.tick_params(axis="y", labelsize=8)
    ax.set_title("Mặt cắt ngang CDM", fontsize=11, fontweight="bold", pad=8)
    for s in ["top", "right", "bottom"]:
        ax.spines[s].set_visible(False)

    def hrect(y0, y1, fc, ec="#666", zo=1):
        if y1 > y0:
            ax.add_patch(mpatches.Rectangle(
                (0, y0), W, y1 - y0, fc=fc, ec=ec, lw=0.7, zorder=zo
            ))

    hrect(z_base,     z_clay_bot, "#BCAAA4", "#795548")
    hrect(z_clay_bot, CDTK,       "#BBDEFB", "#1565C0")
    if h_mat > 0:
        hrect(CDTK,   z_mat,      "#FFA726", "#E65100")
    hrect(z_mat,      z_fill,     "#FFD54F", "#E6BE00")
    hrect(z_fill,     z_road,     "#90A4AE", "#546E7A")

    p_bot = max(z_pile_bot, z_base + 0.1)
    ax.add_patch(mpatches.Rectangle(
        (cx - r, p_bot), D, CDTK - p_bot, fc="#66BB6A", ec="#2E7D32", lw=1.2, zorder=3
    ))
    p_mid = (CDTK + p_bot) / 2
    ax.text(cx, p_mid, f"CDM\nD={D:.2f}m\nLc={Lc:.1f}m",
            ha="center", va="center", fontsize=8.5,
            color="white", fontweight="bold", zorder=4)

    # CDTK = đỉnh cọc + đệm + đắp (đỉnh lớp cát đắp, dưới mặt đường)
    ax.axhline(z_fill, color="#E53935", lw=1.5, ls="--", zorder=5)
    ax.text(W * 0.98, z_fill + 0.08, f"CDTK = {z_fill:+.2f} m",
            ha="right", va="bottom", fontsize=8.5, color="#E53935",
            fontweight="bold", zorder=7,
            bbox=dict(boxstyle="round,pad=0.2", fc="white", ec="#E53935",
                      alpha=0.9, lw=0.8))
    # Đỉnh cọc CDM (nét chấm xanh lá)
    ax.axhline(CDTK, color="#2E7D32", lw=0.9, ls=":", zorder=4, alpha=0.8)
    ax.text(W * 0.98, CDTK - 0.1, f"Đỉnh cọc = {CDTK:+.2f} m",
            ha="right", va="top", fontsize=7.5, color="#2E7D32", zorder=5,
            bbox=dict(boxstyle="round,pad=0.1", fc="white", ec="none", alpha=0.8))

    def layer_text(y0, y1, txt, color):
        if (y1 - y0) > 0.4:
            ax.text(W * 0.03, (y0 + y1) / 2, txt,
                    ha="left", va="center", fontsize=8, color=color, zorder=6)

    layer_text(z_base,     z_clay_bot, "Đất tốt",               "#5D4037")
    layer_text(z_clay_bot, CDTK,       f"Bùn sét  h={h_clay:.1f}m", "#0D47A1")
    if h_mat > 0:
        layer_text(CDTK,   z_mat,      f"Đệm cát  h={h_mat:.1f}m",  "#E65100")
    layer_text(z_mat,      z_fill,     f"Cát đắp  h={h_fill:.1f}m",  "#827717")
    layer_text(z_fill,     z_road,     f"Mặt đường h={h_road:.1f}m", "#37474F")

    dim_yD = (CDTK + z_mat) / 2 if h_mat > 0.2 else CDTK + 0.25
    ax.annotate("", xy=(cx + r, dim_yD), xytext=(cx - r, dim_yD),
                arrowprops=dict(arrowstyle="<->", color="#2E7D32", lw=1.5), zorder=5)
    ax.text(cx, dim_yD + 0.12, f"D={D:.2f}m",
            ha="center", va="bottom", fontsize=8, color="#2E7D32", zorder=5,
            bbox=dict(boxstyle="round,pad=0.1", fc="white", ec="none", alpha=0.8))

    dim_x = cx + r + W * 0.07
    ax.annotate("", xy=(dim_x, p_bot), xytext=(dim_x, CDTK),
                arrowprops=dict(arrowstyle="<->", color="#1A237E", lw=1.5), zorder=5)
    ax.text(dim_x + W * 0.02, p_mid, f"Lc={Lc:.1f}m",
            ha="left", va="center", fontsize=8, color="#1A237E", zorder=5)

    q_top = z_road + 1.0
    for xi in [W * 0.25, W * 0.5, W * 0.75]:
        ax.annotate("", xy=(xi, z_road), xytext=(xi, q_top),
                    arrowprops=dict(arrowstyle="-|>", color="#B71C1C", lw=1.8), zorder=6)
    ax.text(cx, q_top + 0.2, f"q = {q_traffic:.0f} kN/m²",
            ha="center", va="bottom", fontsize=10,
            color="#B71C1C", fontweight="bold", zorder=6)

    plt.tight_layout(pad=0.5)
    return fig


def _draw_cdm_grid(D: float, arr: str, e_ref: float):
    """Sơ đồ bố trí lưới cọc CDM (plan view)."""
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches

    r = D / 2
    centers = []
    for ri in range(4):
        for ci in range(4):
            if arr == "triangle":
                px = ci * e_ref + (e_ref / 2 if ri % 2 else 0)
                py = ri * (e_ref * (3 ** 0.5) / 2)
            else:
                px = ci * e_ref
                py = ri * e_ref
            centers.append((px, py))

    all_px = [p[0] for p in centers]
    all_py = [p[1] for p in centers]
    pad = e_ref * 0.75

    fig, ax = plt.subplots(figsize=(5, 4.5))
    fig.patch.set_facecolor("#FAFAFA")
    ax.set_facecolor("#F0F4F8")
    ax.set_aspect("equal")
    ax.set_xticks([])
    ax.set_yticks([])
    arr_lbl = "Tam giác" if arr == "triangle" else "Hình vuông"
    ax.set_title(f"Sơ đồ lưới {arr_lbl}  (e = {e_ref:.2f} m)",
                 fontsize=11, fontweight="bold", pad=8)
    for s in ax.spines.values():
        s.set_visible(False)

    if arr == "square":
        ax.add_patch(mpatches.Rectangle(
            (-e_ref / 2, -e_ref / 2), e_ref, e_ref,
            fc="#BBDEFB", ec="#1565C0", lw=1.5, ls="--", alpha=0.4, zorder=1
        ))
    else:
        e_h = e_ref * (3 ** 0.5) / 2
        ax.add_patch(mpatches.Polygon(
            [(0, 0), (e_ref, 0), (e_ref / 2, e_h)],
            closed=True, fc="#BBDEFB", ec="#1565C0", lw=1.5, ls="--", alpha=0.35, zorder=1
        ))

    for px, py in centers:
        ax.add_patch(mpatches.Circle(
            (px, py), r, fc="#A5D6A7", ec="#2E7D32", lw=1.2, zorder=3
        ))

    p0x, p0y = centers[0]
    p1x, _   = centers[1]
    d_y = p0y - r - 0.2
    ax.annotate("", xy=(p1x, d_y), xytext=(p0x, d_y),
                arrowprops=dict(arrowstyle="<->", color="#E53935", lw=1.5))
    ax.text((p0x + p1x) / 2, d_y - 0.12, f"e = {e_ref:.2f} m",
            ha="center", va="top", fontsize=8.5, color="#E53935")

    a_val = calc_a(D, e_ref, arr)
    ax.text((min(all_px) + max(all_px)) / 2, max(all_py) + r + 0.3,
            f"a = {a_val * 100:.1f}%",
            ha="center", va="bottom", fontsize=11, color="#1A237E", fontweight="bold")

    ax.set_xlim(min(all_px) - pad, max(all_px) + pad)
    ax.set_ylim(min(all_py) - pad * 1.8, max(all_py) + pad * 0.8)
    plt.tight_layout(pad=0.5)
    return fig


def _draw_soil_column(layers: list[dict], bh_name: str = "") -> go.Figure:
    """Cột địa chất dạng thanh màu theo ký hiệu lớp đất (chiều sâu từ trên xuống)."""
    if not layers:
        fig = go.Figure()
        fig.update_layout(height=500,
                          title=dict(text=f"Cột ĐC: {bh_name}", font=dict(size=10), x=0.5))
        return fig

    def _tc(hex_color: str) -> str:
        h = hex_color.lstrip("#")
        if len(h) < 6:
            return "#212121"
        r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
        return "white" if (0.299 * r + 0.587 * g + 0.114 * b) < 140 else "#212121"

    shapes: list[dict] = []
    annotations: list[dict] = []
    seen: set[float] = set()

    for lay in layers:
        y0    = lay["depth_top_m"]
        y1    = lay["depth_bot_m"]
        sym   = (lay.get("symbol") or "?").strip()
        color = _LAYER_COLORS.get(sym, _LAYER_DEFAULT_COLOR)
        thick = y1 - y0

        shapes.append(dict(
            type="rect", x0=0, y0=y0, x1=1, y1=y1,
            fillcolor=color, line=dict(color="#555", width=0.6),
        ))

        if y0 not in seen:
            annotations.append(dict(
                x=1.08, y=y0, text=f"{y0:.1f}",
                showarrow=False, font=dict(size=8, color="#555"),
                xanchor="left", yanchor="middle",
            ))
            seen.add(y0)

        if thick >= 0.5:
            mid  = (y0 + y1) / 2
            tc   = _tc(color)
            desc = (lay.get("description") or "")
            lbl  = (f"<b>{sym}</b><br>{desc[:20]}"
                    if thick >= 2.5 and desc else f"<b>{sym}</b>")
            annotations.append(dict(
                x=0.5, y=mid, text=lbl,
                showarrow=False, font=dict(size=9, color=tc),
                xanchor="center", yanchor="middle",
            ))

    y_bot = layers[-1]["depth_bot_m"]
    if y_bot not in seen:
        annotations.append(dict(
            x=1.08, y=y_bot, text=f"{y_bot:.1f}",
            showarrow=False, font=dict(size=8, color="#555"),
            xanchor="left", yanchor="middle",
        ))

    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=[0.5], y=[y_bot / 2],
        mode="markers", marker=dict(opacity=0, size=1),
        showlegend=False,
    ))
    fig.update_layout(
        shapes=shapes,
        annotations=annotations,
        title=dict(text=f"Cột ĐC: {bh_name}", font=dict(size=10, color="#1F4E79"), x=0.5),
        height=500,
        margin=dict(l=10, r=55, t=35, b=20),
        xaxis=dict(showticklabels=False, showgrid=False, zeroline=False, range=[-0.05, 1.7]),
        yaxis=dict(
            range=[y_bot + 1, -0.5],
            title="Độ sâu (m)", title_font=dict(size=9),
            tickfont=dict(size=8),
            showgrid=True, gridcolor="#EEE", gridwidth=0.5,
            dtick=5,
        ),
        plot_bgcolor="white",
        paper_bgcolor="#FAFAFA",
        showlegend=False,
    )
    return fig


# ═══════════════════════════════════════════════════════════════════════════════
# EXPORT HELPERS
# ═══════════════════════════════════════════════════════════════════════════════
def _export_excel(scenarios: list[dict], params: dict) -> bytes:
    """Xuất Excel 3 sheet: Đầu vào | So sánh PA | Kết quả."""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.chart import BarChart, Reference
    except ImportError:
        return b""

    wb = openpyxl.Workbook()

    # --- Helpers style ---
    HDR_FILL = PatternFill("solid", fgColor="1F4E79")
    REC_FILL = PatternFill("solid", fgColor="E2EFDA")
    HDR_FONT = Font(bold=True, color="FFFFFF", size=10)
    BD = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    def _hdr(ws, row, col, val, fill=HDR_FILL, font=HDR_FONT):
        c = ws.cell(row=row, column=col, value=val)
        c.fill = fill; c.font = font; c.border = BD
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        return c

    def _val(ws, row, col, val, rec=False, fmt=None):
        c = ws.cell(row=row, column=col, value=val)
        if rec: c.fill = REC_FILL
        c.border = BD
        c.alignment = Alignment(horizontal="center", vertical="center")
        if fmt: c.number_format = fmt
        return c

    # ── Sheet 1: Đầu vào ──────────────────────────────────────────────────────
    ws1 = wb.active
    ws1.title = "Đầu vào"
    ws1.column_dimensions["A"].width = 36
    ws1.column_dimensions["B"].width = 20
    ws1.column_dimensions["C"].width = 16

    title_rows = [
        ("THÔNG SỐ ĐỊA KỸ THUẬT VÀ THIẾT KẾ CDM",),
        (f"Hố khoan: {params.get('bh_name','–')}   |   Zone: {params.get('zone','–')}",),
    ]
    for i, (txt,) in enumerate(title_rows, 1):
        c = ws1.cell(row=i, column=1, value=txt)
        c.font = Font(bold=True, size=11)
        ws1.merge_cells(start_row=i, start_column=1, end_row=i, end_column=3)

    _hdr(ws1, 4, 1, "Thông số"); _hdr(ws1, 4, 2, "Ký hiệu"); _hdr(ws1, 4, 3, "Giá trị")
    geo_rows = [
        ("Hố khoan tham chiếu", "HK", params.get("bh_name","–")),
        ("Bề dày lớp bùn sét", "h₁ (m)", params.get("h_clay", 0)),
        ("Sức kháng cắt VST trung bình", "Su (kN/m²)", params.get("Su", 0)),
        ("Dung trọng tự nhiên TB", "γ (kN/m³)", params.get("gamma", 0)),
        ("Đường kính trụ CDM", "D (mm)", int(params.get("D", 0.8)*1000)),
        ("Chiều dài trụ CDM", "Lc (m)", params.get("Lc", 0)),
        ("Cao độ đỉnh cọc CDM", "Đỉnh cọc (m)", params.get("CDTK", 2.7)),
        ("Bố trí", "–", "Tam giác" if params.get("arrangement","triangle")=="triangle" else "Vuông"),
        ("Cường độ TK hiện trường", "qu,tk (kPa)", params.get("qu", 800)),
        ("Hệ số quy đổi TN→HT", "FS", params.get("FS_lab", 2.0)),
        ("Cc = qu,tk/2", "Cc (kN/m²)", params.get("qu", 800)/2),
        ("Ec = 100·Cc", "Ec (kN/m²)", int(100*params.get("qu",800)/2)),
        ("Es = 250·Su", "Es (kN/m²)", int(250*params.get("Su",11.2))),
        ("Tổng tải trọng gây lún", "q (kN/m²)", params.get("q_total", 75.2)),
    ]
    for r, (n, k, v) in enumerate(geo_rows, 5):
        ws1.cell(row=r, column=1, value=n).border = BD
        ws1.cell(row=r, column=2, value=k).border = BD
        c = ws1.cell(row=r, column=3, value=v); c.border = BD
        c.alignment = Alignment(horizontal="center")

    # ── Sheet 2: So sánh PA ───────────────────────────────────────────────────
    ws2 = wb.create_sheet("So sánh PA")
    hdrs = ["PA", "e (m)", "a", "a (%)", "Ec (kN/m²)", "Es (kN/m²)",
            "Etb (kN/m²)", "S₁ (cm)", "Pcol (kN)", "Qa (kN)", "Đạt SCT"]
    for j, h in enumerate(hdrs, 1):
        _hdr(ws2, 2, j, h)
        ws2.column_dimensions[chr(64+j)].width = 13
    rec_idx = params.get("rec_idx", 0)
    for i, s in enumerate(scenarios):
        is_rec = (i == rec_idx)
        _val(ws2, i+3, 1, f"PA{i+1}", is_rec)
        _val(ws2, i+3, 2, s["e (m)"], is_rec)
        _val(ws2, i+3, 3, s["a"], is_rec, "0.0000")
        _val(ws2, i+3, 4, s["a (%)"], is_rec)
        _val(ws2, i+3, 5, s["Ec (kN/m²)"], is_rec, "#,##0")
        _val(ws2, i+3, 6, s["Es (kN/m²)"], is_rec, "#,##0")
        _val(ws2, i+3, 7, s["Etb (kN/m²)"], is_rec, "#,##0")
        _val(ws2, i+3, 8, s["S₁ (cm)"], is_rec, "0.00")
        _val(ws2, i+3, 9, s["Pcol (kN)"], is_rec, "0.0")
        _val(ws2, i+3, 10, s["Qa (kN)"], is_rec, "0.0")
        _val(ws2, i+3, 11, s["Đạt SCT"], is_rec)

    # Biểu đồ S1
    chart = BarChart()
    chart.title = "Độ lún S₁ theo khoảng cách trụ"
    chart.y_axis.title = "S₁ (cm)"
    chart.x_axis.title = "Khoảng cách e (m)"
    data_ref  = Reference(ws2, min_col=8, min_row=2, max_row=2+len(scenarios))
    cats_ref  = Reference(ws2, min_col=2, min_row=3, max_row=2+len(scenarios))
    chart.add_data(data_ref, titles_from_data=True)
    chart.set_categories(cats_ref)
    chart.shape = 4; chart.width = 14; chart.height = 10
    ws2.add_chart(chart, f"M2")

    # ── Sheet 3: Kết quả chi tiết PA kiến nghị ────────────────────────────────
    ws3 = wb.create_sheet("Kết quả PA KN")
    if scenarios and 0 <= rec_idx < len(scenarios):
        s = scenarios[rec_idx]
        ws3.column_dimensions["A"].width = 36
        ws3.column_dimensions["B"].width = 22
        ws3.column_dimensions["C"].width = 18
        c = ws3.cell(row=1, column=1,
                     value=f"TÍNH TOÁN CHI TIẾT – PA{rec_idx+1}: e = {s['e (m)']} m")
        c.font = Font(bold=True, size=12)
        ws3.merge_cells("A1:C1")

        _hdr(ws3, 2, 1, "Thông số"); _hdr(ws3, 2, 2, "Công thức"); _hdr(ws3, 2, 3, "Giá trị")
        detail_rows = [
            ("Khoảng cách bố trí", "e", f"{s['e (m)']} m"),
            ("Tỷ lệ thay thế", "a = π·D²/(2√3·e²)" if params.get("arrangement","triangle")=="triangle"
             else "a = π·D²/(4·e²)", f"{s['a']:.4f}  ({s['a (%)']:.1f}%)"),
            ("Mô đun trụ CDM", "Ec = 100·Cc", f"{s['Ec (kN/m²)']:,} kN/m²"),
            ("Mô đun đất nền", "Es = 250·Su", f"{s['Es (kN/m²)']:,} kN/m²"),
            ("Mô đun tương đương", "Etb = a·Ec+(1-a)·Es", f"{int(s['Etb (kN/m²)']):,} kN/m²"),
            ("Tải trọng thiết kế", "q", f"{params.get('q_total',0):.1f} kN/m²"),
            ("Chiều dài trụ", "Lc", f"{params.get('Lc',0):.1f} m"),
            ("Độ lún S₁", "q·Lc/Etb×100", f"{s['S₁ (cm)']:.2f} cm"),
            ("Ứng suất đầu cọc", "σ_col = Ec/Etb·q", f"{s['σ_col (kN/m²)']:.1f} kN/m²"),
            ("Lực nén lên trụ", "Pcol = σ_col·Ac", f"{s['Pcol (kN)']:.1f} kN"),
            ("Sức chịu tải cho phép", "Qa (AIT, FS=2)", f"{s['Qa (kN)']:.1f} kN"),
            ("Kết luận SCT", "Pcol < Qa ?", s["Đạt SCT"]),
        ]
        for r, (n, f, v) in enumerate(detail_rows, 3):
            ws3.cell(row=r, column=1, value=n).border = BD
            ws3.cell(row=r, column=2, value=f).border = BD
            c = ws3.cell(row=r, column=3, value=v); c.border = BD
            c.alignment = Alignment(horizontal="center")
            if v in ("Đạt", "Không đạt"):
                c.font = Font(color="375623" if v == "Đạt" else "C00000", bold=True)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _export_word_bytes(scenarios: list[dict], params: dict, rec_idx: int) -> bytes:
    """Tạo Word thuyết minh CDM và trả về bytes."""
    sys.path.insert(0, str(_CDM_SC))
    try:
        import docx_helpers as H
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import math as _m
    except ImportError as e:
        return b""

    H.reset_counters()
    doc = Document()
    H.setup_page(doc)

    D    = params["D"]
    Lc   = params["Lc"]
    qu   = params["qu"]
    Cu   = params["Su"]
    q    = params["q_total"]
    Cc   = qu / 2
    Ec   = 100 * Cc
    Es   = 250 * Cu
    arr  = "tam giác" if params.get("arrangement", "triangle") == "triangle" else "hình vuông"
    rec  = scenarios[rec_idx] if scenarios else {}

    # Trang bìa
    for _ in range(4): doc.add_paragraph()
    H.heading(doc, "THUYẾT MINH TÍNH TOÁN",
              size=16, bold=True, center=True, space_before=0)
    H.heading(doc, "GIA CỐ NỀN ĐẤT YẾU BẰNG CỌC ĐẤT XI MĂNG (CDM)",
              size=15, bold=True, center=True, space_before=4)
    for _ in range(2): doc.add_paragraph()
    H.heading(doc,
              f"D = {int(D*1000)} mm – e = {' / '.join(str(s['e (m)']) for s in scenarios)} m"
              f" – Lc = {Lc:.1f} m – qu = {int(qu)} kPa",
              size=13, bold=False, center=True, space_before=0)
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    H._fmt(p.add_run("TP. Hồ Chí Minh, tháng 5 năm 2026"), size=12)
    doc.add_page_break()

    H.add_toc(doc)

    # I. Cơ sở pháp lý
    H.heading(doc, "I. CƠ SỞ PHÁP LÝ VÀ TIÊU CHUẨN ÁP DỤNG", size=13)
    for item in [
        "– TCVN 9403:2012 – Gia cố nền đất yếu – Phương pháp trụ đất xi măng;",
        "– TCVN 4200:2012 – Đất xây dựng – Phương pháp xác định tính nén lún;",
        "– TCVN 8868:2011 – Thí nghiệm xác định sức kháng cắt không cố kết – không thoát nước;",
        f"– Kết quả khảo sát địa kỹ thuật (hố khoan {params.get('bh_name','HK1')});",
        "– Kết quả thí nghiệm thành phần cấp phối xi măng đất.",
    ]:
        H.body(doc, item, indent=0.5)

    # II. Địa chất
    H.heading(doc, "II. ĐẶC ĐIỂM ĐỊA CHẤT KHU VỰC", size=13)
    H.body(doc, f"Căn cứ kết quả khảo sát hố khoan {params.get('bh_name','–')}, "
           f"zone {params.get('zone','–')}:")
    H.tbl_caption(doc, f"Thông số địa kỹ thuật – {params.get('bh_name','–')}")
    H.make_table(doc,
        headers=["Thông số", "Ký hiệu", "Giá trị"],
        rows=[
            ("Bề dày lớp bùn sét (lớp cần gia cố)", "h₁ (m)", f"{params.get('h_clay',0):.2f}"),
            ("Sức kháng cắt không thoát nước (VST)", "Su (kN/m²)", f"{Cu:.2f}"),
            ("Dung trọng tự nhiên trung bình", "γ (kN/m³)", f"{params.get('gamma',15):.2f}"),
            ("Mô đun biến dạng đất nền", "Es = 250·Su (kN/m²)", f"{int(Es):,}"),
        ])

    # III. Thông số thiết kế
    H.heading(doc, "III. THÔNG SỐ THIẾT KẾ CỌC ĐẤT XI MĂNG", size=13)
    H.tbl_caption(doc, "Thông số hình học và vật liệu trụ CDM")
    H.make_table(doc,
        headers=["Thông số", "Ký hiệu", "Giá trị"],
        rows=[
            ("Đường kính trụ CDM",              "D",          f"{int(D*1000)} mm"),
            (f"Khoảng cách tâm-tâm (lưới {arr})", "e",        " / ".join(f"{s['e (m)']}m" for s in scenarios)),
            ("Chiều dài trụ CDM",               "Lc",         f"{Lc:.1f} m"),
            ("Hàm lượng xi măng",               "x",          f"{params.get('dosage',240)} kg/m³"),
            ("Cường độ nén TK hiện trường",      "qu,tk",      f"{int(qu)} kPa"),
            ("Cường độ kháng cắt thiết kế",      "Cc = qu/2",  f"{int(Cc)} kN/m²"),
            ("Mô đun đàn hồi trụ",              "Ec = 100·Cc", f"{int(Ec):,} kN/m²"),
            ("Tổng tải trọng gây lún",           "q",          f"{q:.1f} kN/m²"),
        ])

    # IV. Phương pháp tính
    H.heading(doc, "IV. PHƯƠNG PHÁP TÍNH TOÁN LÚN (TCVN 9403:2012 – Phụ lục C)", size=13)
    H.body(doc, "Độ lún bản thân khối CDM:")
    H.formula(doc, "S₁ = q × H / Etb    (C.2)")
    H.formula(doc, "Etb = a × Ec + (1 − a) × Es")
    form = ("a = π × D² / (2√3 × e²)  [lưới tam giác]"
            if params.get("arrangement","triangle")=="triangle"
            else "a = π × D² / (4 × e²)  [lưới vuông]")
    H.formula(doc, form)
    H.body(doc, "S₂ = 0 (trụ CDM xuyên qua toàn bộ lớp bùn sét yếu).")

    # V. So sánh phương án
    H.heading(doc, "V. SO SÁNH CÁC PHƯƠNG ÁN KHOẢNG CÁCH TRỤ", size=13)
    H.tbl_caption(doc, f"So sánh S₁ theo khoảng cách trụ (D={int(D*1000)}mm, qu={int(qu)}kPa, q={q:.1f}kN/m²)")
    H.make_table(doc,
        headers=["PA", "e (m)", "a", "Etb (kN/m²)", "S₁ (cm)", "Đạt SCT"],
        rows=[(f"PA{i+1}", str(s["e (m)"]), f"{s['a']:.4f}",
               f"{int(s['Etb (kN/m²)']):,}", f"{s['S₁ (cm)']:.2f}", s["Đạt SCT"])
              for i, s in enumerate(scenarios)],
        highlight_rows={rec_idx})

    # VI. Kết quả PA kiến nghị
    if rec:
        H.heading(doc,
                  f"VI. KẾT QUẢ – PHƯƠNG ÁN KIẾN NGHỊ (e = {rec['e (m)']} m)", size=13)
        for line in [
            f"Tỷ lệ thay thế: a = {rec['a']:.4f}  ({rec['a (%)']:.1f}%)",
            f"Mô đun tương đương: Etb = {int(rec['Etb (kN/m²)']):,} kN/m²",
            f"Độ lún: S₁ = {rec['S₁ (cm)']:.2f} cm  ;  S₂ = 0  →  S = {rec['S₁ (cm)']:.2f} cm",
            f"Sức chịu tải: Pcol = {rec['Pcol (kN)']:.1f} kN  <  Qa = {rec['Qa (kN)']:.1f} kN  →  {rec['Đạt SCT']}",
        ]:
            H.body(doc, line, indent=0.3)

    # VII. Kiểm tra chọc thủng lớp đệm xi măng
    _ld_w   = params.get("loads", {})
    _pr_w   = calc_punching_check(
        D=D, e=rec.get("e (m)", 1.8) if rec else 1.8,
        Hse=_ld_w.get("h_mat", 0.4),
        He=_ld_w.get("h_fill", 1.5),
        quckse=params.get("quckse", 600.0),
        Fs=params.get("Fs_mat", 3.0),
        theta_deg=params.get("theta", 80.0),
        gamma_fill=_ld_w.get("g_fill", 18.0),
        gamma_mat=_ld_w.get("g_mat", 22.5),
        qa=params.get("qa_mat", 0.0),
    )
    H.heading(doc, "VII. KIỂM TRA CHỌC THỦNG LỚP ĐỆM XI MĂNG", size=13)
    H.body(doc, "Phương pháp: ALiCC (PWRI Japan). Điều kiện: τse ≤ τase.")
    H.tbl_caption(doc, "Kết quả kiểm tra chọc thủng lớp đệm xi măng")
    H.make_table(doc,
        headers=["Thông số", "Công thức", "Giá trị"],
        rows=[
            ("Bề dày đệm xi măng Hse",     "—",                              f"{_ld_w.get('h_mat',0.4):.2f} m"),
            ("Ứng suất cắt cho phép",       "τase = quckse/(2·Fs)",           f"{_pr_w['tase_kPa']:.2f} kPa"),
            ("Chiều cao vùng vòm Ho",       "Ho = (e−D)·tan(θ/2)",            f"{_pr_w['Ho_m']:.3f} m"),
            ("Công thức áp dụng",           "So sánh Ho và He",               _pr_w["cong_thuc"]),
            ("Áp lực vùng không gia cố",    "PSoil",                          f"{_pr_w['PSoil_kPa']:.3f} kPa"),
            ("Ứng suất cắt thực tế τse",    "τse = PSoil·A/(π·D·Hse)",        f"{_pr_w['tse_kPa']:.3f} kPa"),
            ("Kết quả kiểm tra",            "τse ≤ τase ?",                   _pr_w["check_CT"]),
        ])

    # VIII. Biểu đồ so sánh phương án
    import matplotlib.pyplot as _plt
    from docx.shared import Inches as _In

    def _bar_to_buf(x_lbls, y_vals, title, y_lbl, hi_idx, ref_line=None):
        _fig, _ax = _plt.subplots(figsize=(8, 3))
        _cols = ["#ED7D31" if i == hi_idx else "#4472C4" for i in range(len(y_vals))]
        _bars = _ax.bar(x_lbls, y_vals, color=_cols)
        _ax.bar_label(_bars, fmt="%.2f", padding=3, fontsize=9)
        if ref_line is not None:
            _ax.axhline(ref_line, color="#E53935", ls="--", lw=1.2,
                        label=f"Giới hạn = {ref_line:.1f}")
            _ax.legend(fontsize=9)
        _ax.set_title(title, fontsize=11)
        _ax.set_ylabel(y_lbl, fontsize=9)
        _ax.tick_params(labelsize=9)
        for _sp in ["top", "right"]:
            _ax.spines[_sp].set_visible(False)
        _buf = io.BytesIO()
        _fig.tight_layout()
        _fig.savefig(_buf, format="png", dpi=150, bbox_inches="tight")
        _plt.close(_fig)
        _buf.seek(0)
        return _buf

    if scenarios:
        H.heading(doc, "VIII. BIỂU ĐỒ SO SÁNH PHƯƠNG ÁN", size=13)
        _xl = [f"e={s['e (m)']}m" for s in scenarios]

        doc.add_picture(
            _bar_to_buf(_xl, [s["S₁ (cm)"] for s in scenarios],
                        f"Độ lún S₁ (cm) – D={int(D*1000)}mm, qu={int(qu)}kPa",
                        "S₁ (cm)", rec_idx),
            width=_In(5.5))
        doc.add_paragraph()
        doc.add_picture(
            _bar_to_buf(_xl, [s["a (%)"] for s in scenarios],
                        "Tỷ lệ thay thế a (%)",
                        "a (%)", rec_idx),
            width=_In(5.5))
        doc.add_paragraph()

        _tse_list = []
        for _s in scenarios:
            _pr_s = calc_punching_check(
                D=D, e=_s["e (m)"],
                Hse=_ld_w.get("h_mat", 0.4),
                He=_ld_w.get("h_fill", 1.5),
                quckse=params.get("quckse", 600.0),
                Fs=params.get("Fs_mat", 3.0),
                theta_deg=params.get("theta", 80.0),
                gamma_fill=_ld_w.get("g_fill", 18.0),
                gamma_mat=_ld_w.get("g_mat", 22.5),
                qa=params.get("qa_mat", 0.0),
            )
            _tse_list.append(_pr_s["tse_kPa"])
        doc.add_picture(
            _bar_to_buf(_xl, _tse_list,
                        "Ứng suất cắt τse – kiểm tra chọc thủng",
                        "τse (kPa)", rec_idx,
                        ref_line=_pr_w["tase_kPa"]),
            width=_In(5.5))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# PROJECT SAVE / LOAD
# ═══════════════════════════════════════════════════════════════════════════════
_SAVE_KEYS = ["cdm_zone","cdm_bh","cdm_h_clay","cdm_Su","cdm_gamma","cdm_elevation",
              "cdm_D","cdm_Lc","cdm_CDTK","cdm_qu","cdm_FS_lab","cdm_arrangement",
              "cdm_cement_type","cdm_dosage","cdm_WC","cdm_spacings","cdm_rec_idx","cdm_loads",
              "cdm_quckse","cdm_Fs_mat","cdm_theta","cdm_qa_mat"]

def _project_to_json() -> bytes:
    return json.dumps({k: st.session_state.get(k) for k in _SAVE_KEYS},
                      ensure_ascii=False, indent=2).encode("utf-8")

def _project_from_dict(data: dict):
    for k, v in data.items():
        if k in _SAVE_KEYS:
            st.session_state[k] = v


# ═══════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════════════════════════════════════════
st.sidebar.markdown("### CDM Design Tool")
st.sidebar.markdown("**Thiết kế cọc đất xi măng**")
st.sidebar.divider()

# Save / Load
with st.sidebar.expander("Lưu / Mở dự án", expanded=False):
    st.download_button("Tải xuống (.cdm)", _project_to_json(),
                       file_name="cdm_project.cdm", mime="application/json")
    _up = st.file_uploader("Mở file dự án (.cdm)", type=["cdm"], key="cdm_upload")
    if _up:
        try:
            _project_from_dict(json.load(_up))
            st.success("Đã tải dự án.")
            st.rerun()
        except Exception as e:
            st.error(f"Lỗi: {e}")

st.sidebar.divider()
_page = st.sidebar.radio("", _PAGES, label_visibility="collapsed")

# Info sidebar dưới cùng
st.sidebar.divider()
_q_now = q_total(_get("cdm_loads"))
_Su    = _get("cdm_Su")
_Lc    = _get("cdm_Lc")
_qu    = _get("cdm_qu")
_rec   = _get("cdm_rec_idx")
_sps   = _get("cdm_spacings")
st.sidebar.caption(
    f"**Zone:** {_get('cdm_zone')}  \n"
    f"**HK:** {_get('cdm_bh') or '–'}  \n"
    f"**q =** {_q_now:.1f} kN/m²  \n"
    f"**Su =** {_Su:.1f} kPa  \n"
    f"**Lc =** {_Lc:.1f} m  \n"
    f"**qu,tk =** {_qu:.0f} kPa"
)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 1 – ĐỊA CHẤT
# ═══════════════════════════════════════════════════════════════════════════════
if _page == "Địa chất":
    st.subheader("Địa chất – Chọn hố khoan")

    col_sel, col_prof = st.columns([1, 2])

    with col_sel:
        zone = st.radio("Zone", list(_ZONE_NAMES.keys()),
                        format_func=lambda z: _ZONE_NAMES[z],
                        index=list(_ZONE_NAMES.keys()).index(_get("cdm_zone")),
                        horizontal=False)
        st.session_state["cdm_zone"] = zone

        bhs = _load_boreholes_by_zone(zone)
        bh_names = [b["name"] for b in bhs]
        if not bh_names:
            st.warning("Zone này chưa có dữ liệu hố khoan.")
            bh_names = ["(trống)"]

        cur_bh = _get("cdm_bh")
        sel_idx = bh_names.index(cur_bh) if cur_bh in bh_names else 0
        bh_name = st.selectbox("Hố khoan", bh_names, index=sel_idx)

        if st.button("Áp dụng hố khoan", type="primary"):
            st.session_state["cdm_bh"] = bh_name
            # Auto-fill soil params
            cp = _clay_params(bh_name, zone)
            if cp:
                st.session_state["cdm_h_clay"]   = cp["h_clay"]
                st.session_state["cdm_elevation"] = cp.get("elevation_m", 0.0)
                st.session_state["cdm_Lc"]        = round(cp["h_clay"] + 1.5, 1)
            su = _su_avg_in_range(zone, 0, cp.get("depth_clay_bot", 30) if cp else 30)
            if su:
                st.session_state["cdm_Su"] = su
            g = _gamma_avg(bh_name)
            if g:
                st.session_state["cdm_gamma"] = g
            st.success(f"Đã áp dụng {bh_name}.")
            st.rerun()

        # Thông tin nhanh
        if _get("cdm_bh"):
            cp = _clay_params(_get("cdm_bh"), zone)
            if cp:
                st.metric("Bề dày lớp bùn sét", f"{cp['h_clay']:.1f} m")
                st.metric("Cao độ đáy lớp bùn",  f"{cp['elev_clay_bot']:.2f} m")
            su = _su_avg_in_range(zone, 0, cp.get("depth_clay_bot",30) if cp else 30)
            if su:
                st.metric("Su trung bình (VST)", f"{su:.1f} kPa")

    with col_prof:
        layers = _load_layers(bh_name)
        if layers:
            st.markdown("**Địa tầng hố khoan**")
            df_lay = pd.DataFrame(layers)
            df_lay.columns = ["Ký hiệu", "Mô tả", "Đỉnh (m)", "Đáy (m)", "Dày (m)",
                              "γ (kN/m³)", "c (kPa)", "φ (°)"]
            st.dataframe(
                df_lay,
                use_container_width=True,
                height=310,
                column_config={
                    "Mô tả":    st.column_config.TextColumn(width="large"),
                    "γ (kN/m³)": st.column_config.NumberColumn(format="%.2f"),
                    "c (kPa)":   st.column_config.NumberColumn(format="%.1f"),
                    "φ (°)":     st.column_config.NumberColumn(format="%.1f"),
                },
            )
        else:
            st.info("Hố khoan này chưa có dữ liệu địa tầng trong CSDL.")

    st.divider()
    # VST profile + soil column
    _su_col, _sc_col = st.columns([2, 1])
    with _su_col:
        df_vst = _load_vst_su(zone)
        if not df_vst.empty:
            _vst_all_locs = sorted(df_vst["loc_name"].unique().tolist())
            _vst_sel_prev = [l for l in _get("cdm_vst_sel") if l in _vst_all_locs]
            _vst_sel = st.multiselect(
                "Chọn hố khoan VST (để trống = toàn bộ)",
                options=_vst_all_locs,
                default=_vst_sel_prev,
                key="_vst_loc_sel",
            )
            st.session_state["cdm_vst_sel"] = _vst_sel
            _vst_pass = _vst_sel if _vst_sel else None
            if _HAS_PLOTLY:
                st.plotly_chart(
                    _chart_su_profile(df_vst, _vst_pass),
                    use_container_width=True,
                )
            else:
                st.dataframe(df_vst)
    with _sc_col:
        _bh_sc = bh_name if bh_name != "(trống)" else None
        if _bh_sc and _HAS_PLOTLY:
            try:
                _sc_layers = _load_layers(_bh_sc)
                st.plotly_chart(
                    _draw_soil_column(_sc_layers, _bh_sc),
                    use_container_width=True,
                    config={"displayModeBar": False},
                )
            except Exception as _e:
                st.warning(f"Không vẽ được cột địa chất: {_e}")
        else:
            st.caption("Chọn hố khoan để xem cột địa chất.")

    # CDM test results
    df_cdm = _load_cdm_tests()
    if not df_cdm.empty:
        with st.expander("Kết quả thí nghiệm CDM (R7)", expanded=False):
            st.dataframe(df_cdm[["bh_name","cement_type","WC_ratio","dosage_kgm3",
                                  "qu_R7_kPa","E50_R7_MPa"]].rename(columns={
                "bh_name":"HK","cement_type":"Xi măng","WC_ratio":"N/XM",
                "dosage_kgm3":"Hàm lượng (kg/m³)","qu_R7_kPa":"qu R7 (kPa)",
                "E50_R7_MPa":"E50 R7 (MPa)"}), use_container_width=True)
            st.caption("Lưu ý: Kết quả R7 (7 ngày). Cần qu_R90 để tính Ec chính xác.")

    # ── 3D / Map toggle ───────────────────────────────────────────────────────
    st.divider()
    if not _HAS_PLOTLY:
        st.caption("Cài `plotly` để xem bản đồ địa chất.")
    else:
        _bhs_all, _ = _load_borehole_3d_data()
        _zones_with_coords = sorted({b["zone"] for b in _bhs_all})
        if not _zones_with_coords:
            st.info("Chưa có tọa độ hố khoan trong CSDL (x_coord_m / y_coord_m = NULL).")
        else:
            # ── Thanh điều khiển ──────────────────────────────────────────────
            _ctrl_a, _ctrl_b = st.columns([3, 2])
            with _ctrl_a:
                _geo_view = st.radio(
                    "Chế độ xem",
                    ["3D địa chất", "Bản đồ vị trí"],
                    index=["3D địa chất", "Bản đồ vị trí"].index(
                        _get("cdm_geo_view")),
                    horizontal=True,
                    key="_geo_view_radio",
                )
                st.session_state["cdm_geo_view"] = _geo_view

            # ── 3D VIEW ───────────────────────────────────────────────────────
            if _geo_view == "3D địa chất":
                with _ctrl_b:
                    _sel_zones = st.multiselect(
                        "Zone", _zones_with_coords,
                        default=_zones_with_coords, key="_3d_cdm_zones",
                    )
                _cb1, _cb2, _cb3 = st.columns(3)
                _show_clay = _cb1.checkbox("Mặt đáy lớp bùn", value=True, key="_3d_cdm_clay")
                _show_cdm  = _cb2.checkbox("Đỉnh trụ CDM", value=False, key="_3d_cdm_top")
                _cdm_top_z = float(_get("cdm_CDTK"))
                if _show_cdm:
                    _cdm_top_z = _cb3.number_input(
                        "Cao độ (m)", value=_cdm_top_z, step=0.1, key="_3d_cdm_top_z")
                _vst_focus = _get("cdm_vst_sel")
                _focus_bh = _vst_focus[0] if len(_vst_focus) == 1 else None
                if _sel_zones:
                    st.plotly_chart(
                        _draw_boreholes_3d(
                            _sel_zones, _show_clay, _show_cdm, _cdm_top_z,
                            focus_bh=_focus_bh,
                        ),
                        use_container_width=True, config={"displayModeBar": True},
                    )
                else:
                    st.info("Chọn ít nhất một zone.")

            # ── MAP VIEW ──────────────────────────────────────────────────────
            else:
                with _ctrl_b:
                    _map_style_label = st.selectbox(
                        "Nền bản đồ",
                        list(_MAP_STYLES.keys()),
                        index=list(_MAP_STYLES.keys()).index(
                            _get("cdm_map_style")
                            if _get("cdm_map_style") in _MAP_STYLES else "Đường phố (OSM)"
                        ),
                        key="_map_style_sel",
                    )
                    st.session_state["cdm_map_style"] = _map_style_label

                # Hiệu chỉnh GPS
                _ref_default = _get("cdm_map_ref_bh") or ""
                _calibrated  = bool(_ref_default)
                with st.expander(
                    ("Hiệu chỉnh GPS — đã áp dụng" if _calibrated
                     else "Hiệu chỉnh GPS — cần thiết lập lần đầu"),
                    expanded=not _calibrated,
                ):
                    st.markdown(
                        "**Cách lấy toạ độ GPS:**  \n"
                        "1. Mở [Google Maps](https://maps.google.com) → tìm vị trí hố khoan trên thực địa  \n"
                        "2. **Chuột phải** vào điểm đó → copy dòng số đầu tiên (ví dụ: `10.782341, 106.718560`)  \n"
                        "3. Chọn hố khoan tương ứng bên dưới → nhập lat/lon → **Áp dụng**"
                    )
                    _bh_names_all = [b["name"] for b in _bhs_all]
                    _ref_idx = _bh_names_all.index(_ref_default) if _ref_default in _bh_names_all else 0
                    _c1, _c2, _c3 = st.columns([2, 1, 1])
                    _ref_bh  = _c1.selectbox("Hố khoan tham chiếu", _bh_names_all,
                                             index=_ref_idx, key="_map_ref_bh_sel")
                    _ref_lat = _c2.number_input("Latitude", value=float(_get("cdm_map_ref_lat")),
                                                format="%.6f", step=0.000100, key="_map_ref_lat")
                    _ref_lon = _c3.number_input("Longitude", value=float(_get("cdm_map_ref_lon")),
                                                format="%.6f", step=0.000100, key="_map_ref_lon")
                    if st.button("Áp dụng hiệu chỉnh", type="primary", key="_map_calib_btn"):
                        st.session_state.update({
                            "cdm_map_ref_bh":  _ref_bh,
                            "cdm_map_ref_lat": _ref_lat,
                            "cdm_map_ref_lon": _ref_lon,
                        })
                        st.rerun()
                    if _calibrated:
                        st.success(
                            f"Đang dùng: {_get('cdm_map_ref_bh')} → "
                            f"lat={_get('cdm_map_ref_lat'):.6f}, "
                            f"lon={_get('cdm_map_ref_lon'):.6f}"
                        )
                _use_gps = True
                _ref_bh  = _get("cdm_map_ref_bh") or (_bh_names_all[0] if _bh_names_all else "")
                _ref_lat = float(_get("cdm_map_ref_lat"))
                _ref_lon = float(_get("cdm_map_ref_lon"))

                _bhs_ll = _project_to_latlon(
                    _bhs_all, _ref_bh, _ref_lat, _ref_lon, use_gps_ref=_use_gps
                )
                if _bhs_ll:
                    st.plotly_chart(
                        _draw_map_2d(_bhs_ll, _MAP_STYLES[_map_style_label]),
                        use_container_width=True,
                        config={"displayModeBar": True, "scrollZoom": True},
                    )
                else:
                    st.warning("Không có hố khoan nào có tọa độ.")


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 2 – THÔNG SỐ
# ═══════════════════════════════════════════════════════════════════════════════
elif _page == "Thông số":
    st.subheader("Thông số CDM")

    # Layout: thông số bên trái, hình minh họa bên phải
    _col_inp, _col_sec = st.columns([3, 2], gap="medium")

    with _col_inp:
        c1, c2, c3, c4 = st.columns(4)

        with c1:
            st.markdown("**Hình học trụ CDM**")
            D    = st.number_input("Đường kính D (m)", 0.5, 1.2, _get("cdm_D"), 0.05)
            Lc   = st.number_input("Chiều dài Lc (m)", 5.0, 60.0, _get("cdm_Lc"), 0.5)
            CDTK = st.number_input("Cao độ đỉnh cọc CDM (m)", -5.0, 10.0, _get("cdm_CDTK"), 0.1)
            _ld0 = _get("cdm_loads")
            _cdtk_cal = CDTK + _ld0.get("h_mat", 0.4) + _ld0.get("h_fill", 1.5)
            st.info(f"CDTK = đỉnh + đệm + đắp = **{_cdtk_cal:+.2f} m**")
            arr  = st.radio("Bố trí lưới", ["triangle", "square"],
                            format_func=lambda x: "Tam giác" if x == "triangle" else "Hình vuông",
                            index=["triangle", "square"].index(_get("cdm_arrangement")))
            st.session_state.update({"cdm_D": D, "cdm_Lc": Lc,
                                      "cdm_CDTK": CDTK, "cdm_arrangement": arr})

        with c2:
            st.markdown("**Vật liệu xi măng đất**")
            cement = st.text_input("Loại xi măng", _get("cdm_cement_type"))
            dosage = st.number_input("Hàm lượng XM (kg/m³)", 100, 400, _get("cdm_dosage"), 10)
            WC     = st.number_input("Tỷ lệ N/XM", 0.5, 2.0, _get("cdm_WC"), 0.1)
            qu     = st.number_input("qu thiết kế HT (kPa)", 100.0, 5000.0, _get("cdm_qu"), 50.0)
            FS_lab = st.number_input("FS quy đổi TN→HT", 1.0, 5.0, _get("cdm_FS_lab"), 0.1)
            Cc = qu / 2
            Ec = 100 * Cc
            st.info(f"Cc = {Cc:.0f} kN/m²  |  Ec = {int(Ec):,} kN/m²")
            st.session_state.update({"cdm_cement_type": cement, "cdm_dosage": dosage,
                                      "cdm_WC": WC, "cdm_qu": qu, "cdm_FS_lab": FS_lab})

        with c3:
            st.markdown("**Thông số địa kỹ thuật**")
            h_clay  = st.number_input("Bề dày lớp bùn h₁ (m)", 1.0, 60.0, _get("cdm_h_clay"), 0.5)
            Su      = st.number_input("Su trung bình (kPa)", 1.0, 100.0, _get("cdm_Su"), 0.5)
            gamma   = st.number_input("γ tự nhiên (kN/m³)", 10.0, 22.0, _get("cdm_gamma"), 0.1)
            Es_show = 250 * Su
            st.info(f"Es = 250×Su = {int(Es_show):,} kN/m²")
            st.session_state.update({"cdm_h_clay": h_clay, "cdm_Su": Su, "cdm_gamma": gamma})

        with c4:
            st.markdown("**Tải trọng gây lún**")
            ld = _get("cdm_loads").copy()
            ld["q_traffic"] = st.number_input("Hoạt tải (kN/m²)", 0.0, 200.0,
                                              ld.get("q_traffic", 20.0), 5.0)
            st.markdown("*Áo đường:*")
            ld["h_road"] = st.number_input("h (m)", 0.0, 2.0, ld.get("h_road", 0.8), 0.1,
                                           key="h_road")
            ld["g_road"] = st.number_input("γ (kN/m³)", 10.0, 30.0, ld.get("g_road", 24.0), 0.5,
                                           key="g_road")
            st.markdown("*Cát đắp:*")
            ld["h_fill"] = st.number_input("h (m)", 0.0, 5.0, ld.get("h_fill", 1.5), 0.1,
                                           key="h_fill")
            ld["g_fill"] = st.number_input("γ (kN/m³)", 10.0, 24.0, ld.get("g_fill", 18.0), 0.5,
                                           key="g_fill")
            st.markdown("*Đệm cát:*")
            ld["h_mat"] = st.number_input("h (m)", 0.0, 2.0, ld.get("h_mat", 0.4), 0.1,
                                          key="h_mat")
            ld["g_mat"] = st.number_input("γ (kN/m³)", 10.0, 26.0, ld.get("g_mat", 22.5), 0.5,
                                          key="g_mat")
            q_tot = q_total(ld)
            st.success(f"**q = {q_tot:.2f} kN/m²**")
            st.session_state["cdm_loads"] = ld

    # ── Kiểm tra chọc thủng – thông số đệm xi măng ───────────────────────────
    with st.expander("Kiểm tra chọc thủng lớp đệm xi măng", expanded=False):
        _km1, _km2, _km3, _km4 = st.columns(4)
        with _km1:
            _quckse = st.number_input("qu đệm xi măng (kPa)", 100.0, 3000.0,
                                      _get("cdm_quckse"), 50.0)
            st.session_state["cdm_quckse"] = _quckse
        with _km2:
            _Fs_mat = st.number_input("Hệ số an toàn Fs", 1.0, 5.0,
                                      _get("cdm_Fs_mat"), 0.5)
            _tase_p = _quckse / (2.0 * _Fs_mat)
            st.info(f"τase = **{_tase_p:.1f} kPa**")
            st.session_state["cdm_Fs_mat"] = _Fs_mat
        with _km3:
            _theta = st.number_input("Góc đàn hồi dẻo θ (°)", 30.0, 89.0,
                                     _get("cdm_theta"), 5.0)
            st.session_state["cdm_theta"] = _theta
        with _km4:
            _qa_mat = st.number_input("qa vùng không GC (kPa)", 0.0, 200.0,
                                      _get("cdm_qa_mat"), 5.0)
            st.session_state["cdm_qa_mat"] = _qa_mat

    # ── Hình minh họa bên phải: mặt cắt (trên) + lưới (dưới) ────────────────
    with _col_sec:
        _ld_s = _get("cdm_loads")
        _fig_sec = _draw_cdm_section(
            D=_get("cdm_D"), Lc=_get("cdm_Lc"), CDTK=_get("cdm_CDTK"),
            h_clay=_get("cdm_h_clay"),
            h_road=_ld_s.get("h_road", 0.8),
            h_fill=_ld_s.get("h_fill", 1.5),
            h_mat=_ld_s.get("h_mat", 0.4),
            q_traffic=_ld_s.get("q_traffic", 20.0),
        )
        st.pyplot(_fig_sec, use_container_width=True)
        _fig_sec.clf()

        _sps_now   = _get("cdm_spacings")
        _rec_now   = min(_get("cdm_rec_idx"), len(_sps_now) - 1) if _sps_now else 0
        _e_ref_now = _sps_now[_rec_now] if _sps_now else 1.6
        _fig_grd = _draw_cdm_grid(_get("cdm_D"), _get("cdm_arrangement"), _e_ref_now)
        st.pyplot(_fig_grd, use_container_width=True)
        _fig_grd.clf()

    # ── Lý thuyết tính toán ──────────────────────────────────────────────────
    _theory_text = _load_theory()
    if _theory_text:
        st.divider()
        with st.expander("Lý thuyết tính toán – TCVN 9403:2012", expanded=False):
            st.markdown(_theory_text)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 3 – SO SÁNH PHƯƠNG ÁN
# ═══════════════════════════════════════════════════════════════════════════════
elif _page == "So sánh PA":
    st.subheader("So sánh các phương án khoảng cách trụ")

    # Input khoảng cách
    with st.expander("Thiết lập khoảng cách trụ", expanded=True):
        mode = st.radio("Chế độ nhập", ["Tự động (min/max/bước)", "Nhập tay"],
                        horizontal=True)
        if mode == "Tự động (min/max/bước)":
            c1, c2, c3 = st.columns(3)
            e_min  = c1.number_input("e min (m)", 1.0, 3.0, 1.2, 0.1)
            e_max  = c2.number_input("e max (m)", 1.0, 4.0, 2.0, 0.1)
            e_step = c3.number_input("Bước (m)",  0.1, 1.0, 0.2, 0.1)
            spacings = [round(e_min + i * e_step, 2)
                        for i in range(int(round((e_max - e_min) / e_step)) + 1)]
        else:
            raw = st.text_input("Danh sách khoảng cách (m, cách nhau bởi dấu phẩy)",
                                ", ".join(str(s) for s in _get("cdm_spacings")))
            try:
                spacings = [float(x.strip()) for x in raw.split(",") if x.strip()]
            except ValueError:
                spacings = _get("cdm_spacings")
        spacings = sorted(set(spacings))
        st.session_state["cdm_spacings"] = spacings

    # Tính toán
    D   = _get("cdm_D")
    Lc  = _get("cdm_Lc")
    qu  = _get("cdm_qu")
    Su  = _get("cdm_Su")
    arr = _get("cdm_arrangement")
    q   = q_total(_get("cdm_loads"))

    scenarios = build_scenarios(D, Lc, qu, Su, q, spacings, arr)
    df = pd.DataFrame(scenarios)

    # Chọn PA kiến nghị
    rec_idx = st.selectbox(
        "Phương án kiến nghị",
        range(len(scenarios)),
        index=min(_get("cdm_rec_idx"), len(scenarios)-1),
        format_func=lambda i: f"PA{i+1}: e={scenarios[i]['e (m)']}m  →  S₁={scenarios[i]['S₁ (cm)']:.2f}cm",
    )
    st.session_state["cdm_rec_idx"] = rec_idx

    # Tính kiểm tra chọc thủng cho mỗi PA
    _ld_punch = _get("cdm_loads")
    _punch = [
        calc_punching_check(
            D=D, e=s["e (m)"],
            Hse=_ld_punch.get("h_mat", 0.4),
            He=_ld_punch.get("h_fill", 1.5),
            quckse=_get("cdm_quckse"),
            Fs=_get("cdm_Fs_mat"),
            theta_deg=_get("cdm_theta"),
            gamma_fill=_ld_punch.get("g_fill", 18.0),
            gamma_mat=_ld_punch.get("g_mat", 22.5),
            qa=_get("cdm_qa_mat"),
        )
        for s in scenarios
    ]
    for i, pr in enumerate(_punch):
        df.at[i, "τse (kPa)"]  = pr["tse_kPa"]
        df.at[i, "τase (kPa)"] = pr["tase_kPa"]
        df.at[i, "Đạt CT"]     = pr["check_CT"]

    # Bảng kết quả
    st.markdown("**Bảng so sánh phương án**")
    display_cols = ["e (m)", "a (%)", "Etb (kN/m²)", "S₁ (cm)",
                    "Pcol (kN)", "Qa (kN)", "Đạt SCT",
                    "τse (kPa)", "τase (kPa)", "Đạt CT"]

    def _row_color(row):
        idx = df.index.get_loc(row.name)
        if idx == rec_idx:
            return ["background-color: #E2EFDA"] * len(row)
        return [""] * len(row)

    st.dataframe(
        df[display_cols].style.apply(_row_color, axis=1)
                               .format({"a (%)": "{:.1f}%",
                                        "Etb (kN/m²)": "{:,.0f}",
                                        "S₁ (cm)": "{:.2f}",
                                        "Pcol (kN)": "{:.1f}",
                                        "Qa (kN)": "{:.1f}"}),
        use_container_width=True,
    )

    if _HAS_PLOTLY:
        ca, cb, cc = st.columns(3)
        with ca:
            st.plotly_chart(_chart_scenarios(df, rec_idx), use_container_width=True)
        with cb:
            st.plotly_chart(_chart_etb(df, rec_idx), use_container_width=True)
        with cc:
            st.plotly_chart(_chart_combined(df, rec_idx), use_container_width=True)

        # Biểu đồ chọc thủng
        _x_sp = [f"e={s['e (m)']}m" for s in scenarios]
        _col_ct = ["#ED7D31" if i == rec_idx else "#4472C4" for i in range(len(scenarios))]
        _tase_v = _punch[0]["tase_kPa"] if _punch else 100.0
        _fig_ct = go.Figure()
        _fig_ct.add_trace(go.Bar(
            name="τse (kPa)", x=_x_sp,
            y=[pr["tse_kPa"] for pr in _punch],
            marker_color=_col_ct,
            text=[f"{pr['tse_kPa']:.1f}" for pr in _punch],
            textposition="outside",
        ))
        _fig_ct.add_hline(
            y=_tase_v, line_dash="dash", line_color="#E53935", line_width=1.5,
            annotation_text=f"τase = {_tase_v:.0f} kPa",
            annotation_position="top right",
        )
        _fig_ct.update_layout(
            title="Kiểm tra chọc thủng đệm xi măng",
            yaxis_title="Ứng suất cắt (kPa)",
            height=340, margin=dict(t=45, b=20),
            showlegend=False,
        )
        st.plotly_chart(_fig_ct, use_container_width=True)

    # ── Phân tích tham số ─────────────────────────────────────────────────────
    st.divider()
    st.markdown("### Phân tích tham số (giữ nguyên e = PA kiến nghị)")

    _e_rec   = scenarios[rec_idx]["e (m)"]
    _ld_base = _get("cdm_loads").copy()
    _hmat_cur = _ld_base.get("h_mat", 0.4)

    _pa1, _pa2, _pa3 = st.columns(3)

    # ── 1. So sánh qu ─────────────────────────────────────────────────────────
    with _pa1:
        st.markdown("**So sánh qu thiết kế**")
        _qu_raw = st.text_input("Các giá trị qu (kPa)",
                                ", ".join(str(v) for v in [200, 300, 400, 500, 600, 700]),
                                key="_qu_list")
        try:
            _qu_list = sorted(set(float(x.strip()) for x in _qu_raw.split(",") if x.strip()))
        except ValueError:
            _qu_list = [200, 300, 400, 500]

        _qu_rows = []
        for _qv in _qu_list:
            _Ec_v  = calc_Ec(_qv)
            _Es_v  = calc_Es(Su)
            _a_v   = calc_a(D, _e_rec, arr)
            _Etb_v = calc_Etb(_a_v, _Ec_v, _Es_v)
            _S1_v  = calc_S1(q, Lc, _Etb_v)
            _sc_v  = calc_sigma_col(q, _Ec_v, _Etb_v)
            _Pcol_v = calc_Pcol(_sc_v, D)
            _bc_v  = calc_bearing(D, Lc, _qv / 2, Su)
            _qu_rows.append({
                "qu (kPa)": int(_qv),
                "Ec (kN/m²)": int(_Ec_v),
                "Etb (kN/m²)": int(_Etb_v),
                "S₁ (cm)": round(_S1_v, 2),
                "Pcol (kN)": round(_Pcol_v, 1),
                "Qa (kN)": _bc_v["Qa"],
                "Đạt SCT": "Đạt" if _Pcol_v < _bc_v["Qa"] else "Không đạt",
            })
        _df_qu = pd.DataFrame(_qu_rows)
        st.dataframe(_df_qu, use_container_width=True, hide_index=True,
                     column_config={"Etb (kN/m²)": st.column_config.NumberColumn(format="%,d"),
                                    "Ec (kN/m²)":  st.column_config.NumberColumn(format="%,d")})
        if _HAS_PLOTLY:
            _col_qu = ["#ED7D31" if r["qu (kPa)"] == int(qu) else "#4472C4"
                       for r in _qu_rows]
            _fig_qu = go.Figure(go.Bar(
                x=[f"{r['qu (kPa)']} kPa" for r in _qu_rows],
                y=_df_qu["S₁ (cm)"],
                marker_color=_col_qu,
                text=_df_qu["S₁ (cm)"].apply(lambda v: f"{v:.2f}"),
                textposition="outside",
            ))
            _fig_qu.update_layout(
                title=f"S₁ theo qu  (e={_e_rec}m, D={D}m)",
                yaxis_title="S₁ (cm)", height=300,
                margin=dict(t=45, b=20),
            )
            st.plotly_chart(_fig_qu, use_container_width=True)

    # ── 2. So sánh D ──────────────────────────────────────────────────────────
    with _pa2:
        st.markdown("**So sánh đường kính D**")
        _D_raw = st.text_input("Các giá trị D (m)", "0.50, 0.60, 0.70, 0.80, 1.00",
                               key="_D_list")
        try:
            _D_list = sorted(set(float(x.strip()) for x in _D_raw.split(",") if x.strip()))
        except ValueError:
            _D_list = [0.5, 0.6, 0.7, 0.8]

        _D_rows = []
        for _Dv in _D_list:
            _Ec_v  = calc_Ec(qu)
            _Es_v  = calc_Es(Su)
            _a_v   = calc_a(_Dv, _e_rec, arr)
            _Etb_v = calc_Etb(_a_v, _Ec_v, _Es_v)
            _S1_v  = calc_S1(q, Lc, _Etb_v)
            _sc_v  = calc_sigma_col(q, _Ec_v, _Etb_v)
            _Pcol_v = calc_Pcol(_sc_v, _Dv)
            _bc_v  = calc_bearing(_Dv, Lc, qu / 2, Su)
            _D_rows.append({
                "D (m)": _Dv,
                "a (%)": round(_a_v * 100, 1),
                "Etb (kN/m²)": int(_Etb_v),
                "S₁ (cm)": round(_S1_v, 2),
                "Pcol (kN)": round(_Pcol_v, 1),
                "Qa (kN)": _bc_v["Qa"],
                "Đạt SCT": "Đạt" if _Pcol_v < _bc_v["Qa"] else "Không đạt",
            })
        _df_D = pd.DataFrame(_D_rows)
        st.dataframe(_df_D, use_container_width=True, hide_index=True,
                     column_config={"a (%)": st.column_config.NumberColumn(format="%.1f%%"),
                                    "Etb (kN/m²)": st.column_config.NumberColumn(format="%,d")})
        if _HAS_PLOTLY:
            _col_D = ["#ED7D31" if abs(_Dv - D) < 0.001 else "#4472C4"
                      for _Dv in _D_list]
            _fig_D = go.Figure()
            _fig_D.add_trace(go.Bar(
                name="a (%)",
                x=[f"{r['D (m)']}m" for r in _D_rows],
                y=_df_D["a (%)"],
                marker_color=_col_D,
                text=_df_D["a (%)"].apply(lambda v: f"{v:.1f}%"),
                textposition="outside",
                yaxis="y1",
            ))
            _fig_D.add_trace(go.Scatter(
                name="S₁ (cm)",
                x=[f"{r['D (m)']}m" for r in _D_rows],
                y=_df_D["S₁ (cm)"],
                mode="lines+markers+text",
                text=_df_D["S₁ (cm)"].apply(lambda v: f"{v:.2f}"),
                textposition="top center",
                yaxis="y2",
                line=dict(color="#E53935", width=2),
                marker=dict(size=8, color="#E53935"),
            ))
            _fig_D.update_layout(
                title=f"a(%) và S₁ theo D  (e={_e_rec}m, qu={int(qu)}kPa)",
                yaxis=dict(title="a (%)"),
                yaxis2=dict(title="S₁ (cm)", overlaying="y", side="right"),
                height=300, margin=dict(t=45, b=20),
                legend=dict(orientation="h"),
            )
            st.plotly_chart(_fig_D, use_container_width=True)

    # ── 3. So sánh h_mat ──────────────────────────────────────────────────────
    with _pa3:
        st.markdown("**So sánh chiều dày đệm cát**")
        _hm_raw = st.text_input("Các giá trị h_mat (m)", "0.20, 0.30, 0.40, 0.50, 0.60, 0.80",
                                key="_hmat_list")
        try:
            _hmat_list = sorted(set(float(x.strip()) for x in _hm_raw.split(",") if x.strip()))
        except ValueError:
            _hmat_list = [0.2, 0.3, 0.4, 0.5]

        _hm_rows = []
        for _hm in _hmat_list:
            _ld_v = _ld_base.copy()
            _ld_v["h_mat"] = _hm
            _q_v   = q_total(_ld_v)
            _Ec_v  = calc_Ec(qu)
            _Es_v  = calc_Es(Su)
            _a_v   = calc_a(D, _e_rec, arr)
            _Etb_v = calc_Etb(_a_v, _Ec_v, _Es_v)
            _S1_v  = calc_S1(_q_v, Lc, _Etb_v)
            _hm_rows.append({
                "h_mat (m)": _hm,
                "q (kN/m²)": round(_q_v, 2),
                "ΔS₁ (cm)": round(_S1_v, 2),
            })
        _df_hm = pd.DataFrame(_hm_rows)
        st.dataframe(_df_hm, use_container_width=True, hide_index=True)
        if _HAS_PLOTLY:
            _col_hm = ["#ED7D31" if abs(_hm - _hmat_cur) < 0.001 else "#4472C4"
                       for _hm in _hmat_list]
            _fig_hm = go.Figure()
            _fig_hm.add_trace(go.Bar(
                name="q (kN/m²)",
                x=[f"{r['h_mat (m)']}m" for r in _hm_rows],
                y=_df_hm["q (kN/m²)"],
                marker_color=_col_hm,
                text=_df_hm["q (kN/m²)"].apply(lambda v: f"{v:.1f}"),
                textposition="outside",
                yaxis="y1",
            ))
            _fig_hm.add_trace(go.Scatter(
                name="S₁ (cm)",
                x=[f"{r['h_mat (m)']}m" for r in _hm_rows],
                y=_df_hm["ΔS₁ (cm)"],
                mode="lines+markers+text",
                text=_df_hm["ΔS₁ (cm)"].apply(lambda v: f"{v:.2f}"),
                textposition="top center",
                yaxis="y2",
                line=dict(color="#E53935", width=2),
                marker=dict(size=8, color="#E53935"),
            ))
            _fig_hm.update_layout(
                title=f"q và S₁ theo h_mat  (e={_e_rec}m, D={D}m)",
                yaxis=dict(title="q (kN/m²)"),
                yaxis2=dict(title="S₁ (cm)", overlaying="y", side="right"),
                height=300, margin=dict(t=45, b=20),
                legend=dict(orientation="h"),
            )
            st.plotly_chart(_fig_hm, use_container_width=True)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 4 – KẾT QUẢ
# ═══════════════════════════════════════════════════════════════════════════════
elif _page == "Kết quả":
    st.subheader("Kết quả chi tiết – Phương án kiến nghị")

    D   = _get("cdm_D")
    Lc  = _get("cdm_Lc")
    qu  = _get("cdm_qu")
    Su  = _get("cdm_Su")
    arr = _get("cdm_arrangement")
    q   = q_total(_get("cdm_loads"))
    sps = _get("cdm_spacings")
    rec = _get("cdm_rec_idx")

    scenarios = build_scenarios(D, Lc, qu, Su, q, sps, arr)
    if not scenarios or rec >= len(scenarios):
        st.warning("Chưa có phương án. Vào tab So sánh PA để tính.")
        st.stop()

    s   = scenarios[rec]
    Cc  = qu / 2
    Ec  = calc_Ec(qu)
    Es  = calc_Es(Su)
    Etb = s["Etb (kN/m²)"]

    st.success(f"Phương án kiến nghị: **PA{rec+1} — e = {s['e (m)']} m**")

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("S₁ (cm)", f"{s['S₁ (cm)']:.2f}")
    c2.metric("Etb (kN/m²)", f"{int(Etb):,}")
    c3.metric("a (%)", f"{s['a (%)']:.1f}%")
    c4.metric("Sức chịu tải", s["Đạt SCT"])

    st.divider()
    col_l, col_r = st.columns(2)

    with col_l:
        st.markdown("**Tính toán độ lún**")
        arr_label = "Tam giác: a = π·D²/(2√3·e²)" if arr == "triangle" else "Vuông: a = π·D²/(4·e²)"
        steps = [
            ("Khoảng cách trụ", "e", f"{s['e (m)']} m"),
            ("Tỷ lệ thay thế", arr_label, f"{s['a']:.4f}  ({s['a (%)']:.1f}%)"),
            ("Mô đun trụ CDM", "Ec = 100·Cc = 100·(qu/2)", f"{int(Ec):,} kN/m²"),
            ("Mô đun đất nền", "Es = 250·Su", f"{int(Es):,} kN/m²"),
            ("Mô đun tương đương", "Etb = a·Ec + (1−a)·Es", f"{int(Etb):,} kN/m²"),
            ("Tải trọng thiết kế", "q", f"{q:.2f} kN/m²"),
            ("Chiều dài trụ", "Lc", f"{Lc:.1f} m"),
            ("Độ lún bản thân CDM", "S₁ = q·Lc/Etb × 100", f"{s['S₁ (cm)']:.2f} cm"),
            ("Độ lún dưới mũi trụ", "S₂ = 0 (trụ xuyên qua lớp bùn)", "0,00 cm"),
            ("Tổng độ lún", "S = S₁ + S₂", f"{s['S₁ (cm)']:.2f} cm"),
        ]
        df_steps = pd.DataFrame(steps, columns=["Thông số", "Công thức", "Giá trị"])
        st.dataframe(df_steps, use_container_width=True, hide_index=True,
                     column_config={"Công thức": st.column_config.TextColumn(width="medium")})

    with col_r:
        st.markdown("**Kiểm tra sức chịu tải (AIT – TCVN 9403:2012 Phụ lục B)**")
        bc = calc_bearing(D, Lc, Cc, Su)
        sct_steps = [
            ("Diện tích cắt ngang", "Ac = π·D²/4", f"{math.pi*D**2/4:.4f} m²"),
            ("Sức chịu tải mũi", "Qult_mũi = 9·Cc·Ac", f"{bc['Qult_col']:.1f} kN"),
            ("Sức chịu tải thân", "Qult_thân = π·D·Lc·Cu", f"{bc['Qult_skin']:.1f} kN"),
            ("Tổng sức chịu tải", "Qult = Qmũi + Qthân", f"{bc['Qult']:.1f} kN"),
            ("Sức chịu tải cho phép", "Qa = Qult / FS = Qult / 2", f"{bc['Qa']:.1f} kN"),
            ("Ứng suất đầu cọc", "σ_col = (Ec/Etb) × q", f"{s['σ_col (kN/m²)']:.1f} kN/m²"),
            ("Lực nén lên 1 trụ", "Pcol = σ_col × Ac", f"{s['Pcol (kN)']:.1f} kN"),
            ("Kiểm tra", "Pcol < Qa ?", s["Đạt SCT"]),
        ]
        df_sct = pd.DataFrame(sct_steps, columns=["Thông số", "Công thức", "Giá trị"])

        def _color_sct(val):
            if val == "Đạt":    return "color: green; font-weight: bold"
            if val == "Không đạt": return "color: red; font-weight: bold"
            return ""

        st.dataframe(df_sct.style.map(_color_sct, subset=["Giá trị"]),
                     use_container_width=True, hide_index=True,
                     column_config={"Công thức": st.column_config.TextColumn(width="medium")})

    st.divider()
    st.markdown("**Kiểm tra chọc thủng lớp đệm xi măng (ALiCC)**")
    _ld_r   = _get("cdm_loads")
    _pr_r   = calc_punching_check(
        D=D, e=s["e (m)"],
        Hse=_ld_r.get("h_mat", 0.4),
        He=_ld_r.get("h_fill", 1.5),
        quckse=_get("cdm_quckse"),
        Fs=_get("cdm_Fs_mat"),
        theta_deg=_get("cdm_theta"),
        gamma_fill=_ld_r.get("g_fill", 18.0),
        gamma_mat=_ld_r.get("g_mat", 22.5),
        qa=_get("cdm_qa_mat"),
    )
    _punch_steps = [
        ("Bề dày đệm xi măng",        "Hse",                                f"{_ld_r.get('h_mat',0.4):.2f} m"),
        ("Chiều cao đất đắp trên đệm", "He = h_fill",                        f"{_ld_r.get('h_fill',1.5):.2f} m"),
        ("Cường độ kháng nén đệm XM",  "quckse",                             f"{_get('cdm_quckse'):.0f} kPa"),
        ("Ứng suất cắt cho phép",      "τase = quckse/(2·Fs)",               f"{_pr_r['tase_kPa']:.2f} kPa"),
        ("Chiều cao vùng vòm đất",     "Ho = (e−D)·tan(θ/2)",                f"{_pr_r['Ho_m']:.3f} m"),
        ("So sánh Ho vs He",           "Chọn công thức",                     _pr_r["cong_thuc"]),
        ("Thể tích đất tác dụng",      "Vsoil",                              f"{_pr_r['Vsoil_m3']:.4f} m³"),
        ("Thể tích đệm XM tác dụng",   "VCGCXM",                             f"{_pr_r['VCGCXM_m3']:.4f} m³"),
        ("Áp lực vùng không gia cố",   "PSoil",                              f"{_pr_r['PSoil_kPa']:.3f} kPa"),
        ("Ứng suất cắt thực tế",       "τse = PSoil·(e²−πD²/4)/(π·D·Hse)",  f"{_pr_r['tse_kPa']:.3f} kPa"),
        ("Kiểm tra",                   "τse ≤ τase ?",                       _pr_r["check_CT"]),
    ]
    _df_punch = pd.DataFrame(_punch_steps, columns=["Thông số", "Công thức", "Giá trị"])
    st.dataframe(
        _df_punch.style.map(_color_sct, subset=["Giá trị"]),
        use_container_width=True, hide_index=True,
        column_config={"Công thức": st.column_config.TextColumn(width="large")},
    )


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 5 – XUẤT
# ═══════════════════════════════════════════════════════════════════════════════
elif _page == "Xuất":
    st.subheader("Xuất báo cáo")

    D   = _get("cdm_D")
    Lc  = _get("cdm_Lc")
    qu  = _get("cdm_qu")
    Su  = _get("cdm_Su")
    arr = _get("cdm_arrangement")
    q   = q_total(_get("cdm_loads"))
    sps = _get("cdm_spacings")
    rec = _get("cdm_rec_idx")

    scenarios = build_scenarios(D, Lc, qu, Su, q, sps, arr)
    rec = min(rec, len(scenarios) - 1) if scenarios else 0

    params = {
        "bh_name":   _get("cdm_bh") or "–",
        "zone":      _get("cdm_zone"),
        "h_clay":    _get("cdm_h_clay"),
        "Su": Su,    "gamma": _get("cdm_gamma"),
        "D": D,      "Lc": Lc, "CDTK": _get("cdm_CDTK"),
        "qu": qu,    "FS_lab": _get("cdm_FS_lab"),
        "dosage":    _get("cdm_dosage"),
        "arrangement": arr,
        "q_total":   q,
        "rec_idx":   rec,
        "loads":     _get("cdm_loads"),
        "quckse":    _get("cdm_quckse"),
        "Fs_mat":    _get("cdm_Fs_mat"),
        "theta":     _get("cdm_theta"),
        "qa_mat":    _get("cdm_qa_mat"),
    }

    st.markdown("### Tóm tắt trước khi xuất")
    if scenarios:
        s = scenarios[rec]
        col1, col2, col3 = st.columns(3)
        col1.metric("Phương án kiến nghị", f"e = {s['e (m)']} m")
        col2.metric("Độ lún S₁", f"{s['S₁ (cm)']:.2f} cm")
        col3.metric("Sức chịu tải", s["Đạt SCT"])

    st.divider()
    c_word, c_excel = st.columns(2)

    with c_word:
        st.markdown("#### Thuyết minh tính toán (Word)")
        st.caption("Tài liệu gồm 6 mục: Cơ sở pháp lý · Địa chất · Thông số · "
                   "Phương pháp · So sánh PA · Kết quả PA kiến nghị.")
        if st.button("Tạo file Word", type="primary"):
            with st.spinner("Đang tạo tài liệu..."):
                word_bytes = _export_word_bytes(scenarios, params, rec)
            if word_bytes:
                st.download_button(
                    "Tải xuống (.docx)",
                    word_bytes,
                    file_name=f"CDM-THUYET-MINH-{params['bh_name']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            else:
                st.error("Lỗi tạo Word. Kiểm tra CDM/scripts/docx_helpers.py.")

    with c_excel:
        st.markdown("#### Tính toán chi tiết (Excel)")
        st.caption("3 sheet: Đầu vào · So sánh PA (bảng + biểu đồ nhúng) · Kết quả PA kiến nghị.")
        if st.button("Tạo file Excel", type="primary"):
            with st.spinner("Đang tạo Excel..."):
                xl_bytes = _export_excel(scenarios, params)
            if xl_bytes:
                st.download_button(
                    "Tải xuống (.xlsx)",
                    xl_bytes,
                    file_name=f"CDM-TINH-TOAN-{params['bh_name']}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            else:
                st.error("Lỗi tạo Excel. Kiểm tra openpyxl đã cài chưa.")

    # ── Lý thuyết tính toán (in PDF) ─────────────────────────────────────────
    st.divider()
    _theory_text = _load_theory()
    if _theory_text:
        with st.expander("Lý thuyết tính toán – đính kèm khi in PDF", expanded=True):
            st.markdown(
                "<style>"
                "@media print {"
                "  details { display: block !important; }"
                "  summary { display: none; }"
                "}"
                "</style>",
                unsafe_allow_html=True,
            )
            st.markdown(_theory_text)
